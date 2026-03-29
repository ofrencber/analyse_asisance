#!/usr/bin/env python3
"""
Comprehensive QA test suite for mcdm_article.py IMRAD generator.
"""
import sys
import os
import traceback
import time

sys.path.insert(0, os.path.dirname(__file__))

import pandas as pd
import numpy as np

# ---------------------------------------------------------------------------
# Mock data factory
# ---------------------------------------------------------------------------
def make_mock(wm, rm, n_alt=8, n_crit=5):
    np.random.seed(42)
    criteria = [f'C{i+1}' for i in range(n_crit)]
    ctypes = {c: ('min' if i == 0 else 'max') for i, c in enumerate(criteria)}
    data = pd.DataFrame(
        np.random.uniform(10, 100, (n_alt, n_crit)),
        columns=criteria,
        index=[f'A{i+1}' for i in range(n_alt)]
    )
    wt = pd.DataFrame({
        'Kriter': criteria,
        'Ağırlık': np.random.dirichlet(np.ones(n_crit))
    })

    rt = None
    rd = {}
    if rm:
        scores = np.random.uniform(0.2, 0.9, n_alt)
        scores.sort()
        scores = scores[::-1]
        rt = pd.DataFrame({
            'Alternatif': [f'A{i+1}' for i in range(n_alt)],
            'Skor': scores,
            'Sıra': range(1, n_alt + 1)
        })

    result = {
        'weights': {'method': wm, 'weight_table': wt},
        'ranking': {
            'method': rm,
            'ranking_table': rt,
            'ranking_details': rd
        } if rm else {
            'method': None,
            'ranking_table': None,
            'ranking_details': {}
        },
        'comparison': {
            'corr_matrix': pd.DataFrame(
                [[1.0, 0.88], [0.88, 1.0]],
                columns=['M1', 'M2'], index=['M1', 'M2']
            ),
            'rank_table': pd.DataFrame({
                'Alternatif': [f'A{i+1}' for i in range(n_alt)],
                'M1': range(1, n_alt + 1),
                'M2': list(range(1, n_alt + 1))
            })
        } if rm else {},
        'sensitivity': {
            'iterations': 200,
            'sigma': 0.12,
            'top_stability': 0.85
        } if rm else {},
        'validation': {'criteria': criteria, 'criteria_types': ctypes},
        'fuzzy_spread': 0.10,
        'selected_data': data,
    }
    return result, data


# ---------------------------------------------------------------------------
# Import the module under test
# ---------------------------------------------------------------------------
from mcdm_article import (
    generate_imrad_docx,
    generate_panel_imrad_docx,
    _load_template,
    _find_template_content,
    _extract_analysis_data,
    _WEIGHT_JUSTIFICATION_TR,
    _WEIGHT_JUSTIFICATION_EN,
    _RANKING_JUSTIFICATION_TR,
    _RANKING_JUSTIFICATION_EN,
)

# ---------------------------------------------------------------------------
# Collect results
# ---------------------------------------------------------------------------
results = []

def record(test_id, name, status, detail="", size=0):
    results.append({
        'test_id': test_id,
        'name': name,
        'status': status,
        'detail': detail,
        'size': size
    })

# ===========================================================================
# TEST 1: All 14 weight methods with TOPSIS ranking
# ===========================================================================
print("=" * 70)
print("TEST 1: All 14 weight methods with TOPSIS ranking")
print("=" * 70)

WEIGHT_METHODS = [
    "Entropy", "CRITIC", "Standart Sapma", "MEREC", "LOPCOW",
    "PCA", "CILOS", "IDOCRIW", "Fuzzy IDOCRIW",
    "AHP", "BWM", "SWARA", "DEMATEL", "SMART"
]

for wm in WEIGHT_METHODS:
    test_name = f"Weight={wm} + TOPSIS"
    try:
        t0 = time.time()
        result, data = make_mock(wm, "TOPSIS")
        docx_bytes = generate_imrad_docx(result, data, lang="TR")
        elapsed = time.time() - t0
        if docx_bytes and len(docx_bytes) > 0:
            sz = len(docx_bytes)
            record("T1", test_name, "PASS", f"{sz:,} bytes, {elapsed:.2f}s", sz)
            print(f"  PASS  {test_name:40s}  {sz:>8,} bytes  {elapsed:.2f}s")
        else:
            record("T1", test_name, "FAIL", "returned None or empty")
            print(f"  FAIL  {test_name:40s}  returned None")
    except Exception as e:
        record("T1", test_name, "ERROR", str(e))
        print(f"  ERROR {test_name:40s}  {e}")
        traceback.print_exc()

# ===========================================================================
# TEST 2: All ranking methods with Entropy
# ===========================================================================
print("\n" + "=" * 70)
print("TEST 2: All ranking methods with Entropy weight")
print("=" * 70)

RANKING_METHODS = [
    "TOPSIS", "VIKOR", "EDAS", "CODAS", "COPRAS", "PROMETHEE", "GRA",
    "SPOTIS", "MULTIMOORA", "MARCOS", "CoCoSo", "MOORA", "SAW", "WPM",
    "MAUT", "WASPAS", "ARAS", "OCRA", "MABAC", "RAWEC", "RAFSI", "ROV",
    "AROMAN", "DNMA",
    "Fuzzy TOPSIS", "Fuzzy VIKOR", "Fuzzy EDAS", "Fuzzy CODAS",
]

template = _load_template()

for rm in RANKING_METHODS:
    test_name = f"Entropy + {rm}"
    try:
        t0 = time.time()
        result, data = make_mock("Entropy", rm)
        docx_bytes = generate_imrad_docx(result, data, lang="TR")
        elapsed = time.time() - t0
        if docx_bytes and len(docx_bytes) > 0:
            sz = len(docx_bytes)
            # Check template content
            wc = _find_template_content("Entropy", template["weight_methods"])
            rc = _find_template_content(rm, template["ranking_methods"])
            wc_status = "HAS_CONTENT" if wc else "EMPTY"
            rc_status = "HAS_CONTENT" if rc else "EMPTY"
            detail = f"{sz:,} bytes, {elapsed:.2f}s, weight_tpl={wc_status}, rank_tpl={rc_status}"
            record("T2", test_name, "PASS", detail, sz)
            print(f"  PASS  {test_name:35s} {sz:>8,}B  wt={wc_status:11s} rt={rc_status:11s}")
        else:
            record("T2", test_name, "FAIL", "returned None or empty")
            print(f"  FAIL  {test_name:35s} returned None")
    except Exception as e:
        record("T2", test_name, "ERROR", str(e))
        print(f"  ERROR {test_name:35s} {e}")
        traceback.print_exc()

# ===========================================================================
# TEST 3: Edge cases
# ===========================================================================
print("\n" + "=" * 70)
print("TEST 3: Edge cases")
print("=" * 70)

# 3a: Weight-only (no ranking method)
test_name = "Weight-only (no ranking)"
try:
    result, data = make_mock("Entropy", None)
    docx_bytes = generate_imrad_docx(result, data, lang="TR")
    if docx_bytes and len(docx_bytes) > 0:
        sz = len(docx_bytes)
        record("T3a", test_name, "PASS", f"{sz:,} bytes", sz)
        print(f"  PASS  {test_name:40s}  {sz:>8,} bytes")
    else:
        record("T3a", test_name, "FAIL", "returned None")
        print(f"  FAIL  {test_name}")
except Exception as e:
    record("T3a", test_name, "ERROR", str(e))
    print(f"  ERROR {test_name}  {e}")
    traceback.print_exc()

# 3b: Panel data with multiple periods
test_name = "Panel data (3 periods)"
try:
    panel = {}
    for period in ["2022", "2023", "2024"]:
        r, d = make_mock("Entropy", "TOPSIS")
        r["selected_data"] = d
        panel[period] = r
    docx_bytes = generate_panel_imrad_docx(panel, lang="TR")
    if docx_bytes and len(docx_bytes) > 0:
        sz = len(docx_bytes)
        record("T3b", test_name, "PASS", f"{sz:,} bytes", sz)
        print(f"  PASS  {test_name:40s}  {sz:>8,} bytes")
    else:
        record("T3b", test_name, "FAIL", "returned None")
        print(f"  FAIL  {test_name}")
except Exception as e:
    record("T3b", test_name, "ERROR", str(e))
    print(f"  ERROR {test_name}  {e}")
    traceback.print_exc()

# 3c: English language output
test_name = "English language (EN)"
try:
    result, data = make_mock("Entropy", "TOPSIS")
    docx_bytes = generate_imrad_docx(result, data, lang="EN")
    if docx_bytes and len(docx_bytes) > 0:
        sz = len(docx_bytes)
        record("T3c", test_name, "PASS", f"{sz:,} bytes", sz)
        print(f"  PASS  {test_name:40s}  {sz:>8,} bytes")
    else:
        record("T3c", test_name, "FAIL", "returned None")
        print(f"  FAIL  {test_name}")
except Exception as e:
    record("T3c", test_name, "ERROR", str(e))
    print(f"  ERROR {test_name}  {e}")
    traceback.print_exc()

# 3d: Empty comparison/sensitivity
test_name = "Empty comparison/sensitivity"
try:
    result, data = make_mock("Entropy", "TOPSIS")
    result['comparison'] = {}
    result['sensitivity'] = {}
    docx_bytes = generate_imrad_docx(result, data, lang="TR")
    if docx_bytes and len(docx_bytes) > 0:
        sz = len(docx_bytes)
        record("T3d", test_name, "PASS", f"{sz:,} bytes", sz)
        print(f"  PASS  {test_name:40s}  {sz:>8,} bytes")
    else:
        record("T3d", test_name, "FAIL", "returned None")
        print(f"  FAIL  {test_name}")
except Exception as e:
    record("T3d", test_name, "ERROR", str(e))
    print(f"  ERROR {test_name}  {e}")
    traceback.print_exc()

# 3e: Very long criteria names
test_name = "Very long criteria names"
try:
    np.random.seed(42)
    long_criteria = [f'VeryLongCriterionName_Category{i+1}_Subcategory_Detail_Extra' for i in range(5)]
    ctypes = {c: 'max' for c in long_criteria}
    data_long = pd.DataFrame(
        np.random.uniform(10, 100, (8, 5)),
        columns=long_criteria,
        index=[f'A{i+1}' for i in range(8)]
    )
    wt_long = pd.DataFrame({
        'Kriter': long_criteria,
        'Ağırlık': np.random.dirichlet(np.ones(5))
    })
    result_long = {
        'weights': {'method': 'Entropy', 'weight_table': wt_long},
        'ranking': {'method': 'TOPSIS', 'ranking_table': pd.DataFrame({
            'Alternatif': [f'A{i+1}' for i in range(8)],
            'Skor': np.random.uniform(0.2, 0.9, 8),
            'Sıra': range(1, 9)
        }), 'ranking_details': {}},
        'comparison': {},
        'sensitivity': {},
        'validation': {'criteria': long_criteria, 'criteria_types': ctypes},
        'fuzzy_spread': 0.10,
        'selected_data': data_long,
    }
    docx_bytes = generate_imrad_docx(result_long, data_long, lang="TR")
    if docx_bytes and len(docx_bytes) > 0:
        sz = len(docx_bytes)
        record("T3e", test_name, "PASS", f"{sz:,} bytes", sz)
        print(f"  PASS  {test_name:40s}  {sz:>8,} bytes")
    else:
        record("T3e", test_name, "FAIL", "returned None")
        print(f"  FAIL  {test_name}")
except Exception as e:
    record("T3e", test_name, "ERROR", str(e))
    print(f"  ERROR {test_name}  {e}")
    traceback.print_exc()

# 3f: Single alternative, single criterion
test_name = "Single alt, single criterion"
try:
    result_s, data_s = make_mock("Entropy", "TOPSIS", n_alt=1, n_crit=1)
    docx_bytes = generate_imrad_docx(result_s, data_s, lang="TR")
    if docx_bytes and len(docx_bytes) > 0:
        sz = len(docx_bytes)
        record("T3f", test_name, "PASS", f"{sz:,} bytes", sz)
        print(f"  PASS  {test_name:40s}  {sz:>8,} bytes")
    else:
        record("T3f", test_name, "FAIL", "returned None")
        print(f"  FAIL  {test_name}")
except Exception as e:
    record("T3f", test_name, "ERROR", str(e))
    print(f"  ERROR {test_name}  {e}")
    traceback.print_exc()

# ===========================================================================
# TEST 4: Content quality check - Entropy + Fuzzy TOPSIS, TR
# ===========================================================================
print("\n" + "=" * 70)
print("TEST 4: Content quality check (Entropy + Fuzzy TOPSIS, TR)")
print("=" * 70)

try:
    from docx import Document as DocxDocument

    result, data = make_mock("Entropy", "Fuzzy TOPSIS")
    docx_bytes = generate_imrad_docx(result, data, lang="TR")

    if docx_bytes:
        # Save to file
        out_path = "/Users/rencber/Desktop/MCDM_Toolbox/_exports/test_article_quality.docx"
        os.makedirs(os.path.dirname(out_path), exist_ok=True)
        with open(out_path, 'wb') as f:
            f.write(docx_bytes)
        print(f"  Saved DOCX: {out_path} ({len(docx_bytes):,} bytes)")

        # Extract all text
        import io
        doc = DocxDocument(io.BytesIO(docx_bytes))

        all_text = []
        for para in doc.paragraphs:
            all_text.append(para.text)

        full_text = "\n".join(all_text)

        # Check IMRAD sections
        imrad_checks = {
            "OZET_section": "ÖZET" in full_text,
            "GIRIS_section": "GİRİŞ" in full_text,
            "YONTEM_section": "YÖNTEM" in full_text,
            "BULGULAR_section": "BULGULAR" in full_text,
            "TARTISMA_section": "TARTIŞMA" in full_text,
            "SONUC_section": "SONUÇ" in full_text,
            "KAYNAKLAR_section": "KAYNAKLAR" in full_text,
            "method_dm_section": "Karar Matrisi" in full_text,
            "method_weight_section": "Kriter Ağırlıklarının Belirlenmesi" in full_text,
            "method_rank_section": "Alternatif Sıralaması" in full_text,
            "method_robust_section": "Sağlamlık Analizi" in full_text,
            "results_weight_section": "Kriter Ağırlıkları" in full_text,
            "results_rank_section": "Sıralama Sonuçları" in full_text,
        }

        # Check placeholders
        placeholder_checks = {
            "has_ALAN_placeholder": "[ALAN]" in full_text,
            "has_PROBLEM_placeholder": "[PROBLEM]" in full_text,
            "has_VERI_KAYNAGI_placeholder": "[VERİ_KAYNAĞI]" in full_text,
        }

        # Check justification
        justification_checks = {
            "entropy_justification": "Shannon bilgi entropisi" in full_text or "entropi" in full_text.lower(),
            "fuzzy_justification": "bulanık" in full_text.lower() or "TFN" in full_text,
            "pipeline_position": "pipeline" in full_text.lower() or "bulanıklaştırma adımı" in full_text,
        }

        # Check tables
        n_tables = len(doc.tables)
        table_check = n_tables > 0

        # Check images (inline shapes)
        n_images = 0
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                n_images += 1

        # Check for highlighted runs (yellow placeholders)
        n_highlighted = 0
        for para in doc.paragraphs:
            for run in para.runs:
                rpr = run._element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}highlight')
                if rpr is not None:
                    n_highlighted += 1

        content_results = {
            **imrad_checks,
            **placeholder_checks,
            **justification_checks,
            "tables_present": table_check,
            "n_tables": n_tables,
            "images_present": n_images > 0,
            "n_images": n_images,
            "n_highlighted_runs": n_highlighted,
            "total_paragraphs": len(doc.paragraphs),
            "total_chars": len(full_text),
        }

        print("\n  --- IMRAD Section Checks ---")
        for k, v in imrad_checks.items():
            status = "PASS" if v else "FAIL"
            print(f"    {status}  {k}")
            record("T4", f"IMRAD:{k}", status)

        print("\n  --- Placeholder Checks ---")
        for k, v in placeholder_checks.items():
            status = "PASS" if v else "FAIL"
            print(f"    {status}  {k}")
            record("T4", f"Placeholder:{k}", status)

        print("\n  --- Justification Checks ---")
        for k, v in justification_checks.items():
            status = "PASS" if v else "FAIL"
            print(f"    {status}  {k}")
            record("T4", f"Justification:{k}", status)

        print(f"\n  --- Document Stats ---")
        print(f"    Tables: {n_tables}")
        print(f"    Images: {n_images}")
        print(f"    Highlighted runs (placeholders): {n_highlighted}")
        print(f"    Paragraphs: {len(doc.paragraphs)}")
        print(f"    Total chars: {len(full_text):,}")

        record("T4", "Tables present", "PASS" if table_check else "FAIL", f"n={n_tables}")
        record("T4", "Images present", "PASS" if n_images > 0 else "FAIL", f"n={n_images}")
        record("T4", "Highlighted placeholders", "PASS" if n_highlighted > 0 else "FAIL", f"n={n_highlighted}")

        # Print full extracted text for review
        print("\n  --- FULL EXTRACTED TEXT (first 5000 chars) ---")
        print(full_text[:5000])
        if len(full_text) > 5000:
            print(f"\n  ... [{len(full_text) - 5000} more chars] ...")

    else:
        record("T4", "Content quality", "FAIL", "generate_imrad_docx returned None")
        print("  FAIL  generate_imrad_docx returned None")

except Exception as e:
    record("T4", "Content quality", "ERROR", str(e))
    print(f"  ERROR  {e}")
    traceback.print_exc()


# ===========================================================================
# TEST 5: Template coverage
# ===========================================================================
print("\n" + "=" * 70)
print("TEST 5: Template coverage - which methods have content")
print("=" * 70)

template = _load_template()

print("\n  --- Weight method template keys found ---")
for k in sorted(template["weight_methods"].keys()):
    length = len(template["weight_methods"][k])
    print(f"    '{k}': {length} chars")

print("\n  --- Ranking method template keys found ---")
for k in sorted(template["ranking_methods"].keys()):
    length = len(template["ranking_methods"][k])
    print(f"    '{k}': {length} chars")

print("\n  --- Weight methods: template content coverage ---")
for wm in WEIGHT_METHODS:
    content = _find_template_content(wm, template["weight_methods"])
    status = "HAS_CONTENT" if content else "NO_CONTENT"
    length = len(content) if content else 0
    flag = "PASS" if content else "WARN"
    record("T5", f"Weight template: {wm}", flag, f"{length} chars")
    print(f"    {flag:4s}  {wm:20s}  {status} ({length} chars)")

print("\n  --- Ranking methods: template content coverage ---")
for rm in RANKING_METHODS:
    content = _find_template_content(rm, template["ranking_methods"])
    status = "HAS_CONTENT" if content else "NO_CONTENT"
    length = len(content) if content else 0
    flag = "PASS" if content else "WARN"
    record("T5", f"Rank template: {rm}", flag, f"{length} chars")
    print(f"    {flag:4s}  {rm:20s}  {status} ({length} chars)")


# ===========================================================================
# TEST 6: Justification coverage
# ===========================================================================
print("\n" + "=" * 70)
print("TEST 6: Justification text coverage")
print("=" * 70)

print("\n  --- Weight justification (TR) ---")
for wm in WEIGHT_METHODS:
    has_tr = wm in _WEIGHT_JUSTIFICATION_TR
    has_en = wm in _WEIGHT_JUSTIFICATION_EN
    status = f"TR={'YES' if has_tr else 'NO':3s} EN={'YES' if has_en else 'NO':3s}"
    flag = "PASS" if has_tr else "WARN"
    record("T6", f"Weight just: {wm}", flag, status)
    print(f"    {flag:4s}  {wm:20s}  {status}")

print("\n  --- Ranking justification (TR) ---")
for rm in RANKING_METHODS:
    base_rm = rm.replace("Fuzzy ", "")
    has_tr = rm in _RANKING_JUSTIFICATION_TR or base_rm in _RANKING_JUSTIFICATION_TR
    has_en = rm in _RANKING_JUSTIFICATION_EN or base_rm in _RANKING_JUSTIFICATION_EN
    status = f"TR={'YES' if has_tr else 'NO':3s} EN={'YES' if has_en else 'NO':3s}"
    flag = "PASS" if has_tr else "WARN"
    record("T6", f"Rank just: {rm}", flag, status)
    print(f"    {flag:4s}  {rm:20s}  {status}")


# ===========================================================================
# SUMMARY
# ===========================================================================
print("\n" + "=" * 70)
print("SUMMARY REPORT")
print("=" * 70)

pass_count = sum(1 for r in results if r['status'] == 'PASS')
fail_count = sum(1 for r in results if r['status'] == 'FAIL')
error_count = sum(1 for r in results if r['status'] == 'ERROR')
warn_count = sum(1 for r in results if r['status'] == 'WARN')
total = len(results)

print(f"\n  Total tests: {total}")
print(f"  PASS:  {pass_count}")
print(f"  FAIL:  {fail_count}")
print(f"  ERROR: {error_count}")
print(f"  WARN:  {warn_count}")

if fail_count > 0 or error_count > 0:
    print("\n  --- FAILURES and ERRORS ---")
    for r in results:
        if r['status'] in ('FAIL', 'ERROR'):
            print(f"    [{r['status']}] {r['test_id']}: {r['name']} — {r['detail']}")

if warn_count > 0:
    print("\n  --- WARNINGS (missing content) ---")
    for r in results:
        if r['status'] == 'WARN':
            print(f"    [WARN] {r['test_id']}: {r['name']} — {r['detail']}")

# File sizes summary for T1/T2
sizes = [r for r in results if r['size'] > 0 and r['test_id'] in ('T1', 'T2')]
if sizes:
    min_s = min(r['size'] for r in sizes)
    max_s = max(r['size'] for r in sizes)
    avg_s = sum(r['size'] for r in sizes) / len(sizes)
    print(f"\n  --- DOCX file sizes (T1+T2) ---")
    print(f"    Min: {min_s:,} bytes")
    print(f"    Max: {max_s:,} bytes")
    print(f"    Avg: {avg_s:,.0f} bytes")

print("\n  DONE.")
