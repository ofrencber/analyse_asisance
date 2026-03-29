from __future__ import annotations

import io
import re
import time
from typing import Any, Dict, List

import numpy as np
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
import streamlit as st
from scipy.stats import spearmanr

import mcdm_engine as me

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt
    from docx.enum.section import WD_SECTION
    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False

st.set_page_config(
    page_title="MCDM Toolbox — Prof. Dr. Ömer Faruk Rençber",
    page_icon="🔬",
    layout="wide",
    initial_sidebar_state="expanded",
)

st.markdown(
    """
    <style>
        :root {
            --bg-main: #FAF8F2;
            --bg-sidebar: #BFC5CF;
            --card-bg: #FFFFFF;
            --text-main: #0F172A;
            --text-muted: #334155;
            --border-strong: #000000;
            --accent: #1D4E89;
            --accent-soft: #EAF3FF;
            --header-text: #FFF5E1;
        }
        .stApp {
            background-color: var(--bg-main);
            color: var(--text-main);
        }
        .main .block-container {
            padding-top: 8.6rem !important;
            padding-bottom: 2rem;
            max-width: 1500px;
        }

        .global-header {
            position: fixed;
            top: 0;
            left: 0;
            width: 100vw;
            min-height: 132px;
            background: linear-gradient(100deg, #113A63 0%, #225D8F 50%, #2D6FA5 100%);
            color: var(--header-text) !important;
            z-index: 999 !important;
            display: grid;
            grid-template-columns: 1.1fr 1.5fr 1.1fr;
            align-items: center;
            gap: 0.5rem;
            padding: 1.15rem 2rem;
            border-bottom: 2px solid #0b2137;
        }
        .header-left {
            display: flex;
            flex-direction: column;
            justify-content: center;
            text-align: left;
        }
        .header-dedication {
            margin: 0;
            font-size: 0.92rem;
            font-style: italic;
            font-weight: 600;
            color: #FFECCE !important;
        }
        .header-title {
            margin: 0;
            font-size: 1.28rem;
            font-weight: 800;
            letter-spacing: 0.45px;
            color: var(--header-text) !important;
        }
        .header-slogan {
            margin: 0;
            font-size: 0.95rem;
            font-style: italic;
            color: #FFECCE !important;
        }
        .header-signature {
            text-align: center;
            font-size: 1.56rem;
            color: var(--header-text) !important;
            font-family: "Georgia", serif;
            font-style: italic;
            font-weight: 700;
        }
        .header-right {
            text-align: right;
            font-size: 0.92rem;
            font-weight: 600;
            color: #FFECCE !important;
        }

        section[data-testid="stSidebar"],
        section[data-testid="stSidebar"] > div,
        aside[data-testid="stSidebar"],
        aside[data-testid="stSidebar"] > div {
            background-color: var(--bg-sidebar) !important;
            border-right: 2px solid var(--border-strong);
        }
        section[data-testid="stSidebar"],
        aside[data-testid="stSidebar"] {
            min-width: 320px !important;
            width: 320px !important;
            transform: translateX(0) !important;
        }
        section[data-testid="stSidebar"] *,
        aside[data-testid="stSidebar"] *,
        [data-testid="stSidebar"] * {
            color: #111827 !important;
        }

        .section-card {
            background: var(--card-bg);
            border-radius: 10px;
            padding: 1rem 1.1rem;
            border: 1px solid var(--border-strong);
            box-shadow: none;
            margin-bottom: 0.9rem;
        }
        [data-testid="stTable"] table {
            background: #FFFFFF !important;
            color: #000000 !important;
            border: 1px solid #9CA3AF !important;
        }
        [data-testid="stTable"] thead th {
            background: #F3F4F6 !important;
            color: #000000 !important;
            border: 1px solid #9CA3AF !important;
        }
        [data-testid="stTable"] tbody td {
            background: #FFFFFF !important;
            color: #000000 !important;
            border: 1px solid #D1D5DB !important;
        }
        .stTable {
            color: #000000 !important;
        }
        [data-testid="stExpander"] {
            border: 1px solid var(--border-strong);
            border-radius: 10px;
            background: #FFFFFF;
            margin-bottom: 0.6rem;
        }
        [data-testid="stExpander"] details {
            border: none !important;
            border-radius: 10px;
            background: #FFFFFF;
        }
        [data-testid="stExpander"] summary {
            border-radius: 10px;
        }
        [data-testid="stFileUploaderDropzone"] {
            border: 1px solid var(--border-strong) !important;
            background: #F3F4F6 !important;
        }
        [data-testid="stFileUploaderDropzone"] button {
            border: 1px solid var(--border-strong) !important;
            background: #4B5563 !important;
            color: #FFFFFF !important;
            border-radius: 8px !important;
            font-weight: 700 !important;
        }

        .stButton > button,
        [data-testid="stDownloadButton"] > button {
            border-radius: 8px !important;
            font-weight: 700 !important;
            border: 1px solid var(--border-strong) !important;
            background: #D1D5DB !important;
            color: #111827 !important;
            transition: all 0.12s ease;
        }
        .stButton > button:hover,
        [data-testid="stDownloadButton"] > button:hover {
            background: #9CA3AF !important;
            color: #111827 !important;
            border: 1px solid var(--border-strong) !important;
        }
        .stButton > button:focus,
        .stButton > button:active,
        [data-testid="stDownloadButton"] > button:focus,
        [data-testid="stDownloadButton"] > button:active {
            color: #111827 !important;
            border: 1px solid var(--border-strong) !important;
            background: #B9C0CA !important;
            box-shadow: none !important;
        }
        .stButton > button[kind="primary"] {
            background: #D1D5DB !important;
            color: #111827 !important;
            border: 1px solid var(--border-strong) !important;
        }
        .sidebar-small-note {
            font-size: 0.80rem;
            line-height: 1.45;
        }

        .stTabs [data-baseweb="tab-list"] {
            gap: 4px;
            background: #E2E8F0;
            border-radius: 10px;
            border: 1px solid var(--border-strong);
            padding: 4px;
        }
        .stTabs [data-baseweb="tab"] {
            background: transparent;
            border-radius: 8px;
            border: none;
            padding: 0.42rem 0.88rem;
            font-size: 0.85rem;
            font-weight: 600;
            color: #1E293B;
        }
        .stTabs [aria-selected="true"] {
            background: #FFFFFF !important;
            color: #0F172A !important;
            font-weight: 800 !important;
            border: 1px solid var(--border-strong) !important;
        }

        .assistant-box {
            background: var(--accent-soft);
            border: 1px solid var(--border-strong);
            border-radius: 10px;
            padding: 1rem;
            margin-bottom: 0.9rem;
        }
        h2, h3, h4, p, span, label, div {
            color: var(--text-main);
        }

        .layer-card {
            border-radius: 8px;
            padding: 0.8rem 1rem;
            margin-bottom: 0.5rem;
            border: 1px solid var(--border-strong);
            line-height: 1.5;
        }
        .layer-desc { background: #EAF4FF; }
        .layer-analytic { background: #EEFDF3; }
        .layer-norm { background: #FFF8E8; }
        .layer-label {
            font-size: 0.69rem;
            font-weight: 700;
            text-transform: uppercase;
            letter-spacing: 0.4px;
            margin-bottom: 0.3rem;
        }
        .layer-text {
            font-size: 0.87rem;
            color: #1F2937;
            margin: 0;
        }

        .tab-assistant {
            background: #EAF4FF;
            border: 1px solid var(--border-strong);
            border-left: 4px solid var(--accent);
            border-radius: 10px;
            padding: 0.9rem 1rem;
            font-size: 0.86rem;
            line-height: 1.62;
            color: #0B1F35;
        }
        .tab-assistant strong { color: #0B1F35; }

        .badge-benefit { background:#DCFCE7; color:#166534; border-radius:4px; padding:0.1rem 0.35rem; font-size:0.65rem; font-weight:700; }
        .badge-cost    { background:#FEE2E2; color:#991B1B; border-radius:4px; padding:0.1rem 0.35rem; font-size:0.65rem; font-weight:700; }

        .network-wrap {
            background:#ffffff;
            border-radius:10px;
            border:1px solid var(--border-strong);
            padding:0.5rem;
        }

        .mc-high  { color:#166534; font-weight:700; }
        .mc-mid   { color:#92400E; font-weight:700; }
        .mc-low   { color:#991B1B; font-weight:700; }

        .stepper-wrap {
            background:#ffffff;
            border:1px solid var(--border-strong);
            border-radius:10px;
            padding:0.65rem 0.8rem;
            margin-bottom:0.8rem;
        }
        .step-item { display:flex; align-items:center; gap:0.45rem; margin:0.35rem 0; font-size:0.82rem; color:#1F2937; }
        .step-dot { width:9px; height:9px; border-radius:50%; display:inline-block; }
        .step-done { background:#16A34A; }
        .step-pending { background:#64748B; }
        .step-active { background:#2563EB; box-shadow:0 0 0 3px rgba(37,99,235,0.16); }
        .kpi-strip { background:#ffffff; border:1px solid var(--border-strong); border-radius:10px; padding:0.75rem 0.9rem; margin:0.35rem 0 1rem 0; }
        .kpi-grid { display:grid; grid-template-columns:repeat(5, minmax(0,1fr)); gap:0.55rem; }
        .kpi-item { background:#F8FAFC; border:1px solid var(--border-strong); border-radius:8px; padding:0.5rem 0.6rem; }
        .kpi-label { font-size:0.70rem; color:#334155; text-transform:uppercase; letter-spacing:.35px; }
        .kpi-value { font-size:0.9rem; font-weight:700; color:#0F172A; margin-top:0.1rem; }
        .assistant-grid { display:grid; grid-template-columns:repeat(3, minmax(0,1fr)); gap:0.6rem; margin-bottom:0.6rem; }
        .assistant-card2 { background:#ffffff; border:1px solid var(--border-strong); border-radius:10px; padding:0.6rem 0.7rem; }
        .assistant-title2 { font-size:0.72rem; font-weight:700; color:#0F172A; text-transform:uppercase; letter-spacing:.3px; margin-bottom:0.25rem; }
        .assistant-body2 { font-size:0.82rem; color:#1F2937; line-height:1.45; }
        .diag-badge { display:inline-block; font-size:0.68rem; font-weight:700; border-radius:999px; padding:0.15rem 0.45rem; margin-left:0.3rem; }
        .diag-badge-good { background:#DCFCE7; color:#166534; }
        .diag-badge-mid { background:#FEF3C7; color:#92400E; }
        .diag-badge-bad { background:#FEE2E2; color:#991B1B; }

        @media (max-width: 980px) {
            .main .block-container { padding-top: 6.9rem !important; }
            .global-header {
                grid-template-columns: 1fr;
                text-align: center;
                padding: 0.5rem 0.8rem;
                min-height: 86px;
            }
            .header-left, .header-right { display: none; }
            .header-signature { font-size: 1.08rem; }
            .header-dedication { display: none; }
            .kpi-grid { grid-template-columns:repeat(2, minmax(0,1fr)); }
            .assistant-grid { grid-template-columns:1fr; }
        }
    </style>
    """,
    unsafe_allow_html=True,
)

def init_state() -> None:
    defaults = {
        "analysis_result": None,
        "clean_data": None,
        "raw_data": None,
        "report_docx": None,
        "editor_df": None,
        "data_source_id": None,
        "prep_done": False,
        "alt_names": {},        
        "crit_dir": {},         
        "crit_include": {},     
        "weight_method_pref": None,
        "ranking_prefs": [],
        "ui_lang": "TR",
    }
    for key, value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = value

init_state()

def tt(tr_text: str, en_text: str) -> str:
    return en_text if st.session_state.get("ui_lang", "TR") == "EN" else tr_text

_COL_TR_EN = {
    "Kriter": "Criterion",
    "Ağırlık": "Weight",
    "HamAğırlık": "RawWeight",
    "Alternatif": "Alternative",
    "Skor": "Score",
    "Sıra": "Rank",
    "Yöntem": "Method",
    "BirincilikOranı": "FirstPlaceRate",
    "OrtalamaSıra": "MeanRank",
    "AğırlıkDeğişimi": "WeightChange",
    "YeniAğırlık": "NewWeight",
    "TopKriterUyumu": "TopCriterionMatch",
    "MaksMutlakFark": "MaxAbsoluteDiff",
    "OrtalamaAğırlık": "MeanWeight",
    "AltYuzde5": "Lower5Pct",
    "UstYuzde95": "Upper95Pct",
    "StdSapma": "StdDev",
    "Bilgiİçeriği": "InformationContent",
    "Etkisi": "Effect",
    "Enİyi": "Best",
    "EnKötü": "Worst",
    "Ortalama": "Average",
    "BirinciAlternatif": "TopAlternative",
    "Değer": "Value",
    "Parametre": "Parameter",
}

_WEIGHT_METHOD_TR_EN = {
    "Standart Sapma": "Standard Deviation",
    "Eşit Ağırlık": "Equal Weights",
    "Manuel Ağırlık": "Manual Weights",
}

def col_key(df: pd.DataFrame, tr_name: str, en_name: str) -> str:
    if tr_name in df.columns:
        return tr_name
    if en_name in df.columns:
        return en_name
    return tr_name

def localize_df(df: pd.DataFrame) -> pd.DataFrame:
    if not isinstance(df, pd.DataFrame) or st.session_state.get("ui_lang") != "EN":
        return df
    return df.rename(columns={k: v for k, v in _COL_TR_EN.items() if k in df.columns})

def localize_df_lang(df: pd.DataFrame, lang: str) -> pd.DataFrame:
    if not isinstance(df, pd.DataFrame) or lang != "EN":
        return df
    return df.rename(columns={k: v for k, v in _COL_TR_EN.items() if k in df.columns})

def method_display_name(name: str) -> str:
    if st.session_state.get("ui_lang") == "EN":
        return _WEIGHT_METHOD_TR_EN.get(name, name)
    return name

def method_internal_name(name: str) -> str:
    if name in _WEIGHT_METHOD_TR_EN:
        return name
    for tr_name, en_name in _WEIGHT_METHOD_TR_EN.items():
        if name == en_name:
            return tr_name
    return name

def sample_dataset() -> pd.DataFrame:
    rng = np.random.default_rng(42)
    n = 16
    df = pd.DataFrame({
        "Alternatif": [f"A{idx:02d}" for idx in range(1, n + 1)],
        "Karlılık": rng.uniform(10, 28, n),
        "Likidite": rng.uniform(0.9, 2.6, n),
        "PazarPayı": rng.uniform(8, 22, n),
        "MüşteriMemnuniyeti": rng.uniform(55, 95, n),
        "KaliteSkoru": rng.uniform(60, 98, n),
        "TeslimSüresi": rng.uniform(2, 15, n),
        "Risk": rng.uniform(1, 12, n),
        "İşletmeMaliyeti": rng.uniform(80, 420, n),
    })
    return df

def guess_direction(col_name: str) -> str:
    lowered = col_name.lower()
    cost_keywords = ["maliyet", "risk", "süre", "borç", "gider", "hata", "şikayet", "kayip", "loss", "cost", "time", "defect"]
    return "Min (Maliyet)" if any(k in lowered for k in cost_keywords) else "Max (Fayda)"

def clean_dataframe(df: pd.DataFrame, missing_strategy: str, clip_outliers: bool) -> pd.DataFrame:
    out = df.copy()
    numeric_cols = out.select_dtypes(include=[np.number]).columns.tolist()
    for col in numeric_cols:
        if out[col].isna().any():
            if missing_strategy in {"Medyan", "Median"}:
                out[col] = out[col].fillna(out[col].median())
            elif missing_strategy in {"Ortalama", "Mean"}:
                out[col] = out[col].fillna(out[col].mean())
            elif missing_strategy in {"Interpolasyon", "Interpolation"}:
                out[col] = out[col].interpolate(method="linear", limit_direction="both")
                out[col] = out[col].fillna(out[col].median())
            else:
                out[col] = out[col].fillna(0)
        if clip_outliers:
            q1, q3 = out[col].quantile(0.25), out[col].quantile(0.75)
            iqr = q3 - q1
            out[col] = np.clip(out[col], q1 - 1.5 * iqr, q3 + 1.5 * iqr)
    return out

def recommend_parameter_defaults(selected_methods: List[str], n_alt: int, n_crit: int, uses_fuzzy: bool) -> Dict[str, float]:
    defaults = {
        "vikor_v": 0.50,
        "waspas_lambda": 0.50,
        "codas_tau": 0.02,
        "cocoso_lambda": 0.50,
        "gra_rho": 0.50,
        "promethee_q": 0.05,
        "promethee_p": 0.30,
        "promethee_s": 0.20,
        "promethee_pref_func": "linear",
        "fuzzy_spread": 0.10,
        "sensitivity_iterations": 800,
        "sensitivity_sigma": 0.12,
    }
    if n_alt <= 8:
        defaults["sensitivity_iterations"] = 500
        defaults["sensitivity_sigma"] = 0.14
    elif n_alt >= 30:
        defaults["sensitivity_iterations"] = 1500
        defaults["sensitivity_sigma"] = 0.10
    if n_crit >= 10:
        defaults["sensitivity_sigma"] = max(0.08, defaults["sensitivity_sigma"] - 0.01)
    if uses_fuzzy:
        defaults["fuzzy_spread"] = 0.08 if n_alt >= 20 else 0.10
    if any("CODAS" in m for m in selected_methods) and n_alt > 20:
        defaults["codas_tau"] = 0.015
    return defaults

def _safe_spearman(x: np.ndarray, y: np.ndarray) -> float:
    try:
        rho, _ = spearmanr(x, y)
        if rho is None or not np.isfinite(rho):
            return 0.0
        return float(rho)
    except Exception:
        return 0.0

def compute_weight_robustness(
    data: pd.DataFrame,
    criteria: List[str],
    criteria_types: Dict[str, str],
    weight_method: str,
    base_weights: Dict[str, float],
    bootstrap_n: int = 160,
) -> Dict[str, Any]:
    base_vec = np.asarray([float(base_weights[c]) for c in criteria], dtype=float)
    concentration = float(np.sum(base_vec ** 2))
    eff_n = float(1.0 / concentration) if concentration > 0 else 0.0
    sorted_w = sorted(base_weights.items(), key=lambda kv: kv[1], reverse=True)
    top1, top1_w = sorted_w[0]
    top2_w = sorted_w[1][1] if len(sorted_w) > 1 else np.nan
    dominance_ratio = float(top1_w / top2_w) if np.isfinite(top2_w) and top2_w > 0 else np.nan

    loo_rows: List[Dict[str, Any]] = []
    if len(data) >= 5:
        for alt in data.index:
            sub = data.drop(index=alt)
            if len(sub) < 4:
                continue
            try:
                w_sub, _ = me.compute_objective_weights(sub, criteria, criteria_types, weight_method)
                sub_vec = np.asarray([float(w_sub[c]) for c in criteria], dtype=float)
                rho = _safe_spearman(base_vec, sub_vec)
                top_sub = max(w_sub, key=w_sub.get)
                loo_rows.append(
                    {
                        "Alternatif": str(alt),
                        "SpearmanRho": rho,
                        "TopKriterUyumu": int(top_sub == top1),
                        "MaksMutlakFark": float(np.max(np.abs(sub_vec - base_vec))),
                    }
                )
            except Exception:
                continue
    loo_df = pd.DataFrame(loo_rows)
    loo_mean = float(loo_df["SpearmanRho"].mean()) if not loo_df.empty else np.nan
    loo_min = float(loo_df["SpearmanRho"].min()) if not loo_df.empty else np.nan
    loo_top_match = float(loo_df["TopKriterUyumu"].mean()) if not loo_df.empty else np.nan

    rng = np.random.default_rng(42)
    boot_rows: List[Dict[str, float]] = []
    n_boot = int(max(80, min(bootstrap_n, 320)))
    for _ in range(n_boot):
        idx = rng.integers(0, len(data), len(data))
        sample = data.iloc[idx]
        try:
            w_b, _ = me.compute_objective_weights(sample, criteria, criteria_types, weight_method)
            boot_rows.append({c: float(w_b[c]) for c in criteria})
        except Exception:
            continue
    boot_df = pd.DataFrame(boot_rows)
    boot_summary_rows: List[Dict[str, Any]] = []
    if not boot_df.empty:
        for c in criteria:
            col = boot_df[c]
            boot_summary_rows.append(
                {
                    "Kriter": c,
                    "OrtalamaAğırlık": float(col.mean()),
                    "StdSapma": float(col.std(ddof=0)),
                    "AltYuzde5": float(col.quantile(0.05)),
                    "UstYuzde95": float(col.quantile(0.95)),
                }
            )
    boot_summary = pd.DataFrame(boot_summary_rows).sort_values("OrtalamaAğırlık", ascending=False) if boot_summary_rows else pd.DataFrame()

    return {
        "base_top_criterion": top1,
        "concentration": concentration,
        "effective_criteria_n": eff_n,
        "dominance_ratio": dominance_ratio,
        "leave_one_out": loo_df,
        "loo_mean_rho": loo_mean,
        "loo_min_rho": loo_min,
        "loo_top_match_rate": loo_top_match,
        "bootstrap_summary": boot_summary,
        "bootstrap_n": len(boot_df),
    }

def build_ranking_robustness_table(result: Dict[str, Any]) -> pd.DataFrame:
    ranking = result.get("ranking", {})
    rt = ranking.get("table")
    primary = ranking.get("method")
    if rt is None or rt.empty or not primary:
        return pd.DataFrame()

    alt_col = col_key(rt, "Alternatif", "Alternative")
    primary_top = str(rt.iloc[0][alt_col])
    sens = result.get("sensitivity") or {}
    primary_stab = sens.get("top_stability")

    rows: List[Dict[str, Any]] = []
    comp = result.get("comparison", {}) or {}
    spdf = comp.get("spearman_matrix")
    tops = comp.get("top_alternatives")
    top_map: Dict[str, str] = {}
    if isinstance(tops, pd.DataFrame) and not tops.empty:
        t_m_col = col_key(tops, "Yöntem", "Method")
        t_a_col = col_key(tops, "BirinciAlternatif", "TopAlternative")
        top_map = dict(zip(tops[t_m_col].astype(str), tops[t_a_col].astype(str)))
    top_map.setdefault(primary, primary_top)

    rho_map: Dict[str, float] = {}
    if isinstance(spdf, pd.DataFrame) and not spdf.empty:
        m_col = col_key(spdf, "Yöntem", "Method")
        mat = spdf.set_index(m_col)
        for m in mat.columns:
            if primary in mat.index and m in mat.columns:
                val = mat.loc[primary, m]
                rho_map[str(m)] = float(val) if np.isfinite(val) else np.nan

    methods = list(top_map.keys()) if top_map else [primary]
    if primary not in methods:
        methods = [primary] + methods

    for m in methods:
        rho = rho_map.get(m, 1.0 if m == primary else np.nan)
        top_alt = top_map.get(m, "")
        same_top = str(top_alt) == str(primary_top)
        if np.isfinite(rho):
            if rho >= 0.85:
                mean_txt = tt("Bu yöntem ana sonuca çok yakın davranıyor.", "This method behaves very close to the main result.")
            elif rho >= 0.70:
                mean_txt = tt("Genel tablo benzer, alt sıralarda fark olabilir.", "The overall picture is similar, with possible differences in lower ranks.")
            else:
                mean_txt = tt("Bu yöntem farklı bir bakış sunuyor; sonucu dikkatle karşılaştırın.", "This method gives a different perspective; compare results carefully.")
        else:
            mean_txt = tt("Bu yöntem için yeterli karşılaştırma verisi oluşmadı.", "There is not enough comparison data for this method.")

        top_txt = (
            tt("Lider alternatif ana yöntemle aynı.", "The top alternative is the same as in the main method.")
            if same_top
            else tt("Lider alternatif farklı. Karar öncesi iki yöntemi birlikte okuyun.", "The top alternative is different. Review both methods together before deciding.")
        )
        if m == primary and primary_stab is not None:
            stab_txt = tt(
                f"Monte Carlo'da birincilik koruma oranı %{float(primary_stab)*100:.1f}.",
                f"Monte Carlo first-place retention is {float(primary_stab)*100:.1f}%.",
            )
        else:
            stab_txt = tt("Bu yöntemde Monte Carlo ayrı çalıştırılmadı.", "Monte Carlo was not run separately for this method.")

        rows.append(
            {
                tt("Yöntem", "Method"): m,
                tt("Ana Yöntemle Uyum (ρ)", "Agreement with Primary (ρ)"): (f"{rho:.3f}" if np.isfinite(rho) else "—"),
                tt("Lider Aynı mı?", "Same Top Alternative?"): (tt("Evet", "Yes") if same_top else tt("Hayır", "No")),
                tt("Ne Anlama Geliyor?", "What Does It Mean?"): f"{mean_txt} {top_txt} {stab_txt}",
            }
        )
    return pd.DataFrame(rows)

def get_method_steps(method: str, stage: str) -> str:
    is_en = st.session_state.get("ui_lang") == "EN"

    if is_en:
        weight_steps = {
            "Entropy": "1) Normalize\n2) Compute entropy\n3) Compute divergence\n4) Normalize final weights",
            "CRITIC": "1) Min-max normalization\n2) Compute standard deviations\n3) Build correlation structure\n4) Generate information-content weights",
            "Standart Sapma": "1) Normalize\n2) Compute criterion standard deviations\n3) Convert relative dispersion to weights",
            "Eşit Ağırlık": "1) Assign equal weight to each criterion\n2) Normalize (sum=1)\n3) Proceed to ranking",
            "Manuel Ağırlık": "1) Collect user-defined criterion importance\n2) Normalize (sum=1)\n3) Proceed to ranking",
            "MEREC": "1) Positive transformation\n2) Log-based performance matrix\n3) Criterion-removal effect\n4) Normalize effects to weights",
            "LOPCOW": "1) Normalize\n2) Compute RMS/std terms\n3) Logarithmic percentage variability\n4) Normalize weights",
            "PCA": "1) Standardize\n2) Extract principal components\n3) Combine loadings and explained variance\n4) Produce criterion weights",
        }
        ranking_steps = {
            "TOPSIS": "1) Vector normalization\n2) Weighted matrix\n3) Define PIS/NIS\n4) Distances and closeness score",
            "VIKOR": "1) Best/worst criterion values\n2) Compute S and R measures\n3) Compromise index Q\n4) Lower Q ranks higher",
            "EDAS": "1) Average solution\n2) PDA/NDA distances\n3) Weighted totals\n4) Combined normalized score",
            "CODAS": "1) Normalize\n2) Distances to negative ideal\n3) Euclidean + Manhattan separation\n4) Normalize H score",
            "COPRAS": "1) Column-based normalization\n2) Weighted normalized matrix\n3) Compute S+ and S-\n4) Rank by Q_i",
            "OCRA": "1) Min-max relative performance\n2) Benefit/cost competitiveness components\n3) Weighted total score\n4) Final ranking",
            "ARAS": "1) Add ideal row\n2) Normalize\n3) Weighted totals\n4) Rank by utility ratio K",
            "WASPAS": "1) Normalize\n2) Compute WSM and WPM\n3) Combine with lambda\n4) Higher combined score is better",
            "MOORA": "1) Vector normalization\n2) Weighting\n3) Benefit-cost difference\n4) Rank by score",
            "MABAC": "1) Min-max normalization\n2) Border approximation area\n3) Distance matrices\n4) Rank by total score",
            "MARCOS": "1) Add ideal/anti-ideal rows\n2) Normalize + weight\n3) Utility ratios\n4) Final utility ranking",
            "CoCoSo": "1) Normalize\n2) Compute S and P components\n3) Ka, Kb, Kc compromises\n4) Rank by composite score",
            "PROMETHEE": "1) Pairwise comparisons\n2) Preference indices\n3) Phi+ / Phi- flows\n4) Complete ranking by net flow",
            "GRA": "1) Min-max normalization\n2) Distance to reference\n3) Grey relational coefficients\n4) Rank by relational grade",
        }
        if stage == "weight":
            return weight_steps.get(method, "Detailed steps are available for this weighting method.")
        if method.startswith("Fuzzy "):
            base = method.replace("Fuzzy ", "")
            return "1) Fuzzify data (TFN)\n2) Build defuzzified representation\n3) " + ranking_steps.get(base, "Core ranking steps") + "\n4) Interpret final ranking"
        return ranking_steps.get(method, "Detailed steps are available for this ranking method.")

    weight_steps = {
        "Entropy": "1) Normalize et\n2) Entropi hesapla\n3) Ayrışma derecesi bul\n4) Ağırlıkları normalize et",
        "CRITIC": "1) Min-max normalize et\n2) Standart sapmaları bul\n3) Korelasyon yapısını çıkar\n4) Bilgi içeriğinden ağırlık üret",
        "Standart Sapma": "1) Normalize et\n2) Kriter std sapmalarını hesapla\n3) Std oranlarını ağırlığa çevir",
        "Eşit Ağırlık": "1) Her kritere eşit önem ver\n2) Toplamı 1'e normalize et\n3) Sıralama aşamasına geç",
        "Manuel Ağırlık": "1) Kullanıcıdan kriter önemlerini al\n2) Toplamı 1'e normalize et\n3) Sıralama aşamasına geç",
        "MEREC": "1) Pozitif dönüşüm\n2) Log tabanlı performans\n3) Kriter çıkarma etkisi\n4) Etkiyi normalize et",
        "LOPCOW": "1) Normalize et\n2) RMS ve std oranlarını hesapla\n3) Logaritmik yüzde değişim gücü\n4) Ağırlıkları normalize et",
        "PCA": "1) Standardize et\n2) Temel bileşenleri çıkar\n3) Yükleri ve açıklanan varyansı birleştir\n4) Kriter ağırlıklarını üret",
    }
    ranking_steps = {
        "TOPSIS": "1) Vektör normalize et\n2) Ağırlıklı matris\n3) PIS/NIS belirle\n4) Uzaklıklar ve yakınlık skoru",
        "VIKOR": "1) En iyi/en kötü değerler\n2) S ve R ölçütleri\n3) Q uzlaşı indeksi\n4) Q küçük olan üst sırada",
        "EDAS": "1) Ortalama çözüm\n2) PDA/NDA\n3) Ağırlıklı toplamlar\n4) Normalize birleşik skor",
        "CODAS": "1) Normalize et\n2) Negatif ideal mesafeler\n3) Öklid+Manhattan ayrışımı\n4) H skorunu normalize et",
        "COPRAS": "1) Sütun bazlı normalize et\n2) Ağırlıklı normalize matris\n3) S+ ve S- hesapla\n4) Q_i ile sırala",
        "OCRA": "1) Min-max göreli puanlar\n2) Fayda ve maliyet rekabet bileşenleri\n3) Ağırlıklı toplam skor\n4) Sıralama",
        "ARAS": "1) İdeal satır ekle\n2) Normalize et\n3) Ağırlıklı toplam\n4) Fayda oranı K ile sırala",
        "WASPAS": "1) Normalize et\n2) WSM ve WPM hesapla\n3) λ ile birleştir\n4) Büyük skor daha iyi",
        "MOORA": "1) Vektör normalize et\n2) Ağırlıklandır\n3) Fayda-maliyet farkı\n4) Skoru sırala",
        "MABAC": "1) Min-max normalize et\n2) Sınır yaklaşım alanı\n3) Uzaklık matrisleri\n4) Toplam skorla sırala",
        "MARCOS": "1) İdeal/anti-ideal ekle\n2) Normalize + ağırlıklandır\n3) Yararlılık oranları\n4) Nihai utility ile sırala",
        "CoCoSo": "1) Normalize et\n2) S ve P bileşenleri\n3) Ka, Kb, Kc uzlaşıları\n4) Bileşik skorla sırala",
        "PROMETHEE": "1) İkili karşılaştırma\n2) Tercih indeksleri\n3) Phi+ / Phi- akışları\n4) Net akış ile tam sıralama",
        "GRA": "1) Min-max normalize et\n2) Referansa uzaklık\n3) Gri katsayılar\n4) İlişkisel dereceyi sırala",
    }
    if stage == "weight":
        return weight_steps.get(method, "Ağırlık yöntemi için detay adımlar mevcut.")
    if method.startswith("Fuzzy "):
        base = method.replace("Fuzzy ", "")
        return "1) Veriyi bulanıklaştır (TFN)\n2) Durulaştırılmış temsil üret\n3) " + ranking_steps.get(base, "Temel sıralama adımları") + "\n4) Sonuçları yorumla"
    return ranking_steps.get(method, "Sıralama yöntemi için detay adımlar mevcut.")

def _extract_detail_tables(details: Dict[str, Any]) -> Dict[str, pd.DataFrame]:
    out: Dict[str, pd.DataFrame] = {}
    for key, val in (details or {}).items():
        if isinstance(val, pd.DataFrame) and not val.empty:
            out[key] = val.copy()
        elif isinstance(val, np.ndarray):
            arr = np.asarray(val)
            if arr.ndim == 1 and arr.size > 0:
                out[key] = pd.DataFrame({key: arr})
            elif arr.ndim == 2 and arr.size > 0:
                out[key] = pd.DataFrame(arr)
    return out

def diag_rec_text(rec: Dict[str, str]) -> tuple[str, str]:
    text = rec.get("text", "")
    action = rec.get("action", "")
    if st.session_state.get("ui_lang") != "EN":
        return text, action
    low = text.lower()
    if "sabit kriter" in low:
        return (
            "Constant criterion/criteria detected. These criteria do not provide discrimination power.",
            "Remove constant criteria from the active analysis set.",
        )
    if "yüksek korelasyon" in low:
        return (
            "High correlation detected among criteria, indicating information redundancy.",
            "Prefer correlation-aware methods such as CRITIC/PCA and reconsider overlapping criteria.",
        )
    if "orta düzey korelasyon" in low:
        return (
            "Moderate correlation detected among criteria.",
            "Use conflict-aware weighting methods and review criterion independence.",
        )
    if "yüksek değişkenlik" in low:
        return (
            "High dispersion detected in criteria values.",
            "Entropy or Standard Deviation weighting can effectively capture this discriminative structure.",
        )
    if "dengeli veri yapısı" in low:
        return (
            "Data profile appears balanced in dispersion and correlation.",
            "Most objective weighting and ranking methods can be applied safely.",
        )
    if "alternatif sayısı az" in low:
        return (
            "The number of alternatives is limited.",
            "Interpret Monte Carlo robustness more cautiously and cross-check with additional methods.",
        )
    if "geniş alternatif seti" in low:
        return (
            "A large alternative set is detected.",
            "Distance-based methods and robustness simulations are particularly informative in this setting.",
        )
    if "tüm kriterler fayda" in low:
        return (
            "All criteria are currently marked as benefit criteria.",
            "Review directions and verify whether at least one criterion should be modeled as cost.",
        )
    return (
        "No additional critical anomaly was detected for this item.",
        "Proceed with the current workflow and validate results with sensitivity analysis.",
    )

def reset_all() -> None:
    st.session_state.clear()
    init_state()

def _diag_score_and_label(diag: Dict[str, Any]) -> tuple[int, str]:
    score = 100
    if diag.get("constant_criteria"):
        score -= 35
    max_corr = float(diag.get("max_corr", 0.0))
    if max_corr >= 0.85:
        score -= 30
    elif max_corr >= 0.70:
        score -= 15
    mean_cv = float(diag.get("mean_cv", 0.0))
    if mean_cv < 0.08:
        score -= 10
    if int(diag.get("n_alt", 0)) < 5:
        score -= 10
    score = int(max(0, min(score, 100)))
    if score >= 80:
        return score, tt("Yüksek Uygunluk", "High Suitability")
    if score >= 60:
        return score, tt("Orta Uygunluk", "Moderate Suitability")
    return score, tt("Düşük Uygunluk", "Low Suitability")

# ── MATEMATİKSEL METİN OLUŞTURUCU ──────────────────────────────────────────────
def get_math_formulation(w_method: str, r_methods: List[str]) -> str:
    text = "Kullanılan Yöntemlerin Matematiksel Altyapısı:\n\n"
    
    if ("Entropi" in w_method) or ("Entropy" in w_method):
        text += "Entropi (Entropy) Ağırlıklandırma Yöntemi:\nShannon bilgi kuramına dayanan bu yöntemde, veri setindeki zıtlık ve çeşitlilik ölçülerek nesnel ağırlıklar elde edilir.\n1. Karar matrisi normalize edilir: P_ij = x_ij / Σ x_ij\n2. Her kriter için entropi değeri (e_j) hesaplanır: e_j = -k * Σ [P_ij * ln(P_ij)] (Burada k=1/ln(m))\n3. Çeşitlilik derecesi bulunur: d_j = 1 - e_j\n4. Nihai objektif ağırlıklar elde edilir: w_j = d_j / Σ d_j\n\n"
    elif "CRITIC" in w_method:
        text += "CRITIC Ağırlıklandırma Yöntemi:\nKriterlerin standart sapmasını ve birbiriyle olan korelasyonlarını (çakışmalarını) dikkate alan objektif bir yöntemdir.\n1. Karar matrisi min-max yöntemiyle doğrusal olarak normalize edilir.\n2. Kriterlerin standart sapması (σ_j) hesaplanır.\n3. Kriterler arası korelasyon matrisi (r_jk) oluşturulur.\n4. Kapsanan bilgi miktarı hesaplanır: C_j = σ_j * Σ (1 - r_jk)\n5. Nihai ağırlıklar elde edilir: w_j = C_j / Σ C_j\n\n"
    elif "Standart Sapma" in w_method:
        text += "Standart Sapma Yöntemi:\n1. Karar matrisi doğrusal olarak normalize edilir.\n2. Her bir kriterin standart sapması (σ_j) hesaplanır.\n3. Ağırlıklar, standart sapmaların oransal payına göre belirlenir: w_j = σ_j / Σ σ_j\n\n"
    
    if r_methods:
        text += f"Alternatiflerin sıralanmasında ve performanslarının değerlendirilmesinde {', '.join(r_methods)} algoritması/algoritmaları uygulanmıştır.\n\n"
        for rm in r_methods:
            if "TOPSIS" in rm and "Fuzzy" not in rm:
                text += "TOPSIS Yöntemi:\n1. Karar matrisi vektör normalizasyon yöntemi ile dönüştürülür: r_ij = x_ij / √(Σ x_ij²)\n2. Ağırlıklı matris oluşturulur: v_ij = w_j * r_ij\n3. Pozitif İdeal (A*) ve Negatif İdeal (A⁻) çözümler belirlenir.\n4. İdeallere olan Öklid uzaklıkları (S_i* ve S_i⁻) hesaplanır.\n5. Göreceli yakınlık katsayısı bulunur: C_i* = S_i⁻ / (S_i* + S_i⁻). Katsayısı 1'e en yakın olan alternatif optimum çözüm olarak kabul edilir.\n\n"
            elif "VIKOR" in rm and "Fuzzy" not in rm:
                text += "VIKOR Yöntemi:\n1. Her kriter için en iyi (f_j*) ve en kötü (f_j⁻) değerler tespit edilir.\n2. Grup faydası (S_i) ve Bireysel pişmanlık (R_i) hesaplanır.\n3. VIKOR uzlaşı indeksi hesaplanır: Q_i = v*(S_i - S*) / (S⁻ - S*) + (1-v)*(R_i - R*) / (R⁻ - R*).\n4. Q_i değeri en küçük (minimum) olan alternatif en iyi uzlaşı çözümü olarak sıralanır.\n\n"
            elif "WASPAS" in rm and "Fuzzy" not in rm:
                text += "WASPAS Yöntemi:\n1. Karar matrisi fayda/maliyet yönlerine göre doğrusal normalize edilir.\n2. Ağırlıklı Toplam Modeli (WSM) ve Ağırlıklı Çarpım Modeli (WPM) çalıştırılır.\n3. Birleşik optimizasyon skoru bulunur: Q_i = λ * WSM_i + (1 - λ) * WPM_i. Q değeri en büyük olan alternatif birinci seçilir.\n\n"
            elif "EDAS" in rm:
                text += "EDAS Yöntemi:\n1. Her bir kriterin aritmetik ortalaması (AV) hesaplanır.\n2. Ortalama değere göre Pozitif Uzaklık (PDA) ve Negatif Uzaklık (NDA) hesaplanır.\n3. Bu uzaklıklar kriter ağırlıklarıyla çarpılarak toplam skorlar (SP ve SN) elde edilir.\n4. Normalize edilmiş değerler birleştirilerek Değerlendirme Skoru (AS_i) bulunur.\n\n"
            elif "COPRAS" in rm and "Fuzzy" not in rm:
                text += "COPRAS Yöntemi:\n1. Karar matrisi sütun toplamlarına bölünerek normalize edilir.\n2. Kriter ağırlıkları ile çarpılarak ağırlıklı normalize matris elde edilir.\n3. Fayda kriterleri için S+ ve maliyet kriterleri için S- toplamları hesaplanır.\n4. Göreli önem düzeyi (Q_i) ile sıralama yapılır; büyük Q_i daha iyidir.\n\n"
            elif "OCRA" in rm and "Fuzzy" not in rm:
                text += "OCRA Yöntemi:\n1. Her kriter için min-max tabanlı göreli performans hesaplanır.\n2. Fayda ve maliyet kriterlerinin rekabet bileşen puanları ayrı üretilir.\n3. Bileşenler ağırlıklarla toplanarak operasyonel rekabet skoru elde edilir.\n4. En yüksek toplam skor en iyi alternatifi verir.\n\n"
            elif "ARAS" in rm and "Fuzzy" not in rm:
                text += "ARAS Yöntemi:\n1. Karar matrisine, en iyi performans göstergelerinden oluşan 'Optimum Alternatif (A0)' eklenir.\n2. Matris normalize edilir ve kriter ağırlıkları ile çarpılır.\n3. Her bir alternatifin optimalite fonksiyonu (S_i) hesaplanır.\n4. Fayda derecesi (K_i = S_i / S_0) formülü ile belirlenip sıralama yapılır.\n\n"
            elif "PROMETHEE" in rm:
                text += "PROMETHEE II Yöntemi:\n1. Alternatifler kriter bazında ikili olarak karşılaştırılır.\n2. Her kriter için tercih fonksiyonu ile 0-1 aralığında tercih derecesi hesaplanır.\n3. Kriter ağırlıklarıyla küresel tercih indeksi elde edilir.\n4. Pozitif akış (phi+) ve negatif akış (phi-) hesaplanır.\n5. Net akış (phi = phi+ - phi-) en büyük olan alternatif en üstte sıralanır.\n\n"
            elif "Fuzzy" in rm:
                text += f"{rm} Yöntemi:\nKlasik (kesin) sayıların ölçüm belirsizliği veya insan hata payı barındırabileceği varsayımıyla, veriler tolerans sınırları içeren bulanık (fuzzy) sayılara dönüştürülür. Değerlendirme, üçgensel bulanık kümeler kullanılarak aralıklar üzerinden yapılır ve sonuçlar durulaştırılarak (defuzzification) kesin sıralamaya çevrilir.\n\n"
            
    return text.strip()

# ── EKRAN VE ÇIKTI FONKSİYONLARI ──────────────────────────────────────────────
def render_3layer(layers: Dict[str, str], title: str = "") -> None:
    if not any(layers.values()): return
    if title: st.markdown(f'<p style="font-size:0.8rem;font-weight:700;color:#718096;text-transform:uppercase;">{title}</p>', unsafe_allow_html=True)
    if layers.get("descriptive"): st.markdown(f'<div class="layer-card layer-desc"><div class="layer-label" style="color:#3182CE">🔵 {tt("Tanımlayıcı — Ne bulundu?", "Descriptive — What was found?")}</div><p class="layer-text">{layers["descriptive"]}</p></div>', unsafe_allow_html=True)
    if layers.get("analytic"): st.markdown(f'<div class="layer-card layer-analytic"><div class="layer-label" style="color:#38A169">🟢 {tt("Analitik — Ne anlama gelir?", "Analytic — What does it mean?")}</div><p class="layer-text">{layers["analytic"]}</p></div>', unsafe_allow_html=True)
    if layers.get("normative"): st.markdown(f'<div class="layer-card layer-norm"><div class="layer-label" style="color:#D69E2E">🟡 {tt("Normatif — Ne yapılmalı?", "Normative — What should be done?")}</div><p class="layer-text">{layers["normative"]}</p></div>', unsafe_allow_html=True)

def load_uploaded_file(uploaded_file) -> pd.DataFrame:
    return pd.read_csv(uploaded_file) if uploaded_file.name.lower().endswith(".csv") else pd.read_excel(uploaded_file)

def add_table_to_doc(doc: Document, df: pd.DataFrame, max_rows: int = 25) -> None:
    if df is None or df.empty: return
    use_df = df.copy().head(max_rows)
    table = doc.add_table(rows=1, cols=len(use_df.columns))
    table.style = "Table Grid"
    for i, col in enumerate(use_df.columns): table.rows[0].cells[i].text = str(col)
    for _, row in use_df.iterrows():
        cells = table.add_row().cells
        for j, value in enumerate(row):
            cells[j].text = f"{value:.4f}" if isinstance(value, (float, np.floating)) else str(value)
    doc.add_paragraph("")

def render_table(df: pd.DataFrame, max_rows: int = 300) -> None:
    if not isinstance(df, pd.DataFrame):
        st.write(df)
        return
    if df.empty:
        st.info(tt("Tablo boş.", "Table is empty."))
        return
    view = df.copy()
    if len(view) > max_rows:
        st.caption(tt(f"Tablo ilk {max_rows} satırla sınırlandı.", f"Table limited to first {max_rows} rows."))
        view = view.head(max_rows)
    _height = min(540, max(180, 36 + len(view) * 28))
    st.dataframe(view, use_container_width=True, height=_height)

def set_docx_language(run, lang_code: str = "tr-TR") -> None:
    rpr = run._element.get_or_add_rPr()
    lang = OxmlElement("w:lang")
    lang.set(qn("w:val"), lang_code)
    rpr.append(lang)

def _html_to_plain(text: str) -> str:
    if not text:
        return ""
    txt = text.replace("<br><br>", "\n\n").replace("<br>", "\n")
    txt = re.sub(r"<[^>]+>", "", txt)
    return txt.strip()

def get_math_formulation_en(w_method: str, r_methods: List[str]) -> str:
    lines: List[str] = ["Mathematical and Algorithmic Framework", ""]
    lines.append(f"Objective weighting method: {method_display_name(w_method)}")
    lines.append(get_method_steps(w_method, "weight"))
    lines.append("")
    if r_methods:
        lines.append(f"Ranking methods applied: {', '.join(r_methods)}")
        lines.append("")
        for rm in r_methods:
            lines.append(f"{rm} method:")
            lines.append(get_method_steps(rm, "ranking"))
            lines.append("")
    return "\n".join(lines).strip()

def _build_doc_sections(result: Dict[str, Any], lang: str) -> Dict[str, Any]:
    base = result.get("report_sections", {})
    if lang != "EN":
        return base
    w_method = result["weights"]["method"]
    r_method = result.get("ranking", {}).get("method")
    findings_parts = [
        _html_to_plain(gen_stat_commentary(result)),
        _html_to_plain(gen_weight_commentary(result)),
    ]
    if r_method:
        findings_parts.append(_html_to_plain(gen_ranking_commentary(result, {})))
    findings_parts.append(_html_to_plain(gen_mc_commentary(result)))
    refs = base.get("Kaynakça", []) or base.get("References", [])
    return {
        "Objective of the Study": (
            "This study aims to support transparent and reproducible decision-making by combining "
            "objective weighting and multi-criteria ranking methods."
        ),
        "Philosophy of the Study": (
            "The analytical philosophy is to derive criterion importance from data structure and then evaluate "
            "alternatives through method-specific ranking logic. This provides both methodological traceability "
            "and practical interpretability for non-expert users."
        ),
        "Methodology": (
            f"Weighting method: {method_display_name(w_method)}. "
            f"Primary ranking method: {r_method if r_method else 'Not applied (weights-only mode)'}."
        ),
        "Findings": "\n\n".join([p for p in findings_parts if p]),
        "References": refs,
    }

def generate_apa_docx(result: Dict[str, Any], selected_data: pd.DataFrame, lang: str = "TR") -> bytes | None:
    if not DOCX_AVAILABLE: return None
    doc = Document()
    for section in doc.sections:
        section.top_margin, section.bottom_margin = Inches(1), Inches(1)
        section.left_margin, section.right_margin = Inches(1), Inches(1)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(
        "Çok Kriterli Karar Verme Akademik Analiz Raporu"
        if lang != "EN"
        else "Multi-Criteria Decision-Making Academic Analysis Report"
    )
    title_run.bold = True
    title_run.font.size = Pt(14)
    doc.add_paragraph("")

    sections = _build_doc_sections(result, lang)
    
    # Metodoloji Bölümüne Matematiksel Formülleri Ekleme
    w_meth = result["weights"]["method"]
    r_meths = []
    if result["ranking"]["table"] is not None:
        r_meths.append(result["ranking"]["method"])
    if "comparison" in result and "spearman_matrix" in result["comparison"]:
        comp_df = result["comparison"]["spearman_matrix"]
        if isinstance(comp_df, pd.DataFrame):
            for c in comp_df.columns:
                if c not in r_meths and c != "Yöntem":
                    r_meths.append(c)
    
    if lang == "EN":
        math_text = get_math_formulation_en(w_meth, r_meths)
        sections["Methodology"] = sections.get("Methodology", "") + "\n\n" + math_text
        heading_order = ["Objective of the Study", "Philosophy of the Study", "Methodology", "Findings", "References"]
        ref_heading = "References"
        meth_heading = "Methodology"
        body_lang = "en-US"
    else:
        math_text = get_math_formulation(w_meth, r_meths)
        sections["Metodoloji"] = sections.get("Metodoloji", "") + "\n\n" + math_text
        heading_order = ["Çalışmanın Amacı", "Çalışmanın Felsefesi", "Metodoloji", "Bulgular", "Kaynakça"]
        ref_heading = "Kaynakça"
        meth_heading = "Metodoloji"
        body_lang = "tr-TR"

    for heading in heading_order:
        p = doc.add_paragraph()
        run = p.add_run(heading)
        run.bold = True
        run.font.size = Pt(12)
        set_docx_language(run, body_lang)
        if heading != ref_heading:
            body = doc.add_paragraph(sections.get(heading, ""))
            body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            if body.runs:
                set_docx_language(body.runs[0], body_lang)
            if heading == meth_heading:
                add_table_to_doc(doc, localize_df_lang(selected_data.reset_index().head(15), lang))
            if heading in {"Bulgular", "Findings"}:
                add_table_to_doc(doc, localize_df_lang(result["weights"]["table"], lang))
                if result["ranking"]["table"] is not None:
                    add_table_to_doc(doc, localize_df_lang(result["ranking"]["table"], lang))
        else:
            for ref in sections.get(ref_heading, []):
                doc.add_paragraph(ref)

    doc.add_paragraph("\n")
    sig = doc.add_paragraph()
    sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sig.add_run('Prof. Dr. Ömer Faruk Rençber').italic = True
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()

def generate_excel(result: Dict[str, Any], selected_data: pd.DataFrame, lang: str = "TR") -> bytes:
    output = io.BytesIO()
    sheet_map = (
        {"decision": "Decision_Matrix", "stats": "Statistics", "weights": "Weights", "ranking": "Ranking"}
        if lang == "EN"
        else {"decision": "Karar_Matrisi", "stats": "Istatistikler", "weights": "Agirliklar", "ranking": "Siralama"}
    )
    with pd.ExcelWriter(output, engine="xlsxwriter") as writer:
        localize_df_lang(selected_data, lang).to_excel(writer, sheet_name=sheet_map["decision"])
        localize_df_lang(result["stats"], lang).to_excel(writer, sheet_name=sheet_map["stats"], index=False)
        localize_df_lang(result["weights"]["table"], lang).to_excel(writer, sheet_name=sheet_map["weights"], index=False)
        if result["ranking"]["table"] is not None:
            localize_df_lang(result["ranking"]["table"], lang).to_excel(writer, sheet_name=sheet_map["ranking"], index=False)
        workbook = writer.book
        fmt = workbook.add_format({"bold": True, "bg_color": "#1d3557", "font_color": "white"})
        for ws in writer.sheets.values():
            ws.freeze_panes(1, 0)
            ws.set_row(0, 22, fmt)
            ws.set_column(0, 18, 18)
    return output.getvalue()

_THEME = dict(paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="#F8FAFC", margin=dict(l=20, r=20, t=44, b=20))

def fig_weight_bar(weight_df: pd.DataFrame) -> go.Figure:
    fig = px.bar(weight_df.sort_values("Ağırlık", ascending=False), x="Kriter", y="Ağırlık", text_auto=".4f", color="Ağırlık", color_continuous_scale="Blues")
    fig.update_layout(
        height=400,
        title=tt("Kriter Ağırlıkları", "Criterion Weights"),
        coloraxis_colorbar=dict(title=tt("Ağırlık", "Weight")),
        **_THEME,
    )
    return fig

def fig_rank_bar(rank_df: pd.DataFrame) -> go.Figure:
    fig = px.bar(rank_df.sort_values("Skor"), x="Skor", y="Alternatif", orientation="h", text_auto=".4f", color="Skor", color_continuous_scale="Blues")
    fig.update_layout(
        height=450,
        title=tt("Nihai Sıralama Skorları", "Final Ranking Scores"),
        coloraxis_colorbar=dict(title=tt("Skor", "Score")),
        **_THEME,
    )
    return fig

def fig_box_plots(data: pd.DataFrame, criteria: List[str]) -> go.Figure:
    colors = ["#2E7D9E", "#1B365D", "#48BB78", "#ED8936", "#9F7AEA", "#F56565", "#38B2AC", "#ECC94B"]
    fig = go.Figure()
    for i, c in enumerate(criteria):
        hex_c = colors[i % len(colors)]
        r, g, b = int(hex_c[1:3], 16), int(hex_c[3:5], 16), int(hex_c[5:7], 16)
        
        fig.add_trace(go.Box(
            y=data[c], name=c, boxpoints="all", jitter=0.35, pointpos=-1.6,
            marker=dict(size=5, color=hex_c, opacity=0.65),
            line=dict(color=hex_c),
            fillcolor=f"rgba({r},{g},{b},0.2)",
        ))
    fig.update_layout(height=420, title=tt("Kriter Dağılım Profilleri (Kutu + Nokta)", "Criterion Distribution Profiles (Box + Points)"), showlegend=False, **_THEME)
    return fig

def fig_weight_radar(weight_df: pd.DataFrame) -> go.Figure:
    cats = weight_df["Kriter"].tolist()
    vals = weight_df["Ağırlık"].tolist()
    fig = go.Figure(go.Scatterpolar(
        r=vals + [vals[0]], theta=cats + [cats[0]],
        fill="toself", fillcolor="rgba(27,54,93,0.12)",
        line=dict(color="#1B365D", width=2),
        marker=dict(size=7, color="#2E7D9E"),
        name=tt("Ağırlık", "Weight"),
    ))
    fig.update_layout(
        height=420, title=tt("Ağırlık Dağılımı — Radar Görünümü", "Weight Distribution — Radar View"),
        polar=dict(bgcolor="#F8FAFC", radialaxis=dict(visible=True, gridcolor="#E2E8F0"), angularaxis=dict(gridcolor="#E2E8F0")),
        **_THEME,
    )
    return fig

def fig_parallel_coords(ranking_table: pd.DataFrame, contrib_table: pd.DataFrame, criteria: List[str]) -> go.Figure:
    merged = ranking_table.merge(contrib_table, on="Alternatif", how="inner")
    avail = [c for c in criteria if c in merged.columns]
    if not avail:
        return go.Figure()
    dims = [dict(label=c, values=merged[c].tolist()) for c in avail]
    dims.append(dict(label=tt("Skor", "Score"), values=merged["Skor"].tolist()))
    fig = go.Figure(go.Parcoords(
        line=dict(color=merged["Skor"].tolist(), colorscale="Blues", showscale=True, colorbar=dict(title=tt("Skor", "Score"))),
        dimensions=dims,
    ))
    fig.update_layout(height=460, title=tt("Paralel Koordinat — Alternatif Profilleri", "Parallel Coordinates — Alternative Profiles"), **_THEME)
    return fig

def fig_network_alternatives(ranking_table: pd.DataFrame) -> go.Figure:
    n = len(ranking_table)
    if n == 0:
        return go.Figure()
    alts   = ranking_table["Alternatif"].tolist()
    scores = ranking_table["Skor"].tolist()
    ranks  = ranking_table["Sıra"].tolist() if "Sıra" in ranking_table.columns else list(range(1, n + 1))

    angles = [2 * np.pi * i / n for i in range(n)]
    xs = [np.cos(a) for a in angles]
    ys = [np.sin(a) for a in angles]

    cx, cy = 0.0, 0.0
    edge_x, edge_y = [], []
    for i in range(n):
        edge_x += [xs[i], cx, None]
        edge_y += [ys[i], cy, None]

    for i in range(n):
        j = (i + 1) % n
        edge_x += [xs[i], xs[j], None]
        edge_y += [ys[i], ys[j], None]

    fig = go.Figure()
    fig.add_trace(go.Scatter(
        x=edge_x, y=edge_y, mode="lines",
        line=dict(color="#CBD5E0", width=1.0), hoverinfo="none", showlegend=False,
    ))
    fig.add_trace(go.Scatter(
        x=[cx], y=[cy], mode="markers",
        marker=dict(size=18, color="#1B365D", symbol="star"),
        text=[tt("En Yüksek", "Highest")], hoverinfo="text", showlegend=False,
    ))
    node_sizes = [max(20, 55 - (r - 1) * 3) for r in ranks]
    fig.add_trace(go.Scatter(
        x=xs, y=ys, mode="markers+text",
        marker=dict(
            size=node_sizes,
            color=scores, colorscale="Blues", showscale=True,
            colorbar=dict(title=tt("Skor", "Score"), thickness=12),
            line=dict(color="#1B365D", width=1.5),
        ),
        text=alts, textposition="top center",
        hovertemplate=f"<b>%{{text}}</b><br>{tt('Skor', 'Score')}: %{{marker.color:.4f}}<extra></extra>",
        showlegend=False,
    ))
    fig.update_layout(
        height=520, title=tt("Alternatif Ağ Diyagramı (Dairesel Yerleşim)", "Alternative Network Diagram (Circular Layout)"),
        xaxis=dict(showgrid=False, zeroline=False, showticklabels=False, range=[-1.5, 1.5]),
        yaxis=dict(showgrid=False, zeroline=False, showticklabels=False, range=[-1.5, 1.5]),
        **_THEME,
    )
    return fig

def fig_mc_stability_bubble(mc_df: pd.DataFrame) -> go.Figure:
    sdf = mc_df.sort_values("BirincilikOranı", ascending=False).copy()
    sdf["Yüzde"] = sdf["BirincilikOranı"] * 100
    fig = go.Figure(go.Scatter(
        x=sdf["OrtalamaSıra"], y=sdf["Yüzde"],
        mode="markers+text",
        text=sdf["Alternatif"],
        textposition="top center",
        marker=dict(
            size=sdf["Yüzde"].clip(lower=5) * 1.2,
            color=sdf["Yüzde"],
            colorscale="Blues", showscale=True,
            colorbar=dict(title=tt("Birincilik %", "First-place %")),
            line=dict(color="#1B365D", width=1),
        ),
        hovertemplate=f"<b>%{{text}}</b><br>{tt('Ort. Sıra', 'Mean Rank')}: %{{x:.2f}}<br>{tt('Birincilik', 'First-place')}: %{{y:.1f}}%<extra></extra>",
    ))
    fig.update_layout(
        height=460, title=tt("Monte Carlo — Birincilik Oranı vs Ortalama Sıra (Balon)", "Monte Carlo — First-Place Rate vs Mean Rank (Bubble)"),
        xaxis_title=tt("Ortalama Sıra (düşük = daha iyi)", "Mean Rank (lower = better)"),
        yaxis_title=tt("Birincilik Oranı (%)", "First-Place Rate (%)"),
        **_THEME,
    )
    fig.update_xaxes(autorange="reversed")
    return fig

def fig_mc_rank_bar(mc_df: pd.DataFrame) -> go.Figure:
    sdf = mc_df.sort_values("BirincilikOranı", ascending=False)
    colors_mc = ["#1B365D" if v >= 0.5 else "#2E7D9E" if v >= 0.2 else "#90CDF4" for v in sdf["BirincilikOranı"]]
    fig = go.Figure(go.Bar(
        x=sdf["Alternatif"], y=sdf["BirincilikOranı"] * 100,
        text=[f"%{v*100:.1f}" for v in sdf["BirincilikOranı"]],
        textposition="outside",
        marker_color=colors_mc,
    ))
    fig.update_layout(
        height=420, title=tt("Monte Carlo — Birincilik Oranı (%)", "Monte Carlo — First-Place Rate (%)"),
        yaxis_title=tt("Birincilik Oranı (%)", "First-Place Rate (%)"), xaxis_title="",
        **_THEME,
    )
    return fig

def fig_local_sensitivity(local_df: pd.DataFrame) -> go.Figure:
    cat_order = ["-20%", "-10%", "+10%", "+20%"]
    local_df = local_df.copy()
    crit_col = "Kriter"
    wchg_col = "AğırlıkDeğişimi"
    if st.session_state.get("ui_lang") == "EN":
        local_df = local_df.rename(columns={"Kriter": "Criterion", "AğırlıkDeğişimi": "WeightChange"})
        crit_col = "Criterion"
        wchg_col = "WeightChange"
    local_df[wchg_col] = pd.Categorical(local_df[wchg_col], categories=cat_order, ordered=True)
    fig = px.line(
        local_df.sort_values([crit_col, wchg_col]),
        x=wchg_col, y="SpearmanRho", color=crit_col,
        markers=True, title=tt("Lokal Duyarlılık — Ağırlık Değişimine Tepki", "Local Sensitivity — Response to Weight Changes"),
    )
    fig.update_traces(line_width=2, marker_size=8)
    fig.update_layout(height=420, **_THEME)
    return fig

def render_tab_assistant(commentary: str, key: str = "") -> None:
    with st.expander(tt("🤖 Analiz Asistanı — Akademik Yorum", "🤖 Analysis Assistant — Academic Commentary"), expanded=False):
        st.markdown(f'<div class="tab-assistant">{commentary}</div>', unsafe_allow_html=True)

def _top_cv_crit(result: Dict[str, Any]) -> tuple[str, float]:
    data = result["selected_data"]
    cv = (data.std() / data.mean().abs())
    return cv.idxmax(), float(cv.max())

def _max_corr_pair(result: Dict[str, Any]) -> tuple[str, str, float]:
    corr = result.get("correlation_matrix")
    if corr is None or corr.shape[0] < 2:
        return ("—", "—", 0.0)
    mask = ~np.eye(len(corr), dtype=bool)
    vals = corr.abs().where(mask).stack()
    if vals.empty: return ("—", "—", 0.0)
    pair = vals.idxmax()
    return pair[0], pair[1], float(vals.max())

def gen_stat_commentary(result: Dict[str, Any]) -> str:
    n_alt = result["selected_data"].shape[0]
    n_crit = result["selected_data"].shape[1]
    cv_crit, cv_val = _top_cv_crit(result)
    c1, c2, max_rho = _max_corr_pair(result)
    if st.session_state.get("ui_lang") == "EN":
        line1 = f"Your decision matrix includes <strong>{n_alt} alternatives</strong> and <strong>{n_crit} criteria</strong>, which is suitable for statistical analysis."
        line2 = f"<strong>{cv_crit}</strong> has the highest coefficient of variation (CV ≈ {cv_val:.2f}), indicating strong discriminative power for ranking."
        if max_rho > 0.75:
            line3 = f"High correlation between <strong>{c1}</strong> and <strong>{c2}</strong> (ρ = {max_rho:.2f}) suggests information overlap; CRITIC weighting or criterion reduction should be considered."
        else:
            line3 = f"Maximum inter-criterion correlation is {max_rho:.2f}; multicollinearity risk is limited and does not constrain method selection."
        return f"{line1}<br><br>{line2}<br><br>{line3}"
    line1 = f"Karar matrisiniz <strong>{n_alt} alternatif</strong> ve <strong>{n_crit} kriter</strong> içermekte olup istatistiksel yapı analiz için uygundur."
    line2 = f"<strong>{cv_crit}</strong> kriteri varyasyon katsayısı açısından en yüksek farklılaşmayı sergilemekte (CV ≈ {cv_val:.2f}); bu kriter sıralama kararını en güçlü biçimde ayrıştıran boyut konumundadır."
    if max_rho > 0.75:
        line3 = f"<strong>{c1}</strong> ile <strong>{c2}</strong> arasındaki yüksek korelasyon (ρ = {max_rho:.2f}) bilgi tekrarına işaret etmekte; CRITIC ağırlıklandırması veya kriter elenmesi akademik raporlama için değerlendirilmelidir."
    else:
        line3 = f"Kriterler arası maksimum korelasyon {max_rho:.2f} düzeyinde olup çoklu doğrusallık riski ihmal edilebilir — ağırlıklandırma yöntemi seçiminde herhangi bir metodolojik kısıt bulunmamaktadır."
    return f"{line1}<br><br>{line2}<br><br>{line3}"

def gen_weight_commentary(result: Dict[str, Any]) -> str:
    w_method = result["weights"]["method"]
    w_method_disp = method_display_name(w_method)
    w_vals   = result["weights"]["values"]
    top_k    = max(w_vals, key=w_vals.get)
    top_v    = w_vals[top_k]
    n_crit   = len(w_vals)
    concentration = sum(v**2 for v in w_vals.values())
    if st.session_state.get("ui_lang") == "EN":
        line1 = f"<strong>{w_method_disp}</strong> derived objective weights from data and assigned the highest importance to <strong>{top_k}</strong> (w = {top_v:.4f})."
        if top_v > 0.35:
            line2 = "This concentration indicates notable dependence on a single criterion; this dependency should be reported explicitly."
        else:
            line2 = f"Weight distribution is relatively balanced (Concentration Index ≈ {concentration:.2f}), supporting methodological robustness."
        line3 = f"Sensitivity checks should prioritize perturbations on <strong>{top_k}</strong>, and the rationale for choosing <strong>{w_method_disp}</strong> should be clearly documented."
        return f"{line1}<br><br>{line2}<br><br>{line3}"
    line1 = f"<strong>{w_method}</strong> yöntemi, {n_crit} kriter arasında ağırlıkları verinin kendi iç yapısından türeterek <strong>{top_k}</strong> kriterine en yüksek ağırlığı atamıştır (w = {top_v:.4f})."
    if top_v > 0.35:
        line2 = f"Bu denli yüksek bir yoğunlaşma, kararın büyük ölçüde tek bir kritere dayandığını göstermektedir; sonuçların yorumlanmasında bu bağımlılık açıkça belirtilmelidir."
    else:
        line2 = f"Ağırlık dağılımı görece dengeli görünmekte (Yoğunlaşma İndeksi ≈ {concentration:.2f}), bu da birden fazla kriterin karar üzerinde eşanlı belirleyici role sahip olduğuna ve metodolojik sağlamlığa işaret etmektedir."
    line3 = f"Duyarlılık analizinde {top_k} ağırlığındaki değişime özellikle odaklanılmalı; raporun metodoloji bölümünde {w_method} yönteminin seçim gerekçesi net olarak sunulmalıdır."
    return f"{line1}<br><br>{line2}<br><br>{line3}"

def gen_ranking_commentary(result: Dict[str, Any], alt_names: Dict[str, str] = None) -> str:
    rt = result["ranking"]["table"]
    if rt is None or rt.empty:
        return tt("Sıralama tablosu mevcut değil.", "Ranking table is not available.")
    method   = result["ranking"]["method"]
    alt_col  = col_key(rt, "Alternatif", "Alternative")
    score_col = col_key(rt, "Skor", "Score")
    top_alt  = rt.iloc[0][alt_col]
    top_disp = (alt_names or {}).get(str(top_alt), str(top_alt))
    n_alt    = len(rt)
    score_range = float(rt[score_col].max()) - float(rt[score_col].min())
    if st.session_state.get("ui_lang") == "EN":
        line1 = f"<strong>{method}</strong> ranks <strong>{top_disp}</strong> first among {n_alt} alternatives."
        if score_range < 0.05:
            line2 = f"The score range is narrow ({score_range:.4f}), indicating ranking uncertainty; Monte Carlo robustness and multi-method comparison are strongly recommended."
        else:
            line2 = f"The score range ({score_range:.4f}) indicates clear separation among alternatives and supports methodological confidence."
        line3 = "Parallel coordinates and network diagrams are recommended for transparent academic reporting."
        return f"{line1}<br><br>{line2}<br><br>{line3}"
    line1 = f"<strong>{method}</strong> yöntemi uygulaması sonucunda <strong>{top_disp}</strong> alternatifi {n_alt} seçenek içinde birinci sıraya yerleşmiştir."
    if score_range < 0.05:
        line2 = f"Skor aralığı {score_range:.4f} gibi dar bir bant içinde kaldığından sıralama belirsizliği yüksektir; Monte Carlo sağlamlık testi ve çoklu yöntem karşılaştırması kesinlikle önerilmektedir."
    else:
        line2 = f"Skor aralığının {score_range:.4f} olması alternatiflerin birbirinden net biçimde ayrıştığını göstermekte, bu durum metodolojik güveni artırmaktadır."
    line3 = f"Akademik yayında bu bulguyu desteklemek için paralel koordinat grafiği ve ağ diyagramı görsellerinin sunulması, hakemlerin metodoloji itirazlarını büyük ölçüde azaltacaktır."
    return f"{line1}<br><br>{line2}<br><br>{line3}"

def gen_comparison_commentary(result: Dict[str, Any]) -> str:
    comp = result.get("comparison", {})
    if not comp or "spearman_matrix" not in comp:
        return tt("Yöntem karşılaştırması verisi mevcut değil.", "Method comparison data is not available.")
    spdf = comp["spearman_matrix"]
    if not isinstance(spdf, pd.DataFrame) or spdf.shape[0] < 2:
        return tt("Karşılaştırma için en az iki yöntem gereklidir.", "At least two methods are required for comparison.")
    method_col = col_key(spdf, "Yöntem", "Method")
    mat = spdf.set_index(method_col)
    mask = ~np.eye(mat.shape[0], dtype=bool)
    mean_rho = float(mat.where(mask).stack().mean())
    n_meth = mat.shape[0]
    if st.session_state.get("ui_lang") == "EN":
        line1 = f"Average Spearman correlation across <strong>{n_meth} methods</strong> is <strong>ρ = {mean_rho:.3f}</strong>."
        if mean_rho >= 0.85:
            line2 = "High agreement indicates stable rankings regardless of method family."
        elif mean_rho >= 0.70:
            line2 = "Moderate agreement suggests consensus at upper ranks with some divergence in lower positions."
        else:
            line2 = "Low agreement indicates methodological sensitivity; revisit weighting and normalization assumptions."
        line3 = "Include the Spearman matrix in appendices for transparent methodological justification."
        return f"{line1}<br><br>{line2}<br><br>{line3}"
    line1 = f"Toplam <strong>{n_meth} yöntem</strong> arasındaki ortalama Spearman korelasyonu <strong>ρ = {mean_rho:.3f}</strong> olarak hesaplanmıştır."
    if mean_rho >= 0.85: line2 = f"Bu yüksek uyum, seçilen yöntemlerden bağımsız olarak sıralamanın kararlı kaldığını kanıtlamakta ve bulguların metodolojik sağlamlığını güçlü biçimde desteklemektedir."
    elif mean_rho >= 0.70: line2 = f"Orta düzeydeki bu uyum, yöntemlerin genel sıralamada mutabık kaldığını ancak alt sıralarda görüş ayrılıkları yaşandığını göstermektedir."
    else: line2 = f"Düşük yöntemler-arası uyum ciddi bir metodolojik soruya işaret etmektedir; kriter ağırlıkları veya normalleştirme yaklaşımı yeniden gözden geçirilmeli."
    line3 = f"Spearman matrisi tablo olarak metodoloji ekine eklendiğinde hakem incelemesinde güçlü bir referans niteliği kazanmaktadır."
    return f"{line1}<br><br>{line2}<br><br>{line3}"

def gen_mc_commentary(result: Dict[str, Any]) -> str:
    sens = result.get("sensitivity")
    if not sens:
        return tt("Sağlamlık analizi verisi mevcut değil.", "Robustness analysis data is not available.")
    stab = float(sens.get("top_stability", 0.0))
    mc_df = sens.get("monte_carlo_summary")
    n_iter = int(sens.get("n_iterations", 0))
    top_alt = "—"
    if isinstance(mc_df, pd.DataFrame) and not mc_df.empty:
        mc_alt_col = col_key(mc_df, "Alternatif", "Alternative")
        top_alt = mc_df.iloc[0][mc_alt_col]
    if st.session_state.get("ui_lang") == "EN":
        line1 = f"Monte Carlo simulation ({n_iter:,} scenarios) shows the leading alternative <strong>{top_alt}</strong> remains first in <strong>{stab*100:.1f}%</strong> of perturbations."
        if stab >= 0.80:
            line2 = "This indicates strong robustness against weight uncertainty."
        elif stab >= 0.60:
            line2 = "This indicates moderate robustness; multi-method consensus checks are advisable."
        else:
            line2 = "This indicates low robustness; reconsider weighting assumptions and criterion structure."
        line3 = "Use local sensitivity plots to identify criteria that drive rank instability."
        return f"{line1}<br><br>{line2}<br><br>{line3}"
    line1 = f"Monte Carlo simülasyonu ({n_iter:,} senaryo) sonucunda lider alternatif <strong>{top_alt}</strong> ağırlık şoklarının <strong>%{stab*100:.1f}</strong>'inde birinci sırasını korumuştur."
    if stab >= 0.80: line2 = f"Bu yüksek kararlılık oranı, karar sonucunun ağırlık belirsizliğine karşı güçlü bir direnç sergilediğini ortaya koymaktadır; bulgular güvenle raporlanabilir."
    elif stab >= 0.60: line2 = f"Orta düzey kararlılık, ağırlık kabullerine bağlı olarak sonucun değişebildiğini göstermektedir; çoklu yöntem uzlaşısı eklenmesi önerilir."
    else: line2 = f"Düşük kararlılık oranı modelin ağırlık şoklarına kırılgan olduğuna işaret etmektedir; ağırlıklandırma yöntemi veya kriter yapısı mutlaka yeniden değerlendirilmelidir."
    line3 = f"Lokal duyarlılık grafiği, hangi kriterlerin sıralamayı en çok sarstığını görsel olarak ortaya koymakta; bu bulgunun 'Sınırlılıklar' bölümünde tartışılması önerilir."
    return f"{line1}<br><br>{line2}<br><br>{line3}"

# ---------------------------------------------------------
# GLOBAL BANNER
# ---------------------------------------------------------
st.markdown(
    f"""
    <div class="global-header">
        <div class="header-left">
            <p class="header-dedication">{tt("Çocuklarıma ithaf", "Dedicated to my children")}</p>
        </div>
        <div class="header-signature">Prof. Dr. Ömer Faruk Rençber</div>
        <div class="header-right">www.ofrencber.com</div>
    </div>
    """,
    unsafe_allow_html=True,
)

# ---------------------------------------------------------
# SOL PANEL: KONTROL, AMAC VE TEMİZLİK
# ---------------------------------------------------------
with st.sidebar:
    st.markdown('<div style="height:118px;"></div>', unsafe_allow_html=True)
    st.markdown(
        f"<p class='sidebar-small-note' style='margin:0 0 0.2rem 0; font-weight:700;'>{tt('Dil / Language', 'Language / Dil')}</p>",
        unsafe_allow_html=True,
    )
    _lang_idx = 1 if st.session_state.get("ui_lang") == "EN" else 0
    _lang_pick = st.radio(
        "Language",
        ["Türkçe", "English"],
        index=_lang_idx,
        horizontal=True,
        label_visibility="collapsed",
    )
    _new_lang = "EN" if _lang_pick == "English" else "TR"
    if _new_lang != st.session_state.get("ui_lang"):
        st.session_state["ui_lang"] = _new_lang
        st.rerun()

    with st.expander(tt("📘 Uygulama Rehberi ve Metodolojik Yardım", "📘 User Guide and Methodological Help"), expanded=False):
        _guide_tab, _method_tab = st.tabs([tt("📘 Rehber", "📘 Guide"), tt("🧭 Metodoloji", "🧭 Methodology")])
        with _guide_tab:
            st.markdown(
                f"<p class='sidebar-small-note' style='margin:0;'>{tt('Bu sistemde önce kullanım amacını belirleyin, sonra veri girişini yapın. Kriter yönlerini (fayda/maliyet) doğruladıktan sonra uygun ağırlıklandırma ve sıralama yöntemlerini seçerek analizi çalıştırın. Sonuç ekranında ağırlıklar, sıralama, yöntem karşılaştırması ve duyarlılık bulgularını birlikte yorumlayın.', 'Start by selecting the usage objective, then upload your dataset. After validating criterion directions (benefit/cost), choose weighting and ranking methods and run the analysis. In the results view, interpret weights, ranking, method comparison, and sensitivity findings together.')}</p>",
                unsafe_allow_html=True,
            )
        with _method_tab:
            st.markdown(
                f"<p class='sidebar-small-note' style='margin:0;'>{tt('Objektif ağırlıklandırma yöntemleri kriter önemini verinin yapısından üretir. Sıralama yöntemleri ise bu ağırlıkları kullanarak alternatif performanslarını farklı matematiksel bakış açılarıyla değerlendirir. Belirsizlik içeren durumlarda fuzzy yöntemler tercih edilerek ölçüm oynaklığı modele dahil edilir.', 'Objective weighting methods derive criterion importance from data structure. Ranking methods evaluate alternatives with different mathematical perspectives using these weights. Under uncertainty, fuzzy methods are recommended to model measurement variability explicitly.')}</p>",
                unsafe_allow_html=True,
            )

    is_data_loaded = st.session_state.get("raw_data") is not None

    if st.button(tt("🔄 Yeni Analize Başla (Sıfırla)", "🔄 Start New Analysis (Reset)"), use_container_width=True):
        reset_all()
        st.rerun()

    with st.expander(f"{tt('📂 1. Veri Girişi', '📂 1. Data Input')} {'✅' if is_data_loaded else '⏳'}", expanded=not is_data_loaded):
        uploaded = st.file_uploader(tt("CSV veya XLSX yükleyin", "Upload CSV or XLSX"), type=["csv", "xlsx"], label_visibility="collapsed")
        if st.button(tt("📘 Örnek Veri Kullan", "📘 Use Sample Data"), use_container_width=True):
            st.session_state["raw_data"] = sample_dataset()
            st.session_state["data_source_id"] = "sample_data"
            st.session_state.prep_done = False
            st.rerun()

        if uploaded is not None and st.session_state.get("data_source_id") != uploaded.name:
            st.session_state["raw_data"] = load_uploaded_file(uploaded)
            st.session_state["data_source_id"] = uploaded.name
            st.session_state.prep_done = False
            st.rerun()

    if is_data_loaded:
        with st.expander(tt("🧹 3. Veri Ön İşleme", "🧹 3. Data Preprocessing"), expanded=False):
            missing_strategy = st.selectbox(
                tt("Eksik veri doldurma", "Missing value strategy"),
                [
                    tt("Medyan", "Median"),
                    tt("Ortalama", "Mean"),
                    tt("Interpolasyon", "Interpolation"),
                    tt("Sıfır", "Zero"),
                ],
            )
            clip_outliers = st.checkbox(tt("Aykırı Değerleri (Outlier) Temizle", "Clean Outliers"), value=False)
            st.caption(tt("Ön işleme yalnız sayısal kriter sütunlarına uygulanır.", "Preprocessing is applied only to numeric criterion columns."))
    else:
        missing_strategy, clip_outliers = tt("Medyan", "Median"), False

    st.markdown("---")
    _step_data_done = is_data_loaded
    _step_prep_done = bool(st.session_state.get("prep_done"))
    _step_result_done = st.session_state.get("analysis_result") is not None
    _step_method_active = _step_data_done and _step_prep_done and not _step_result_done
    st.markdown(
        f"""
        <div class="stepper-wrap">
            <div class="step-item"><span class="step-dot {'step-done' if _step_data_done else 'step-pending'}"></span> 1) {tt('Veri girişi', 'Data input')}</div>
            <div class="step-item"><span class="step-dot {'step-done' if _step_prep_done else ('step-active' if _step_data_done else 'step-pending')}"></span> 2) {tt('Kriter doğrulama', 'Criteria validation')}</div>
            <div class="step-item"><span class="step-dot {'step-done' if _step_result_done else ('step-active' if _step_method_active else 'step-pending')}"></span> 3) {tt('Yöntem ve analiz', 'Method and analysis')}</div>
            <div class="step-item"><span class="step-dot {'step-done' if _step_result_done else 'step-pending'}"></span> 4) {tt('Sonuç ve rapor', 'Results and report')}</div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    st.markdown("---")
    st.markdown(
        """
        <div style="text-align:center; margin-top:24px; font-family:'Georgia',serif;
                    font-style:italic; font-weight:bold; color:#1d3557; font-size:1.0rem;">
            Prof. Dr. Ömer Faruk Rençber
        </div>
        <div style="text-align:center; margin-top:4px; font-size:0.75rem; color:#718096;">
            <a href="https://www.ofrencber.com" target="_blank"
               style="color:#2E7D9E; text-decoration:none;">
                🌐 www.ofrencber.com
            </a>
        </div>
        """,
        unsafe_allow_html=True,
    )

# ---------------------------------------------------------
# ANA GÖVDE
# ---------------------------------------------------------
raw_data = st.session_state.get("raw_data")
_needs_ranking_default = bool(st.session_state.get("needs_ranking", True))
_purpose_options = [tt("Sadece Ağırlık Tespiti", "Weights Only"), tt("Ağırlık + Sıralama", "Weights + Ranking")]
_purpose_default_idx = 1 if _needs_ranking_default else 0
_purpose_left, _purpose_right = st.columns([1.35, 1])
with _purpose_right:
    with st.expander(tt("🎯 2. Kullanım Amacı", "🎯 2. Usage Objective"), expanded=True):
        _purpose_choice = st.radio(
            tt("Hedefiniz nedir?", "What is your objective?"),
            _purpose_options,
            index=_purpose_default_idx,
            label_visibility="collapsed",
        )
needs_ranking = _purpose_choice == tt("Ağırlık + Sıralama", "Weights + Ranking")
st.session_state["needs_ranking"] = needs_ranking

if raw_data is None:
    st.markdown(
        f"""
        <div style="text-align:center;padding:5rem 1rem;">
            <h3 style="color:#1B365D;font-weight:700;font-size:1.6rem;">{tt("Karar Destek Sistemine Hoş Geldiniz", "Welcome to the Decision Support System")}</h3>
            <p style="color:#64748B;font-size:1.05rem;">{tt("Bilimsel ve tarafsız analizler için <b>sol panelden veri setinizi yükleyerek</b> ve amacınızı seçerek başlayabilirsiniz.", "For scientific and objective analyses, start by <b>uploading your dataset from the left panel</b> and selecting your objective.")}</p>
        </div>
        """, unsafe_allow_html=True
    )
    st.stop()

# ── Veri Hazırlığı ──
working = raw_data.copy()
working = clean_dataframe(working, missing_strategy, clip_outliers)
working.index = [f"A{idx+1}" for idx in range(len(working))]
numeric_cols = working.select_dtypes(include=[np.number]).columns.tolist()

if len(numeric_cols) < 2:
    st.error(tt("Hata: Analiz için en az iki sayısal sütun (kriter) gereklidir.", "Error: At least two numeric criterion columns are required for analysis."))
    st.stop()

_existing_crits = set(st.session_state.get("crit_dir", {}).keys())
if _existing_crits != set(numeric_cols):
    st.session_state["crit_dir"]     = {c: (guess_direction(c) == "Max (Fayda)") for c in numeric_cols}
    st.session_state["crit_include"] = {c: True for c in numeric_cols}


# ── ASİSTAN BÖLÜMÜ (Ayrı Kapanır Pencere) ──
_sel_temp = [c for c in numeric_cols if st.session_state["crit_include"].get(c, True)]
if len(_sel_temp) >= 2:
    _diag = me.generate_data_diagnostics(working, _sel_temp, {c: ("max" if st.session_state["crit_dir"].get(c, True) else "min") for c in _sel_temp})
    _score, _label = _diag_score_and_label(_diag)
    _badge_cls = "diag-badge-good" if _score >= 80 else ("diag-badge-mid" if _score >= 60 else "diag-badge-bad")
    _rec_items = _diag.get("recommendations", [])[:3]
    _sugg_weight = method_display_name(_diag.get("suggested_weight") or "Entropy")
    _sugg_rank = _diag.get("suggested_ranking") or "TOPSIS / VIKOR"
    while len(_rec_items) < 3:
        _rec_items.append(
            {
                "icon": "•",
                "text": tt("Ek kritik bulgu yok.", "No additional critical finding."),
                "action": tt("Mevcut akışla devam edebilirsiniz.", "You may proceed with the current workflow."),
            }
        )

    with st.expander(tt("🧭 Ön İnceleme Sonuçları", "🧭 Preliminary Review Results"), expanded=True):
        st.markdown(
            f"""
            <div class="assistant-box">
                <strong>{tt("Veri Uygunluk Skoru:", "Data Suitability Score:")}</strong> {_score}/100
                <span class="diag-badge {_badge_cls}">{_label}</span>
                <br><br>
                <div class="assistant-grid">
                    <div class="assistant-card2">
                        <div class="assistant-title2">{tt("Profil", "Profile")}</div>
                        <div class="assistant-body2">{_diag.get('n_alt', 0)} {tt("alternatif", "alternatives")} · {_diag.get('n_crit', 0)} {tt("kriter", "criteria")}<br>{tt("Ort. CV", "Avg CV")}: {_diag.get('mean_cv', 0.0):.2f} · {tt("Maks. |ρ|", "Max |ρ|")}: {_diag.get('max_corr', 0.0):.2f}</div>
                    </div>
                    <div class="assistant-card2">
                        <div class="assistant-title2">{tt("Ağırlık Önerisi", "Weighting Recommendation")}</div>
                        <div class="assistant-body2">{tt("Önerilen yöntem:", "Recommended method:")} <strong>{_sugg_weight}</strong><br>{tt("Bu seçim veri yapısına göre otomatik üretildi.", "This recommendation is generated automatically from your data structure.")}</div>
                    </div>
                    <div class="assistant-card2">
                        <div class="assistant-title2">{tt("Sıralama Önerisi", "Ranking Recommendation")}</div>
                        <div class="assistant-body2">{tt("Önerilen yöntem:", "Recommended method:")} <strong>{_sugg_rank}</strong><br>{tt("Uzlaşı veya mesafe temelli tercih için başlangıç noktası.", "A starting point for compromise- or distance-based ranking.")}</div>
                    </div>
                </div>
            </div>
            """,
            unsafe_allow_html=True,
        )
        for idx, rec in enumerate(_rec_items, start=1):
            _r_text, _r_action = diag_rec_text(rec)
            st.markdown(
                f"""
                <div class="assistant-card2">
                    <div class="assistant-title2">{rec.get("icon", "•")} {tt("Bulgu", "Finding")} {idx}</div>
                    <div class="assistant-body2">{_r_text}<br><strong>{tt("Öneri:", "Recommendation:")}</strong> {_r_action}</div>
                </div>
                """,
                unsafe_allow_html=True,
            )
        st.markdown(f"**{tt('Öneri Uygulama Sonuçları', 'Recommendation Application Results')}**")
        st.markdown(
            f"""
            <div class="assistant-box" style="background:#EDF7FF;">
                <div class="assistant-title2">{tt("Ne yaptık?", "What did we do?")}</div>
                <div class="assistant-body2">
                    {tt(
                        "Veri yapısını önce tarafsız biçimde taradık: değişkenlik, kriterler arası benzerlik ve alternatif sayısının analize uygunluğunu kontrol ettik. Sonrasında veri karakterine en uygun objektif ağırlıklandırma ve sıralama yaklaşımını önerdik.",
                        "We first screened the data structure objectively: variation, inter-criterion similarity, and whether the number of alternatives is suitable for analysis. Then we recommended the most suitable objective weighting and ranking approaches for this data profile."
                    )}
                </div>
                <br>
                <div class="assistant-title2">{tt("Ne bulduk?", "What did we find?")}</div>
                <div class="assistant-body2">
                    {tt(
                        "Bu ön inceleme, hangi kriterlerin bilgiyi tekrar ettiği, hangilerinin ayırt edici güç taşıdığı ve kararın hangi yöntem felsefesiyle daha güvenilir okunabileceği konusunda yol haritası verir. Bu nedenle öneriler, sadece teknik seçenek değil; kararın nasıl daha savunulabilir hale geleceğini gösteren metodolojik bir çerçevedir.",
                        "This preliminary review provides a roadmap on which criteria repeat information, which criteria carry discriminative power, and which method philosophy gives a more reliable reading. So these recommendations are not only technical options; they are a methodological framework showing how to make the decision more defensible."
                    )}
                    <br><br>
                    <strong>{tt("Önerilen Ağırlık Felsefesi:", "Recommended Weighting Philosophy:")}</strong> {_sugg_weight}
                    <br>
                    <strong>{tt("Önerilen Sıralama Felsefesi:", "Recommended Ranking Philosophy:")}</strong> {_sugg_rank}
                </div>
            </div>
            """,
            unsafe_allow_html=True,
        )


# ── VERİ ÖN İZLEME VE KRİTER YAPILANDIRMASI BÖLÜMÜ (Ayrı Kapanır Pencere) ──
with st.expander(tt("📊 Veri Ön İzleme ve Kriter Yapılandırması", "📊 Data Preview and Criteria Configuration"), expanded=False):
    col_prev, col_conf = st.columns([1.3, 1])

    with col_prev:
        st.markdown(f"**📋 {tt('Veri Ön İzleme (İlk 10 Satır)', 'Data Preview (First 10 Rows)')}**")
        render_table(working.head(10))

    with col_conf:
        st.markdown(f"**⚙️ {tt('Kriter Seçimi ve Yönü (Fayda / Maliyet)', 'Criteria Selection and Direction (Benefit / Cost)')}**")
        st.markdown(
            '<p style="font-size:0.75rem;color:#718096;margin:0 0 0.5rem 0">'
            + tt("Dahil/çıkar için onay kutusu · Yön için hedefe tıklayın.", "Use checkbox to include/exclude · click to set direction.")
            + '</p>',
            unsafe_allow_html=True,
        )
        for c in numeric_cols:
            row_c = st.columns([0.15, 1.8, 2.05])
            with row_c[0]:
                st.session_state["crit_include"][c] = st.checkbox(
                    "", value=st.session_state["crit_include"].get(c, True), key=f"inc_{c}"
                )
            with row_c[1]:
                st.markdown(
                    f'<div style="padding-top:0.45rem; font-size:0.87rem; font-weight:600; color:#2D3748;">{c}</div>',
                    unsafe_allow_html=True
                )
            with row_c[2]:
                current_dir = tt("⬆ Fayda", "⬆ Benefit") if st.session_state["crit_dir"].get(c, True) else tt("⬇ Maliyet", "⬇ Cost")
                choice = st.radio(
                    "", [tt("⬆ Fayda", "⬆ Benefit"), tt("⬇ Maliyet", "⬇ Cost")],
                    index=0 if current_dir == tt("⬆ Fayda", "⬆ Benefit") else 1,
                    key=f"dir_{c}", horizontal=True, label_visibility="collapsed"
                )
                st.session_state["crit_dir"][c] = (choice == tt("⬆ Fayda", "⬆ Benefit"))

criteria = [c for c in numeric_cols if st.session_state["crit_include"].get(c, True)]
criteria_types = {c: ("max" if st.session_state["crit_dir"].get(c, True) else "min") for c in criteria}

if len(criteria) < 2:
    st.error(tt("En az 2 kriter seçmelisiniz.", "You must select at least 2 criteria."))
    st.stop()

if not st.session_state.prep_done:
    if st.button(tt("✅ Veri Ön İşleme Bitti (Yöntem Seçimine Geç)", "✅ Preprocessing Complete (Proceed to Method Selection)"), type="primary", use_container_width=True):
        st.session_state.prep_done = True
        st.rerun()
    st.stop()
else:
    if st.button(tt("🔄 Kriter Ayarlarına Geri Dön", "🔄 Back to Criteria Settings"), use_container_width=True):
        st.session_state.prep_done = False
        st.session_state["analysis_result"] = None
        st.rerun()

# ---------------------------------------------------------
# YÖNTEM SEÇİMİ (ALT PANEL)
# ---------------------------------------------------------
st.divider()
st.markdown(f"<h3 style='font-size: 1.2rem;'>⚙️ {tt('Yöntem Seçimi ve Karşılaştırma', 'Method Selection and Comparison')}</h3>", unsafe_allow_html=True)

methods_internal = me.OBJECTIVE_WEIGHT_METHODS
methods_display = [method_display_name(m) for m in methods_internal]
_default_weight = st.session_state.get("weight_method_pref")
_default_weight_display = method_display_name(_default_weight) if _default_weight else methods_display[0]
_weight_idx = methods_display.index(_default_weight_display) if _default_weight_display in methods_display else 0

weight_mode = st.radio(
    f"**1. {tt('Ağırlıklandırma Modu', 'Weighting Mode')}**",
    [tt("🎯 Objektif Ağırlık", "🎯 Objective Weights"), tt("⚖️ Eşit Ağırlık", "⚖️ Equal Weights"), tt("✍️ Manuel Ağırlık", "✍️ Manual Weights")],
    horizontal=True,
)

manual_weights: Dict[str, float] | None = None
manual_weights_valid = True
if "Objektif" in weight_mode or "Objective" in weight_mode:
    weight_method_display = st.radio(
        f"**{tt('1.1 Ağırlıklandırma Yöntemini Seçin (Sadece Bir Tane):', '1.1 Select Weighting Method (Choose One):')}**",
        methods_display,
        index=_weight_idx,
        horizontal=True,
        help=tt(
            "Kriterlerin sonuca etki gücünü (önemini) nesnel verilerle belirler. Entropi verideki çeşitliliği, CRITIC ise çeşitlilik ve zıtlığı (korelasyonu) baz alır.",
            "Determines criterion importance objectively. Entropy uses dispersion; CRITIC combines dispersion and conflict (correlation).",
        )
    )
    weight_method = method_internal_name(weight_method_display)
    weight_mode_key = "objective"
elif "Eşit" in weight_mode or "Equal" in weight_mode:
    weight_method = "Eşit Ağırlık"
    weight_mode_key = "equal"
    st.caption(tt("Tüm kriterlere eşit önem atanır.", "All criteria are assigned equal importance."))
else:
    weight_method = "Manuel Ağırlık"
    weight_mode_key = "manual"
    with st.expander(tt("✍️ Manuel Ağırlık Girişi", "✍️ Manual Weight Input"), expanded=True):
        st.caption(tt("Kriter değerlerini serbest girin; sistem otomatik olarak 1'e normalize eder.", "Enter raw values freely; the system automatically normalizes them to sum to 1."))
        _mcols = st.columns(min(3, max(1, len(criteria))))
        manual_raw = {}
        for i, c in enumerate(criteria):
            with _mcols[i % len(_mcols)]:
                manual_raw[c] = st.number_input(
                    f"{c}",
                    min_value=0.0,
                    value=float(st.session_state.get(f"manual_w_{c}", 1.0)),
                    step=0.1,
                    key=f"manual_w_{c}",
                )
        _sum_raw = sum(manual_raw.values())
        st.caption(f"{tt('Toplam', 'Total')}: {_sum_raw:.6f}")
        if _sum_raw <= 0:
            st.warning(tt("Toplam sıfır olamaz. Sistem geçici olarak eşit ağırlık uygular.", "Sum cannot be zero. System temporarily applies equal weights."))
            manual_weights = {c: 1.0 / len(criteria) for c in criteria}
            manual_weights_valid = False
        else:
            manual_weights = {c: float(v) for c, v in manual_raw.items()}
            if not np.isclose(_sum_raw, 1.0, rtol=0.0, atol=1e-6):
                st.warning(
                    tt(
                        "Manuel ağırlıklarda toplam 1 olmalıdır. Lütfen değerleri toplam 1 olacak şekilde düzeltin.",
                        "Manual weights must sum to 1. Please adjust values so the total equals 1.",
                    )
                )
                manual_weights_valid = False
        render_table(pd.DataFrame({"Kriter": list(manual_weights.keys()), "Ağırlık": list(manual_weights.values())}))

ranking_methods_selected = []
primary_rank_method: str | None = None
vikor_v = float(st.session_state.get("vikor_v", 0.5))
waspas_lambda = float(st.session_state.get("waspas_lambda", 0.5))
codas_tau = float(st.session_state.get("codas_tau", 0.02))
cocoso_lambda = float(st.session_state.get("cocoso_lambda", 0.5))
gra_rho = float(st.session_state.get("gra_rho", 0.5))
promethee_pref_func = st.session_state.get("promethee_pref_func", "linear")
promethee_q = float(st.session_state.get("promethee_q", 0.05))
promethee_p = float(st.session_state.get("promethee_p", 0.30))
promethee_s = float(st.session_state.get("promethee_s", 0.20))
fuzzy_spread = float(st.session_state.get("fuzzy_spread", 0.10))
sensitivity_iterations = int(st.session_state.get("sensitivity_iterations", 400))
sensitivity_sigma = float(st.session_state.get("sensitivity_sigma", 0.12))

if needs_ranking:
    layer_choice = st.radio(
        f"**2. {tt('Analiz Katmanı:', 'Analysis Layer:')}**",
        [tt("🎯 Klasik Mantık (Kesin Değerler)", "🎯 Classical (Crisp Values)"), tt("🌫️ Fuzzy Mantık (Bulanık/Esnek Değerler)", "🌫️ Fuzzy (Uncertain/Flexible Values)")],
        horizontal=True,
        help=tt("Klasik katman kesin sayısal değerlerle, Fuzzy katman ise belirsizlik toleransı (spread) ile çalışır.", "Classical uses crisp numeric values; fuzzy uses uncertainty tolerance (spread)."),
    )
    st.write(f"**3. {tt('Sıralama Yöntemlerini Seçin (Karşılaştırmak İçin Birden Fazla Seçebilirsiniz):', 'Select Ranking Methods (You Can Choose Multiple for Comparison):')}**")
    if "Klasik" in layer_choice or "Classical" in layer_choice:
        all_ranks = me.CLASSICAL_MCDM_METHODS
        def_method = "TOPSIS"
    else:
        all_ranks = me.FUZZY_MCDM_METHODS
        def_method = "Fuzzy TOPSIS"

    _pref_ranks = st.session_state.get("ranking_prefs") or []
    _has_pref_in_layer = any(m in all_ranks for m in _pref_ranks)
    rank_cols = st.columns(4)
    for i, rank_method in enumerate(all_ranks):
        with rank_cols[i % 4]:
            if _pref_ranks and _has_pref_in_layer:
                default_val = rank_method in _pref_ranks
            else:
                default_val = (rank_method == def_method)
            if st.checkbox(rank_method, value=default_val, key=f"rank_cb_{rank_method}", help=tt(f"{rank_method} yöntemini analize dahil etmek için tıklayınız. Birden fazla seçim sonuçları karşılaştırmanızı sağlar.", f"Click to include {rank_method} in the analysis. Multiple selections enable comparison.")):
                ranking_methods_selected.append(rank_method)
                
    if not ranking_methods_selected:
        st.warning(tt("Lütfen en az bir sıralama yöntemi seçiniz.", "Please select at least one ranking method."))
        st.stop()

    if len(ranking_methods_selected) == 1:
        primary_rank_method = ranking_methods_selected[0]
        st.caption(f"{tt('Ana sıralama yöntemi', 'Primary ranking method')}: {primary_rank_method}")
    else:
        primary_rank_method = st.selectbox(
            tt("🎯 Ana Çalıştırılacak Sıralama Yöntemi", "🎯 Primary Ranking Method to Run"),
            ranking_methods_selected,
            index=0,
            help=tt("Çıktı ekranında ana sıralama tablosu ve lider alternatif bu yönteme göre oluşturulur.", "The main ranking table and leading alternative are produced with this method."),
        )

    uses_vikor = ("VIKOR" in ranking_methods_selected) or ("Fuzzy VIKOR" in ranking_methods_selected)
    uses_waspas = ("WASPAS" in ranking_methods_selected) or ("Fuzzy WASPAS" in ranking_methods_selected)
    uses_codas = ("CODAS" in ranking_methods_selected) or ("Fuzzy CODAS" in ranking_methods_selected)
    uses_cocoso = ("CoCoSo" in ranking_methods_selected) or ("Fuzzy CoCoSo" in ranking_methods_selected)
    uses_gra = ("GRA" in ranking_methods_selected) or ("Fuzzy GRA" in ranking_methods_selected)
    uses_promethee = ("PROMETHEE" in ranking_methods_selected) or ("Fuzzy PROMETHEE" in ranking_methods_selected)
    uses_fuzzy = any(m.startswith("Fuzzy") for m in ranking_methods_selected)
    rec = recommend_parameter_defaults(ranking_methods_selected, len(working), len(criteria), uses_fuzzy)

    with st.expander(tt("⚙️ 4. Parametre Girişi (Kapanır)", "⚙️ 4. Parameter Input (Collapsible)"), expanded=False):
        param_mode = st.radio(
            tt("Parametre Modu", "Parameter Mode"),
            [tt("🎯 Önerilen Varsayılan", "🎯 Recommended Default"), tt("✍️ Manuel Değiştir", "✍️ Manual Override")],
            horizontal=True,
            key="param_mode_choice",
        )
        if not any([uses_vikor, uses_waspas, uses_codas, uses_cocoso, uses_gra, uses_promethee, uses_fuzzy]):
            st.caption(tt("Seçili sıralama yöntemleri ek parametre gerektirmiyor.", "Selected ranking methods do not require extra parameters."))
        if ("Önerilen" in param_mode) or ("Recommended" in param_mode):
            vikor_v = float(rec["vikor_v"])
            waspas_lambda = float(rec["waspas_lambda"])
            codas_tau = float(rec["codas_tau"])
            cocoso_lambda = float(rec["cocoso_lambda"])
            gra_rho = float(rec["gra_rho"])
            promethee_pref_func = str(rec.get("promethee_pref_func", "linear"))
            promethee_q = float(rec.get("promethee_q", 0.05))
            promethee_p = float(rec.get("promethee_p", 0.30))
            promethee_s = float(rec.get("promethee_s", 0.20))
            fuzzy_spread = float(rec["fuzzy_spread"])
            sensitivity_iterations = int(rec["sensitivity_iterations"])
            sensitivity_sigma = float(rec["sensitivity_sigma"])
            st.caption(tt("Önerilen değerler otomatik uygulanır. İsterseniz Manuel Değiştir moduna geçebilirsiniz.", "Recommended values are applied automatically. Switch to Manual Override if needed."))
            rec_rows = []
            if uses_vikor: rec_rows.append({tt("Parametre", "Parameter"): "VIKOR v", tt("Değer", "Value"): vikor_v})
            if uses_waspas: rec_rows.append({tt("Parametre", "Parameter"): "WASPAS λ", tt("Değer", "Value"): waspas_lambda})
            if uses_codas: rec_rows.append({tt("Parametre", "Parameter"): "CODAS τ", tt("Değer", "Value"): codas_tau})
            if uses_cocoso: rec_rows.append({tt("Parametre", "Parameter"): "CoCoSo λ", tt("Değer", "Value"): cocoso_lambda})
            if uses_gra: rec_rows.append({tt("Parametre", "Parameter"): "GRA ρ", tt("Değer", "Value"): gra_rho})
            if uses_promethee:
                rec_rows.extend([
                    {tt("Parametre", "Parameter"): "PROMETHEE pref_func", tt("Değer", "Value"): promethee_pref_func},
                    {tt("Parametre", "Parameter"): "PROMETHEE q", tt("Değer", "Value"): promethee_q},
                    {tt("Parametre", "Parameter"): "PROMETHEE p", tt("Değer", "Value"): promethee_p},
                    {tt("Parametre", "Parameter"): "PROMETHEE s", tt("Değer", "Value"): promethee_s},
                ])
            if uses_fuzzy: rec_rows.append({tt("Parametre", "Parameter"): "Fuzzy spread", tt("Değer", "Value"): fuzzy_spread})
            rec_rows.append({tt("Parametre", "Parameter"): tt("Monte Carlo iterasyon", "Monte Carlo iterations"), tt("Değer", "Value"): sensitivity_iterations})
            rec_rows.append({tt("Parametre", "Parameter"): "Monte Carlo sigma", tt("Değer", "Value"): sensitivity_sigma})
            rec_df = pd.DataFrame(rec_rows)
            render_table(rec_df)
        else:
            if uses_fuzzy:
                fuzzy_spread = st.slider(tt("Fuzzy — spread (belirsizlik bandı)", "Fuzzy — spread (uncertainty band)"), 0.01, 0.50, float(fuzzy_spread), 0.01)
            if uses_vikor:
                vikor_v = st.slider(tt("VIKOR — v (uzlaşı katsayısı)", "VIKOR — v (compromise factor)"), 0.0, 1.0, float(vikor_v), 0.01)
            if uses_waspas:
                waspas_lambda = st.slider(tt("WASPAS — λ (hibrit katsayı)", "WASPAS — λ (hybrid coefficient)"), 0.0, 1.0, float(waspas_lambda), 0.01)
            if uses_codas:
                codas_tau = st.slider(tt("CODAS — τ (eşik değeri)", "CODAS — τ (threshold)"), 0.0, 0.20, float(codas_tau), 0.005)
            if uses_cocoso:
                cocoso_lambda = st.slider(tt("CoCoSo — λ (birleşim katsayısı)", "CoCoSo — λ (aggregation coefficient)"), 0.0, 1.0, float(cocoso_lambda), 0.01)
            if uses_gra:
                gra_rho = st.slider(tt("GRA — ρ (ayırt edici katsayı)", "GRA — ρ (distinguishing coefficient)"), 0.10, 0.90, float(gra_rho), 0.01)
            if uses_promethee:
                promethee_pref_func = st.selectbox(
                    tt("PROMETHEE tercih fonksiyonu", "PROMETHEE preference function"),
                    ["linear", "usual", "u_shape", "v_shape", "level", "gaussian"],
                    index=["linear", "usual", "u_shape", "v_shape", "level", "gaussian"].index(str(promethee_pref_func) if str(promethee_pref_func) in ["linear", "usual", "u_shape", "v_shape", "level", "gaussian"] else "linear"),
                )
                if promethee_pref_func in {"u_shape", "level", "linear"}:
                    promethee_q = st.slider("PROMETHEE q", 0.0, 1.0, float(promethee_q), 0.01)
                if promethee_pref_func in {"v_shape", "level", "linear"}:
                    promethee_p = st.slider("PROMETHEE p", 0.0, 1.0, float(promethee_p), 0.01)
                if promethee_pref_func == "gaussian":
                    promethee_s = st.slider("PROMETHEE s", 0.01, 1.0, float(promethee_s), 0.01)
            st.markdown("---")
            sensitivity_iterations = st.slider(tt("Monte Carlo iterasyon", "Monte Carlo iterations"), 100, 5000, int(sensitivity_iterations), 100)
            sensitivity_sigma = st.slider(tt("Monte Carlo sigma", "Monte Carlo sigma"), 0.01, 0.50, float(sensitivity_sigma), 0.01)

    st.session_state["vikor_v"] = float(vikor_v)
    st.session_state["waspas_lambda"] = float(waspas_lambda)
    st.session_state["codas_tau"] = float(codas_tau)
    st.session_state["cocoso_lambda"] = float(cocoso_lambda)
    st.session_state["gra_rho"] = float(gra_rho)
    st.session_state["promethee_pref_func"] = str(promethee_pref_func)
    st.session_state["promethee_q"] = float(promethee_q)
    st.session_state["promethee_p"] = float(promethee_p)
    st.session_state["promethee_s"] = float(promethee_s)
    st.session_state["fuzzy_spread"] = float(fuzzy_spread)
    st.session_state["sensitivity_iterations"] = int(sensitivity_iterations)
    st.session_state["sensitivity_sigma"] = float(sensitivity_sigma)

    index_options = [tt("Satır Sırası (A1, A2...)", "Row Order (A1, A2...)")] + raw_data.select_dtypes(include=["object", "string"]).columns.tolist()
    idx_col = st.selectbox(
        f"**5. {tt('Alternatif İsimlendirme:', 'Alternative Naming:')}** {tt('Değerlendirdiğiniz alternatiflerin ID Bilgileri veya İsimleri hangi sütunda?', 'Which column contains alternative IDs or names?')}",
        index_options,
        help=tt("Analiz sonuçlarında, grafiklerde ve tablolarda A1, A2 yerine gerçek alternatif isimlerinin görünmesini istiyorsanız, isimlerin bulunduğu uygun sütunu seçin.", "Select the column containing names if you want real alternative names instead of A1, A2 in outputs.")
    )

    if idx_col != tt("Satır Sırası (A1, A2...)", "Row Order (A1, A2...)"):
        working = working.set_index(idx_col)
    else:
        with st.expander(tt("✏️ Alternatif İsimlerini Özelleştir (İsteğe Bağlı)", "✏️ Customize Alternative Names (Optional)"), expanded=False):
            st.markdown(
                '<p style="font-size:0.8rem;color:#718096;margin:0 0 0.5rem 0">'
                + tt("Analizde görünen kısa kodların (A1, A2…) yerine gerçek isimler yazın. Boş bırakırsanız orijinal isim korunur.",
                     "Replace short codes (A1, A2...) with real names. Leave blank to keep originals.")
                + '</p>',
                unsafe_allow_html=True,
            )
            current_alts = list(working.index)
            n_cols_alt = min(4, len(current_alts))
            alt_col_grid = st.columns(n_cols_alt)
            for i, alt in enumerate(current_alts):
                with alt_col_grid[i % n_cols_alt]:
                    entered = st.text_input(
                        f"{alt}", value=st.session_state["alt_names"].get(str(alt), ""),
                        key=f"altname_{i}", placeholder=tt("Özel isim…", "Custom name..."),
                    )
                    if entered:
                        st.session_state["alt_names"][str(alt)] = entered
                    elif str(alt) in st.session_state["alt_names"] and not entered:
                        del st.session_state["alt_names"][str(alt)]

if st.button(tt("🚀 Analiz Zamanı", "🚀 Run Analysis"), type="primary", use_container_width=True):
    if ("Manuel" in weight_mode or "Manual" in weight_mode) and (not manual_weights_valid):
        st.error(tt("Manuel ağırlık toplamı 1 olmalıdır. Lütfen düzeltip tekrar deneyin.", "Manual weight total must be 1. Please correct and try again."))
        st.stop()
    with st.spinner(tt("Laboratuvar çalışıyor, veriler işleniyor...", "Running analysis... processing data...")):
        st.session_state.docx_buffer = None
        main_rank = primary_rank_method if primary_rank_method else (ranking_methods_selected[0] if ranking_methods_selected else None)
        comp_ranks = ranking_methods_selected if len(ranking_methods_selected) > 1 else []
        if main_rank and main_rank not in ranking_methods_selected:
            st.error(tt("Ana sıralama yöntemi seçimi geçersiz. Lütfen yöntemleri yeniden seçin.", "Invalid primary ranking selection. Please reselect methods."))
            st.stop()
        st.session_state["weight_method_pref"] = weight_method
        st.session_state["ranking_prefs"] = ranking_methods_selected

        config = me.AnalysisConfig(
            criteria=criteria, criteria_types=criteria_types,
            weight_method=weight_method,
            weight_mode=weight_mode_key,
            manual_weights=manual_weights,
            ranking_method=main_rank,
            compare_methods=comp_ranks,
            vikor_v=vikor_v,
            waspas_lambda=waspas_lambda,
            codas_tau=codas_tau,
            cocoso_lambda=cocoso_lambda,
            gra_rho=gra_rho,
            promethee_pref_func=promethee_pref_func,
            promethee_q=promethee_q,
            promethee_p=promethee_p,
            promethee_s=promethee_s,
            fuzzy_spread=fuzzy_spread,
            sensitivity_iterations=sensitivity_iterations,
            sensitivity_sigma=sensitivity_sigma,
        )

        start = time.time()
        try:
            result = me.run_full_analysis(working, config)
            if main_rank:
                actual_method = result.get("ranking", {}).get("method")
                if actual_method != main_rank:
                    raise ValueError(tt(f"Yöntem uyumsuzluğu: seçilen '{main_rank}', çalışan '{actual_method}'.", f"Method mismatch: selected '{main_rank}', executed '{actual_method}'."))
                rt_chk = result.get("ranking", {}).get("table")
                if rt_chk is None or rt_chk.empty:
                    raise ValueError(tt("Sıralama çıktısı boş döndü.", "Ranking output is empty."))
                if rt_chk["Skor"].isna().any() or (~np.isfinite(rt_chk["Skor"])).any():
                    raise ValueError(tt(f"{main_rank} yöntemi geçersiz skor üretti (NaN/Inf).", f"{main_rank} produced invalid scores (NaN/Inf)."))
                if "Sıra" not in rt_chk.columns or rt_chk["Sıra"].isna().any():
                    raise ValueError(tt(f"{main_rank} yöntemi sıra kolonunu üretemedi.", f"{main_rank} did not produce a valid rank column."))
        except Exception as exc:
            st.error(f"{tt('Analiz Hatası', 'Analysis Error')}: {exc}")
            st.stop()

        result["analysis_time"] = time.time() - start
        result["selected_data"] = working[criteria].copy()
        if weight_mode_key == "objective":
            try:
                result["weight_robustness"] = compute_weight_robustness(
                    data=working[criteria].copy(),
                    criteria=criteria,
                    criteria_types=criteria_types,
                    weight_method=weight_method,
                    base_weights=result["weights"]["values"],
                    bootstrap_n=180,
                )
            except Exception as _w_exc:
                result["weight_robustness"] = {"error": str(_w_exc)}
        else:
            result["weight_robustness"] = {
                "info": tt(
                    "Sağlamlık testi objektif ağırlık yöntemleri için tasarlanmıştır.",
                    "Robustness test is designed for objective weighting methods.",
                )
            }
        st.session_state["analysis_result"] = result
        if DOCX_AVAILABLE:
            st.session_state["report_docx"] = generate_apa_docx(
                result,
                working[criteria].copy(),
                lang=st.session_state.get("ui_lang", "TR"),
            )

# ---------------------------------------------------------
# SONUÇLARIN GÖSTERİMİ
# ---------------------------------------------------------
result = st.session_state.get("analysis_result")
if result is None:
    st.stop()

st.markdown(f"<h3 style='font-size: 1.2rem; margin-top:1rem;'>📥 {tt('Analiz Sonuçları ve Raporlar', 'Analysis Results and Reports')}</h3>", unsafe_allow_html=True)

weight_df     = result["weights"]["table"]
ranking_table = result.get("ranking", {}).get("table")
alt_names     = st.session_state.get("alt_names", {})

_lead_alt = "—"
if isinstance(ranking_table, pd.DataFrame) and not ranking_table.empty:
    _alt_col = col_key(ranking_table, "Alternatif", "Alternative")
    _raw_lead = str(ranking_table.iloc[0][_alt_col])
    _lead_alt = alt_names.get(_raw_lead, _raw_lead)
_rank_method = result.get("ranking", {}).get("method") or "—"
_sensitivity_payload = result.get("sensitivity") or {}
_stability = _sensitivity_payload.get("top_stability")
_stability_txt = f"%{float(_stability) * 100:.1f}" if _stability is not None else "—"
_runtime = result.get("analysis_time")
_runtime_txt = f"{float(_runtime):.2f} {tt('sn', 'sec')}" if _runtime is not None else "—"
_shape = result.get("selected_data", pd.DataFrame()).shape
st.markdown(
    f"""
    <div class="kpi-strip">
        <div class="kpi-grid">
            <div class="kpi-item"><div class="kpi-label">{tt('Lider Alternatif', 'Leading Alternative')}</div><div class="kpi-value">{_lead_alt}</div></div>
            <div class="kpi-item"><div class="kpi-label">{tt('Sıralama Yöntemi', 'Ranking Method')}</div><div class="kpi-value">{_rank_method}</div></div>
            <div class="kpi-item"><div class="kpi-label">{tt('Kararlılık', 'Stability')}</div><div class="kpi-value">{_stability_txt}</div></div>
            <div class="kpi-item"><div class="kpi-label">{tt('Analiz Süresi', 'Analysis Time')}</div><div class="kpi-value">{_runtime_txt}</div></div>
            <div class="kpi-item"><div class="kpi-label">{tt('Veri Boyutu', 'Data Size')}</div><div class="kpi-value">{_shape[0]}×{_shape[1]}</div></div>
        </div>
    </div>
    """,
    unsafe_allow_html=True,
)

with st.expander(tt("🧮 Hesaplama Detaylarını Gör", "🧮 View Calculation Details"), expanded=False):
    if "show_calc_details" not in st.session_state:
        st.session_state["show_calc_details"] = False
    if st.button(tt("🔍 Hesaplama Adımlarını Göster / Gizle", "🔍 Show / Hide Calculation Steps"), use_container_width=True, key="btn_toggle_calc_details"):
        st.session_state["show_calc_details"] = not st.session_state["show_calc_details"]
    if st.session_state["show_calc_details"]:
        dt1, dt2, dt3 = st.tabs([tt("⚖️ Ağırlık Adımları", "⚖️ Weighting Steps"), tt("🏆 Sıralama Adımları", "🏆 Ranking Steps"), tt("🎲 Duyarlılık Adımları", "🎲 Sensitivity Steps")])

        with dt1:
            wm = result["weights"]["method"]
            st.markdown(f"**{tt('Yöntem', 'Method')}:** `{method_display_name(wm)}`")
            st.code(get_method_steps(wm, "weight"), language="text")
            w_frames = _extract_detail_tables(result["weights"].get("details", {}))
            if w_frames:
                w_key = st.selectbox(tt("Ağırlık detay tablosu", "Weighting detail table"), list(w_frames.keys()), key="calc_weight_tbl")
                render_table(localize_df(w_frames[w_key].head(250)))
            else:
                st.info(tt("Ağırlık yöntemi için detay tablo bulunamadı.", "No detail table found for the weighting method."))

        with dt2:
            rm = result.get("ranking", {}).get("method")
            if rm:
                st.markdown(f"**{tt('Yöntem', 'Method')}:** `{rm}`")
                st.code(get_method_steps(rm, "ranking"), language="text")
                r_frames = _extract_detail_tables(result.get("ranking", {}).get("details", {}))
                if r_frames:
                    r_key = st.selectbox(tt("Sıralama detay tablosu", "Ranking detail table"), list(r_frames.keys()), key="calc_rank_tbl")
                    render_table(localize_df(r_frames[r_key].head(250)))
                else:
                    st.info(tt("Sıralama yöntemi için detay tablo bulunamadı.", "No detail table found for the ranking method."))
            else:
                st.info(tt("Bu analizde sıralama yöntemi seçilmedi.", "No ranking method was selected in this analysis."))

        with dt3:
            sens = result.get("sensitivity")
            if sens:
                st.markdown(f"**{tt('Duyarlılık yaklaşımı', 'Sensitivity approach')}:** {tt('Lokal senaryolar + Monte Carlo bozulma analizi', 'Local scenarios + Monte Carlo perturbation analysis')}")
                st.code(
                    tt(
                        "1) Baz ağırlıkla referans sıralama\n2) Her kriterde ±%10/±%20 lokal değişim\n3) Spearman ile sıra benzerliği\n4) Monte Carlo ile birincilik oranı",
                        "1) Baseline ranking with base weights\n2) ±10%/±20% local perturbations per criterion\n3) Rank similarity via Spearman\n4) First-place frequency via Monte Carlo",
                    ),
                    language="text",
                )
                local_df = sens.get("local_scenarios")
                mc_df = sens.get("monte_carlo_summary")
                if isinstance(local_df, pd.DataFrame) and not local_df.empty:
                    st.markdown(f"**{tt('Lokal senaryolar', 'Local scenarios')}**")
                    render_table(localize_df(local_df.head(250)))
                if isinstance(mc_df, pd.DataFrame) and not mc_df.empty:
                    st.markdown(f"**{tt('Monte Carlo özeti', 'Monte Carlo summary')}**")
                    render_table(localize_df(mc_df.head(250)))
            else:
                st.info(tt("Duyarlılık analizi bu çalışmada üretilmedi.", "Sensitivity analysis was not generated in this study."))

tabs_list = [
    tt("📊 İstatistik & Dağılım", "📊 Statistics & Distribution"),
    tt("⚖️ Ağırlıklar", "⚖️ Weights"),
    tt("🛡️ Ağırlık Sağlamlık", "🛡️ Weight Robustness"),
]
if needs_ranking:
    tabs_list.extend([tt("🏆 Sıralama & Ağ", "🏆 Ranking & Network"), tt("🔁 Karşılaştırma", "🔁 Comparison"), tt("🎲 Monte Carlo Dayanıklılık", "🎲 Monte Carlo Robustness")])
tabs_list.append(tt("📥 Çıktı İndir", "📥 Download Outputs"))
tabs = st.tabs(tabs_list)

with tabs[0]:
    render_tab_assistant(gen_stat_commentary(result), key="stat")
    s1, s2 = st.columns([1, 1.3])
    with s1:
        st.markdown('<div class="section-card">', unsafe_allow_html=True)
        st.subheader(tt("Betimleyici İstatistikler", "Descriptive Statistics"))
        _stats_disp = localize_df(result["stats"]).copy()
        _num_cols = _stats_disp.select_dtypes(include=[np.number]).columns
        _stats_disp[_num_cols] = _stats_disp[_num_cols].round(3)
        render_table(_stats_disp)
        st.markdown("</div>", unsafe_allow_html=True)
    with s2:
        st.markdown('<div class="section-card">', unsafe_allow_html=True)
        st.plotly_chart(fig_box_plots(result["selected_data"], criteria), use_container_width=True)
        st.markdown("</div>", unsafe_allow_html=True)

    corr_df = result.get("correlation_matrix")
    if isinstance(corr_df, pd.DataFrame):
        st.markdown('<div class="section-card">', unsafe_allow_html=True)
        fig_corr = px.imshow(
            corr_df, text_auto=".2f", aspect="auto",
            color_continuous_scale="RdBu_r", zmin=-1, zmax=1,
            title=tt("Kriterler Arası Korelasyon Haritası", "Correlation Heatmap Among Criteria"),
        )
        fig_corr.update_layout(height=420, **_THEME)
        st.plotly_chart(fig_corr, use_container_width=True)
        st.markdown("</div>", unsafe_allow_html=True)

with tabs[1]:
    render_tab_assistant(gen_weight_commentary(result), key="weight")
    with st.expander(tt("Tablolar ve Şekiller", "Tables and Figures"), expanded=True):
        weight_col = col_key(weight_df, "Ağırlık", "Weight")
        top3_weights = weight_df.sort_values(weight_col, ascending=False).head(3).reset_index(drop=True)
        w1, w2 = st.columns(2)
        with w1:
            st.markdown('<div class="section-card">', unsafe_allow_html=True)
            st.markdown(f"**{tt('İlk 3 Önemli Kriter', 'Top 3 Important Criteria')}**")
            _top3_disp = localize_df(top3_weights)
            _top3_w_col = col_key(_top3_disp, "Ağırlık", "Weight")
            if _top3_w_col in _top3_disp.columns:
                _top3_disp[_top3_w_col] = pd.to_numeric(_top3_disp[_top3_w_col], errors="coerce").round(6)
            render_table(_top3_disp)
            st.subheader(tt("Ağırlık Tablosu", "Weight Table"))
            _weight_disp = localize_df(weight_df)
            _weight_disp_col = col_key(_weight_disp, "Ağırlık", "Weight")
            if _weight_disp_col in _weight_disp.columns:
                _weight_disp[_weight_disp_col] = pd.to_numeric(_weight_disp[_weight_disp_col], errors="coerce").round(6)
            render_table(_weight_disp)
            st.markdown("</div>", unsafe_allow_html=True)
        with w2:
            st.markdown('<div class="section-card">', unsafe_allow_html=True)
            st.plotly_chart(fig_weight_bar(weight_df), use_container_width=True)
            st.markdown("</div>", unsafe_allow_html=True)

        st.markdown('<div class="section-card">', unsafe_allow_html=True)
        st.plotly_chart(fig_weight_radar(weight_df), use_container_width=True)
        st.markdown("</div>", unsafe_allow_html=True)

    
with tabs[2]:
    _wr = result.get("weight_robustness", {}) or {}
    st.markdown('<div class="section-card">', unsafe_allow_html=True)
    st.subheader(tt("Ağırlık Sağlamlık Testleri", "Weight Robustness Tests"))
    if _wr.get("info"):
        st.info(str(_wr.get("info")))
    elif _wr.get("error"):
        st.warning(tt("Ağırlık sağlamlık testleri üretilemedi.", "Weight robustness tests could not be generated."))
    else:
        _loo_mean = _wr.get("loo_mean_rho")
        _loo_min = _wr.get("loo_min_rho")
        _top_match = _wr.get("loo_top_match_rate")
        _eff_n = _wr.get("effective_criteria_n")
        _dom = _wr.get("dominance_ratio")

        m1, m2, m3, m4 = st.columns(4)
        m1.metric(tt("Ortalama Uyum (LOO)", "Average Agreement (LOO)"), (f"{float(_loo_mean):.3f}" if _loo_mean is not None and np.isfinite(_loo_mean) else "—"))
        m2.metric(tt("En Düşük Uyum (LOO)", "Minimum Agreement (LOO)"), (f"{float(_loo_min):.3f}" if _loo_min is not None and np.isfinite(_loo_min) else "—"))
        m3.metric(tt("Lider Kriter Tutarlılığı", "Top-Criterion Consistency"), (f"%{float(_top_match)*100:.1f}" if _top_match is not None and np.isfinite(_top_match) else "—"))
        m4.metric(tt("Etkin Kriter Sayısı", "Effective Number of Criteria"), (f"{float(_eff_n):.2f}" if _eff_n is not None and np.isfinite(_eff_n) else "—"))

        _simple_msg = ""
        if _loo_mean is not None and np.isfinite(_loo_mean) and _top_match is not None and np.isfinite(_top_match):
            if (_loo_mean >= 0.90) and (_top_match >= 0.75):
                _simple_msg = tt("Ağırlıklar güçlü biçimde kararlı görünüyor. Bir alternatifi çıkardığınızda bile kriter önceliği büyük ölçüde korunuyor.", "Weights look strongly stable. Even when one alternative is removed, criterion priorities remain largely unchanged.")
            elif (_loo_mean >= 0.75) and (_top_match >= 0.50):
                _simple_msg = tt("Ağırlıklar orta düzeyde kararlı. Genel yön korunuyor ancak bazı kriterlerde öncelik sırası değişebilir.", "Weights are moderately stable. The general direction is preserved, but priority order may shift for some criteria.")
            else:
                _simple_msg = tt("Ağırlıklar kırılgan olabilir. Veri değiştiğinde kriter öncelikleri hızlı değişiyor; sonucu temkinli yorumlamak gerekir.", "Weights may be fragile. Criterion priorities change quickly when data changes, so interpret results cautiously.")
        _felsefe_msg = tt("Yöntem felsefesi açısından bu testler şunu söyler: 'seçtiğimiz objektif ağırlıklandırma, veri biraz oynadığında aynı hikâyeyi anlatıyor mu?'", "From a method philosophy perspective, these tests answer: 'Does the chosen objective weighting tell the same story when data changes slightly?'")
        st.markdown(f"**{_simple_msg}**<br><br>{_felsefe_msg}", unsafe_allow_html=True)
        if _dom is not None and np.isfinite(_dom):
            st.caption(tt(f"Lider/ikinci kriter ağırlık oranı: {_dom:.2f}. Oran yükseldikçe karar tek kritere daha bağımlı hale gelir.", f"Top/second criterion weight ratio: {_dom:.2f}. As this ratio grows, the decision depends more on a single criterion."))

        with st.expander(tt("Detay Tabloları", "Detailed Tables"), expanded=False):
            _loo_df = _wr.get("leave_one_out")
            if isinstance(_loo_df, pd.DataFrame) and not _loo_df.empty:
                st.markdown(f"**{tt('Leave-one-out özeti', 'Leave-one-out summary')}**")
                _loo_disp = localize_df(_loo_df.copy())
                _rho_col = col_key(_loo_disp, "SpearmanRho", "SpearmanRho")
                _maxdiff_col = col_key(_loo_disp, "MaksMutlakFark", "MaxAbsoluteDiff")
                if _rho_col in _loo_disp.columns:
                    _loo_disp[_rho_col] = pd.to_numeric(_loo_disp[_rho_col], errors="coerce").round(3)
                if _maxdiff_col in _loo_disp.columns:
                    _loo_disp[_maxdiff_col] = pd.to_numeric(_loo_disp[_maxdiff_col], errors="coerce").round(4)
                render_table(_loo_disp)
            _boot_df = _wr.get("bootstrap_summary")
            if isinstance(_boot_df, pd.DataFrame) and not _boot_df.empty:
                st.markdown(f"**{tt('Bootstrap kriter özeti', 'Bootstrap criterion summary')}**")
                _boot_disp = localize_df(_boot_df.copy())
                _mean_col = col_key(_boot_disp, "OrtalamaAğırlık", "MeanWeight")
                _std_col = col_key(_boot_disp, "StdSapma", "StdDev")
                _l5_col = col_key(_boot_disp, "AltYuzde5", "Lower5Pct")
                _u95_col = col_key(_boot_disp, "UstYuzde95", "Upper95Pct")
                for _c in [_mean_col, _std_col, _l5_col, _u95_col]:
                    if _c in _boot_disp.columns:
                        _boot_disp[_c] = pd.to_numeric(_boot_disp[_c], errors="coerce").round(4)
                render_table(_boot_disp)
    st.markdown("</div>", unsafe_allow_html=True)

curr_idx = 3
if needs_ranking:
    with tabs[curr_idx]:
        render_tab_assistant(gen_ranking_commentary(result, alt_names), key="rank")
        if ranking_table is None or ranking_table.empty:
            st.info(tt("Sıralama tablosu oluşturulamadı.", "Ranking table could not be generated."))
        else:
            _rt_disp = ranking_table.copy()
            _rt_alt_col = col_key(_rt_disp, "Alternatif", "Alternative")
            _rt_score_col = col_key(_rt_disp, "Skor", "Score")
            if alt_names:
                _rt_disp[_rt_alt_col] = _rt_disp[_rt_alt_col].astype(str).map(
                    lambda x: alt_names.get(x, x)
                )

            _rank_alt_col = col_key(ranking_table, "Alternatif", "Alternative")
            top_disp = alt_names.get(str(ranking_table.iloc[0][_rank_alt_col]), str(ranking_table.iloc[0][_rank_alt_col]))
            st.info(f"🏆 {tt('Lider Alternatif', 'Leading Alternative')}: **{top_disp}** — {tt('Yöntem', 'Method')}: {result.get('ranking', {}).get('method')}")
            st.markdown('<div class="section-card">', unsafe_allow_html=True)
            st.markdown(f"**{tt('İlk 5 Alternatif', 'Top 5 Alternatives')}**")
            _rt_top5 = localize_df(_rt_disp.head(5))
            _rt_top5_score_col = col_key(_rt_top5, "Skor", "Score")
            if _rt_top5_score_col in _rt_top5.columns:
                _rt_top5[_rt_top5_score_col] = pd.to_numeric(_rt_top5[_rt_top5_score_col], errors="coerce").round(4)
            render_table(_rt_top5)
            st.markdown("</div>", unsafe_allow_html=True)

            r1, r2 = st.columns(2)
            with r1:
                st.markdown('<div class="section-card">', unsafe_allow_html=True)
                st.subheader(tt("Nihai Sıralama", "Final Ranking"))
                _rt_table_disp = localize_df(_rt_disp)
                _rt_table_score_col = col_key(_rt_table_disp, "Skor", "Score")
                if _rt_table_score_col in _rt_table_disp.columns:
                    _rt_table_disp[_rt_table_score_col] = pd.to_numeric(_rt_table_disp[_rt_table_score_col], errors="coerce").round(4)
                render_table(_rt_table_disp)
                st.markdown("</div>", unsafe_allow_html=True)
            with r2:
                st.markdown('<div class="section-card">', unsafe_allow_html=True)
                st.plotly_chart(fig_rank_bar(_rt_disp), use_container_width=True)
                st.markdown("</div>", unsafe_allow_html=True)

            st.markdown('<div class="section-card network-wrap">', unsafe_allow_html=True)
            st.plotly_chart(fig_network_alternatives(_rt_disp), use_container_width=True)
            st.markdown("</div>", unsafe_allow_html=True)

            contrib = result.get("contribution_table")
            if isinstance(contrib, pd.DataFrame):
                _contrib_disp = contrib.copy()
                _contrib_alt_col = col_key(_contrib_disp, "Alternatif", "Alternative")
                if alt_names:
                    _contrib_disp[_contrib_alt_col] = _contrib_disp[_contrib_alt_col].astype(str).map(
                        lambda x: alt_names.get(x, x)
                    )
                st.markdown('<div class="section-card">', unsafe_allow_html=True)
                st.plotly_chart(fig_parallel_coords(_rt_disp, _contrib_disp, criteria), use_container_width=True)
                st.markdown("</div>", unsafe_allow_html=True)

    curr_idx += 1

    with tabs[curr_idx]:
        render_tab_assistant(gen_comparison_commentary(result), key="comp")
        comp = result.get("comparison", {})
        if comp and "rank_table" in comp:
            c1, c2 = st.columns(2)
            with c1:
                st.markdown('<div class="section-card">', unsafe_allow_html=True)
                st.subheader(tt("Yöntem Sıralama Karşılaştırması", "Method Ranking Comparison"))
                _rank_comp = comp["rank_table"].copy()
                _rank_comp_alt_col = col_key(_rank_comp, "Alternatif", "Alternative")
                if alt_names:
                    if _rank_comp_alt_col in _rank_comp.columns:
                        _rank_comp[_rank_comp_alt_col] = _rank_comp[_rank_comp_alt_col].astype(str).map(
                            lambda x: alt_names.get(x, x)
                        )
                render_table(localize_df(_rank_comp))
                st.markdown("</div>", unsafe_allow_html=True)
            with c2:
                if isinstance(comp.get("spearman_matrix"), pd.DataFrame):
                    st.markdown('<div class="section-card">', unsafe_allow_html=True)
                    spdf = comp["spearman_matrix"]
                    _method_col = col_key(spdf, "Yöntem", "Method")
                    fig_sp = px.imshow(
                        spdf.set_index(_method_col), text_auto=".2f", aspect="auto",
                        color_continuous_scale="RdYlGn", zmin=-1, zmax=1,
                        title=tt("Spearman Uyum Matrisi", "Spearman Agreement Matrix"),
                    )
                    fig_sp.update_layout(height=420, **_THEME)
                    st.plotly_chart(fig_sp, use_container_width=True)
                    st.markdown("</div>", unsafe_allow_html=True)

            _method_rob_df = build_ranking_robustness_table(result)
            if isinstance(_method_rob_df, pd.DataFrame) and not _method_rob_df.empty:
                st.markdown('<div class="section-card">', unsafe_allow_html=True)
                st.subheader(tt("Yöntem Bazlı Sağlamlık Yorumu", "Method-Based Robustness Interpretation"))
                st.caption(
                    tt(
                        "Bu tablo, her yöntemin ana yöntemle ne kadar benzer düşündüğünü ve bunun pratikte ne anlama geldiğini basit dille açıklar.",
                        "This table explains, in simple language, how similarly each method thinks compared with the primary method and what it means in practice.",
                    )
                )
                render_table(_method_rob_df)
                st.markdown("</div>", unsafe_allow_html=True)
        else:
            st.info(tt("Karşılaştırma için birden fazla yöntem seçilmelidir.", "Select more than one method for comparison."))

    curr_idx += 1

    with tabs[curr_idx]:
        render_tab_assistant(gen_mc_commentary(result), key="mc")
        sens = result.get("sensitivity")
        if not sens:
            st.info(tt("Sağlamlık analizi verisi oluşturulamadı. Lütfen sıralama yöntemi seçerek analizi çalıştırın.", "Robustness data could not be generated. Please run analysis with a ranking method."))
        else:
            mc_df  = sens.get("monte_carlo_summary")
            loc_df = sens.get("local_scenarios")
            stab   = float(sens.get("top_stability", 0.0))
            n_iter = int(sens.get("n_iterations", 0)) if "n_iterations" in sens else "—"

            mc1, mc2, mc3 = st.columns(3)
            mc1.metric(tt("Lider Kararlılığı", "Leader Stability"), f"%{stab*100:.1f}", delta="Monte Carlo")
            mc2.metric(tt("Simülasyon Sayısı", "Simulation Count"), f"{n_iter:,}" if isinstance(n_iter, int) else n_iter)
            if isinstance(mc_df, pd.DataFrame) and not mc_df.empty:
                _mc_alt_col = col_key(mc_df, "Alternatif", "Alternative")
                runner_up = mc_df.iloc[1][_mc_alt_col] if len(mc_df) > 1 else "—"
                runner_disp = alt_names.get(str(runner_up), str(runner_up))
                mc3.metric(tt("2. Sıra Alternatif", "Runner-up Alternative"), runner_disp)

            if isinstance(mc_df, pd.DataFrame) and not mc_df.empty:
                _mc_disp = mc_df.copy()
                _mc_disp_alt_col = col_key(_mc_disp, "Alternatif", "Alternative")
                if alt_names:
                    _mc_disp[_mc_disp_alt_col] = _mc_disp[_mc_disp_alt_col].astype(str).map(
                        lambda x: alt_names.get(x, x)
                    )

                mg1, mg2 = st.columns(2)
                with mg1:
                    st.markdown('<div class="section-card">', unsafe_allow_html=True)
                    st.plotly_chart(fig_mc_rank_bar(_mc_disp), use_container_width=True)
                    st.markdown("</div>", unsafe_allow_html=True)
                with mg2:
                    st.markdown('<div class="section-card">', unsafe_allow_html=True)
                    st.plotly_chart(fig_mc_stability_bubble(_mc_disp), use_container_width=True)
                    st.markdown("</div>", unsafe_allow_html=True)

                if isinstance(loc_df, pd.DataFrame) and not loc_df.empty:
                    st.markdown('<div class="section-card">', unsafe_allow_html=True)
                    st.plotly_chart(fig_local_sensitivity(loc_df), use_container_width=True)
                    st.markdown("</div>", unsafe_allow_html=True)

                st.markdown('<div class="section-card">', unsafe_allow_html=True)
                st.subheader(tt("Simülasyon Özet Tablosu", "Simulation Summary Table"))
                _mc_table_disp = localize_df(_mc_disp)
                _fp_col = col_key(_mc_table_disp, "BirincilikOranı", "FirstPlaceRate")
                _mr_col = col_key(_mc_table_disp, "OrtalamaSıra", "MeanRank")
                if _fp_col in _mc_table_disp.columns:
                    _mc_table_disp[_fp_col] = (pd.to_numeric(_mc_table_disp[_fp_col], errors="coerce") * 100.0).round(2)
                    _mc_table_disp = _mc_table_disp.rename(columns={_fp_col: f"{_fp_col} (%)"})
                if _mr_col in _mc_table_disp.columns:
                    _mc_table_disp[_mr_col] = pd.to_numeric(_mc_table_disp[_mr_col], errors="coerce").round(2)
                render_table(_mc_table_disp)
                st.markdown("</div>", unsafe_allow_html=True)

            _method_rob_df = build_ranking_robustness_table(result)
            if isinstance(_method_rob_df, pd.DataFrame) and not _method_rob_df.empty:
                st.markdown('<div class="section-card">', unsafe_allow_html=True)
                st.markdown(f"**{tt('Her Yöntem İçin Basit Yorum', 'Simple Interpretation for Each Method')}**")
                render_table(_method_rob_df)
                st.markdown("</div>", unsafe_allow_html=True)

    curr_idx += 1

with tabs[curr_idx]:
    _out_lang = st.session_state.get("ui_lang", "TR")
    dl1, dl2 = st.columns(2)
    with dl1:
        st.download_button(
            tt("📊 Tüm Sonuçları İndir (Excel)", "📊 Download All Results (Excel)"),
            data=generate_excel(result, result["selected_data"], lang=_out_lang),
            file_name=tt("MCDM_Sonuclari.xlsx", "MCDM_Results.xlsx"),
            use_container_width=True,
        )
    with dl2:
        if DOCX_AVAILABLE:
            _doc_bytes = generate_apa_docx(result, result["selected_data"], lang=_out_lang)
            st.session_state["report_docx"] = _doc_bytes
            st.download_button(
                tt("📄 Akademik Raporu İndir — APA Word", "📄 Download Academic Report — APA Word"),
                data=_doc_bytes,
                file_name=tt("MCDM_Akademik_Rapor.docx", "MCDM_Academic_Report.docx"),
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
        else:
            st.warning(tt("Word çıktısı için python-docx kurulu olmalıdır.", "python-docx must be installed for Word output."))
            
st.markdown(
    f"""
    <div style="text-align:right; padding: 16px 20px 6px 20px; font-family:'Georgia',serif;
                color:#1d3557; font-size:15px; font-weight:600; font-style:italic;">
        Prof. Dr. Ömer Faruk Rençber
    </div>
    <div style="text-align:center; padding: 4px 0 24px 0; font-size:0.80rem; color:#718096;">
        <a href="https://www.ofrencber.com" target="_blank"
           style="color:#2E7D9E; text-decoration:none; font-weight:600;">
            🌐 www.ofrencber.com
        </a>
        &nbsp;·&nbsp; MCDM Toolbox Professional &nbsp;·&nbsp; {tt("Akademik Karar Destek Sistemi", "Academic Decision Support System")}
    </div>
    """,
    unsafe_allow_html=True,
)
