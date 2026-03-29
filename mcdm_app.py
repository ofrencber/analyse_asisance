from __future__ import annotations

import base64
import html
import inspect
import io
import re
import time
from pathlib import Path
from typing import Any, Dict, List

import numpy as np
import pandas as pd
import streamlit as st

import mcdm_access as access
import mcdm_article

APP_DIR = Path(__file__).resolve().parent
MAX_UPLOAD_SIZE_MB = 20
MAX_UPLOAD_BYTES = MAX_UPLOAD_SIZE_MB * 1024 * 1024
MAX_UPLOAD_ROWS = 50_000
MAX_UPLOAD_COLS = 120
MAX_PANEL_YEAR_SELECTION = 10
MAX_SENSITIVITY_ITERATIONS = 1500
SESSION_INACTIVITY_SECONDS = 4 * 3600  # 4 saat hareketsizlikte otomatik çıkış

SPSS_AVAILABLE: bool | None = None
DOCX_AVAILABLE: bool | None = None
MPL_AVAILABLE: bool | None = None

Document = None
WD_ALIGN_PARAGRAPH = None
OxmlElement = None
qn = None
Inches = None
Pt = None
plt = None
px = None
go = None
spearmanr = None
me = None

def _ensure_pyreadstat() -> bool:
    global SPSS_AVAILABLE
    if SPSS_AVAILABLE is not None:
        return bool(SPSS_AVAILABLE)
    try:
        import pyreadstat  # noqa: F401
    except Exception:
        SPSS_AVAILABLE = False
    else:
        SPSS_AVAILABLE = True
    return bool(SPSS_AVAILABLE)

def _ensure_docx_support() -> bool:
    global DOCX_AVAILABLE, Document, WD_ALIGN_PARAGRAPH, OxmlElement, qn, Inches, Pt
    if DOCX_AVAILABLE is not None:
        return bool(DOCX_AVAILABLE)
    try:
        from docx import Document as _Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH as _WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement as _OxmlElement
        from docx.oxml.ns import qn as _qn
        from docx.shared import Inches as _Inches, Pt as _Pt
    except Exception:
        DOCX_AVAILABLE = False
        return False

    Document = _Document
    WD_ALIGN_PARAGRAPH = _WD_ALIGN_PARAGRAPH
    OxmlElement = _OxmlElement
    qn = _qn
    Inches = _Inches
    Pt = _Pt
    DOCX_AVAILABLE = True
    return True

def _ensure_matplotlib_support() -> bool:
    global MPL_AVAILABLE, plt
    if MPL_AVAILABLE is not None:
        return bool(MPL_AVAILABLE)
    try:
        import matplotlib

        matplotlib.use("Agg")
        import matplotlib.pyplot as _plt
    except Exception:
        MPL_AVAILABLE = False
        plt = None
        return False

    plt = _plt
    MPL_AVAILABLE = True
    return True

def _ensure_plotly_support() -> bool:
    global px, go
    if px is not None and go is not None:
        return True
    try:
        import plotly.express as _px
        import plotly.graph_objects as _go
    except Exception:
        px = None
        go = None
        return False
    px = _px
    go = _go
    return True

def _ensure_spearman_support() -> bool:
    global spearmanr
    if spearmanr is not None:
        return True
    try:
        from scipy.stats import spearmanr as _spearmanr
    except Exception:
        spearmanr = None
        return False
    spearmanr = _spearmanr
    return True

def _ensure_mcdm_engine() -> bool:
    global me
    if me is not None:
        return True
    try:
        import mcdm_engine as _me
    except Exception:
        me = None
        return False
    me = _me
    return True

st.set_page_config(
    page_title="MCDM Karar Destek Sistemi — Prof. Dr. Ömer Faruk Rençber",
    page_icon="🔬",
    layout="wide",
    initial_sidebar_state="expanded",
)

@st.cache_data(show_spinner=False)
def _get_app_css() -> str:
    return """
    <style>
        @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800&family=Space+Mono:wght@400;700&display=swap');

        :root {
            --bg-main: #F0F5FA;
            --bg-sidebar: #E5EDF5;
            --card-bg: #FFFFFF;
            --text-main: #1C1C1E;
            --text-muted: #556070;
            --text-light: #7A8A9A;
            --border: #C8D8E8;
            --border-strong: #B0C8DC;
            --accent: #1F4A73;
            --accent-hover: #173754;
            --accent-soft: #EAF0F5;
            --button-soft: #87CEEB;
            --button-soft-hover: #5BB8E0;
            --button-soft-text: #000000;
            --header-bg-from: #1C3252;
            --header-bg-to: #263D5A;
            --header-text: #EEF2FA;
            --success-bg: #DFF0E6;
            --success-text: #1A5C3A;
            --warn-bg: #FEF0CC;
            --warn-text: #7A5018;
            --danger-bg: #FAE0E0;
            --danger-text: #7B1E1E;
        }

        .stApp {
            background:
                radial-gradient(circle at 0% 0%, rgba(135, 206, 235, 0.12) 0%, rgba(135, 206, 235, 0.00) 28%),
                radial-gradient(circle at 100% 0%, rgba(31, 74, 115, 0.08) 0%, rgba(31, 74, 115, 0.00) 26%),
                linear-gradient(180deg, #F6FAFE 0%, var(--bg-main) 50%, #EEF4FA 100%) !important;
            color: var(--text-main) !important;
            font-family: 'Inter', -apple-system, BlinkMacSystemFont, sans-serif;
        }

        .main .block-container {
            padding-top: 0.18rem !important;
            padding-bottom: 2.5rem;
            max-width: 1500px;
        }

        /* ────── Streamlit'in kendi header'ını gizle ────── */
        header[data-testid="stHeader"] { display: none !important; }
        .stDeployButton { display: none !important; }
        #MainMenu { display: none !important; }
        footer { display: none !important; }
        [data-testid="collapsedControl"] { display: none !important; }
        button[aria-label="Collapse sidebar"] { display: none !important; }
        button[aria-label="Expand sidebar"] { display: none !important; }

        /* ────── GLOBAL HEADER (yalnızca sağ panel) ────── */
        .global-header {
            position: relative;
            width: 100%;
            min-height: 108px;
            background:
                linear-gradient(135deg, rgba(255,255,255,0.06) 0%, rgba(255,255,255,0.0) 42%),
                linear-gradient(120deg, #132742 0%, #1E3A5F 52%, #12253D 100%);
            color: #EEF2FA !important;
            display: flex;
            flex-direction: column;
            justify-content: center;
            gap: 0.35rem;
            padding: 0.9rem 2.2rem 0.8rem 2.2rem;
            margin: 0 0 0.95rem 0;
            border-radius: 0 0 22px 22px;
            border-bottom: 1px solid rgba(184, 154, 92, 0.48);
            box-shadow: 0 10px 28px rgba(5, 13, 24, 0.28);
        }
        .header-title {
            font-size: 1.14rem;
            font-weight: 700;
            letter-spacing: 0.12em;
            text-transform: uppercase;
            color: #E7EDF7 !important;
            line-height: 1.2;
            text-align: center;
            text-shadow: 0 1px 10px rgba(0, 0, 0, 0.18);
        }
        .header-professor {
            width: 100%;
            font-size: 1.78rem;
            font-weight: 700;
            color: #FFFFFF !important;
            font-family: 'Monotype Corsiva', 'Apple Chancery', 'URW Chancery L', 'Brush Script MT', cursive;
            font-style: italic;
            line-height: 1.12;
            text-align: center;
            letter-spacing: 0.01em;
            text-shadow: 0 3px 18px rgba(0, 0, 0, 0.24);
        }
        .header-meta {
            display: flex;
            justify-content: space-between;
            align-items: center;
            gap: 1rem;
            margin-top: 0.15rem;
            padding-top: 0.45rem;
            border-top: 1px solid rgba(255,255,255,0.14);
        }
        .header-dedication {
            font-size: 0.82rem;
            font-style: italic;
            color: #D6E0EE !important;
            letter-spacing: 0.03em;
        }
        .header-url {
            font-size: 0.82rem;
            font-weight: 600;
            color: #DCE5F2 !important;
            letter-spacing: 0.05em;
        }
        .header-url a {
            color: #DCE5F2 !important;
            text-decoration: none;
            transition: color 0.15s ease;
        }
        .header-url a:hover {
            color: #FFFFFF !important;
        }

        /* ────── SIDEBAR BRANDING ────── */
        .sidebar-brand {
            text-align: center;
            padding: 0 0.35rem 0.5rem 0.35rem;
            border-bottom: 1px solid var(--border-strong);
            margin-bottom: 0.5rem;
            margin-top: -0.75rem;
        }
        .sidebar-brand-logo {
            display: block;
            width: 100%;
            max-width: 282px;
            margin: 0 auto;
            height: auto;
            filter: drop-shadow(0 10px 18px rgba(0, 0, 0, 0.16));
        }
        .sidebar-brand-name {
            margin-top: 0.45rem;
            font-size: 0.88rem;
            font-weight: 700;
            color: #1C3252 !important;
            font-family: "Georgia", "Times New Roman", serif;
            font-style: italic;
            letter-spacing: 0.01em;
            white-space: nowrap;
            overflow: hidden;
            text-overflow: ellipsis;
        }
        .sidebar-brand-title {
            font-size: 1.55rem;
            font-weight: 900;
            letter-spacing: 2px;
            color: #1C3252 !important;
            line-height: 1.15;
            font-family: 'Inter', sans-serif;
        }
        .sidebar-brand-sub {
            font-size: 0.68rem;
            font-weight: 600;
            letter-spacing: 0.4px;
            color: #5C5650 !important;
            text-transform: uppercase;
            margin-top: 0.1rem;
        }

        /* ────── SIDEBAR ────── */
        section[data-testid="stSidebar"],
        section[data-testid="stSidebar"] > div,
        aside[data-testid="stSidebar"],
        aside[data-testid="stSidebar"] > div {
            background-color: var(--bg-sidebar) !important;
            border-right: 1px solid var(--border-strong) !important;
        }
        section[data-testid="stSidebar"],
        aside[data-testid="stSidebar"] {
            min-width: 315px !important;
            width: 315px !important;
        }
        section[data-testid="stSidebar"][aria-expanded="false"],
        aside[data-testid="stSidebar"][aria-expanded="false"] {
            min-width: 315px !important;
            max-width: 315px !important;
            width: 315px !important;
            transform: translateX(0) !important;
            margin-left: 0 !important;
        }
        section[data-testid="stSidebar"][aria-expanded="false"] > div,
        aside[data-testid="stSidebar"][aria-expanded="false"] > div {
            min-width: 315px !important;
            max-width: 315px !important;
            width: 315px !important;
        }
        section[data-testid="stSidebar"] *,
        aside[data-testid="stSidebar"] *,
        [data-testid="stSidebar"] * {
            color: var(--text-main) !important;
        }
        /* Sidebar global font sizes — compact & readable */
        section[data-testid="stSidebar"] label,
        section[data-testid="stSidebar"] .stRadio label,
        section[data-testid="stSidebar"] .stCheckbox label,
        section[data-testid="stSidebar"] .stSelectbox label,
        section[data-testid="stSidebar"] p,
        section[data-testid="stSidebar"] span {
            font-size: 0.75rem !important;
            line-height: 1.25 !important;
        }
        section[data-testid="stSidebar"] .stButton > button {
            font-size: 0.73rem !important;
            padding: 0.25rem 0.6rem !important;
            min-height: 1.8rem !important;
        }
        section[data-testid="stSidebar"] .stExpander {
            margin-bottom: 0.3rem !important;
        }
        section[data-testid="stSidebar"] [data-testid="stExpander"] summary {
            padding: 0.25rem 0.5rem !important;
            min-height: 1.6rem !important;
        }
        section[data-testid="stSidebar"] .element-container {
            margin-bottom: 0.15rem !important;
        }

        /* Sidebar section labels */
        .sb-section-label {
            font-size: 0.7rem;
            font-weight: 700;
            letter-spacing: 0.5px;
            text-transform: uppercase;
            color: #5C5650 !important;
            margin: 0.6rem 0 0.25rem 0;
            border-bottom: 1px solid var(--border);
            padding-bottom: 0.15rem;
        }

        /* Compact guide/method expanders in sidebar */
        section[data-testid="stSidebar"] [data-testid="stExpander"] summary {
            font-size: 0.73rem !important;
            padding: 0.3rem 0.5rem !important;
            min-height: 1.8rem !important;
            background: rgba(90,141,200,0.08) !important;
            border-radius: 6px !important;
            border: 1px solid var(--border) !important;
        }
        section[data-testid="stSidebar"] [data-testid="stExpander"] summary p {
            font-size: 0.73rem !important;
        }

        /* Sidebar footer */
        .sidebar-footer {
            text-align: center;
            padding: 0.6rem 0.4rem 0.4rem 0.4rem;
            margin-top: 0.5rem;
            border-top: 1px solid var(--border-strong);
            font-size: 0.68rem;
            font-weight: 600;
            color: #5C5650 !important;
            letter-spacing: 0.2px;
        }
        .sidebar-footer-mail {
            display: block;
            margin-top: 0.18rem;
            font-size: 0.64rem;
            font-weight: 500;
            color: #5C5650 !important;
            text-decoration: none !important;
            letter-spacing: 0;
        }

        /* ────── CARDS ────── */
        .section-card {
            background: var(--card-bg);
            border-radius: 10px;
            padding: 1rem 1.2rem;
            border: 1px solid var(--border);
            margin-bottom: 0.8rem;
        }

        /* ────── BUTTONS — açık gökyüzü mavisi, siyah yazı ────── */
        .stButton > button {
            border-radius: 8px !important;
            font-weight: 600 !important;
            border: 1px solid rgba(70, 150, 210, 0.40) !important;
            background: #87CEEB !important;
            color: #000000 !important;
            padding: 0.44rem 1rem !important;
            font-size: 0.875rem !important;
            letter-spacing: 0.15px !important;
            box-shadow: 0 2px 8px rgba(70, 140, 200, 0.15) !important;
            transition: background 0.12s ease, box-shadow 0.12s ease, border-color 0.12s ease !important;
        }
        .stDownloadButton > button,
        [data-testid="stDownloadButton"] > button {
            border-radius: 8px !important;
            font-weight: 700 !important;
            border: 1px solid rgba(18, 39, 62, 0.18) !important;
            background: linear-gradient(180deg, #274F78 0%, var(--accent) 100%) !important;
            color: #FFFFFF !important;
            padding: 0.44rem 1rem !important;
            font-size: 0.875rem !important;
            letter-spacing: 0.15px !important;
            box-shadow: 0 8px 18px rgba(21, 43, 66, 0.16) !important;
            transition: background 0.12s ease, box-shadow 0.12s ease, border-color 0.12s ease !important;
        }
        /* Force download button text white — override generic .stButton black text */
        .stDownloadButton > button *,
        [data-testid="stDownloadButton"] > button * {
            color: #FFFFFF !important;
        }
        .stButton > button:hover,
        .stButton > button[kind="secondary"]:hover {
            background: #5BB8E0 !important;
            color: #000000 !important;
            border-color: rgba(70, 150, 210, 0.60) !important;
            box-shadow: 0 4px 12px rgba(70, 140, 200, 0.25) !important;
        }
        .stDownloadButton > button:hover,
        [data-testid="stDownloadButton"] > button:hover {
            background: linear-gradient(180deg, #1E4367 0%, var(--accent-hover) 100%) !important;
            color: #FFFFFF !important;
            box-shadow: 0 10px 20px rgba(21, 43, 66, 0.22) !important;
            border-color: rgba(18, 39, 62, 0.26) !important;
        }
        .stButton > button:focus,
        .stButton > button:active,
        [data-testid="stDownloadButton"] > button:focus,
        [data-testid="stDownloadButton"] > button:active {
            box-shadow: 0 0 0 0.16rem rgba(70, 150, 210, 0.30) !important;
            outline: none !important;
        }
        .stButton > button[kind="primary"] {
            background: linear-gradient(180deg, #274F78 0%, var(--accent) 100%) !important;
            color: #FFFFFF !important;
            border: 1px solid rgba(18, 39, 62, 0.18) !important;
            box-shadow: 0 8px 18px rgba(21, 43, 66, 0.16) !important;
        }
        .stButton > button[kind="primary"]:hover {
            background: linear-gradient(180deg, #1E4367 0%, var(--accent-hover) 100%) !important;
            color: #FFFFFF !important;
            border-color: rgba(18, 39, 62, 0.26) !important;
            box-shadow: 0 10px 20px rgba(21, 43, 66, 0.22) !important;
        }
        .st-key-step1_continue_btn button,
        .st-key-btn_use_upload_data_main button,
        .st-key-btn_use_manual_data_main button,
        .st-key-btn_prep_complete_main button,
        .st-key-btn_run_analysis_main button {
            background: #C62828 !important;
            color: #FFFFFF !important;
            -webkit-text-fill-color: #FFFFFF !important;
            border: 1px solid #C62828 !important;
            box-shadow: 0 8px 18px rgba(198, 40, 40, 0.22) !important;
        }
        .st-key-step1_continue_btn button *,
        .st-key-btn_use_upload_data_main button *,
        .st-key-btn_use_manual_data_main button *,
        .st-key-btn_prep_complete_main button *,
        .st-key-btn_run_analysis_main button * {
            color: #FFFFFF !important;
            fill: #FFFFFF !important;
            -webkit-text-fill-color: #FFFFFF !important;
        }
        .st-key-step1_continue_btn button:hover,
        .st-key-btn_use_upload_data_main button:hover,
        .st-key-btn_use_manual_data_main button:hover,
        .st-key-btn_prep_complete_main button:hover,
        .st-key-btn_run_analysis_main button:hover {
            background: #B71C1C !important;
            color: #FFFFFF !important;
            -webkit-text-fill-color: #FFFFFF !important;
            border-color: #B71C1C !important;
            box-shadow: 0 10px 20px rgba(183, 28, 28, 0.24) !important;
        }
        .st-key-step1_continue_btn button:hover *,
        .st-key-btn_use_upload_data_main button:hover *,
        .st-key-btn_use_manual_data_main button:hover *,
        .st-key-btn_prep_complete_main button:hover *,
        .st-key-btn_run_analysis_main button:hover * {
            color: #FFFFFF !important;
            fill: #FFFFFF !important;
            -webkit-text-fill-color: #FFFFFF !important;
        }

        /* ────── FILE UPLOADER ────── */
        [data-testid="stFileUploaderDropzone"] {
            border: 1.5px dashed var(--border-strong) !important;
            background: var(--accent-soft) !important;
            border-radius: 10px !important;
        }
        [data-testid="stFileUploaderDropzone"] button {
            background: var(--accent) !important;
            color: #FFFFFF !important;
            border: none !important;
            border-radius: 7px !important;
            font-weight: 600 !important;
        }

        /* ────── TABLES — light background ────── */
        [data-testid="stTable"] table {
            background: #FFFFFF !important;
            color: var(--text-main) !important;
            border: 1px solid var(--border) !important;
            border-radius: 8px !important;
        }
        [data-testid="stTable"] thead th {
            background: #EEF2F8 !important;
            color: #1A2A3A !important;
            border: 1px solid var(--border) !important;
            font-weight: 600 !important;
            font-size: 0.84rem !important;
        }
        [data-testid="stTable"] tbody td {
            background: #FFFFFF !important;
            color: var(--text-main) !important;
            border: 1px solid #EDE8DF !important;
        }
        .stTable { color: var(--text-main) !important; }

        /* ────── EXPANDERS ────── */
        [data-testid="stExpander"] {
            border: 1px solid var(--border) !important;
            border-radius: 10px !important;
            background: var(--card-bg) !important;
            margin-bottom: 0.65rem !important;
            box-shadow: none !important;
            transition: border-color 0.12s ease !important;
        }
        [data-testid="stExpander"]:hover {
            border-color: var(--accent) !important;
        }
        [data-testid="stExpander"] details {
            border: none !important;
            border-radius: 10px !important;
            background: var(--card-bg) !important;
        }
        [data-testid="stExpander"] summary {
            border-radius: 10px !important;
            font-weight: 600 !important;
            color: var(--text-main) !important;
            font-size: 0.9rem !important;
            padding: 0.65rem 0.8rem !important;
        }

        /* ────── TABS ────── */
        .stTabs [data-baseweb="tab-list"] {
            gap: 3px;
            background: #D6E8F5;
            border-radius: 10px;
            border: 1px solid var(--border-strong);
            padding: 4px;
        }
        .stTabs [data-baseweb="tab"] {
            background: transparent;
            border-radius: 8px;
            border: none;
            padding: 0.38rem 0.82rem;
            font-size: 0.83rem;
            font-weight: 500;
            color: var(--text-main);
            transition: all 0.12s ease;
        }
        .stTabs [data-baseweb="tab"]:hover {
            background: rgba(135, 206, 235, 0.45) !important;
        }
        .stTabs [aria-selected="true"] {
            background: #87CEEB !important;
            color: #000000 !important;
            font-weight: 700 !important;
            border: none !important;
        }

        /* ────── ASSISTANT / INFO BOXES ────── */
        .assistant-box {
            background: var(--accent-soft);
            border: 1px solid var(--border);
            border-left: 3px solid var(--accent);
            border-radius: 10px;
            padding: 1rem 1.1rem;
            margin-bottom: 0.85rem;
        }
        .tab-assistant {
            background: var(--accent-soft);
            border: 1px solid var(--border);
            border-left: 4px solid var(--accent);
            border-radius: 10px;
            padding: 0.9rem 1rem;
            font-size: 0.86rem;
            line-height: 1.68;
            color: var(--text-main);
        }
        .tab-assistant strong { color: #1A2A3A; }

        /* ────── COMMENTARY BOX ────── */
        .commentary-box {
            background: #F0F6E8;
            border: 1px solid #C8D8B4;
            border-left: 3px solid #5A7A3A;
            border-radius: 8px;
            padding: 0.65rem 0.9rem;
            font-size: 0.84rem;
            line-height: 1.62;
            color: var(--text-main);
            margin-top: 0.4rem;
        }

        /* ────── LAYER CARDS ────── */
        .layer-card {
            border-radius: 8px;
            padding: 0.75rem 1rem;
            margin-bottom: 0.45rem;
            border: 1px solid var(--border);
            line-height: 1.52;
        }
        .layer-desc { background: #EDF4FF; }
        .layer-analytic { background: #EDF7F0; }
        .layer-norm { background: #FFF8EC; }
        .layer-label {
            font-size: 0.69rem;
            font-weight: 700;
            text-transform: uppercase;
            letter-spacing: 0.4px;
            margin-bottom: 0.25rem;
        }
        .layer-text { font-size: 0.87rem; color: var(--text-main); margin: 0; }

        /* ────── BADGES ────── */
        .badge-benefit { background:#D4EFE2; color:#1A5C40; border-radius:5px; padding:0.1rem 0.35rem; font-size:0.65rem; font-weight:700; border-left:3px solid #1A5C40; }
        .badge-benefit::before { content:"▲ "; font-size:0.55rem; }
        .badge-cost    { background:#FAE0E0; color:#7B1E1E; border-radius:5px; padding:0.1rem 0.35rem; font-size:0.65rem; font-weight:700; border-left:3px solid #7B1E1E; }
        .badge-cost::before { content:"▼ "; font-size:0.55rem; }
        .sidebar-version-badge { display:inline-block; margin-top:0.35rem; padding:0.15rem 0.55rem; font-size:0.58rem; font-weight:700; letter-spacing:0.5px; text-transform:uppercase; color:#FFFFFF; background:linear-gradient(135deg, var(--accent) 0%, #2B6CA0 100%); border-radius:999px; }
        .section-card-result { background:linear-gradient(135deg, #F7FBF7 0%, #FFFFFF 100%); border-radius:10px; padding:1rem 1.2rem; border:1px solid #C3DCC3; border-left:4px solid #2E7D32; margin-bottom:0.8rem; }

        /* ────── NETWORK ────── */
        .network-wrap {
            background: #FFFFFF;
            border-radius: 10px;
            border: 1px solid var(--border);
            padding: 0.5rem;
        }

        .mc-high  { color: #1A5C3A; font-weight: 700; }
        .mc-mid   { color: #7A5018; font-weight: 700; }
        .mc-low   { color: #7B1E1E; font-weight: 700; }

        /* ────── STEPPER ────── */
        .stepper-wrap {
            background: #FFFFFF;
            border: 1px solid var(--border);
            border-radius: 10px;
            padding: 0.55rem 0.75rem;
            margin-bottom: 0.6rem;
            display: flex;
            flex-direction: column;
            gap: 0;
        }
        .step-item { display:flex; align-items:center; gap:0.4rem; margin:0.2rem 0; font-size:0.74rem; color: var(--text-main); }
        .step-dot { width:9px; height:9px; border-radius:50%; display:inline-block; }
        .step-done { background:#3D7A5E; }
        .step-pending { background:#BFBAB4; }
        .step-active { background:var(--accent); box-shadow:0 0 0 3px rgba(61,100,145,0.2); }

        /* ────── KPI STRIP ────── */
        .kpi-strip { background:#FFFFFF; border:1px solid var(--border); border-radius:10px; padding:0.8rem 1rem; margin:0.35rem 0 1rem 0; }
        .kpi-grid { display:grid; grid-template-columns:repeat(5, minmax(0,1fr)); gap:0.55rem; }
        .kpi-item { background:var(--card-bg); border:1px solid var(--border); border-radius:8px; padding:0.5rem 0.7rem; }
        .kpi-label { font-size:0.70rem; color: var(--text-muted); text-transform:uppercase; letter-spacing:.35px; }
        .kpi-value { font-size:0.92rem; font-weight:700; color: var(--text-main); margin-top:0.1rem; }

        /* ────── ASSISTANT GRID ────── */
        .assistant-grid { display:grid; grid-template-columns:repeat(3, minmax(0,1fr)); gap:0.45rem; margin-bottom:0.45rem; }
        .assistant-card2 { background:#FBFCFD; border:1px solid #E3E8EE; border-radius:8px; padding:0.52rem 0.6rem; }
        .assistant-title2 { font-size:0.68rem; font-weight:700; color:#58697A; text-transform:uppercase; letter-spacing:.26px; margin-bottom:0.18rem; }
        .assistant-body2 { font-size:0.78rem; color:#273645; line-height:1.45; }

        /* ────── TRACKING PANELS ────── */
        .tracking-panel {
            border-radius: 10px;
            padding: 0.58rem 0.68rem;
            border: 1px solid #E5EAF0;
            background: #FAFBFC;
            margin: 0.2rem 0 0.7rem 0;
        }
        .tracking-panel-blue  { border-color:#E2EAF3; background:#FBFCFE; }
        .tracking-panel-green { border-color:#E1EADF; background:#FBFDFC; }
        .tracking-panel-amber { border-color:#EEE4D2; background:#FEFCF8; }
        .tracking-panel-rose  { border-color:#EEDDDD; background:#FEFBFB; }
        .tracking-panel-compact .tracking-grid { margin-top:0.12rem; }
        .tracking-panel-guided .tracking-grid { margin-top:0.45rem; }
        .tracking-title {
            font-size: 0.72rem;
            font-weight: 700;
            letter-spacing: 0.28px;
            text-transform: uppercase;
            color: #516374;
            margin-bottom: 0.1rem;
        }
        .tracking-subtitle {
            font-size: 0.74rem;
            line-height: 1.45;
            color: #718292;
        }
        .tracking-grid {
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(140px, 1fr));
            gap: 0.42rem;
        }
        .tracking-card {
            background: #FFFFFF;
            border: 1px solid #E6EBF0;
            border-left: 2px solid #C4D2DE;
            border-radius: 8px;
            padding: 0.46rem 0.56rem;
            min-height: 0;
        }
        .tracking-card-blue {
            border-left-color: #7EA8C8;
        }
        .tracking-card-green {
            border-left-color: #739B7F;
        }
        .tracking-card-amber {
            border-left-color: #B8914A;
        }
        .tracking-card-rose {
            border-left-color: #BF7D7D;
        }
        .tracking-card-slate {
            border-left-color: #95A3B3;
        }
        .tracking-label {
            font-size: 0.62rem;
            font-weight: 700;
            text-transform: uppercase;
            letter-spacing: 0.22px;
            color: #7B8B99;
        }
        .tracking-value {
            font-size: 0.79rem;
            font-weight: 600;
            line-height: 1.38;
            color: #283645;
            margin-top: 0.12rem;
            word-break: break-word;
        }
        .tracking-note {
            margin-top: 0.42rem;
            padding: 0.34rem 0.44rem;
            border-radius: 7px;
            border: 1px solid #E3E8EE;
            background: #FFFFFF;
            font-size: 0.72rem;
            line-height: 1.4;
            color: #667687;
        }

        /* ────── DIAG BADGES ────── */
        .diag-badge { display:inline-block; font-size:0.68rem; font-weight:700; border-radius:999px; padding:0.15rem 0.45rem; margin-left:0.3rem; }
        .diag-badge-good { background:#D4EFE2; color:#1A5C40; }
        .diag-badge-mid  { background:#FEF0CC; color:#7A5018; }
        .diag-badge-bad  { background:#FAE0E0; color:#7B1E1E; }

        .sidebar-small-note { font-size: 0.80rem; line-height: 1.5; color: var(--text-main); }

        .analysis-mini-banner {
            display: flex;
            justify-content: space-between;
            align-items: center;
            gap: 1rem;
            padding: 0.58rem 0.95rem;
            margin: 0 0 0.85rem 0;
            border-radius: 12px;
            background: linear-gradient(120deg, #17324D 0%, #234768 52%, #16314C 100%);
            border: 1px solid rgba(184, 154, 92, 0.26);
            box-shadow: 0 8px 20px rgba(5, 13, 24, 0.12);
        }
        .analysis-mini-banner-left {
            font-size: 0.9rem;
            font-weight: 800;
            letter-spacing: 0.04em;
            color: #F3F6FB !important;
            white-space: nowrap;
        }
        .analysis-mini-banner-right {
            font-size: 0.84rem;
            font-weight: 600;
            color: #DCE6F3 !important;
            text-align: right;
            white-space: nowrap;
        }

        h1, h2, h3, h4, h5, h6 { color: var(--text-main) !important; }
        p, span, label, div { color: var(--text-main); }

        /* ────── RESPONSIVE ────── */
        @media (max-width: 980px) {
            .main .block-container { padding-top: 1rem !important; }
            .global-header {
                padding: 0.8rem 1rem;
                min-height: 132px;
            }
            .header-title { font-size: 0.9rem; }
            .header-professor { font-size: 1.26rem; }
            .header-meta {
                flex-direction: column;
                align-items: flex-start;
                gap: 0.28rem;
            }
            .analysis-mini-banner {
                flex-direction: column;
                align-items: flex-start;
                padding: 0.62rem 0.8rem;
            }
            .analysis-mini-banner-left,
            .analysis-mini-banner-right {
                white-space: normal;
                text-align: left;
            }
            .kpi-grid { grid-template-columns:repeat(2, minmax(0,1fr)); }
            .assistant-grid { grid-template-columns:1fr; }
            .tracking-grid { grid-template-columns:1fr; }
        }
    </style>
    """

st.markdown(_get_app_css(), unsafe_allow_html=True)

def init_state() -> None:
    defaults = {
        "analysis_result": None,
        "clean_data": None,
        "raw_data": None,
        "report_docx": None,
        "editor_df": None,
        "data_source_id": None,
        "prep_done": False,
        "alt_names": {},        
        "crit_dir": {},         
        "crit_include": {},     
        "weight_method_pref": None,
        "ranking_prefs": [],
        "ui_lang": "TR",
        "analysis_scope": "single",
        "panel_year_column": None,
        "panel_entity_column": None,
        "panel_weight_strategy": "yearly",
        "panel_selected_years": [],
        "panel_selected_years_all": [],
        "panel_selected_years_col": None,
        "panel_results": None,
        "panel_years": [],
        "panel_view_choice": None,
        "panel_run_warnings": [],
        "step1_done": False,
        "direction_notice": None,
        "download_blob_cache": {},
        "download_blob_sig": None,
        "run_heavy_robustness": False,
        "study_title": "",
        "manual_entry_df": None,
        "manual_criteria_names": "",
        "manual_name_inputs_mode": False,
        "manual_criteria_inputs_seed": "",
        "manual_bulk_paste_text": "",
        "manual_paste_has_header": True,
        "data_entry_mode": "upload",
        "pending_data": None,
        "pending_data_source_id": None,
        "shown_step_hints": set(),
        "impute_mode_open": False,
        "missing_strategy_saved": "Sil",
        "clip_outliers_saved": False,
        "show_step_guidance": False,
    }
    for key, value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = value

init_state()

def tt(tr_text: str, en_text: str) -> str:
    return en_text if st.session_state.get("ui_lang", "TR") == "EN" else tr_text

def _show_step_guidance_enabled() -> bool:
    return bool(st.session_state.get("show_step_guidance", False))

def _show_step_caption(
    summary_tr: str,
    summary_en: str,
    detail_tr: str | None = None,
    detail_en: str | None = None,
) -> None:
    st.caption(tt(summary_tr, summary_en))
    if detail_tr and _show_step_guidance_enabled():
        st.caption(tt(detail_tr, detail_en or detail_tr))

_TR_HINTS = (
    " ve ",
    " ile ",
    " için ",
    " kriter",
    " yöntem",
    " ağırlık",
    " sıralama",
    " çalışma",
    " bulgu",
    " çözüm",
    " olarak ",
    " göre ",
)
_EN_HINTS = (
    " the ",
    " and ",
    " with ",
    " method",
    " criterion",
    " weight",
    " ranking",
    " analysis",
    " result",
    " should ",
    " therefore ",
    " according ",
)

def _looks_turkish_text(text: str) -> bool:
    raw = str(text or "")
    if not raw.strip():
        return False
    if any(ch in raw for ch in "çğıöşüÇĞİÖŞÜ"):
        return True
    low = f" {raw.lower()} "
    return any(tok in low for tok in _TR_HINTS)

def _looks_english_text(text: str) -> bool:
    raw = str(text or "")
    if not raw.strip():
        return False
    low = f" {raw.lower()} "
    return any(tok in low for tok in _EN_HINTS)

def _lang_text(text: str, lang: str, fallback: str) -> str:
    val = str(text or "").strip()
    if not val:
        return fallback
    if lang == "EN":
        return fallback if _looks_turkish_text(val) else val
    if _looks_english_text(val) and not _looks_turkish_text(val):
        return fallback
    return val

def _method_fallback_lines(method_label: str, lang: str) -> tuple[str, str]:
    if lang == "EN":
        return (
            f"{method_label} evaluates the problem through its own decision logic.",
            f"The findings from {method_label} should be interpreted within this method's assumptions.",
        )
    return (
        f"{method_label} yöntemi problemi kendi karar mantığıyla değerlendirir.",
        f"{method_label} bulguları bu yöntemin varsayımları içinde yorumlanmalıdır.",
    )

_SOFTWARE_CITATION = (
    "Rençber, Ö. F. (2026). MCDM Analysis Assistance (Version 1.1) "
    "[Computer software]. https://mcdm-assistance.streamlit.app/"
)

def _reference_key(text: Any) -> str:
    return re.sub(r"\s+", " ", str(text or "").strip()).casefold()

def _normalize_references(refs: List[Any], lang: str) -> List[str]:
    out: List[str] = []
    seen: set[str] = set()
    for ref in refs or []:
        txt = str(ref).strip()
        if not txt:
            continue
        if lang == "EN":
            txt = txt.replace("Kaynak:", "Source:")
        else:
            txt = txt.replace("Source:", "Kaynak:")
        key = _reference_key(txt)
        if key and key not in seen:
            seen.add(key)
            out.append(txt)
    if _reference_key(_SOFTWARE_CITATION) not in seen:
        out.append(_SOFTWARE_CITATION)
    return out

def _method_help_text(method_name: str) -> str:
    lang = "EN" if st.session_state.get("ui_lang") == "EN" else "TR"
    base_method = str(method_name or "").replace("Fuzzy ", "").strip()
    fallback = {
        "Eşit Ağırlık": {
            "simple": "Eşit ağırlık yaklaşımı, tüm kriterlere aynı başlangıç önemini verir.",
            "academic": "Bu seçim, veri temelli farklılaştırma yapmadan dengeli ve normatif bir başlangıç çerçevesi kurar.",
        },
        "Manuel Ağırlık": {
            "simple": "Manuel ağırlık yaklaşımı, kriter önemini doğrudan kullanıcı tercihine göre tanımlar.",
            "academic": "Bu seçim, karar verici yargısını modele açık biçimde taşıdığı için sonuçların öznel önceliklerle birlikte okunmasını gerektirir.",
        },
    }
    ph = (
        me.METHOD_PHILOSOPHY.get(str(method_name), {})
        or me.METHOD_PHILOSOPHY.get(base_method, {})
        or fallback.get(str(method_name), {})
        or fallback.get(base_method, {})
    )
    method_label = method_display_name(str(method_name or base_method or tt("Seçilen yöntem", "Selected method")))
    simple_fb, academic_fb = _method_fallback_lines(method_label, lang)
    simple = _lang_text(ph.get("simple", ""), lang, simple_fb)
    academic = _lang_text(ph.get("academic", ""), lang, academic_fb)
    return f"{simple}\n{academic}"

_COL_TR_EN = {
    "Kriter": "Criterion",
    "Ağırlık": "Weight",
    "ÖnemSırası": "ImportanceRank",
    "HamAğırlık": "RawWeight",
    "Alternatif": "Alternative",
    "Skor": "Score",
    "Sıra": "Rank",
    "Yöntem": "Method",
    "BirincilikOranı": "FirstPlaceRate",
    "OrtalamaSıra": "MeanRank",
    "AğırlıkDeğişimi": "WeightChange",
    "YeniAğırlık": "NewWeight",
    "TopKriterUyumu": "TopCriterionMatch",
    "MaksMutlakFark": "MaxAbsoluteDiff",
    "OrtalamaAğırlık": "MeanWeight",
    "AltYuzde5": "Lower5Pct",
    "UstYuzde95": "Upper95Pct",
    "StdSapma": "StdDev",
    "Bilgiİçeriği": "InformationContent",
    "Etkisi": "Effect",
    "Enİyi": "Best",
    "EnKötü": "Worst",
    "Ortalama": "Average",
    "BirinciAlternatif": "TopAlternative",
    "Değer": "Value",
    "Parametre": "Parameter",
    "Yıl": "Year",
    "AğırlıkYöntemi": "WeightMethod",
    "AğırlıkStratejisi": "WeightStrategy",
    "SıralamaYöntemi": "RankingMethod",
    "LiderAlternatif": "TopAlternative",
    "LiderSkor": "TopScore",
    "LiderKararlılığı": "LeaderStability",
    "GözlenenYılSayısı": "ObservedYears",
    "OrtalamaSkor": "MeanScore",
    "AçıklananVaryansOranı": "ExplainedVarianceRatio",
    "Bileşen": "Component",
    "DeltaX": "DeltaX",
    "DeltaY": "DeltaY",
    "EnİyiAlternatif": "BestAlternative",
    "EntropyAğırlığı": "EntropyWeight",
    "CILOSAğırlığı": "CILOSWeight",
    "IDOCRIWAğırlığı": "IDOCRIWWeight",
    "OrtalamaEntropy": "MeanEntropy",
    "OrtalamaCILOS": "MeanCILOS",
    "Senaryo": "Scenario",
}

_WEIGHT_METHOD_TR_EN = {
    "Standart Sapma": "Standard Deviation",
    "Eşit Ağırlık": "Equal Weights",
    "Manuel Ağırlık": "Manual Weights",
}

def col_key(df: pd.DataFrame, tr_name: str, en_name: str) -> str:
    if tr_name in df.columns:
        return tr_name
    if en_name in df.columns:
        return en_name
    return tr_name

def localize_df(df: pd.DataFrame) -> pd.DataFrame:
    if not isinstance(df, pd.DataFrame) or st.session_state.get("ui_lang") != "EN":
        return df
    return df.rename(columns={k: v for k, v in _COL_TR_EN.items() if k in df.columns})

def localize_df_lang(df: pd.DataFrame, lang: str) -> pd.DataFrame:
    if not isinstance(df, pd.DataFrame) or lang != "EN":
        return df
    return df.rename(columns={k: v for k, v in _COL_TR_EN.items() if k in df.columns})

def method_display_name(name: str) -> str:
    if st.session_state.get("ui_lang") == "EN":
        return _WEIGHT_METHOD_TR_EN.get(name, name)
    return name

def method_internal_name(name: str) -> str:
    if name in _WEIGHT_METHOD_TR_EN:
        return name
    for tr_name, en_name in _WEIGHT_METHOD_TR_EN.items():
        if name == en_name:
            return tr_name
    return name

def weight_method_groups() -> List[tuple[str, List[str]]]:
    return [
        (
            tt("Bilgi ve dağılım temelli", "Information and dispersion based"),
            ["Entropy", "Standart Sapma", "LOPCOW"],
        ),
        (
            tt("İlişki ve yapı temelli", "Relation and structure based"),
            ["CRITIC", "PCA"],
        ),
        (
            tt("Etki ve hibrit temelli", "Impact and hybrid based"),
            ["MEREC", "CILOS", "IDOCRIW", "Fuzzy IDOCRIW"],
        ),
        (
            tt("Denge ve simetri temelli", "Balance and symmetry based"),
            ["SPC"],
        ),
    ]

def _wm_single_select_cb(selected_name: str, all_methods: List[str]) -> None:
    """Weight method checkbox on_change: enforce single selection."""
    if st.session_state.get(f"weight_cb_{selected_name}", False):
        for _m in all_methods:
            if _m != selected_name:
                st.session_state[f"weight_cb_{_m}"] = False
        st.session_state["weight_method_pref"] = selected_name
    else:
        _selected = [m for m in all_methods if st.session_state.get(f"weight_cb_{m}", False)]
        st.session_state["weight_method_pref"] = _selected[0] if _selected else None

def ranking_method_groups(layer_key: str) -> List[tuple[str, List[str]]]:
    base_groups: List[tuple[str, List[str]]] = [
        (
            tt("İdeal ve uzlaşı odaklı", "Ideal and compromise oriented"),
            ["TOPSIS", "VIKOR", "MARCOS", "ARAS", "SPOTIS"],
        ),
        (
            tt("Sapma ve mesafe odaklı", "Deviation and distance oriented"),
            ["EDAS", "CODAS", "MABAC", "GRA", "RAFSI"],
        ),
        (
            tt("Fayda toplulaştırma odaklı", "Utility aggregation oriented"),
            ["SAW", "WPM", "MAUT", "WASPAS", "CoCoSo", "ROV", "AROMAN", "DNMA", "WISP"],
        ),
        (
            tt("Ağırlıksız / Kendi kendine ağırlıklı", "Self-weighting / Weight-free"),
            ["PSI"],
        ),
        (
            tt("Göreli üstünlük ve rekabet odaklı", "Relative dominance and competitiveness oriented"),
            ["COPRAS", "OCRA", "MOORA", "MULTIMOORA", "PROMETHEE", "RAWEC"],
        ),
    ]
    if layer_key != "fuzzy":
        return base_groups
    fuzzy_methods = set(me.FUZZY_MCDM_METHODS)
    return [
        (label, [f"Fuzzy {method}" for method in methods if f"Fuzzy {method}" in fuzzy_methods])
        for label, methods in base_groups
    ]

def sample_dataset() -> pd.DataFrame:
    rng = np.random.default_rng(42)
    n = 16
    df = pd.DataFrame({
        "Alternatif": [f"A{idx:02d}" for idx in range(1, n + 1)],
        "Karlılık": rng.uniform(10, 28, n),
        "Likidite": rng.uniform(0.9, 2.6, n),
        "PazarPayı": rng.uniform(8, 22, n),
        "MüşteriMemnuniyeti": rng.uniform(55, 95, n),
        "KaliteSkoru": rng.uniform(60, 98, n),
        "TeslimSüresi": rng.uniform(2, 15, n),
        "Risk": rng.uniform(1, 12, n),
        "İşletmeMaliyeti": rng.uniform(80, 420, n),
    })
    return df

def sample_dataset_en() -> pd.DataFrame:
    rng = np.random.default_rng(142)
    n = 16
    df = pd.DataFrame({
        "Alternative": [f"A{idx:02d}" for idx in range(1, n + 1)],
        "Profitability": rng.uniform(10, 28, n),
        "Liquidity": rng.uniform(0.9, 2.6, n),
        "MarketShare": rng.uniform(8, 22, n),
        "CustomerSatisfaction": rng.uniform(55, 95, n),
        "QualityScore": rng.uniform(60, 98, n),
        "DeliveryTime": rng.uniform(2, 15, n),
        "Risk": rng.uniform(1, 12, n),
        "OperatingCost": rng.uniform(80, 420, n),
    })
    return df

def sample_panel_dataset_en() -> pd.DataFrame:
    rng = np.random.default_rng(2026)
    countries = [
        "United States", "Canada", "Mexico", "Brazil", "Argentina",
        "United Kingdom", "Germany", "France", "Italy", "Spain",
        "Netherlands", "Sweden", "Norway", "Poland", "Turkey",
        "Russia", "China", "Japan", "South Korea", "India",
        "Indonesia", "Malaysia", "Thailand", "Vietnam", "Philippines",
        "Australia", "New Zealand", "South Africa", "Egypt", "Saudi Arabia",
    ]
    years = [2020, 2021, 2022, 2023, 2024]
    rows: List[Dict[str, Any]] = []

    for c_idx, country in enumerate(countries):
        country_factor = 0.88 + (c_idx / (len(countries) - 1)) * 0.24
        for y_idx, year in enumerate(years):
            progress = y_idx / (len(years) - 1)
            rows.append({
                "Country": country,
                "Year": year,
                "GDPGrowth": rng.uniform(1.2, 6.8) * country_factor + 0.35 * progress,
                "Inflation": rng.uniform(1.8, 9.8) * (1.08 - 0.12 * progress),
                "Unemployment": rng.uniform(3.0, 12.0) * (1.04 - 0.08 * progress),
                "RAndDSpending": rng.uniform(0.6, 4.2) * country_factor + 0.12 * progress,
                "DigitalReadiness": rng.uniform(45, 92) + 2.0 * progress + 2.5 * country_factor,
                "RenewableEnergyShare": rng.uniform(8, 58) + 3.0 * progress,
                "PublicDebt": rng.uniform(28, 130) * (1.02 - 0.05 * progress),
            })

    return pd.DataFrame(rows)

def _xl_text(lang: str, tr_text: str, en_text: str) -> str:
    return en_text if lang == "EN" else tr_text

def _make_unique_names(names: List[str], fallback_prefix: str) -> List[str]:
    out: List[str] = []
    seen: Dict[str, int] = {}
    for idx, raw in enumerate(names, start=1):
        base = str(raw or "").strip() or f"{fallback_prefix}_{idx}"
        count = seen.get(base, 0) + 1
        seen[base] = count
        out.append(base if count == 1 else f"{base}_{count}")
    return out

def _default_manual_criteria(count: int, lang: str) -> List[str]:
    prefix = "Criterion" if lang == "EN" else "Kriter"
    return [f"{prefix}_{idx}" for idx in range(1, count + 1)]

def _parse_manual_criteria(raw: str, count: int, lang: str) -> List[str]:
    fallback = "Criterion" if lang == "EN" else "Kriter"
    items = [part.strip() for part in str(raw or "").split(",") if part.strip()]
    while len(items) < count:
        items.append(f"{fallback}_{len(items) + 1}")
    return _make_unique_names(items[:count], fallback)

def _seed_manual_entry_df(
    existing: pd.DataFrame | None,
    row_count: int,
    entity_col: str,
    criteria_cols: List[str],
    year_col: str | None = None,
) -> pd.DataFrame:
    row_count = int(max(2, row_count))
    all_cols = [entity_col] + ([year_col] if year_col else []) + list(criteria_cols)
    out = pd.DataFrame(index=range(row_count))
    if isinstance(existing, pd.DataFrame):
        existing = existing.copy()
    for col in all_cols:
        if isinstance(existing, pd.DataFrame) and col in existing.columns:
            values = existing[col].tolist()[:row_count]
            if len(values) < row_count:
                if col == entity_col:
                    values += [f"A{idx:02d}" for idx in range(len(values) + 1, row_count + 1)]
                else:
                    values += [np.nan] * (row_count - len(values))
            out[col] = values
        elif col == entity_col:
            out[col] = [f"A{idx:02d}" for idx in range(1, row_count + 1)]
        elif year_col and col == year_col:
            out[col] = [2024] * row_count
        else:
            out[col] = pd.Series([np.nan] * row_count, dtype="float")
    return out

def _manual_schema_columns(entity_col: str, criteria_cols: List[str], year_col: str | None = None) -> List[str]:
    return [entity_col] + ([year_col] if year_col else []) + list(criteria_cols)

def _normalize_header_label(value: Any) -> str:
    return re.sub(r"[\W_]+", "", str(value or "").strip().lower(), flags=re.UNICODE)

def _looks_numeric_token(value: Any) -> bool:
    token = str(value or "").strip()
    if not token:
        return False
    try:
        float(token.replace(",", "."))
    except Exception:
        return False
    return True

def _parse_manual_paste_text(
    raw_text: str,
    entity_col: str,
    criteria_cols: List[str],
    year_col: str | None = None,
    *,
    has_header: bool = True,
) -> pd.DataFrame:
    raw_text = str(raw_text or "").strip()
    if not raw_text:
        raise ValueError(
            tt(
                "Yapistirilacak veri bulunamadi. Excel veya tablodan bir blok kopyalayip buraya yapistirin.",
                "No pasted data was found. Copy a block from Excel or another table and paste it here.",
            )
        )

    expected_cols = _manual_schema_columns(entity_col, criteria_cols, year_col)
    expected_norm = [_normalize_header_label(col) for col in expected_cols]
    parse_orders = [0, None] if has_header else [None, 0]

    for sep in ("\t", ";", ","):
        for header in parse_orders:
            try:
                parsed = pd.read_csv(io.StringIO(raw_text), sep=sep, header=header, engine="python")
            except Exception:
                continue
            if not isinstance(parsed, pd.DataFrame) or parsed.empty:
                continue
            parsed = parsed.dropna(axis=1, how="all").dropna(how="all").reset_index(drop=True)
            if parsed.empty:
                continue

            if len(parsed.columns) == len(expected_cols) + 1:
                first_col = str(parsed.columns[0]).strip().lower()
                if first_col.startswith("unnamed") or first_col in {"", "index"}:
                    parsed = parsed.iloc[:, 1:].copy()

            if len(parsed.columns) != len(expected_cols):
                continue

            if header == 0:
                header_values = [str(col).strip() for col in parsed.columns]
                header_norm = [_normalize_header_label(col) for col in header_values]
                numeric_header_cells = sum(_looks_numeric_token(col) for col in header_values)
                exact_header_match = header_norm == expected_norm
                likely_data_row = numeric_header_cells >= max(1, len(header_values) - 1)
                if not exact_header_match and has_header and likely_data_row:
                    continue

            parsed.columns = expected_cols
            return _seed_manual_entry_df(
                parsed,
                max(2, len(parsed)),
                entity_col,
                criteria_cols,
                year_col,
            )

    raise ValueError(
        tt(
            f"Yapistirilan tablo beklenen yapıya uymuyor. Beklenen sutun sirasi: {', '.join(expected_cols)}",
            f"The pasted table does not match the expected structure. Expected column order: {', '.join(expected_cols)}",
        )
    )

def _prepare_manual_entry_df(
    edited_df: pd.DataFrame,
    entity_col: str,
    criteria_cols: List[str],
    year_col: str | None = None,
) -> pd.DataFrame:
    if not isinstance(edited_df, pd.DataFrame) or edited_df.empty:
        raise ValueError(tt("Manuel giriş tablosu boş görünüyor.", "The manual entry table appears to be empty."))
    out = edited_df.copy()
    missing_cols = [col for col in [entity_col] + ([year_col] if year_col else []) + list(criteria_cols) if col not in out.columns]
    if missing_cols:
        raise ValueError(
            tt(
                f"Eksik sütunlar var: {', '.join(missing_cols)}",
                f"Some columns are missing: {', '.join(missing_cols)}",
            )
        )
    out = out[[entity_col] + ([year_col] if year_col else []) + list(criteria_cols)].copy()
    out[entity_col] = out[entity_col].astype(str).replace({"nan": "", "None": ""}).str.strip()
    numeric_targets = list(criteria_cols) + ([year_col] if year_col else [])
    for col in numeric_targets:
        out[col] = pd.to_numeric(out[col], errors="coerce")
    keep_mask = out[entity_col].ne("") | out[criteria_cols].notna().any(axis=1)
    if year_col:
        keep_mask = keep_mask | out[year_col].notna()
    out = out.loc[keep_mask].copy()
    if out.empty:
        raise ValueError(
            tt(
                "Kaydedilecek satır bulunamadı. En az bir alternatif ve sayısal değer girin.",
                "No rows are ready to save. Enter at least one alternative and some numeric values.",
            )
        )
    out[entity_col] = _make_unique_labels(out[entity_col].tolist(), fallback_prefix="A")
    if not out[criteria_cols].notna().any().any():
        raise ValueError(
            tt(
                "En az bir kriter sütununda sayısal veri girmeniz gerekiyor.",
                "Enter numeric values in at least one criterion column.",
            )
        )
    return out.reset_index(drop=True)

def _sync_manual_criteria_names_state(count: int, lang: str) -> None:
    raw = str(st.session_state.get("manual_criteria_names", "") or "").strip()
    if not raw:
        raw = ", ".join(_default_manual_criteria(count, lang))
    parsed = _parse_manual_criteria(raw, count, lang)
    st.session_state["manual_criteria_names"] = ", ".join(parsed)

def _sync_manual_criteria_input_keys(count: int, lang: str) -> None:
    parsed = _parse_manual_criteria(st.session_state.get("manual_criteria_names", ""), count, lang)
    seed = "||".join(parsed)
    prev_seed = str(st.session_state.get("manual_criteria_inputs_seed", "") or "")
    if seed != prev_seed:
        for idx, name in enumerate(parsed, start=1):
            st.session_state[f"manual_criteria_name_{idx}"] = name
        stale_idx = count + 1
        while f"manual_criteria_name_{stale_idx}" in st.session_state:
            del st.session_state[f"manual_criteria_name_{stale_idx}"]
            stale_idx += 1
        st.session_state["manual_criteria_inputs_seed"] = seed

def _upload_limit_mb() -> int:
    try:
        return int(st.get_option("server.maxUploadSize") or 200)
    except Exception:
        return 200

def _stage_data_source(df: pd.DataFrame, source_id: str) -> None:
    st.session_state["pending_data"] = df.copy()
    st.session_state["pending_data_source_id"] = source_id

def _clear_pending_data_source() -> None:
    st.session_state["pending_data"] = None
    st.session_state["pending_data_source_id"] = None

def _data_ready_button_label() -> str:
    return tt("✅ Verilerim tamam. Diğer adıma geçelim", "✅ My data is ready. Continue to the next step")

def _show_step_hint_once(token: str, message_tr: str, message_en: str, *, icon: str = "ℹ️") -> None:
    shown = set(st.session_state.get("shown_step_hints", set()) or set())
    if token in shown:
        return
    shown.add(token)
    st.session_state["shown_step_hints"] = shown
    show_guidance = _show_step_guidance_enabled()
    message = tt(message_tr, message_en)
    flow_steps = [
        tt("Veri hazır", "Data ready"),
        tt("Kriter doğrula", "Validate criteria"),
        tt("Yöntem ve sağlamlık", "Methods and robustness"),
        tt("Sonuçları incele", "Review results"),
    ]
    active_step = {
        "stage_objective": 1,
        "stage_prep": 2,
        "stage_methods": 3,
        "stage_results": 4,
    }.get(token, 1)
    steps_html_parts: List[str] = []
    for idx, label in enumerate(flow_steps, start=1):
        if idx < active_step:
            state_class = "is-done"
            marker = "✓"
        elif idx == active_step:
            state_class = "is-active"
            marker = str(idx)
        else:
            state_class = "is-pending"
            marker = str(idx)
        connector_html = '<div class="mcdm-flow-connector"></div>' if idx < len(flow_steps) else ""
        steps_html_parts.append(
            f"""
            <div class="mcdm-flow-step {state_class}">
                <div class="mcdm-flow-marker">{marker}</div>
                <div class="mcdm-flow-label">{_safe_html_text(label)}</div>
            </div>
            {connector_html}
            """
        )
    hint_id = f"mcdm-flow-hint-{re.sub(r'[^a-zA-Z0-9_-]+', '-', token)}"
    head_margin = "0.58rem" if show_guidance else "0.3rem"
    body_html = f'<div class="mcdm-flow-body">{_safe_html_text(message)}</div>' if show_guidance else ""
    flow_html = f"""
    <style>
    @keyframes mcdmFlowHintIn {{
        from {{ opacity:0; transform:translateX(26px) translateY(-6px); }}
        to {{ opacity:1; transform:translateX(0) translateY(0); }}
    }}
    @keyframes mcdmFlowHintOut {{
        from {{ opacity:1; transform:translateX(0) translateY(0); }}
        to {{ opacity:0; transform:translateX(34px) translateY(-4px); visibility:hidden; }}
    }}
    #{hint_id} {{
        position: fixed;
        top: 5.4rem;
        right: 1rem;
        width: min(312px, calc(100vw - 1.4rem));
        background: rgba(255,255,255,0.96);
        color: #243344;
        border: 1px solid rgba(157, 171, 186, 0.34);
        border-radius: 14px;
        box-shadow: 0 14px 30px rgba(18, 32, 48, 0.12);
        padding: 0.7rem 0.75rem;
        z-index: 999999;
        pointer-events: none;
        animation: mcdmFlowHintIn 0.22s ease-out, mcdmFlowHintOut 0.45s ease-in 4.8s forwards;
    }}
    #{hint_id} .mcdm-flow-head {{
        display:flex;
        gap:0.58rem;
        align-items:flex-start;
        margin-bottom:{head_margin};
    }}
    #{hint_id} .mcdm-flow-badge {{
        width:1.9rem;
        height:1.9rem;
        border-radius:999px;
        display:flex;
        align-items:center;
        justify-content:center;
        font-size:0.9rem;
        background:#F2F5F8;
        border:1px solid #E1E7ED;
        flex-shrink:0;
    }}
    #{hint_id} .mcdm-flow-title {{
        font-size:0.64rem;
        font-weight:700;
        letter-spacing:0.24px;
        text-transform:uppercase;
        color:#6C7B89;
        margin-bottom:0.18rem;
    }}
    #{hint_id} .mcdm-flow-body {{
        font-size:0.77rem;
        line-height:1.4;
        font-weight:600;
        color:#253646;
    }}
    #{hint_id} .mcdm-flow-route {{
        display:flex;
        align-items:center;
        gap:0.18rem;
        flex-wrap:wrap;
    }}
    #{hint_id} .mcdm-flow-step {{
        display:flex;
        align-items:center;
        gap:0.38rem;
        padding:0.26rem 0.4rem;
        border-radius:999px;
        background:#F8FAFB;
        border:1px solid #E4EAF0;
    }}
    #{hint_id} .mcdm-flow-step.is-done {{
        background:#F3F8F4;
        border-color:#D9E8DC;
    }}
    #{hint_id} .mcdm-flow-step.is-active {{
        background:#FBF5EA;
        border-color:#EAD6B1;
        box-shadow:0 0 0 1px rgba(212, 179, 121, 0.12) inset;
    }}
    #{hint_id} .mcdm-flow-step.is-pending {{
        opacity:0.74;
    }}
    #{hint_id} .mcdm-flow-marker {{
        width:1.08rem;
        height:1.08rem;
        border-radius:999px;
        display:flex;
        align-items:center;
        justify-content:center;
        font-size:0.62rem;
        font-weight:700;
        background:#EDF2F6;
        color:#4C6175;
        flex-shrink:0;
    }}
    #{hint_id} .mcdm-flow-label {{
        font-size:0.68rem;
        font-weight:600;
        letter-spacing:0.1px;
        white-space:nowrap;
        color:#435466;
    }}
    #{hint_id} .mcdm-flow-connector {{
        width:0.75rem;
        height:2px;
        background:linear-gradient(90deg, rgba(128,143,159,0.3), rgba(128,143,159,0.08));
        border-radius:999px;
        flex-shrink:0;
    }}
    @media (max-width: 920px) {{
        #{hint_id} {{
            top: auto;
            right: 0.7rem;
            left: 0.7rem;
            bottom: 0.8rem;
            width: auto;
        }}
        #{hint_id} .mcdm-flow-route {{
            row-gap:0.35rem;
        }}
    }}
    </style>
    <div id="{hint_id}">
        <div class="mcdm-flow-head">
            <div class="mcdm-flow-badge">{_safe_html_text(icon)}</div>
            <div>
                <div class="mcdm-flow-title">{_safe_html_text(tt("Akış", "Flow"))}</div>
                {body_html}
            </div>
        </div>
        <div class="mcdm-flow-route">
            {''.join(steps_html_parts)}
        </div>
    </div>
    """
    try:
        st.html(flow_html)
    except Exception:
        try:
            st.toast(message, icon=icon)
        except Exception:
            pass

def _activate_data_source(df: pd.DataFrame, source_id: str, entry_mode: str) -> None:
    st.session_state["raw_data"] = df.copy()
    st.session_state["data_source_id"] = source_id
    _clear_pending_data_source()
    st.session_state["analysis_result"] = None
    st.session_state["panel_results"] = None
    st.session_state["clean_data"] = None
    st.session_state["report_docx"] = None
    st.session_state["prep_done"] = False
    st.session_state["step1_done"] = False
    st.session_state["alt_names"] = {}
    st.session_state["crit_dir"] = {}
    st.session_state["crit_include"] = {}
    st.session_state["weight_method_pref"] = None
    st.session_state["ranking_prefs"] = []
    st.session_state["direction_notice"] = None
    st.session_state["download_blob_cache"] = {}
    st.session_state["download_blob_sig"] = None
    st.session_state["study_title"] = ""
    st.session_state["shown_step_hints"] = set()
    st.session_state["panel_year_column"] = None
    st.session_state["panel_entity_column"] = None
    st.session_state["panel_selected_years"] = []
    st.session_state["panel_selected_years_all"] = []
    st.session_state["panel_selected_years_col"] = None
    st.session_state["panel_years"] = []
    st.session_state["panel_view_choice"] = None
    st.session_state["panel_run_warnings"] = []
    st.session_state["analysis_scope"] = "panel" if _guess_year_columns(df) else "single"
    access.track_event(
        "data_source_loaded",
        {
            "entry_mode": entry_mode,
            "rows": int(df.shape[0]),
            "columns": int(df.shape[1]),
            "source_id": str(source_id),
        },
    )

def _manual_preset_catalog(lang: str) -> List[Dict[str, Any]]:
    entity_default = "Alternative" if lang == "EN" else "Alternatif"
    year_default = "Year" if lang == "EN" else "Yıl"
    return [
        {
            "key": "small",
            "label": _xl_text(lang, "Hızlı 5x4", "Quick 5x4"),
            "rows": 5,
            "criteria_count": 4,
            "entity_col": entity_default,
            "criteria_names": _default_manual_criteria(4, lang),
            "has_year": False,
            "year_col": year_default,
        },
        {
            "key": "medium",
            "label": _xl_text(lang, "Orta 10x6", "Medium 10x6"),
            "rows": 10,
            "criteria_count": 6,
            "entity_col": entity_default,
            "criteria_names": _default_manual_criteria(6, lang),
            "has_year": False,
            "year_col": year_default,
        },
        {
            "key": "panel",
            "label": _xl_text(lang, "Panel", "Panel"),
            "rows": 12,
            "criteria_count": 5,
            "entity_col": _xl_text(lang, "Ülke", "Country"),
            "criteria_names": _default_manual_criteria(5, lang),
            "has_year": True,
            "year_col": year_default,
        },
    ]

def _apply_manual_preset(preset: Dict[str, Any], lang: str) -> None:
    entity_col = str(preset.get("entity_col") or _xl_text(lang, "Alternatif", "Alternative"))
    criteria_names = [str(name).strip() for name in preset.get("criteria_names", []) if str(name).strip()]
    criteria_count = int(preset.get("criteria_count", len(criteria_names) or 4))
    if not criteria_names:
        criteria_names = _default_manual_criteria(criteria_count, lang)
    year_col = str(preset.get("year_col") or _xl_text(lang, "Yıl", "Year")).strip()
    has_year = bool(preset.get("has_year"))
    rows = int(max(2, preset.get("rows", 8)))

    st.session_state["manual_row_count"] = rows
    st.session_state["manual_criteria_count"] = criteria_count
    st.session_state["manual_entity_col"] = entity_col
    st.session_state["manual_has_year"] = has_year
    st.session_state["manual_year_col"] = year_col
    st.session_state["manual_criteria_names"] = ", ".join(criteria_names)
    st.session_state["manual_entry_df"] = _seed_manual_entry_df(
        None,
        rows,
        entity_col,
        criteria_names,
        year_col if has_year else None,
    )

def _fill_manual_demo_data(
    df: pd.DataFrame,
    entity_col: str,
    criteria_cols: List[str],
    year_col: str | None = None,
) -> pd.DataFrame:
    out = df.copy()
    rng = np.random.default_rng(42)
    n_rows = len(out)
    if entity_col in out.columns:
        values = out[entity_col].astype(str).replace({"nan": "", "None": ""}).str.strip().tolist()
        filled = []
        for idx, value in enumerate(values, start=1):
            filled.append(value if value else f"A{idx:02d}")
        out[entity_col] = filled
    if year_col and year_col in out.columns:
        years = [2022, 2023, 2024]
        out[year_col] = [years[idx % len(years)] for idx in range(n_rows)]
    for idx, col in enumerate(criteria_cols, start=1):
        if col not in out.columns:
            continue
        base = rng.uniform(10 * idx, 22 * idx, n_rows)
        trend = np.linspace(0, idx * 3, n_rows)
        out[col] = np.round(base + trend, 2)
    return out

def _manual_editor_stats(
    df: pd.DataFrame,
    entity_col: str,
    criteria_cols: List[str],
    year_col: str | None = None,
) -> Dict[str, Any]:
    if not isinstance(df, pd.DataFrame) or df.empty:
        return {"filled_rows": 0, "criteria_count": len(criteria_cols), "fill_rate": 0.0}
    entity_ok = df[entity_col].astype(str).replace({"nan": "", "None": ""}).str.strip().ne("") if entity_col in df.columns else pd.Series(False, index=df.index)
    numeric_df = df[criteria_cols].apply(pd.to_numeric, errors="coerce") if criteria_cols else pd.DataFrame(index=df.index)
    row_mask = entity_ok | numeric_df.notna().any(axis=1)
    if year_col and year_col in df.columns:
        row_mask = row_mask | pd.to_numeric(df[year_col], errors="coerce").notna()
    filled_rows = int(row_mask.sum())
    total_numeric_slots = int(max(1, len(criteria_cols) * max(filled_rows, 1)))
    filled_numeric_slots = int(numeric_df.loc[row_mask].notna().sum().sum()) if not numeric_df.empty else 0
    fill_rate = float(filled_numeric_slots / total_numeric_slots) if total_numeric_slots > 0 else 0.0
    return {
        "filled_rows": filled_rows,
        "criteria_count": len(criteria_cols),
        "fill_rate": fill_rate,
    }

def _render_upload_data_source_section(lang: str) -> None:
    _max_mb = _upload_limit_mb()
    st.caption(
        tt(
            "Beklenen format: ilk sütun alternatif etiketi olabilir; analizde kullanılacak kriter sütunları sayısal olmalıdır. CSV, XLSX ve SPSS `.sav` desteklenir.",
            "Expected format: the first column may contain alternative labels; criterion columns used in the analysis must be numeric. CSV, XLSX, and SPSS `.sav` are supported.",
        )
    )
    st.caption(
        tt(
            f"Yükleme sınırı: yaklaşık {_max_mb} MB. Daha büyük dosyalarda sütunları sadeleştirmeniz önerilir.",
            f"Upload limit: about {_max_mb} MB. For larger files, reduce unnecessary columns before upload.",
        )
    )

    _show_tpl = st.checkbox(tt("📐 Örnek veri formatı ve şablonu göster", "📐 Show sample data format and template"), key="chk_show_template")
    if _show_tpl:
        render_table(_input_format_notes(lang, panel=False))
        st.markdown(f"**{tt('Tek dönem örnek görünüm', 'Single-period example preview')}**")
        render_table(sample_dataset().head(5) if lang != "EN" else sample_dataset_en().head(5))
        _tpl_col1, _tpl_col2 = st.columns(2)
        with _tpl_col1:
            _single_tpl_clicked = st.download_button(
                tt("⬇️ Tek dönem şablonu indir", "⬇️ Download single-period template"),
                data=generate_input_template_excel(lang=lang, panel=False),
                file_name=tt("MCDM_Ornek_Format.xlsx", "MCDM_Sample_Format.xlsx"),
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                use_container_width=True,
                on_click="ignore",
            )
            if _single_tpl_clicked:
                access.track_event(
                    "template_downloaded",
                    {"template_type": "single", "lang": lang},
                )
        with _tpl_col2:
            _panel_tpl_clicked = st.download_button(
                tt("⬇️ Panel şablonu indir", "⬇️ Download panel template"),
                data=generate_input_template_excel(lang="EN", panel=True),
                file_name="MCDM_Panel_Template.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                use_container_width=True,
                on_click="ignore",
            )
            if _panel_tpl_clicked:
                access.track_event(
                    "template_downloaded",
                    {"template_type": "panel", "lang": "EN"},
                )

    uploaded = st.file_uploader(
        tt("CSV, XLSX veya SAV yükleyin", "Upload CSV, XLSX, or SAV"),
        type=["csv", "xlsx", "sav"],
        label_visibility="collapsed",
        key="main_data_file_uploader",
    )
    sample_col_1, sample_col_2, sample_col_3 = st.columns(3)
    with sample_col_1:
        if st.button(tt("📘 Örnek Veri (TR)", "📘 Sample Data (TR)"), use_container_width=True, key="btn_sample_tr_main"):
            _stage_data_source(sample_dataset(), "sample_data_tr")
            st.rerun()
    with sample_col_2:
        if st.button(tt("📗 Örnek Veri (EN)", "📗 Sample Data (EN)"), use_container_width=True, key="btn_sample_en_main"):
            _stage_data_source(sample_dataset_en(), "sample_data_en")
            st.rerun()
    with sample_col_3:
        if st.button(tt("📙 Panel Veri (EN)", "📙 Panel Data (EN)"), use_container_width=True, key="btn_sample_panel_en_main"):
            _stage_data_source(sample_panel_dataset_en(), "sample_panel_en")
            st.rerun()

    if uploaded is not None:
        _upload_sig = f"upload::{uploaded.name}::{getattr(uploaded, 'size', 0)}"
        if (
            st.session_state.get("data_source_id") != _upload_sig
            and st.session_state.get("pending_data_source_id") != _upload_sig
        ):
            try:
                loaded_df = load_uploaded_file(uploaded)
            except ValueError as upload_exc:
                st.error(str(upload_exc))
            except Exception as upload_exc:
                st.error(tt("Dosya okunamadı. Lütfen dosya formatını kontrol edip tekrar deneyin.", "File could not be read. Please verify the file format and try again."))
                st.caption(tt(f"Hata kodu: {_safe_error_code(upload_exc)}", f"Error code: {_safe_error_code(upload_exc)}"))
            else:
                _stage_data_source(loaded_df, _upload_sig)
                st.rerun()

    _pending_df = st.session_state.get("pending_data")
    _pending_id = str(st.session_state.get("pending_data_source_id") or "")
    if isinstance(_pending_df, pd.DataFrame) and _pending_id and _pending_id != "manual_entry":
        st.success(
            tt(
                f"Seçilen veri hazır: {_pending_df.shape[0]} satır, {_pending_df.shape[1]} sütun.",
                f"Selected data is ready: {_pending_df.shape[0]} rows, {_pending_df.shape[1]} columns.",
            )
        )
        _show_preview = st.checkbox(tt("👀 Yüklenen veriyi önizle", "👀 Preview selected data"), key="chk_preview_upload")
        if _show_preview:
            render_table(_pending_df.head(8))
        if st.button(_data_ready_button_label(), key="btn_use_upload_data_main", use_container_width=True):
            _activate_data_source(_pending_df, _pending_id, "upload")
            st.rerun()

def _render_data_input_workspace_body(lang: str) -> None:
    _mode = st.radio(
        tt("Veri kaynağı", "Data source"),
        options=["upload", "manual"],
        index=0 if st.session_state.get("data_entry_mode", "upload") == "upload" else 1,
        format_func=lambda mode: tt("📤 Dosya / Örnek Veri", "📤 File / Sample Data") if mode == "upload" else tt("✍️ Manuel Giriş", "✍️ Manual Entry"),
        horizontal=True,
        key="data_entry_mode",
    )
    if _mode == "upload":
        _render_upload_data_source_section(lang)
    else:
        st.caption(
            tt(
                "Manuel giriş seçildi. Dosya yükleme alanı gizlendi; isterseniz aşağıdaki bölümü açıp tabloyu oluşturabilirsiniz.",
                "Manual entry is selected. The file upload area is hidden; open the section below to build the table if needed.",
            )
        )
        _render_manual_entry_workspace(lang)

def _render_data_input_workspace(lang: str, is_data_loaded: bool) -> None:
    _title = f"📂 {tt('Veri Girişi', 'Data Input')} {'✅' if is_data_loaded else '⏳'}"
    if is_data_loaded:
        with st.expander(_title, expanded=False):
            _render_data_input_workspace_body(lang)
    else:
        st.markdown(f"### {_title}")
        st.caption(
            tt(
                "Bu alan veri seçilene kadar açık kalır. Dosya yükleyin veya manuel girişinizi tamamlayıp aşağıdaki onay butonuyla ilerleyin.",
                "This area stays open until data is selected. Upload a file or complete manual entry, then continue with the confirmation button below.",
            )
        )
        _render_data_input_workspace_body(lang)

def _render_manual_entry_workspace(lang: str) -> None:
    st.markdown(f"#### {tt('✍️ Manuel tablo girişi', '✍️ Manual table entry')}")
    with st.container():
        st.caption(
            tt(
                "Dosya yuklemeden ilerlemek icin bu alani kullanin. En hizli akis: yapıyı bir kez kurun, sonra Excel'den toplu yapistirin veya tabloyu kaydedip devam edin.",
                "Use this area to proceed without uploading a file. Fastest flow: set the structure once, then bulk-paste from Excel or save the grid and continue.",
            )
        )
        _controls_col, _editor_col = st.columns([0.95, 1.8], gap="large")

        with _controls_col:
            st.markdown(f"**{tt('1. Hızlı kurulum', '1. Quick setup')}**")
            _sync_manual_criteria_names_state(int(st.session_state.get("manual_criteria_count", 4) or 4), lang)
            with st.form("manual_structure_form"):
                _quick_col1, _quick_col2 = st.columns(2)
                with _quick_col1:
                    _manual_rows = int(
                        st.number_input(
                            tt("Alternatif sayısı", "Number of alternatives"),
                            min_value=2,
                            max_value=500,
                            value=int(st.session_state.get("manual_row_count", 8) or 8),
                            step=1,
                            key="manual_row_count",
                        )
                    )
                with _quick_col2:
                    _manual_crit_count = int(
                        st.number_input(
                            tt("Kriter sayısı", "Criterion count"),
                            min_value=2,
                            max_value=30,
                            value=int(st.session_state.get("manual_criteria_count", 4) or 4),
                            step=1,
                            key="manual_criteria_count",
                        )
                    )
                st.caption(
                    tt(
                        "Hiz icin tablo artik sabit satirli calisir. Daha fazla satir gerekirse asagidaki hizli ekleme dugmelerini kullanin.",
                        "For speed, the editor now uses a fixed row count. Use the quick add-row buttons below when you need more rows.",
                    )
                )

                st.markdown(f"**{tt('2. Yapıyı özelleştir', '2. Customize structure')}**")
                _cfg_col1, _cfg_col2 = st.columns(2)
                with _cfg_col1:
                    _manual_entity_col = st.text_input(
                        tt("Etiket sütunu", "Label column"),
                        value=st.session_state.get("manual_entity_col", tt("Alternatif", "Alternative")),
                        key="manual_entity_col",
                    ).strip() or tt("Alternatif", "Alternative")
                with _cfg_col2:
                    _manual_has_year = st.checkbox(
                        tt("Panel için yıl sütunu", "Year column for panel"),
                        value=bool(st.session_state.get("manual_has_year", False)),
                        key="manual_has_year",
                    )

                _manual_year_col = None
                if _manual_has_year:
                    _manual_year_col = st.text_input(
                        tt("Yıl sütunu adı", "Year column name"),
                        value=st.session_state.get("manual_year_col", tt("Yıl", "Year")),
                        key="manual_year_col",
                    ).strip() or tt("Yıl", "Year")

                _sync_manual_criteria_names_state(_manual_crit_count, lang)
                _sync_manual_criteria_input_keys(_manual_crit_count, lang)
                _name_cols = st.columns(2)
                _typed_names: List[str] = []
                for idx in range(_manual_crit_count):
                    with _name_cols[idx % 2]:
                        _typed_names.append(
                            st.text_input(
                                f"{tt('Kriter', 'Criterion')} {idx + 1}",
                                key=f"manual_criteria_name_{idx + 1}",
                            ).strip()
                        )
                _apply_structure = st.form_submit_button(
                    tt("🧱 Yapıyı tabloya uygula", "🧱 Apply structure to table"),
                    use_container_width=True,
                )

            if _apply_structure:
                _manual_criteria_now = _parse_manual_criteria(", ".join(_typed_names), _manual_crit_count, lang)
                st.session_state["manual_criteria_names"] = ", ".join(_manual_criteria_now)
                st.session_state["manual_criteria_inputs_seed"] = "||".join(_manual_criteria_now)
                st.session_state["manual_entry_df"] = _seed_manual_entry_df(
                    st.session_state.get("manual_entry_df"),
                    _manual_rows,
                    _manual_entity_col,
                    _manual_criteria_now,
                    _manual_year_col,
                )
                st.rerun()

            _manual_rows = int(st.session_state.get("manual_row_count", 8) or 8)
            _manual_crit_count = int(st.session_state.get("manual_criteria_count", 4) or 4)
            _manual_entity_col = str(st.session_state.get("manual_entity_col", tt("Alternatif", "Alternative"))).strip() or tt("Alternatif", "Alternative")
            _manual_has_year = bool(st.session_state.get("manual_has_year", False))
            _manual_year_col = str(st.session_state.get("manual_year_col", tt("Yıl", "Year"))).strip() if _manual_has_year else None
            _manual_criteria = _parse_manual_criteria(
                st.session_state.get("manual_criteria_names", ""),
                _manual_crit_count,
                lang,
            )

            st.info(
                tt(
                    f"Mevcut yapı: {_manual_rows} satır, {_manual_crit_count} kriter, etiket sütunu '{_manual_entity_col}'.",
                    f"Current setup: {_manual_rows} rows, {_manual_crit_count} criteria, label column '{_manual_entity_col}'.",
                )
            )
            _action_cols = st.columns(4)
            with _action_cols[0]:
                if st.button(tt("🧪 Örnek değer yükle", "🧪 Load sample values"), key="btn_fill_manual_sample_main", use_container_width=True):
                    _base_manual = _seed_manual_entry_df(
                        st.session_state.get("manual_entry_df"),
                        _manual_rows,
                        _manual_entity_col,
                        _manual_criteria,
                        _manual_year_col,
                    )
                    st.session_state["manual_entry_df"] = _fill_manual_demo_data(
                        _base_manual,
                        _manual_entity_col,
                        _manual_criteria,
                        _manual_year_col,
                    )
                    st.rerun()
            with _action_cols[1]:
                if st.button(tt("🧹 Tabloyu sıfırla", "🧹 Reset table"), key="btn_reset_manual_data_main", use_container_width=True):
                    st.session_state["manual_entry_df"] = _seed_manual_entry_df(
                        None,
                        _manual_rows,
                        _manual_entity_col,
                        _manual_criteria,
                        _manual_year_col,
                    )
                    st.rerun()
            with _action_cols[2]:
                if st.button(tt("➕ 5 satır", "➕ Add 5 rows"), key="btn_add_5_manual_rows", use_container_width=True):
                    _new_rows = _manual_rows + 5
                    st.session_state["manual_row_count"] = _new_rows
                    st.session_state["manual_entry_df"] = _seed_manual_entry_df(
                        st.session_state.get("manual_entry_df"),
                        _new_rows,
                        _manual_entity_col,
                        _manual_criteria,
                        _manual_year_col,
                    )
                    st.rerun()
            with _action_cols[3]:
                if st.button(tt("➕ 20 satır", "➕ Add 20 rows"), key="btn_add_20_manual_rows", use_container_width=True):
                    _new_rows = _manual_rows + 20
                    st.session_state["manual_row_count"] = _new_rows
                    st.session_state["manual_entry_df"] = _seed_manual_entry_df(
                        st.session_state.get("manual_entry_df"),
                        _new_rows,
                        _manual_entity_col,
                        _manual_criteria,
                        _manual_year_col,
                    )
                    st.rerun()
            st.caption(
                tt(
                    "Ipuçu: Yapıyı bir kez kurduktan sonra en hizli akis toplu yapistir, ikinci en hizli akis ise tabloyu duzenleyip Kaydet demektir.",
                    "Tip: After setting the structure once, the fastest workflow is bulk paste, and the second fastest is editing the grid and clicking Save.",
                )
            )

        with _editor_col:
            _manual_seed = _seed_manual_entry_df(
                st.session_state.get("manual_entry_df"),
                _manual_rows,
                _manual_entity_col,
                _manual_criteria,
                _manual_year_col,
            )
            _manual_col_cfg: Dict[str, Any] = {
                _manual_entity_col: st.column_config.TextColumn(
                    tt("Etiket", "Label"),
                    help=tt("Alternatif, ülke veya firma adı yazın.", "Enter alternative, country, or company label."),
                )
            }
            if _manual_year_col:
                _manual_col_cfg[_manual_year_col] = st.column_config.NumberColumn(
                    _manual_year_col,
                    format="%d",
                    step=1,
                )
            for _crit in _manual_criteria:
                _manual_col_cfg[_crit] = st.column_config.NumberColumn(_crit, format="%.4f")

            st.markdown(f"**{tt('3. Hızlı veri girişi', '3. Faster data entry')}**")
            st.caption(
                tt(
                    "Buyuk tablolar icin once toplu yapistir kullanin. Hucre editorunde ise degisiklikler ancak Kaydet dugmesine bastiginizda islenir.",
                    "For larger tables, use bulk paste first. In the grid editor, changes are applied only when you click Save.",
                )
            )
            _paste_tab, _grid_tab = st.tabs(
                [
                    tt("⚡ Toplu yapıştır", "⚡ Bulk paste"),
                    tt("🧮 Hücre editörü", "🧮 Grid editor"),
                ]
            )
            with _paste_tab:
                _expected_cols = _manual_schema_columns(_manual_entity_col, _manual_criteria, _manual_year_col)
                _sample_row = ["A01"] + ([str(2024)] if _manual_year_col else []) + [f"{10.0 + idx:.2f}" for idx in range(len(_manual_criteria))]
                with st.form("manual_paste_form"):
                    st.checkbox(
                        tt("İlk satır başlık içeriyor", "First row includes headers"),
                        key="manual_paste_has_header",
                    )
                    st.text_area(
                        tt("Excel veya tablodan veri yapıştırın", "Paste data from Excel or another table"),
                        key="manual_bulk_paste_text",
                        height=220,
                        placeholder="\n".join(["\t".join(_expected_cols), "\t".join(_sample_row)]),
                    )
                    _apply_bulk_paste = st.form_submit_button(
                        tt("⚡ Yapıştırılan veriyi tabloya uygula", "⚡ Apply pasted data to the table"),
                        use_container_width=True,
                    )

                if _apply_bulk_paste:
                    try:
                        _parsed_manual = _parse_manual_paste_text(
                            st.session_state.get("manual_bulk_paste_text", ""),
                            _manual_entity_col,
                            _manual_criteria,
                            _manual_year_col,
                            has_header=bool(st.session_state.get("manual_paste_has_header", True)),
                        )
                    except ValueError as manual_paste_exc:
                        st.error(str(manual_paste_exc))
                    else:
                        st.session_state["manual_entry_df"] = _parsed_manual
                        st.session_state["manual_row_count"] = max(int(len(_parsed_manual)), 2)
                        access.track_event(
                            "manual_bulk_paste_applied",
                            {
                                "rows": int(len(_parsed_manual)),
                                "columns": int(len(_parsed_manual.columns)),
                            },
                        )
                        st.rerun()

                st.caption(
                    tt(
                        "Excel kopyalari genellikle sekme ayracli olur; bu alan sekme, noktalı virgul ve virgul ayraclarini dener.",
                        "Excel copies are usually tab-delimited; this field tries tab, semicolon, and comma separators.",
                    )
                )
                with st.container():
                    render_table(_manual_seed.head(8))
                if st.button(_data_ready_button_label(), key="btn_use_manual_data_from_paste_main", use_container_width=True):
                    try:
                        _prepared_manual = _prepare_manual_entry_df(
                            _manual_seed,
                            _manual_entity_col,
                            _manual_criteria,
                            _manual_year_col,
                        )
                    except ValueError as manual_exc:
                        st.error(str(manual_exc))
                    else:
                        st.session_state["manual_entry_df"] = _prepared_manual
                        _activate_data_source(_prepared_manual, "manual_entry", "manual")
                        st.rerun()

            with _grid_tab:
                with st.form("manual_editor_form"):
                    _manual_edited = st.data_editor(
                        _manual_seed,
                        hide_index=True,
                        num_rows="fixed",
                        use_container_width=True,
                        height=470,
                        key="manual_entry_editor",
                        column_config=_manual_col_cfg,
                    )
                    _grid_btn_col1, _grid_btn_col2 = st.columns(2)
                    with _grid_btn_col1:
                        _save_manual_grid = st.form_submit_button(
                            tt("💾 Tablo düzenlemelerini kaydet", "💾 Save table edits"),
                            use_container_width=True,
                        )
                    with _grid_btn_col2:
                        _save_and_continue_manual = st.form_submit_button(
                            tt("✅ Kaydet ve devam et", "✅ Save and continue"),
                            type="primary",
                            use_container_width=True,
                        )

                if _save_manual_grid or _save_and_continue_manual:
                    st.session_state["manual_entry_df"] = _manual_edited.copy()
                    if _save_and_continue_manual:
                        try:
                            _prepared_manual = _prepare_manual_entry_df(
                                _manual_edited,
                                _manual_entity_col,
                                _manual_criteria,
                                _manual_year_col,
                            )
                        except ValueError as manual_exc:
                            st.error(str(manual_exc))
                        else:
                            access.track_event(
                                "manual_grid_saved",
                                {
                                    "rows": int(len(_manual_edited)),
                                    "columns": int(len(_manual_edited.columns)),
                                },
                            )
                            st.session_state["manual_entry_df"] = _prepared_manual
                            _activate_data_source(_prepared_manual, "manual_entry", "manual")
                            st.rerun()
                    else:
                        access.track_event(
                            "manual_grid_saved",
                            {
                                "rows": int(len(_manual_edited)),
                                "columns": int(len(_manual_edited.columns)),
                            },
                        )
                        st.rerun()

                st.caption(
                    tt(
                        "Bu editor hiz icin sabit satirli calisir. Daha fazla satir gerekirse soldaki +5 veya +20 dugmelerini kullanin.",
                        "This editor uses a fixed row count for speed. If you need more rows, use the +5 or +20 buttons on the left.",
                    )
                )

            _manual_stats = _manual_editor_stats(_manual_seed, _manual_entity_col, _manual_criteria, _manual_year_col)
            _metric_col1, _metric_col2, _metric_col3 = st.columns(3)
            _metric_col1.metric(tt("Dolu satır", "Filled rows"), _manual_stats["filled_rows"])
            _metric_col2.metric(tt("Kriter", "Criteria"), _manual_stats["criteria_count"])
            _metric_col3.metric(tt("Sayısal doluluk", "Numeric fill"), f"%{_manual_stats['fill_rate']*100:.0f}")

            if _manual_year_col:
                st.caption(
                    tt(
                        "Panel veri için her satırda yıl değeri bulunduğunu kontrol edin.",
                        "For panel data, make sure each row includes a year value.",
                    )
                )
            else:
                st.caption(
                    tt(
                        "Etiket sütunu boş kalırsa sistem otomatik olarak A01, A02, A03... üretir.",
                        "If the label column is blank, the app will auto-generate A01, A02, A03...",
                    )
                )

def _input_format_notes(lang: str, panel: bool = False) -> pd.DataFrame:
    if panel:
        rows = [
            {
                _xl_text(lang, "Başlık", "Item"): _xl_text(lang, "Kimlik sütunu", "Entity column"),
                _xl_text(lang, "Açıklama", "Description"): _xl_text(lang, "Her satırda alternatif/ülke/firma adı yer alabilir.", "Each row can include an alternative/country/company name."),
            },
            {
                _xl_text(lang, "Başlık", "Item"): _xl_text(lang, "Yıl sütunu", "Year column"),
                _xl_text(lang, "Açıklama", "Description"): _xl_text(lang, "Panel analiz için dönem/yıl bilgisini ayrı sütunda verin.", "Provide period/year information in a separate column for panel analysis."),
            },
            {
                _xl_text(lang, "Başlık", "Item"): _xl_text(lang, "Kriterler", "Criteria"),
                _xl_text(lang, "Açıklama", "Description"): _xl_text(lang, "Analizde kullanılacak kriter sütunları sayısal olmalıdır.", "Criterion columns used in the analysis must be numeric."),
            },
        ]
    else:
        rows = [
            {
                _xl_text(lang, "Başlık", "Item"): _xl_text(lang, "İlk sütun", "First column"),
                _xl_text(lang, "Açıklama", "Description"): _xl_text(lang, "Alternatif adı metin olabilir; sistem bunu etiket olarak kullanır.", "Alternative names can be text; the app uses them as labels."),
            },
            {
                _xl_text(lang, "Başlık", "Item"): _xl_text(lang, "Kriter sayısı", "Number of criteria"),
                _xl_text(lang, "Açıklama", "Description"): _xl_text(lang, "Analiz için en az 2 sayısal kriter gerekir.", "At least 2 numeric criteria are required."),
            },
            {
                _xl_text(lang, "Başlık", "Item"): _xl_text(lang, "Maliyet/Fayda", "Cost/Benefit"),
                _xl_text(lang, "Açıklama", "Description"): _xl_text(lang, "Kriter yönlerini sonraki adımda fayda veya maliyet olarak işaretleyebilirsiniz.", "You can mark criterion directions as benefit or cost in the next step."),
            },
        ]
    return pd.DataFrame(rows)

def generate_input_template_excel(lang: str = "TR", panel: bool = False) -> bytes:
    output = io.BytesIO()
    sample = sample_panel_dataset_en().head(20) if panel else (sample_dataset_en() if lang == "EN" else sample_dataset())
    sheet_data = _xl_text(lang, "Ornek_Veri", "Sample_Data")
    sheet_notes = _xl_text(lang, "Format_Notlari", "Format_Notes")
    with pd.ExcelWriter(output, engine="xlsxwriter") as writer:
        sample.to_excel(writer, sheet_name=sheet_data, index=False)
        _input_format_notes(lang, panel=panel).to_excel(writer, sheet_name=sheet_notes, index=False)
        workbook = writer.book
        header_fmt = workbook.add_format({"bold": True, "bg_color": "#1d3557", "font_color": "white", "border": 1})
        cell_fmt = workbook.add_format({"border": 1})
        wrap_fmt = workbook.add_format({"border": 1, "text_wrap": True, "valign": "top"})
        for sheet_name, df in [(sheet_data, sample), (sheet_notes, _input_format_notes(lang, panel=panel))]:
            ws = writer.sheets[sheet_name]
            ws.freeze_panes(1, 0)
            ws.set_row(0, 22, header_fmt)
            for idx, col in enumerate(df.columns):
                max_len = max(len(str(col)), int(df[col].astype(str).map(len).max() if not df.empty else 0))
                width = min(40, max(14, max_len + 2))
                ws.set_column(idx, idx, width, wrap_fmt if sheet_name == sheet_notes else cell_fmt)
    return output.getvalue()

def guess_direction(col_name: str) -> str:
    lowered = col_name.lower()
    cost_keywords = ["maliyet", "risk", "süre", "borç", "gider", "hata", "şikayet", "kayip", "loss", "cost", "time", "defect"]
    return "▼ Min (Maliyet)" if any(k in lowered for k in cost_keywords) else "▲ Max (Fayda)"

def clean_dataframe(df: pd.DataFrame, missing_strategy: str, clip_outliers: bool) -> pd.DataFrame:
    out = df.copy()
    numeric_cols = out.select_dtypes(include=[np.number]).columns.tolist()
    if missing_strategy in {"Sil", "Drop"}:
        out = out.dropna(subset=numeric_cols)
    else:
        for col in numeric_cols:
            if out[col].isna().any():
                if missing_strategy in {"Medyan", "Median"}:
                    out[col] = out[col].fillna(out[col].median())
                elif missing_strategy in {"Ortalama", "Mean"}:
                    out[col] = out[col].fillna(out[col].mean())
                elif missing_strategy in {"Interpolasyon", "Interpolation"}:
                    out[col] = out[col].interpolate(method="linear", limit_direction="both")
                    out[col] = out[col].fillna(out[col].median())
                else:
                    out[col] = out[col].fillna(0)
    for col in numeric_cols:
        if clip_outliers:
            q1, q3 = out[col].quantile(0.25), out[col].quantile(0.75)
            iqr = q3 - q1
            out[col] = np.clip(out[col], q1 - 1.5 * iqr, q3 + 1.5 * iqr)
    return out

def recommend_parameter_defaults(selected_methods: List[str], n_alt: int, n_crit: int, uses_fuzzy: bool) -> Dict[str, float]:
    defaults = {
        "vikor_v": 0.50,
        "waspas_lambda": 0.50,
        "codas_tau": 0.02,
        "cocoso_lambda": 0.50,
        "gra_rho": 0.50,
        "promethee_q": 0.05,
        "promethee_p": 0.30,
        "promethee_s": 0.20,
        "promethee_pref_func": "linear",
        "fuzzy_spread": 0.10,
        "sensitivity_iterations": 800,
        "sensitivity_sigma": 0.12,
    }
    if n_alt <= 8:
        defaults["sensitivity_iterations"] = 500
        defaults["sensitivity_sigma"] = 0.14
    elif n_alt >= 30:
        defaults["sensitivity_iterations"] = 1500
        defaults["sensitivity_sigma"] = 0.10
    if n_crit >= 10:
        defaults["sensitivity_sigma"] = max(0.08, defaults["sensitivity_sigma"] - 0.01)
    if uses_fuzzy:
        defaults["fuzzy_spread"] = 0.08 if n_alt >= 20 else 0.10
    if any("CODAS" in m for m in selected_methods) and n_alt > 20:
        defaults["codas_tau"] = 0.015
    return defaults

def _safe_spearman(x: np.ndarray, y: np.ndarray) -> float:
    if not _ensure_spearman_support():
        return 0.0
    try:
        rho, _ = spearmanr(x, y)
        if rho is None or not np.isfinite(rho):
            return 0.0
        return float(rho)
    except Exception:
        return 0.0

@st.cache_data(show_spinner=False)
def _encoded_image_b64(path_str: str) -> str | None:
    path = Path(path_str)
    if not path.exists():
        return None
    try:
        with open(path, "rb") as image_file:
            return base64.b64encode(image_file.read()).decode("utf-8")
    except Exception:
        return None

@st.cache_data(show_spinner=False)
def compute_weight_robustness(
    data: pd.DataFrame,
    criteria: List[str],
    criteria_types: Dict[str, str],
    weight_method: str,
    base_weights: Dict[str, float],
    bootstrap_n: int = 160,
    fuzzy_spread: float = 0.10,
) -> Dict[str, Any]:
    base_vec = np.asarray([float(base_weights[c]) for c in criteria], dtype=float)
    concentration = float(np.sum(base_vec ** 2))
    eff_n = float(1.0 / concentration) if concentration > 0 else 0.0
    sorted_w = sorted(base_weights.items(), key=lambda kv: kv[1], reverse=True)
    top1, top1_w = sorted_w[0]
    top2_w = sorted_w[1][1] if len(sorted_w) > 1 else np.nan
    dominance_ratio = float(top1_w / top2_w) if np.isfinite(top2_w) and top2_w > 0 else np.nan

    loo_rows: List[Dict[str, Any]] = []
    if len(data) >= 5:
        for alt in data.index:
            sub = data.drop(index=alt)
            if len(sub) < 4:
                continue
            try:
                w_sub, _ = me.compute_objective_weights(sub, criteria, criteria_types, weight_method, fuzzy_spread=fuzzy_spread)
                sub_vec = np.asarray([float(w_sub[c]) for c in criteria], dtype=float)
                rho = _safe_spearman(base_vec, sub_vec)
                top_sub = max(w_sub, key=w_sub.get)
                loo_rows.append(
                    {
                        "Alternatif": str(alt),
                        "SpearmanRho": rho,
                        "TopKriterUyumu": int(top_sub == top1),
                        "MaksMutlakFark": float(np.max(np.abs(sub_vec - base_vec))),
                    }
                )
            except Exception:
                continue
    loo_df = pd.DataFrame(loo_rows)
    loo_mean = float(loo_df["SpearmanRho"].mean()) if not loo_df.empty else np.nan
    loo_min = float(loo_df["SpearmanRho"].min()) if not loo_df.empty else np.nan
    loo_top_match = float(loo_df["TopKriterUyumu"].mean()) if not loo_df.empty else np.nan

    rng = np.random.default_rng(42)
    boot_rows: List[Dict[str, float]] = []
    n_boot = int(max(80, min(bootstrap_n, 320)))
    for _ in range(n_boot):
        idx = rng.integers(0, len(data), len(data))
        sample = data.iloc[idx]
        try:
            w_b, _ = me.compute_objective_weights(sample, criteria, criteria_types, weight_method, fuzzy_spread=fuzzy_spread)
            boot_rows.append({c: float(w_b[c]) for c in criteria})
        except Exception:
            continue
    boot_df = pd.DataFrame(boot_rows)
    boot_summary_rows: List[Dict[str, Any]] = []
    if not boot_df.empty:
        for c in criteria:
            col = boot_df[c]
            boot_summary_rows.append(
                {
                    "Kriter": c,
                    "OrtalamaAğırlık": float(col.mean()),
                    "StdSapma": float(col.std(ddof=1)),
                    "AltYuzde5": float(col.quantile(0.05)),
                    "UstYuzde95": float(col.quantile(0.95)),
                }
            )
    boot_summary = pd.DataFrame(boot_summary_rows).sort_values("OrtalamaAğırlık", ascending=False) if boot_summary_rows else pd.DataFrame()

    return {
        "base_top_criterion": top1,
        "concentration": concentration,
        "effective_criteria_n": eff_n,
        "dominance_ratio": dominance_ratio,
        "leave_one_out": loo_df,
        "loo_mean_rho": loo_mean,
        "loo_min_rho": loo_min,
        "loo_top_match_rate": loo_top_match,
        "bootstrap_summary": boot_summary,
        "bootstrap_n": len(boot_df),
    }

@st.cache_data(show_spinner=False)
def run_full_analysis_cached(data_slice: pd.DataFrame, config_payload: Dict[str, Any]) -> Dict[str, Any]:
    return me.run_full_analysis(data_slice, me.AnalysisConfig(**config_payload))

def build_ranking_robustness_table(result: Dict[str, Any]) -> pd.DataFrame:
    ranking = result.get("ranking", {})
    rt = ranking.get("table")
    primary = ranking.get("method")
    if rt is None or rt.empty or not primary:
        return pd.DataFrame()

    alt_col = col_key(rt, "Alternatif", "Alternative")
    primary_top = str(rt.iloc[0][alt_col])
    sens = result.get("sensitivity") or {}
    primary_stab = sens.get("top_stability")

    rows: List[Dict[str, Any]] = []
    comp = result.get("comparison", {}) or {}
    spdf = comp.get("spearman_matrix")
    tops = comp.get("top_alternatives")
    top_map: Dict[str, str] = {}
    if isinstance(tops, pd.DataFrame) and not tops.empty:
        t_m_col = col_key(tops, "Yöntem", "Method")
        t_a_col = col_key(tops, "BirinciAlternatif", "TopAlternative")
        top_map = dict(zip(tops[t_m_col].astype(str), tops[t_a_col].astype(str)))
    top_map.setdefault(primary, primary_top)

    rho_map: Dict[str, float] = {}
    if isinstance(spdf, pd.DataFrame) and not spdf.empty:
        m_col = col_key(spdf, "Yöntem", "Method")
        mat = spdf.set_index(m_col)
        for m in mat.columns:
            if primary in mat.index and m in mat.columns:
                val = mat.loc[primary, m]
                rho_map[str(m)] = float(val) if np.isfinite(val) else np.nan

    methods = list(top_map.keys()) if top_map else [primary]
    if primary not in methods:
        methods = [primary] + methods

    for m in methods:
        rho = rho_map.get(m, 1.0 if m == primary else np.nan)
        top_alt = top_map.get(m, "")
        same_top = str(top_alt) == str(primary_top)
        if np.isfinite(rho):
            if rho >= 0.85:
                mean_txt = tt("Bu yöntem ana sonuca çok yakın davranıyor.", "This method behaves very close to the main result.")
            elif rho >= 0.70:
                mean_txt = tt("Genel tablo benzer, alt sıralarda fark olabilir.", "The overall picture is similar, with possible differences in lower ranks.")
            else:
                mean_txt = tt("Bu yöntem farklı bir bakış sunuyor; sonucu dikkatle karşılaştırın.", "This method gives a different perspective; compare results carefully.")
        else:
            mean_txt = tt("Bu yöntem için yeterli karşılaştırma verisi oluşmadı.", "There is not enough comparison data for this method.")

        top_txt = (
            tt("Lider alternatif ana yöntemle aynı.", "The top alternative is the same as in the main method.")
            if same_top
            else tt("Lider alternatif farklı. Karar öncesi iki yöntemi birlikte okuyun.", "The top alternative is different. Review both methods together before deciding.")
        )
        if m == primary and primary_stab is not None:
            stab_txt = tt(
                f"Monte Carlo'da birincilik koruma oranı %{float(primary_stab)*100:.1f}.",
                f"Monte Carlo first-place retention is {float(primary_stab)*100:.1f}%.",
            )
        else:
            stab_txt = tt("Bu yöntemde Monte Carlo ayrı çalıştırılmadı.", "Monte Carlo was not run separately for this method.")

        rows.append(
            {
                tt("Yöntem", "Method"): m,
                tt("Ana Yöntemle Uyum (ρ)", "Agreement with Primary (ρ)"): (f"{rho:.3f}" if np.isfinite(rho) else "—"),
                tt("Lider Aynı mı?", "Same Top Alternative?"): (tt("Evet", "Yes") if same_top else tt("Hayır", "No")),
                tt("Ne Anlama Geliyor?", "What Does It Mean?"): f"{mean_txt} {top_txt} {stab_txt}",
            }
        )
    return pd.DataFrame(rows)

def get_method_steps(method: str, stage: str) -> str:
    is_en = st.session_state.get("ui_lang") == "EN"

    if is_en:
        weight_steps = {
            "Entropy": "1) Normalize\n2) Compute entropy\n3) Compute divergence\n4) Normalize final weights",
            "CRITIC": "1) Min-max normalization\n2) Compute standard deviations\n3) Build correlation structure\n4) Generate information-content weights",
            "Standart Sapma": "1) Normalize\n2) Compute criterion standard deviations\n3) Convert relative dispersion to weights",
            "Eşit Ağırlık": "1) Assign equal weight to each criterion\n2) Normalize (sum=1)\n3) Proceed to ranking",
            "Manuel Ağırlık": "1) Collect raw user-defined importance values\n2) Normalize valid ratios automatically\n3) Proceed to ranking",
            "MEREC": "1) Positive transformation\n2) Log-based performance matrix\n3) Criterion-removal effect\n4) Normalize effects to weights",
            "LOPCOW": "1) Normalize\n2) Compute RMS/std terms\n3) Logarithmic percentage variability\n4) Normalize weights",
            "PCA": "1) Standardize\n2) Extract principal components\n3) Combine loadings and explained variance\n4) Produce criterion weights",
            "IDOCRIW": "1) Compute entropy weights\n2) Build CILOS loss matrix\n3) Integrate entropy and CILOS components\n4) Normalize final hybrid weights",
            "Fuzzy IDOCRIW": "1) Build lower/middle/upper fuzzy scenarios\n2) Compute entropy and CILOS in each scenario\n3) Average hybrid scenario weights\n4) Normalize final fuzzy IDOCRIW weights",
        }
        ranking_steps = {
            "TOPSIS": "1) Vector normalization\n2) Weighted matrix\n3) Define PIS/NIS\n4) Distances and closeness score",
            "VIKOR": "1) Best/worst criterion values\n2) Compute S and R measures\n3) Compromise index Q\n4) Lower Q ranks higher",
            "EDAS": "1) Average solution\n2) PDA/NDA distances\n3) Weighted totals\n4) Combined normalized score",
            "CODAS": "1) Normalize\n2) Distances to negative ideal\n3) Euclidean + Manhattan separation\n4) Normalize H score",
            "COPRAS": "1) Column-based normalization\n2) Weighted normalized matrix\n3) Compute S+ and S-\n4) Rank by Q_i",
            "OCRA": "1) Min-max relative performance\n2) Benefit/cost competitiveness components\n3) Weighted total score\n4) Final ranking",
            "ARAS": "1) Add ideal row\n2) Normalize\n3) Weighted totals\n4) Rank by utility ratio K",
            "SAW": "1) Benefit/cost normalization\n2) Weighted additive utility\n3) Aggregate total utility\n4) Rank by utility score",
            "WPM": "1) Benefit/cost normalization\n2) Weighted multiplicative utility\n3) Compute product utility\n4) Rank by utility score",
            "MAUT": "1) Min-max utility transformation\n2) Weighted utility matrix\n3) Aggregate expected utility\n4) Rank by utility score",
            "WASPAS": "1) Normalize\n2) Compute WSM and WPM\n3) Combine with lambda\n4) Higher combined score is better",
            "MOORA": "1) Vector normalization\n2) Weighting\n3) Benefit-cost difference\n4) Rank by score",
            "MABAC": "1) Min-max normalization\n2) Border approximation area\n3) Distance matrices\n4) Rank by total score",
            "MARCOS": "1) Add ideal/anti-ideal rows\n2) Normalize + weight\n3) Utility ratios\n4) Final utility ranking",
            "CoCoSo": "1) Normalize\n2) Compute S and P components\n3) Ka, Kb, Kc compromises\n4) Rank by composite score",
            "PROMETHEE": "1) Pairwise comparisons\n2) Preference indices\n3) Phi+ / Phi- flows\n4) Complete ranking by net flow",
            "GRA": "1) Min-max normalization\n2) Distance to reference\n3) Grey relational coefficients\n4) Rank by relational grade",
        }
        if stage == "weight":
            return weight_steps.get(method, "Detailed steps are available for this weighting method.")
        if method.startswith("Fuzzy "):
            base = method.replace("Fuzzy ", "")
            return "1) Fuzzify data (TFN)\n2) Build defuzzified representation\n3) " + ranking_steps.get(base, "Core ranking steps") + "\n4) Interpret final ranking"
        return ranking_steps.get(method, "Detailed steps are available for this ranking method.")

    weight_steps = {
        "Entropy": "1) Normalize et\n2) Entropi hesapla\n3) Ayrışma derecesi bul\n4) Ağırlıkları normalize et",
        "CRITIC": "1) Min-max normalize et\n2) Standart sapmaları bul\n3) Korelasyon yapısını çıkar\n4) Bilgi içeriğinden ağırlık üret",
        "Standart Sapma": "1) Normalize et\n2) Kriter std sapmalarını hesapla\n3) Std oranlarını ağırlığa çevir",
        "Eşit Ağırlık": "1) Her kritere eşit önem ver\n2) Toplamı 1'e normalize et\n3) Sıralama aşamasına geç",
        "Manuel Ağırlık": "1) Kullanıcıdan ham önem puanlarını al\n2) Geçerli oranları otomatik normalize et\n3) Sıralama aşamasına geç",
        "MEREC": "1) Pozitif dönüşüm\n2) Log tabanlı performans\n3) Kriter çıkarma etkisi\n4) Etkiyi normalize et",
        "LOPCOW": "1) Normalize et\n2) RMS ve std oranlarını hesapla\n3) Logaritmik yüzde değişim gücü\n4) Ağırlıkları normalize et",
        "PCA": "1) Standardize et\n2) Temel bileşenleri çıkar\n3) Yükleri ve açıklanan varyansı birleştir\n4) Kriter ağırlıklarını üret",
        "IDOCRIW": "1) Entropi ağırlıklarını hesapla\n2) CILOS kayıp matrisini kur\n3) Entropi ve CILOS'u birleştir\n4) Hibrit ağırlıkları normalize et",
        "Fuzzy IDOCRIW": "1) Alt/orta/üst bulanık senaryoları kur\n2) Her senaryoda Entropi ve CILOS hesapla\n3) Senaryo hibrit ağırlıklarını ortala\n4) Nihai bulanık IDOCRIW ağırlıklarını normalize et",
    }
    ranking_steps = {
        "TOPSIS": "1) Vektör normalize et\n2) Ağırlıklı matris\n3) PIS/NIS belirle\n4) Uzaklıklar ve yakınlık skoru",
        "VIKOR": "1) En iyi/en kötü değerler\n2) S ve R ölçütleri\n3) Q uzlaşı indeksi\n4) Q küçük olan üst sırada",
        "EDAS": "1) Ortalama çözüm\n2) PDA/NDA\n3) Ağırlıklı toplamlar\n4) Normalize birleşik skor",
        "CODAS": "1) Normalize et\n2) Negatif ideal mesafeler\n3) Öklid+Manhattan ayrışımı\n4) H skorunu normalize et",
        "COPRAS": "1) Sütun bazlı normalize et\n2) Ağırlıklı normalize matris\n3) S+ ve S- hesapla\n4) Q_i ile sırala",
        "OCRA": "1) Min-max göreli puanlar\n2) Fayda ve maliyet rekabet bileşenleri\n3) Ağırlıklı toplam skor\n4) Sıralama",
        "ARAS": "1) İdeal satır ekle\n2) Normalize et\n3) Ağırlıklı toplam\n4) Fayda oranı K ile sırala",
        "SAW": "1) Fayda/maliyet normalize et\n2) Ağırlıklı toplamsal fayda\n3) Toplam faydayı hesapla\n4) Fayda skoruna göre sırala",
        "WPM": "1) Fayda/maliyet normalize et\n2) Ağırlıklı çarpımsal fayda\n3) Çarpımsal faydayı hesapla\n4) Fayda skoruna göre sırala",
        "MAUT": "1) Min-max fayda dönüşümü\n2) Ağırlıklı fayda matrisi\n3) Toplam beklenen fayda\n4) Fayda skoruna göre sırala",
        "WASPAS": "1) Normalize et\n2) WSM ve WPM hesapla\n3) λ ile birleştir\n4) Büyük skor daha iyi",
        "MOORA": "1) Vektör normalize et\n2) Ağırlıklandır\n3) Fayda-maliyet farkı\n4) Skoru sırala",
        "MABAC": "1) Min-max normalize et\n2) Sınır yaklaşım alanı\n3) Uzaklık matrisleri\n4) Toplam skorla sırala",
        "MARCOS": "1) İdeal/anti-ideal ekle\n2) Normalize + ağırlıklandır\n3) Yararlılık oranları\n4) Nihai utility ile sırala",
        "CoCoSo": "1) Normalize et\n2) S ve P bileşenleri\n3) Ka, Kb, Kc uzlaşıları\n4) Bileşik skorla sırala",
        "PROMETHEE": "1) İkili karşılaştırma\n2) Tercih indeksleri\n3) Phi+ / Phi- akışları\n4) Net akış ile tam sıralama",
        "GRA": "1) Min-max normalize et\n2) Referansa uzaklık\n3) Gri katsayılar\n4) İlişkisel dereceyi sırala",
    }
    if stage == "weight":
        return weight_steps.get(method, "Ağırlık yöntemi için detay adımlar mevcut.")
    if method.startswith("Fuzzy "):
        base = method.replace("Fuzzy ", "")
        return "1) Veriyi bulanıklaştır (TFN)\n2) Lower/Middle/Upper senaryolarını hesapla\n3) " + ranking_steps.get(base, "Temel sıralama adımları") + "\n4) Senaryo sonuçlarını birleştir ve yorumla"
    return ranking_steps.get(method, "Sıralama yöntemi için detay adımlar mevcut.")

def _extract_detail_tables(details: Dict[str, Any]) -> Dict[str, pd.DataFrame]:
    out: Dict[str, pd.DataFrame] = {}
    for key, val in (details or {}).items():
        if isinstance(val, pd.DataFrame) and not val.empty:
            out[key] = val.copy()
        elif isinstance(val, np.ndarray):
            arr = np.asarray(val)
            if arr.ndim == 1 and arr.size > 0:
                out[key] = pd.DataFrame({key: arr})
            elif arr.ndim == 2 and arr.size > 0:
                out[key] = pd.DataFrame(arr)
    return out

def diag_rec_text(rec: Dict[str, str]) -> tuple[str, str]:
    if st.session_state.get("ui_lang") == "EN" and rec.get("text_en"):
        return rec.get("text_en", ""), rec.get("action_en", "")
    text = rec.get("text", "")
    action = rec.get("action", "")
    if st.session_state.get("ui_lang") != "EN":
        return text, action
    low = text.lower()
    if "sabit kriter" in low:
        return (
            "Constant criterion/criteria detected. These criteria do not provide discrimination power.",
            "Remove constant criteria from the active analysis set.",
        )
    if "yüksek korelasyon" in low:
        return (
            "High correlation detected among criteria, indicating information redundancy.",
            "Prefer correlation-aware methods such as CRITIC/PCA and reconsider overlapping criteria.",
        )
    if "orta düzey korelasyon" in low:
        return (
            "Moderate correlation detected among criteria.",
            "Use conflict-aware weighting methods and review criterion independence.",
        )
    if "yüksek değişkenlik" in low:
        return (
            "High dispersion detected in criteria values.",
            "Entropy or Standard Deviation weighting can effectively capture this discriminative structure.",
        )
    if "dengeli veri yapısı" in low:
        return (
            "Data profile appears balanced in dispersion and correlation.",
            "Most objective weighting and ranking methods can be applied safely.",
        )
    if "alternatif sayısı az" in low:
        return (
            "The number of alternatives is limited.",
            "Interpret Monte Carlo robustness more cautiously and cross-check with additional methods.",
        )
    if "geniş alternatif seti" in low:
        return (
            "A large alternative set is detected.",
            "Distance-based methods and robustness simulations are particularly informative in this setting.",
        )
    if "tüm kriterler fayda" in low:
        return (
            "All criteria are currently marked as benefit criteria.",
            "Review directions and verify whether at least one criterion should be modeled as cost.",
        )
    return (
        "No additional critical anomaly was detected for this item.",
        "Proceed with the current workflow and validate results with sensitivity analysis.",
    )

def reset_all() -> None:
    preserved = {
        key: st.session_state.get(key)
        for key in access.preserved_session_keys()
        if key in st.session_state
    }
    st.session_state.clear()
    st.session_state.update(preserved)
    init_state()

def _diag_score_and_label(diag: Dict[str, Any]) -> tuple[int, str]:
    score = 100
    if diag.get("constant_criteria"):
        score -= 35
    if diag.get("non_positive_criteria"):
        score -= 8
    max_corr = float(diag.get("max_corr", 0.0))
    if max_corr >= 0.85:
        score -= 30
    elif max_corr >= 0.70:
        score -= 15
    mean_cv = float(diag.get("mean_cv", 0.0))
    if mean_cv < 0.08:
        score -= 10
    outlier_ratio = float(diag.get("outlier_ratio", 0.0))
    if outlier_ratio >= 0.12:
        score -= 15
    elif outlier_ratio >= 0.06:
        score -= 8
    if float(diag.get("mean_abs_skew", 0.0)) >= 1.0:
        score -= 8
    if float(diag.get("alt_crit_ratio", 99.0)) < 2.0:
        score -= 10
    if int(diag.get("n_alt", 0)) < 5:
        score -= 10
    score = int(max(0, min(score, 100)))
    if score >= 80:
        return score, tt("Yüksek Uygunluk", "High Suitability")
    if score >= 60:
        return score, tt("Orta Uygunluk", "Moderate Suitability")
    return score, tt("Düşük Uygunluk", "Low Suitability")

# ── MATEMATİKSEL METİN OLUŞTURUCU ──────────────────────────────────────────────
def get_math_formulation(w_method: str, r_methods: List[str]) -> str:
    text = "Kullanılan Yöntemlerin Matematiksel Altyapısı:\n\n"
    
    if ("Entropi" in w_method) or ("Entropy" in w_method):
        text += "Entropi (Entropy) Ağırlıklandırma Yöntemi:\nShannon bilgi kuramına dayanan bu yöntemde, veri setindeki zıtlık ve çeşitlilik ölçülerek nesnel ağırlıklar elde edilir.\n1. Karar matrisi normalize edilir: P_ij = x_ij / Σ x_ij\n2. Her kriter için entropi değeri (e_j) hesaplanır: e_j = -k * Σ [P_ij * ln(P_ij)] (Burada k=1/ln(m))\n3. Çeşitlilik derecesi bulunur: d_j = 1 - e_j\n4. Nihai objektif ağırlıklar elde edilir: w_j = d_j / Σ d_j\n\n"
    elif "CRITIC" in w_method:
        text += "CRITIC Ağırlıklandırma Yöntemi:\nKriterlerin standart sapmasını ve birbiriyle olan korelasyonlarını (çakışmalarını) dikkate alan objektif bir yöntemdir.\n1. Karar matrisi min-max yöntemiyle doğrusal olarak normalize edilir.\n2. Kriterlerin standart sapması (σ_j) hesaplanır.\n3. Kriterler arası korelasyon matrisi (r_jk) oluşturulur.\n4. Kapsanan bilgi miktarı hesaplanır: C_j = σ_j * Σ (1 - r_jk)\n5. Nihai ağırlıklar elde edilir: w_j = C_j / Σ C_j\n\n"
    elif "Standart Sapma" in w_method:
        text += "Standart Sapma Yöntemi:\n1. Karar matrisi doğrusal olarak normalize edilir.\n2. Her bir kriterin standart sapması (σ_j) hesaplanır.\n3. Ağırlıklar, standart sapmaların oransal payına göre belirlenir: w_j = σ_j / Σ σ_j\n\n"
    
    if r_methods:
        text += f"Alternatiflerin sıralanmasında ve performanslarının değerlendirilmesinde {', '.join(r_methods)} algoritması/algoritmaları uygulanmıştır.\n\n"
        for rm in r_methods:
            if "TOPSIS" in rm and "Fuzzy" not in rm:
                text += "TOPSIS Yöntemi:\n1. Karar matrisi vektör normalizasyon yöntemi ile dönüştürülür: r_ij = x_ij / √(Σ x_ij²)\n2. Ağırlıklı matris oluşturulur: v_ij = w_j * r_ij\n3. Pozitif İdeal (A*) ve Negatif İdeal (A⁻) çözümler belirlenir.\n4. İdeallere olan Öklid uzaklıkları (S_i* ve S_i⁻) hesaplanır.\n5. Göreceli yakınlık katsayısı bulunur: C_i* = S_i⁻ / (S_i* + S_i⁻). Katsayısı 1'e en yakın olan alternatif optimum çözüm olarak kabul edilir.\n\n"
            elif "VIKOR" in rm and "Fuzzy" not in rm:
                text += "VIKOR Yöntemi:\n1. Her kriter için en iyi (f_j*) ve en kötü (f_j⁻) değerler tespit edilir.\n2. Grup faydası (S_i) ve Bireysel pişmanlık (R_i) hesaplanır.\n3. VIKOR uzlaşı indeksi hesaplanır: Q_i = v*(S_i - S*) / (S⁻ - S*) + (1-v)*(R_i - R*) / (R⁻ - R*).\n4. Q_i değeri en küçük (minimum) olan alternatif en iyi uzlaşı çözümü olarak sıralanır.\n\n"
            elif "WASPAS" in rm and "Fuzzy" not in rm:
                text += "WASPAS Yöntemi:\n1. Karar matrisi fayda/maliyet yönlerine göre doğrusal normalize edilir.\n2. Ağırlıklı Toplam Modeli (WSM) ve Ağırlıklı Çarpım Modeli (WPM) çalıştırılır.\n3. Birleşik optimizasyon skoru bulunur: Q_i = λ * WSM_i + (1 - λ) * WPM_i. Q değeri en büyük olan alternatif birinci seçilir.\n\n"
            elif "EDAS" in rm:
                text += "EDAS Yöntemi:\n1. Her bir kriterin aritmetik ortalaması (AV) hesaplanır.\n2. Ortalama değere göre Pozitif Uzaklık (PDA) ve Negatif Uzaklık (NDA) hesaplanır.\n3. Bu uzaklıklar kriter ağırlıklarıyla çarpılarak toplam skorlar (SP ve SN) elde edilir.\n4. Normalize edilmiş değerler birleştirilerek Değerlendirme Skoru (AS_i) bulunur.\n\n"
            elif "COPRAS" in rm and "Fuzzy" not in rm:
                text += "COPRAS Yöntemi:\n1. Karar matrisi sütun toplamlarına bölünerek normalize edilir.\n2. Kriter ağırlıkları ile çarpılarak ağırlıklı normalize matris elde edilir.\n3. Fayda kriterleri için S+ ve maliyet kriterleri için S- toplamları hesaplanır.\n4. Göreli önem düzeyi (Q_i) ile sıralama yapılır; büyük Q_i daha iyidir.\n\n"
            elif "OCRA" in rm and "Fuzzy" not in rm:
                text += "OCRA Yöntemi:\n1. Her kriter için min-max tabanlı göreli performans hesaplanır.\n2. Fayda ve maliyet kriterlerinin rekabet bileşen puanları ayrı üretilir.\n3. Bileşenler ağırlıklarla toplanarak operasyonel rekabet skoru elde edilir.\n4. En yüksek toplam skor en iyi alternatifi verir.\n\n"
            elif "ARAS" in rm and "Fuzzy" not in rm:
                text += "ARAS Yöntemi:\n1. Karar matrisine, en iyi performans göstergelerinden oluşan 'Optimum Alternatif (A0)' eklenir.\n2. Matris normalize edilir ve kriter ağırlıkları ile çarpılır.\n3. Her bir alternatifin optimalite fonksiyonu (S_i) hesaplanır.\n4. Fayda derecesi (K_i = S_i / S_0) formülü ile belirlenip sıralama yapılır.\n\n"
            elif "SAW" in rm and "Fuzzy" not in rm:
                text += "SAW (Simple Additive Weighting) Yöntemi:\n1. Kriterler fayda/maliyet yönüne göre normalize edilir.\n2. Normalize değerler kriter ağırlıkları ile çarpılır.\n3. Her alternatif için ağırlıklı toplam fayda skoru hesaplanır.\n4. En yüksek toplam fayda değerine sahip alternatif en iyi olarak seçilir.\n\n"
            elif "WPM" in rm and "Fuzzy" not in rm:
                text += "WPM (Weighted Product Model) Yöntemi:\n1. Kriterler fayda/maliyet yönüne göre normalize edilir.\n2. Normalize değerler kriter ağırlıkları kadar üs alınır.\n3. Her alternatif için çarpımsal fayda skoru hesaplanır.\n4. En yüksek çarpımsal fayda skoruna sahip alternatif üstte sıralanır.\n\n"
            elif "MAUT" in rm and "Fuzzy" not in rm:
                text += "MAUT (Multi-Attribute Utility Theory) Yöntemi:\n1. Her kriter doğrusal fayda fonksiyonu ile ortak ölçeğe dönüştürülür.\n2. Fayda değerleri kriter ağırlıkları ile çarpılır.\n3. Alternatifler için toplam beklenen fayda hesaplanır.\n4. Toplam faydası en yüksek alternatif en uygun seçenek kabul edilir.\n\n"
            elif "PROMETHEE" in rm:
                text += "PROMETHEE II Yöntemi:\n1. Alternatifler kriter bazında ikili olarak karşılaştırılır.\n2. Her kriter için tercih fonksiyonu ile 0-1 aralığında tercih derecesi hesaplanır.\n3. Kriter ağırlıklarıyla küresel tercih indeksi elde edilir.\n4. Pozitif akış (phi+) ve negatif akış (phi-) hesaplanır.\n5. Net akış (phi = phi+ - phi-) en büyük olan alternatif en üstte sıralanır.\n\n"
            elif "Fuzzy" in rm:
                text += f"{rm} Yöntemi:\nKlasik (kesin) sayıların ölçüm belirsizliği veya insan hata payı barındırabileceği varsayımıyla, veriler tolerans sınırları içeren bulanık (fuzzy) sayılara dönüştürülür. Değerlendirme, üçgensel bulanık kümeler kullanılarak aralıklar üzerinden yapılır ve sonuçlar durulaştırılarak (defuzzification) kesin sıralamaya çevrilir.\n\n"
            
    return text.strip()

# ── EKRAN VE ÇIKTI FONKSİYONLARI ──────────────────────────────────────────────
def render_3layer(layers: Dict[str, str], title: str = "") -> None:
    if not any(layers.values()):
        return
    if title:
        st.markdown(
            f'<p style="font-size:0.8rem;font-weight:700;color:#718096;text-transform:uppercase;">{_safe_html_text(title)}</p>',
            unsafe_allow_html=True,
        )
    if layers.get("descriptive"):
        st.markdown(
            f'<div class="layer-card layer-desc"><div class="layer-label" style="color:#3182CE">🔵 {tt("Tanımlayıcı — Ne bulundu?", "Descriptive — What was found?")}</div><p class="layer-text">{_safe_plain_commentary_html(layers["descriptive"])}</p></div>',
            unsafe_allow_html=True,
        )
    if layers.get("analytic"):
        st.markdown(
            f'<div class="layer-card layer-analytic"><div class="layer-label" style="color:#38A169">🟢 {tt("Analitik — Ne anlama gelir?", "Analytic — What does it mean?")}</div><p class="layer-text">{_safe_plain_commentary_html(layers["analytic"])}</p></div>',
            unsafe_allow_html=True,
        )
    if layers.get("normative"):
        st.markdown(
            f'<div class="layer-card layer-norm"><div class="layer-label" style="color:#D69E2E">🟡 {tt("Normatif — Ne yapılmalı?", "Normative — What should be done?")}</div><p class="layer-text">{_safe_plain_commentary_html(layers["normative"])}</p></div>',
            unsafe_allow_html=True,
        )

def load_uploaded_file(uploaded_file) -> pd.DataFrame:
    file_size = int(getattr(uploaded_file, "size", 0) or 0)
    if file_size > MAX_UPLOAD_BYTES:
        raise ValueError(
            tt(
                f"Dosya boyutu sınırı aşıldı. En fazla {MAX_UPLOAD_SIZE_MB} MB yükleyebilirsiniz.",
                f"File size limit exceeded. You can upload up to {MAX_UPLOAD_SIZE_MB} MB.",
            )
        )
    file_name = str(getattr(uploaded_file, "name", "") or "").lower()
    if file_name.endswith(".csv"):
        df = pd.read_csv(uploaded_file)
    elif file_name.endswith(".xlsx"):
        df = pd.read_excel(uploaded_file)
    elif file_name.endswith(".sav"):
        if not _ensure_pyreadstat():
            raise ValueError(
                tt(
                    "`.sav` dosyaları için `pyreadstat` kurulmalıdır. `pip install pyreadstat` veya `pip install -r requirements.txt` çalıştırın.",
                    "`.sav` files require `pyreadstat`. Run `pip install pyreadstat` or `pip install -r requirements.txt`.",
                )
            )
        try:
            uploaded_file.seek(0)
        except Exception:
            pass
        df = pd.read_spss(uploaded_file)
    else:
        raise ValueError(
            tt(
                "Desteklenmeyen dosya türü. Lütfen CSV, XLSX veya SAV yükleyin.",
                "Unsupported file type. Please upload CSV, XLSX, or SAV.",
            )
        )
    if not isinstance(df, pd.DataFrame) or df.empty:
        raise ValueError(tt("Dosya boş veya okunamadı.", "The file is empty or could not be read."))
    n_rows, n_cols = df.shape
    if n_rows > MAX_UPLOAD_ROWS:
        raise ValueError(
            tt(
                f"Satır sayısı sınırı aşıldı ({n_rows:,}). En fazla {MAX_UPLOAD_ROWS:,} satır destekleniyor.",
                f"Row limit exceeded ({n_rows:,}). A maximum of {MAX_UPLOAD_ROWS:,} rows is supported.",
            )
        )
    if n_cols > MAX_UPLOAD_COLS:
        raise ValueError(
            tt(
                f"Sütun sayısı sınırı aşıldı ({n_cols}). En fazla {MAX_UPLOAD_COLS} sütun destekleniyor.",
                f"Column limit exceeded ({n_cols}). A maximum of {MAX_UPLOAD_COLS} columns is supported.",
            )
        )
    return df

def _load_user_guide_text() -> str:
    guide_path = APP_DIR / "KULLANIM_KILAVUZU.md"
    try:
        return guide_path.read_text(encoding="utf-8").strip()
    except Exception:
        return ""

def _panel_label(value: Any) -> str:
    if pd.isna(value):
        return ""
    if isinstance(value, (int, np.integer)):
        return str(int(value))
    if isinstance(value, (float, np.floating)) and float(value).is_integer():
        return str(int(value))
    return str(value).strip()

def _guess_year_columns(df: pd.DataFrame) -> List[str]:
    if not isinstance(df, pd.DataFrame) or df.empty:
        return []
    scored: List[tuple[int, str]] = []
    for col in df.columns:
        series = df[col]
        col_low = str(col).strip().lower()
        score = 0
        if any(token in col_low for token in ["year", "yil", "yıl", "donem", "dönem", "period"]):
            score += 4
        numeric = pd.to_numeric(series, errors="coerce")
        non_na = numeric.dropna()
        if len(non_na) >= max(2, int(len(series) * 0.6)):
            unique_vals = sorted(set(non_na.astype(float).tolist()))
            if 2 <= len(unique_vals) <= min(60, max(2, len(series))):
                if all(abs(v - round(v)) < 1e-9 for v in unique_vals):
                    min_v, max_v = min(unique_vals), max(unique_vals)
                    if 1900 <= min_v <= 2100 and 1900 <= max_v <= 2100:
                        score += 3
        if score > 0:
            scored.append((score, str(col)))
    scored.sort(key=lambda item: (-item[0], item[1]))
    return [col for _, col in scored]

def _guess_entity_columns(df: pd.DataFrame, year_col: str | None = None) -> List[str]:
    if not isinstance(df, pd.DataFrame) or df.empty:
        return []
    scored: List[tuple[int, str]] = []
    year_col_s = str(year_col) if year_col is not None else None
    for col in df.columns:
        col_s = str(col)
        if year_col_s is not None and col_s == year_col_s:
            continue
        series = df[col]
        if pd.api.types.is_numeric_dtype(series):
            continue
        score = 0
        col_low = col_s.strip().lower()
        if any(token in col_low for token in [
            "alternatif", "alternative", "ulke", "ülke", "country", "entity",
            "firma", "company", "name", "isim", "ad", "id",
        ]):
            score += 4
        non_na = series.dropna().astype(str).str.strip()
        non_na = non_na[non_na != ""]
        if not non_na.empty:
            uniq = int(non_na.nunique())
            uniq_ratio = float(uniq / max(len(non_na), 1))
            if uniq_ratio >= 0.5:
                score += 2
            elif uniq_ratio >= 0.2:
                score += 1
            if uniq >= max(3, int(len(non_na) * 0.15)):
                score += 1
        if score > 0:
            scored.append((score, col_s))
    scored.sort(key=lambda item: (-item[0], item[1]))
    return [col for _, col in scored]

def _make_unique_labels(values: List[Any], fallback_prefix: str = "A") -> List[str]:
    labels: List[str] = []
    seen: Dict[str, int] = {}
    for idx, value in enumerate(values, start=1):
        raw = _panel_label(value)
        base = raw if raw else f"{fallback_prefix}{idx}"
        count = seen.get(base, 0) + 1
        seen[base] = count
        labels.append(base if count == 1 else f"{base} ({count})")
    return labels

def _sorted_panel_years(series: pd.Series) -> List[str]:
    values = {_panel_label(v) for v in series.dropna().tolist() if _panel_label(v)}
    def _sort_key(label: str) -> tuple[int, float | str]:
        try:
            return (0, float(label))
        except Exception:
            return (1, label)
    return sorted(values, key=_sort_key)

def _panel_mask(df: pd.DataFrame, year_col: str, year_label: str) -> pd.Series:
    return df[year_col].map(_panel_label) == year_label

def _short_error_text(exc: Exception, max_len: int = 180) -> str:
    msg = str(exc).replace("\n", " | ").strip()
    if not msg:
        return "Unknown error"
    return msg if len(msg) <= max_len else (msg[: max_len - 3] + "...")

def _safe_error_code(exc: Exception) -> str:
    return type(exc).__name__ if exc is not None else "Error"

def _set_docx_run_style(
    run,
    lang_code: str = "tr-TR",
    font_name: str = "Times New Roman",
    font_size: float = 12,
    *,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), font_name)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    set_docx_language(run, lang_code)

def _style_docx_paragraph(paragraph, lang_code: str = "tr-TR", font_size: float = 12) -> None:
    for run in paragraph.runs:
        _set_docx_run_style(run, lang_code=lang_code, font_size=font_size)
    paragraph.paragraph_format.space_after = Pt(6)

def _configure_apa_doc(doc: Document) -> None:
    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style.font.size = Pt(12)

def add_table_to_doc(doc: Document, df: pd.DataFrame, max_rows: int = 25, lang_code: str = "tr-TR") -> None:
    if df is None or df.empty:
        return
    use_df = df.copy().head(max_rows)
    table = doc.add_table(rows=1, cols=len(use_df.columns))
    table.style = "Table Grid"
    for i, col in enumerate(use_df.columns):
        table.rows[0].cells[i].text = str(col)
        for paragraph in table.rows[0].cells[i].paragraphs:
            for run in paragraph.runs:
                _set_docx_run_style(run, lang_code=lang_code, bold=True)
    _is_tr = lang_code.startswith("tr")
    for _, row in use_df.iterrows():
        cells = table.add_row().cells
        for j, value in enumerate(row):
            if isinstance(value, (float, np.floating)) and pd.notna(value):
                s = f"{value:.3f}"
                cells[j].text = s.replace(".", ",") if _is_tr else s
            elif isinstance(value, float) and pd.isna(value):
                cells[j].text = ""
            else:
                cells[j].text = str(value)
            for paragraph in cells[j].paragraphs:
                for run in paragraph.runs:
                    _set_docx_run_style(run, lang_code=lang_code)
    doc.add_paragraph("")

def render_table(df: pd.DataFrame, max_rows: int = 300) -> None:
    if not isinstance(df, pd.DataFrame):
        st.write(df)
        return
    if df.empty:
        st.info(tt("Tablo boş.", "Table is empty."))
        return
    view = df.copy()
    if len(view) > max_rows:
        st.caption(tt(f"Tablo ilk {max_rows} satırla sınırlandı.", f"Table limited to first {max_rows} rows."))
        view = view.head(max_rows)
    _height = min(540, max(180, 36 + len(view) * 28))
    view = view.replace([np.inf, -np.inf], np.nan).fillna("")
    _html = view.to_html(index=False, escape=True)
    st.markdown(
        f"""
        <div style="max-height:{_height}px; overflow:auto; border:1px solid #D1D5DB; border-radius:8px; background:#FFFFFF;">
            <style>
                .mcdm-inline-table {{
                    width:100%;
                    border-collapse:collapse;
                    font-size:0.86rem;
                    color:#111827;
                }}
                .mcdm-inline-table thead th {{
                    position:sticky;
                    top:0;
                    background:#F3F4F6;
                    color:#111827;
                    border:1px solid #D1D5DB;
                    padding:6px 8px;
                    text-align:left;
                    white-space:nowrap;
                }}
                .mcdm-inline-table tbody td {{
                    background:#FFFFFF;
                    color:#111827;
                    border:1px solid #E5E7EB;
                    padding:6px 8px;
                    white-space:nowrap;
                }}
            </style>
            {_html.replace('<table border="1" class="dataframe">', '<table class="mcdm-inline-table">')}
        </div>
        """,
        unsafe_allow_html=True,
    )

def set_docx_language(run, lang_code: str = "tr-TR") -> None:
    rpr = run._element.get_or_add_rPr()
    lang = OxmlElement("w:lang")
    lang.set(qn("w:val"), lang_code)
    rpr.append(lang)

def _fmt(value: float, lang: str, decimals: int = 3) -> str:
    """Locale-aware float formatter: TR uses comma, EN uses period."""
    s = f"{value:.{decimals}f}"
    if lang != "EN":
        s = s.replace(".", ",")
    return s

def _format_df_numbers(df: pd.DataFrame, lang: str, decimals: int = 3) -> pd.DataFrame:
    """Return copy of df with float columns formatted as locale-aware strings."""
    out = df.copy()
    for col in out.columns:
        if pd.api.types.is_float_dtype(out[col]):
            out[col] = out[col].apply(
                lambda v: _fmt(float(v), lang, decimals) if pd.notna(v) else ""
            )
    return out

def _html_to_plain(text: str) -> str:
    if not text:
        return ""
    txt = text.replace("<br><br>", "\n\n").replace("<br>", "\n")
    txt = re.sub(r"<[^>]+>", "", txt)
    return txt.strip()

def _safe_html_text(value: Any) -> str:
    return html.escape("" if value is None else str(value), quote=True)

def _safe_plain_commentary_html(text: str) -> str:
    plain = _html_to_plain(text or "")
    return _safe_html_text(plain).replace("\n", "<br>")

def _preview_list_text(values: Any, max_items: int = 4) -> str:
    if values is None:
        iterable: List[Any] = []
    elif isinstance(values, (str, bytes)):
        iterable = [values]
    else:
        iterable = list(values)
    cleaned: List[str] = []
    for value in iterable:
        txt = str(value or "").strip()
        if txt and txt not in cleaned:
            cleaned.append(txt)
    if not cleaned:
        return tt("Belirlenmedi", "Not set")
    head = ", ".join(cleaned[:max_items])
    if len(cleaned) > max_items:
        head += tt(f" +{len(cleaned) - max_items} daha", f" +{len(cleaned) - max_items} more")
    return head

def _data_source_label(source_id: Any) -> str:
    sid = str(source_id or "")
    if sid == "manual_entry":
        return tt("Manuel giriş", "Manual entry")
    if sid.startswith("upload::"):
        return tt("Dosya yükleme", "File upload")
    if "sample_panel" in sid:
        return tt("Örnek panel veri", "Sample panel data")
    if "sample_data" in sid:
        return tt("Örnek veri", "Sample data")
    return tt("Hazır veri", "Selected data")

def _render_tracking_panel(
    title: str,
    subtitle: str,
    items: List[tuple[Any, ...]],
    *,
    note: str | None = None,
    tone: str = "blue",
    expanded: bool | None = None,
    icon: str = "🧭",
    show_title: bool | None = None,
    show_context: bool | None = None,
) -> None:
    tone_class = {
        "blue": "tracking-panel-blue",
        "green": "tracking-panel-green",
        "amber": "tracking-panel-amber",
        "rose": "tracking-panel-rose",
    }.get(str(tone or "blue"), "tracking-panel-blue")
    if show_title is None:
        show_title = expanded is None
    if show_context is None:
        show_context = _show_step_guidance_enabled()
    cards_html_parts: List[str] = []
    for item in items:
        if len(item) >= 3:
            label, value, card_tone = item[0], item[1], str(item[2] or "slate")
        else:
            label, value = item[0], item[1]
            card_tone = "slate"
        card_tone_class = {
            "blue": "tracking-card-blue",
            "green": "tracking-card-green",
            "amber": "tracking-card-amber",
            "rose": "tracking-card-rose",
            "slate": "tracking-card-slate",
        }.get(card_tone, "tracking-card-slate")
        cards_html_parts.append(
            f"""
            <div class="tracking-card {card_tone_class}">
                <div class="tracking-label">{_safe_html_text(label)}</div>
                <div class="tracking-value">{_safe_html_text(value)}</div>
            </div>
            """
        )
    cards_html = "".join(cards_html_parts)
    subtitle_html = (
        f'<div class="tracking-subtitle">{_safe_html_text(subtitle)}</div>'
        if show_context and subtitle and str(subtitle).strip()
        else ""
    )
    note_html = (
        f'<div class="tracking-note">{_safe_html_text(note)}</div>'
        if show_context and note and str(note).strip()
        else ""
    )
    title_html = f'<div class="tracking-title">{_safe_html_text(title)}</div>' if show_title else ""
    panel_mode_class = "tracking-panel-guided" if show_context else "tracking-panel-compact"
    panel_html = (
        f"""
        <div class="tracking-panel {tone_class} {panel_mode_class}">
            {title_html}
            {subtitle_html}
            <div class="tracking-grid">{cards_html}</div>
            {note_html}
        </div>
        """
    )
    if expanded is None:
        st.html(panel_html.strip())
    else:
        st.markdown(f"**{icon} {title}**")
        with st.container():
            st.html(panel_html.strip())

def _current_ui_stage() -> str:
    if st.session_state.get("analysis_result") is not None:
        return "results"
    if bool(st.session_state.get("prep_done")):
        return "step3"
    if bool(st.session_state.get("step1_done")):
        return "step2"
    return "step1"

def _render_analysis_mini_banner() -> None:
    st.markdown(
        """
        <div class="analysis-mini-banner">
            <div class="analysis-mini-banner-left">MCDM-Karar Destek Sistemi</div>
            <div class="analysis-mini-banner-right">Prof. Dr. Ömer Faruk Rençber</div>
        </div>
        """,
        unsafe_allow_html=True,
    )

def get_math_formulation_en(w_method: str, r_methods: List[str]) -> str:
    lines: List[str] = ["Mathematical and Algorithmic Framework", ""]
    lines.append(f"Objective weighting method: {method_display_name(w_method)}")
    lines.append(get_method_steps(w_method, "weight"))
    lines.append("")
    if r_methods:
        lines.append(f"Ranking methods applied: {', '.join(r_methods)}")
        lines.append("")
        for rm in r_methods:
            lines.append(f"{rm} method:")
            lines.append(get_method_steps(rm, "ranking"))
            lines.append("")
    return "\n".join(lines).strip()

def _build_doc_sections(result: Dict[str, Any], lang: str) -> Dict[str, Any]:
    base = result.get("report_sections", {})
    w_method = result["weights"]["method"]
    w_method_disp = _WEIGHT_METHOD_TR_EN.get(w_method, w_method) if lang == "EN" else w_method
    r_method = result.get("ranking", {}).get("method")
    r_method_disp = r_method or ("Not applied" if lang == "EN" else "Uygulanmadı")
    data = result.get("selected_data", pd.DataFrame())
    n_alt, n_crit = data.shape if isinstance(data, pd.DataFrame) else (0, 0)
    comp = result.get("comparison", {}) or {}
    sens = result.get("sensitivity") or {}
    spdf = comp.get("spearman_matrix")
    n_comp = int(spdf.shape[0]) if isinstance(spdf, pd.DataFrame) and not spdf.empty else 0
    n_iter = int(sens.get("n_iterations", 0)) if sens else 0
    w_ph = result.get("method_philosophy", {}).get("weight", {}) or {}
    r_ph = result.get("method_philosophy", {}).get("ranking", {}) or {}
    findings_parts = [
        ("Data Profile", "Veri Profili"),
        ("Weight Findings", "Ağırlık Bulguları"),
    ]
    findings_texts = [
        _html_to_plain(gen_stat_commentary(result)),
        _html_to_plain(gen_weight_commentary(result)),
    ]
    if r_method:
        findings_parts.append(("Ranking Findings", "Sıralama Bulguları"))
        findings_texts.append(_html_to_plain(gen_ranking_commentary(result, {})))
    if n_comp >= 2:
        findings_parts.append(("Method Comparison", "Yöntem Karşılaştırması"))
        findings_texts.append(_html_to_plain(gen_comparison_commentary(result)))
    if sens:
        findings_parts.append(("Robustness", "Sağlamlık"))
        findings_texts.append(_html_to_plain(gen_mc_commentary(result)))
    refs = _normalize_references(base.get("Kaynakça", []) or base.get("References", []), lang)

    w_simple_fb, w_academic_fb = _method_fallback_lines(w_method_disp or ("Selected method" if lang == "EN" else "Seçilen yöntem"), lang)
    r_simple_fb, r_academic_fb = _method_fallback_lines(r_method_disp or ("Selected method" if lang == "EN" else "Seçilen yöntem"), lang)
    w_simple = _lang_text(w_ph.get("simple", ""), lang, w_simple_fb)
    w_academic = _lang_text(w_ph.get("academic", ""), lang, w_academic_fb)
    r_simple = _lang_text(r_ph.get("simple", ""), lang, r_simple_fb)
    r_academic = _lang_text(r_ph.get("academic", ""), lang, r_academic_fb)

    if lang == "EN":
        findings_body = []
        for (en_title, _), text in zip(findings_parts, findings_texts):
            if text:
                findings_body.append(f"{en_title}\n{text}")
        philosophy_lines = [
            f"Weighting approach: {w_method_disp}. {w_simple}",
        ]
        if w_academic:
            philosophy_lines.append(w_academic)
        if r_method:
            philosophy_lines.append(f"Ranking approach: {r_method_disp}. {r_simple}")
            if r_academic:
                philosophy_lines.append(r_academic)
        return {
            "Objective of the Study": (
                f"This report evaluates {n_alt} alternatives across {n_crit} criteria. "
                f"The scope includes criterion weighting with {w_method_disp}"
                + (f", final ranking with {r_method_disp}" if r_method else "")
                + (f", comparison across {n_comp} methods" if n_comp >= 2 else "")
                + (f", and robustness checks based on {n_iter:,} Monte Carlo scenarios." if n_iter > 0 else ".")
            ),
            "Philosophy of the Study": "\n\n".join([line for line in philosophy_lines if line.strip()]),
            "Methodology": (
                "The workflow was kept application-focused: data profile review, weighting, ranking"
                + (", cross-method comparison" if n_comp >= 2 else "")
                + (", and robustness testing" if sens else "")
                + ". This report intentionally excludes theoretical formulas and focuses on the chosen methods, their findings, and the supporting tables."
            ),
            "Findings": "\n\n".join(findings_body),
            "References": refs,
        }

    findings_body = []
    for (_, tr_title), text in zip(findings_parts, findings_texts):
        if text:
            findings_body.append(f"{tr_title}\n{text}")
    philosophy_lines = [
        f"Ağırlıklandırma yaklaşımı: {w_method_disp}. {w_simple}",
    ]
    if w_academic:
        philosophy_lines.append(w_academic)
    if r_method:
        philosophy_lines.append(f"Sıralama yaklaşımı: {r_method_disp}. {r_simple}")
        if r_academic:
            philosophy_lines.append(r_academic)
    return {
        "Çalışmanın Amacı": (
            f"Bu rapor {n_alt} alternatif ve {n_crit} kriterden oluşan veri setini değerlendirmektedir. "
            f"Kapsamda {w_method_disp} ile kriter ağırlıkları"
            + (f", {r_method_disp} ile nihai sıralama" if r_method else "")
            + (f", {n_comp} yöntem arasında karşılaştırma" if n_comp >= 2 else "")
            + (f" ve {n_iter:,} Monte Carlo senaryosuna dayalı sağlamlık incelemesi yer almaktadır." if n_iter > 0 else ".")
        ),
        "Çalışmanın Felsefesi": "\n\n".join([line for line in philosophy_lines if line.strip()]),
        "Metodoloji": (
            "İş akışı uygulama odaklı tutulmuştur: veri profilinin okunması, ağırlıkların üretilmesi, sıralamanın kurulması"
            + (", yöntemler arası karşılaştırma" if n_comp >= 2 else "")
            + (", sağlamlık testleri" if sens else "")
            + ". Bu raporda teorik formüller özellikle çıkarılmış; yalnızca seçilen yöntemler, bulgular ve bunları destekleyen tablolar korunmuştur."
        ),
        "Bulgular": "\n\n".join(findings_body),
        "Kaynakça": refs,
    }

def _reference_notice_lines(lang: str) -> List[str]:
    if lang == "EN":
        return [
            "The references listed below are recommended sources. Responsibility for checking their accuracy, suitability, and correct use in a study belongs to the author of that study; users are advised to review all references before use.",
            "The methods presented in this system are free of charge. If you benefit from this system or my related publications in your academic or professional work, appropriate citation is kindly requested as a matter of academic courtesy.",
        ]
    return [
        "Aşağıda listelenen kaynaklar önerilen kaynaklar niteliğindedir. Bu kaynakların doğruluğunu, uygunluğunu ve bir çalışmada doğru kullanımını kontrol etme sorumluluğu o çalışmanın yazarına aittir; kullanıcıların tüm kaynakları kullanmadan önce gözden geçirmesi önerilir.",
        "Bu sistemde sunulan yöntemler ücretsizdir. Akademik veya mesleki çalışmalarınızda bu sistemden ya da ilgili yayınlarımdan yararlanmanız halinde, akademik nezaket gereği uygun atıf yapmanız kibarca rica olunur.",
    ]

def _reference_notice_html(lang: str) -> str:
    title = "Important Note" if lang == "EN" else "Önemli Not"
    body = "<br><br>".join(_reference_notice_lines(lang))
    return f'<div class="commentary-box"><strong>{title}:</strong><br>{body}</div>'

def _docx_method_name(method: str | None, lang: str) -> str:
    if not method:
        return "Not Applied" if lang == "EN" else "Uygulanmadı"
    if lang == "EN":
        return _WEIGHT_METHOD_TR_EN.get(method, method)
    return method

def _docx_heading_map(lang: str) -> Dict[str, str]:
    if lang == "EN":
        return {
            "objective": "Objective of the Study",
            "scope": "Scope of the Study",
            "philosophy": "Philosophy of the Method",
            "tables": "Tables and Interpretations",
            "conclusion": "Conclusion",
            "references": "References",
        }
    return {
        "objective": "Çalışmanın Amacı",
        "scope": "Çalışmanın Kapsamı",
        "philosophy": "Yöntemin Felsefesi",
        "tables": "Tablolar ve Yorumlar",
        "conclusion": "Sonuç",
        "references": "Kaynakça",
    }

def _docx_preview_df(selected_data: pd.DataFrame, lang: str) -> pd.DataFrame:
    if not isinstance(selected_data, pd.DataFrame) or selected_data.empty:
        return pd.DataFrame()
    preview = selected_data.copy()
    if "Alternatif" not in preview.columns and "Alternative" not in preview.columns and not isinstance(preview.index, pd.RangeIndex):
        preview = preview.reset_index()
        first_col = str(preview.columns[0])
        if first_col == "index":
            preview = preview.rename(columns={first_col: "Alternative" if lang == "EN" else "Alternatif"})
    return localize_df_lang(preview.head(15), lang)

def _docx_mean_spearman(comp: Dict[str, Any]) -> float | None:
    spdf = comp.get("spearman_matrix")
    if not isinstance(spdf, pd.DataFrame) or spdf.empty or spdf.shape[0] < 2:
        return None
    method_col = col_key(spdf, "Yöntem", "Method")
    mat = spdf.set_index(method_col)
    vals = mat.where(~np.eye(mat.shape[0], dtype=bool)).stack()
    if len(vals) == 0:
        return None
    vals = pd.to_numeric(vals, errors="coerce").dropna()
    return float(vals.mean()) if len(vals) else None

def _docx_comparison_methods(comp: Dict[str, Any]) -> List[str]:
    rank_table = comp.get("rank_table")
    if not isinstance(rank_table, pd.DataFrame) or rank_table.empty:
        return []
    alt_col = col_key(rank_table, "Alternatif", "Alternative")
    return [str(c) for c in rank_table.columns if str(c) != alt_col]

def _docx_detail_interpretation(method: str | None, lang: str) -> str:
    if lang == "EN":
        mapping = {
            "PROMETHEE": "The PROMETHEE flow table makes the internal structure of the ranking visible through positive, negative, and net flows, thereby clarifying whether the leading alternative dominates pairwise comparisons broadly or only selectively.",
            "TOPSIS": "The TOPSIS distance table supports the ranking by jointly reporting closeness to the ideal solution and distance from the negative ideal, which helps explain whether leadership is driven by proximity to the optimum rather than by a single dominant criterion.",
            "VIKOR": "The VIKOR compromise table reveals how the selected leader balances group utility, individual regret, and compromise performance, making it possible to evaluate whether the result reflects a stable compromise rather than a one-sided optimum.",
            "EDAS": "The EDAS detail table documents positive and negative deviations from the average solution, showing whether the leading alternative preserves its position by producing advantage while limiting adverse deviation.",
            "CODAS": "The CODAS distance table interprets the result through separation from the negative ideal, enabling the reader to see whether the leading alternative remains strong under both Euclidean and taxicab distance logic.",
            "GRA": "The GRA detail table demonstrates how strongly each alternative is associated with the reference profile, allowing the ranking to be interpreted through relational closeness rather than raw magnitude alone.",
            "MARCOS": "The MARCOS detail table frames the result with simultaneous reference to ideal and anti-ideal benchmarks, which strengthens the interpretability of the utility ratio underlying the final order.",
            "CoCoSo": "The CoCoSo detail table shows how additive and multiplicative compromise logics converge in the final score, making the composite nature of the leadership pattern explicit.",
            "MABAC": "The MABAC detail table shows how far each alternative stands from the border approximation area, which helps explain whether the final ranking is supported by stable criterion-level positioning.",
        }
    else:
        mapping = {
            "PROMETHEE": "PROMETHEE akış tablosu, pozitif, negatif ve net akış bileşenleri üzerinden sıralamanın iç yapısını görünür kılmakta; böylece lider alternatifin ikili karşılaştırmalarda geniş tabanlı mı yoksa sınırlı mı üstünlük kurduğu daha açık biçimde izlenebilmektedir.",
            "TOPSIS": "TOPSIS uzaklık tablosu, ideal çözüme yakınlık ile negatif idealden uzaklık bilgisini birlikte sunarak liderliğin tek bir baskın kritere değil, optimum profile görece yakınlığa dayanıp dayanmadığını açıklamaktadır.",
            "VIKOR": "VIKOR uzlaşı tablosu, seçilen liderin grup faydası, bireysel pişmanlık ve uzlaşı performansı arasında nasıl bir denge kurduğunu ortaya koymakta; böylece sonucun tek yönlü bir optimumdan mı yoksa savunulabilir bir uzlaşı çözümünden mi beslendiği anlaşılmaktadır.",
            "EDAS": "EDAS detay tablosu, ortalama çözüme göre pozitif ve negatif sapmaları birlikte raporlayarak lider alternatifin avantaj üretirken olumsuz sapmayı ne ölçüde sınırladığını göstermektedir.",
            "CODAS": "CODAS uzaklık tablosu, sonucu negatif idealden ayrışma mantığıyla yorumlamakta ve lider alternatifin hem Öklid hem Manhattan uzaklığı altında ne ölçüde güçlü kaldığını görünür kılmaktadır.",
            "GRA": "GRA detay tablosu, her alternatifin referans profile olan ilişkisel yakınlığını ortaya koymakta; böylece sıralama ham büyüklükten çok referans örüntüye benzerlik üzerinden yorumlanabilmektedir.",
            "MARCOS": "MARCOS detay tablosu, ideal ve anti-ideal referanslarla eş zamanlı karşılaştırma sağlayarak nihai yararlılık oranının hangi eksende oluştuğunu daha açık hale getirmektedir.",
            "CoCoSo": "CoCoSo detay tablosu, toplamsal ve çarpımsal uzlaşı mantıklarının nihai skorda nasıl birleştiğini göstererek liderlik deseninin bileşik yapısını açıklamaktadır.",
            "MABAC": "MABAC detay tablosu, alternatiflerin sınır yaklaşım alanına göre konumunu gösterdiği için nihai sıralamanın kriter bazlı yerleşim bakımından ne kadar tutarlı olduğunu desteklemektedir.",
        }
    if method and method.startswith("Fuzzy "):
        return (
            "The fuzzy extension preserves the logic of the selected method while explicitly incorporating uncertainty through fuzzy representations, which broadens interpretability under imprecise evaluations."
            if lang == "EN"
            else "Bulanık uzantı, seçilen yöntemin temel mantığını korurken belirsizliği bulanık temsiller üzerinden açık biçimde modele kattığı için yorum gücünü kesin olmayan değerlendirmeler altında genişletmektedir."
        )
    return mapping.get(method or "", "")

def _weight_bar_figure_bytes(weight_df: pd.DataFrame, lang: str) -> bytes | None:
    """Horizontal bar chart of criterion weights; returns PNG bytes or None."""
    if weight_df is None or weight_df.empty:
        return None
    if not _ensure_matplotlib_support():
        return None
    try:
        c_col = col_key(weight_df, "Kriter", "Criterion")
        w_col = col_key(weight_df, "Ağırlık", "Weight")
        if c_col not in weight_df.columns or w_col not in weight_df.columns:
            return None
        df = weight_df[[c_col, w_col]].copy()
        df[w_col] = pd.to_numeric(df[w_col], errors="coerce")
        df = df.dropna().sort_values(w_col, ascending=True)
        fig, ax = plt.subplots(figsize=(6, max(2.5, len(df) * 0.38)))
        ax.barh(df[c_col].astype(str), df[w_col].astype(float), color="#2E75B6")
        ax.set_xlabel("Ağırlık" if lang != "EN" else "Weight", fontsize=9)
        ax.tick_params(axis="both", labelsize=8)
        ax.spines["top"].set_visible(False)
        ax.spines["right"].set_visible(False)
        plt.tight_layout()
        buf = io.BytesIO()
        fig.savefig(buf, format="png", dpi=150, bbox_inches="tight")
        plt.close(fig)
        return buf.getvalue()
    except Exception:
        return None

def _ranking_bar_figure_bytes(ranking_df: pd.DataFrame, lang: str) -> bytes | None:
    """Vertical bar chart of alternative scores; returns PNG bytes or None."""
    if ranking_df is None or ranking_df.empty:
        return None
    if not _ensure_matplotlib_support():
        return None
    try:
        alt_col = col_key(ranking_df, "Alternatif", "Alternative")
        score_col = col_key(ranking_df, "Skor", "Score")
        if alt_col not in ranking_df.columns or score_col not in ranking_df.columns:
            return None
        df = ranking_df[[alt_col, score_col]].copy()
        df[score_col] = pd.to_numeric(df[score_col], errors="coerce")
        df = df.dropna().sort_values(score_col, ascending=False)
        fig, ax = plt.subplots(figsize=(max(4, len(df) * 0.6), 4))
        ax.bar(df[alt_col].astype(str), df[score_col].astype(float), color="#C9A227")
        ax.set_ylabel("Skor" if lang != "EN" else "Score", fontsize=9)
        ax.tick_params(axis="x", labelsize=8, rotation=45)
        ax.tick_params(axis="y", labelsize=8)
        ax.spines["top"].set_visible(False)
        ax.spines["right"].set_visible(False)
        plt.tight_layout()
        buf = io.BytesIO()
        fig.savefig(buf, format="png", dpi=150, bbox_inches="tight")
        plt.close(fig)
        return buf.getvalue()
    except Exception:
        return None

def _weight_radar_figure_bytes(weight_df: pd.DataFrame, lang: str) -> bytes | None:
    """Polar radar chart of criterion weights; returns PNG bytes or None."""
    if weight_df is None or weight_df.empty:
        return None
    if not _ensure_matplotlib_support():
        return None
    try:
        c_col = col_key(weight_df, "Kriter", "Criterion")
        w_col = col_key(weight_df, "Ağırlık", "Weight")
        if c_col not in weight_df.columns or w_col not in weight_df.columns:
            return None
        df = weight_df[[c_col, w_col]].copy()
        df[w_col] = pd.to_numeric(df[w_col], errors="coerce")
        df = df.dropna()
        if len(df) < 3:
            return None
        labels = df[c_col].astype(str).tolist()
        values = df[w_col].astype(float).tolist()
        N = len(labels)
        angles = [n / float(N) * 2 * np.pi for n in range(N)]
        values_closed = values + [values[0]]
        angles_closed = angles + [angles[0]]
        fig, ax = plt.subplots(figsize=(5, 5), subplot_kw=dict(polar=True))
        ax.plot(angles_closed, values_closed, color="#2E75B6", linewidth=2)
        ax.fill(angles_closed, values_closed, color="#2E75B6", alpha=0.25)
        ax.set_xticks(angles)
        ax.set_xticklabels(labels, fontsize=7)
        ax.set_yticklabels([])
        ax.spines["polar"].set_visible(True)
        ax.grid(color="grey", linestyle="--", linewidth=0.5, alpha=0.5)
        plt.tight_layout()
        buf = io.BytesIO()
        fig.savefig(buf, format="png", dpi=150, bbox_inches="tight")
        plt.close(fig)
        return buf.getvalue()
    except Exception:
        return None

def _mc_stability_figure_bytes(mc_df: pd.DataFrame, lang: str) -> bytes | None:
    """Horizontal bar chart of Monte Carlo first-place rates; returns PNG bytes or None."""
    if mc_df is None or mc_df.empty:
        return None
    if not _ensure_matplotlib_support():
        return None
    try:
        if "Alternatif" not in mc_df.columns or "BirincilikOranı" not in mc_df.columns:
            return None
        df = mc_df[["Alternatif", "BirincilikOranı"]].copy()
        df["BirincilikOranı"] = pd.to_numeric(df["BirincilikOranı"], errors="coerce")
        df = df.dropna().sort_values("BirincilikOranı", ascending=True)
        fig, ax = plt.subplots(figsize=(6, max(2.5, len(df) * 0.38)))
        bars = ax.barh(df["Alternatif"].astype(str), df["BirincilikOranı"] * 100, color="#4CAF50")
        for bar, val in zip(bars, df["BirincilikOranı"] * 100):
            ax.text(bar.get_width() + 0.5, bar.get_y() + bar.get_height() / 2,
                    f"{val:.1f}%", va="center", fontsize=7)
        ax.set_xlabel("Birincilik Oranı (%)" if lang != "EN" else "First-Place Rate (%)", fontsize=9)
        ax.set_xlim(0, 105)
        ax.tick_params(axis="both", labelsize=8)
        ax.spines["top"].set_visible(False)
        ax.spines["right"].set_visible(False)
        plt.tight_layout()
        buf = io.BytesIO()
        fig.savefig(buf, format="png", dpi=150, bbox_inches="tight")
        plt.close(fig)
        return buf.getvalue()
    except Exception:
        return None

def _sensitivity_heatmap_figure_bytes(local_df: pd.DataFrame, lang: str) -> bytes | None:
    """Heatmap of Spearman ρ by criterion × weight-change scenario; returns PNG bytes or None."""
    if local_df is None or local_df.empty:
        return None
    if not _ensure_matplotlib_support():
        return None
    try:
        required = {"Kriter", "AğırlıkDeğişimi", "SpearmanRho"}
        if not required.issubset(local_df.columns):
            return None
        pivot = local_df.pivot_table(index="Kriter", columns="AğırlıkDeğişimi",
                                     values="SpearmanRho", aggfunc="mean")
        if pivot.empty:
            return None
        fig, ax = plt.subplots(figsize=(max(4, pivot.shape[1] * 1.1),
                                        max(3, pivot.shape[0] * 0.45)))
        im = ax.imshow(pivot.values, cmap="RdYlGn", vmin=0.5, vmax=1.0, aspect="auto")
        ax.set_xticks(range(pivot.shape[1]))
        ax.set_xticklabels(pivot.columns.tolist(), fontsize=8)
        ax.set_yticks(range(pivot.shape[0]))
        ax.set_yticklabels(pivot.index.tolist(), fontsize=7)
        xlabel = "Ağırlık Değişimi" if lang != "EN" else "Weight Change"
        ylabel = "Kriter" if lang != "EN" else "Criterion"
        ax.set_xlabel(xlabel, fontsize=9)
        ax.set_ylabel(ylabel, fontsize=9)
        for r in range(pivot.shape[0]):
            for c in range(pivot.shape[1]):
                val = pivot.values[r, c]
                if not np.isnan(val):
                    ax.text(c, r, f"{val:.2f}", ha="center", va="center",
                            fontsize=6, color="black")
        plt.colorbar(im, ax=ax, fraction=0.03, pad=0.04)
        plt.tight_layout()
        buf = io.BytesIO()
        fig.savefig(buf, format="png", dpi=150, bbox_inches="tight")
        plt.close(fig)
        return buf.getvalue()
    except Exception:
        return None

def _doc_add_table_block(
    doc: "Document",
    label: str,
    df: pd.DataFrame,
    interpretation: str,
    lang_code: str,
) -> None:
    """SSCI-style table block: caption ABOVE → table → italic interpretation BELOW."""
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_docx_run_style(caption.add_run(label), lang_code, bold=True)
    add_table_to_doc(doc, df, lang_code=lang_code)
    if interpretation and interpretation.strip():
        interp = doc.add_paragraph()
        interp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _set_docx_run_style(interp.add_run(interpretation.strip()), lang_code, italic=True)
    doc.add_paragraph("")

def _doc_add_figure_block(
    doc: "Document",
    label: str,
    img_bytes: bytes,
    interpretation: str,
    lang_code: str,
) -> None:
    """SSCI-style figure block: image → caption BELOW → italic interpretation."""
    buf = io.BytesIO(img_bytes)
    doc.add_picture(buf, width=Inches(5.0))
    last_para = doc.paragraphs[-1]
    last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_docx_run_style(caption.add_run(label), lang_code, bold=True)
    if interpretation and interpretation.strip():
        interp = doc.add_paragraph()
        interp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _set_docx_run_style(interp.add_run(interpretation.strip()), lang_code, italic=True)
    doc.add_paragraph("")

def _build_academic_doc_sections(result: Dict[str, Any], selected_data: pd.DataFrame, lang: str) -> Dict[str, Any]:
    headings = _docx_heading_map(lang)
    refs = _normalize_references(
        (result.get("report_sections", {}) or {}).get("Kaynakça", []) or (result.get("report_sections", {}) or {}).get("References", []),
        lang,
    )
    data = selected_data if isinstance(selected_data, pd.DataFrame) else result.get("selected_data", pd.DataFrame())
    validation = result.get("validation", {}) or {}
    shape = validation.get("shape", (0, 0))
    n_alt = int(data.shape[0]) if isinstance(data, pd.DataFrame) and not data.empty else int(shape[0] or 0)
    n_crit = int(data.shape[1]) if isinstance(data, pd.DataFrame) and not data.empty else int(shape[1] or 0)
    direction = validation.get("direction_summary", {}) or {}
    benefit_n = len(direction.get("benefit", []) or [])
    cost_n = len(direction.get("cost", []) or [])

    weight_method = result.get("weights", {}).get("method")
    ranking_method = result.get("ranking", {}).get("method")
    weight_name = _docx_method_name(weight_method, lang)
    ranking_name = _docx_method_name(ranking_method, lang) if ranking_method else None
    w_ph = (result.get("method_philosophy", {}) or {}).get("weight", {}) or {}
    r_ph = (result.get("method_philosophy", {}) or {}).get("ranking", {}) or {}
    w_simple_fb, w_academic_fb = _method_fallback_lines(weight_name or ("Selected method" if lang == "EN" else "Seçilen yöntem"), lang)
    r_simple_fb, r_academic_fb = _method_fallback_lines(ranking_name or ("Selected method" if lang == "EN" else "Seçilen yöntem"), lang)
    w_simple = _lang_text(w_ph.get("simple", ""), lang, w_simple_fb)
    w_academic = _lang_text(w_ph.get("academic", ""), lang, w_academic_fb)
    r_simple = _lang_text(r_ph.get("simple", ""), lang, r_simple_fb)
    r_academic = _lang_text(r_ph.get("academic", ""), lang, r_academic_fb)

    weight_df = result.get("weights", {}).get("table")
    ranking_df = result.get("ranking", {}).get("table")
    comp = result.get("comparison", {}) or {}
    sens = result.get("sensitivity") or {}
    comp_methods = _docx_comparison_methods(comp)
    mean_rho = _docx_mean_spearman(comp)
    stability = sens.get("top_stability")
    n_iter = int(sens.get("n_iterations", 0)) if sens else 0

    top_criterion = ""
    top_weight = None
    top3_share = None
    if isinstance(weight_df, pd.DataFrame) and not weight_df.empty:
        c_col = col_key(weight_df, "Kriter", "Criterion")
        w_col = col_key(weight_df, "Ağırlık", "Weight")
        w_sorted = weight_df.sort_values(w_col, ascending=False).reset_index(drop=True)
        top_criterion = str(w_sorted.iloc[0][c_col])
        top_weight = float(pd.to_numeric(w_sorted.iloc[0][w_col], errors="coerce"))
        top3_share = float(pd.to_numeric(w_sorted[w_col], errors="coerce").head(3).sum())

    top_alt = ""
    second_alt = ""
    top_score = None
    score_gap = None
    if isinstance(ranking_df, pd.DataFrame) and not ranking_df.empty:
        alt_col = col_key(ranking_df, "Alternatif", "Alternative")
        score_col = col_key(ranking_df, "Skor", "Score")
        r_sorted = ranking_df.sort_values(col_key(ranking_df, "Sıra", "Rank")).reset_index(drop=True)
        top_alt = str(r_sorted.iloc[0][alt_col])
        top_score = float(pd.to_numeric(r_sorted.iloc[0][score_col], errors="coerce"))
        if len(r_sorted) > 1:
            second_alt = str(r_sorted.iloc[1][alt_col])
            second_score = float(pd.to_numeric(r_sorted.iloc[1][score_col], errors="coerce"))
            _lower_is_better = isinstance(ranking_name, str) and ("VIKOR" in ranking_name)
            score_gap = (second_score - top_score) if _lower_is_better else (top_score - second_score)

    if lang == "EN":
        objective = (
            f"The purpose of this report is to examine a decision problem composed of {n_alt} alternatives and {n_crit} criteria "
            f"by deriving criterion weights through {weight_name}"
            + (f" and constructing the final ranking through {ranking_name}" if ranking_name else "")
            + ". The analytical focus is to identify the criteria that structure the problem and"
            + (f" to present the leading alternative produced by {ranking_name} within an academically defensible frame." if ranking_name else " to establish the importance structure of the criteria within an academically defensible frame.")
        )

        scope_parts = [f"The scope of the study is limited to a decision matrix containing {n_alt} alternatives and {n_crit} criteria."]
        if benefit_n or cost_n:
            scope_parts.append(f"The criterion set consists of {benefit_n} benefit-oriented and {cost_n} cost-oriented variables.")
        scope_parts.append(
            f"The analytical design is based on the user-selected {weight_name} weighting approach"
            + (f" together with the {ranking_name} ranking procedure." if ranking_name else ".")
        )
        if len(comp_methods) >= 2:
            scope_parts.append(
                f"In addition, the agreement among the user-selected methods ({', '.join(comp_methods)}) is reported through a Spearman rank-correlation table."
            )
        if n_iter > 0 and stability is not None:
            scope_parts.append(f"Result robustness against weight perturbation is further examined through {n_iter:,} Monte Carlo scenarios.")
        scope = " ".join(scope_parts)

        philosophy_parts = []
        philosophy_parts.append(w_academic or f"{weight_name} serves as the weighting backbone of the analysis.")
        if ranking_name:
            philosophy_parts.append(r_academic or f"{ranking_name} provides the final ranking logic of the analysis.")
            philosophy_parts.append("Taken together, the selected weighting and ranking stages separate criterion importance from alternative dominance and thereby improve interpretability at the reporting stage.")
        philosophy_parts.append("Accordingly, the findings should be interpreted within the epistemic assumptions of the selected method set rather than as method-independent truths.")
        philosophy = "\n\n".join(philosophy_parts)

        _ti_dm = (
            f"The decision matrix presents the raw performance values for {n_alt} alternatives across {n_crit} criteria. "
            + (f"Of these, {benefit_n} are benefit criteria (higher is better) and {cost_n} are cost criteria (lower is better). " if benefit_n or cost_n else "")
            + "These values form the input to all subsequent weighting and ranking computations."
        )
        _ti_weight = ""
        if top_criterion and top_weight is not None and top3_share is not None:
            _ti_weight = (
                f"The weight table shows that {top_criterion} is the dominant criterion with a weight of {top_weight:.3f}. "
                f"The cumulative weight of the first three criteria reaches {top3_share:.1%}, suggesting that the decision structure concentrates around a limited number of decisive dimensions."
            )
        _ti_ranking = ""
        if ranking_name and top_alt and top_score is not None:
            _ti_ranking = f"According to {ranking_name}, {top_alt} ranks first with a score of {top_score:.3f}."
            if second_alt and score_gap is not None:
                _ti_ranking += f" The score gap of {score_gap:.3f} over {second_alt} indicates the degree to which the leading alternative differentiates itself from the immediate follower."
        _ti_detail = _docx_detail_interpretation(ranking_method, lang)
        if _ti_detail and _ti_ranking:
            _ti_ranking = _ti_ranking + " " + _ti_detail
        elif _ti_detail:
            _ti_ranking = _ti_detail
        _ti_comparison = ""
        if mean_rho is not None and len(comp_methods) >= 2:
            consistency = "high" if mean_rho >= 0.85 else "moderate" if mean_rho >= 0.70 else "limited"
            _ti_comparison = f"The mean Spearman agreement across the user-selected methods is {mean_rho:.3f}, which points to {consistency} structural consistency among the reported rankings."
        _ti_mc = ""
        if stability is not None:
            robustness = "strong" if float(stability) >= 0.75 else "moderate" if float(stability) >= 0.50 else "limited"
            _ti_mc = f"The Monte Carlo analysis yields a first-place retention rate of {float(stability):.1%} for the leading alternative, indicating {robustness} robustness against weight perturbation."
        table_parts = [t for t in [_ti_weight, _ti_ranking, _ti_comparison, _ti_mc] if t]
        tables_text = "\n\n".join(table_parts)

        conclusion_parts = []
        if ranking_name and top_alt:
            conclusion_parts.append(
                f"In conclusion, the weighting structure derived through {weight_name} and the ranking logic established by {ranking_name} jointly identify {top_alt} as the leading alternative."
            )
        else:
            conclusion_parts.append(
                f"In conclusion, the {weight_name} weighting structure identifies the main criteria that shape the decision problem."
            )
        if top_criterion and top_weight is not None:
            conclusion_parts.append(
                f"The prominence of {top_criterion} confirms that the decision is primarily driven by criterion-specific differentiation rather than by uniform criterion effects."
            )
        if stability is not None:
            conclusion_parts.append(
                "The robustness evidence supports the transfer of the result into practice, provided that the final decision is still read together with contextual expert judgment and data recency."
            )
        else:
            conclusion_parts.append(
                "The result should be transferred into practice together with contextual expert judgment and problem-specific constraints."
            )
        conclusion = " ".join(conclusion_parts)
    else:
        objective = (
            f"Bu raporun amacı, {n_alt} alternatif ve {n_crit} kriterden oluşan karar problemini "
            f"{weight_name} ile kriter ağırlıklarını türeterek"
            + (f" ve {ranking_name} ile nihai sıralamayı kurarak" if ranking_name else "")
            + " sistematik biçimde incelemektir. Analitik odak, karar yapısını sürükleyen kriterleri görünür kılmak ve"
            + (f" {ranking_name} altında öne çıkan alternatifi akademik olarak savunulabilir bir çerçevede ortaya koymaktır." if ranking_name else " kriter önem yapısını akademik olarak savunulabilir bir çerçevede ortaya koymaktır.")
        )

        scope_parts = [f"Çalışmanın kapsamı, {n_alt} alternatif ve {n_crit} kriter içeren karar matrisi ile sınırlıdır."]
        if benefit_n or cost_n:
            scope_parts.append(f"Kriter seti {benefit_n} fayda yönlü ve {cost_n} maliyet yönlü değişkenden oluşmaktadır.")
        scope_parts.append(
            f"Analitik kurgu kullanıcı tarafından seçilen {weight_name} ağırlıklandırmasına"
            + (f" ve {ranking_name} sıralama yaklaşımına dayanmaktadır." if ranking_name else " dayanmaktadır.")
        )
        if len(comp_methods) >= 2:
            scope_parts.append(
                f"Ek olarak kullanıcı tarafından seçilen {', '.join(comp_methods)} yöntemleri arasındaki uyum, Spearman sıra korelasyonu tablosu üzerinden raporlanmıştır."
            )
        if n_iter > 0 and stability is not None:
            scope_parts.append(f"Sonuçların ağırlık bozulmalarına karşı dayanıklılığı ayrıca {n_iter:,} Monte Carlo senaryosu ile incelenmiştir.")
        scope = " ".join(scope_parts)

        philosophy_parts = []
        philosophy_parts.append(w_academic or f"{weight_name} bu raporda ağırlıklandırma omurgasını oluşturmaktadır.")
        if ranking_name:
            philosophy_parts.append(r_academic or f"{ranking_name} bu raporda nihai sıralama mantığını üretmektedir.")
            philosophy_parts.append("Seçilen ağırlıklandırma ve sıralama aşamalarının birlikte kullanılması, kriter önemini alternatif üstünlüğünden ayırarak yorum izlenebilirliğini artırmaktadır.")
        philosophy_parts.append("Bu nedenle bulgular, yöntemden bağımsız mutlak doğrular olarak değil, seçilen yöntem setinin epistemik varsayımları içinde okunmalıdır.")
        philosophy = "\n\n".join(philosophy_parts)

        _ti_dm = (
            f"Karar matrisi, {n_alt} alternatifin {n_crit} kriter üzerindeki ham performans değerlerini içermektedir. "
            + (f"Bu kriterlerden {benefit_n} tanesi fayda yönlü (yüksek değer daha iyi), {cost_n} tanesi maliyet yönlüdür (düşük değer daha iyi). " if benefit_n or cost_n else "")
            + "Bu değerler, sonraki tüm ağırlıklandırma ve sıralama hesaplamalarına girdi oluşturmaktadır."
        )
        _ti_weight = ""
        if top_criterion and top_weight is not None and top3_share is not None:
            _ti_weight = (
                f"Ağırlık tablosu, {top_criterion} kriterinin {_fmt(top_weight, 'TR')} ağırlık ile en baskın boyut olduğunu göstermektedir. "
                f"İlk üç kriterin toplam ağırlığı %{top3_share * 100:.1f} düzeyine ulaşmakta; bu görünüm karar yapısının sınırlı sayıda belirleyici eksen etrafında yoğunlaştığını düşündürmektedir."
            )
        _ti_ranking = ""
        if ranking_name and top_alt and top_score is not None:
            _ti_ranking = f"{ranking_name} sonuçlarına göre {top_alt}, {_fmt(top_score, 'TR')} skoruyla ilk sırada yer almıştır."
            if second_alt and score_gap is not None:
                _ti_ranking += f" İkinci sıradaki {second_alt} ile oluşan {_fmt(score_gap, 'TR')} puanlık fark, lider alternatifin en yakın rakibine göre ne ölçüde ayrıştığını göstermektedir."
        _ti_detail = _docx_detail_interpretation(ranking_method, lang)
        if _ti_detail and _ti_ranking:
            _ti_ranking = _ti_ranking + " " + _ti_detail
        elif _ti_detail:
            _ti_ranking = _ti_detail
        _ti_comparison = ""
        if mean_rho is not None and len(comp_methods) >= 2:
            consistency = "yüksek" if mean_rho >= 0.85 else "orta" if mean_rho >= 0.70 else "sınırlı"
            _ti_comparison = f"Kullanıcı tarafından seçilen yöntemler arasındaki ortalama Spearman uyumu {_fmt(mean_rho, 'TR')} düzeyindedir; bu değer yöntemler arası {consistency} yapısal tutarlılığa işaret etmektedir."
        _ti_mc = ""
        if stability is not None:
            robustness = "güçlü" if float(stability) >= 0.75 else "orta" if float(stability) >= 0.50 else "sınırlı"
            _ti_mc = f"Monte Carlo incelemesinde lider alternatifin birinci sırayı koruma oranı %{float(stability) * 100:.1f} olarak hesaplanmıştır. Bu bulgu, sonucun ağırlık bozulmalarına karşı {robustness} bir dayanıklılık sergilediğini göstermektedir."
        table_parts = [t for t in [_ti_weight, _ti_ranking, _ti_comparison, _ti_mc] if t]
        tables_text = "\n\n".join(table_parts)

        conclusion_parts = []
        if ranking_name and top_alt:
            conclusion_parts.append(
                f"Sonuç olarak, {weight_name} ile üretilen ağırlık yapısı ve {ranking_name} ile kurulan sıralama mantığı birlikte değerlendirildiğinde {top_alt} alternatifi en güçlü seçenek olarak öne çıkmaktadır."
            )
        else:
            conclusion_parts.append(
                f"Sonuç olarak, {weight_name} ile elde edilen ağırlık yapısı karar problemini taşıyan başlıca kriterleri açık biçimde ortaya koymaktadır."
            )
        if top_criterion and top_weight is not None:
            conclusion_parts.append(
                f"{top_criterion} kriterinin baskınlığı, kararın homojen bir yapıdan çok kriterler arası ayırt edicilik üzerinden şekillendiğini doğrulamaktadır."
            )
        if stability is not None:
            conclusion_parts.append(
                "Sağlamlık bulguları, sonucun uygulamaya aktarılabilirliğini desteklemektedir; bununla birlikte nihai karar yine bağlamsal uzman görüşü ve veri güncelliği ile birlikte değerlendirilmelidir."
            )
        else:
            conclusion_parts.append(
                "Bulguların uygulamaya aktarımı, bağlamsal uzman görüşü ve problem alanına özgü sınırlılıklar ile birlikte ele alınmalıdır."
            )
        conclusion = " ".join(conclusion_parts)

    return {
        headings["objective"]: objective,
        headings["scope"]: scope,
        headings["philosophy"]: philosophy,
        headings["tables"]: tables_text,
        headings["conclusion"]: conclusion,
        headings["references"]: refs,
        "table_interp": {
            "decision_matrix": _ti_dm,
            "weights": _ti_weight,
            "ranking": _ti_ranking,
            "comparison": _ti_comparison,
            "monte_carlo": _ti_mc,
        },
        "figure_interp": {
            "weights": _ti_weight,
            "ranking": _ti_ranking,
        },
    }

def _preferred_doc_detail_table(result: Dict[str, Any], lang: str) -> tuple[str, pd.DataFrame] | None:
    details = result.get("ranking", {}).get("details", {}) or {}
    preferred = [
        ("promethee_flows", ("PROMETHEE Flow Table", "PROMETHEE Akış Tablosu")),
        ("distance_table", ("TOPSIS Distance Table", "TOPSIS Uzaklık Tablosu")),
        ("vikor_table", ("VIKOR Compromise Table", "VIKOR Uzlaşı Tablosu")),
        ("edas_table", ("EDAS Distance Table", "EDAS Sapma Tablosu")),
        ("codas_table", ("CODAS Distance Table", "CODAS Uzaklık Tablosu")),
        ("gra_table", ("GRA Detail Table", "GRA Detay Tablosu")),
        ("marcos_table", ("MARCOS Detail Table", "MARCOS Detay Tablosu")),
        ("cocoso_table", ("CoCoSo Detail Table", "CoCoSo Detay Tablosu")),
        ("mabac_table", ("MABAC Detail Table", "MABAC Detay Tablosu")),
    ]
    for key, labels in preferred:
        df = details.get(key)
        if isinstance(df, pd.DataFrame) and not df.empty:
            return (labels[0] if lang == "EN" else labels[1], df)
    return None

def generate_apa_docx(result: Dict[str, Any], selected_data: pd.DataFrame, lang: str = "TR") -> bytes | None:
    if not _ensure_docx_support():
        return None
    doc = Document()
    _configure_apa_doc(doc)
    for section in doc.sections:
        section.top_margin, section.bottom_margin = Inches(1), Inches(1)
        section.left_margin, section.right_margin = Inches(1), Inches(1)

    is_tr = lang != "EN"
    body_lang = "tr-TR" if is_tr else "en-US"

    def _add_heading(text: str) -> None:
        p = doc.add_paragraph()
        _set_docx_run_style(p.add_run(text), body_lang, bold=True)

    def _add_body(text: str) -> None:
        for block in [b.strip() for b in str(text or "").split("\n\n") if b.strip()]:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            _set_docx_run_style(p.add_run(block), body_lang)

    # ── Title ────────────────────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_docx_run_style(
        title_p.add_run(
            "Çok Kriterli Karar Verme Akademik Analiz Raporu"
            if is_tr else
            "Multi-Criteria Decision-Making Academic Analysis Report"
        ),
        body_lang, bold=True,
    )
    doc.add_paragraph("")

    # ── Build narrative sections ──────────────────────────────────────────────
    sections = _build_academic_doc_sections(result, selected_data, lang)
    heading_map = _docx_heading_map(lang)
    ti = sections.get("table_interp", {})
    fi = sections.get("figure_interp", {})

    # ── 1. Objective ─────────────────────────────────────────────────────────
    _add_heading(heading_map["objective"])
    _add_body(sections[heading_map["objective"]])

    # ── 2. Scope ─────────────────────────────────────────────────────────────
    _add_heading(heading_map["scope"])
    _add_body(sections[heading_map["scope"]])

    # ── 3. Philosophy ─────────────────────────────────────────────────────────
    _add_heading(heading_map["philosophy"])
    _add_body(sections[heading_map["philosophy"]])

    # ── 4. Findings: Tables & Figures ─────────────────────────────────────────
    findings_heading = "Bulgular" if is_tr else "Findings"
    _add_heading(findings_heading)

    tbl_num = [0]
    fig_num = [0]

    def _tbl_label(title: str) -> str:
        tbl_num[0] += 1
        return (f"Tablo {tbl_num[0]}: {title}" if is_tr else f"Table {tbl_num[0]}: {title}")

    def _fig_label(title: str) -> str:
        fig_num[0] += 1
        return (f"Şekil {fig_num[0]}: {title}" if is_tr else f"Figure {fig_num[0]}: {title}")

    # Table: Decision Matrix
    dm_df = _docx_preview_df(selected_data, lang)
    if not dm_df.empty:
        _doc_add_table_block(
            doc,
            _tbl_label("Karar Matrisi Özeti" if is_tr else "Decision Matrix Snapshot"),
            dm_df,
            ti.get("decision_matrix", ""),
            body_lang,
        )

    # Table + Figure: Weights
    weight_df_raw = (result.get("weights") or {}).get("table")
    if isinstance(weight_df_raw, pd.DataFrame) and not weight_df_raw.empty:
        weight_df = localize_df_lang(weight_df_raw, lang)
        _doc_add_table_block(
            doc,
            _tbl_label("Kriter Ağırlık Tablosu" if is_tr else "Criterion Weight Table"),
            weight_df,
            ti.get("weights", ""),
            body_lang,
        )
        w_img = _weight_bar_figure_bytes(weight_df_raw, lang)
        if w_img:
            _doc_add_figure_block(
                doc,
                _fig_label("Kriter Ağırlıkları" if is_tr else "Criterion Weights"),
                w_img,
                fi.get("weights", ""),
                body_lang,
            )
        w_radar = _weight_radar_figure_bytes(weight_df_raw, lang)
        if w_radar:
            _doc_add_figure_block(
                doc,
                _fig_label("Kriter Ağırlık Dağılımı (Radar)" if is_tr else "Criterion Weight Distribution (Radar)"),
                w_radar,
                fi.get("weights", ""),
                body_lang,
            )

    # Table + Figure: Ranking
    ranking_df_raw = (result.get("ranking") or {}).get("table")
    if isinstance(ranking_df_raw, pd.DataFrame) and not ranking_df_raw.empty:
        ranking_df = localize_df_lang(ranking_df_raw, lang)
        _doc_add_table_block(
            doc,
            _tbl_label("Sıralama Tablosu" if is_tr else "Ranking Table"),
            ranking_df,
            ti.get("ranking", ""),
            body_lang,
        )
        r_img = _ranking_bar_figure_bytes(ranking_df_raw, lang)
        if r_img:
            _doc_add_figure_block(
                doc,
                _fig_label("Alternatif Skorları" if is_tr else "Alternative Scores"),
                r_img,
                fi.get("ranking", ""),
                body_lang,
            )

    # Table: Method detail (TOPSIS, VIKOR, PROMETHEE, etc.)
    detail_info = _preferred_doc_detail_table(result, lang)
    if detail_info is not None:
        detail_title, detail_df = detail_info
        detail_interp = _docx_detail_interpretation((result.get("ranking") or {}).get("method"), lang)
        _doc_add_table_block(
            doc,
            _tbl_label(detail_title),
            localize_df_lang(detail_df, lang),
            detail_interp,
            body_lang,
        )

    # Table: Method comparison (Spearman)
    comp_df = (result.get("comparison") or {}).get("spearman_matrix")
    if isinstance(comp_df, pd.DataFrame) and not comp_df.empty and comp_df.shape[0] >= 2:
        _doc_add_table_block(
            doc,
            _tbl_label("Yöntem Karşılaştırma Tablosu (Spearman)" if is_tr else "Method Comparison Table (Spearman)"),
            localize_df_lang(comp_df, lang),
            ti.get("comparison", ""),
            body_lang,
        )

    # Table + Figures: Monte Carlo & Sensitivity
    sensitivity_data = result.get("sensitivity") or {}
    mc_df = sensitivity_data.get("monte_carlo_summary")
    if isinstance(mc_df, pd.DataFrame) and not mc_df.empty:
        _doc_add_table_block(
            doc,
            _tbl_label("Monte Carlo Duyarlılık Özeti" if is_tr else "Monte Carlo Sensitivity Summary"),
            localize_df_lang(mc_df, lang),
            ti.get("monte_carlo", ""),
            body_lang,
        )
        mc_img = _mc_stability_figure_bytes(mc_df, lang)
        if mc_img:
            _doc_add_figure_block(
                doc,
                _fig_label("Monte Carlo Stabilite Analizi" if is_tr else "Monte Carlo Stability Analysis"),
                mc_img,
                (
                    "Her alternatifin Monte Carlo simülasyonunda birincilik oranını göstermektedir."
                    if is_tr else
                    "First-place rate of each alternative across Monte Carlo simulations."
                ),
                body_lang,
            )
    local_df = sensitivity_data.get("local_sensitivity")
    if isinstance(local_df, pd.DataFrame) and not local_df.empty:
        sens_img = _sensitivity_heatmap_figure_bytes(local_df, lang)
        if sens_img:
            _doc_add_figure_block(
                doc,
                _fig_label("Yerel Duyarlılık Analizi (Spearman ρ)" if is_tr else "Local Sensitivity Analysis (Spearman ρ)"),
                sens_img,
                (
                    "Kriter ağırlıklarındaki değişimlerin sıralama tutarlılığına etkisi (Spearman korelasyonu)."
                    if is_tr else
                    "Impact of criterion weight perturbations on ranking consistency (Spearman correlation)."
                ),
                body_lang,
            )

    # ── 5. Conclusion ─────────────────────────────────────────────────────────
    _add_heading(heading_map["conclusion"])
    _add_body(sections[heading_map["conclusion"]])

    # ── 6. References ─────────────────────────────────────────────────────────
    for note_line in _reference_notice_lines(lang):
        p = doc.add_paragraph()
        _set_docx_run_style(p.add_run(note_line), body_lang, italic=True)
    _add_heading(heading_map["references"])
    for ref in sections.get(heading_map["references"], []):
        ref_p = doc.add_paragraph()
        ref_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        ref_p.paragraph_format.left_indent = Inches(0.5)
        ref_p.paragraph_format.first_line_indent = Inches(-0.5)
        _set_docx_run_style(ref_p.add_run(ref), body_lang)

    # ── Signature ─────────────────────────────────────────────────────────────
    doc.add_paragraph("\n")
    sig = doc.add_paragraph()
    sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_docx_run_style(sig.add_run("Prof. Dr. Ömer Faruk Rençber"), body_lang, italic=True)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()

def _export_study_title(result: Dict[str, Any], lang: str) -> str:
    manual_title = str(st.session_state.get("study_title", "") or "").strip()
    if manual_title:
        return manual_title
    weight_method = (result.get("weights") or {}).get("method") or _xl_text(lang, "Ağırlık", "Weights")
    ranking_method = (result.get("ranking") or {}).get("method")
    if ranking_method:
        return _xl_text(lang, f"{weight_method} - {ranking_method} Analiz Raporu", f"{weight_method} - {ranking_method} Analysis Report")
    return _xl_text(lang, f"{weight_method} Ağırlık Analizi Raporu", f"{weight_method} Weight Analysis Report")

def _export_file_name(title: str, lang: str, ext: str) -> str:
    base = re.sub(r"[^\w.-]+", "_", str(title or "").strip(), flags=re.UNICODE).strip("._")
    if not base:
        base = _xl_text(lang, "MCDM_Sonuclari", "MCDM_Results")
    return f"{base}.{ext}"

def _render_export_download_button(
    label: str,
    data: bytes | bytearray,
    file_name: str,
    mime: str,
    key: str,
) -> bool:
    return st.download_button(
        label,
        data=bytes(data),
        file_name=file_name,
        mime=mime,
        use_container_width=True,
        on_click="ignore",
        key=key,
    )

def _safe_sheet_name(name: str, used_names: set[str]) -> str:
    clean = re.sub(r"[\[\]:*?/\\]", "_", str(name or "Sheet")).strip() or "Sheet"
    clean = clean[:31]
    candidate = clean
    counter = 2
    while candidate in used_names:
        suffix = f"_{counter}"
        candidate = f"{clean[: 31 - len(suffix)]}{suffix}"
        counter += 1
    used_names.add(candidate)
    return candidate

def _excel_formats(workbook):
    return {
        "cover_title": workbook.add_format({"bold": True, "font_size": 18, "align": "center", "valign": "vcenter", "font_color": "#FFFFFF", "bg_color": "#16324F"}),
        "cover_subtitle": workbook.add_format({"italic": True, "font_size": 11, "align": "center", "valign": "vcenter", "font_color": "#EAF1F7", "bg_color": "#1F4A73"}),
        "cover_badge": workbook.add_format({"bold": True, "font_size": 13, "align": "center", "valign": "vcenter", "font_color": "#FFFFFF", "bg_color": "#C9A227"}),
        "cover_label_card": workbook.add_format({"bold": True, "font_size": 11, "align": "center", "valign": "vcenter", "font_color": "#FFFFFF", "bg_color": "#2E75B6", "border": 1}),
        "cover_value_card": workbook.add_format({"font_size": 11, "align": "left", "valign": "vcenter", "bg_color": "#F2F2F2", "border": 1, "text_wrap": True}),
        "sheet_title": workbook.add_format({"bold": True, "font_size": 14, "align": "center", "valign": "vcenter", "font_color": "#FFFFFF", "bg_color": "#1F3864"}),
        "sheet_subtitle": workbook.add_format({"italic": True, "font_size": 10, "align": "center", "valign": "vcenter", "font_color": "#1F1F1F", "bg_color": "#E9EFF7"}),
        "section": workbook.add_format({"bold": True, "font_size": 12, "align": "left", "valign": "vcenter", "font_color": "#16324F", "bg_color": "#DCE9F5", "border": 1}),
        "header": workbook.add_format({"bold": True, "bg_color": "#1d3557", "font_color": "#FFFFFF", "border": 1, "text_wrap": True, "valign": "vcenter"}),
        "cell": workbook.add_format({"border": 1, "valign": "top"}),
        "cell_wrap": workbook.add_format({"border": 1, "text_wrap": True, "valign": "top"}),
        "label": workbook.add_format({"bold": True, "border": 1, "bg_color": "#EFF5FB"}),
        "value": workbook.add_format({"border": 1, "text_wrap": True}),
        "note": workbook.add_format({"italic": True, "font_color": "#516173"}),
        "text_block": workbook.add_format({"text_wrap": True, "valign": "top", "border": 1}),
        "citation": workbook.add_format({"italic": True, "font_size": 9, "font_color": "#516173", "bg_color": "#F7F9FC", "border": 1, "text_wrap": True, "valign": "vcenter"}),
        "highlight_good": workbook.add_format({"bold": True, "font_color": "#1F5130", "bg_color": "#D9EAD3", "border": 1}),
        "highlight_warn": workbook.add_format({"bold": True, "font_color": "#7F1D1D", "bg_color": "#F4CCCC", "border": 1}),
    }

def _excel_method_chain(weight_method: str | None, ranking_method: str | None, lang: str) -> str:
    w_label = str(weight_method or "").strip() or _xl_text(lang, "Belirtilmedi", "Not specified")
    r_label = str(ranking_method or "").strip()
    if not r_label:
        return f"{w_label} {_xl_text(lang, 'Analizi', 'Analysis')}"
    return f"{w_label} -> {r_label}"

def _write_sheet_banner(ws, title: str, subtitle: str, formats, *, end_col: int = 9) -> int:
    end_col = max(5, int(end_col))
    ws.merge_range(0, 0, 0, end_col, title, formats["sheet_title"])
    ws.merge_range(1, 0, 1, end_col, subtitle, formats["sheet_subtitle"])
    ws.set_row(0, 24)
    ws.set_row(1, 20)
    return 3

def _write_citation_block(ws, start_row: int, text: str, formats, *, end_col: int = 9) -> int:
    end_col = max(3, int(end_col))
    ws.merge_range(start_row, 0, start_row, end_col, text, formats["citation"])
    ws.set_row(start_row, 34)
    return start_row + 2

def _set_worksheet_widths(ws, df: pd.DataFrame, startcol: int = 0, index: bool = False, max_width: int = 38) -> None:
    if index:
        idx_sample = df.index.astype(str).tolist()[:100]
        idx_width = min(max_width, max(12, max([len(str(df.index.name or ""))] + [len(val.replace("\n", " ")) for val in idx_sample]) + 2))
        ws.set_column(startcol, startcol, idx_width)
        startcol += 1
    for offset, col in enumerate(df.columns):
        sample = df[col].head(100).astype(str).tolist()
        max_len = max([len(str(col))] + [len(val.replace("\n", " ")) for val in sample]) if sample else len(str(col))
        ws.set_column(startcol + offset, startcol + offset, min(max_width, max(12, max_len + 2)))

def _write_df_block(
    writer,
    sheet_name: str,
    df: pd.DataFrame,
    startrow: int,
    title: str,
    formats,
    *,
    index: bool = False,
    freeze: bool = False,
) -> Dict[str, int]:
    ws = writer.sheets[sheet_name]
    cols_n = max(1, df.shape[1] + (1 if index else 0))
    ws.merge_range(startrow, 0, startrow, max(4, cols_n - 1), title, formats["section"])
    data_row = startrow + 2
    if df.empty:
        ws.write(data_row, 0, "—", formats["note"])
        return {"header_row": data_row, "data_row": data_row + 1, "nrows": 0, "ncols": cols_n, "next_row": data_row + 3}
    df.to_excel(writer, sheet_name=sheet_name, startrow=data_row, startcol=0, index=index)
    ws.set_row(data_row, 24, formats["header"])
    ws.autofilter(data_row, 0, data_row + len(df), cols_n - 1)
    _set_worksheet_widths(ws, df, startcol=0, index=index)
    if freeze:
        ws.freeze_panes(data_row + 1, 0)
    return {
        "header_row": data_row,
        "data_row": data_row + 1,
        "nrows": len(df),
        "ncols": cols_n,
        "next_row": data_row + len(df) + 3,
    }

def _clean_excel_text(text: Any) -> str:
    raw = html.unescape(str(text or ""))
    raw = re.sub(r"<br\s*/?>", "\n", raw, flags=re.IGNORECASE)
    raw = re.sub(r"</(p|div|li)>", "\n", raw, flags=re.IGNORECASE)
    raw = re.sub(r"<li[^>]*>", "- ", raw, flags=re.IGNORECASE)
    raw = re.sub(r"<[^>]+>", "", raw)
    raw = raw.replace("**", "").replace("__", "").replace("*", "")
    raw = re.sub(r"[ \t]+", " ", raw)
    raw = re.sub(r"\n{3,}", "\n\n", raw)
    return raw.strip()

def _write_layer_insight_block(
    ws,
    start_row: int,
    title: str,
    layers: Dict[str, str],
    lang: str,
    formats,
    *,
    end_col: int = 11,
) -> int:
    valid_rows = [
        (_xl_text(lang, "Betimsel çıkarım", "Descriptive inference"), _clean_excel_text((layers or {}).get("descriptive", ""))),
        (_xl_text(lang, "Analitik yorum", "Analytic interpretation"), _clean_excel_text((layers or {}).get("analytic", ""))),
        (_xl_text(lang, "Karar önerisi", "Decision implication"), _clean_excel_text((layers or {}).get("normative", ""))),
    ]
    valid_rows = [(label, text) for label, text in valid_rows if text]
    if not valid_rows:
        return start_row

    ws.merge_range(start_row, 0, start_row, end_col, title, formats["section"])
    row = start_row + 2
    for label, text in valid_rows:
        ws.write(row, 0, label, formats["label"])
        ws.merge_range(row, 1, row + 2, end_col, text, formats["text_block"])
        ws.set_row(row, 24)
        ws.set_row(row + 1, 24)
        ws.set_row(row + 2, 24)
        row += 4
    return row

def _write_bullet_block(
    ws,
    start_row: int,
    title: str,
    lines: List[str],
    formats,
    *,
    end_col: int = 11,
) -> int:
    clean_lines = [_clean_excel_text(line) for line in lines if _clean_excel_text(line)]
    if not clean_lines:
        return start_row
    ws.merge_range(start_row, 0, start_row, end_col, title, formats["section"])
    body = "\n".join(f"{idx}. {line}" for idx, line in enumerate(clean_lines, start=1))
    span = max(3, min(12, len(clean_lines) + 2))
    ws.merge_range(start_row + 2, 0, start_row + 1 + span, end_col, body, formats["text_block"])
    for row_idx in range(start_row + 2, start_row + 2 + span):
        ws.set_row(row_idx, 24)
    return start_row + span + 3

def _weight_layers_for_excel(result: Dict[str, Any]) -> Dict[str, str]:
    weights = (result.get("weights") or {}).get("values") or {}
    if not weights:
        return {}
    method = str((result.get("weights") or {}).get("method") or "").strip()
    selected_data = result.get("selected_data", pd.DataFrame())
    n_alt = len(selected_data) if isinstance(selected_data, pd.DataFrame) else 0
    return me.generate_3layer_weight(method, weights, max(n_alt, 1))

def _ranking_layers_for_excel(method_name: str, ranking_table: pd.DataFrame, weights: Dict[str, float]) -> Dict[str, str]:
    if not str(method_name or "").strip() or not isinstance(ranking_table, pd.DataFrame) or ranking_table.empty:
        return {}
    internal_table = ranking_table.copy()
    internal_table = internal_table.rename(columns={"Alternative": "Alternatif", "Score": "Skor", "Rank": "Sıra"})
    return me.generate_3layer_ranking(str(method_name), internal_table, weights or {})

def _comparison_layers_for_excel(result: Dict[str, Any]) -> Dict[str, str]:
    comparison = result.get("comparison") or {}
    base_method = str((result.get("ranking") or {}).get("method") or "").strip()
    if not comparison or not base_method:
        return {}
    return me.generate_3layer_comparison(comparison, base_method)

def _sensitivity_layers_for_excel(result: Dict[str, Any]) -> Dict[str, str]:
    sensitivity = result.get("sensitivity") or {}
    if not sensitivity:
        return {}
    return me.generate_3layer_sensitivity(sensitivity)

def _vikor_compromise_export_df(method_name: str, ranking_details: Dict[str, Any], lang: str) -> pd.DataFrame:
    method_label = str(method_name or "").strip()
    if "VIKOR" not in method_label.upper():
        return pd.DataFrame()
    vikor_table = ranking_details.get("vikor_table")
    conditions = ranking_details.get("compromise_conditions") or {}
    if not isinstance(vikor_table, pd.DataFrame) or vikor_table.empty:
        return pd.DataFrame()

    table = vikor_table.copy()
    alt_col = col_key(table, "Alternatif", "Alternative")
    for metric_col in ["S", "R", "Q"]:
        if metric_col in table.columns:
            table[metric_col] = pd.to_numeric(table[metric_col], errors="coerce")
    table = table.sort_values("Q", ascending=True).reset_index(drop=True)
    top_alt = str(table.iloc[0][alt_col])
    second_alt = str(table.iloc[1][alt_col]) if len(table) > 1 else top_alt
    q1 = float(pd.to_numeric(table.iloc[0]["Q"], errors="coerce")) if "Q" in table.columns else np.nan
    q2 = float(pd.to_numeric(table.iloc[1]["Q"], errors="coerce")) if len(table) > 1 and "Q" in table.columns else q1
    dq = 1.0 / max(len(table) - 1, 1) if len(table) > 1 else np.nan
    gap = q2 - q1 if np.isfinite(q1) and np.isfinite(q2) else np.nan
    s_rank = int(pd.to_numeric(table["S"], errors="coerce").rank(method="min", ascending=True).iloc[0]) if "S" in table.columns else 1
    r_rank = int(pd.to_numeric(table["R"], errors="coerce").rank(method="min", ascending=True).iloc[0]) if "R" in table.columns else 1
    advantage_ok = bool(conditions.get("acceptable_advantage", np.isfinite(gap) and np.isfinite(dq) and gap >= dq))
    stability_ok = bool(conditions.get("acceptable_stability", (s_rank == 1) or (r_rank == 1)))

    if advantage_ok and stability_ok:
        conclusion = _xl_text(lang, f"{top_alt} güçlü bir uzlaşı çözümüdür.", f"{top_alt} is a strong compromise solution.")
    elif stability_ok:
        conclusion = _xl_text(
            lang,
            f"{top_alt} lider görünmektedir; ancak {second_alt} ile avantaj farkı sınırlıdır.",
            f"{top_alt} remains the leader, but the advantage over {second_alt} is limited.",
        )
    elif advantage_ok:
        conclusion = _xl_text(
            lang,
            f"{top_alt} Q açısından ayrışmaktadır; ancak S/R kararlılığı tam destek vermemektedir.",
            f"{top_alt} separates on Q, but the S/R stability condition is not fully supported.",
        )
    else:
        conclusion = _xl_text(
            lang,
            f"{top_alt} ile {second_alt} birlikte değerlendirilmelidir; uzlaşı sinyali zayıftır.",
            f"{top_alt} and {second_alt} should be evaluated together; the compromise signal is weak.",
        )

    return pd.DataFrame(
        [
            {
                _xl_text(lang, "Koşul", "Condition"): _xl_text(lang, "Kabul edilebilir avantaj", "Acceptable advantage"),
                _xl_text(lang, "Kanıt", "Evidence"): (
                    _xl_text(lang, "Yeterli alternatif yok; avantaj koşulu otomatik kabul edildi.", "There is only one alternative; the advantage condition is accepted by default.")
                    if len(table) <= 1 else
                    _xl_text(
                        lang,
                        f"DQ = 1/(n-1) = {dq:.4f}; Q2 - Q1 = {gap:.4f} ({second_alt} - {top_alt})",
                        f"DQ = 1/(n-1) = {dq:.4f}; Q2 - Q1 = {gap:.4f} ({second_alt} - {top_alt})",
                    )
                ),
                _xl_text(lang, "Durum", "Status"): _xl_text(lang, "Sağlandı", "Satisfied") if advantage_ok else _xl_text(lang, "Sağlanmadı", "Not satisfied"),
            },
            {
                _xl_text(lang, "Koşul", "Condition"): _xl_text(lang, "Kabul edilebilir kararlılık", "Acceptable stability"),
                _xl_text(lang, "Kanıt", "Evidence"): _xl_text(
                    lang,
                    f"{top_alt} için S sırası = {s_rank}, R sırası = {r_rank}. En az birinde 1. sıra gerekir.",
                    f"For {top_alt}, S rank = {s_rank} and R rank = {r_rank}. The leader should be first in at least one of them.",
                ),
                _xl_text(lang, "Durum", "Status"): _xl_text(lang, "Sağlandı", "Satisfied") if stability_ok else _xl_text(lang, "Sağlanmadı", "Not satisfied"),
            },
            {
                _xl_text(lang, "Koşul", "Condition"): _xl_text(lang, "Sonuç", "Conclusion"),
                _xl_text(lang, "Kanıt", "Evidence"): conclusion,
                _xl_text(lang, "Durum", "Status"): _xl_text(lang, "Güçlü", "Strong") if (advantage_ok and stability_ok) else _xl_text(lang, "İhtiyatlı", "Cautious"),
            },
        ]
    )

def _comparison_method_inference_df(result: Dict[str, Any], lang: str) -> pd.DataFrame:
    comparison = result.get("comparison") or {}
    method_tables = comparison.get("method_tables") or {}
    method_details = comparison.get("method_details") or {}
    weights = (result.get("weights") or {}).get("values") or {}
    rows: List[Dict[str, Any]] = []
    for method_name, method_table in method_tables.items():
        if not isinstance(method_table, pd.DataFrame) or method_table.empty:
            continue
        internal_table = method_table.copy().rename(columns={"Alternative": "Alternatif", "Score": "Skor", "Rank": "Sıra"})
        top_alt_col = col_key(internal_table, "Alternatif", "Alternative")
        score_col = col_key(internal_table, "Skor", "Score")
        top_alt = str(internal_table.iloc[0][top_alt_col])
        top_score_val = pd.to_numeric(internal_table.iloc[0][score_col], errors="coerce")
        layers = _ranking_layers_for_excel(str(method_name), internal_table, weights)
        insight_text = _clean_excel_text(layers.get("analytic") or layers.get("descriptive") or "")
        if "VIKOR" in str(method_name).upper():
            vikor_eval = _vikor_compromise_export_df(str(method_name), method_details.get(method_name, {}) or {}, lang)
            if isinstance(vikor_eval, pd.DataFrame) and not vikor_eval.empty:
                status_col = _xl_text(lang, "Durum", "Status")
                condition_col = _xl_text(lang, "Koşul", "Condition")
                conclusion_row = vikor_eval[vikor_eval[condition_col] == _xl_text(lang, "Sonuç", "Conclusion")]
                if not conclusion_row.empty and status_col in conclusion_row.columns:
                    insight_text = f"{insight_text} {_xl_text(lang, 'Uzlaşı durumu', 'Compromise status')}: {str(conclusion_row.iloc[0][status_col])}.".strip()
        rows.append(
            {
                _xl_text(lang, "Yöntem", "Method"): str(method_name),
                _xl_text(lang, "Lider alternatif", "Leading alternative"): top_alt,
                _xl_text(lang, "Lider skor", "Leading score"): round(float(top_score_val), 4) if pd.notna(top_score_val) else np.nan,
                _xl_text(lang, "İç çıkarsama", "Inference"): insight_text,
            }
        )
    return pd.DataFrame(rows)

def _single_summary_findings_lines(
    result: Dict[str, Any],
    weight_df: pd.DataFrame,
    ranking_df: pd.DataFrame,
    comparison_df: pd.DataFrame,
    sensitivity_df: pd.DataFrame,
    lang: str,
) -> List[str]:
    lines: List[str] = []
    weight_method = str((result.get("weights") or {}).get("method") or "").strip()
    ranking_method = str((result.get("ranking") or {}).get("method") or "").strip()
    if isinstance(weight_df, pd.DataFrame) and not weight_df.empty:
        crit_col = col_key(weight_df, "Kriter", "Criterion")
        top_crit = str(weight_df.iloc[0][crit_col])
        pct_col = _xl_text(lang, "Ağırlık %", "Weight %")
        if pct_col in weight_df.columns:
            top_pct = pd.to_numeric(weight_df.iloc[0][pct_col], errors="coerce")
            if pd.notna(top_pct):
                lines.append(
                    _xl_text(
                        lang,
                        f"{weight_method} ağırlıklandırması altında en baskın kriter {top_crit} olarak belirlenmiştir (%{float(top_pct):.1f}).",
                        f"Under {weight_method}, the most influential criterion is {top_crit} ({float(top_pct):.1f}%).",
                    )
                )
    if isinstance(ranking_df, pd.DataFrame) and not ranking_df.empty:
        alt_col = col_key(ranking_df, "Alternatif", "Alternative")
        score_col = col_key(ranking_df, "Skor", "Score")
        top_alt = str(ranking_df.iloc[0][alt_col])
        top_score = pd.to_numeric(ranking_df.iloc[0][score_col], errors="coerce")
        last_alt = str(ranking_df.iloc[-1][alt_col])
        if pd.notna(top_score):
            lines.append(
                _xl_text(
                    lang,
                    f"{ranking_method} sonucunda {top_alt} {float(top_score):.4f} skoruyla lider alternatif olmuştur; son sırada {last_alt} yer almaktadır.",
                    f"Under {ranking_method}, {top_alt} becomes the leading alternative with a score of {float(top_score):.4f}, while {last_alt} remains at the bottom.",
                )
            )
    vikor_eval = _vikor_compromise_export_df(ranking_method, (result.get("ranking") or {}).get("details", {}) or {}, lang)
    if isinstance(vikor_eval, pd.DataFrame) and not vikor_eval.empty:
        evidence_col = _xl_text(lang, "Kanıt", "Evidence")
        conclusion_mask = vikor_eval[_xl_text(lang, "Koşul", "Condition")] == _xl_text(lang, "Sonuç", "Conclusion")
        if conclusion_mask.any():
            lines.append(str(vikor_eval.loc[conclusion_mask, evidence_col].iloc[0]))
    if isinstance(comparison_df, pd.DataFrame) and not comparison_df.empty:
        mean_rho = _docx_mean_spearman(result.get("comparison") or {})
        if mean_rho is not None:
            level = (
                _xl_text(lang, "yüksek", "high") if mean_rho >= 0.85 else
                _xl_text(lang, "orta", "moderate") if mean_rho >= 0.70 else
                _xl_text(lang, "düşük", "low")
            )
            lines.append(
                _xl_text(
                    lang,
                    f"Yöntemler arası ortalama Spearman uyumu ρ = {mean_rho:.3f} düzeyinde olup {level} metodolojik tutarlılık göstermektedir.",
                    f"The average Spearman agreement across methods is ρ = {mean_rho:.3f}, indicating {level} methodological consistency.",
                )
            )
    if isinstance(sensitivity_df, pd.DataFrame) and not sensitivity_df.empty:
        fp_col = _xl_text(lang, "BirincilikOranı", "FirstPlaceRate")
        alt_col = col_key(sensitivity_df, "Alternatif", "Alternative")
        if fp_col in sensitivity_df.columns and alt_col in sensitivity_df.columns:
            top_alt = str(sensitivity_df.iloc[0][alt_col])
            stability = pd.to_numeric(sensitivity_df.iloc[0][fp_col], errors="coerce")
            if pd.notna(stability):
                lines.append(
                    _xl_text(
                        lang,
                        f"Duyarlılık analizinde {top_alt} alternatifinin birincilik oranı %{float(stability) * 100:.1f} olarak gözlenmiştir.",
                        f"In the sensitivity analysis, {top_alt} retains first place in {float(stability) * 100:.1f}% of the scenarios.",
                    )
                )
    report_findings = _clean_excel_text(str((_build_doc_sections(result, lang) or {}).get(_xl_text(lang, "Bulgular", "Findings"), "") or ""))
    if report_findings:
        first_sentence = re.split(r"(?<=[.!?])\s+", report_findings, maxsplit=1)[0].strip()
        lines.append(first_sentence[:240] + ("..." if len(first_sentence) > 240 else ""))
    return lines[:6]

def _decision_matrix_export_df(selected_data: pd.DataFrame, lang: str) -> pd.DataFrame:
    if not isinstance(selected_data, pd.DataFrame):
        return pd.DataFrame()
    preview = selected_data.copy()
    if "Alternatif" not in preview.columns and "Alternative" not in preview.columns and not isinstance(preview.index, pd.RangeIndex):
        preview = preview.reset_index()
        first_col = str(preview.columns[0])
        preview = preview.rename(columns={first_col: ("Alternative" if lang == "EN" else "Alternatif")})
    return localize_df_lang(preview, lang)

def _weight_analysis_export_df(result: Dict[str, Any], lang: str) -> pd.DataFrame:
    weight_df = (result.get("weights") or {}).get("table")
    if not isinstance(weight_df, pd.DataFrame) or weight_df.empty:
        return pd.DataFrame()
    out = weight_df.copy()
    crit_col = col_key(out, "Kriter", "Criterion")
    weight_col = col_key(out, "Ağırlık", "Weight")
    type_label = _xl_text(lang, "Tip", "Type")
    pct_label = _xl_text(lang, "Ağırlık %", "Weight %")
    visual_label = _xl_text(lang, "Görsel Ağırlık", "Visual Weight")
    criteria_types = result.get("criteria_types", {}) or {}
    out[type_label] = out[crit_col].astype(str).map(
        lambda key: _xl_text(lang, "Fayda", "Benefit") if criteria_types.get(key, "max") == "max" else _xl_text(lang, "Maliyet", "Cost")
    )
    out[pct_label] = pd.to_numeric(out[weight_col], errors="coerce") * 100
    out[visual_label] = out[pct_label].map(
        lambda x: (f"{'█' * max(1, int(round(float(x) / 4)))}  {float(x):.1f}%") if pd.notna(x) else ""
    )
    return localize_df_lang(out, lang)

def _report_text_sections_df(result: Dict[str, Any], lang: str) -> pd.DataFrame:
    sections = _build_doc_sections(result, lang)
    order = [
        _xl_text(lang, "Çalışmanın Amacı", "Objective of the Study"),
        _xl_text(lang, "Çalışmanın Felsefesi", "Philosophy of the Study"),
        _xl_text(lang, "Metodoloji", "Methodology"),
        _xl_text(lang, "Bulgular", "Findings"),
    ]
    rows: List[Dict[str, Any]] = []
    label_map = {
        _xl_text(lang, "Çalışmanın Amacı", "Objective of the Study"): _xl_text(lang, "Çalışmanın Amacı", "Objective of the Study"),
        _xl_text(lang, "Çalışmanın Felsefesi", "Philosophy of the Study"): _xl_text(lang, "Çalışmanın Felsefesi", "Philosophy of the Study"),
        _xl_text(lang, "Metodoloji", "Methodology"): _xl_text(lang, "Metodoloji", "Methodology"),
        _xl_text(lang, "Bulgular", "Findings"): _xl_text(lang, "Bulgular", "Findings"),
    }
    for key in order:
        text = str(sections.get(key, "") or "").strip()
        if text:
            rows.append(
                {
                    _xl_text(lang, "Bölüm", "Section"): label_map.get(key, key),
                    _xl_text(lang, "Metin", "Text"): text,
                }
            )
    return pd.DataFrame(rows)

def _coerce_excel_frame(value: Any, name: str, lang: str) -> pd.DataFrame | None:
    if isinstance(value, pd.DataFrame) and not value.empty:
        return localize_df_lang(value.copy(), lang)
    if isinstance(value, pd.Series) and not value.empty:
        out = value.reset_index()
        out.columns = [_xl_text(lang, "Gösterge", "Metric"), _xl_text(lang, "Değer", "Value")]
        return out
    if isinstance(value, np.ndarray):
        arr = np.asarray(value)
        if arr.ndim == 1 and arr.size > 0:
            return pd.DataFrame({name: arr})
        if arr.ndim == 2 and arr.size > 0:
            return pd.DataFrame(arr)
    if isinstance(value, dict):
        scalar_items = {k: v for k, v in value.items() if np.isscalar(v) or v is None}
        if scalar_items:
            return pd.DataFrame(
                {
                    _xl_text(lang, "Gösterge", "Metric"): list(scalar_items.keys()),
                    _xl_text(lang, "Değer", "Value"): list(scalar_items.values()),
                }
            )
    return None

def _single_export_detail_tables(result: Dict[str, Any], lang: str) -> List[tuple[str, pd.DataFrame]]:
    tables: List[tuple[str, pd.DataFrame]] = []
    corr_matrix = result.get("correlation_matrix")
    if isinstance(corr_matrix, pd.DataFrame) and not corr_matrix.empty:
        tables.append((_xl_text(lang, "Korelasyon_Matrisi", "Correlation_Matrix"), corr_matrix.reset_index()))
    pca_info = result.get("pca") or {}
    for key, title in [
        ("explained_variance", _xl_text(lang, "PCA_Aciklanan_Varyans", "PCA_Explained_Variance")),
        ("explained_ratio", _xl_text(lang, "PCA_Varyans_Orani", "PCA_Variance_Ratio")),
        ("loadings", _xl_text(lang, "PCA_Yukler", "PCA_Loadings")),
        ("score_df", _xl_text(lang, "PCA_Skorlar", "PCA_Scores")),
    ]:
        df = _coerce_excel_frame(pca_info.get(key), key, lang)
        if isinstance(df, pd.DataFrame) and not df.empty:
            tables.append((title, df))
    weight_details = (result.get("weights") or {}).get("details", {}) or {}
    for key, value in weight_details.items():
        if key == "weight_table":
            continue
        df = _coerce_excel_frame(value, key, lang)
        if isinstance(df, pd.DataFrame) and not df.empty:
            tables.append((f"{(result.get('weights') or {}).get('method', 'Weights')}_{key}", df))
    ranking_details = (result.get("ranking") or {}).get("details", {}) or {}
    for key, value in ranking_details.items():
        if key in {"result_table", "score_direction"}:
            continue
        df = _coerce_excel_frame(value, key, lang)
        if isinstance(df, pd.DataFrame) and not df.empty:
            tables.append((f"{(result.get('ranking') or {}).get('method', 'Ranking')}_{key}", df))
    comparison = result.get("comparison") or {}
    for key in ["rank_table", "score_table", "spearman_matrix", "top_alternatives"]:
        df = _coerce_excel_frame(comparison.get(key), key, lang)
        if isinstance(df, pd.DataFrame) and not df.empty:
            tables.append((f"Comparison_{key}", df))
    for method_name, value in (comparison.get("method_tables") or {}).items():
        df = _coerce_excel_frame(value, method_name, lang)
        if isinstance(df, pd.DataFrame) and not df.empty:
            tables.append((f"Compare_{method_name}", df))
    sensitivity = result.get("sensitivity") or {}
    for key in ["monte_carlo_summary", "local_sensitivity"]:
        df = _coerce_excel_frame(sensitivity.get(key), key, lang)
        if isinstance(df, pd.DataFrame) and not df.empty:
            tables.append((f"Sensitivity_{key}", df))
    robustness = result.get("weight_robustness") or {}
    for key in ["leave_one_out", "bootstrap_summary"]:
        df = _coerce_excel_frame(robustness.get(key), key, lang)
        if isinstance(df, pd.DataFrame) and not df.empty:
            tables.append((f"WeightRobust_{key}", df))
    return tables

def _write_single_cover_sheet(writer, sheet_name: str, result: Dict[str, Any], selected_data: pd.DataFrame, lang: str, formats) -> None:
    ws = writer.sheets[sheet_name]
    title = _export_study_title(result, lang)
    weight_method = (result.get("weights") or {}).get("method")
    ranking_method = (result.get("ranking") or {}).get("method")
    subtitle = f"{_excel_method_chain(weight_method, ranking_method, lang)} {_xl_text(lang, 'Analiz Raporu', 'Analysis Report')}"
    ws.merge_range("B2:J3", title, formats["cover_title"])
    ws.merge_range("B4:J4", subtitle, formats["cover_badge"])
    n_alt, n_crit = selected_data.shape if isinstance(selected_data, pd.DataFrame) else (0, 0)
    source_id = str(result.get("data_source_id") or "")
    source_label = source_id or _xl_text(lang, "Belirtilmedi", "Not specified")
    if source_id == "manual_entry":
        source_label = _xl_text(lang, "Manuel tablo girişi", "Manual table entry")
    elif source_id == "sample_data_tr":
        source_label = "Sample Data (TR)"
    elif source_id == "sample_data_en":
        source_label = "Sample Data (EN)"
    elif source_id == "sample_panel_en":
        source_label = "Panel Sample Data (EN)"
    meta_rows = [
        (_xl_text(lang, "Yöntem", "Method"), _excel_method_chain(weight_method, ranking_method, lang)),
        (_xl_text(lang, "Örneklem", "Sample"), f"{n_alt} {_xl_text(lang, 'alternatif', 'alternatives')}"),
        (_xl_text(lang, "Veri kaynağı", "Data source"), source_label),
        (_xl_text(lang, "Kriter sayısı", "Number of criteria"), n_crit),
        (_xl_text(lang, "Analiz kapsamı", "Analysis scope"), _xl_text(lang, "Tek dönem / kesitsel analiz", "Single-period / cross-sectional analysis")),
        (_xl_text(lang, "Oluşturulma zamanı", "Generated at"), pd.Timestamp.now().strftime("%Y-%m-%d %H:%M")),
    ]
    row = 7
    for label, value in meta_rows:
        ws.write(row, 1, label, formats["cover_label_card"])
        ws.merge_range(row, 3, row, 9, value, formats["cover_value_card"])
        row += 2
    sections = _build_doc_sections(result, lang)
    ws.write(row + 1, 1, _xl_text(lang, "Çalışmanın amacı", "Objective"), formats["section"])
    ws.merge_range(row + 2, 1, row + 6, 9, str(sections.get(_xl_text(lang, "Çalışmanın Amacı", "Objective of the Study"), "") or ""), formats["text_block"])
    _write_citation_block(
        ws,
        row + 8,
        f"{_xl_text(lang, 'Önerilen atıf', 'Suggested citation')}: {_SOFTWARE_CITATION}",
        formats,
        end_col=9,
    )
    ws.set_row(1, 30)
    ws.set_row(2, 34)
    ws.set_column(1, 8, 18)
    ws.set_column(3, 9, 22)

def _write_raw_data_sheet(writer, sheet_name: str, df: pd.DataFrame, title: str, subtitle: str, lang: str, formats) -> None:
    ws = writer.sheets[sheet_name]
    start_row = _write_sheet_banner(ws, title, subtitle, formats, end_col=max(8, df.shape[1] if isinstance(df, pd.DataFrame) else 8))
    meta = _write_df_block(
        writer,
        sheet_name,
        df,
        start_row,
        _xl_text(lang, "Ham Veri Tablosu", "Raw Data Table"),
        formats,
        index=False,
        freeze=True,
    )
    if isinstance(df, pd.DataFrame) and not df.empty:
        _apply_heatmap(ws, meta, df)

def _write_vikor_compromise_sheet(
    writer,
    sheet_name: str,
    method_name: str,
    ranking_details: Dict[str, Any],
    result: Dict[str, Any],
    lang: str,
    formats,
) -> None:
    ws = writer.sheets[sheet_name]
    start_row = _write_sheet_banner(
        ws,
        _xl_text(lang, "Uzlaşı Koşulları", "Compromise Conditions"),
        _xl_text(lang, "VIKOR liderinin kabul edilebilir avantaj ve kararlılık koşullarına göre değerlendirmesi", "Evaluation of the VIKOR leader under acceptable advantage and stability conditions"),
        formats,
        end_col=11,
    )
    condition_df = _vikor_compromise_export_df(method_name, ranking_details, lang)
    row = start_row
    if isinstance(condition_df, pd.DataFrame) and not condition_df.empty:
        meta = _write_df_block(
            writer,
            sheet_name,
            condition_df,
            row,
            _xl_text(lang, "Uzlaşı koşulu değerlendirme tablosu", "Compromise condition assessment table"),
            formats,
            index=False,
            freeze=False,
        )
        status_col = _xl_text(lang, "Durum", "Status")
        if status_col in condition_df.columns:
            status_idx = list(condition_df.columns).index(status_col)
            for idx, value in enumerate(condition_df[status_col].astype(str).tolist()):
                fmt = formats["highlight_good"] if any(token in value.casefold() for token in ["sağlandı", "satisfied", "güçlü", "strong"]) else formats["highlight_warn"]
                ws.write(meta["data_row"] + idx, status_idx, value, fmt)
        row = meta["next_row"]
    row = _write_layer_insight_block(
        ws,
        row,
        _xl_text(lang, "VIKOR yöntemi iç çıkarsama", "VIKOR method inference"),
        _ranking_layers_for_excel(method_name, (result.get("ranking") or {}).get("table"), (result.get("weights") or {}).get("values") or {}),
        lang,
        formats,
        end_col=11,
    )
    _write_citation_block(
        ws,
        row + 1,
        f"{_xl_text(lang, 'Önerilen atıf', 'Suggested citation')}: {_SOFTWARE_CITATION}",
        formats,
        end_col=11,
    )
    ws.set_column(0, 11, 18)

def _single_summary_key_rows(result: Dict[str, Any], weight_df: pd.DataFrame, ranking_df: pd.DataFrame, comparison_df: pd.DataFrame, sensitivity_df: pd.DataFrame, lang: str) -> List[tuple[str, Any]]:
    weight_method = (result.get("weights") or {}).get("method") or "—"
    ranking_method = (result.get("ranking") or {}).get("method") or _xl_text(lang, "Uygulanmadı", "Not applied")
    top_alt = "—"
    top_score = "—"
    if isinstance(ranking_df, pd.DataFrame) and not ranking_df.empty:
        alt_col = col_key(ranking_df, "Alternatif", "Alternative")
        score_col = col_key(ranking_df, "Skor", "Score")
        top_alt = str(ranking_df.iloc[0][alt_col])
        score_val = pd.to_numeric(ranking_df.iloc[0][score_col], errors="coerce")
        top_score = f"{float(score_val):.4f}" if pd.notna(score_val) else "—"
    top_criterion = "—"
    if isinstance(weight_df, pd.DataFrame) and not weight_df.empty:
        crit_col = col_key(weight_df, "Kriter", "Criterion")
        pct_col = _xl_text(lang, "Ağırlık %", "Weight %")
        top_criterion = str(weight_df.iloc[0][crit_col])
        if pct_col in weight_df.columns:
            pct_val = pd.to_numeric(weight_df.iloc[0][pct_col], errors="coerce")
            if pd.notna(pct_val):
                top_criterion = f"{top_criterion} ({float(pct_val):.1f}%)"
    agreement = "—"
    if isinstance(comparison_df, pd.DataFrame) and not comparison_df.empty:
        metric_values = pd.to_numeric(comparison_df.select_dtypes(include=[np.number]).stack(), errors="coerce")
        metric_values = metric_values.dropna()
        if not metric_values.empty:
            agreement = f"{float(metric_values.mean()):.3f}"
    stability = "—"
    if isinstance(sensitivity_df, pd.DataFrame) and not sensitivity_df.empty:
        fp_col = _xl_text(lang, "BirincilikOranı", "FirstPlaceRate")
        if fp_col in sensitivity_df.columns:
            fp_val = pd.to_numeric(sensitivity_df.iloc[0][fp_col], errors="coerce")
            if pd.notna(fp_val):
                stability = f"%{float(fp_val) * 100:.1f}"
    return [
        (_xl_text(lang, "Uygulanan analiz", "Applied analysis"), _excel_method_chain(weight_method, ranking_method if ranking_method != _xl_text(lang, "Uygulanmadı", "Not applied") else "", lang)),
        (_xl_text(lang, "Öne çıkan alternatif", "Leading alternative"), f"{top_alt} | {_xl_text(lang, 'Skor', 'Score')}: {top_score}"),
        (_xl_text(lang, "En baskın kriter", "Dominant criterion"), top_criterion),
        (_xl_text(lang, "Yöntem uyumu", "Method agreement"), agreement),
        (_xl_text(lang, "Kararlılık göstergesi", "Robustness signal"), stability),
    ]

def _write_single_summary_sheet(
    writer,
    sheet_name: str,
    result: Dict[str, Any],
    weight_df: pd.DataFrame,
    ranking_df: pd.DataFrame,
    comparison_df: pd.DataFrame,
    sensitivity_df: pd.DataFrame,
    lang: str,
    formats,
) -> None:
    ws = writer.sheets[sheet_name]
    start_row = _write_sheet_banner(
        ws,
        _xl_text(lang, "Özet Bulgular", "Summary Findings"),
        _xl_text(lang, "Karar özeti, en güçlü/zayıf alternatifler ve temel metodolojik sinyaller", "Decision summary, strongest/weakest alternatives, and key methodological signals"),
        formats,
        end_col=11,
    )
    row = start_row
    for label, value in _single_summary_key_rows(result, weight_df, ranking_df, comparison_df, sensitivity_df, lang):
        ws.write(row, 0, label, formats["cover_label_card"])
        ws.merge_range(row, 1, row, 5, value, formats["cover_value_card"])
        row += 2

    summary_row = row + 1
    if isinstance(ranking_df, pd.DataFrame) and not ranking_df.empty:
        rank_col = col_key(ranking_df, "Sıra", "Rank")
        ranking_sorted = ranking_df.sort_values(rank_col, ascending=True).reset_index(drop=True)
        top_df = ranking_sorted.head(5).copy()
        bottom_df = ranking_sorted.tail(5).copy()
        top_meta = _write_df_block(writer, sheet_name, top_df, summary_row, _xl_text(lang, "En Güçlü 5 Alternatif", "Top 5 Alternatives"), formats, index=False, freeze=False)
        ws.merge_range(summary_row, 7, summary_row, 11, _xl_text(lang, "En Zayıf 5 Alternatif", "Bottom 5 Alternatives"), formats["section"])
        bottom_df.to_excel(writer, sheet_name=sheet_name, startrow=summary_row + 2, startcol=7, index=False)
        ws.set_row(summary_row + 2, 24, formats["header"])
        _set_worksheet_widths(ws, bottom_df, startcol=7, index=False)
        ws.autofilter(summary_row + 2, 7, summary_row + 2 + len(bottom_df), 7 + max(0, bottom_df.shape[1] - 1))
        _apply_data_bar(ws, top_meta, top_df, col_key(top_df, "Skor", "Score"), color="#5E81AC")
        bottom_meta = {"data_row": summary_row + 3, "nrows": len(bottom_df)}
        if col_key(bottom_df, "Skor", "Score") in bottom_df.columns:
            score_idx = 7 + list(bottom_df.columns).index(col_key(bottom_df, "Skor", "Score"))
            ws.conditional_format(summary_row + 3, score_idx, summary_row + 2 + len(bottom_df), score_idx, {"type": "data_bar", "bar_color": "#D38B5D"})
        _insert_chart(
            writer.book,
            ws,
            sheet_name,
            {"data_row": top_meta["data_row"], "nrows": min(top_meta["nrows"], 5), "header_row": top_meta["header_row"]},
            top_df.head(5),
            category_col=col_key(top_df, "Alternatif", "Alternative"),
            value_col=col_key(top_df, "Skor", "Score"),
            title=_xl_text(lang, "İlk 5 Alternatif Skoru", "Top 5 Alternative Scores"),
            chart_type="bar",
            anchor_col=13,
            anchor_row=3,
        )
        row = max(top_meta["next_row"], summary_row + len(bottom_df) + 6, row)

    if isinstance(weight_df, pd.DataFrame) and not weight_df.empty:
        top_weight_cols = [c for c in [col_key(weight_df, "Kriter", "Criterion"), _xl_text(lang, "Tip", "Type"), _xl_text(lang, "Ağırlık %", "Weight %")] if c in weight_df.columns]
        weight_focus_df = weight_df[top_weight_cols].head(5).copy() if top_weight_cols else weight_df.head(5).copy()
        weight_meta = _write_df_block(writer, sheet_name, weight_focus_df, row, _xl_text(lang, "En Etkili Kriterler", "Most Influential Criteria"), formats, index=False, freeze=False)
        pct_col = _xl_text(lang, "Ağırlık %", "Weight %")
        if pct_col in weight_focus_df.columns:
            _apply_data_bar(ws, weight_meta, weight_focus_df, pct_col, color="#7CB342")
        row = weight_meta["next_row"]

    row = _write_bullet_block(
        ws,
        row,
        _xl_text(lang, "Otomatik temel bulgular", "Automated key findings"),
        _single_summary_findings_lines(result, weight_df, ranking_df, comparison_df, sensitivity_df, lang),
        formats,
        end_col=11,
    )

    sections = _build_doc_sections(result, lang)
    findings_text = _clean_excel_text(str(sections.get(_xl_text(lang, "Bulgular", "Findings"), "") or ""))
    row = _write_citation_block(
        ws,
        row + 1,
        f"{_xl_text(lang, 'Önerilen atıf', 'Suggested citation')}: {_SOFTWARE_CITATION}",
        formats,
        end_col=11,
    )
    if findings_text:
        ws.merge_range(row, 0, row, 11, _xl_text(lang, "Temel yorum", "Key interpretation"), formats["section"])
        ws.merge_range(row + 2, 0, row + 7, 11, findings_text, formats["text_block"])
    ws.set_column(0, 11, 18)

def _panel_raw_data_export_df(panel_results: Dict[str, Dict[str, Any]], lang: str) -> pd.DataFrame:
    frames: List[pd.DataFrame] = []
    year_label_col = _xl_text(lang, "Yıl", "Year")
    for year_label, year_result in panel_results.items():
        frame = _decision_matrix_export_df(year_result.get("selected_data", pd.DataFrame()), lang)
        if isinstance(frame, pd.DataFrame) and not frame.empty:
            frame.insert(0, year_label_col, str(year_label))
            frames.append(frame)
    return pd.concat(frames, ignore_index=True) if frames else pd.DataFrame()

def _write_panel_summary_findings_sheet(
    writer,
    sheet_name: str,
    panel_results: Dict[str, Dict[str, Any]],
    summary_df: pd.DataFrame,
    rank_df: pd.DataFrame,
    weight_df: pd.DataFrame,
    lang: str,
    formats,
) -> None:
    ws = writer.sheets[sheet_name]
    start_row = _write_sheet_banner(
        ws,
        _xl_text(lang, "Özet Bulgular", "Summary Findings"),
        _xl_text(lang, "Dönemler arası liderlik, ortalama sıralar ve baskın kriterler", "Cross-period leadership, average ranks, and dominant criteria"),
        formats,
        end_col=11,
    )
    row = start_row
    if isinstance(summary_df, pd.DataFrame) and not summary_df.empty:
        top_period = str(summary_df.iloc[0][col_key(summary_df, "Yıl", "Year")])
        top_leader = str(summary_df.iloc[0][col_key(summary_df, "LiderAlternatif", "TopAlternative")])
        ws.write(row, 0, _xl_text(lang, "İlk gözlenen dönem lideri", "First observed period leader"), formats["cover_label_card"])
        ws.merge_range(row, 1, row, 5, f"{top_period} - {top_leader}", formats["cover_value_card"])
        row += 2
    if isinstance(rank_df, pd.DataFrame) and not rank_df.empty:
        top_rank_df = rank_df.head(5).copy()
        bottom_rank_df = rank_df.tail(5).copy()
        top_meta = _write_df_block(writer, sheet_name, top_rank_df, row, _xl_text(lang, "Ortalama Sıraya Göre İlk 5", "Top 5 by Average Rank"), formats, index=False, freeze=False)
        bottom_title = _xl_text(lang, "Ortalama Sıraya Göre Son 5", "Bottom 5 by Average Rank")
        side_start_col = max(7, top_meta["ncols"] + 1)
        bottom_span_end = side_start_col + max(4, len(bottom_rank_df.columns) - 1)
        if bottom_span_end <= 11:
            ws.merge_range(row, side_start_col, row, 11, bottom_title, formats["section"])
            bottom_rank_df.to_excel(writer, sheet_name=sheet_name, startrow=row + 2, startcol=side_start_col, index=False)
            ws.set_row(row + 2, 24, formats["header"])
            _set_worksheet_widths(ws, bottom_rank_df, startcol=side_start_col, index=False)
            row = max(top_meta["next_row"], row + len(bottom_rank_df) + 6)
        else:
            row = top_meta["next_row"]
            bottom_meta = _write_df_block(writer, sheet_name, bottom_rank_df, row, bottom_title, formats, index=False, freeze=False)
            row = bottom_meta["next_row"]
    if isinstance(weight_df, pd.DataFrame) and not weight_df.empty:
        top_weight_df = weight_df.head(6).copy()
        weight_meta = _write_df_block(writer, sheet_name, top_weight_df, row, _xl_text(lang, "Ortalama Ağırlığa Göre Öne Çıkan Kriterler", "Leading Criteria by Average Weight"), formats, index=False, freeze=False)
        avg_weight_col = col_key(top_weight_df, "OrtalamaAğırlık", "AverageWeight")
        if avg_weight_col in top_weight_df.columns:
            _apply_data_bar(ws, weight_meta, top_weight_df, avg_weight_col, color="#7CB342")
        row = weight_meta["next_row"]
    row = _write_citation_block(
        ws,
        row + 1,
        f"{_xl_text(lang, 'Önerilen atıf', 'Suggested citation')}: {_SOFTWARE_CITATION}",
        formats,
        end_col=11,
    )
    ws.set_column(0, 11, 18)

def _excel_col_label(col_idx: int) -> str:
    col_idx = int(col_idx)
    label = ""
    while col_idx >= 0:
        col_idx, rem = divmod(col_idx, 26)
        label = chr(65 + rem) + label
        col_idx -= 1
    return label

def _excel_cell(row_idx: int, col_idx: int) -> str:
    return f"{_excel_col_label(col_idx)}{int(row_idx) + 1}"

def _apply_data_bar(
    ws,
    meta: Dict[str, int],
    df: pd.DataFrame,
    column_name: str,
    *,
    color: str = "#5A8FCB",
) -> None:
    if meta.get("nrows", 0) <= 0 or column_name not in df.columns:
        return
    col_idx = list(df.columns).index(column_name)
    ws.conditional_format(
        meta["data_row"],
        col_idx,
        meta["data_row"] + meta["nrows"] - 1,
        col_idx,
        {"type": "data_bar", "bar_color": color},
    )

def _apply_heatmap(
    ws,
    meta: Dict[str, int],
    df: pd.DataFrame,
    *,
    start_col_name: str | None = None,
    colors: tuple[str, str, str] = ("#F8D7DA", "#FFF3CD", "#D1E7DD"),
) -> None:
    if meta.get("nrows", 0) <= 0 or df.empty:
        return
    numeric_cols = [c for c in df.columns if pd.api.types.is_numeric_dtype(df[c])]
    if not numeric_cols:
        return
    if start_col_name and start_col_name in df.columns:
        start_col = list(df.columns).index(start_col_name)
        end_col = max(start_col, len(df.columns) - 1)
    else:
        start_col = min(list(df.columns).index(numeric_cols[0]), len(df.columns) - 1)
        end_col = max(list(df.columns).index(c) for c in numeric_cols)
    ws.conditional_format(
        meta["data_row"],
        start_col,
        meta["data_row"] + meta["nrows"] - 1,
        end_col,
        {
            "type": "3_color_scale",
            "min_color": colors[0],
            "mid_color": colors[1],
            "max_color": colors[2],
        },
    )

def _insert_chart(
    workbook,
    ws,
    sheet_name: str,
    meta: Dict[str, int],
    df: pd.DataFrame,
    *,
    category_col: str,
    value_col: str,
    title: str,
    chart_type: str = "column",
    anchor_col: int = 11,
    anchor_row: int | None = None,
    x_scale: float = 1.12,
    y_scale: float = 1.05,
) -> None:
    if meta.get("nrows", 0) <= 0 or category_col not in df.columns or value_col not in df.columns:
        return
    cat_idx = list(df.columns).index(category_col)
    val_idx = list(df.columns).index(value_col)
    chart = workbook.add_chart({"type": chart_type})
    chart.add_series(
        {
            "name": title,
            "categories": [sheet_name, meta["data_row"], cat_idx, meta["data_row"] + meta["nrows"] - 1, cat_idx],
            "values": [sheet_name, meta["data_row"], val_idx, meta["data_row"] + meta["nrows"] - 1, val_idx],
            "fill": {"color": "#7FB3D5"},
            "border": {"color": "#1F4A73"},
        }
    )
    chart.set_title({"name": title})
    chart.set_legend({"none": True})
    chart.set_style(10)
    ws.insert_chart(_excel_cell(anchor_row if anchor_row is not None else meta["header_row"], anchor_col), chart, {"x_scale": x_scale, "y_scale": y_scale})

def generate_excel(result: Dict[str, Any], selected_data: pd.DataFrame, lang: str = "TR") -> bytes:
    output = io.BytesIO()
    used_names: set[str] = set()
    cover_name = _safe_sheet_name(_xl_text(lang, "Kapak", "Cover"), used_names)
    raw_name = _safe_sheet_name(_xl_text(lang, "Ham Veri", "Raw Data"), used_names)

    decision_df = _decision_matrix_export_df(selected_data, lang)
    stats_raw = result.get("stats")
    stats_df = localize_df_lang(stats_raw.copy(), lang) if isinstance(stats_raw, pd.DataFrame) else pd.DataFrame()
    weight_df = _weight_analysis_export_df(result, lang)
    ranking_raw = (result.get("ranking") or {}).get("table")
    ranking_df = localize_df_lang(ranking_raw.copy(), lang) if isinstance(ranking_raw, pd.DataFrame) else pd.DataFrame()
    comparison = result.get("comparison") or {}
    comparison_raw = comparison.get("spearman_matrix")
    comparison_df = localize_df_lang(comparison_raw.copy(), lang) if isinstance(comparison_raw, pd.DataFrame) else pd.DataFrame()
    sensitivity_raw = (result.get("sensitivity") or {}).get("monte_carlo_summary")
    sensitivity_df = localize_df_lang(sensitivity_raw.copy(), lang) if isinstance(sensitivity_raw, pd.DataFrame) else pd.DataFrame()
    local_sensitivity_raw = (result.get("sensitivity") or {}).get("local_scenarios")
    local_sensitivity_df = localize_df_lang(local_sensitivity_raw.copy(), lang) if isinstance(local_sensitivity_raw, pd.DataFrame) else pd.DataFrame()
    ranking_details = (result.get("ranking") or {}).get("details", {}) or {}
    comparison_method_inference_df = localize_df_lang(_comparison_method_inference_df(result, lang), lang)

    weight_method = str((result.get("weights") or {}).get("method") or _xl_text(lang, "Ağırlık", "Weight")).strip()
    ranking_method = str((result.get("ranking") or {}).get("method") or "").strip()
    weight_sheet_name = _safe_sheet_name(f"{weight_method} {_xl_text(lang, 'Analizi', 'Analysis')}", used_names)
    ranking_sheet_name = _safe_sheet_name(f"{ranking_method} {_xl_text(lang, 'Sonuçları', 'Results')}", used_names) if ranking_method else None
    compromise_sheet_name = _safe_sheet_name(_xl_text(lang, "Uzlaşı Koşulları", "Compromise Conditions"), used_names) if ("VIKOR" in ranking_method.upper()) else None
    comparison_sheet_name = _safe_sheet_name(_xl_text(lang, "Yöntem Karşılaştırması", "Method Comparison"), used_names) if isinstance(comparison_df, pd.DataFrame) and not comparison_df.empty else None
    sensitivity_sheet_name = _safe_sheet_name(_xl_text(lang, "Duyarlılık Analizi", "Sensitivity Analysis"), used_names) if (
        (isinstance(sensitivity_df, pd.DataFrame) and not sensitivity_df.empty)
        or (isinstance(local_sensitivity_df, pd.DataFrame) and not local_sensitivity_df.empty)
    ) else None
    summary_sheet_name = _safe_sheet_name(_xl_text(lang, "Özet Bulgular", "Summary Findings"), used_names)

    weight_detail_tables: List[tuple[str, pd.DataFrame]] = []
    for key, value in ((result.get("weights") or {}).get("details", {}) or {}).items():
        if key == "weight_table":
            continue
        df = _coerce_excel_frame(value, key, lang)
        if isinstance(df, pd.DataFrame) and not df.empty:
            weight_detail_tables.append((key, df))

    ranking_detail_tables: List[tuple[str, pd.DataFrame]] = []
    for key, value in ((result.get("ranking") or {}).get("details", {}) or {}).items():
        if key in {"result_table", "score_direction"}:
            continue
        df = _coerce_excel_frame(value, key, lang)
        if isinstance(df, pd.DataFrame) and not df.empty:
            ranking_detail_tables.append((key, df))

    comparison_detail_tables: List[tuple[str, pd.DataFrame]] = []
    for key in ["rank_table", "score_table", "top_alternatives"]:
        df = _coerce_excel_frame(comparison.get(key), key, lang)
        if isinstance(df, pd.DataFrame) and not df.empty:
            comparison_detail_tables.append((key, df))
    for method_name, value in (comparison.get("method_tables") or {}).items():
        df = _coerce_excel_frame(value, method_name, lang)
        if isinstance(df, pd.DataFrame) and not df.empty:
            comparison_detail_tables.append((method_name, df))

    with pd.ExcelWriter(output, engine="xlsxwriter") as writer:
        workbook = writer.book
        formats = _excel_formats(workbook)
        sheet_names = [cover_name, raw_name, weight_sheet_name]
        if ranking_sheet_name:
            sheet_names.append(ranking_sheet_name)
        if compromise_sheet_name:
            sheet_names.append(compromise_sheet_name)
        if comparison_sheet_name:
            sheet_names.append(comparison_sheet_name)
        if sensitivity_sheet_name:
            sheet_names.append(sensitivity_sheet_name)
        sheet_names.append(summary_sheet_name)
        for sheet_name in sheet_names:
            writer.sheets[sheet_name] = workbook.add_worksheet(sheet_name)

        _write_single_cover_sheet(writer, cover_name, result, selected_data, lang, formats)
        _write_raw_data_sheet(
            writer,
            raw_name,
            decision_df,
            _xl_text(lang, "Ham Veri", "Raw Data"),
            _xl_text(lang, "Analizde kullanılan karar matrisi ve temizlenmiş giriş tablosu", "Decision matrix and cleaned input table used in the analysis"),
            lang,
            formats,
        )

        weight_ws = writer.sheets[weight_sheet_name]
        weight_ws.set_column(10, 18, 15)
        weight_row = _write_sheet_banner(
            weight_ws,
            f"{weight_method} {_xl_text(lang, 'Analizi', 'Analysis')}",
            _xl_text(lang, "Kriter ağırlıklarının dağılımı ve yöntem özel hesaplama çıktıları", "Criterion-weight distribution and method-specific calculation outputs"),
            formats,
            end_col=11,
        )
        weight_meta = _write_df_block(writer, weight_sheet_name, weight_df, weight_row, _xl_text(lang, "Kriter Ağırlık Tablosu", "Criterion Weight Table"), formats, index=False, freeze=True)
        if weight_meta["nrows"] > 0:
            _apply_data_bar(weight_ws, weight_meta, weight_df, col_key(weight_df, "Ağırlık", "Weight"), color="#5A8FCB")
            pct_col = _xl_text(lang, "Ağırlık %", "Weight %")
            if pct_col in weight_df.columns:
                _apply_data_bar(weight_ws, weight_meta, weight_df, pct_col, color="#7CB342")
            _insert_chart(
                workbook,
                weight_ws,
                weight_sheet_name,
                {"data_row": weight_meta["data_row"], "nrows": min(weight_meta["nrows"], 10), "header_row": weight_meta["header_row"]},
                weight_df.head(10),
                category_col=col_key(weight_df, "Kriter", "Criterion"),
                value_col=col_key(weight_df, "Ağırlık", "Weight"),
                title=_xl_text(lang, "Kriter Ağırlıkları", "Criterion Weights"),
                chart_type="column",
                anchor_col=10,
                anchor_row=3,
            )
        weight_row = max(weight_meta["next_row"], 28)
        for key, df in weight_detail_tables[:4]:
            meta = _write_df_block(writer, weight_sheet_name, df, weight_row, str(key), formats, index=False, freeze=False)
            if "corr" in str(key).lower() or "matrix" in str(key).lower():
                _apply_heatmap(weight_ws, meta, df)
            weight_row = meta["next_row"]
        weight_row = _write_layer_insight_block(
            weight_ws,
            weight_row,
            _xl_text(lang, "Ağırlık yöntemi iç çıkarsama", "Weight-method inference"),
            _weight_layers_for_excel(result),
            lang,
            formats,
            end_col=11,
        )
        _write_citation_block(weight_ws, weight_row + 1, f"{_xl_text(lang, 'Önerilen atıf', 'Suggested citation')}: {_SOFTWARE_CITATION}", formats, end_col=11)

        if ranking_sheet_name:
            ranking_ws = writer.sheets[ranking_sheet_name]
            ranking_ws.set_column(10, 18, 15)
            ranking_row = _write_sheet_banner(
                ranking_ws,
                f"{ranking_method} {_xl_text(lang, 'Sonuçları', 'Results')}",
                _xl_text(lang, "Alternatif skorları, sıralar ve yöntem özel çıktı tabloları", "Alternative scores, ranks, and method-specific output tables"),
                formats,
                end_col=11,
            )
            if isinstance(ranking_df, pd.DataFrame) and not ranking_df.empty:
                rank_meta = _write_df_block(writer, ranking_sheet_name, ranking_df, ranking_row, _xl_text(lang, "Sıralama Tablosu", "Ranking Table"), formats, index=False, freeze=True)
                _apply_data_bar(ranking_ws, rank_meta, ranking_df, col_key(ranking_df, "Skor", "Score"), color="#5E81AC")
                _insert_chart(
                    workbook,
                    ranking_ws,
                    ranking_sheet_name,
                    {"data_row": rank_meta["data_row"], "nrows": min(rank_meta["nrows"], 10), "header_row": rank_meta["header_row"]},
                    ranking_df.head(10),
                    category_col=col_key(ranking_df, "Alternatif", "Alternative"),
                    value_col=col_key(ranking_df, "Skor", "Score"),
                    title=_xl_text(lang, "İlk Alternatiflerin Skorları", "Scores of Top Alternatives"),
                    chart_type="bar",
                    anchor_col=10,
                    anchor_row=3,
                )
                ranking_row = max(rank_meta["next_row"], 28)
            if isinstance(stats_df, pd.DataFrame) and not stats_df.empty:
                stats_meta = _write_df_block(writer, ranking_sheet_name, stats_df, ranking_row, _xl_text(lang, "Tanımlayıcı İstatistikler", "Descriptive Statistics"), formats, index=False, freeze=False)
                _apply_heatmap(ranking_ws, stats_meta, stats_df)
                ranking_row = stats_meta["next_row"]
            for key, df in ranking_detail_tables[:4]:
                meta = _write_df_block(writer, ranking_sheet_name, df, ranking_row, str(key), formats, index=False, freeze=False)
                if "matrix" in str(key).lower():
                    _apply_heatmap(ranking_ws, meta, df)
                ranking_row = meta["next_row"]
            ranking_row = _write_layer_insight_block(
                ranking_ws,
                ranking_row,
                _xl_text(lang, "Sıralama yöntemi iç çıkarsama", "Ranking-method inference"),
                _ranking_layers_for_excel(ranking_method, (result.get("ranking") or {}).get("table"), (result.get("weights") or {}).get("values") or {}),
                lang,
                formats,
                end_col=11,
            )
            _write_citation_block(ranking_ws, ranking_row + 1, f"{_xl_text(lang, 'Önerilen atıf', 'Suggested citation')}: {_SOFTWARE_CITATION}", formats, end_col=11)

        if compromise_sheet_name:
            _write_vikor_compromise_sheet(
                writer,
                compromise_sheet_name,
                ranking_method,
                ranking_details,
                result,
                lang,
                formats,
            )

        if comparison_sheet_name:
            comparison_ws = writer.sheets[comparison_sheet_name]
            comparison_row = _write_sheet_banner(
                comparison_ws,
                _xl_text(lang, "Yöntem Karşılaştırması", "Method Comparison"),
                _xl_text(lang, "Farklı yöntemlerin sıralama uyumu ve karşılaştırmalı sonuçları", "Ranking agreement and comparative outcomes across methods"),
                formats,
                end_col=11,
            )
            comp_meta = _write_df_block(writer, comparison_sheet_name, comparison_df, comparison_row, _xl_text(lang, "Spearman Uyum Matrisi", "Spearman Agreement Matrix"), formats, index=False, freeze=True)
            _apply_heatmap(comparison_ws, comp_meta, comparison_df)
            comparison_row = comp_meta["next_row"]
            for key, df in comparison_detail_tables[:4]:
                meta = _write_df_block(writer, comparison_sheet_name, df, comparison_row, str(key), formats, index=False, freeze=False)
                if "matrix" in str(key).lower():
                    _apply_heatmap(comparison_ws, meta, df)
                comparison_row = meta["next_row"]
            comparison_row = _write_layer_insight_block(
                comparison_ws,
                comparison_row,
                _xl_text(lang, "Yöntem karşılaştırması yorumu", "Method-comparison interpretation"),
                _comparison_layers_for_excel(result),
                lang,
                formats,
                end_col=11,
            )
            if isinstance(comparison_method_inference_df, pd.DataFrame) and not comparison_method_inference_df.empty:
                comparison_meta = _write_df_block(
                    writer,
                    comparison_sheet_name,
                    comparison_method_inference_df,
                    comparison_row,
                    _xl_text(lang, "Yöntem bazlı iç çıkarsamalar", "Method-specific inferences"),
                    formats,
                    index=False,
                    freeze=False,
                )
                comparison_row = comparison_meta["next_row"]
            _write_citation_block(comparison_ws, comparison_row + 1, f"{_xl_text(lang, 'Önerilen atıf', 'Suggested citation')}: {_SOFTWARE_CITATION}", formats, end_col=11)

        if sensitivity_sheet_name:
            sensitivity_ws = writer.sheets[sensitivity_sheet_name]
            sensitivity_row = _write_sheet_banner(
                sensitivity_ws,
                _xl_text(lang, "Duyarlılık Analizi", "Sensitivity Analysis"),
                _xl_text(lang, "Monte Carlo sonuçları ve kriter bazlı yerel duyarlılık çıktıları", "Monte Carlo outputs and criterion-level local sensitivity results"),
                formats,
                end_col=11,
            )
            if isinstance(sensitivity_df, pd.DataFrame) and not sensitivity_df.empty:
                sens_meta = _write_df_block(writer, sensitivity_sheet_name, sensitivity_df, sensitivity_row, _xl_text(lang, "Monte Carlo Özeti", "Monte Carlo Summary"), formats, index=False, freeze=True)
                fp_col = _xl_text(lang, "BirincilikOranı", "FirstPlaceRate")
                if fp_col in sensitivity_df.columns:
                    _apply_data_bar(sensitivity_ws, sens_meta, sensitivity_df, fp_col, color="#7CB342")
                _apply_heatmap(sensitivity_ws, sens_meta, sensitivity_df)
                sensitivity_row = sens_meta["next_row"]
            if isinstance(local_sensitivity_df, pd.DataFrame) and not local_sensitivity_df.empty:
                local_meta = _write_df_block(writer, sensitivity_sheet_name, local_sensitivity_df, sensitivity_row, _xl_text(lang, "Yerel Duyarlılık", "Local Sensitivity"), formats, index=False, freeze=False)
                _apply_heatmap(sensitivity_ws, local_meta, local_sensitivity_df)
                sensitivity_row = local_meta["next_row"]
            sensitivity_row = _write_layer_insight_block(
                sensitivity_ws,
                sensitivity_row,
                _xl_text(lang, "Duyarlılık yorumu", "Sensitivity interpretation"),
                _sensitivity_layers_for_excel(result),
                lang,
                formats,
                end_col=11,
            )
            _write_citation_block(sensitivity_ws, sensitivity_row + 1, f"{_xl_text(lang, 'Önerilen atıf', 'Suggested citation')}: {_SOFTWARE_CITATION}", formats, end_col=11)

        _write_single_summary_sheet(
            writer,
            summary_sheet_name,
            result,
            weight_df,
            ranking_df,
            comparison_df,
            sensitivity_df,
            lang,
            formats,
        )

    return output.getvalue()

def _panel_summary_rows(panel_results: Dict[str, Dict[str, Any]]) -> pd.DataFrame:
    rows: List[Dict[str, Any]] = []
    for year_label, year_result in panel_results.items():
        ranking_table = year_result.get("ranking", {}).get("table")
        top_alt = "—"
        top_score = np.nan
        if isinstance(ranking_table, pd.DataFrame) and not ranking_table.empty:
            alt_col = col_key(ranking_table, "Alternatif", "Alternative")
            score_col = col_key(ranking_table, "Skor", "Score")
            top_alt = str(ranking_table.iloc[0][alt_col])
            top_score = float(pd.to_numeric(ranking_table.iloc[0][score_col], errors="coerce"))
        strategy_key = str(year_result.get("panel_weight_strategy", "yearly")).strip().lower()
        strategy_label = "global" if strategy_key == "global" else "yearly"
        rows.append(
            {
                "Yıl": year_label,
                "AğırlıkYöntemi": year_result.get("weights", {}).get("method"),
                "AğırlıkStratejisi": strategy_label,
                "SıralamaYöntemi": year_result.get("ranking", {}).get("method") or "—",
                "LiderAlternatif": top_alt,
                "LiderSkor": top_score,
                "LiderKararlılığı": (year_result.get("sensitivity") or {}).get("top_stability"),
            }
        )
    return pd.DataFrame(rows)

def _panel_rank_matrix(panel_results: Dict[str, Dict[str, Any]]) -> pd.DataFrame:
    by_year: Dict[str, pd.Series] = {}
    for year_label, year_result in panel_results.items():
        ranking_table = year_result.get("ranking", {}).get("table")
        if not isinstance(ranking_table, pd.DataFrame) or ranking_table.empty:
            continue
        alt_col = col_key(ranking_table, "Alternatif", "Alternative")
        rank_col = col_key(ranking_table, "Sıra", "Rank")
        if alt_col not in ranking_table.columns or rank_col not in ranking_table.columns:
            continue
        tmp = ranking_table[[alt_col, rank_col]].copy()
        tmp[alt_col] = tmp[alt_col].astype(str).str.strip()
        tmp = tmp[tmp[alt_col] != ""]
        tmp[rank_col] = pd.to_numeric(tmp[rank_col], errors="coerce")
        tmp = tmp.dropna(subset=[rank_col])
        if tmp.empty:
            continue
        series = tmp.groupby(alt_col, dropna=False)[rank_col].min()
        by_year[str(year_label)] = series
    if not by_year:
        return pd.DataFrame()
    matrix = pd.concat(by_year, axis=1, sort=True)
    year_cols = list(matrix.columns)
    matrix["OrtalamaSıra"] = matrix[year_cols].mean(axis=1, skipna=True)
    matrix["GözlenenYılSayısı"] = matrix[year_cols].notna().sum(axis=1).astype(int)
    matrix = matrix.sort_values(["OrtalamaSıra", "GözlenenYılSayısı"], ascending=[True, False]).reset_index()
    matrix = matrix.rename(columns={"index": "Alternatif"})
    for col in year_cols + ["OrtalamaSıra"]:
        if col in matrix.columns:
            matrix[col] = pd.to_numeric(matrix[col], errors="coerce").round(2)
    return matrix

def _panel_score_matrix(panel_results: Dict[str, Dict[str, Any]]) -> pd.DataFrame:
    by_year: Dict[str, pd.Series] = {}
    for year_label, year_result in panel_results.items():
        ranking_table = year_result.get("ranking", {}).get("table")
        if not isinstance(ranking_table, pd.DataFrame) or ranking_table.empty:
            continue
        alt_col = col_key(ranking_table, "Alternatif", "Alternative")
        score_col = col_key(ranking_table, "Skor", "Score")
        if alt_col not in ranking_table.columns or score_col not in ranking_table.columns:
            continue
        tmp = ranking_table[[alt_col, score_col]].copy()
        tmp[alt_col] = tmp[alt_col].astype(str).str.strip()
        tmp = tmp[tmp[alt_col] != ""]
        tmp[score_col] = pd.to_numeric(tmp[score_col], errors="coerce")
        tmp = tmp.dropna(subset=[score_col])
        if tmp.empty:
            continue
        series = tmp.groupby(alt_col, dropna=False)[score_col].mean()
        by_year[str(year_label)] = series
    if not by_year:
        return pd.DataFrame()
    matrix = pd.concat(by_year, axis=1, sort=True)
    year_cols = list(matrix.columns)
    matrix["OrtalamaSkor"] = matrix[year_cols].mean(axis=1, skipna=True)
    matrix["GözlenenYılSayısı"] = matrix[year_cols].notna().sum(axis=1).astype(int)
    matrix = matrix.sort_values(["OrtalamaSkor", "GözlenenYılSayısı"], ascending=[False, False]).reset_index()
    matrix = matrix.rename(columns={"index": "Alternatif"})
    for col in year_cols + ["OrtalamaSkor"]:
        if col in matrix.columns:
            matrix[col] = pd.to_numeric(matrix[col], errors="coerce").round(4)
    return matrix

def _panel_weight_matrix(panel_results: Dict[str, Dict[str, Any]]) -> pd.DataFrame:
    by_year: Dict[str, pd.Series] = {}
    for year_label, year_result in panel_results.items():
        wt = (year_result.get("weights") or {}).get("table")
        if not isinstance(wt, pd.DataFrame) or wt.empty:
            continue
        crit_col = col_key(wt, "Kriter", "Criterion")
        w_col = col_key(wt, "Ağırlık", "Weight")
        if crit_col not in wt.columns or w_col not in wt.columns:
            continue
        tmp = wt[[crit_col, w_col]].copy()
        tmp[crit_col] = tmp[crit_col].astype(str).str.strip()
        tmp = tmp[tmp[crit_col] != ""]
        tmp[w_col] = pd.to_numeric(tmp[w_col], errors="coerce")
        tmp = tmp.dropna(subset=[w_col])
        if tmp.empty:
            continue
        series = tmp.groupby(crit_col, dropna=False)[w_col].mean()
        by_year[str(year_label)] = series
    if not by_year:
        return pd.DataFrame()
    matrix = pd.concat(by_year, axis=1, sort=True)
    year_cols = list(matrix.columns)
    matrix["OrtalamaAğırlık"] = matrix[year_cols].mean(axis=1, skipna=True)
    matrix = matrix.sort_values("OrtalamaAğırlık", ascending=False).reset_index()
    matrix = matrix.rename(columns={"index": "Kriter"})
    for col in year_cols + ["OrtalamaAğırlık"]:
        if col in matrix.columns:
            matrix[col] = pd.to_numeric(matrix[col], errors="coerce").round(4)
    return matrix

def _write_panel_cover_sheet(writer, sheet_name: str, panel_results: Dict[str, Dict[str, Any]], lang: str, formats) -> None:
    ws = writer.sheets[sheet_name]
    first_result = next(iter(panel_results.values()))
    year_labels = list(panel_results.keys())
    title = str(st.session_state.get("study_title", "") or "").strip() or _xl_text(lang, "Panel Analiz Raporu", "Panel Analysis Report")
    subtitle = f"{_excel_method_chain((first_result.get('weights') or {}).get('method'), (first_result.get('ranking') or {}).get('method'), lang)} {_xl_text(lang, 'Panel Analiz Raporu', 'Panel Analysis Report')}"
    ws.merge_range("B2:J3", title, formats["cover_title"])
    ws.merge_range("B4:J4", subtitle, formats["cover_badge"])
    rank_matrix = _panel_rank_matrix(panel_results)
    meta_rows = [
        (_xl_text(lang, "Yöntem", "Method"), _excel_method_chain((first_result.get("weights") or {}).get("method"), (first_result.get("ranking") or {}).get("method"), lang)),
        (_xl_text(lang, "Dönem sayısı", "Number of periods"), len(year_labels)),
        (_xl_text(lang, "Dönemler", "Periods"), ", ".join(map(str, year_labels))),
        (_xl_text(lang, "İzlenen alternatif sayısı", "Tracked alternatives"), int(rank_matrix.shape[0]) if isinstance(rank_matrix, pd.DataFrame) else 0),
        (_xl_text(lang, "Oluşturulma zamanı", "Generated at"), pd.Timestamp.now().strftime("%Y-%m-%d %H:%M")),
    ]
    row = 7
    for label, value in meta_rows:
        ws.write(row, 1, label, formats["cover_label_card"])
        ws.merge_range(row, 3, row, 9, value, formats["cover_value_card"])
        row += 2
    first_sections = _build_doc_sections(first_result, lang)
    ws.write(row + 1, 1, _xl_text(lang, "Çalışmanın amacı", "Objective"), formats["section"])
    ws.merge_range(
        row + 2,
        1,
        row + 6,
        9,
        str(first_sections.get(_xl_text(lang, "Çalışmanın Amacı", "Objective of the Study"), "") or ""),
        formats["text_block"],
    )
    _write_citation_block(
        ws,
        row + 8,
        f"{_xl_text(lang, 'Önerilen atıf', 'Suggested citation')}: {_SOFTWARE_CITATION}",
        formats,
        end_col=9,
    )
    ws.set_column(1, 8, 18)

def generate_panel_excel(panel_results: Dict[str, Dict[str, Any]], lang: str = "TR") -> bytes:
    output = io.BytesIO()
    used_names: set[str] = set()
    cover_name = _safe_sheet_name(_xl_text(lang, "Kapak", "Cover"), used_names)
    raw_name = _safe_sheet_name(_xl_text(lang, "Ham Veri", "Raw Data"), used_names)
    summary_name = _safe_sheet_name(_xl_text(lang, "Panel_Ozet", "Panel_Summary"), used_names)
    matrix_name = _safe_sheet_name(_xl_text(lang, "Panel_Matris", "Panel_Matrix"), used_names)
    findings_name = _safe_sheet_name(_xl_text(lang, "Özet Bulgular", "Summary Findings"), used_names)

    summary_df = localize_df_lang(_panel_summary_rows(panel_results), lang)
    rank_df = localize_df_lang(_panel_rank_matrix(panel_results), lang)
    score_df = localize_df_lang(_panel_score_matrix(panel_results), lang)
    weight_df = localize_df_lang(_panel_weight_matrix(panel_results), lang)
    raw_df = _panel_raw_data_export_df(panel_results, lang)

    with pd.ExcelWriter(output, engine="xlsxwriter") as writer:
        workbook = writer.book
        formats = _excel_formats(workbook)
        for sheet_name in [cover_name, raw_name, summary_name, matrix_name, findings_name]:
            writer.sheets[sheet_name] = workbook.add_worksheet(sheet_name)

        _write_panel_cover_sheet(writer, cover_name, panel_results, lang, formats)
        _write_raw_data_sheet(
            writer,
            raw_name,
            raw_df,
            _xl_text(lang, "Ham Veri", "Raw Data"),
            _xl_text(lang, "Panel analizinde kullanılan dönemsel karar matrisi", "Period-wise decision matrix used in the panel analysis"),
            lang,
            formats,
        )
        summary_ws = writer.sheets[summary_name]
        summary_ws.set_column(11, 19, 15)
        summary_ws.freeze_panes(2, 0)
        summary_start = _write_sheet_banner(
            summary_ws,
            _xl_text(lang, "Panel Sonuç Özeti", "Panel Result Summary"),
            _xl_text(lang, "Dönem liderleri, skor eğilimleri ve genel panel görünümü", "Period leaders, score trends, and the overall panel view"),
            formats,
            end_col=11,
        )
        matrix_ws = writer.sheets[matrix_name]
        matrix_ws.set_column(0, 12, 15)
        matrix_ws.freeze_panes(2, 0)
        matrix_start = _write_sheet_banner(
            matrix_ws,
            _xl_text(lang, "Panel Matrisleri", "Panel Matrices"),
            _xl_text(lang, "Yıllara göre sıralama, skor ve ağırlık değişim tabloları", "Year-by-year ranking, score, and weight change tables"),
            formats,
            end_col=11,
        )

        summary_meta = _write_df_block(
            writer,
            summary_name,
            summary_df,
            summary_start,
            _xl_text(lang, "Panel Sonuç Özeti Tablosu", "Panel Summary Table"),
            formats,
            index=False,
            freeze=False,
        )
        _apply_data_bar(summary_ws, summary_meta, summary_df, col_key(summary_df, "LiderSkor", "TopScore"), color="#5E81AC")
        _apply_data_bar(summary_ws, summary_meta, summary_df, col_key(summary_df, "LiderKararlılığı", "LeaderStability"), color="#7CB342")
        _insert_chart(
            workbook,
            summary_ws,
            summary_name,
            summary_meta,
            summary_df,
            category_col=col_key(summary_df, "Yıl", "Year"),
            value_col=col_key(summary_df, "LiderSkor", "TopScore"),
            title=_xl_text(lang, "Yıllara Göre Lider Skoru", "Leader Score by Period"),
            chart_type="line",
            anchor_col=11,
            anchor_row=1,
        )
        if isinstance(score_df, pd.DataFrame) and not score_df.empty and col_key(score_df, "Alternatif", "Alternative") in score_df.columns:
            _insert_chart(
                workbook,
                summary_ws,
                summary_name,
                {"data_row": summary_meta["data_row"], "nrows": min(8, len(score_df)), "header_row": summary_meta["header_row"]},
                score_df.head(8),
                category_col=col_key(score_df, "Alternatif", "Alternative"),
                value_col=col_key(score_df, "OrtalamaSkor", "AverageScore"),
                title=_xl_text(lang, "Ortalama Skora Göre İlk Alternatifler", "Top Alternatives by Average Score"),
                chart_type="bar",
                anchor_col=11,
                anchor_row=20,
            )
        _write_citation_block(summary_ws, max(summary_meta["next_row"], 34), f"{_xl_text(lang, 'Önerilen atıf', 'Suggested citation')}: {_SOFTWARE_CITATION}", formats, end_col=11)

        matrix_row = matrix_start
        rank_meta = _write_df_block(
            writer,
            matrix_name,
            rank_df,
            matrix_row,
            _xl_text(lang, "Alternatiflerin Yıllara Göre Sıraları", "Alternative Ranks by Period"),
            formats,
            index=False,
            freeze=False,
        )
        _apply_heatmap(matrix_ws, rank_meta, rank_df)
        matrix_row = rank_meta["next_row"]

        score_meta = _write_df_block(
            writer,
            matrix_name,
            score_df,
            matrix_row,
            _xl_text(lang, "Alternatiflerin Yıllara Göre Skorları", "Alternative Scores by Period"),
            formats,
            index=False,
            freeze=False,
        )
        _apply_heatmap(matrix_ws, score_meta, score_df)
        matrix_row = score_meta["next_row"]

        weight_meta = _write_df_block(
            writer,
            matrix_name,
            weight_df,
            matrix_row,
            _xl_text(lang, "Kriter Ağırlıklarının Yıllara Göre Değişimi", "Criterion Weights by Period"),
            formats,
            index=False,
            freeze=False,
        )
        _apply_heatmap(matrix_ws, weight_meta, weight_df)
        _write_citation_block(matrix_ws, weight_meta["next_row"], f"{_xl_text(lang, 'Önerilen atıf', 'Suggested citation')}: {_SOFTWARE_CITATION}", formats, end_col=11)
        _write_panel_summary_findings_sheet(
            writer,
            findings_name,
            panel_results,
            summary_df,
            rank_df,
            weight_df,
            lang,
            formats,
        )

    return output.getvalue()

def generate_panel_apa_docx(panel_results: Dict[str, Dict[str, Any]], lang: str = "TR") -> bytes | None:
    if not _ensure_docx_support():
        return None
    doc = Document()
    _configure_apa_doc(doc)
    body_lang = "en-US" if lang == "EN" else "tr-TR"
    for section in doc.sections:
        section.top_margin, section.bottom_margin = Inches(1), Inches(1)
        section.left_margin, section.right_margin = Inches(1), Inches(1)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_text = (
        "Panel Data Academic Analysis Report"
        if lang == "EN"
        else "Panel Veri Akademik Analiz Raporu"
    )
    _set_docx_run_style(title.add_run(title_text), body_lang, bold=True)
    doc.add_paragraph("")

    refs_union: List[str] = []
    seen_refs: set[str] = set()
    heading_map = _docx_heading_map(lang)
    section_order = [
        heading_map["objective"],
        heading_map["scope"],
        heading_map["philosophy"],
        heading_map["tables"],
        heading_map["conclusion"],
    ]
    for year_label, year_result in panel_results.items():
        y_head = doc.add_paragraph()
        _set_docx_run_style(
            y_head.add_run(
                f"Year {year_label}" if lang == "EN" else f"Yıl {year_label}"
            ),
            body_lang,
            bold=True,
        )
        sections = _build_academic_doc_sections(year_result, year_result.get("selected_data", pd.DataFrame()), lang)
        for heading in section_order:
            p = doc.add_paragraph()
            _set_docx_run_style(p.add_run(heading), body_lang, bold=True)
            section_text = str(sections.get(heading, "") or "").strip()
            for block in [b.strip() for b in section_text.split("\n\n") if b.strip()]:
                body = doc.add_paragraph()
                body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                _set_docx_run_style(body.add_run(block), body_lang)
            if heading == heading_map["tables"]:
                preview_title = "Decision Matrix Snapshot" if lang == "EN" else "Karar Matrisi Özeti"
                preview_head = doc.add_paragraph()
                _set_docx_run_style(preview_head.add_run(preview_title), body_lang, bold=True)
                add_table_to_doc(doc, _docx_preview_df(year_result.get("selected_data", pd.DataFrame()), lang), lang_code=body_lang)
                weight_title = "Weight Table" if lang == "EN" else "Ağırlık Tablosu"
                weight_head = doc.add_paragraph()
                _set_docx_run_style(weight_head.add_run(weight_title), body_lang, bold=True)
                add_table_to_doc(doc, localize_df_lang(year_result["weights"]["table"], lang), lang_code=body_lang)
                ranking_table = year_result.get("ranking", {}).get("table")
                if isinstance(ranking_table, pd.DataFrame) and not ranking_table.empty:
                    rank_title = "Ranking Table" if lang == "EN" else "Sıralama Tablosu"
                    rank_head = doc.add_paragraph()
                    _set_docx_run_style(rank_head.add_run(rank_title), body_lang, bold=True)
                    add_table_to_doc(doc, localize_df_lang(ranking_table, lang), lang_code=body_lang)
                detail_info = _preferred_doc_detail_table(year_result, lang)
                if detail_info is not None:
                    detail_title, detail_df = detail_info
                    detail_head = doc.add_paragraph()
                    _set_docx_run_style(detail_head.add_run(detail_title), body_lang, bold=True)
                    add_table_to_doc(doc, localize_df_lang(detail_df, lang), lang_code=body_lang)
                comp_df = (year_result.get("comparison") or {}).get("spearman_matrix")
                if isinstance(comp_df, pd.DataFrame) and not comp_df.empty and comp_df.shape[0] >= 2:
                    comp_title = "Method Comparison Table" if lang == "EN" else "Yöntem Karşılaştırma Tablosu"
                    comp_head = doc.add_paragraph()
                    _set_docx_run_style(comp_head.add_run(comp_title), body_lang, bold=True)
                    add_table_to_doc(doc, localize_df_lang(comp_df, lang), lang_code=body_lang)
                mc_df = (year_result.get("sensitivity") or {}).get("monte_carlo_summary")
                if isinstance(mc_df, pd.DataFrame) and not mc_df.empty:
                    mc_title = "Monte Carlo Summary" if lang == "EN" else "Monte Carlo Özeti"
                    mc_head = doc.add_paragraph()
                    _set_docx_run_style(mc_head.add_run(mc_title), body_lang, bold=True)
                    add_table_to_doc(doc, localize_df_lang(mc_df, lang), lang_code=body_lang)
        for ref in sections.get(heading_map["references"], []):
            if ref not in seen_refs:
                seen_refs.add(ref)
                refs_union.append(ref)
        doc.add_paragraph("")

    for note_line in _reference_notice_lines(lang):
        note_p = doc.add_paragraph()
        _set_docx_run_style(note_p.add_run(note_line), body_lang, italic=True)
    ref_heading = doc.add_paragraph()
    _set_docx_run_style(ref_heading.add_run(heading_map["references"]), body_lang, bold=True)
    for ref in refs_union:
        ref_p = doc.add_paragraph()
        ref_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        ref_p.paragraph_format.left_indent = Inches(0.5)
        ref_p.paragraph_format.first_line_indent = Inches(-0.5)
        _set_docx_run_style(ref_p.add_run(ref), body_lang)

    sig = doc.add_paragraph()
    sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_docx_run_style(sig.add_run("Prof. Dr. Ömer Faruk Rençber"), body_lang, italic=True)
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()

def _download_signature(result: Dict[str, Any] | None, panel_results: Dict[str, Dict[str, Any]] | None) -> tuple:
    study_title = str(st.session_state.get("study_title", "") or "").strip()
    if isinstance(panel_results, dict) and panel_results:
        years = tuple(str(y) for y in panel_results.keys())
        base_time = result.get("analysis_time") if isinstance(result, dict) else None
        return ("panel", years, base_time, study_title)
    if not isinstance(result, dict):
        return ("none", study_title)
    wm = (result.get("weights") or {}).get("method")
    rm = (result.get("ranking") or {}).get("method")
    shape = tuple(result.get("selected_data", pd.DataFrame()).shape)
    return ("single", wm, rm, result.get("analysis_time"), shape, study_title)

def _render_report_download_controls_core(lang: str) -> None:
    result = st.session_state.get("analysis_result")
    panel_results = st.session_state.get("panel_results")
    if not isinstance(result, dict):
        return

    is_panel_download = isinstance(panel_results, dict) and bool(panel_results)
    download_sig = _download_signature(result, panel_results)
    if st.session_state.get("download_blob_sig") != download_sig:
        st.session_state["download_blob_sig"] = download_sig
        st.session_state["download_blob_cache"] = {}

    blob_cache = dict(st.session_state.get("download_blob_cache") or {})
    excel_key = f"excel::{lang}"
    excel_title = (
        _export_study_title(result, lang)
        if not is_panel_download
        else (str(st.session_state.get("study_title", "") or "").strip() or _xl_text(lang, "Panel Analiz Raporu", "Panel Analysis Report"))
    )
    excel_filename = _export_file_name(excel_title, lang, "xlsx")

    if excel_key not in blob_cache:
        try:
            if is_panel_download:
                blob_cache[excel_key] = generate_panel_excel(panel_results, lang=lang)
            else:
                selected_data = result.get("selected_data", pd.DataFrame())
                blob_cache[excel_key] = generate_excel(result, selected_data, lang=lang)
        except Exception as exc:
            st.error(tt("Excel çıktısı oluşturulamadı.", "Excel output could not be created."))
            st.caption(tt(f"Hata kodu: {_safe_error_code(exc)}", f"Error code: {_safe_error_code(exc)}"))
            st.session_state["download_blob_cache"] = blob_cache
            return
    st.session_state["download_blob_cache"] = blob_cache

    doc_key = f"docx::{lang}"
    doc_bytes = None
    docx_enabled = _ensure_docx_support()
    if docx_enabled:
        if doc_key not in blob_cache:
            try:
                selected_data = result.get("selected_data", pd.DataFrame())
                blob_cache[doc_key] = (
                    generate_panel_apa_docx(panel_results, lang=lang)
                    if is_panel_download
                    else generate_apa_docx(result, selected_data, lang=lang)
                )
            except Exception as exc:
                st.error(tt("Word çıktısı oluşturulamadı.", "Word output could not be created."))
                st.caption(tt(f"Hata kodu: {_safe_error_code(exc)}", f"Error code: {_safe_error_code(exc)}"))
                st.session_state["report_docx"] = None
                st.session_state["download_blob_cache"] = blob_cache
                return
        doc_bytes = blob_cache.get(doc_key)
        st.session_state["report_docx"] = doc_bytes
        st.session_state["download_blob_cache"] = blob_cache

    # --- IMRAD article generation (HTML) ---
    imrad_key = f"imrad::{lang}"
    imrad_bytes = None
    if imrad_key not in blob_cache:
        try:
            selected_data = result.get("selected_data", pd.DataFrame())
            _imrad_result = mcdm_article.generate_imrad_docx(result, selected_data, lang=lang)
            if isinstance(_imrad_result, str):
                _imrad_result = _imrad_result.encode("utf-8")
            blob_cache[imrad_key] = _imrad_result
        except Exception as _exc:
            blob_cache[imrad_key] = None
            st.session_state["_imrad_gen_error"] = str(_exc)
    imrad_bytes = blob_cache.get(imrad_key)
    st.session_state["download_blob_cache"] = blob_cache

    dl1, dl2, dl3 = st.columns(3)
    with dl1:
        _excel_clicked = _render_export_download_button(
            tt("📊 Tüm Sonuçları İndir (Excel)", "📊 Download All Results (Excel)"),
            blob_cache[excel_key],
            excel_filename,
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            key=f"download_excel_results_{lang}",
        )
        if _excel_clicked:
            access.track_event("report_downloaded", {"format": "excel", "lang": lang, "is_panel": is_panel_download})
        st.caption(tt("Tablolar, grafikler ve ham veriler. İş raporları için.", "Tables, charts and raw data. For business reports."))

    with dl2:
        if docx_enabled:
            docx_name = tt("MCDM_Akademik_Rapor.docx", "MCDM_Academic_Report.docx")
            _docx_clicked = _render_export_download_button(
                tt("📄 Akademik Rapor — APA Word", "📄 Academic Report — APA Word"),
                doc_bytes,
                docx_name,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key=f"download_docx_results_{lang}",
            )
            if _docx_clicked:
                access.track_event("report_downloaded", {"format": "docx", "lang": lang, "is_panel": is_panel_download})
            st.caption(tt("Yorumlu rapor. Matematiğe hakim olmayan okuyucular için.", "Interpretive report for non-technical readers."))
        else:
            st.warning(tt("Word çıktısı için python-docx kurulu olmalıdır.", "python-docx must be installed for Word output."))

    with dl3:
        if imrad_bytes:
            imrad_name = tt("MCDM_IMRAD_Makale.html", "MCDM_IMRAD_Article.html")
            _imrad_clicked = _render_export_download_button(
                tt("📝 IMRAD Makale (HTML)", "📝 IMRAD Article (HTML)"),
                imrad_bytes,
                imrad_name,
                "text/html",
                key=f"download_imrad_results_{lang}",
            )
            if _imrad_clicked:
                access.track_event("report_downloaded", {"format": "imrad_html", "lang": lang, "is_panel": is_panel_download})
            st.caption(tt("Formüller ve gerekçeler dahil makale taslağı. Tarayıcıda açılır, yazdırılabilir.", "Article draft with formulas. Opens in browser, printable."))
        else:
            _gen_err = st.session_state.get("_imrad_gen_error", "")
            st.caption(tt(f"IMRAD makale taslağı oluşturulamadı. {_gen_err}", f"IMRAD draft could not be generated. {_gen_err}"))

if hasattr(st, "fragment"):
    _render_report_download_controls = st.fragment(_render_report_download_controls_core)
else:
    _render_report_download_controls = _render_report_download_controls_core

def _run_single_analysis_bundle(
    data_slice: pd.DataFrame,
    criteria: List[str],
    criteria_types: Dict[str, str],
    config,
    weight_mode_key: str,
    weight_method: str,
    main_rank: str | None,
) -> Dict[str, Any]:
    config_payload = dict(config.__dict__)
    result = run_full_analysis_cached(data_slice, config_payload)
    run_heavy_robustness = bool(getattr(config, "run_heavy_robustness", False))
    if main_rank:
        actual_method = result.get("ranking", {}).get("method")
        if actual_method != main_rank:
            raise ValueError(tt(f"Yöntem uyumsuzluğu: seçilen '{main_rank}', çalışan '{actual_method}'.", f"Method mismatch: selected '{main_rank}', executed '{actual_method}'."))
        rt_chk = result.get("ranking", {}).get("table")
        if rt_chk is None or rt_chk.empty:
            raise ValueError(tt("Sıralama çıktısı boş döndü.", "Ranking output is empty."))
        if rt_chk["Skor"].isna().any() or (~np.isfinite(rt_chk["Skor"])).any():
            raise ValueError(tt(f"{main_rank} yöntemi geçersiz skor üretti (NaN/Inf).", f"{main_rank} produced invalid scores (NaN/Inf)."))
        if "Sıra" not in rt_chk.columns or rt_chk["Sıra"].isna().any():
            raise ValueError(tt(f"{main_rank} yöntemi sıra kolonunu üretemedi.", f"{main_rank} did not produce a valid rank column."))
    result["selected_data"] = data_slice[criteria].copy()
    if weight_mode_key == "objective" and run_heavy_robustness:
        try:
            result["weight_robustness"] = compute_weight_robustness(
                data=data_slice[criteria].copy(),
                criteria=criteria,
                criteria_types=criteria_types,
                weight_method=weight_method,
                base_weights=result["weights"]["values"],
                bootstrap_n=180,
                fuzzy_spread=float(getattr(config, "fuzzy_spread", 0.10)),
            )
        except Exception as _w_exc:
            result["weight_robustness"] = {
                "error": tt(
                    "Ağırlık sağlamlık testleri üretilemedi.",
                    "Weight robustness tests could not be generated.",
                ),
                "code": _safe_error_code(_w_exc),
            }
    elif weight_mode_key == "objective":
        result["weight_robustness"] = {
            "info": tt(
                "Ağır sağlamlık testleri bu çalışmada kapatıldı.",
                "Full robustness tests are disabled for this run.",
            )
        }
    else:
        result["weight_robustness"] = {
            "info": tt(
                "Sağlamlık testi objektif ağırlık yöntemleri için tasarlanmıştır.",
                "Robustness test is designed for objective weighting methods.",
            )
        }
    return result

_THEME = dict(paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="#F8FAFC", margin=dict(l=20, r=20, t=44, b=20))

def fig_weight_bar(weight_df: pd.DataFrame) -> go.Figure:
    fig = px.bar(weight_df.sort_values("Ağırlık", ascending=False), x="Kriter", y="Ağırlık", text_auto=".4f", color="Ağırlık", color_continuous_scale="Blues")
    fig.update_layout(
        height=400,
        title=tt("Kriter Ağırlıkları", "Criterion Weights"),
        coloraxis_colorbar=dict(title=tt("Ağırlık", "Weight")),
        **_THEME,
    )
    return fig

def fig_rank_bar(rank_df: pd.DataFrame) -> go.Figure:
    fig = px.bar(rank_df.sort_values("Skor", ascending=True), x="Skor", y="Alternatif", orientation="h", text_auto=".4f", color="Skor", color_continuous_scale="Blues")
    fig.update_layout(
        height=450,
        title=tt("Nihai Sıralama Skorları", "Final Ranking Scores"),
        coloraxis_colorbar=dict(title=tt("Skor", "Score")),
        **_THEME,
    )
    return fig

def fig_box_plots(data: pd.DataFrame, criteria: List[str]) -> go.Figure:
    colors = ["#2E7D9E", "#1B365D", "#48BB78", "#ED8936", "#9F7AEA", "#F56565", "#38B2AC", "#ECC94B"]
    fig = go.Figure()
    for i, c in enumerate(criteria):
        hex_c = colors[i % len(colors)]
        r, g, b = int(hex_c[1:3], 16), int(hex_c[3:5], 16), int(hex_c[5:7], 16)
        
        fig.add_trace(go.Box(
            y=data[c], name=c, boxpoints="all", jitter=0.35, pointpos=-1.6,
            marker=dict(size=5, color=hex_c, opacity=0.65),
            line=dict(color=hex_c),
            fillcolor=f"rgba({r},{g},{b},0.2)",
        ))
    fig.update_layout(height=420, title=tt("Kriter Dağılım Profilleri (Kutu + Nokta)", "Criterion Distribution Profiles (Box + Points)"), showlegend=False, **_THEME)
    return fig

def fig_weight_radar(weight_df: pd.DataFrame) -> go.Figure:
    cats = weight_df["Kriter"].tolist()
    vals = weight_df["Ağırlık"].tolist()
    fig = go.Figure(go.Scatterpolar(
        r=vals + [vals[0]], theta=cats + [cats[0]],
        fill="toself", fillcolor="rgba(27,54,93,0.12)",
        line=dict(color="#1B365D", width=2),
        marker=dict(size=7, color="#2E7D9E"),
        name=tt("Ağırlık", "Weight"),
    ))
    fig.update_layout(
        height=420, title=tt("Ağırlık Dağılımı — Radar Görünümü", "Weight Distribution — Radar View"),
        polar=dict(bgcolor="#F8FAFC", radialaxis=dict(visible=True, gridcolor="#E2E8F0"), angularaxis=dict(gridcolor="#E2E8F0")),
        **_THEME,
    )
    return fig

def fig_parallel_coords(ranking_table: pd.DataFrame, contrib_table: pd.DataFrame, criteria: List[str]) -> go.Figure:
    merged = ranking_table.merge(contrib_table, on="Alternatif", how="inner")
    avail = [c for c in criteria if c in merged.columns]
    if not avail:
        return go.Figure()
    dims = [dict(label=c, values=merged[c].tolist()) for c in avail]
    dims.append(dict(label=tt("Skor", "Score"), values=merged["Skor"].tolist()))
    fig = go.Figure(go.Parcoords(
        line=dict(color=merged["Skor"].tolist(), colorscale="Blues", showscale=True, colorbar=dict(title=tt("Skor", "Score"))),
        dimensions=dims,
    ))
    fig.update_layout(height=460, title=tt("Paralel Koordinat — Alternatif Profilleri", "Parallel Coordinates — Alternative Profiles"), **_THEME)
    return fig

def fig_network_alternatives(ranking_table: pd.DataFrame) -> go.Figure:
    n = len(ranking_table)
    if n == 0:
        return go.Figure()
    alts   = ranking_table["Alternatif"].tolist()
    scores = ranking_table["Skor"].tolist()
    ranks  = ranking_table["Sıra"].tolist() if "Sıra" in ranking_table.columns else list(range(1, n + 1))

    angles = [2 * np.pi * i / n for i in range(n)]
    xs = [np.cos(a) for a in angles]
    ys = [np.sin(a) for a in angles]

    cx, cy = 0.0, 0.0
    edge_x, edge_y = [], []
    for i in range(n):
        edge_x += [xs[i], cx, None]
        edge_y += [ys[i], cy, None]

    for i in range(n):
        j = (i + 1) % n
        edge_x += [xs[i], xs[j], None]
        edge_y += [ys[i], ys[j], None]

    fig = go.Figure()
    fig.add_trace(go.Scatter(
        x=edge_x, y=edge_y, mode="lines",
        line=dict(color="#CBD5E0", width=1.0), hoverinfo="none", showlegend=False,
    ))
    fig.add_trace(go.Scatter(
        x=[cx], y=[cy], mode="markers",
        marker=dict(size=18, color="#1B365D", symbol="star"),
        text=[tt("En Yüksek", "Highest")], hoverinfo="text", showlegend=False,
    ))
    node_sizes = [max(20, 55 - (r - 1) * 3) for r in ranks]
    fig.add_trace(go.Scatter(
        x=xs, y=ys, mode="markers+text",
        marker=dict(
            size=node_sizes,
            color=scores, colorscale="Blues", showscale=True,
            colorbar=dict(title=tt("Skor", "Score"), thickness=12),
            line=dict(color="#1B365D", width=1.5),
        ),
        text=alts, textposition="top center",
        hovertemplate=f"<b>%{{text}}</b><br>{tt('Skor', 'Score')}: %{{marker.color:.4f}}<extra></extra>",
        showlegend=False,
    ))
    fig.update_layout(
        height=520, title=tt("Alternatif Ağ Diyagramı (Dairesel Yerleşim)", "Alternative Network Diagram (Circular Layout)"),
        xaxis=dict(showgrid=False, zeroline=False, showticklabels=False, range=[-1.5, 1.5]),
        yaxis=dict(showgrid=False, zeroline=False, showticklabels=False, range=[-1.5, 1.5]),
        **_THEME,
    )
    return fig

def fig_mc_stability_bubble(mc_df: pd.DataFrame) -> go.Figure:
    sdf = mc_df.sort_values("BirincilikOranı", ascending=False).copy()
    sdf["Yüzde"] = sdf["BirincilikOranı"] * 100
    fig = go.Figure(go.Scatter(
        x=sdf["OrtalamaSıra"], y=sdf["Yüzde"],
        mode="markers+text",
        text=sdf["Alternatif"],
        textposition="top center",
        marker=dict(
            size=sdf["Yüzde"].clip(lower=5) * 1.2,
            color=sdf["Yüzde"],
            colorscale="Blues", showscale=True,
            colorbar=dict(title=tt("Birincilik %", "First-place %")),
            line=dict(color="#1B365D", width=1),
        ),
        hovertemplate=f"<b>%{{text}}</b><br>{tt('Ort. Sıra', 'Mean Rank')}: %{{x:.2f}}<br>{tt('Birincilik', 'First-place')}: %{{y:.1f}}%<extra></extra>",
    ))
    fig.update_layout(
        height=460, title=tt("Monte Carlo — Birincilik Oranı vs Ortalama Sıra (Balon)", "Monte Carlo — First-Place Rate vs Mean Rank (Bubble)"),
        xaxis_title=tt("Ortalama Sıra (düşük = daha iyi)", "Mean Rank (lower = better)"),
        yaxis_title=tt("Birincilik Oranı (%)", "First-Place Rate (%)"),
        **_THEME,
    )
    fig.update_xaxes(autorange="reversed")
    return fig

def fig_mc_rank_bar(mc_df: pd.DataFrame) -> go.Figure:
    sdf = mc_df.sort_values("BirincilikOranı", ascending=False)
    colors_mc = ["#1B365D" if v >= 0.5 else "#2E7D9E" if v >= 0.2 else "#90CDF4" for v in sdf["BirincilikOranı"]]
    fig = go.Figure(go.Bar(
        x=sdf["Alternatif"], y=sdf["BirincilikOranı"] * 100,
        text=[f"%{v*100:.1f}" for v in sdf["BirincilikOranı"]],
        textposition="outside",
        marker_color=colors_mc,
    ))
    fig.update_layout(
        height=420, title=tt("Monte Carlo — Birincilik Oranı (%)", "Monte Carlo — First-Place Rate (%)"),
        yaxis_title=tt("Birincilik Oranı (%)", "First-Place Rate (%)"), xaxis_title="",
        **_THEME,
    )
    return fig

def fig_local_sensitivity(local_df: pd.DataFrame) -> go.Figure:
    cat_order = ["-20%", "-10%", "+10%", "+20%"]
    local_df = local_df.copy()
    crit_col = "Kriter"
    wchg_col = "AğırlıkDeğişimi"
    if st.session_state.get("ui_lang") == "EN":
        local_df = local_df.rename(columns={"Kriter": "Criterion", "AğırlıkDeğişimi": "WeightChange"})
        crit_col = "Criterion"
        wchg_col = "WeightChange"
    local_df[wchg_col] = pd.Categorical(local_df[wchg_col], categories=cat_order, ordered=True)
    fig = px.line(
        local_df.sort_values([crit_col, wchg_col]),
        x=wchg_col, y="SpearmanRho", color=crit_col,
        markers=True, title=tt("Lokal Duyarlılık — Ağırlık Değişimine Tepki", "Local Sensitivity — Response to Weight Changes"),
    )
    fig.update_traces(line_width=2, marker_size=8)
    fig.update_layout(height=420, **_THEME)
    return fig

def _map_alt_names(values, alt_names: Dict[str, str] | None) -> List[str]:
    if not alt_names:
        return [str(v) for v in values]
    return [alt_names.get(str(v), str(v)) for v in values]

def _map_alt_names_in_df(df: pd.DataFrame, alt_names: Dict[str, str] | None) -> pd.DataFrame:
    out = df.copy()
    if not alt_names:
        return out
    for col in ["Alternatif", "Alternative", "BirinciAlternatif", "TopAlternative"]:
        if col in out.columns:
            out[col] = out[col].astype(str).map(lambda x: alt_names.get(x, x))
    return out

def _map_alt_names_in_matrix(df: pd.DataFrame, alt_names: Dict[str, str] | None) -> pd.DataFrame:
    out = df.copy()
    if not alt_names:
        return out
    out.index = [alt_names.get(str(v), str(v)) for v in out.index]
    out.columns = [alt_names.get(str(v), str(v)) for v in out.columns]
    return out

def fig_topsis_distance_scatter(distance_df: pd.DataFrame, alt_names: Dict[str, str] | None = None) -> go.Figure:
    sdf = distance_df.copy().sort_values("Skor", ascending=False)
    sdf["AltLabel"] = _map_alt_names(sdf["Alternatif"].tolist(), alt_names)
    fig = px.scatter(
        sdf,
        x="D+",
        y="D-",
        text="AltLabel",
        color="Skor",
        color_continuous_scale="Blues",
        title=tt("TOPSIS Uzaklık Haritası", "TOPSIS Distance Map"),
    )
    fig.update_traces(marker=dict(size=13, line=dict(color="#1B365D", width=1.2)), textposition="top center")
    fig.update_layout(
        height=440,
        xaxis_title=tt("İdeal Çözüme Uzaklık D+ (düşük daha iyi)", "Distance to Ideal D+ (lower is better)"),
        yaxis_title=tt("Negatif İdeale Uzaklık D- (yüksek daha iyi)", "Distance to Negative Ideal D- (higher is better)"),
        **_THEME,
    )
    fig.update_xaxes(autorange="reversed")
    return fig

def fig_vikor_components(vikor_df: pd.DataFrame, alt_names: Dict[str, str] | None = None) -> go.Figure:
    sdf = vikor_df.copy().sort_values("Q", ascending=True)
    labels = _map_alt_names(sdf["Alternatif"].tolist(), alt_names)
    fig = go.Figure()
    fig.add_trace(go.Bar(name="S", x=labels, y=sdf["S"], marker_color="#7AA6D1"))
    fig.add_trace(go.Bar(name="R", x=labels, y=sdf["R"], marker_color="#315D8A"))
    fig.add_trace(go.Bar(name="Q", x=labels, y=sdf["Q"], marker_color="#17324D"))
    fig.update_layout(
        barmode="group",
        height=440,
        title=tt("VIKOR Bileşen Dengesi (S, R, Q)", "VIKOR Component Balance (S, R, Q)"),
        xaxis_title="",
        yaxis_title=tt("Bileşen Değeri", "Component Value"),
        **_THEME,
    )
    return fig

def fig_edas_balance(edas_df: pd.DataFrame, alt_names: Dict[str, str] | None = None) -> go.Figure:
    sdf = edas_df.copy().sort_values("Skor", ascending=False)
    labels = _map_alt_names(sdf["Alternatif"].tolist(), alt_names)
    fig = go.Figure()
    fig.add_trace(go.Bar(name="NSP", x=labels, y=sdf["NSP"], marker_color="#2E7D9E"))
    fig.add_trace(go.Bar(name="NSN", x=labels, y=sdf["NSN"], marker_color="#1B365D"))
    fig.update_layout(
        barmode="group",
        height=440,
        title=tt("EDAS Pozitif/Negatif Sapma Dengesi", "EDAS Positive/Negative Distance Balance"),
        xaxis_title="",
        yaxis_title=tt("Normalize Bileşen", "Normalized Component"),
        **_THEME,
    )
    return fig

def fig_codas_distance_map(codas_df: pd.DataFrame, alt_names: Dict[str, str] | None = None) -> go.Figure:
    sdf = codas_df.copy().sort_values("H", ascending=False)
    sdf["AltLabel"] = _map_alt_names(sdf["Alternatif"].tolist(), alt_names)
    fig = px.scatter(
        sdf,
        x="E",
        y="T",
        text="AltLabel",
        color="H",
        size=np.abs(sdf["H"]).clip(lower=0.05),
        color_continuous_scale="Blues",
        title=tt("CODAS Negatif İdeale Uzaklık Haritası", "CODAS Distance-from-Negative-Ideal Map"),
    )
    fig.update_traces(marker=dict(line=dict(color="#1B365D", width=1.1)), textposition="top center")
    fig.update_layout(
        height=440,
        xaxis_title=tt("Öklid Uzaklığı E", "Euclidean Distance E"),
        yaxis_title=tt("Manhattan Uzaklığı T", "Taxicab Distance T"),
        **_THEME,
    )
    return fig

def fig_promethee_flows(prom_df: pd.DataFrame, alt_names: Dict[str, str] | None = None) -> go.Figure:
    sdf = prom_df.copy().sort_values("PhiNet", ascending=False)
    labels = _map_alt_names(sdf["Alternatif"].tolist(), alt_names)
    fig = go.Figure()
    fig.add_trace(go.Bar(name="Phi+", x=labels, y=sdf["PhiPlus"], marker_color="#2E7D9E"))
    fig.add_trace(go.Bar(name="Phi-", x=labels, y=sdf["PhiMinus"], marker_color="#9BB9D6"))
    fig.add_trace(go.Scatter(name="PhiNet", x=labels, y=sdf["PhiNet"], mode="lines+markers", line=dict(color="#10283F", width=3), marker=dict(size=8)))
    fig.update_layout(
        barmode="group",
        height=460,
        title=tt("PROMETHEE Akış Ayrışımı", "PROMETHEE Flow Decomposition"),
        xaxis_title="",
        yaxis_title=tt("Akış Değeri", "Flow Value"),
        **_THEME,
    )
    return fig

def fig_promethee_gaia(
    gaia_alt_df: pd.DataFrame,
    gaia_crit_df: pd.DataFrame | None = None,
    gaia_decision_df: pd.DataFrame | None = None,
    alt_names: Dict[str, str] | None = None,
) -> go.Figure:
    fig = go.Figure()
    if not isinstance(gaia_alt_df, pd.DataFrame) or gaia_alt_df.empty:
        fig.update_layout(height=520, title=tt("PROMETHEE GAIA Düzlemi", "PROMETHEE GAIA Plane"), **_THEME)
        return fig

    sdf = gaia_alt_df.copy()
    alt_col = col_key(sdf, "Alternatif", "Alternative")
    if alt_col not in sdf.columns or "GAIA1" not in sdf.columns or "GAIA2" not in sdf.columns:
        fig.update_layout(height=520, title=tt("PROMETHEE GAIA Düzlemi", "PROMETHEE GAIA Plane"), **_THEME)
        return fig

    sdf["AltLabel"] = _map_alt_names(sdf[alt_col].astype(str).tolist(), alt_names)
    color_col = "PhiNet" if "PhiNet" in sdf.columns else ("Skor" if "Skor" in sdf.columns else None)
    marker_kwargs: Dict[str, Any] = dict(size=12, line=dict(color="#0B2239", width=1.0), symbol="circle")
    if color_col:
        marker_kwargs.update(dict(color=sdf[color_col], colorscale="RdBu", cmid=0, colorbar=dict(title=color_col)))
    else:
        marker_kwargs.update(dict(color="#2E7D9E"))

    fig.add_trace(
        go.Scatter(
            x=sdf["GAIA1"],
            y=sdf["GAIA2"],
            mode="markers+text",
            text=sdf["AltLabel"],
            textposition="top center",
            marker=marker_kwargs,
            name=tt("Alternatifler", "Alternatives"),
        )
    )

    max_abs = float(np.nanmax(np.abs(np.r_[sdf["GAIA1"].to_numpy(dtype=float), sdf["GAIA2"].to_numpy(dtype=float)]))) if len(sdf) else 0.0

    if isinstance(gaia_crit_df, pd.DataFrame) and not gaia_crit_df.empty and {"GAIA1", "GAIA2"}.issubset(gaia_crit_df.columns):
        cdf = gaia_crit_df.copy()
        crit_col = col_key(cdf, "Kriter", "Criterion")
        if crit_col not in cdf.columns:
            crit_col = cdf.columns[0]
        for _, row in cdf.iterrows():
            cx = float(pd.to_numeric(row.get("GAIA1"), errors="coerce"))
            cy = float(pd.to_numeric(row.get("GAIA2"), errors="coerce"))
            if not (np.isfinite(cx) and np.isfinite(cy)):
                continue
            cname = str(row.get(crit_col, "C"))
            fig.add_trace(
                go.Scatter(
                    x=[0.0, cx],
                    y=[0.0, cy],
                    mode="lines",
                    line=dict(color="#1B365D", width=1.6, dash="dot"),
                    showlegend=False,
                    hoverinfo="skip",
                )
            )
            fig.add_trace(
                go.Scatter(
                    x=[cx],
                    y=[cy],
                    mode="text",
                    text=[cname],
                    textposition="top center",
                    textfont=dict(color="#1B365D", size=11),
                    showlegend=False,
                    hovertemplate=f"{cname}<extra></extra>",
                )
            )
            max_abs = max(max_abs, abs(cx), abs(cy))

    if isinstance(gaia_decision_df, pd.DataFrame) and not gaia_decision_df.empty and {"DeltaX", "DeltaY"}.issubset(gaia_decision_df.columns):
        dx = float(pd.to_numeric(gaia_decision_df.iloc[0]["DeltaX"], errors="coerce"))
        dy = float(pd.to_numeric(gaia_decision_df.iloc[0]["DeltaY"], errors="coerce"))
        if np.isfinite(dx) and np.isfinite(dy):
            fig.add_trace(
                go.Scatter(
                    x=[0.0, dx],
                    y=[0.0, dy],
                    mode="lines+text",
                    text=["", tt("Karar Ekseni (Δ)", "Decision Axis (Δ)")],
                    textposition="top right",
                    line=dict(color="#B22222", width=3),
                    name=tt("Karar Ekseni", "Decision Axis"),
                )
            )
            max_abs = max(max_abs, abs(dx), abs(dy))

    lim = max(0.25, max_abs * 1.25)
    fig.update_xaxes(range=[-lim, lim], zeroline=True, zerolinewidth=1.0, zerolinecolor="#A8B7C7", title=tt("GAIA 1", "GAIA 1"), scaleanchor="y", scaleratio=1)
    fig.update_yaxes(range=[-lim, lim], zeroline=True, zerolinewidth=1.0, zerolinecolor="#A8B7C7", title=tt("GAIA 2", "GAIA 2"))
    fig.update_layout(height=520, title=tt("PROMETHEE GAIA Düzlemi", "PROMETHEE GAIA Plane"), **_THEME)
    return fig

def fig_preference_heatmap(pref_df: pd.DataFrame, alt_names: Dict[str, str] | None = None, title: str | None = None) -> go.Figure:
    disp = _map_alt_names_in_matrix(pref_df, alt_names)
    fig = px.imshow(
        disp,
        text_auto=".2f",
        aspect="auto",
        color_continuous_scale="Blues",
        zmin=0,
        zmax=1,
        title=title or tt("Tercih Matrisi", "Preference Matrix"),
    )
    fig.update_layout(height=max(420, len(disp) * 28 + 120), **_THEME)
    return fig

def _live_detail_commentary(method: str, detail_df: pd.DataFrame, alt_names: Dict[str, str], lang: str) -> str:
    """Tablodaki gerçek değerleri okuyarak yönteme özgü dinamik yorum üretir."""
    if detail_df is None or not isinstance(detail_df, pd.DataFrame) or detail_df.empty:
        return ""
    is_en = (lang == "EN")
    df = detail_df.copy()
    alt_col  = next((c for c in df.columns if str(c).lower() in ("alternatif", "alternative")), None)
    skor_col = next((c for c in df.columns if str(c).lower() in ("skor", "score")), None)
    if skor_col is None:
        return ""
    if alt_col and alt_names:
        df[alt_col] = df[alt_col].astype(str).map(lambda x: alt_names.get(x, x))
    scores = pd.to_numeric(df[skor_col], errors="coerce")
    if scores.isna().all():
        return ""
    best_idx  = int(scores.idxmax())
    worst_idx = int(scores.idxmin())
    best_name  = str(df.loc[best_idx,  alt_col]) if alt_col else str(best_idx)
    worst_name = str(df.loc[worst_idx, alt_col]) if alt_col else str(worst_idx)
    best_name = _safe_html_text(best_name)
    worst_name = _safe_html_text(worst_name)
    best_score  = float(scores.loc[best_idx])
    worst_score = float(scores.loc[worst_idx])
    spread = best_score - worst_score
    n = len(df)
    method_safe = _safe_html_text(method)

    def _v(col):
        """Satır best_idx için o sütundaki yuvarlı değeri döndür."""
        if col not in df.columns:
            return None
        return round(float(pd.to_numeric(df.loc[best_idx, col], errors="coerce")), 4)

    def _vw(col):
        """Satır worst_idx için yuvarlı değer."""
        if col not in df.columns:
            return None
        return round(float(pd.to_numeric(df.loc[worst_idx, col], errors="coerce")), 4)

    def _col(name):
        """Sütun ismi tam veya kısmi eşleşme."""
        for c in df.columns:
            if str(c) == name:
                return c
        return None

    base = method.replace("Fuzzy ", "")
    fz   = ("Fuzzy" in method)

    # ── TOPSIS ──────────────────────────────────────────────────────────────────
    if base == "TOPSIS":
        dp, dm = _v("D+"), _v("D-")
        wp, wm = _vw("D+"), _vw("D-")
        if is_en:
            calc = ("TOPSIS computes two Euclidean distances for each alternative: <b>D⁺</b> (distance to the ideal solution) "
                    "and <b>D⁻</b> (distance to the negative-ideal). "
                    "The closeness coefficient <b>C* = D⁻ / (D⁺ + D⁻)</b> is the final score, ranging from 0 to 1.")
            found = (f"<b>{best_name}</b> achieved the highest C* = <b>{best_score:.4f}</b>"
                     + (f" (D⁺={dp}, D⁻={dm})" if dp is not None else "") + ". "
                     + f"<b>{worst_name}</b> scored lowest at C* = <b>{worst_score:.4f}</b>"
                     + (f" (D⁺={wp}, D⁻={wm})" if wp is not None else "") + ". "
                     + f"Score spread across {n} alternatives: <b>{spread:.4f}</b> — "
                     + ("the field is tightly packed; small differences matter." if spread < 0.15 else
                        "alternatives are well-differentiated." if spread > 0.35 else
                        "moderate differentiation between alternatives."))
            watch = ("Watch: a high C* driven by a very low D⁺ alone (D⁻ also low) may indicate the alternative is mediocre overall. "
                     "A robust leader has both low D⁺ <em>and</em> high D⁻ simultaneously." if fz is False else
                     "In the fuzzy extension, C* reflects triangular fuzzy distances; interpret spread relative to the crisp version.")
        else:
            calc = ("TOPSIS her alternatif için iki Öklid uzaklığı hesaplar: <b>D⁺</b> (ideal çözüme uzaklık) "
                    "ve <b>D⁻</b> (negatif-ideal çözüme uzaklık). "
                    "Kapanım katsayısı <b>C* = D⁻ / (D⁺ + D⁻)</b> 0 ile 1 arasında değer alır ve nihai skoru verir.")
            found = (f"<b>{best_name}</b> en yüksek C* = <b>{best_score:.4f}</b> değerini elde etti"
                     + (f" (D⁺={dp}, D⁻={dm})" if dp is not None else "") + ". "
                     + f"<b>{worst_name}</b> en düşük C* = <b>{worst_score:.4f}</b> ile son sıradadır"
                     + (f" (D⁺={wp}, D⁻={wm})" if wp is not None else "") + ". "
                     + f"{n} alternatif arasındaki skor aralığı <b>{spread:.4f}</b> — "
                     + ("alternatifler birbirine çok yakın; küçük farklılıklar önemlidir." if spread < 0.15 else
                        "alternatifler belirgin şekilde ayrışmaktadır." if spread > 0.35 else
                        "orta düzeyde ayrışma gözlemlenmektedir."))
            watch = ("Dikkat: yüksek C* yalnızca düşük D⁺'dan kaynaklanıyorsa ve D⁻ de düşükse, "
                     "alternatif genel olarak orta performanslı olabilir. "
                     "Güçlü bir lider hem düşük D⁺ <em>hem de</em> yüksek D⁻ değerini birlikte taşır.")

    # ── VIKOR ────────────────────────────────────────────────────────────────────
    elif base == "VIKOR":
        s_v, r_v, q_v = _v("S"), _v("R"), _v("Q")
        if is_en:
            calc = ("<b>S</b> (group utility): weighted sum of normalized deviations from the ideal — lower is better. "
                    "<b>R</b> (individual regret): maximum weighted deviation for the worst-performing criterion — lower is better. "
                    "<b>Q</b> (compromise score): combines S and R via λ (balance weight, default 0.5) to rank alternatives. "
                    "VIKOR declares a stable compromise when the top-ranked alternative satisfies both the <em>advantage</em> "
                    "condition (ΔQ ≥ 1/(n−1)) and the <em>stability</em> condition.")
            found = (f"<b>{best_name}</b> leads with Q = <b>{q_v if q_v is not None else best_score:.4f}</b>"
                     + (f", S = {s_v}, R = {r_v}" if s_v is not None else "") + ". "
                     + f"<b>{worst_name}</b> has the highest compromise cost (Q = <b>{_vw('Q') if _col('Q') else worst_score:.4f}</b>). "
                     + ("ΔQ between 1st and 2nd is narrow — verify the advantage condition before declaring a stable compromise."
                        if spread < 1.0 / max(n - 1, 1) + 0.05 else
                        "The advantage condition appears satisfied; the leading alternative is a stable compromise."))
            watch = ("S measures group perspective, R measures the worst-case single criterion. "
                     "If R of the leader is high despite low Q, the ranking depends heavily on the λ parameter — consider sensitivity analysis.")
        else:
            calc = ("<b>S</b> (grup faydası): ağırlıklı normalize sapmaların toplamı — düşük daha iyidir. "
                    "<b>R</b> (bireysel pişmanlık): en kötü kriterteki maksimum ağırlıklı sapma — düşük daha iyidir. "
                    "<b>Q</b> (uzlaşı skoru): S ve R, λ parametresiyle (varsayılan 0.5) birleştirilerek alternatifleri sıralar. "
                    "VIKOR, birinci alternatifin hem <em>avantaj</em> koşulunu (ΔQ ≥ 1/(n−1)) "
                    "hem de <em>istikrar</em> koşulunu sağlaması durumunda kararlı uzlaşı ilan eder.")
            found = (f"<b>{best_name}</b> en düşük Q = <b>{q_v if q_v is not None else best_score:.4f}</b> ile öne çıkıyor"
                     + (f" (S={s_v}, R={r_v})" if s_v is not None else "") + ". "
                     + f"<b>{worst_name}</b> en yüksek uzlaşı maliyetine sahip (Q = <b>{_vw('Q') if _col('Q') else worst_score:.4f}</b>). "
                     + ("1. ve 2. sıra arasındaki ΔQ dar — kararlı uzlaşı ilan etmeden önce avantaj koşulunu doğrulayın."
                        if spread < 1.0 / max(n - 1, 1) + 0.05 else
                        "Avantaj koşulu sağlanıyor görünüyor; lider alternatif kararlı bir uzlaşıyı temsil ediyor."))
            watch = ("S grup bakış açısını, R ise en kötü kriterdeki tek taraflı başarısızlığı ölçer. "
                     "Lider düşük Q'ya rağmen yüksek R taşıyorsa, sıralama λ parametresine duyarlı demektir — duyarlılık analizi önerilebilir.")

    # ── EDAS ─────────────────────────────────────────────────────────────────────
    elif base == "EDAS":
        sp, sn = _v("SP"), _v("SN")
        nsp, nsn = _v("NSP"), _v("NSN")
        if is_en:
            calc = ("EDAS measures distances from the <em>average</em> solution (not ideal/anti-ideal). "
                    "<b>SP</b> = sum of positive deviations weighted — shows how much an alternative exceeds the average. "
                    "<b>SN</b> = sum of negative deviations weighted — shows how much it falls below the average. "
                    "<b>NSP</b> and <b>NSN</b> are normalised to [0,1]; the appraisal score <b>AS = 0.5·(NSP + NSN)</b>.")
            found = (f"<b>{best_name}</b> achieved AS = <b>{best_score:.4f}</b>"
                     + (f" (NSP={nsp}, NSN={nsn})" if nsp is not None else "")
                     + (f" with SP={sp}, SN={sn}" if sp is not None else "") + ". "
                     + ("NSP close to 1 signals strong above-average performance; NSN close to 1 means below-average deviations are minimal. "
                        if nsp is not None and nsp > 0.8 else
                        "Check whether leadership is driven by SP or by NSN — a balanced profile is more robust. ")
                     + f"<b>{worst_name}</b> scored AS = <b>{worst_score:.4f}</b>.")
            watch = ("An alternative with high NSP but low NSN may rank well despite significant negative deviations in some criteria — "
                     "inspect the raw PDA/NDA matrices for hidden weaknesses.")
        else:
            calc = ("EDAS alternatifleri <em>ortalama</em> çözüme (ideal/anti-ideal değil) göre değerlendirir. "
                    "<b>SP</b> = ağırlıklı pozitif sapmalar toplamı — alternatifin ortalamanın ne kadar üstünde kaldığını gösterir. "
                    "<b>SN</b> = ağırlıklı negatif sapmalar toplamı — ortalamanın altına ne kadar düştüğünü gösterir. "
                    "<b>NSP</b> ve <b>NSN</b> [0,1]'e normalize edilir; değerlendirme skoru <b>AS = 0,5·(NSP + NSN)</b>.")
            found = (f"<b>{best_name}</b> AS = <b>{best_score:.4f}</b> ile öne çıkıyor"
                     + (f" (NSP={nsp}, NSN={nsn})" if nsp is not None else "")
                     + (f", SP={sp}, SN={sn}" if sp is not None else "") + ". "
                     + ("NSP 1'e yakın: alternatifte ortalamanın belirgin üstünde bir performans var. "
                        if nsp is not None and nsp > 0.8 else
                        "Liderliğin SP mi yoksa NSN'den mi kaynaklandığını kontrol edin — dengeli profil daha güvenilirdir. ")
                     + f"<b>{worst_name}</b> AS = <b>{worst_score:.4f}</b> ile son sırada yer alıyor.")
            watch = ("Yüksek NSP ama düşük NSN: bazı kriterlerde ciddi negatif sapmalar gizlenmiş olabilir. "
                     "Ham PDA/NDA matrislerini inceleyerek gizli zayıflıkları ortaya çıkarabilirsiniz.")

    # ── CODAS ────────────────────────────────────────────────────────────────────
    elif base == "CODAS":
        e_v, t_v, h_v = _v("E"), _v("T"), _v("H")
        if is_en:
            calc = ("CODAS ranks alternatives by their distance from the <em>negative-ideal</em> (worst) solution — "
                    "greater distance means better performance. "
                    "<b>E</b> = Euclidean distance (primary); <b>T</b> = Taxicab (Manhattan) distance (tie-breaker, active when |E_i − E_j| < ψ = 0.02). "
                    "The assessment score <b>H</b> is derived from a pairwise comparison matrix.")
            found = (f"<b>{best_name}</b> leads with H = <b>{h_v if h_v is not None else best_score:.4f}</b>"
                     + (f" (E={e_v}, T={t_v})" if e_v is not None else "") + ". "
                     + ("Both Euclidean and Taxicab distances are high, confirming robust separation from the worst. "
                        if e_v is not None and t_v is not None and e_v > 0.3 and t_v > 0.3 else
                        "Check T — if T is low relative to E, the leader's advantage narrows under the Taxicab criterion. ")
                     + f"Score spread: <b>{spread:.4f}</b>.")
            watch = ("The ψ = 0.02 threshold is the boundary between Euclidean and Taxicab tiebreaking. "
                     "When E distances are very close, T dominates — pairs near this threshold deserve attention.")
        else:
            calc = ("CODAS alternatifleri <em>negatif-ideal</em> (en kötü) çözümden uzaklıklarına göre sıralar — "
                    "daha uzak olan daha iyidir. "
                    "<b>E</b> = Öklid uzaklığı (birincil ölçüt); <b>T</b> = Manhattan uzaklığı (eşit durumda devreye girer, eşik ψ = 0.02). "
                    "Değerlendirme skoru <b>H</b>, ikili karşılaştırma matrisinden türetilir.")
            found = (f"<b>{best_name}</b> H = <b>{h_v if h_v is not None else best_score:.4f}</b> ile lider"
                     + (f" (E={e_v}, T={t_v})" if e_v is not None else "") + ". "
                     + ("Hem Öklid hem Manhattan uzaklıkları yüksek — negatif idealden güçlü ve tutarlı ayrışma var. "
                        if e_v is not None and t_v is not None and e_v > 0.3 and t_v > 0.3 else
                        "T'yi kontrol edin — E'ye göre düşükse, lider Manhattan ölçütü altında daha az avantajlıdır. ")
                     + f"Skor aralığı: <b>{spread:.4f}</b>.")
            watch = ("ψ = 0.02 eşiği, Öklid ve Manhattan geçiş sınırıdır. "
                     "E mesafeleri birbirine çok yakın olan alternatif çiftlerinde T belirleyici olur — bu çiftleri ayrıca inceleyin.")

    # ── PROMETHEE ────────────────────────────────────────────────────────────────
    elif base == "PROMETHEE":
        pp, pm, pn = _v("PhiPlus"), _v("PhiMinus"), _v("PhiNet")
        if is_en:
            calc = ("PROMETHEE II builds a complete ranking via pairwise comparisons using a preference function H(d). "
                    "<b>Φ⁺</b> (positive flow): average preference of this alternative over all others — outranking strength. "
                    "<b>Φ⁻</b> (negative flow): average preference of all others over this alternative — vulnerability. "
                    "<b>Φnet = Φ⁺ − Φ⁻</b> is the net flow and the ranking criterion.")
            found = (f"<b>{best_name}</b> leads with Φnet = <b>{pn if pn is not None else best_score:.4f}</b>"
                     + (f" (Φ⁺={pp}, Φ⁻={pm})" if pp is not None else "") + ". "
                     + ("High Φ⁺ and low Φ⁻ together confirm dominance from both sides. "
                        if pp is not None and pm is not None and pp > 0.5 and pm < 0.3 else
                        "Note whether Φnet is driven by high Φ⁺ (broad outranking) or low Φ⁻ (few losses). ")
                     + f"<b>{worst_name}</b> has Φnet = <b>{_vw('PhiNet') if _col('PhiNet') else worst_score:.4f}</b>.")
            watch = ("Φ⁺ answers 'how much does this alt beat others?'; Φ⁻ answers 'how much is it beaten?'. "
                     "A leader with moderate Φ⁺ but very low Φ⁻ is a 'safe' consensus candidate rather than a dominator.")
        else:
            calc = ("PROMETHEE II ikili karşılaştırmalar ve tercih fonksiyonu H(d) ile tam sıralama oluşturur. "
                    "<b>Φ⁺</b> (pozitif akış): bu alternatifin diğerlerini ne kadar güçlü geçtiğinin ortalaması. "
                    "<b>Φ⁻</b> (negatif akış): diğerlerin bu alternatifi ne kadar güçlü geçtiğinin ortalaması. "
                    "<b>Φnet = Φ⁺ − Φ⁻</b> net akış değeri ve sıralama ölçütüdür.")
            found = (f"<b>{best_name}</b> Φnet = <b>{pn if pn is not None else best_score:.4f}</b> ile öne çıkıyor"
                     + (f" (Φ⁺={pp}, Φ⁻={pm})" if pp is not None else "") + ". "
                     + ("Yüksek Φ⁺ ve düşük Φ⁻ birlikte: lider hem saldırı hem savunma bakımından güçlü. "
                        if pp is not None and pm is not None and pp > 0.5 and pm < 0.3 else
                        "Φnet'in yüksek Φ⁺'tan mı yoksa düşük Φ⁻'dan mı geldiğini inceleyin. ")
                     + f"<b>{worst_name}</b> Φnet = <b>{_vw('PhiNet') if _col('PhiNet') else worst_score:.4f}</b>.")
            watch = ("Φ⁺ 'bu alternatif diğerlerini ne kadar geçiyor?' sorusunu, Φ⁻ ise 'diğerleri bunu ne kadar geçiyor?' sorusunu yanıtlar. "
                     "Orta Φ⁺ ama çok düşük Φ⁻: bu alternatif baskın değil ama güvenli bir uzlaşı adayıdır.")

    # ── GRA ──────────────────────────────────────────────────────────────────────
    elif base == "GRA":
        grg = _v("GRG") or best_score
        wgrg = _vw("GRG") or worst_score
        if is_en:
            calc = ("Grey Relational Analysis (GRA) measures how closely each alternative matches a reference sequence (ideal profile). "
                    "The <b>Grey Relational Coefficient (GRC)</b> for each criterion ranges in (0, 1] — "
                    "ζ = 0.5 (distinguishing coefficient) controls contrast sensitivity. "
                    "The <b>GRG</b> (Grey Relational Grade) is the weighted average of GRCs.")
            found = (f"<b>{best_name}</b> achieved GRG = <b>{grg:.4f}</b>, indicating the closest match to the reference ideal. "
                     f"<b>{worst_name}</b> scored GRG = <b>{wgrg:.4f}</b>. "
                     f"GRG spread: <b>{spread:.4f}</b> — "
                     + ("alternatives are tightly clustered; marginal GRC differences in key criteria are decisive." if spread < 0.1 else
                        "clear differentiation relative to the reference profile." if spread > 0.25 else
                        "moderate spread along the reference profile."))
            watch = ("GRC values near 1.0 for all criteria indicate near-perfect alignment with the ideal; "
                     "GRC near 0.33 (ζ/(1+ζ) at maximum distance) signals poor alignment. "
                     "Inspect per-criterion GRCs to find which criteria drag the leader's grade down.")
        else:
            calc = ("Gri İlişkisel Analiz (GRA), her alternatifin referans diziye (ideal profile) ne kadar yakın olduğunu ölçer. "
                    "Her kriter için <b>Gri İlişkisel Katsayı (GRC)</b> (0, 1] aralığında yer alır — "
                    "ζ = 0.5 (ayrım katsayısı) kontrast duyarlılığını kontrol eder. "
                    "<b>GRG</b> (Gri İlişkisel Derece), GRC değerlerinin ağırlıklı ortalamasıdır.")
            found = (f"<b>{best_name}</b> GRG = <b>{grg:.4f}</b> ile referans ideale en yakın alternatif. "
                     f"<b>{worst_name}</b> GRG = <b>{wgrg:.4f}</b>. "
                     f"GRG aralığı: <b>{spread:.4f}</b> — "
                     + ("alternatifler birbirine çok yakın; belirleyici fark kriterlerin GRC değerlerinde gizlidir." if spread < 0.1 else
                        "referans profile göre belirgin ayrışma gözlemleniyor." if spread > 0.25 else
                        "orta düzey ayrışma."))
            watch = ("Tüm kriterlerde GRC ≈ 1.0 → idealin neredeyse tam karşılanması; "
                     "GRC ≈ 0.33 → maksimum uzaklıkta zayıf uyum. "
                     "Liderin GRG'sini düşüren kriterleri kriter bazlı GRC tablosundan tespit edin.")

    # ── MARCOS ───────────────────────────────────────────────────────────────────
    elif base == "MARCOS":
        km, kp, ui = _v("K-"), _v("K+"), _v("U_i")
        if is_en:
            calc = ("MARCOS evaluates each alternative against <em>both</em> the ideal (AI) and anti-ideal (AAI) simultaneously. "
                    "<b>K⁻ = Sᵢ / S_AAI</b>: utility relative to the anti-ideal — higher means the alternative is farther from the worst. "
                    "<b>K⁺ = Sᵢ / S_AI</b>: utility relative to the ideal — higher means the alternative is closer to the best. "
                    "The final <b>Uᵢ</b> (utility function) combines K⁻ and K⁺ via f(K⁻) and f(K⁺) weighting fractions.")
            found = (f"<b>{best_name}</b> achieved Uᵢ = <b>{ui if ui is not None else best_score:.4f}</b>"
                     + (f" (K⁻={km}, K⁺={kp})" if km is not None else "") + ". "
                     + ("K⁺ close to 1 indicates near-ideal performance; "
                        if kp is not None and kp > 0.8 else
                        "K⁺ below 0.5 means still significant distance from ideal despite leading; " if kp is not None and kp < 0.5 else "")
                     + f"<b>{worst_name}</b> scored Uᵢ = <b>{_vw('U_i') if _col('U_i') else worst_score:.4f}</b>.")
            watch = ("K⁻ > 1 means the alternative outperforms the anti-ideal by more than 100% — strong floor-clearing. "
                     "K⁺ > 1 is theoretically possible only when S_i > S_AI (scores can exceed the augmented ideal row due to weighting).")
        else:
            calc = ("MARCOS her alternatifi <em>aynı anda</em> ideal (AI) ve anti-ideal (AAI) referanslara göre değerlendirir. "
                    "<b>K⁻ = Sᵢ / S_AAI</b>: anti-ideale göre fayda — yüksek olması en kötüden uzaklaşıldığını gösterir. "
                    "<b>K⁺ = Sᵢ / S_AI</b>: ideale göre fayda — yüksek olması en iyiye yaklaşıldığını gösterir. "
                    "Nihai <b>Uᵢ</b> (fayda fonksiyonu), K⁻ ve K⁺'yı f(K⁻) ve f(K⁺) ağırlık kesirleriyle birleştirir.")
            found = (f"<b>{best_name}</b> Uᵢ = <b>{ui if ui is not None else best_score:.4f}</b> ile lider"
                     + (f" (K⁻={km}, K⁺={kp})" if km is not None else "") + ". "
                     + ("K⁺ 1'e yakın — lider neredeyse ideal performansta. "
                        if kp is not None and kp > 0.8 else
                        "K⁺ < 0.5 — liderin ideale hâlâ uzaklığı var; diğer kriterler baskı altında olabilir. " if kp is not None and kp < 0.5 else "")
                     + f"<b>{worst_name}</b> Uᵢ = <b>{_vw('U_i') if _col('U_i') else worst_score:.4f}</b>.")
            watch = ("K⁻ > 1: anti-ideali %100'ün üzerinde geride bırakmak — zemin güçlü. "
                     "K⁺ > 1: teorik olarak mümkün ama nadir; ağırlıklı skorum S_AI'yi aşıyor demektir.")

    # ── CoCoSo ───────────────────────────────────────────────────────────────────
    elif base == "CoCoSo":
        ka, kb, kc = _v("K_a"), _v("K_b"), _v("K_c")
        if is_en:
            calc = ("CoCoSo combines three compromise strategies from the weighted sum (Sᵢ) and weighted power (Pᵢ): "
                    "<b>Kₐ</b> = arithmetic–geometric compromise; "
                    "<b>K_b</b> = sum-relative comparison (Sᵢ/ΣSⱼ + Pᵢ/ΣPⱼ); "
                    "<b>K_c</b> = max-relative (λ-balanced). "
                    "Final score = ∛(Kₐ·K_b·K_c) + (Kₐ+K_b+K_c)/3.")
            found = (f"<b>{best_name}</b> scores Kₐ={ka}, K_b={kb}, K_c={kc} — "
                     + ("all three strategies agree, confirming robust leadership. " if all(v is not None for v in [ka, kb, kc]) else "")
                     + f"Final score = <b>{best_score:.4f}</b>. "
                     + f"<b>{worst_name}</b> final = <b>{worst_score:.4f}</b>.")
            watch = ("If Kₐ and K_c are high but K_b is low, the alternative performs well in absolute terms "
                     "but ranks lower within the field — check Sᵢ relative to other alternatives.")
        else:
            calc = ("CoCoSo ağırlıklı toplam (Sᵢ) ve ağırlıklı güç (Pᵢ) üzerinden üç uzlaşı stratejisini birleştirir: "
                    "<b>Kₐ</b> = aritmetik-geometrik uzlaşı; "
                    "<b>K_b</b> = toplam-göreli karşılaştırma (Sᵢ/ΣSⱼ + Pᵢ/ΣPⱼ); "
                    "<b>K_c</b> = maksimum-göreli (λ dengeli). "
                    "Nihai skor = ∛(Kₐ·K_b·K_c) + (Kₐ+K_b+K_c)/3.")
            found = (f"<b>{best_name}</b>: Kₐ={ka}, K_b={kb}, K_c={kc} — "
                     + ("üç strateji de aynı fikirde, liderlik tutarlı ve güçlü. " if all(v is not None for v in [ka, kb, kc]) else "")
                     + f"Nihai skor = <b>{best_score:.4f}</b>. "
                     + f"<b>{worst_name}</b> nihai skor = <b>{worst_score:.4f}</b>.")
            watch = ("Kₐ ve K_c yüksek ama K_b düşükse: alternatif mutlak anlamda iyi ama alan içindeki göreli konumu daha zayıf. "
                     "Sᵢ değerini diğer alternatiflerle karşılaştırarak kontrol edin.")

    # ── ARAS ─────────────────────────────────────────────────────────────────────
    elif base == "ARAS":
        si_v, ki_v = _v("S_i"), _v("K_i")
        s0 = round(float(pd.to_numeric(df["S_i"], errors="coerce").max()), 4) if "S_i" in df.columns else None
        if is_en:
            calc = ("ARAS augments the decision matrix with an optimal (ideal) row A₀, then normalises and weights all rows. "
                    "<b>Sᵢ</b> = weighted normalised performance of alternative i; "
                    "<b>S₀</b> = performance of the ideal row. "
                    "<b>Kᵢ = Sᵢ / S₀</b> is the optimality degree, ranging (0, 1] — closer to 1 means near-optimal.")
            found = (f"<b>{best_name}</b>: Sᵢ = {si_v}, Kᵢ = <b>{ki_v if ki_v is not None else best_score:.4f}</b>"
                     + (f" (S₀ = {s0})" if s0 else "") + ". "
                     + (f"Kᵢ = {ki_v:.2f} → achieves {ki_v*100:.1f}% of ideal performance. " if ki_v else "")
                     + f"<b>{worst_name}</b>: Kᵢ = <b>{_vw('K_i') if _col('K_i') else worst_score:.4f}</b>.")
            watch = ("Kᵢ < 0.5 across all alternatives signals that none reaches half the ideal — consider redefining the criterion scale or reviewing cost/benefit directions.")
        else:
            calc = ("ARAS karar matrisine optimal (ideal) bir A₀ satırı ekler, ardından tüm satırları normalize eder ve ağırlıklandırır. "
                    "<b>Sᵢ</b> = alternatifin ağırlıklı normalize performansı; "
                    "<b>S₀</b> = ideal satırın performansı. "
                    "<b>Kᵢ = Sᵢ / S₀</b> optimallik derecesidir; (0, 1] aralığında — 1'e yakın idealdir.")
            found = (f"<b>{best_name}</b>: Sᵢ = {si_v}, Kᵢ = <b>{ki_v if ki_v is not None else best_score:.4f}</b>"
                     + (f" (S₀ = {s0})" if s0 else "") + ". "
                     + (f"Kᵢ = {ki_v:.2f} → ideal performansın %{ki_v*100:.1f}'ine ulaşıyor. " if ki_v else "")
                     + f"<b>{worst_name}</b>: Kᵢ = <b>{_vw('K_i') if _col('K_i') else worst_score:.4f}</b>.")
            watch = ("Tüm alternatiflerde Kᵢ < 0.5 ise hiçbiri idealin yarısına ulaşamıyor — kriter ölçeğini veya maliyet/fayda yönlerini gözden geçirmeniz önerilir.")

    # ── MABAC ────────────────────────────────────────────────────────────────────
    elif base == "MABAC":
        if is_en:
            calc = ("MABAC defines a <em>Border Approximation Area (BAA)</em> for each criterion using the geometric mean of all alternatives. "
                    "Each cell in the distance matrix = wⱼ·(xᵢⱼ_norm − gⱼ): positive → alternative is above the BAA (upper approximation), "
                    "negative → below the BAA (lower approximation). "
                    "The final score <b>G</b> is the row sum of all criterion distances.")
            found = (f"<b>{best_name}</b> scored G = <b>{best_score:.4f}</b> — positioned above the BAA on more criteria. "
                     f"<b>{worst_name}</b> scored G = <b>{worst_score:.4f}</b>. "
                     + ("A negative G means the alternative falls below the BAA on a net basis — structurally weak positioning. "
                        if worst_score < 0 else ""))
            watch = ("Inspect per-criterion distances: a high G might be driven by a single criterion. "
                     "Consistent positive distances across criteria signal more robust leadership.")
        else:
            calc = ("MABAC her kriter için tüm alternatiflerin geometrik ortalamasından <em>Sınır Yaklaşım Alanı (BAA)</em> belirler. "
                    "Mesafe matrisinin her hücresi = wⱼ·(xᵢⱼ_norm − gⱼ): pozitif → BAA'nın üstünde (üst yaklaşım), "
                    "negatif → altında (alt yaklaşım). "
                    "Nihai skor <b>G</b>, tüm kriter mesafelerinin satır toplamıdır.")
            found = (f"<b>{best_name}</b> G = <b>{best_score:.4f}</b> ile daha fazla kriterde BAA üstünde konumlanıyor. "
                     f"<b>{worst_name}</b> G = <b>{worst_score:.4f}</b>. "
                     + ("Negatif G: alternatif net olarak BAA'nın altında — yapısal konumlanma zayıf. "
                        if worst_score < 0 else ""))
            watch = ("Kriter bazlı mesafeleri inceleyin: yüksek G tek bir kritere bağlıysa gerçek liderlik kırılgandır. "
                     "Tüm kriterlerde tutarlı pozitif mesafe, daha güvenilir liderliğin göstergesidir.")

    # ── Genel fallback ────────────────────────────────────────────────────────────
    else:
        if is_en:
            calc  = f"This table shows the intermediate calculation steps for the <b>{method_safe}</b> method."
            found = (f"<b>{best_name}</b> achieved the highest score = <b>{best_score:.4f}</b>; "
                     f"<b>{worst_name}</b> the lowest = <b>{worst_score:.4f}</b>. "
                     f"Score spread: <b>{spread:.4f}</b> across {n} alternatives.")
            watch = "Review the method documentation to interpret each column's contribution to the final ranking."
        else:
            calc  = f"Bu tablo <b>{method_safe}</b> yönteminin ara hesaplama adımlarını göstermektedir."
            found = (f"<b>{best_name}</b> en yüksek skoru = <b>{best_score:.4f}</b> aldı; "
                     f"<b>{worst_name}</b> en düşük = <b>{worst_score:.4f}</b>. "
                     f"{n} alternatif arasında skor aralığı: <b>{spread:.4f}</b>.")
            watch = "Her sütunun nihai sıralamaya katkısını anlamak için yöntem dokümantasyonunu inceleyiniz."

    fuzzy_note = ""
    if fz:
        fuzzy_note = (" <em>(Bulanık uzantı: değerler üçgensel bulanık sayılardan durulaştırılmıştır — "
                      "sonuçları kesin ölçümler yerine aralık temelli değerlendirmeler olarak yorumlayın.)</em>"
                      if not is_en else
                      " <em>(Fuzzy extension: values are defuzzified from triangular fuzzy numbers — "
                      "interpret results as interval-based estimates rather than precise measurements.)</em>")

    label_calc  = "Ne Hesaplandı?" if not is_en else "What Was Calculated?"
    label_found = "Ne Bulundu?" if not is_en else "What Was Found?"
    label_watch = "Neye Dikkat?" if not is_en else "What to Watch?"

    return (f"<b>📐 {label_calc}</b> {calc}{fuzzy_note}<br><br>"
            f"<b>🔍 {label_found}</b> {found}<br><br>"
            f"<b>⚠️ {label_watch}</b> {watch}")

def render_method_specific_insights(result: Dict[str, Any], alt_names: Dict[str, str]) -> None:
    ranking = result.get("ranking", {}) or {}
    method = ranking.get("method")
    details = ranking.get("details", {}) or {}
    if not method or not details:
        st.info(tt("Yöntem-özel içgörü üretilemedi.", "Method-specific insights are not available."))
        return

    base_method = method.replace("Fuzzy ", "")
    shown = False

    _lang = st.session_state.get("ui_lang", "TR")

    if base_method == "PROMETHEE":
        flows = details.get("promethee_flows")
        pref = details.get("promethee_pref_matrix")
        gaia_alt = details.get("promethee_gaia_alternatives")
        gaia_crit = details.get("promethee_gaia_criteria")
        gaia_axes = details.get("promethee_gaia_axes")
        gaia_dec = details.get("promethee_gaia_decision_axis")
        if isinstance(flows, pd.DataFrame) and not flows.empty:
            shown = True
            st.markdown(f"##### 🔀 {tt('PROMETHEE Akış Tablosu', 'PROMETHEE Flow Table')}")
            _live_c = _live_detail_commentary(method, flows, alt_names, _lang)
            if _live_c:
                st.markdown(f'<div class="commentary-box">{_safe_plain_commentary_html(_live_c)}</div>', unsafe_allow_html=True)
            _flows_disp = _map_alt_names_in_df(flows, alt_names)
            for col in ["PhiPlus", "PhiMinus", "PhiNet", "Skor"]:
                if col in _flows_disp.columns:
                    _flows_disp[col] = pd.to_numeric(_flows_disp[col], errors="coerce").round(4)
            render_table(localize_df(_flows_disp))
            fg1, fg2 = st.columns(2)
            with fg1:
                st.plotly_chart(fig_promethee_flows(flows, alt_names), use_container_width=True)
                st.markdown(f'<div class="commentary-box">{tt("PROMETHEE akış grafiği, her alternatifin diğerlerini ne ölçüde geçtiğini (Phi+) ve ne ölçüde geçildiğini (Phi-) birlikte gösterir. PhiNet çizgisi sıralamayı doğrudan belirler; yüksek net akış, ikili karşılaştırmalarda yapısal üstünlüğe işaret eder.", "The PROMETHEE flow chart shows how much each alternative outranks others (Phi+) and is outranked by others (Phi-) simultaneously. The PhiNet line directly determines the ranking; a high net flow indicates structural superiority in pairwise comparisons.")}</div>', unsafe_allow_html=True)
            with fg2:
                if isinstance(pref, pd.DataFrame) and not pref.empty:
                    st.plotly_chart(
                        fig_preference_heatmap(pref, alt_names, tt("PROMETHEE Tercih Matrisi", "PROMETHEE Preference Matrix")),
                        use_container_width=True,
                    )
                    st.markdown(f'<div class="commentary-box">{tt("Tercih matrisi ısı haritası, her alternatifin diğerine karşı ne kadar güçlü tercih ürettiğini gösterir. Satırdaki koyu hücreler baskınlığı, sütundaki koyu hücreler ise zayıf kalınan eşleşmeleri ortaya çıkarır.", "The preference matrix heatmap shows how strongly each alternative is preferred over another. Dark cells across a row indicate dominance, while dark cells down a column reveal pairings where the alternative is weak.")}</div>', unsafe_allow_html=True)
            if isinstance(gaia_alt, pd.DataFrame) and not gaia_alt.empty:
                st.markdown(f"##### 🧭 {tt('PROMETHEE GAIA Düzlemi', 'PROMETHEE GAIA Plane')}")
                if isinstance(gaia_axes, pd.DataFrame) and not gaia_axes.empty and "AçıklananVaryansOranı" in gaia_axes.columns:
                    _exp_vals = pd.to_numeric(gaia_axes["AçıklananVaryansOranı"], errors="coerce").fillna(0.0).tolist()
                    if len(_exp_vals) >= 2:
                        st.caption(
                            tt(
                                f"GAIA1+GAIA2 açıklanan varyans: %{(_exp_vals[0] + _exp_vals[1]) * 100:.1f}",
                                f"GAIA1+GAIA2 explained variance: {(_exp_vals[0] + _exp_vals[1]) * 100:.1f}%",
                            )
                        )
                st.plotly_chart(fig_promethee_gaia(gaia_alt, gaia_crit, gaia_dec, alt_names), use_container_width=True)
                st.markdown(
                        f'<div class="commentary-box">{tt("GAIA nasıl okunur? 1) Noktalar alternatifleri gösterir; birbirine yakın noktalar benzer davranır. 2) Oklar kriter yönünü gösterir; bir alternatif okun ucuna yakınsa o kriterde daha güçlüdür, ters yöndeyse daha zayıftır. 3) Ok ne kadar uzunsa kriterin ayırt ediciliği o kadar yüksektir. 4) Kırmızı Δ oku modelin genel tercih yönüdür; bu yöne yakın alternatifler genelde üst sıralarda yer alır.", "How to read GAIA: 1) Points are alternatives; nearby points behave similarly. 2) Arrows are criterion directions; an alternative near an arrow tip is stronger on that criterion, opposite direction is weaker. 3) Longer arrows mean stronger discriminating power. 4) The red Δ arrow is the overall preference direction; alternatives close to it usually rank higher.")}</div>',
                        unsafe_allow_html=True,
                    )
        elif isinstance(pref, pd.DataFrame) and not pref.empty:
            shown = True
            st.plotly_chart(
                fig_preference_heatmap(pref, alt_names, tt("PROMETHEE Tercih Matrisi", "PROMETHEE Preference Matrix")),
                use_container_width=True,
            )
            if isinstance(gaia_alt, pd.DataFrame) and not gaia_alt.empty:
                st.plotly_chart(fig_promethee_gaia(gaia_alt, gaia_crit, gaia_dec, alt_names), use_container_width=True)

    elif base_method == "TOPSIS":
        dist_df = details.get("distance_table")
        if isinstance(dist_df, pd.DataFrame) and not dist_df.empty:
            shown = True
            st.markdown(f"##### 📐 {tt('TOPSIS Uzaklık Tablosu', 'TOPSIS Distance Table')}")
            _live_c = _live_detail_commentary(method, dist_df, alt_names, _lang)
            if _live_c:
                st.markdown(f'<div class="commentary-box">{_safe_plain_commentary_html(_live_c)}</div>', unsafe_allow_html=True)
            _dist_disp = _map_alt_names_in_df(dist_df, alt_names)
            for col in ["D+", "D-", "Skor"]:
                if col in _dist_disp.columns:
                    _dist_disp[col] = pd.to_numeric(_dist_disp[col], errors="coerce").round(4)
            render_table(localize_df(_dist_disp))
            st.plotly_chart(fig_topsis_distance_scatter(dist_df, alt_names), use_container_width=True)
            st.markdown(f'<div class="commentary-box">{tt("TOPSIS uzaklık haritasında sağ üst bölgeye yaklaşan alternatifler ideala daha yakın ve negatif ideale daha uzaktır. Bu görünüm liderin yalnızca yüksek skor değil, aynı zamanda güçlü ayrışma ürettiğini göstermeye yarar.", "In the TOPSIS distance map, alternatives closer to the upper-right region are nearer to the ideal and farther from the negative ideal. This view shows whether the leader not only scores high but also separates strongly from the rest.")}</div>', unsafe_allow_html=True)

    elif base_method == "VIKOR":
        vikor_df = details.get("vikor_table")
        if isinstance(vikor_df, pd.DataFrame) and not vikor_df.empty:
            shown = True
            st.markdown(f"##### ⚖️ {tt('VIKOR Uzlaşı Tablosu', 'VIKOR Compromise Table')}")
            _live_c = _live_detail_commentary(method, vikor_df, alt_names, _lang)
            if _live_c:
                st.markdown(f'<div class="commentary-box">{_safe_plain_commentary_html(_live_c)}</div>', unsafe_allow_html=True)
            _vikor_disp = _map_alt_names_in_df(vikor_df, alt_names)
            for col in ["S", "R", "Q", "Skor"]:
                if col in _vikor_disp.columns:
                    _vikor_disp[col] = pd.to_numeric(_vikor_disp[col], errors="coerce").round(4)
            render_table(localize_df(_vikor_disp))
            st.plotly_chart(fig_vikor_components(vikor_df, alt_names), use_container_width=True)
            st.markdown(f'<div class="commentary-box">{tt("VIKOR bileşen grafiği grup faydası (S), bireysel pişmanlık (R) ve uzlaşı skoru (Q) arasındaki dengeyi görünür kılar. Düşük Q ile birlikte makul S ve R değerleri, lider alternatifin uzlaşı çözümü olarak daha savunulabilir olduğunu gösterir.", "The VIKOR component chart makes the balance among group utility (S), individual regret (R), and compromise score (Q) visible. A low Q together with reasonable S and R values indicates the leading alternative is more defensible as a compromise solution.")}</div>', unsafe_allow_html=True)

    elif base_method == "EDAS":
        edas_df = details.get("edas_table")
        if isinstance(edas_df, pd.DataFrame) and not edas_df.empty:
            shown = True
            st.markdown(f"##### 📏 {tt('EDAS Sapma Tablosu', 'EDAS Distance Table')}")
            _live_c = _live_detail_commentary(method, edas_df, alt_names, _lang)
            if _live_c:
                st.markdown(f'<div class="commentary-box">{_safe_plain_commentary_html(_live_c)}</div>', unsafe_allow_html=True)
            _edas_disp = _map_alt_names_in_df(edas_df, alt_names)
            for col in ["SP", "SN", "NSP", "NSN", "Skor"]:
                if col in _edas_disp.columns:
                    _edas_disp[col] = pd.to_numeric(_edas_disp[col], errors="coerce").round(4)
            render_table(localize_df(_edas_disp))
            st.plotly_chart(fig_edas_balance(edas_df, alt_names), use_container_width=True)
            st.markdown(f'<div class="commentary-box">{tt("EDAS grafiği, her alternatifin ortalama çözüme göre pozitif ve negatif sapma dengesini gösterir. Yüksek NSP ve yüksek NSN birlikte görülüyorsa alternatif hem avantaj yaratıyor hem de olumsuz sapmayı sınırlıyor demektir.", "The EDAS chart shows the balance of positive and negative distances from the average solution for each alternative. When high NSP and high NSN appear together, the alternative both creates advantage and limits adverse deviation.")}</div>', unsafe_allow_html=True)

    elif base_method == "CODAS":
        codas_df = details.get("codas_table")
        if isinstance(codas_df, pd.DataFrame) and not codas_df.empty:
            shown = True
            st.markdown(f"##### 🧭 {tt('CODAS Uzaklık Tablosu', 'CODAS Distance Table')}")
            _live_c = _live_detail_commentary(method, codas_df, alt_names, _lang)
            if _live_c:
                st.markdown(f'<div class="commentary-box">{_safe_plain_commentary_html(_live_c)}</div>', unsafe_allow_html=True)
            _codas_disp = _map_alt_names_in_df(codas_df, alt_names)
            for col in ["E", "T", "H", "Skor"]:
                if col in _codas_disp.columns:
                    _codas_disp[col] = pd.to_numeric(_codas_disp[col], errors="coerce").round(4)
            render_table(localize_df(_codas_disp))
            st.plotly_chart(fig_codas_distance_map(codas_df, alt_names), use_container_width=True)
            st.markdown(f'<div class="commentary-box">{tt("CODAS haritası, alternatiflerin negatif idealden hem Öklid hem Manhattan uzaklığıyla nasıl ayrıştığını gösterir. Sağ üst bölgeye taşınan ve yüksek H skoru alan alternatifler riskli kötü senaryolardan daha güçlü biçimde ayrışır.", "The CODAS map shows how alternatives separate from the negative ideal in both Euclidean and Taxicab distance terms. Alternatives moving toward the upper-right region with high H scores separate more strongly from adverse worst-case conditions.")}</div>', unsafe_allow_html=True)

    if not shown:
        _detail_frames = _extract_detail_tables(details)
        _detail_frames.pop("result_table", None)
        if _detail_frames:
            st.markdown(f"##### 🧾 {tt('Yöntem-Özel Detay Tablosu', 'Method-Specific Detail Table')}")
            _first_key = next(iter(_detail_frames))
            _first_df = _detail_frames[_first_key]
            _live_c = _live_detail_commentary(method, _first_df, alt_names, _lang)
            if _live_c:
                st.markdown(f'<div class="commentary-box">{_safe_plain_commentary_html(_live_c)}</div>', unsafe_allow_html=True)
            render_table(localize_df(_map_alt_names_in_df(_first_df, alt_names).head(250)))
        else:
            st.info(tt("Bu yöntem için ek yöntem-özel görsel bulunmuyor.", "No additional method-specific visualization is available for this method."))

def render_tab_assistant(commentary: str, key: str = "") -> None:
    with st.expander(tt("💬 Analiz Asistanı — Yorum", "💬 Analysis Assistant — Commentary"), expanded=True):
        st.markdown(f'<div class="tab-assistant">{_safe_plain_commentary_html(commentary)}</div>', unsafe_allow_html=True)
    st.caption(tt("Detaylar için lütfen aşağıdaki tablo ve şekilleri inceleyin.", "For details, please review the tables and figures below."))

def _top_cv_crit(result: Dict[str, Any]) -> tuple[str, float]:
    data = result["selected_data"]
    cv = (data.std() / data.mean().abs())
    return cv.idxmax(), float(cv.max())

def _max_corr_pair(result: Dict[str, Any]) -> tuple[str, str, float]:
    corr = result.get("correlation_matrix")
    if corr is None or corr.shape[0] < 2:
        return ("—", "—", 0.0)
    mask = ~np.eye(len(corr), dtype=bool)
    vals = corr.abs().where(mask).stack()
    if vals.empty:
        return ("—", "—", 0.0)
    pair = vals.idxmax()
    return pair[0], pair[1], float(vals.max())

def _compose_commentary(simple: str, academic: str, action: str, example: str = "") -> str:
    is_en = st.session_state.get("ui_lang") == "EN"
    main_commentary = academic or simple
    labels = {
        "commentary": "Commentary" if is_en else "Yorum",
        "example": "Short Example" if is_en else "Kısa Örnek",
        "action": "What To Do Next" if is_en else "Ne Yapmalı",
    }
    parts = [f"<strong>{labels['commentary']}:</strong> {main_commentary}"]
    if example:
        parts.append(f"<strong>{labels['example']}:</strong> {example}")
    if action:
        parts.append(f"<strong>{labels['action']}:</strong> {action}")
    return "<br><br>".join(parts)

def _compose_structured_commentary(did: str, why: str, found: str, example: str = "", action: str = "") -> str:
    is_en = st.session_state.get("ui_lang") == "EN"
    labels = {
        "did": "What Was Done" if is_en else "Ne Yapıldı",
        "why": "Why It Was Done" if is_en else "Neden Yapıldı",
        "found": "What Was Found" if is_en else "Ne Bulundu",
        "example": "Short Example" if is_en else "Kısa Örnek",
        "action": "How To Read It" if is_en else "Nasıl Okunmalı",
    }
    parts = [
        f"<strong>{labels['did']}:</strong> {did}",
        f"<strong>{labels['why']}:</strong> {why}",
        f"<strong>{labels['found']}:</strong> {found}",
    ]
    if example:
        parts.append(f"<strong>{labels['example']}:</strong> {example}")
    if action:
        parts.append(f"<strong>{labels['action']}:</strong> {action}")
    return "<br><br>".join(parts)

def gen_weight_robustness_commentary(result: Dict[str, Any]) -> str:
    wr = result.get("weight_robustness", {}) or {}
    if not wr:
        return tt("Ağırlık sağlamlık verisi mevcut değil.", "Weight robustness data is not available.")
    if wr.get("info"):
        return str(wr.get("info"))
    if wr.get("error"):
        return tt("Ağırlık sağlamlık testleri üretilemedi.", "Weight robustness tests could not be generated.")

    loo_mean = wr.get("loo_mean_rho")
    loo_min = wr.get("loo_min_rho")
    top_match = wr.get("loo_top_match_rate")
    eff_n = wr.get("effective_criteria_n")
    dom = wr.get("dominance_ratio")
    loo_df = wr.get("leave_one_out")
    boot_df = wr.get("bootstrap_summary")

    worst_alt = "—"
    worst_rho = np.nan
    if isinstance(loo_df, pd.DataFrame) and not loo_df.empty:
        rho_col = col_key(loo_df, "SpearmanRho", "SpearmanRho")
        alt_col = col_key(loo_df, "Alternatif", "Alternative")
        worst_idx = pd.to_numeric(loo_df[rho_col], errors="coerce").idxmin()
        if pd.notna(worst_idx):
            worst_alt = str(loo_df.loc[worst_idx, alt_col])
            worst_rho = float(pd.to_numeric(loo_df.loc[worst_idx, rho_col], errors="coerce"))

    uncertain_crit = "—"
    uncertain_span = np.nan
    if isinstance(boot_df, pd.DataFrame) and not boot_df.empty:
        crit_col = col_key(boot_df, "Kriter", "Criterion")
        low_col = col_key(boot_df, "AltYuzde5", "Lower5Pct")
        high_col = col_key(boot_df, "UstYuzde95", "Upper95Pct")
        boot_tmp = boot_df.copy()
        boot_tmp["_span"] = pd.to_numeric(boot_tmp[high_col], errors="coerce") - pd.to_numeric(boot_tmp[low_col], errors="coerce")
        idx = boot_tmp["_span"].idxmax()
        if pd.notna(idx):
            uncertain_crit = str(boot_tmp.loc[idx, crit_col])
            uncertain_span = float(pd.to_numeric(boot_tmp.loc[idx, "_span"], errors="coerce"))

    is_en = st.session_state.get("ui_lang") == "EN"
    did = (
        "We tested the weight vector in two ways: first, each alternative was removed one by one (leave-one-out); second, the data was repeatedly resampled to build bootstrap confidence intervals for each criterion weight."
        if is_en else
        "Ağırlık vektörü iki yoldan sınandı: önce her alternatif tek tek çıkarıldı (leave-one-out), ardından veri tekrar örneklenerek her kriter ağırlığı için bootstrap güven aralıkları üretildi."
    )
    why = (
        "The goal is to see whether the reported criterion priorities depend too much on a single row or on small sample fluctuations."
        if is_en else
        "Amaç, raporlanan kriter önceliklerinin tek bir satıra ya da küçük örnek oynamalarına aşırı bağlı olup olmadığını görmektir."
    )

    if loo_mean is not None and np.isfinite(loo_mean) and top_match is not None and np.isfinite(top_match):
        if loo_mean >= 0.90 and top_match >= 0.75:
            stability_text = (
                "Overall robustness is high."
                if is_en else
                "Genel sağlamlık düzeyi yüksektir."
            )
            action = (
                "You can report the weight structure with stronger confidence; still keep the most uncertain criterion under observation in the discussion section."
                if is_en else
                "Ağırlık yapısını daha güçlü biçimde raporlayabilirsiniz; yine de tartışma bölümünde en belirsiz kriteri ayrıca izleyin."
            )
        elif loo_mean >= 0.75 and top_match >= 0.50:
            stability_text = (
                "Overall robustness is moderate."
                if is_en else
                "Genel sağlamlık düzeyi ortadır."
            )
            action = (
                "Present the weights together with a caution note, and emphasize which criterion becomes unstable under data perturbation."
                if is_en else
                "Ağırlıkları temkin notuyla birlikte sunun ve veri oynadığında hangi kriterin daha çabuk bozulduğunu özellikle belirtin."
            )
        else:
            stability_text = (
                "Overall robustness is low."
                if is_en else
                "Genel sağlamlık düzeyi düşüktür."
            )
            action = (
                "Avoid presenting the weights as fixed truth; compare alternative weighting methods and treat the current profile as sensitive."
                if is_en else
                "Ağırlıkları kesin gerçek gibi sunmayın; alternatif ağırlık yöntemleriyle karşılaştırın ve mevcut profili duyarlı kabul edin."
            )
    else:
        stability_text = tt("Kararlılık özeti üretilemedi.", "Stability summary could not be generated.")
        action = tt("Veri yapısını yeniden kontrol edin.", "Recheck the data structure.")

    found = (
        f"{stability_text} Mean leave-one-out agreement is {float(loo_mean):.3f}, worst-case agreement is {float(loo_min):.3f}, and top-criterion consistency is {float(top_match)*100:.1f}%. "
        f"The effective number of criteria is {float(eff_n):.2f}. "
        + (
            f"The top-to-second criterion weight ratio is {float(dom):.2f}, which suggests a {'concentrated' if float(dom) >= 1.80 else 'balanced'} weight profile. "
            if dom is not None and np.isfinite(dom) else ""
        )
        + (
            f"The weakest leave-one-out scenario appears when {worst_alt} is removed (ρ = {float(worst_rho):.3f}). "
            if np.isfinite(worst_rho) else ""
        )
        + (
            f"The widest bootstrap uncertainty band belongs to {uncertain_crit} (span ≈ {float(uncertain_span):.4f}), so the first movement is most likely to appear there."
            if np.isfinite(uncertain_span) else ""
        )
        if is_en else
        f"{stability_text} Leave-one-out ortalama uyumu {float(loo_mean):.3f}, en zayıf senaryo uyumu {float(loo_min):.3f} ve lider kriter tutarlılığı %{float(top_match)*100:.1f} düzeyindedir. "
        f"Etkin kriter sayısı {float(eff_n):.2f} olarak görünmektedir. "
        + (
            f"Lider/ikinci kriter ağırlık oranı {float(dom):.2f} olduğu için ağırlık yapısı {'yoğunlaşmış' if float(dom) >= 1.80 else 'daha dengeli'} okunabilir. "
            if dom is not None and np.isfinite(dom) else ""
        )
        + (
            f"En zayıf leave-one-out senaryosu {worst_alt} çıkarıldığında oluşmuştur (ρ = {float(worst_rho):.3f}). "
            if np.isfinite(worst_rho) else ""
        )
        + (
            f"Bootstrap tarafında en geniş belirsizlik bandı {uncertain_crit} kriterinde görülmektedir (aralık ≈ {float(uncertain_span):.4f}); bu nedenle ilk oynama en çok burada beklenir."
            if np.isfinite(uncertain_span) else ""
        )
    )
    example = (
        f"For example, if removing one alternative still leaves the same criterion on top, the model keeps telling the same story. If the story changes immediately, the weighting result depends too much on that observation. Here the most fragile removal is {worst_alt}."
        if is_en else
        f"Örneğin bir alternatif çıkarıldığında en önemli kriter yine aynı kalıyorsa model aynı hikâyeyi anlatmaya devam eder. Hikâye hemen değişiyorsa ağırlık sonucu o gözleme fazla bağımlıdır. Bu çalışmada en hassas çıkarma senaryosu {worst_alt} ile görülmüştür."
    )
    return _compose_structured_commentary(did, why, found, example, action)

def _ranking_method_commentary(result: Dict[str, Any], top_alt: str, top_disp: str) -> tuple[str, str, str]:
    ranking = result.get("ranking", {}) or {}
    method = str(ranking.get("method") or "")
    base_method = method.replace("Fuzzy ", "")
    details = ranking.get("details", {}) or {}
    is_en = st.session_state.get("ui_lang") == "EN"

    if base_method == "TOPSIS" and isinstance(details.get("distance_table"), pd.DataFrame):
        dist_df = details["distance_table"]
        row = dist_df[dist_df["Alternatif"].astype(str) == str(top_alt)]
        if not row.empty:
            d_plus = float(row.iloc[0]["D+"])
            d_minus = float(row.iloc[0]["D-"])
            if is_en:
                return (
                    f"{method} looks for the option closest to the ideal and farthest from the weakest profile. On that basis, {top_disp} stands out.",
                    f"For the leading option, D+ = {d_plus:.4f} and D- = {d_minus:.4f}. This means it is relatively close to the desired profile while staying far from the undesirable profile.",
                    "If the first two options are close, read this map together with the Monte Carlo and comparison tabs before making a final decision.",
                )
            return (
                f"{method}, en iyi profile en yakın ve en zayıf profile en uzak seçeneği arar. Bu ölçüte göre {top_disp} öne çıkıyor.",
                f"Lider seçenek için D+ = {d_plus:.4f} ve D- = {d_minus:.4f}. Yani {top_disp}, istenen profile görece yakın, zayıf profile ise görece uzaktır.",
                "İlk iki seçenek birbirine yakınsa son kararı vermeden önce bu haritayı Monte Carlo ve yöntem karşılaştırması sekmeleriyle birlikte okuyun.",
            )

    if base_method == "VIKOR" and isinstance(details.get("vikor_table"), pd.DataFrame):
        vikor_df = details["vikor_table"]
        row = vikor_df[vikor_df["Alternatif"].astype(str) == str(top_alt)]
        if not row.empty:
            s_val = float(row.iloc[0]["S"])
            r_val = float(row.iloc[0]["R"])
            q_val = float(row.iloc[0]["Q"])
            if is_en:
                return (
                    f"{method} tries to find the most balanced compromise. Here, {top_disp} appears as the most acceptable middle ground.",
                    f"The leading option has S = {s_val:.4f}, R = {r_val:.4f}, and Q = {q_val:.4f}. A lower Q means the option balances collective benefit and individual regret more successfully in this dataset.",
                    "This result is especially useful when you want a defensible compromise rather than an aggressive winner-takes-all choice.",
                )
            return (
                f"{method}, en dengeli uzlaşı çözümünü bulmaya çalışır. Bu veri setinde {top_disp} en kabul edilebilir orta yol olarak görünüyor.",
                f"Lider seçenek için S = {s_val:.4f}, R = {r_val:.4f} ve Q = {q_val:.4f}. Düşük Q değeri, bu alternatifin grup faydası ile bireysel pişmanlık arasında daha iyi denge kurduğunu gösterir.",
                "Keskin bir kazanan yerine savunulabilir bir uzlaşı arıyorsanız bu sonuç daha anlamlıdır.",
            )

    if base_method == "PROMETHEE" and isinstance(details.get("promethee_flows"), pd.DataFrame):
        flows = details["promethee_flows"]
        row = flows[flows["Alternatif"].astype(str) == str(top_alt)]
        if not row.empty:
            phi_plus = float(row.iloc[0]["PhiPlus"])
            phi_minus = float(row.iloc[0]["PhiMinus"])
            phi_net = float(row.iloc[0]["PhiNet"])
            if is_en:
                return (
                    f"{method} compares options pair by pair. In these direct matchups, {top_disp} defeats the others more often than it is defeated.",
                    f"For the leading option, Phi+ = {phi_plus:.4f}, Phi- = {phi_minus:.4f}, and PhiNet = {phi_net:.4f}. A high net flow means structural superiority across pairwise comparisons in this dataset.",
                    "Use the flow chart and preference matrix together if you want to see not only who leads, but against whom that lead is strong or weak.",
                )
            return (
                f"{method}, seçenekleri ikili karşılaştırmalarla okur. Bu eşleşmelerde {top_disp}, diğerlerini geçtiği durumlarda daha baskın görünüyor.",
                f"Lider seçenek için Phi+ = {phi_plus:.4f}, Phi- = {phi_minus:.4f} ve PhiNet = {phi_net:.4f}. Yüksek net akış, bu veri setinde ikili karşılaştırmaların genelinde yapısal üstünlük olduğunu gösterir.",
                "Sadece kimin önde olduğuna değil, hangi rakiplere karşı güçlü veya zayıf olduğuna da bakmak istiyorsanız akış grafiği ile tercih matrisini birlikte okuyun.",
            )

    if base_method == "EDAS" and isinstance(details.get("edas_table"), pd.DataFrame):
        edas_df = details["edas_table"]
        row = edas_df[edas_df["Alternatif"].astype(str) == str(top_alt)]
        if not row.empty:
            nsp = float(row.iloc[0]["NSP"])
            nsn = float(row.iloc[0]["NSN"])
            if is_en:
                return (
                    f"{method} checks who stays favorably above the average profile while avoiding negative deviation. On that basis, {top_disp} is the strongest candidate.",
                    f"For the leading option, NSP = {nsp:.4f} and NSN = {nsn:.4f}. The joint strength of these two components indicates a balanced performance relative to the dataset average.",
                    "This method is useful when you want to compare each option against the general market or sample average, not only against the best case.",
                )
            return (
                f"{method}, ortalama profile göre kimin avantajlı kaldığını ve kimin olumsuz sapmayı daha iyi sınırladığını ölçer. Bu açıdan {top_disp} en güçlü adaydır.",
                f"Lider seçenek için NSP = {nsp:.4f} ve NSN = {nsn:.4f}. Bu iki bileşenin birlikte güçlü olması, veri seti ortalamasına göre dengeli bir üstünlüğe işaret eder.",
                "En iyi duruma değil, genel ortalamaya göre kimin iyi kaldığını görmek istiyorsanız bu yöntem özellikle faydalıdır.",
            )

    if base_method == "CODAS" and isinstance(details.get("codas_table"), pd.DataFrame):
        codas_df = details["codas_table"]
        row = codas_df[codas_df["Alternatif"].astype(str) == str(top_alt)]
        if not row.empty:
            h_val = float(row.iloc[0]["H"])
            e_val = float(row.iloc[0]["E"])
            t_val = float(row.iloc[0]["T"])
            if is_en:
                return (
                    f"{method} rewards options that move far away from the weakest profile. On this criterion, {top_disp} creates the strongest separation.",
                    f"For the leading option, E = {e_val:.4f}, T = {t_val:.4f}, and H = {h_val:.4f}. A high H score means the option stays farther from adverse scenarios in multiple distance senses.",
                    "This reading is especially useful when you care about avoiding weak outcomes, not only chasing the best average score.",
                )
            return (
                f"{method}, en zayıf profile uzaklaşan seçenekleri ödüllendirir. Bu açıdan {top_disp} en güçlü ayrışmayı üretmiştir.",
                f"Lider seçenek için E = {e_val:.4f}, T = {t_val:.4f} ve H = {h_val:.4f}. Yüksek H skoru, olumsuz senaryolardan birden fazla uzaklık ölçüsünde uzak kalındığını gösterir.",
                "Yalnızca ortalama performansı değil, zayıf sonuçlardan kaçınmayı önemsiyorsanız bu okuma daha değerlidir.",
            )

    ph = result.get("method_philosophy", {}).get("ranking", {}) or {}
    method_lbl = method_display_name(method)
    _, academic_fb_en = _method_fallback_lines(method_lbl or "Selected method", "EN")
    _, academic_fb_tr = _method_fallback_lines(method_lbl or "Seçilen yöntem", "TR")
    if is_en:
        return (
            f"{method} places {top_disp} in the first position for this dataset.",
            _lang_text(
                ph.get("academic", ""),
                "EN",
                "The selected method interprets the same data through its own ranking logic and identifies the strongest candidate accordingly."
                if not academic_fb_en else academic_fb_en,
            ),
            "Read this result together with the score gap, comparison, and robustness outputs before turning it into a final decision.",
        )
    return (
        f"{method} bu veri setinde {top_disp} alternatifini ilk sıraya yerleştirmiştir.",
        _lang_text(
            ph.get("academic", ""),
            "TR",
            "Seçilen yöntem aynı veriyi kendi karar mantığıyla okuyarak en güçlü adayı belirlemiştir."
            if not academic_fb_tr else academic_fb_tr,
        ),
        "Bu sonucu son karara dönüştürmeden önce skor farkı, yöntem karşılaştırması ve sağlamlık çıktılarıyla birlikte okuyun.",
    )

def gen_stat_commentary(result: Dict[str, Any]) -> str:
    n_alt = result["selected_data"].shape[0]
    n_crit = result["selected_data"].shape[1]
    cv_crit, cv_val = _top_cv_crit(result)
    c1, c2, max_rho = _max_corr_pair(result)
    if st.session_state.get("ui_lang") == "EN":
        simple = f"This dataset compares <strong>{n_alt} options</strong> across <strong>{n_crit} criteria</strong>. The criterion that separates the options most clearly is <strong>{cv_crit}</strong>."
        example = (
            f"For a non-technical reading: if two criteria move almost together, the model may be hearing the same story twice. Here the strongest overlap is between <strong>{c1}</strong> and <strong>{c2}</strong>."
            if max_rho > 0.75 else
            "For a non-technical reading: the criteria are not strongly repeating each other, so the model can look at the problem from multiple angles."
        )
        if max_rho > 0.75:
            academic = f"<strong>{cv_crit}</strong> has the highest coefficient of variation (Coeff. of Variation ≈{cv_val:.2f}), so it carries the strongest separating signal. At the same time, the high correlation between <strong>{c1}</strong> and <strong>{c2}</strong> (|ρ| = {max_rho:.2f}) suggests partial information overlap."
            action = "If these two criteria are conceptually similar, consider a correlation-aware weighting method such as CRITIC or simplify the active criterion set."
        else:
            academic = f"<strong>{cv_crit}</strong> has the highest coefficient of variation (Coeff. of Variation ≈{cv_val:.2f}), meaning it contributes the strongest discrimination in the ranking problem. Maximum inter-criterion correlation remains limited at |ρ| = {max_rho:.2f}, so redundancy risk is controlled."
            action = "The current data profile is appropriate for objective weighting and multi-method reading without an obvious structural warning."
        return _compose_commentary(simple, academic, action, example)
    simple = f"Bu veri setinde <strong>{n_alt} seçenek</strong>, <strong>{n_crit} kritere</strong> göre karşılaştırılıyor. Seçenekleri birbirinden en çok ayıran kriter <strong>{cv_crit}</strong> görünüyor."
    example = (
        f"Basit bir örnekle: <strong>{c1}</strong> ve <strong>{c2}</strong> neredeyse birlikte hareket ediyorsa sistem aynı bilgiyi iki kez okuyor olabilir."
        if max_rho > 0.75 else
        "Basit bir örnekle: kriterler birbirinin kopyası gibi davranmıyorsa sistem probleme farklı açılardan bakabilir."
    )
    if max_rho > 0.75:
        academic = f"<strong>{cv_crit}</strong> kriteri varyasyon katsayısı bakımından en yüksek ayrıştırıcı gücü sergiliyor (Varyasyon Katsayısı ≈{cv_val:.2f}). Bununla birlikte <strong>{c1}</strong> ile <strong>{c2}</strong> arasındaki yüksek ilişki (|ρ| = {max_rho:.2f}), kısmi bilgi tekrarına işaret ediyor."
        action = "Bu iki kriter kavramsal olarak da benzerse CRITIC gibi korelasyon duyarlı bir ağırlıklandırma tercih edin veya aktif kriter setini sadeleştirin."
    else:
        academic = f"<strong>{cv_crit}</strong> kriteri varyasyon katsayısı bakımından en güçlü ayrıştırıcı sinyali taşıyor (Varyasyon Katsayısı ≈{cv_val:.2f}). Kriterler arası en yüksek ilişki |ρ| = {max_rho:.2f} düzeyinde kaldığı için tekrar riski sınırlı görünüyor."
        action = "Mevcut veri profili, belirgin bir yapısal uyarı vermeden objektif ağırlıklandırma ve çoklu yöntem okumasına uygundur."
    return _compose_commentary(simple, academic, action, example)

def gen_weight_commentary(result: Dict[str, Any]) -> str:
    w_method = result["weights"]["method"]
    w_method_disp = method_display_name(w_method)
    w_vals = result["weights"]["values"]
    top_k = max(w_vals, key=w_vals.get)
    top_v = float(w_vals[top_k])
    concentration = sum(v**2 for v in w_vals.values())
    mode = result.get("weights", {}).get("details", {}).get("mode", "objective")
    if st.session_state.get("ui_lang") == "EN":
        if mode == "equal":
            simple = "All criteria were given equal importance. In other words, the model was instructed not to privilege any one criterion over the others."
            academic = "This is a normative baseline rather than a data-derived weighting scheme. It is useful when you want the ranking to reflect balanced treatment across all criteria."
            example = "Example: if price, quality, and speed are all treated equally, none of them gets special leverage in the final score."
            action = "Use this mode when you want a neutral baseline, then compare it with an objective weighting result to see whether the story changes."
        elif mode == "manual":
            simple = f"The final ranking follows the priorities you entered manually. The strongest emphasis was placed on <strong>{top_k}</strong>."
            academic = f"Manual weighting preserves user judgment and then normalizes it for computation. The highest declared emphasis is <strong>{top_k}</strong> with normalized weight {top_v:.4f}."
            example = "Example: if you care much more about reliability than cost, the system will deliberately push reliability to the front of the decision."
            action = "Use this mode when domain preference is intentional, but validate the outcome with sensitivity analysis because the result depends directly on your preference structure."
        else:
            simple = f"The system says the most influential criterion is <strong>{top_k}</strong>. That means changes in this criterion will shape the final ranking more than the others."
            example = f"Example: if <strong>{top_k}</strong> worsens for a leading option, that option may lose ground faster than expected."
            if top_v > 0.35:
                academic = f"<strong>{w_method_disp}</strong> assigns the highest weight to <strong>{top_k}</strong> (w = {top_v:.4f}). This is a concentrated structure, so the final ranking may depend strongly on one dominant signal."
                action = f"Interpret the ranking with explicit attention to <strong>{top_k}</strong> and examine how sensitive the result is when this weight changes."
            else:
                academic = f"<strong>{w_method_disp}</strong> assigns the highest weight to <strong>{top_k}</strong> (w = {top_v:.4f}), while the overall concentration index remains moderate at ≈ {concentration:.2f}. This indicates a relatively balanced weighting profile."
                action = f"The weighting structure is reasonably distributed; still, prioritize <strong>{top_k}</strong> in sensitivity checks because it remains the main driver."
        return _compose_commentary(simple, academic, action, example)
    if mode == "equal":
        simple = "Tüm kriterlere eşit önem verildi. Yani sistem, hiçbir kriteri diğerlerinden özellikle daha baskın kabul etmedi."
        academic = "Bu yaklaşım veriden türetilmiş değil, bilinçli bir denge varsayımıdır. Tüm kriterlerin eşit söz hakkına sahip olduğu nötr bir başlangıç senaryosu sunar."
        example = "Örnek: fiyat, kalite ve hız aynı önemde kabul edilirse son puanda hiçbiri özel ayrıcalık kazanmaz."
        action = "Bu modu nötr referans senaryosu olarak kullanın; ardından objektif ağırlıklandırma ile karşılaştırıp hikayenin değişip değişmediğine bakın."
    elif mode == "manual":
        simple = f"Son sıralama, sizin girdiğiniz önceliklere göre şekillendi. En güçlü vurgu <strong>{top_k}</strong> üzerinde kaldı."
        academic = f"Manuel ağırlıklandırma kullanıcı tercihlerini korur ve hesaplama için normalize eder. En yüksek kullanıcı önceliği <strong>{top_k}</strong> kriterinde {top_v:.4f} normalize ağırlıkla görülmektedir."
        example = "Örnek: güvenilirliği maliyetten çok daha önemli görüyorsanız sistem bu tercihi özellikle öne taşır."
        action = "Alan bilgisini bilinçli biçimde yansıtmak istiyorsanız bu mod uygundur; ancak sonuç doğrudan tercih yapınıza bağlı olduğu için duyarlılık analizi ile birlikte okunmalıdır."
    else:
        simple = f"Sistem en çok <strong>{top_k}</strong> kriterini önemsemiş görünüyor. Yani nihai sıralamayı en fazla etkileyen başlık şu anda bu kriter."
        example = f"Örnek: <strong>{top_k}</strong> değeri kötüleşirse önde görünen bir seçenek beklenenden daha hızlı geriye düşebilir."
        if top_v > 0.35:
            academic = f"<strong>{w_method}</strong> yöntemi <strong>{top_k}</strong> kriterine en yüksek ağırlığı atamıştır (w = {top_v:.4f}). Bu yoğun yapı, nihai sıralamanın tek bir baskın sinyale duyarlı olabileceğini düşündürür."
            action = f"Sonucu yorumlarken <strong>{top_k}</strong> kriterini özellikle izleyin ve bu ağırlık değiştiğinde sıralamanın ne kadar oynadığını mutlaka kontrol edin."
        else:
            academic = f"<strong>{w_method}</strong> yöntemi <strong>{top_k}</strong> kriterini en güçlü sinyal olarak işaret etse de genel yoğunlaşma indeksi ≈ {concentration:.2f} düzeyinde kalmıştır. Bu, ağırlık yapısının görece dengeli olduğunu gösterir."
            action = f"Ağırlık profili dengeli görünüyor; yine de ana sürükleyici kriter olarak <strong>{top_k}</strong> üzerinde duyarlılık kontrolü yapmak yerinde olacaktır."
    return _compose_commentary(simple, academic, action, example)

def gen_ranking_commentary(result: Dict[str, Any], alt_names: Dict[str, str] = None) -> str:
    rt = result["ranking"]["table"]
    if rt is None or rt.empty:
        return tt("Sıralama tablosu mevcut değil.", "Ranking table is not available.")
    alt_col  = col_key(rt, "Alternatif", "Alternative")
    score_col = col_key(rt, "Skor", "Score")
    top_alt  = rt.iloc[0][alt_col]
    top_disp = (alt_names or {}).get(str(top_alt), str(top_alt))
    n_alt    = len(rt)
    score_range = float(rt[score_col].max()) - float(rt[score_col].min())
    second_disp = top_disp
    gap = 0.0
    if len(rt) > 1:
        second_alt = rt.iloc[1][alt_col]
        second_disp = (alt_names or {}).get(str(second_alt), str(second_alt))
        gap = float(rt.iloc[0][score_col]) - float(rt.iloc[1][score_col])
    method_simple, method_academic, method_action = _ranking_method_commentary(result, str(top_alt), top_disp)
    if st.session_state.get("ui_lang") == "EN":
        simple = f"Among <strong>{n_alt} options</strong>, the current leader is <strong>{top_disp}</strong>. {method_simple}"
        example = (
            f"A practical example: if the gap between the first and second options is small, the final choice may still depend on field knowledge. Here the top-two gap is only {gap:.4f} between <strong>{top_disp}</strong> and <strong>{second_disp}</strong>."
            if gap < 0.05 else
            f"A practical example: when the leader pulls away clearly, the model is not just saying 'first place' but also 'noticeably ahead'. Here <strong>{top_disp}</strong> separates from <strong>{second_disp}</strong> by {gap:.4f}."
        )
        if score_range < 0.05:
            academic = f"{method_academic} The overall score range remains narrow at {score_range:.4f}, so ranking uncertainty is still material."
        else:
            academic = f"{method_academic} The overall score range is {score_range:.4f}, which indicates meaningful separation across alternatives."
        action = (
            f"{method_action} Because the top-two distance is small, confirm the final choice with Monte Carlo robustness and method comparison."
            if gap < 0.05 else
            f"{method_action} Use the ranking table and method-specific visual to communicate why the leading option is substantively ahead."
        )
        return _compose_commentary(simple, academic, action, example)
    simple = f"<strong>{n_alt} seçenek</strong> içinde şu an öne çıkan alternatif <strong>{top_disp}</strong>. {method_simple}"
    example = (
        f"Pratik bir örnek: ilk iki seçenek birbirine çok yakınsa son karar hâlâ saha bilgisine bağlı kalabilir. Burada <strong>{top_disp}</strong> ile <strong>{second_disp}</strong> arasındaki fark yalnızca {gap:.4f}."
        if gap < 0.05 else
        f"Pratik bir örnek: lider seçenek belirgin biçimde ayrıştığında model sadece 'birinci' demiyor, aynı zamanda 'gözle görülür şekilde önde' diyor. Burada <strong>{top_disp}</strong>, <strong>{second_disp}</strong> karşısında {gap:.4f} fark yaratıyor."
    )
    if score_range < 0.05:
        academic = f"{method_academic} Genel skor aralığı {score_range:.4f} gibi dar bir bantta kaldığı için sıralama belirsizliği hâlâ anlamlı düzeydedir."
    else:
        academic = f"{method_academic} Genel skor aralığı {score_range:.4f} düzeyindedir; bu da alternatifler arasında anlamlı bir ayrışma olduğunu destekler."
    action = (
        f"{method_action} İlk iki seçenek yakın olduğu için nihai kararı vermeden önce Monte Carlo ve yöntem karşılaştırması sonuçlarını birlikte okuyun."
        if gap < 0.05 else
        f"{method_action} Liderliğin neden oluştuğunu göstermek için sıralama tablosu ile yöntem-özel grafiği birlikte kullanın."
    )
    return _compose_commentary(simple, academic, action, example)

def gen_comparison_commentary(result: Dict[str, Any]) -> str:
    comp = result.get("comparison", {})
    if not comp or "spearman_matrix" not in comp:
        return tt("Yöntem karşılaştırması verisi mevcut değil.", "Method comparison data is not available.")
    spdf = comp["spearman_matrix"]
    if not isinstance(spdf, pd.DataFrame) or spdf.shape[0] < 2:
        return tt("Karşılaştırma için en az iki yöntem gereklidir.", "At least two methods are required for comparison.")
    method_col = col_key(spdf, "Yöntem", "Method")
    mat = spdf.set_index(method_col)
    mask = ~np.eye(mat.shape[0], dtype=bool)
    off_diag = mat.where(mask).stack()
    mean_rho = float(off_diag.mean()) if len(off_diag) > 0 else 1.0
    n_meth = mat.shape[0]
    if st.session_state.get("ui_lang") == "EN":
        simple = f"We compared <strong>{n_meth} methods</strong> to see whether they tell a similar story. The average agreement is <strong>ρ = {mean_rho:.3f}</strong>."
        example = (
            "In plain terms: if different methods still point to a similar winner, confidence in the decision grows."
            if mean_rho >= 0.70 else
            "In plain terms: if methods disagree strongly, the result depends on the method family you chose."
        )
        if mean_rho >= 0.85:
            academic = "This is high inter-method agreement, indicating that the ranking structure is stable across methodological families."
            action = "You can report the leading result with stronger confidence, while still attaching the Spearman matrix for transparency."
        elif mean_rho >= 0.70:
            academic = "Agreement is moderate. This usually means the upper ranks are relatively stable while lower positions remain method-sensitive."
            action = "Report the common message across methods and avoid over-claiming fine differences in the lower part of the ranking."
        else:
            academic = "Agreement is low, which indicates methodological sensitivity rather than a single clear ranking narrative."
            action = "Revisit weighting and normalization choices, and give more attention to compromise-oriented methods before drawing a firm conclusion."
        return _compose_commentary(simple, academic, action, example)
    simple = f"<strong>{n_meth} yöntemi</strong> yan yana koyduğumuzda ortalama uyum <strong>ρ = {mean_rho:.3f}</strong> çıktı. Yani farklı yöntemler karar hikayesini ne kadar benzer okuyor, bunu görüyoruz."
    example = (
        "Basitçe: farklı yöntemler yine benzer kazananı gösteriyorsa sonuca güven artar."
        if mean_rho >= 0.70 else
        "Basitçe: yöntemler birbirinden çok farklı konuşuyorsa sonuç, seçtiğiniz yöntem ailesine daha bağımlı hale gelir."
    )
    if mean_rho >= 0.85:
        academic = "Bu yüksek yöntemler-arası uyum, sıralama yapısının yöntem ailesinden bağımsız biçimde kararlı kaldığını gösterir."
        action = "Ana bulguyu daha güçlü savunabilirsiniz; yine de Spearman matrisini şeffaflık için rapora ekleyin."
    elif mean_rho >= 0.70:
        academic = "Uyum orta düzeydedir. Bu durumda özellikle üst sıralar görece kararlı kalırken alt sıralarda yöntem duyarlılığı devam edebilir."
        action = "Yöntemlerin ortak mesajını öne çıkarın; alt sıralardaki ince farkları kesin gerçek gibi sunmayın."
    else:
        academic = "Uyum düşüktür; bu durum tek bir kesin sıralama anlatısından çok metodolojik duyarlılığa işaret eder."
        action = "Ağırlıklandırma ve normalleştirme tercihlerini yeniden gözden geçirin; kesin sonuç vermeden önce uzlaşı odaklı yöntemlere daha fazla dikkat edin."
    return _compose_commentary(simple, academic, action, example)

def gen_mc_commentary(result: Dict[str, Any]) -> str:
    sens = result.get("sensitivity")
    if not sens:
        return tt("Sağlamlık analizi verisi mevcut değil.", "Robustness analysis data is not available.")
    stab = float(sens.get("top_stability", 0.0))
    mc_df = sens.get("monte_carlo_summary")
    n_iter = int(sens.get("n_iterations", 0))
    base_top = str(sens.get("base_top", "—"))
    stable_leader = base_top
    challenger = "—"
    if isinstance(mc_df, pd.DataFrame) and not mc_df.empty:
        mc_alt_col = col_key(mc_df, "Alternatif", "Alternative")
        stable_leader = str(mc_df.iloc[0][mc_alt_col])
        if len(mc_df) > 1:
            challenger = str(mc_df.iloc[1][mc_alt_col])
    local_df = sens.get("local_scenarios")
    sensitive_crit = "—"
    sensitive_delta = "—"
    sensitive_rho = np.nan
    if isinstance(local_df, pd.DataFrame) and not local_df.empty:
        rho_col = col_key(local_df, "SpearmanRho", "SpearmanRho")
        crit_col = col_key(local_df, "Kriter", "Criterion")
        delta_col = col_key(local_df, "AğırlıkDeğişimi", "WeightChange")
        idx = pd.to_numeric(local_df[rho_col], errors="coerce").idxmin()
        if pd.notna(idx):
            sensitive_crit = str(local_df.loc[idx, crit_col])
            sensitive_delta = str(local_df.loc[idx, delta_col])
            sensitive_rho = float(pd.to_numeric(local_df.loc[idx, rho_col], errors="coerce"))

    is_en = st.session_state.get("ui_lang") == "EN"
    did = (
        f"We reran the ranking under {n_iter:,} Monte Carlo weight perturbations and also checked local ±10% / ±20% changes for each criterion."
        if is_en else
        f"Sıralama, {n_iter:,} adet Monte Carlo ağırlık bozma senaryosu altında yeniden çalıştırıldı; ayrıca her kriter için yerel ±%10 / ±%20 değişim etkisi de kontrol edildi."
    )
    why = (
        "The goal is to see whether the reported winner survives reasonable weight uncertainty and to identify which criterion is most likely to destabilize the ranking."
        if is_en else
        "Amaç, raporlanan kazananın makul ağırlık belirsizliği altında korunup korunmadığını ve sıralamayı en çok hangi kriterin bozduğunu görmektir."
    )
    if stab >= 0.80:
        stability_text = tt("Dayanıklılık düzeyi yüksektir.", "Robustness is high.")
        action = (
            "You can present the leader with stronger confidence; still keep the local sensitivity result as supporting evidence in the discussion."
            if is_en else
            "Lider alternatifi daha güçlü güvenle sunabilirsiniz; yine de yerel duyarlılık sonucunu tartışma bölümünde destekleyici kanıt olarak koruyun."
        )
    elif stab >= 0.60:
        stability_text = tt("Dayanıklılık düzeyi ortadır.", "Robustness is moderate.")
        action = (
            "Present the leader with caution and report the challenger together with the most sensitive criterion."
            if is_en else
            "Lider alternatifi temkinli sunun ve en hassas kriterle birlikte yakın rakibi de raporlayın."
        )
    else:
        stability_text = tt("Dayanıklılık düzeyi düşüktür.", "Robustness is low.")
        action = (
            f"Do not rely on a single winner yet; compare {stable_leader} and {challenger}, and revisit the weighting logic."
            if is_en else
            f"Henüz tek bir kazanana güvenmeyin; {stable_leader} ile {challenger} seçeneklerini birlikte karşılaştırın ve ağırlık mantığını yeniden gözden geçirin."
        )

    found = (
        f"{stability_text} The original leader {base_top} remains first in {stab*100:.1f}% of the simulations, while the most frequent simulated leader is {stable_leader}. "
        + (
            f"The main challenger is {challenger}. "
            if challenger != "—" else ""
        )
        + (
            f"The most sensitive local scenario occurs in criterion {sensitive_crit} at {sensitive_delta}, where rank agreement falls to ρ = {sensitive_rho:.3f}. "
            if np.isfinite(sensitive_rho) else ""
        )
        + "This means the reported ranking is being tested both globally and criterion by criterion."
        if is_en else
        f"{stability_text} Başlangıç lideri {base_top}, simülasyonların %{stab*100:.1f} bölümünde birinci kalmıştır; simülasyonlarda en sık öne çıkan lider ise {stable_leader} olmuştur. "
        + (
            f"Başlıca yakın rakip {challenger} olarak görünmektedir. "
            if challenger != "—" else ""
        )
        + (
            f"Yerel olarak en hassas senaryo {sensitive_crit} kriterinde {sensitive_delta} değişimde oluşmuş ve sıra uyumu ρ = {sensitive_rho:.3f} düzeyine kadar düşmüştür. "
            if np.isfinite(sensitive_rho) else ""
        )
        + "Bu tablo, sıralamanın hem genel belirsizlik altında hem de kriter bazında tek tek sınandığını gösterir."
    )
    example = (
        f"For example, if a hospital stays first in 85 out of 100 simulations, we say the ranking is durable. If another option frequently takes over when one criterion moves, the conclusion should be read more cautiously. Here the most pressure comes from {sensitive_crit}."
        if is_en else
        f"Örneğin bir hastane 100 simülasyonun 85'inde yine birinci kalıyorsa sıralama dayanıklıdır deriz. Bir kriter oynadığında başka bir seçenek sık sık öne geçiyorsa sonuç daha temkinli okunmalıdır. Bu çalışmada en büyük baskı {sensitive_crit} kriterinden gelmektedir."
    )
    return _compose_structured_commentary(did, why, found, example, action)

def _render_email_verification_wall(auth_settings: access.AuthSettings) -> None:
    """E-posta doğrulanmamış kullanıcıları durduran ekran (ilk giriş grace sona erdikten sonra)."""
    st.markdown(
        f"""
        <div class="tracking-panel tracking-panel-blue" style="margin-top:1rem; padding:1.1rem 1.15rem;">
            <div class="tracking-title">✉️ {tt("E-posta Dogrulamasi Gerekli", "Email Verification Required")}</div>
            <div class="tracking-subtitle" style="margin-top:0.4rem;">
                {tt(
                    "Hesabınızı doğrulamak için kayıt e-postanıza gelen bağlantıya tıklayın. "
                    "Doğrulama tamamlandıktan sonra tekrar giriş yapın.",
                    "Please click the link sent to your registration email to verify your account. "
                    "After verification is complete, sign in again.",
                )}
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )
    st.info(
        tt(
            "Doğrulama e-postası gelmedi mi? Spam klasörünüzü kontrol edin veya kimlik sağlayıcınızdan "
            "tekrar gönderme isteyin.",
            "Didn't receive the verification email? Check your spam folder or request a resend "
            "from your identity provider.",
        )
    )
    if st.button(tt("🚪 Cikis Yap ve Tekrar Giris Yap", "🚪 Sign Out and Sign In Again"), key="btn_verify_logout"):
        access.logout_user()
        st.rerun()

def _render_name_collection_screen(user: access.CurrentUser) -> None:
    """Kullanıcıdan Ad-Soyad toplar (ilk girişte isim email'e eşleşiyorsa gösterilir)."""
    st.markdown(
        f"""
        <div class="tracking-panel tracking-panel-green" style="margin-top:1rem; padding:1.1rem 1.15rem;">
            <div class="tracking-title">👋 {tt("Hosgeldiniz!", "Welcome!")}</div>
            <div class="tracking-subtitle" style="margin-top:0.4rem;">
                {tt(
                    "MCDM Karar Destek Sistemi'a hoş geldiniz. Devam etmek için ad ve soyadınızı girin.",
                    "Welcome to MCDM Karar Destek Sistemi. Please enter your full name to continue.",
                )}
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )
    with st.form("name_collection_form", border=False):
        entered_name = st.text_input(
            tt("Ad Soyad", "Full Name"),
            placeholder=tt("Örn: Ahmet Yılmaz", "E.g. John Smith"),
            max_chars=80,
        )
        submitted = st.form_submit_button(
            tt("✅ Devam Et", "✅ Continue"),
            type="primary",
            use_container_width=True,
        )
        if submitted:
            clean = entered_name.strip()
            if len(clean) < 3 or "@" in clean:
                st.error(tt("Lütfen geçerli bir ad soyad girin (en az 3 karakter).", "Please enter a valid full name (at least 3 characters)."))
            else:
                access.update_display_name(user.auth_subject, clean, access.get_analytics_settings())
                st.session_state["_mcdm_name_collected"] = True
                st.rerun()
    if st.button(tt("🚪 Cikis Yap", "🚪 Sign Out"), key="btn_name_logout"):
        access.logout_user()
        st.rerun()

def _render_auth_gate(auth_settings: access.AuthSettings) -> None:
    # Hero banner — giriş yapmamış kullanıcılara gösterilir (yıldızlı gece gökyüzü)

    # CSS: Kayıt Ol butonu sarı gradient, Giriş Yap butonu outline
    st.markdown(
        """<style>
        /* Auth gate — Kayıt Ol (col 1) = sarı CTA */
        [data-testid="stMain"] > div > div > div > [data-testid="stVerticalBlock"]
          > [data-testid="stHorizontalBlock"]
          > [data-testid="column"]:first-child .stButton > button {
            background: linear-gradient(135deg, #F59E0B 0%, #F97316 100%) !important;
            color: #111827 !important;
            font-weight: 800 !important;
            font-size: 1.05rem !important;
            padding: 0.8rem 1.2rem !important;
            border: none !important;
            border-radius: 10px !important;
            width: 100% !important;
            box-shadow: 0 4px 20px rgba(249,115,22,0.45) !important;
            letter-spacing: 0.01em !important;
        }
        [data-testid="stMain"] > div > div > div > [data-testid="stVerticalBlock"]
          > [data-testid="stHorizontalBlock"]
          > [data-testid="column"]:first-child .stButton > button:hover {
            background: linear-gradient(135deg, #FBBF24 0%, #FB923C 100%) !important;
            box-shadow: 0 6px 28px rgba(249,115,22,0.6) !important;
            transform: translateY(-1px);
        }
        /* Auth gate — Giriş Yap (col 2) = outline */
        [data-testid="stMain"] > div > div > div > [data-testid="stVerticalBlock"]
          > [data-testid="stHorizontalBlock"]
          > [data-testid="column"]:nth-child(2) .stButton > button {
            background: transparent !important;
            color: #CBD5E1 !important;
            font-weight: 600 !important;
            font-size: 0.95rem !important;
            padding: 0.8rem 1rem !important;
            border: 1.5px solid rgba(203,213,225,0.35) !important;
            border-radius: 10px !important;
            width: 100% !important;
        }
        [data-testid="stMain"] > div > div > div > [data-testid="stVerticalBlock"]
          > [data-testid="stHorizontalBlock"]
          > [data-testid="column"]:nth-child(2) .stButton > button:hover {
            border-color: rgba(203,213,225,0.65) !important;
            color: #F1F5F9 !important;
            background: rgba(255,255,255,0.06) !important;
        }
        </style>""",
        unsafe_allow_html=True,
    )

    _subtitle = tt("Çok Kriterli Karar Destek Sistemi", "Multi-Criteria Decision Support System")
    _desc = tt(
        "Akademik düzeyde ağırlıklandırma ve sıralama analizleri gerçekleştirin. "
        "Verilerinizi yükleyin, yönteminizi seçin ve dakikalar içinde yayına hazır raporlar oluşturun.",
        "Run academic-grade weighting and ranking analyses. "
        "Upload your data, choose your method, and generate publication-ready reports in minutes."
    )
    _dedication = tt(
        "Çocuklarım M. Eymen ve H. Serra'ya İthafen",
        "Dedicated to My Children M. Eymen and H. Serra"
    )
    _cta_text = tt(
        "⬇&nbsp; Ücretsiz hesap oluşturun ve hemen kullanmaya başlayın!",
        "⬇&nbsp; Create your free account and start using it right away!"
    )
    _h1_sub = tt("Karar Destek Sistemi", "Decision Support System")

    _stat_label_total  = tt("Toplam Yöntem", "Total Methods")
    _stat_label_klasik = tt("Klasik Sıralama", "Classical Ranking")
    _stat_label_fuzzy  = tt("Bulanık Yöntem", "Fuzzy Methods")
    _stat_label_obj    = tt("Objektif Ağırlık", "Objective Weighting")
    _stat_label_subj   = tt("Sübjektif Ağırlık", "Subjective Weighting")

    _card_base = (
        "border-radius:12px;padding:0.65rem 1.1rem;text-align:center;"
        "min-width:90px;flex:1;"
    )
    _num_base = "font-size:1.9rem;font-weight:800;line-height:1.1;margin-bottom:0.15rem;"
    _lbl_base = "font-size:0.68rem;letter-spacing:0.07em;text-transform:uppercase;color:#94A3B8;"

    _stats_html = f"""
<div style="display:flex;gap:0.65rem;flex-wrap:wrap;margin-bottom:1.5rem;">
  <div style="{_card_base}background:rgba(249,115,22,0.18);border:1px solid rgba(249,115,22,0.38);">
    <div style="{_num_base}color:#F97316;text-shadow:0 0 18px rgba(249,115,22,0.5);">57</div>
    <div style="{_lbl_base}">{_stat_label_total}</div>
  </div>
  <div style="{_card_base}background:rgba(96,165,250,0.12);border:1px solid rgba(96,165,250,0.28);">
    <div style="{_num_base}color:#60A5FA;">24</div>
    <div style="{_lbl_base}">{_stat_label_klasik}</div>
  </div>
  <div style="{_card_base}background:rgba(167,139,250,0.12);border:1px solid rgba(167,139,250,0.28);">
    <div style="{_num_base}color:#A78BFA;">24</div>
    <div style="{_lbl_base}">{_stat_label_fuzzy}</div>
  </div>
  <div style="{_card_base}background:rgba(52,211,153,0.12);border:1px solid rgba(52,211,153,0.28);">
    <div style="{_num_base}color:#34D399;">9</div>
    <div style="{_lbl_base}">{_stat_label_obj}</div>
  </div>
  <div style="{_card_base}background:rgba(251,146,60,0.12);border:1px solid rgba(251,146,60,0.28);">
    <div style="{_num_base}color:#FB923C;">5</div>
    <div style="{_lbl_base}">{_stat_label_subj}</div>
  </div>
</div>"""

    st.markdown(
        f"""<div style="background-color:#020b18;background-image:radial-gradient(1px 1px at 5% 10%,rgba(255,255,255,.95) 0%,transparent 100%),radial-gradient(1.5px 1.5px at 12% 22%,rgba(255,255,255,.8) 0%,transparent 100%),radial-gradient(1px 1px at 18% 5%,rgba(255,255,255,.9) 0%,transparent 100%),radial-gradient(2px 2px at 25% 35%,rgba(255,255,255,.6) 0%,transparent 100%),radial-gradient(1px 1px at 30% 15%,rgba(200,220,255,.9) 0%,transparent 100%),radial-gradient(1.5px 1.5px at 38% 48%,rgba(255,255,255,.7) 0%,transparent 100%),radial-gradient(1px 1px at 43% 8%,rgba(255,255,255,.85) 0%,transparent 100%),radial-gradient(2px 2px at 50% 28%,rgba(180,200,255,.8) 0%,transparent 100%),radial-gradient(1px 1px at 55% 60%,rgba(255,255,255,.7) 0%,transparent 100%),radial-gradient(1.5px 1.5px at 62% 18%,rgba(255,255,255,.9) 0%,transparent 100%),radial-gradient(1px 1px at 68% 42%,rgba(255,255,255,.75) 0%,transparent 100%),radial-gradient(2px 2px at 74% 12%,rgba(200,215,255,.85) 0%,transparent 100%),radial-gradient(1px 1px at 80% 55%,rgba(255,255,255,.8) 0%,transparent 100%),radial-gradient(1.5px 1.5px at 85% 30%,rgba(255,255,255,.65) 0%,transparent 100%),radial-gradient(1px 1px at 90% 8%,rgba(255,255,255,.9) 0%,transparent 100%),radial-gradient(2px 2px at 95% 45%,rgba(180,210,255,.7) 0%,transparent 100%),radial-gradient(1px 1px at 8% 72%,rgba(255,255,255,.6) 0%,transparent 100%),radial-gradient(1.5px 1.5px at 15% 85%,rgba(255,255,255,.75) 0%,transparent 100%),radial-gradient(1px 1px at 22% 65%,rgba(255,255,255,.85) 0%,transparent 100%),radial-gradient(1px 1px at 35% 78%,rgba(200,220,255,.7) 0%,transparent 100%),radial-gradient(2px 2px at 48% 88%,rgba(255,255,255,.6) 0%,transparent 100%),radial-gradient(1px 1px at 58% 75%,rgba(255,255,255,.8) 0%,transparent 100%),radial-gradient(1.5px 1.5px at 70% 82%,rgba(255,255,255,.7) 0%,transparent 100%),radial-gradient(1px 1px at 78% 68%,rgba(180,200,255,.85) 0%,transparent 100%),radial-gradient(1px 1px at 88% 90%,rgba(255,255,255,.65) 0%,transparent 100%),radial-gradient(1.5px 1.5px at 93% 72%,rgba(255,255,255,.8) 0%,transparent 100%),radial-gradient(ellipse at 70% 20%,rgba(20,50,100,.5) 0%,transparent 55%),radial-gradient(ellipse at 15% 60%,rgba(10,30,70,.4) 0%,transparent 45%),linear-gradient(180deg,#020810 0%,#040e1f 40%,#061228 100%);border-radius:16px;padding:2.8rem 2.5rem 2rem 2.5rem;margin-bottom:0.5rem;position:relative;overflow:hidden;">
<div style="font-size:0.78rem;font-weight:500;letter-spacing:0.12em;color:#60A5FA;text-transform:uppercase;margin-bottom:0.25rem;">✦ &nbsp; {_subtitle}</div>
<div style="font-size:0.92rem;color:#94A3B8;font-style:italic;margin-bottom:0.7rem;letter-spacing:0.02em;">Prof. Dr. Ömer Faruk Rençber</div>
<h1 style="font-size:2.5rem;font-weight:800;margin:0 0 0.7rem 0;line-height:1.15;"><span style="color:#F97316;text-shadow:0 0 28px rgba(249,115,22,0.45);">MCDM</span><span style="display:inline-block;margin-left:0.38rem;font-size:0.26em;font-weight:700;letter-spacing:0.08em;text-transform:uppercase;color:#FFFFFF;text-shadow:0 0 24px rgba(96,165,250,0.2);vertical-align:middle;">{_h1_sub}</span></h1>
<p style="font-size:1.02rem;color:#CBD5E1;max-width:680px;line-height:1.7;margin:0 0 1.4rem 0;">{_desc}</p>
{_stats_html}
<div style="border-top:1px solid rgba(255,255,255,0.08);padding-top:1rem;display:flex;align-items:center;justify-content:space-between;flex-wrap:wrap;gap:0.6rem;">
  <span style="font-size:0.95rem;font-weight:700;color:#FCD34D;letter-spacing:0.01em;">{_cta_text}</span>
  <span style="font-size:0.75rem;color:#475569;font-style:italic;">✦ {_dedication}</span>
</div>
</div>""",
        unsafe_allow_html=True,
    )

    # Giriş / Kayıt butonları — col1: sarı CTA, col2: outline login
    _btn_col1, _btn_col2, _btn_spacer = st.columns([1.6, 1, 0.8])
    with _btn_col1:
        st.button(
            tt("🚀 Ücretsiz Kayıt Ol · Hemen Başla", "🚀 Sign Up Free · Get Started Now"),
            on_click=access.login_user,
            args=[auth_settings.signup_provider],
            use_container_width=True,
            key="btn_auth_signup_gate",
        )
    with _btn_col2:
        st.button(
            tt("🔐 Giriş Yap", "🔐 Sign In"),
            on_click=access.login_user,
            args=[auth_settings.provider],
            use_container_width=True,
            key="btn_auth_login_gate",
        )

    # YouTube & Instagram
    _vid = tt("Tanıtım Videosu", "Demo Video")
    st.markdown(
        f'<div style="margin-top:0.8rem;display:flex;flex-wrap:wrap;gap:0.5rem;align-items:center;">'
        f'<a href="https://youtu.be/jp4oih6_Nec" target="_blank" style="display:inline-block;background:#DC2626;color:#FFFFFF;text-decoration:none;border-radius:8px;padding:0.45rem 0.9rem;font-size:0.82rem;font-weight:600;">🎥 {_vid}</a>'
        f'<a href="https://www.instagram.com/mcdm_dss/" target="_blank" style="display:inline-block;background:#C13584;color:#FFFFFF;text-decoration:none;border-radius:8px;padding:0.45rem 0.9rem;font-size:0.82rem;font-weight:600;">📸 @mcdm_dss</a>'
        f'</div>',
        unsafe_allow_html=True,
    )
    st.caption(tt(auth_settings.privacy_notice_tr, auth_settings.privacy_notice_en))
    st.stop()

def _render_user_session_card(auth_settings: access.AuthSettings, current_user: access.CurrentUser) -> None:
    with st.expander(tt("👤 Üye Bilgileri", "👤 Member Info"), expanded=not current_user.is_logged_in):
        if not current_user.is_logged_in:
            st.markdown(
                f"""
                <div class="tracking-panel tracking-panel-blue" style="margin-bottom:0.5rem; padding:0.75rem 1rem;">
                    <div class="tracking-title">{tt("Giris Yapilmadi", "Not Signed In")}</div>
                </div>
                """,
                unsafe_allow_html=True,
            )
            if auth_settings.enabled and auth_settings.provider:
                if st.button(tt("🔐 Giris Yap", "🔐 Sign In"), use_container_width=True, key="btn_sidebar_login"):
                    access.login_user(auth_settings.provider)
                if auth_settings.signup_provider and auth_settings.signup_provider != auth_settings.provider:
                    if st.button(tt("✉️ Kayit Ol", "✉️ Sign Up"), use_container_width=True, key="btn_sidebar_signup"):
                        access.login_user(auth_settings.signup_provider)
            return

        role_label = tt("Yonetici", "Admin") if access.is_admin_email(current_user.email, auth_settings) else tt("Uye", "Member")
        st.markdown(
            f"""
            <div class="tracking-panel tracking-panel-green tracking-panel-compact">
                <div class="tracking-title">{tt("Oturum", "Session")}</div>
                <div class="tracking-grid">
                    <div class="tracking-card tracking-card-green">
                        <div class="tracking-label">{tt("Kullanici", "User")}</div>
                        <div class="tracking-value">{html.escape(current_user.name or current_user.email)}</div>
                    </div>
                    <div class="tracking-card tracking-card-slate">
                        <div class="tracking-label">{tt("E-posta", "Email")}</div>
                        <div class="tracking-value">{html.escape(current_user.email)}</div>
                    </div>
                    <div class="tracking-card tracking-card-blue">
                        <div class="tracking-label">{tt("Rol", "Role")}</div>
                        <div class="tracking-value">{role_label}</div>
                    </div>
                </div>
            </div>
            """,
            unsafe_allow_html=True,
        )
        if st.button(tt("🚪 Cikis Yap", "🚪 Log Out"), use_container_width=True, key="btn_auth_logout_sidebar"):
            access.track_event("logout_clicked", {"source": "sidebar"})
            access.logout_user()
            st.rerun()

def _render_admin_usage_panel(auth_settings: access.AuthSettings, current_user: access.CurrentUser) -> None:
    if not current_user.is_logged_in or not access.is_admin_email(current_user.email, auth_settings):
        return
    with st.expander(tt("📈 Uyelik ve Kullanim Ozeti", "📈 Membership and Usage Summary"), expanded=False):
        summary = access.get_admin_summary(limit=25)
        if not summary:
            st.info(
                tt(
                    "Kullanim veritabani baglantisi henuz hazir degil. Secrets ve veritabani ayarlarini kontrol edin.",
                    "The usage database connection is not ready yet. Check your secrets and database settings.",
                )
            )
            _tracking_error = access.tracking_error()
            if _tracking_error:
                st.caption(tt(f"Son hata: {_tracking_error}", f"Latest error: {_tracking_error}"))
            return

        m1, m2 = st.columns(2)
        m3, m4 = st.columns(2)
        with m1:
            st.metric(tt("Toplam uye", "Total members"), f"{summary['total_users']:,}")
        with m2:
            st.metric(tt("Toplam oturum", "Total sessions"), f"{summary['total_sessions']:,}")
        with m3:
            st.metric(tt("Tamamlanan analiz", "Completed analyses"), f"{summary['total_analyses']:,}")
        with m4:
            st.metric(tt("Son 7 gun aktif uye", "Active users in 7 days"), f"{summary['last_7d_users']:,}")

        recent_df = pd.DataFrame(summary.get("recent_users") or [])
        if not recent_df.empty:
            recent_df = recent_df.rename(
                columns={
                    "user_email": tt("E-posta", "Email"),
                    "last_seen_at": tt("Son gorulme", "Last seen"),
                    "analyses_completed": tt("Analiz", "Analyses"),
                    "session_count": tt("Oturum", "Sessions"),
                }
            )
            st.markdown(f"**{tt('Son kullanicilar', 'Recent users')}**")
            render_table(recent_df)

        events_df = pd.DataFrame(summary.get("events_by_type") or [])
        if not events_df.empty:
            events_df = events_df.rename(
                columns={
                    "event_type": tt("Olay", "Event"),
                    "total": tt("Toplam", "Total"),
                }
            )
            st.markdown(f"**{tt('Olay dagilimi', 'Event distribution')}**")
            render_table(events_df)

        _tracking_error = access.tracking_error()
        if _tracking_error:
            st.caption(tt(f"Son izleme hatasi: {_tracking_error}", f"Latest tracking error: {_tracking_error}"))

# ---------------------------------------------------------
# GLOBAL BANNER
# ---------------------------------------------------------
_auth_settings = access.get_auth_settings()
_analytics_settings = access.get_analytics_settings()
_current_user = access.get_current_user()

# SOL PANEL — Auth durumuna göre içerik
# ---------------------------------------------------------
missing_strategy, clip_outliers = "Sil", False  # defaults (logged-out veya data yok)

with st.sidebar:
    _logo_path = APP_DIR / "logo.png"
    _sidebar_logo_b64 = _encoded_image_b64(str(_logo_path))
    if _sidebar_logo_b64:
        st.markdown(
            f"""
            <div class="sidebar-brand">
                <img src="data:image/png;base64,{_sidebar_logo_b64}" class="sidebar-brand-logo" alt="MCDM logo">
                <div class="sidebar-brand-name">Prof. Dr. Ömer Faruk Rençber</div>
            </div>
            """,
            unsafe_allow_html=True,
        )
    else:
        st.markdown(
            """
            <div class="sidebar-brand">
                <div class="sidebar-brand-name">Prof. Dr. Ömer Faruk Rençber</div>
            </div>
            """,
            unsafe_allow_html=True,
        )

    _lang_idx = 1 if st.session_state.get("ui_lang") == "EN" else 0
    _lang_pick = st.radio(
        "Language",
        ["Türkçe", "English"],
        index=_lang_idx,
        horizontal=True,
        label_visibility="collapsed",
    )
    _new_lang = "EN" if _lang_pick == "English" else "TR"
    if _new_lang != st.session_state.get("ui_lang"):
        st.session_state["ui_lang"] = _new_lang
        st.rerun()

    st.markdown("<hr style='margin: 0.5rem 0'>", unsafe_allow_html=True)

    if not _current_user.is_logged_in:
        # ── Giriş yapılmamış: sadece kullanıcı kartı + yardım ──
        _render_user_session_card(_auth_settings, _current_user)
        st.markdown("<hr style='margin: 0.5rem 0'>", unsafe_allow_html=True)
        with st.expander(tt("📘 Uygulama Rehberi", "📘 User Guide"), expanded=True):
            st.markdown(
                f"<p style='font-size:0.78rem; line-height:1.55; color:#2C2C2C; margin:0;'>"
                f"{tt('1️⃣ Giriş yap.<br>2️⃣ Veri yükle (Excel, CSV, SPSS).<br>3️⃣ Yöntemini seç.<br>4️⃣ Analiz et ve raporu indir.',
                      '1️⃣ Sign in.<br>2️⃣ Upload data (Excel, CSV, SPSS).<br>3️⃣ Choose your method.<br>4️⃣ Run analysis and download report.')}"
                f"</p>",
                unsafe_allow_html=True,
            )
        with st.expander(tt("🧭 Neden MCDM Karar Destek Sistemi?", "🧭 Why MCDM Karar Destek Sistemi?"), expanded=False):
            st.markdown(
                f"<p style='font-size:0.78rem; line-height:1.55; color:#2C2C2C; margin:0;'>"
                f"{tt('✅ 24+ klasik ve bulanık sıralama yöntemi<br>✅ 9 objektif ağırlık yöntemi<br>✅ Monte Carlo duyarlılık analizi<br>✅ SSCI akademik rapor (Word + Excel)<br>✅ Türkçe ve İngilizce arayüz',
                      '✅ 24+ classical and fuzzy ranking methods<br>✅ 9 objective weighting methods<br>✅ Monte Carlo sensitivity analysis<br>✅ SSCI academic report (Word + Excel)<br>✅ Turkish and English interface')}"
                f"</p>",
                unsafe_allow_html=True,
            )
        st.markdown(
            f"<div class='sidebar-footer'>"
            "Prof. Dr. Ömer Faruk Rençber"
            "<a class='sidebar-footer-mail' href='mailto:dr.ofrencber@gaziantep.edu.tr'>dr.ofrencber@gaziantep.edu.tr</a>"
            f"<div style='margin-top:0.3rem;font-size:0.60rem;font-weight:500;line-height:1.35;'>{_SOFTWARE_CITATION}</div>"
            "</div>",
            unsafe_allow_html=True,
        )

    else:
        # ── Giriş yapılmış: tam sidebar ──

        _render_user_session_card(_auth_settings, _current_user)
        _render_admin_usage_panel(_auth_settings, _current_user)

        with st.expander(tt("📘 Uygulama Rehberi", "📘 User Guide"), expanded=False):
            st.markdown(
                f"<p style='font-size:0.78rem; line-height:1.55; color:#2C2C2C; margin:0;'>"
                f"{tt('1️⃣ Sağ paneldeki veri girişi alanından dosya yükleyin, örnek veriyi kullanın veya manuel tabloya geçin.<br>'
                      '2️⃣ Sağ panelde analiz amacını seçin.<br>'
                      '3️⃣ Kriterlerin yönünü (fayda/maliyet) doğrulayın.<br>'
                      '4️⃣ Ağırlıklandırma ve sıralama yöntemlerini belirleyin.<br>'
                      '5️⃣ &#34;Analiz Zamanı&#34; butonuna basın.<br>'
                      '6️⃣ Sonuçları sekmeler arasında inceleyin ve raporu indirin.',
                      '1️⃣ In the right-side data input area, upload a file, use sample data, or switch to manual entry.<br>'
                      '2️⃣ Select your analysis objective in the right panel.<br>'
                      '3️⃣ Confirm criterion directions (benefit/cost).<br>'
                      '4️⃣ Select weighting and ranking methods.<br>'
                      '5️⃣ Click &#34;Run Analysis&#34;.<br>'
                      '6️⃣ Explore results across tabs and download the report.')}"
                f"</p>",
                unsafe_allow_html=True,
            )

        with st.expander(tt("🧭 Metodolojik Yardım", "🧭 Methodological Help"), expanded=False):
            st.markdown(
                f"<p style='font-size:0.78rem; line-height:1.55; color:#2C2C2C; margin:0;'>"
                f"{tt('<b>Objektif Ağırlık Yöntemleri:</b> Kriter önemini verinin kendi yapısından türetir; araştırmacı müdahalesi gerekmez.<br><br>'
                      '<b>Klasik ÇKKV:</b> Kesin sayısal veriler için tasarlanmış; TOPSIS, VIKOR, EDAS vb. 17 yöntem.<br><br>'
                      '<b>Fuzzy ÇKKV:</b> Ölçüm belirsizliği varsa tercih edilir; üçgensel bulanık sayılar kullanır.<br><br>'
                      '<b>Monte Carlo:</b> Ağırlıklar rastgele bozularak sıralamanın ne kadar kararlı olduğu test edilir.',
                      '<b>Objective Weighting:</b> Derives criterion importance from data structure; no researcher intervention needed.<br><br>'
                      '<b>Classical MCDM:</b> Designed for crisp numerical data; 17 methods incl. TOPSIS, VIKOR, EDAS.<br><br>'
                      '<b>Fuzzy MCDM:</b> Preferred when measurement uncertainty exists; uses triangular fuzzy numbers.<br><br>'
                      '<b>Monte Carlo:</b> Weights are randomly perturbed to test how stable the ranking is.')}"
                f"</p>",
                unsafe_allow_html=True,
            )
            # Karar ağacı (decision tree) — Metodolojik Yardım içinde
            _dt_text = tt(
                "<br><b>📊 AĞIRLIK YÖNTEMİ SEÇİM REHBERİ</b><br>"
                "• Uzman görüşü yok → <b>Objektif</b> (Entropy, CRITIC, MEREC)<br>"
                "• Korelasyon yüksek → <b>CRITIC</b> veya <b>PCA</b><br>"
                "• Kriter çıkarılma etkisi → <b>MEREC</b><br>"
                "• Hibrit → <b>IDOCRIW</b><br>"
                "• Uzman + az kriter → <b>AHP</b><br>"
                "• Uzman + çok kriter → <b>BWM</b><br><br>"
                "<b>🏆 SIRALAMA YÖNTEMİ SEÇİM REHBERİ</b><br>"
                "• Genel amaçlı → <b>TOPSIS</b> veya <b>EDAS</b><br>"
                "• Uzlaşı → <b>VIKOR</b><br>"
                "• Rank reversal direnci → <b>SPOTIS</b> veya <b>RAFSI</b><br>"
                "• Çok bileşenli → <b>MULTIMOORA</b><br>"
                "• Belirsizlik → Fuzzy varyantları<br>",
                "<br><b>📊 WEIGHTING METHOD GUIDE</b><br>"
                "• No expert → <b>Objective</b> (Entropy, CRITIC, MEREC)<br>"
                "• High correlation → <b>CRITIC</b> or <b>PCA</b><br>"
                "• Removal effect → <b>MEREC</b><br>"
                "• Hybrid → <b>IDOCRIW</b><br>"
                "• Expert + few criteria → <b>AHP</b><br>"
                "• Expert + many criteria → <b>BWM</b><br><br>"
                "<b>🏆 RANKING METHOD GUIDE</b><br>"
                "• General purpose → <b>TOPSIS</b> or <b>EDAS</b><br>"
                "• Compromise → <b>VIKOR</b><br>"
                "• Rank reversal resistant → <b>SPOTIS</b> or <b>RAFSI</b><br>"
                "• Multi-component → <b>MULTIMOORA</b><br>"
                "• Uncertainty → Fuzzy variants<br>",
            )
            st.markdown(f"<p style='font-size:0.73rem; line-height:1.55; color:#2C2C2C; margin:0;'>{_dt_text}</p>", unsafe_allow_html=True)

        # ── Terimler Sözlüğü ──
        _show_glossary = st.checkbox(tt("📖 Terimler Sözlüğü", "📖 Glossary"), key="show_glossary")
        if _show_glossary:
            _gl = tt(
                "<b>ÇKKV:</b> Çok Kriterli Karar Verme.<br>"
                "<b>Karar Matrisi:</b> Satırlar alternatif, sütunlar kriter.<br>"
                "<b>Fayda (Max):</b> Büyük=iyi. <b>Maliyet (Min):</b> Küçük=iyi.<br>"
                "<b>Normalizasyon:</b> Farklı ölçekleri karşılaştırılabilir yapma.<br>"
                "<b>Ağırlık (w):</b> Kriter göreli önemi, toplamı 1.<br>"
                "<b>TFN:</b> Üçgensel bulanık sayı (alt, orta, üst).<br>"
                "<b>Spread:</b> Bulanık bant genişliği; 0.10=±%10.<br>"
                "<b>PIS/NIS:</b> Pozitif/Negatif ideal çözüm.<br>"
                "<b>CC:</b> Yakınlık katsayısı (0-1).<br>"
                "<b>Spearman ρ:</b> Sıra uyumu; 1.0=tam uyum.<br>"
                "<b>Monte Carlo:</b> Rastgele bozulmayla kararlılık testi.<br>"
                "<b>Kararlılık:</b> MC'de liderin birinciliği koruma %'si.<br>"
                "<b>Pertürbasyon:</b> Ağırlığa ±%10/±%20 değişim.<br>"
                "<b>Rank Reversal:</b> Yeni alternatif ekleyince sıra değişmesi.<br>"
                "<b>GAIA:</b> PROMETHEE sonuçlarının 2D görselleştirmesi.<br>"
                "<b>Borda:</b> MULTIMOORA'da 3 bileşen sıra toplamı.<br>"
                "<b>BAA:</b> MABAC'ta sınır yaklaşım alanı.<br>",
                "<b>MCDM:</b> Multi-Criteria Decision Making.<br>"
                "<b>Decision Matrix:</b> Rows=alternatives, columns=criteria.<br>"
                "<b>Benefit (Max):</b> Higher=better. <b>Cost (Min):</b> Lower=better.<br>"
                "<b>Normalization:</b> Scaling to comparable range.<br>"
                "<b>Weight (w):</b> Criterion importance, sums to 1.<br>"
                "<b>TFN:</b> Triangular fuzzy number (lower, mid, upper).<br>"
                "<b>Spread:</b> Fuzzy bandwidth; 0.10=±10%.<br>"
                "<b>PIS/NIS:</b> Positive/Negative ideal solution.<br>"
                "<b>CC:</b> Closeness coefficient (0-1).<br>"
                "<b>Spearman ρ:</b> Rank agreement; 1.0=perfect.<br>"
                "<b>Monte Carlo:</b> Stability test via random perturbation.<br>"
                "<b>Stability:</b> % of MC runs where leader stays #1.<br>"
                "<b>Perturbation:</b> ±10%/±20% weight change.<br>"
                "<b>Rank Reversal:</b> Rankings change when adding alternatives.<br>"
                "<b>GAIA:</b> 2D projection of PROMETHEE results.<br>"
                "<b>Borda:</b> Sum of 3 component ranks in MULTIMOORA.<br>"
                "<b>BAA:</b> Border Approximation Area in MABAC.<br>",
            )
            st.markdown(f"<div style='font-size:0.72rem; line-height:1.55; color:#2C2C2C;'>{_gl}</div>", unsafe_allow_html=True)

        # ── SSS / FAQ ──
        _show_faq = st.checkbox(tt("❓ Sıkça Sorulan Sorular", "❓ FAQ"), key="show_faq")
        if _show_faq:
            _faq = tt(
                "<b>S: Hangi ağırlık yöntemini seçmeliyim?</b><br>Uzman yok → Entropy/CRITIC/MEREC. Sistem veri ön tanısında öneri sunar.<br><br>"
                "<b>S: Hangi sıralama yöntemi?</b><br>Genel: TOPSIS/EDAS. Uzlaşı: VIKOR. Belirsizlik: Fuzzy. Birden fazla seçip karşılaştırın.<br><br>"
                "<b>S: İki yöntem farklı lider gösteriyorsa?</b><br>Normal. Spearman ρ≥0.85 → uyumlu. ρ<0.70 → kriter yönlerini kontrol edin.<br><br>"
                "<b>S: Kararlılık <%60 ise?</b><br>Kriter azaltın, korelasyonlu kriterleri birleştirin, farklı ağırlık deneyin.<br><br>"
                "<b>S: Excel mi, Rapor mu, Makale mi?</b><br>İş raporu→Excel. Anlamak→APA Rapor. Tez/makale→IMRAD.<br><br>"
                "<b>S: Fuzzy ne zaman?</b><br>Anket/tahmin verisi, ölçüm hatası varsa. Spread=0.10 genellikle yeterli.<br><br>"
                "<b>S: Veri limiti?</b><br>50.000 satır, 120 sütun, 20 MB.<br><br>"
                "<b>S: Panel veri ne zaman?</b><br>Aynı alternatifleri yıllar arası karşılaştırmak için. Yıl/dönem sütunu gerekli.<br>",
                "<b>Q: Which weighting method?</b><br>No expert→Entropy/CRITIC/MEREC. System suggests in Preliminary Review.<br><br>"
                "<b>Q: Which ranking method?</b><br>General: TOPSIS/EDAS. Compromise: VIKOR. Uncertainty: Fuzzy. Select multiple to compare.<br><br>"
                "<b>Q: Two methods show different leaders?</b><br>Normal. Spearman ρ≥0.85→agreement. ρ<0.70→check criterion directions.<br><br>"
                "<b>Q: Stability <60%?</b><br>Reduce criteria, merge correlated ones, try different weighting.<br><br>"
                "<b>Q: Excel, Report, or Article?</b><br>Business→Excel. Understanding→APA Report. Thesis/paper→IMRAD.<br><br>"
                "<b>Q: When to use Fuzzy?</b><br>Survey/estimate data, measurement error. Spread=0.10 usually sufficient.<br><br>"
                "<b>Q: Data limits?</b><br>50,000 rows, 120 columns, 20 MB.<br><br>"
                "<b>Q: Panel data when?</b><br>Comparing same alternatives across years. Needs year/period column.<br>",
            )
            st.markdown(f"<div style='font-size:0.72rem; line-height:1.55; color:#2C2C2C;'>{_faq}</div>", unsafe_allow_html=True)

        is_data_loaded = st.session_state.get("raw_data") is not None

        if st.button(tt("🔄 Yeni Analize Başla (Sıfırla)", "🔄 Start New Analysis (Reset)"), use_container_width=True):
            access.track_event(
                "analysis_reset",
                {
                    "had_data": bool(is_data_loaded),
                    "had_result": st.session_state.get("analysis_result") is not None,
                },
            )
            reset_all()
            st.rerun()
        st.caption(
            tt(
                "Veri yükleme alanı sağ panelde yer alır. Veri geldikten sonra bu bölüm otomatik olarak kapanır.",
                "The data input area is in the right panel. Once data is loaded, that section collapses automatically.",
            )
        )

        if is_data_loaded:
            with st.expander(f"🧹 {tt('Veri Ön İşleme', 'Data Preprocessing')}", expanded=False):
                _method_options = [
                    tt("Medyan", "Median"),
                    tt("Ortalama", "Mean"),
                    tt("Interpolasyon", "Interpolation"),
                    tt("Sıfır", "Zero"),
                ]
                _saved_missing = str(st.session_state.get("missing_strategy_saved", "Sil") or "Sil")
                _saved_to_widget = {
                    "Medyan": tt("Medyan", "Median"),
                    "Median": tt("Medyan", "Median"),
                    "Ortalama": tt("Ortalama", "Mean"),
                    "Mean": tt("Ortalama", "Mean"),
                    "Interpolasyon": tt("Interpolasyon", "Interpolation"),
                    "Interpolation": tt("Interpolasyon", "Interpolation"),
                    "Sıfır": tt("Sıfır", "Zero"),
                    "Zero": tt("Sıfır", "Zero"),
                }
                _default_method = _saved_to_widget.get(_saved_missing, _method_options[0])
                if st.session_state.get("impute_method_select") not in _method_options:
                    st.session_state["impute_method_select"] = _default_method
                if "cb_clip_outliers" not in st.session_state:
                    st.session_state["cb_clip_outliers"] = bool(st.session_state.get("clip_outliers_saved", False))

                impute_checked = st.checkbox(
                    tt("Eksik Veri Tamamla", "Impute Missing Values"),
                    value=st.session_state["impute_mode_open"],
                    key="cb_impute_mode",
                )
                if impute_checked != st.session_state["impute_mode_open"]:
                    st.session_state["impute_mode_open"] = impute_checked
                    st.rerun()

                if impute_checked:
                    _impute_method = st.selectbox(
                        tt("Tamamlama yöntemi", "Imputation method"),
                        _method_options,
                        key="impute_method_select",
                    )
                else:
                    _impute_method = st.session_state.get("impute_method_select", _default_method)

                _clip_outliers_selected = st.checkbox(
                    tt("Aykırı Değerleri (Outlier) Temizle", "Clean Outliers"),
                    value=bool(st.session_state.get("cb_clip_outliers", st.session_state.get("clip_outliers_saved", False))),
                    key="cb_clip_outliers",
                )
                st.caption(tt("Yalnız sayısal sütunlara uygulanır.", "Applied to numeric columns only."))
                if st.button(tt("✅ Ön İşlemeyi Uygula", "✅ Apply Preprocessing"), use_container_width=True, key="btn_apply_preprocessing"):
                    st.session_state["missing_strategy_saved"] = _impute_method if impute_checked else "Sil"
                    st.session_state["clip_outliers_saved"] = bool(_clip_outliers_selected)
                    st.session_state["prep_done"] = False
                    st.session_state["analysis_result"] = None
                    st.session_state["panel_results"] = None
                    st.session_state["report_docx"] = None
                    st.session_state["download_blob_cache"] = {}
                    st.session_state["download_blob_sig"] = None
                    st.session_state["clean_data"] = None
                    st.rerun()

                _active_missing = st.session_state.get("missing_strategy_saved", "Sil")
                _missing_label = _active_missing if _active_missing not in {"Sil", "Drop"} else tt("Kapalı / Sil", "Off / Drop")
                st.caption(tt(f"Aktif eksik veri ayarı: {_missing_label}", f"Active missing-data setting: {_missing_label}"))
                st.caption(
                    tt(
                        f"Aktif uç değer temizliği: {'Açık' if st.session_state.get('clip_outliers_saved', False) else 'Kapalı'}",
                        f"Active outlier cleaning: {'On' if st.session_state.get('clip_outliers_saved', False) else 'Off'}",
                    )
                )
                missing_strategy = st.session_state.get("missing_strategy_saved", "Sil")
                clip_outliers = bool(st.session_state.get("clip_outliers_saved", False))

        _step_data_done = is_data_loaded
        _step_prep_done = bool(st.session_state.get("prep_done"))
        _step_result_done = st.session_state.get("analysis_result") is not None
        _step_method_active = _step_data_done and _step_prep_done and not _step_result_done
        st.markdown("<div style='margin-top:0.5rem;'></div>", unsafe_allow_html=True)
        st.markdown(
            f"""
            <div class="stepper-wrap">
                <div class="step-item"><span class="step-dot {'step-done' if _step_data_done else 'step-pending'}"></span> {tt('1. Adım', 'Step 1')} · {tt('Veri girişi', 'Data input')}{' ✅' if _step_data_done else ''}</div>
                <div class="step-item"><span class="step-dot {'step-done' if _step_prep_done else ('step-active' if _step_data_done else 'step-pending')}"></span> {tt('2. Adım', 'Step 2')} · {tt('Kriter doğrulama', 'Criteria validation')}{' ✅' if _step_prep_done else ''}</div>
                <div class="step-item"><span class="step-dot {'step-done' if _step_result_done else ('step-active' if _step_method_active else 'step-pending')}"></span> {tt('3. Adım', 'Step 3')} · {tt('Yöntem ve analiz', 'Method and analysis')}{' ✅' if _step_result_done else ''}</div>
                <div class="step-item"><span class="step-dot {'step-done' if _step_result_done else 'step-pending'}"></span> {tt('4. Adım', 'Step 4')} · {tt('Sonuç ve rapor', 'Results and report')}{' ✅' if _step_result_done else ''}</div>
            </div>
            """,
            unsafe_allow_html=True,
        )

        st.markdown(
            f"<div class='sidebar-footer'>"
            "Prof. Dr. Ömer Faruk Rençber"
            "<a class='sidebar-footer-mail' href='mailto:dr.ofrencber@gaziantep.edu.tr'>dr.ofrencber@gaziantep.edu.tr</a>"
            f"<div style='margin-top:0.3rem;font-size:0.60rem;font-weight:500;line-height:1.35;'>{_SOFTWARE_CITATION}</div>"
            "</div>",
            unsafe_allow_html=True,
        )

# ---------------------------------------------------------

# AUTH GATE — Giriş zorunluysa login ekranı
# ---------------------------------------------------------
if _auth_settings.require_login and not _current_user.is_logged_in:
    _render_auth_gate(_auth_settings)

if _current_user.is_logged_in:
    access.bootstrap_user_session(_current_user, _analytics_settings)

# --- Hareketsizlik zaman aşımı (4 saat) ---
if _current_user.is_logged_in:
    _now_ts = time.time()
    _last_activity_ts = st.session_state.get("_mcdm_last_activity_at", _now_ts)
    if (_now_ts - _last_activity_ts) > SESSION_INACTIVITY_SECONDS:
        access.track_event("inactivity_logout", {"idle_seconds": int(_now_ts - _last_activity_ts)})
        access.logout_user()
        st.rerun()
    st.session_state["_mcdm_last_activity_at"] = _now_ts

# --- E-posta doğrulama kontrolü ---
if _current_user.is_logged_in and not _current_user.email_verified:
    _grace_given = st.session_state.get("_mcdm_email_grace_given", False)
    if not _grace_given:
        _grace_granted = access.consume_email_verification_grace(
            _current_user.auth_subject,
            _analytics_settings,
        )
        if _grace_granted:
            st.session_state["_mcdm_email_grace_given"] = True
            st.warning(
                tt(
                    "✉️ E-posta adresiniz henüz doğrulanmamış. Lütfen kayıt e-postanızdaki bağlantıya tıklayın. "
                    "Bu oturumda devam edebilirsiniz; ancak sonraki girişte doğrulama zorunlu olacaktır.",
                    "✉️ Your email address is not yet verified. Please click the link in your registration email. "
                    "You may continue this session, but verification will be required on your next sign-in.",
                )
            )
        else:
            _render_email_verification_wall(_auth_settings)
            st.stop()

# --- Ad-Soyad toplama (gerekirse) ---
if _current_user.is_logged_in:
    _needs_name = (
        not _current_user.name
        or _current_user.name.lower() == _current_user.email.lower()
        or "@" in _current_user.name
    )
    if _needs_name and not st.session_state.get("_mcdm_name_collected"):
        _render_name_collection_screen(_current_user)
        st.stop()

# ---------------------------------------------------------

# ANA GÖVDE
# ---------------------------------------------------------
raw_data = st.session_state.get("raw_data")
if raw_data is None:
    st.markdown(
        f"""
        <div style="text-align:center;padding:4.5rem 1rem 3rem 1rem;">
            <h3 style="color:#1B365D;font-weight:700;font-size:1.6rem;margin-bottom:0.45rem;">{tt("Karar Destek Sistemine Hoş Geldiniz", "Welcome to the Decision Support System")}</h3>
            <p style="color:#64748B;font-size:1.02rem;max-width:760px;margin:0 auto;line-height:1.65;">{tt("Başlamak için aşağıdaki veri girişi alanından dosya yükleyin, örnek veriyi kullanın veya manuel tablo seçeneğine geçin. Veri yüklendiğinde üstteki veri girişi alanı kapanır ve analiz adımları burada görünür.", "To begin, use the data input area below to upload a file, use sample data, or switch to manual entry. Once data is loaded, the data input section collapses and the analysis steps appear here.")}</p>
        </div>
        """,
        unsafe_allow_html=True,
    )
    st.markdown(
        f"""
        <div style="display:flex;justify-content:center;gap:0.65rem;flex-wrap:wrap;margin-top:-1.4rem;margin-bottom:1.4rem;">
            <a href="https://youtu.be/jp4oih6_Nec" target="_blank"
               style="display:inline-block;background:#DC2626;color:#FFFFFF;text-decoration:none;padding:0.6rem 1.05rem;border-radius:8px;font-size:0.92rem;font-weight:600;">
               {tt("🎥 Uygulama Tanıtım Videosu (YouTube) · New-Yeni", "🎥 Application Demo Video (YouTube) · New-Yeni")}
            </a>
            <a href="https://www.instagram.com/mcdm_dss/" target="_blank"
               style="display:inline-block;background:#C13584;color:#FFFFFF;text-decoration:none;padding:0.6rem 1.05rem;border-radius:8px;font-size:0.92rem;font-weight:600;">
               {tt("📸 Instagram: @mcdm_dss", "📸 Instagram: @mcdm_dss")}
            </a>
        </div>
        """,
        unsafe_allow_html=True,
    )
_render_data_input_workspace(st.session_state.get("ui_lang", "TR"), raw_data is not None)
raw_data = st.session_state.get("raw_data")
if raw_data is None:
    st.stop()
_render_analysis_mini_banner()
if not _ensure_mcdm_engine():
    st.error(tt("Analiz motoru yuklenemedi. Lutfen kurulumunuzu kontrol edin.", "Analysis engine could not be loaded. Please check the installation."))
    st.stop()
_ui_stage = _current_ui_stage()
_loaded_shape = raw_data.shape if isinstance(raw_data, pd.DataFrame) else (0, 0)
_loaded_numeric_cols = raw_data.select_dtypes(include=[np.number]).columns.tolist() if isinstance(raw_data, pd.DataFrame) else []
_loaded_year_candidates = _guess_year_columns(raw_data) if isinstance(raw_data, pd.DataFrame) else []
_loaded_entity_candidates = _guess_entity_columns(raw_data, _loaded_year_candidates[0] if _loaded_year_candidates else None) if isinstance(raw_data, pd.DataFrame) else []
_render_tracking_panel(
    tt("Veri Durum Özeti", "Data Status Snapshot"),
    tt(
        "Veri doğrulandı. Şimdi çalışma yönünü seçebilirsiniz.",
        "Data validated. You can now choose the study direction.",
    ),
    [
        (tt("Veri kaynağı", "Data source"), _data_source_label(st.session_state.get("data_source_id")), "green"),
        (tt("Boyut", "Shape"), f"{_loaded_shape[0]} {tt('satır', 'rows')} · {_loaded_shape[1]} {tt('sütun', 'columns')}", "blue"),
        (tt("Sayısal aday", "Numeric candidates"), f"{len(_loaded_numeric_cols)} {tt('sütun', 'columns')}", "blue"),
        (
            tt("Yapı sinyali", "Structure hint"),
            tt("Panel veri olasılığı yüksek", "Likely panel-ready")
            if _loaded_year_candidates
            else tt("Tek dönem / düz tablo", "Single-period / flat table"),
            "amber" if _loaded_year_candidates else "slate",
        ),
        (
            tt("Etiket tahmini", "Label guess"),
            _preview_list_text(_loaded_entity_candidates[:2]) if _loaded_entity_candidates else tt("Otomatik etiketleme", "Automatic labels"),
            "slate",
        ),
    ],
    note=tt(
        "Sıradaki adım: amaç ve veri yapısı seçimi.",
        "Next step: select the objective and data structure.",
    ),
    tone="blue",
    expanded=_ui_stage == "step1",
    icon="📂",
)
if not st.session_state.get("step1_done"):
    _show_step_hint_once(
        "stage_objective",
        "Veri hazır. Şimdi analiz amacınızı ve veri yapınızı belirleyin.",
        "Your data is ready. Now define the analysis objective and data structure.",
        icon="📂",
    )
_needs_ranking_default = bool(st.session_state.get("needs_ranking", True))
_purpose_options = [
    tt("⚖️ Kriterleri önem düzeyine göre sıralamak (Yalnızca Ağırlık)", "⚖️ Rank criteria by importance level (Weights Only)"),
    tt("🏆 Alternatifleri kriterler ile birlikte sıralamak (Ağırlık + Sıralama)", "🏆 Rank alternatives together with criteria (Weights + Ranking)"),
]
_purpose_default_idx = 1 if _needs_ranking_default else 0
_step1_done = bool(st.session_state.get("step1_done"))
_step1_label = f"🎯 {tt('1. Adım: Analizi Amacınız Ne?', 'Step 1: What Is Your Analysis Objective?')}{' ✅' if _step1_done else ''}"

_has_results = bool(st.session_state.get("analysis_result"))
with st.expander(tt("⚙️ Analiz Kurulumu (1-2-3. Adımlar)", "⚙️ Analysis Setup (Steps 1-3)"), expanded=not _has_results):
    st.markdown(f"#### {_step1_label}")
    with st.container():
        _show_step_caption(
            "Analiz amacı ve veri yapısını seçin.",
            "Choose the analysis objective and data structure.",
            "Tek dönem veya panel seçimi sonraki hesaplama ve çıktı akışını belirler.",
            "Single-period or panel selection determines the downstream computation and output flow.",
        )
        st.markdown(
            f"<p style='font-size:0.82rem; color:#5C5650; margin:0 0 0.5rem 0;'>"
            f"{tt('Lütfen bu analizde yapmak istediğiniz işlemi seçin:', 'Please select what you want to accomplish in this analysis:')}"
            f"</p>",
            unsafe_allow_html=True,
        )
        _purpose_choice = st.radio(
            tt("Amacınız nedir?", "What is your objective?"),
            _purpose_options,
            index=_purpose_default_idx,
            label_visibility="collapsed",
            horizontal=True,
        )
        _scope_options = [
            tt("📅 Yıl verisi yok / önemsiz", "📅 No / irrelevant year data"),
            tt("🗂️ Panel Veri", "🗂️ Panel Data"),
        ]
        _scope_default = 1 if st.session_state.get("analysis_scope") == "panel" else 0
        _scope_choice = st.radio(
            tt("Veri yapınız nedir?", "What is your data structure?"),
            _scope_options,
            index=_scope_default,
            horizontal=True,
        )

        # ── Panel veri ayarları (sadece panel seçiliyse) ──
        if _scope_choice == _scope_options[1] and isinstance(raw_data, pd.DataFrame):
            _yr_candidates = _guess_year_columns(raw_data)
            _yr_col_opts   = [str(c) for c in raw_data.columns]
            if _yr_col_opts:
                _yr_col_default = st.session_state.get("panel_year_column")
                if _yr_col_default not in _yr_col_opts:
                    _yr_col_default = _yr_candidates[0] if _yr_candidates else _yr_col_opts[0]
                # Başlık + yıl sütunu etiketi tek HTML satırında
                st.markdown(
                    f'<div style="background:#EEF5FB; border:1px solid #C0D8EE; border-radius:7px;'
                    f'padding:0.25rem 0.55rem; margin:0.3rem 0 0 0; display:flex; align-items:center; gap:0.5rem;">'
                    f'<span style="font-size:0.67rem; font-weight:700; color:#1F4A73; white-space:nowrap;">'
                    f'🗂️ {tt("Panel Ayarları", "Panel Settings")}</span>'
                    f'<span style="font-size:0.64rem; color:#5A7A9A;">— {tt("Yıl sütunu:", "Year column:")}</span>'
                    f'</div>',
                    unsafe_allow_html=True,
                )
                # Selectbox etiket gizli, hemen altında kompakt
                st.markdown('<div style="margin-top:-0.4rem;"></div>', unsafe_allow_html=True)
                _panel_col_inner = st.selectbox(
                    tt("Yıl sütunu", "Year column"),
                    _yr_col_opts,
                    index=_yr_col_opts.index(_yr_col_default),
                    key="panel_year_col_select",
                    label_visibility="collapsed",
                )
                st.session_state["panel_year_column"] = _panel_col_inner
                _entity_opts = [c for c in _yr_col_opts if c != _panel_col_inner]
                if _entity_opts:
                    _entity_candidates = _guess_entity_columns(raw_data, _panel_col_inner)
                    _entity_default = st.session_state.get("panel_entity_column")
                    if _entity_default not in _entity_opts:
                        _entity_default = _entity_candidates[0] if _entity_candidates else _entity_opts[0]
                    st.session_state["panel_entity_column"] = _entity_default
                else:
                    st.session_state["panel_entity_column"] = None
                _panel_strategy_opts = [
                    tt("📆 Yıl Bazlı Ağırlık", "📆 Year-Specific Weights"),
                    tt("🌐 Global Ağırlık", "🌐 Global Weights"),
                ]
                _panel_strategy_default = 1 if st.session_state.get("panel_weight_strategy") == "global" else 0
                _panel_strategy_choice = st.radio(
                    tt("Ağırlık stratejisi", "Weight strategy"),
                    _panel_strategy_opts,
                    index=_panel_strategy_default,
                    horizontal=True,
                    help=tt(
                        "❓ Yıl Bazlı: Her yılın kendi dağılımından ağırlık üretir; dönem içi yapıyı yakalar.\n"
                        "❓ Global: Seçili tüm yılları tek havuzda değerlendirip tek ağırlık vektörü üretir; yıllar arası trend karşılaştırmasında metodolojik tutarlılık sağlar.\n"
                        "Not: Bu strateji objektif ağırlık modunda etkindir.",
                        "❓ Year-Specific: Recomputes weights from each year's own distribution; captures within-period structure.\n"
                        "❓ Global: Builds one weight vector from all selected years pooled together; provides methodological consistency for cross-year trend comparison.\n"
                        "Note: This strategy is active in objective weighting mode.",
                    ),
                )
                st.session_state["panel_weight_strategy"] = "global" if _panel_strategy_choice == _panel_strategy_opts[1] else "yearly"
                _det_years = _sorted_panel_years(raw_data[_panel_col_inner])
                if _det_years:
                    _known_years = st.session_state.get("panel_selected_years_all", [])
                    _known_col = st.session_state.get("panel_selected_years_col")
                    _years_changed = set(_known_years) != set(_det_years)
                    _year_col_changed = _known_col != _panel_col_inner
                    _yr_key_prefix = f"panel_yr_cb::{_panel_col_inner}::"
                    if ("panel_selected_years" not in st.session_state) or _years_changed or _year_col_changed:
                        st.session_state["panel_selected_years"] = []
                        st.session_state["panel_selected_years_all"] = list(_det_years)
                        st.session_state["panel_selected_years_col"] = _panel_col_inner
                        for _yr in _det_years:
                            st.session_state[f"{_yr_key_prefix}{_yr}"] = False

                    _sel_all_col, _clr_all_col = st.columns(2, gap="small")
                    with _sel_all_col:
                        _select_all_clicked = st.button(
                            tt("✅ Tümünü Seç", "✅ Select All"),
                            key="panel_years_select_all",
                            use_container_width=True,
                        )
                    with _clr_all_col:
                        _clear_all_clicked = st.button(
                            tt("🧹 Tümünü Temizle", "🧹 Clear All"),
                            key="panel_years_clear_all",
                            use_container_width=True,
                        )

                    if _select_all_clicked:
                        st.session_state["panel_selected_years"] = list(_det_years)
                        st.session_state["panel_selected_years_all"] = list(_det_years)
                        st.session_state["panel_selected_years_col"] = _panel_col_inner
                        for _yr in _det_years:
                            st.session_state[f"{_yr_key_prefix}{_yr}"] = True
                        st.rerun()

                    if _clear_all_clicked:
                        st.session_state["panel_selected_years"] = []
                        st.session_state["panel_selected_years_all"] = list(_det_years)
                        st.session_state["panel_selected_years_col"] = _panel_col_inner
                        for _yr in _det_years:
                            st.session_state[f"{_yr_key_prefix}{_yr}"] = False
                        st.rerun()

                    _sel_yrs_now = []
                    st.markdown(
                        f'<p style="font-size:0.64rem; color:#4A6070; margin:0.1rem 0 0.05rem 0; line-height:1.2;">'
                        f'{tt("Dönemler:", "Periods:")}</p>',
                        unsafe_allow_html=True,
                    )
                    _yr_sel_cols = st.columns(min(10, len(_det_years)), gap="small")
                    for _yi, _yr in enumerate(_det_years):
                        with _yr_sel_cols[_yi % min(10, len(_det_years))]:
                            _yr_key = f"{_yr_key_prefix}{_yr}"
                            if _yr_key not in st.session_state:
                                st.session_state[_yr_key] = _yr in st.session_state.get("panel_selected_years", [])
                            if st.checkbox(str(_yr), key=_yr_key):
                                _sel_yrs_now.append(_yr)
                    st.session_state["panel_selected_years"] = _sel_yrs_now
                    st.session_state["panel_selected_years_col"] = _panel_col_inner
                    if not _sel_yrs_now:
                        st.warning(tt("En az bir dönem seçin.", "Select at least one period."))
                    elif len(_sel_yrs_now) < 2:
                        st.markdown(
                            f'<p style="font-size:0.63rem; color:#7A5018; margin:0.05rem 0 0 0;">'
                            f'⚠ {tt("Karşılaştırma için ≥2 dönem önerilir.", "≥2 periods recommended.")}</p>',
                            unsafe_allow_html=True,
                        )

        _flow_steps = [
            (tt("Amaç belirleme",        "Define Objective"),        "1",   _step1_done, "#2E7D52", "#D4EFE2"),
            (tt("Veri ön hazırlık",      "Data Preparation"),        "2",   False,       "#1F5F9A", "#DAEAF7"),
            (tt("Ağırlık metodu belirle","Set Weighting Method"),    "3.1", False,       "#B5681E", "#FDE8CC"),
        ]
        if _purpose_choice == _purpose_options[1]:
            _flow_steps.insert(3, (tt("Sıralama metodu belirle", "Set Ranking Method"), "3.2", False, "#9A3030", "#FAD9D9"))
        _flow_steps.append((tt("Analiz yap ve sonuçları yorumla", "Run Analysis & Interpret"), "4", False, "#6B4FA0", "#EAE0F7"))

        _step_nodes = []
        for _idx, (_label, _num, _done, _clr, _bg) in enumerate(_flow_steps):
            _is_last = _idx == len(_flow_steps) - 1
            if _done:
                _box_bg  = _clr
                _num_clr = "rgba(255,255,255,0.80)"
                _lbl_clr = "#FFFFFF"
                _bdr_clr = _clr
                _num_txt = "✓"
            else:
                _box_bg  = _bg
                _num_clr = _clr
                _lbl_clr = _clr
                _bdr_clr = _clr + "99"
                _num_txt = _num

            _node_html = (
                f'<div style="flex:1; min-width:90px; max-width:200px;">'
                f'<div style="background:{_box_bg}; border:1.5px solid {_bdr_clr}; border-radius:10px;'
                f'padding:0.35rem 0.55rem 0.4rem 0.55rem; text-align:center;'
                f'box-shadow:0 2px 8px rgba(0,0,0,0.07);">'
                f'<div style="font-size:0.60rem; font-weight:800; color:{_num_clr}; letter-spacing:0.4px; line-height:1.2; margin-bottom:2px;">{_num_txt}</div>'
                f'<div style="font-size:0.67rem; font-weight:600; color:{_lbl_clr}; line-height:1.35;">{_label}</div>'
                f'</div>'
                f'</div>'
            )
            _step_nodes.append(_node_html)

            if not _is_last:
                _step_nodes.append(
                    '<div style="width:14px; flex-shrink:0; display:flex; align-items:center; justify-content:center;">'
                    '<div style="width:100%; height:2px; background:linear-gradient(90deg,#B0C8DC,#C8D8E8);"></div>'
                    '</div>'
                )

        st.markdown(
            f'<div style="display:flex; align-items:stretch; gap:0; width:100%;'
            f'background:linear-gradient(135deg,#F6FAFE,#EEF4FA);'
            f'border:1px solid #C8D8E8; border-radius:12px;'
            f'padding:0.6rem 0.75rem; margin:0.5rem 0 0.25rem 0;">'
            f'{"".join(_step_nodes)}</div>',
            unsafe_allow_html=True,
        )
    needs_ranking = _purpose_choice == _purpose_options[1]
    st.session_state["needs_ranking"] = needs_ranking
    panel_mode = _scope_choice == _scope_options[1]
    st.session_state["analysis_scope"] = "panel" if panel_mode else "single"

    panel_year_col = None
    panel_entity_col = None
    panel_weight_strategy = "yearly"
    if panel_mode and isinstance(raw_data, pd.DataFrame):
        panel_weight_strategy = "global" if st.session_state.get("panel_weight_strategy") == "global" else "yearly"
        st.session_state["panel_weight_strategy"] = panel_weight_strategy
        _stored_yr_col = st.session_state.get("panel_year_column")
        _all_cols = [str(c) for c in raw_data.columns]
        if _stored_yr_col in _all_cols:
            panel_year_col = _stored_yr_col
        elif _all_cols:
            panel_year_col = _guess_year_columns(raw_data)[0] if _guess_year_columns(raw_data) else _all_cols[0]
            st.session_state["panel_year_column"] = panel_year_col
        _entity_opts = [c for c in _all_cols if c != panel_year_col]
        _stored_entity_col = st.session_state.get("panel_entity_column")
        if _stored_entity_col in _entity_opts:
            panel_entity_col = _stored_entity_col
        elif _entity_opts:
            _guessed_entities = _guess_entity_columns(raw_data, panel_year_col)
            panel_entity_col = _guessed_entities[0] if _guessed_entities else _entity_opts[0]
            st.session_state["panel_entity_column"] = panel_entity_col

    _step1_scope_summary = tt("Panel veri", "Panel data") if panel_mode else tt("Tek dönem", "Single-period")
    _step1_objective_summary = (
        tt("Ağırlık + sıralama", "Weights + ranking")
        if needs_ranking
        else tt("Yalnızca ağırlık", "Weights only")
    )
    _step1_period_summary = (
        _preview_list_text(st.session_state.get("panel_selected_years", []), 4)
        if panel_mode
        else tt("Panel kullanılmıyor", "Panel not used")
    )
    _step1_strategy_summary = (
        tt("Yıl bazlı", "Year-specific")
        if panel_mode and panel_weight_strategy == "yearly"
        else (tt("Global", "Global") if panel_mode else tt("Gerekli değil", "Not required"))
    )
    _step1_output_summary = (
        tt("Ağırlık tablosu + alternatif sıralaması", "Weight table + alternative ranking")
        if needs_ranking
        else tt("Ağırlık tablosu", "Weight table only")
    )
    _render_tracking_panel(
        tt("Çalışma Yönü Özeti", "Study Direction Summary"),
        tt(
            "Bu seçimler yöntem ve rapor akışını belirler.",
            "These choices determine the method and reporting flow.",
        ),
        [
            (tt("Analiz amacı", "Objective"), _step1_objective_summary, "blue"),
            (tt("Veri yapısı", "Data structure"), _step1_scope_summary, "blue"),
            (tt("Panel dönemleri", "Panel periods"), _step1_period_summary, "amber" if panel_mode else "slate"),
            (tt("Ağırlık stratejisi", "Weight strategy"), _step1_strategy_summary, "amber" if panel_mode else "slate"),
            (tt("Beklenen çıktı", "Expected output"), _step1_output_summary, "green"),
        ],
        note=(
            tt(
                "Sıradaki adım: seçimi sabitleyin.",
                "Next step: lock this setup.",
            )
            if not st.session_state.get("step1_done")
            else tt(
                "Sıradaki adım: kriterleri doğrulayın.",
                "Next step: validate the criteria.",
            )
        ),
        tone="blue" if not st.session_state.get("step1_done") else "green",
        expanded=_ui_stage == "step1",
        icon="🎯",
    )
    if st.session_state.get("step1_done") and not st.session_state.get("prep_done"):
        _show_step_hint_once(
            "stage_prep",
            "Amaç seçimi tamamlandı. Şimdi kriterleri ve yönlerini doğrulayın.",
            "Objective selection is complete. Now validate the criteria and their directions.",
            icon="🧭",
        )

    if raw_data is not None and not st.session_state.get("step1_done"):
        st.caption(tt("Hazırsanız devam edin.", "Continue when ready."))
        if st.button(
            tt("✨ Veri hazırlığına geç", "✨ Continue to data preparation"),
            use_container_width=True,
            key="step1_continue_btn",
        ):
            st.session_state["step1_done"] = True
            st.rerun()
        st.stop()

    # ── Veri Hazırlığı ──
    working = raw_data.copy()
    working = clean_dataframe(working, missing_strategy, clip_outliers)
    if panel_mode and panel_year_col and panel_year_col in raw_data.columns:
        working[panel_year_col] = raw_data[panel_year_col]
    st.session_state["alt_names"] = {}
    if (not panel_mode) and isinstance(raw_data, pd.DataFrame):
        _entity_candidates = _guess_entity_columns(raw_data)
        if _entity_candidates:
            _single_entity_col = _entity_candidates[0]
            working.index = _make_unique_labels(raw_data[_single_entity_col].tolist(), fallback_prefix="A")
        else:
            working.index = [f"A{idx+1}" for idx in range(len(working))]
    else:
        working.index = [f"A{idx+1}" for idx in range(len(working))]
    numeric_cols = working.select_dtypes(include=[np.number]).columns.tolist()
    if panel_mode and panel_year_col in numeric_cols:
        numeric_cols = [c for c in numeric_cols if c != panel_year_col]

    if len(numeric_cols) < 2:
        st.error(tt("Hata: Analiz için en az iki sayısal sütun (kriter) gereklidir.", "Error: At least two numeric criterion columns are required for analysis."))
        st.stop()

    _existing_crits = set(st.session_state.get("crit_dir", {}).keys())
    if _existing_crits != set(numeric_cols):
        st.session_state["crit_dir"]     = {c: True for c in numeric_cols}  # varsayılan: tümü Fayda
        st.session_state["crit_include"] = {c: True for c in numeric_cols}

    # ── TEK BİRLEŞİK HAZIRLIK PANELİ ──
    # Diagnostics hesabı (render öncesi)
    _sel_temp = [c for c in numeric_cols if st.session_state["crit_include"].get(c, True)]
    _has_diag = len(_sel_temp) >= 2
    if _has_diag:
        _diag = me.generate_data_diagnostics(working, _sel_temp, {c: ("max" if st.session_state["crit_dir"].get(c, True) else "min") for c in _sel_temp})
        _score, _label = _diag_score_and_label(_diag)
        _badge_cls = "diag-badge-good" if _score >= 80 else ("diag-badge-mid" if _score >= 60 else "diag-badge-bad")
        _rec_items = _diag.get("recommendations", [])[:3]
        _sugg_weight = method_display_name(_diag.get("suggested_weight") or "Entropy")
        _sugg_rank_methods = _diag.get("suggested_ranking_methods") or ["TOPSIS", "VIKOR", "EDAS"]
        _sugg_rank = ", ".join(_sugg_rank_methods[:3])
        _weight_reason = _diag.get("weight_rationale_en") if st.session_state.get("ui_lang") == "EN" else _diag.get("weight_rationale_tr")
        _rank_reason = _diag.get("ranking_rationale_en") if st.session_state.get("ui_lang") == "EN" else _diag.get("ranking_rationale_tr")
        st.session_state["_sugg_rank"] = _sugg_rank
        st.session_state["_sugg_rank_methods"] = _sugg_rank_methods[:3]
        while len(_rec_items) < 3:
            _rec_items.append({"icon": "•",
                "text": tt("Ek kritik bulgu yok.", "No additional critical finding."),
                "action": tt("Mevcut akışla devam edebilirsiniz.", "You may proceed with the current workflow.")})
    else:
        _sugg_weight = method_display_name("Entropy")
        _sugg_rank_methods = ["TOPSIS", "VIKOR", "EDAS"]
        _sugg_rank = ", ".join(_sugg_rank_methods)
        _weight_reason = tt("Varsayılan olarak Entropy başlangıç önerisi kullanılıyor.", "Entropy is used as the default initial suggestion.")
        _rank_reason = tt("Varsayılan olarak TOPSIS, VIKOR ve EDAS başlangıç önerileri kullanılıyor.", "TOPSIS, VIKOR, and EDAS are used as the default initial suggestions.")

    _prep_label = f"🔍 {tt('2. Adım: Ön inceleme - Veri Ön Hazırlığı', 'Step 2: Preliminary Review - Data Preparation')}{' ✅' if st.session_state.get('prep_done') else ''}"
    st.markdown(f"#### {_prep_label}")
    with st.container():
        _show_step_caption(
            "Kriterleri ve ön temizliği netleştirin.",
            "Clarify the criteria and preprocessing choices.",
            "Yanlış yönlü veya analize girmeyecek kriter bırakmamaya dikkat edin.",
            "Avoid leaving wrongly directed or excluded criteria in the analysis set.",
        )

        # ── 1) Ön İnceleme Sonuçları ──
        if _has_diag:
            with st.container():
                st.markdown(
                    f"""<div class="assistant-grid">
                        <div class="assistant-card2">
                            <div class="assistant-title2">📊 {tt("Veri Profili", "Data Profile")}</div>
                            <div class="assistant-body2">
                                {_diag.get('n_alt', 0)} {tt("alternatif", "alternatives")} · {_diag.get('n_crit', 0)} {tt("kriter", "criteria")}<br>
                                {tt("Ort. Varyasyon Katsayısı", "Avg. Coeff. of Variation")}: <strong>{_diag.get('mean_cv', 0.0):.2f}</strong> &nbsp;·&nbsp;
                                {tt("Maks. |ρ|", "Max |ρ|")}: <strong>{_diag.get('max_corr', 0.0):.2f}</strong>
                            </div>
                        </div>
                        <div class="assistant-card2">
                            <div class="assistant-title2">⚖️ {tt("Ağırlık Önerisi", "Weighting Suggestion")}</div>
                            <div class="assistant-body2">
                                {tt("Önerilen:", "Suggested:")} <strong>{_sugg_weight}</strong><br>
                                <span style="font-size:0.76rem; color:#556070;">{_weight_reason}</span>
                            </div>
                        </div>
                        <div class="assistant-card2">
                            <div class="assistant-title2">🏆 {tt("Sıralama Önerisi", "Ranking Suggestion")}</div>
                            <div class="assistant-body2">
                                {tt("Önerilen:", "Suggested:")} <strong>{_sugg_rank}</strong><br>
                                <span style="font-size:0.76rem; color:#556070;">{_rank_reason}</span>
                            </div>
                        </div>
                    </div>""",
                    unsafe_allow_html=True,
                )
                with st.container():
                    for idx, rec in enumerate(_rec_items, start=1):
                        _r_text, _r_action = diag_rec_text(rec)
                        st.markdown(
                            f"""<div class="assistant-card2" style="margin-bottom:0.4rem;">
                                <div class="assistant-title2">{rec.get("icon","•")} {tt("Bulgu","Finding")} {idx}</div>
                                <div class="assistant-body2">{_r_text}<br>
                                <strong>{tt("Öneri:","Recommendation:")}</strong> {_r_action}</div>
                            </div>""",
                            unsafe_allow_html=True,
                        )
                    st.markdown(
                        f"""<div class="assistant-box" style="background:#F0F6FF; margin-top:0.5rem;">
                            <div class="assistant-title2">💡 {tt("Önerilen Yöntem Felsefesi","Recommended Method Philosophy")}</div>
                            <div class="assistant-body2">
                                <strong>{tt("Ağırlık:","Weighting:")}</strong> {_sugg_weight} &nbsp;·&nbsp;
                                <strong>{tt("Sıralama:","Ranking:")}</strong> {_sugg_rank}
                            </div>
                        </div>""",
                        unsafe_allow_html=True,
                    )

        # ── 2) Veri Ön İzleme ──
        with st.container():
            render_table(working.head(3))

        # ── 3) Kriter Yapılandırması ──
        with st.container():
            st.markdown("""<style>
            .ct-wrap {
                border:1px solid #D0D8E4;
                border-radius:8px;
                overflow:hidden;
                margin-top:0.15rem;
            }
            .ct-head {
                display:grid;
                grid-template-columns:34px minmax(0, 1fr) 92px 92px;
                align-items:center;
                background:#EEF3F8;
                padding:0.22rem 0.55rem;
                font-size:0.66rem;
                font-weight:700;
                color:#4A6070;
                text-transform:uppercase;
                letter-spacing:0.4px;
                border-bottom:1px solid #C8D4E0;
            }
            .ct-head > span:nth-child(1),
            .ct-head > span:nth-child(3),
            .ct-head > span:nth-child(4) {
                text-align:center;
            }
            .ct-wrap .stHorizontalBlock {
                border-bottom:1px solid #E8EEF4 !important;
                margin:0 !important;
                padding:0.08rem 0.45rem !important;
                align-items:center !important;
                min-height:2.05rem !important;
            }
            .ct-wrap .stHorizontalBlock:last-of-type { border-bottom:none !important; }
            .ct-wrap .stHorizontalBlock:nth-of-type(odd) { background:#FFFFFF; }
            .ct-wrap .stHorizontalBlock:nth-of-type(even) { background:#F7FAFE; }
            .ct-wrap [data-testid="column"] { padding:0 0.10rem !important; }
            .ct-wrap .element-container { margin:0 !important; padding:0 !important; }
            .ct-wrap .stCheckbox {
                margin:0 !important;
                min-height:1.35rem !important;
                display:flex !important;
                align-items:center !important;
                justify-content:center !important;
            }
            .ct-wrap .stCheckbox > label {
                padding:0 !important;
                margin:0 !important;
                min-height:auto !important;
            }
            .ct-wrap .stCheckbox p { font-size:0 !important; }
            .ct-crit-name {
                font-size:0.77rem;
                font-weight:600;
                color:#1C2A38;
                line-height:1.25;
                white-space:nowrap;
                overflow:hidden;
                text-overflow:ellipsis;
                padding-top:0.06rem;
            }
            .ct-wrap .stButton,
            [class*="st-key-dir_benefit_"],
            [class*="st-key-dir_cost_"] {
                display:flex !important;
                align-items:center !important;
            }
            .ct-wrap .stButton > button,
            [class*="st-key-dir_benefit_"] button,
            [class*="st-key-dir_cost_"] button {
                width:100% !important;
                min-height:1.58rem !important;
                padding:0.10rem 0.22rem !important;
                font-size:0.67rem !important;
                font-weight:600 !important;
                border-radius:6px !important;
                border:1px solid #CCD6E2 !important;
                background:#F3F6F9 !important;
                color:#566678 !important;
                box-shadow:none !important;
                letter-spacing:0 !important;
            }
            .ct-wrap .stButton > button:hover,
            [class*="st-key-dir_benefit_"] button:hover,
            [class*="st-key-dir_cost_"] button:hover {
                background:#EAF0F5 !important;
                border-color:#B7C6D6 !important;
                color:#2F4358 !important;
            }
            .ct-wrap .stButton > button[kind="primary"],
            [class*="st-key-dir_benefit_"] button[kind="primary"],
            [class*="st-key-dir_cost_"] button[kind="primary"] {
                background:#C62828 !important;
                color:#FFFFFF !important;
                -webkit-text-fill-color:#FFFFFF !important;
                border-color:#C62828 !important;
                box-shadow:0 4px 10px rgba(198, 40, 40, 0.18) !important;
            }
            .ct-wrap .stButton > button[kind="primary"] *,
            [class*="st-key-dir_benefit_"] button[kind="primary"] *,
            [class*="st-key-dir_cost_"] button[kind="primary"] * {
                color:#FFFFFF !important;
                fill:#FFFFFF !important;
                -webkit-text-fill-color:#FFFFFF !important;
            }
            .ct-wrap .stButton > button[kind="primary"]:hover,
            [class*="st-key-dir_benefit_"] button[kind="primary"]:hover,
            [class*="st-key-dir_cost_"] button[kind="primary"]:hover {
                background:#B71C1C !important;
                color:#FFFFFF !important;
                -webkit-text-fill-color:#FFFFFF !important;
                border-color:#B71C1C !important;
            }
            .ct-wrap .stButton > button[kind="primary"]:hover *,
            [class*="st-key-dir_benefit_"] button[kind="primary"]:hover *,
            [class*="st-key-dir_cost_"] button[kind="primary"]:hover * {
                color:#FFFFFF !important;
                fill:#FFFFFF !important;
                -webkit-text-fill-color:#FFFFFF !important;
            }
            </style>""", unsafe_allow_html=True)

            _dir_benefit = tt("Fayda", "Benefit")
            _dir_cost = tt("Maliyet", "Cost")
            _direction_notice = st.session_state.pop("direction_notice", None)
            if _direction_notice:
                try:
                    st.toast(_direction_notice, icon="✅")
                except Exception:
                    st.info(_direction_notice)

            # Tablo başlığı (HTML div)
            st.markdown(
                f'<div class="ct-wrap"><div class="ct-head">'
                f'<span>✓</span>'
                f'<span>{tt("Kriter", "Criterion")}</span>'
                f'<span>{tt("Fayda", "Benefit")}</span>'
                f'<span>{tt("Maliyet", "Cost")}</span>'
                f'</div></div>',
                unsafe_allow_html=True,
            )
            # Veri satırları
            st.markdown('<div class="ct-wrap">', unsafe_allow_html=True)
            for _c in numeric_cols:
                _rc = st.columns([0.42, 4.5, 1.15, 1.15], gap="small")
                with _rc[0]:
                    st.session_state["crit_include"][_c] = st.checkbox(
                        f"{tt('Kriteri dahil et', 'Include criterion')}: {_c}",
                        value=st.session_state["crit_include"].get(_c, True),
                        key=f"inc_{_c}",
                        label_visibility="collapsed",
                    )
                with _rc[1]:
                    st.markdown(f'<div class="ct-crit-name">{_safe_html_text(_c)}</div>', unsafe_allow_html=True)
                _is_benefit = st.session_state["crit_dir"].get(_c, True)
                with _rc[2]:
                    if st.button(
                        _dir_benefit,
                        key=f"dir_benefit_{_c}",
                        use_container_width=True,
                        type="primary" if _is_benefit else "secondary",
                    ):
                        st.session_state["crit_dir"][_c] = True
                        st.session_state["direction_notice"] = f"{tt('Seçilen', 'Selected')}: {_c} - {_dir_benefit}"
                        st.rerun()
                with _rc[3]:
                    if st.button(
                        _dir_cost,
                        key=f"dir_cost_{_c}",
                        use_container_width=True,
                        type="primary" if not _is_benefit else "secondary",
                    ):
                        st.session_state["crit_dir"][_c] = False
                        st.session_state["direction_notice"] = f"{tt('Seçilen', 'Selected')}: {_c} - {_dir_cost}"
                        st.rerun()
            st.markdown('</div>', unsafe_allow_html=True)

        if not st.session_state.get("prep_done"):
            st.divider()

    if not st.session_state.get("prep_done"):
        st.caption(tt("Hazırsanız yöntem seçimine geçin.", "Proceed to method selection when ready."))
        if st.button(
            tt("✅ Veri Ön İşleme Bitti (Yöntem Seçimine Geç)", "✅ Preprocessing Complete (Proceed to Method Selection)"),
            use_container_width=True,
            key="btn_prep_complete_main",
        ):
            st.session_state.prep_done = True
            st.rerun()

    criteria = [c for c in numeric_cols if st.session_state["crit_include"].get(c, True)]
    criteria_types = {c: ("max" if st.session_state["crit_dir"].get(c, True) else "min") for c in criteria}
    _benefit_count = sum(1 for c in criteria if criteria_types.get(c) == "max")
    _cost_count = sum(1 for c in criteria if criteria_types.get(c) == "min")
    _missing_label_summary = missing_strategy if missing_strategy not in {"Sil", "Drop"} else tt("Kapalı / Sil", "Off / Drop")
    _prep_scope_summary = (
        f"{len(st.session_state.get('panel_selected_years', []))} {tt('dönem', 'periods')} · {_preview_list_text(st.session_state.get('panel_selected_years', []), 4)}"
        if panel_mode
        else tt("Tek veri dilimi", "Single data slice")
    )
    _render_tracking_panel(
        tt("Hazırlık Özeti", "Preparation Summary"),
        tt(
            "Analize girecek kriter ve temizlik yapısı.",
            "Criteria and cleaning setup for the analysis.",
        ),
        [
            (tt("Dahil kriterler", "Included criteria"), f"{len(criteria)} · {_preview_list_text(criteria, 4)}", "green"),
            (tt("Yön dağılımı", "Direction split"), f"{_benefit_count} {tt('fayda', 'benefit')} · {_cost_count} {tt('maliyet', 'cost')}", "blue"),
            (tt("Hariç kriterler", "Excluded criteria"), str(max(0, len(numeric_cols) - len(criteria))), "amber" if len(criteria) < len(numeric_cols) else "slate"),
            (
                tt("Temizlik ayarı", "Cleaning setup"),
                f"{tt('Eksik', 'Missing')}: {_missing_label_summary} · {tt('Aykırı', 'Outlier')}: {tt('Açık', 'On') if clip_outliers else tt('Kapalı', 'Off')}",
                "amber",
            ),
            (tt("Analiz kapsamı", "Analysis scope"), _prep_scope_summary, "blue"),
        ],
        note=(
            tt(
                "En az 2 kriter bırakın; hazırsanız yöntem seçimine geçin.",
                "Keep at least 2 criteria; continue to method selection when ready.",
            )
            if len(criteria) >= 2
            else tt(
                "Devam için en az 2 kriter seçin.",
                "Select at least 2 criteria to continue.",
            )
        ),
        tone="amber" if not st.session_state.get("prep_done") else "green",
        expanded=_ui_stage == "step2",
        icon="🔍",
    )

    if len(criteria) < 2:
        st.error(tt("En az 2 kriter seçmelisiniz.", "You must select at least 2 criteria."))
        st.stop()

    if not st.session_state.prep_done:
        st.stop()
    else:
        if st.button(tt("🔄 Kriter Ayarlarına Geri Dön", "🔄 Back to Criteria Settings"), use_container_width=True):
            st.session_state.prep_done = False
            st.session_state["analysis_result"] = None
            st.rerun()

    # ---------------------------------------------------------
    # YÖNTEM SEÇİMİ (ALT PANEL)
    # ---------------------------------------------------------
    st.divider()
    _step3_title = tt("⚙️ 3. Adım: Yöntem Seçimi ve Karşılaştırma", "⚙️ Step 3: Method Selection and Comparison")
    st.markdown(f"#### {_step3_title}")
    with st.container():
        _show_step_caption(
            "Yöntem kurulumunu tamamlayın.",
            "Complete the method setup.",
            "Bu seçimler sonuçların yorum ve rapor akışını belirler.",
            "These selections determine the interpretation and reporting flow.",
        )

        weight_mode = ""
        weight_method: str | None = None
        weight_mode_key = "objective"
        manual_weights: Dict[str, float] | None = None
        manual_weights_valid = True
        st.markdown(f"#### {tt('🎯 1. Ağırlık Belirleme', '🎯 1. Weight Determination')}")
        with st.container():
            _show_step_caption(
                "Ağırlık mantığını seçin.",
                "Choose the weighting logic.",
                "Değerleri veriden bağlayarak veya uzman görüşüyle sisteme dahil edebilirsiniz.",
                "Integrate values through data or expert opinion."
            )

            methods_internal = me.OBJECTIVE_WEIGHT_METHODS
            _weight_groups = weight_method_groups()

            options = [
                tt("🎯 Objektif Ağırlık", "🎯 Objective Weights"), 
                tt("🧠 Subjektif Ağırlık", "🧠 Subjective Weights"), 
                tt("⚖️ Eşit Ağırlık", "⚖️ Equal Weights"), 
                tt("✍️ Manuel Ağırlık", "✍️ Manual Weights")
            ]

            st.markdown(
                """
                <style>
                div[data-testid="stPills"] {
                    display: flex;
                    flex-wrap: wrap;
                    gap: 12px;
                    padding-bottom: 0.5rem;
                }
                div[data-testid="stPills"] button {
                    background-color: #1e293b !important;
                    border: 1px solid #334155 !important;
                    color: #94a3b8 !important;
                    border-radius: 8px !important;
                    padding: 0.5rem 1.5rem !important;
                    font-size: 1.05rem !important;
                    transition: all 0.3s ease !important;
                }
                div[data-testid="stPills"] button[aria-pressed="true"], div[data-testid="stPills"] button[data-selected="true"] {
                    background: linear-gradient(135deg, #10B981 0%, #059669 100%) !important;
                    color: white !important;
                    border-color: #10B981 !important;
                    box-shadow: 0 4px 12px rgba(16, 185, 129, 0.4) !important;
                    font-weight: 600 !important;
                }
                div[data-testid="stPills"] button:hover {
                    transform: translateY(-2px) !important;
                    border-color: #059669 !important;
                    color: white !important;
                }
                </style>
                """, unsafe_allow_html=True
            )

            weight_mode = st.pills(
                f"**{tt('Ağırlıklandırma Modülü', 'Weighting Module')}**",
                options,
                selection_mode="single",
                default=options[0],
                help=tt(
                    "Bu seçim, ağırlıkların veriden mi, uzman görüşünden mi, eşit dağılımdan mı yoksa manuel olarak mı üretileceğini belirler.",
                    "This choice determines the source of weights: data, expert evaluation, equal distribution, or manual input."
                ),
            )
            if not weight_mode:
                weight_mode = options[0]

            if "Objektif" in weight_mode or "Objective" in weight_mode:
                weight_mode_key = "objective"
                st.caption(tt("Tüm yöntemler aşağıda kümelenmiş halde gösterilir. Tek seçim yapılır.", "All methods are grouped below. A single selection is used."))
                # Checkbox durumlarını ilk açılışta ayarla
                for _m in methods_internal:
                    if f"weight_cb_{_m}" not in st.session_state:
                        st.session_state[f"weight_cb_{_m}"] = False
                for group_label, group_methods in _weight_groups:
                    _filtered_methods = [m for m in group_methods if m in methods_internal]
                    if not _filtered_methods:
                        continue
                    st.markdown(f"**{group_label}**")
                    method_cols = st.columns(3)
                    for i, method_name in enumerate(_filtered_methods):
                        with method_cols[i % len(method_cols)]:
                            _label = method_display_name(method_name)
                            _cb_key = f"weight_cb_{method_name}"
                            _w_help = tt(
                                f"{_method_help_text(method_name)}\nBu yöntemi seçerek devam edebilirsiniz.",
                                f"{_method_help_text(method_name)}\nYou can proceed by selecting this method.",
                            )
                            st.checkbox(
                                _label, key=_cb_key, help=_w_help,
                                on_change=_wm_single_select_cb, args=(method_name, methods_internal),
                            )
                # Seçili yöntemi güncelle
                _all_weight_checked = [m for m in methods_internal if st.session_state.get(f"weight_cb_{m}", False)]
                if _all_weight_checked:
                    weight_method = _all_weight_checked[0]
                    st.session_state["weight_method_pref"] = weight_method
                    # Seçili yöntem felsefesi kutusu
                    _wh = _method_help_text(weight_method).split("\n", 1)
                    _wh_simple   = _wh[0].strip()
                    _wh_academic = _wh[1].strip() if len(_wh) > 1 else ""
                    _academic_html = (
                        f'<div style="font-size:0.70rem; color:#4A6070; line-height:1.5; margin-top:0.15rem;">{_wh_academic}</div>'
                        if _wh_academic else ""
                    )
                    st.markdown(
                        f'<div style="background:#EAF3FB; border-left:3px solid #5A9CC5; border-radius:0 8px 8px 0; '
                        f'padding:0.45rem 0.75rem; margin-top:0.5rem;">'
                        f'<div style="font-size:0.72rem; font-weight:700; color:#1F4A73; margin-bottom:0.18rem;">'
                        f'💡 {method_display_name(weight_method)}</div>'
                        f'<div style="font-size:0.71rem; color:#2A3E54; line-height:1.55;">{_wh_simple}</div>'
                        f'{_academic_html}'
                        f'</div>',
                        unsafe_allow_html=True,
                    )
                else:
                    st.session_state["weight_method_pref"] = None
                    st.markdown(
                        f"<p style='color:#B91C1C;font-size:0.78rem;font-weight:700;margin:0.45rem 0 0 0;'>❗ {tt('Lütfen seçiminizi yapın.', 'Please make your selection.')}</p>",
                        unsafe_allow_html=True,
                    )
            elif "Eşit" in weight_mode or "Equal" in weight_mode:
                weight_method = "Eşit Ağırlık"
                weight_mode_key = "equal"
                st.session_state["weight_method_pref"] = weight_method
                st.caption(tt("Tüm kriterlere eşit önem atanır.", "All criteria are assigned equal importance."))
                _weight_help_lines = _method_help_text(weight_method).split("\n", 1)
                st.caption(_weight_help_lines[0])
                if len(_weight_help_lines) > 1:
                    st.caption(_weight_help_lines[1])
            elif "Subjektif" in weight_mode or "Subjective" in weight_mode:
                weight_method = "Subjektif (Uzman) Ağırlık"
                weight_mode_key = "manual"
                st.session_state["weight_method_pref"] = weight_method
                import ui_subjective
                ui_subjective.render_subjective_component(criteria)

                _sub_weights = st.session_state.get("subjective_manual_weights")
                if _sub_weights:
                    manual_weights = _sub_weights
                    manual_weights_valid = True

                    _manual_preview = pd.DataFrame(
                        {
                            "Kriter": criteria,
                            tt("Türetilen Ağırlık", "Derived Weight"): [_sub_weights.get(c, np.nan) for c in criteria],
                        }
                    )
                    render_table(_manual_preview)
                else:
                    manual_weights = None
                    manual_weights_valid = False
            else:
                weight_method = "Manuel Ağırlık"
                weight_mode_key = "manual"
                st.session_state["weight_method_pref"] = weight_method
                _weight_help_lines = _method_help_text(weight_method).split("\n", 1)
                st.caption(_weight_help_lines[0])
                if len(_weight_help_lines) > 1:
                    st.caption(_weight_help_lines[1])

                _mcols = st.columns(min(3, max(1, len(criteria))))
                manual_raw_text: Dict[str, str] = {}
                for i, c in enumerate(criteria):
                    with _mcols[i % len(_mcols)]:
                        manual_raw_text[c] = st.text_input(
                            c,
                            value=str(st.session_state.get(f"manual_w_{c}", "")),
                            placeholder=tt("örn. 7.5", "e.g. 7.5"),
                            key=f"manual_w_{c}",
                        ).strip()

                st.caption(
                    tt(
                        "Seçili kriterler için ham önem puanı girin. Değerlerin toplamının 1 olması gerekmez; sistem analiz anında oranları otomatik normalize eder. Bu alan, AHP veya benzeri bir yaklaşımla dışarıda belirlediğiniz göreli önemleri esnek biçimde girmeniz için tasarlanmıştır.",
                        "Enter raw importance values for the selected criteria. Their total does not need to be 1; the system automatically normalizes the ratios during analysis. This area is designed for flexibly entering relative priorities produced externally through AHP or a similar approach.",
                    )
                )

                missing_manual: List[str] = []
                invalid_manual: List[str] = []
                parsed_manual: Dict[str, float] = {}
                for c, raw_val in manual_raw_text.items():
                    if not raw_val:
                        missing_manual.append(c)
                        continue
                    try:
                        val = float(raw_val.replace(",", "."))
                        if not np.isfinite(val) or val <= 0:
                            raise ValueError
                        parsed_manual[c] = val
                    except Exception:
                        invalid_manual.append(c)

                if missing_manual:
                    st.info(
                        tt(
                            "Analize geçmeden önce tüm seçili kriterler için değer girin.",
                            "Enter values for all selected criteria before running the analysis.",
                        )
                    )
                    manual_weights_valid = False
                if invalid_manual:
                    st.warning(
                        tt(
                            f"Geçerli pozitif sayı beklenen kriterler: {', '.join(invalid_manual)}.",
                            f"Valid positive numbers are required for: {', '.join(invalid_manual)}.",
                        )
                    )
                    manual_weights_valid = False

                if parsed_manual:
                    _manual_total = float(sum(parsed_manual.values()))
                    manual_weights = {c: parsed_manual[c] for c in criteria if c in parsed_manual}
                    _manual_preview = pd.DataFrame(
                        {
                            "Kriter": criteria,
                            tt("Girilen Değer", "Entered Value"): [parsed_manual.get(c, np.nan) for c in criteria],
                            tt("Normalize Ağırlık", "Normalized Weight"): [
                                (parsed_manual[c] / _manual_total) if c in parsed_manual and _manual_total > 0 else np.nan
                                for c in criteria
                            ],
                        }
                    )
                    render_table(_manual_preview)
                else:
                    manual_weights = None
                    manual_weights_valid = False

        ranking_methods_selected = []
        primary_rank_method: str | None = None
        vikor_v = float(st.session_state.get("vikor_v", 0.5))
        waspas_lambda = float(st.session_state.get("waspas_lambda", 0.5))
        codas_tau = float(st.session_state.get("codas_tau", 0.02))
        cocoso_lambda = float(st.session_state.get("cocoso_lambda", 0.5))
        gra_rho = float(st.session_state.get("gra_rho", 0.5))
        promethee_pref_func = st.session_state.get("promethee_pref_func", "linear")
        promethee_q = float(st.session_state.get("promethee_q", 0.05))
        promethee_p = float(st.session_state.get("promethee_p", 0.30))
        promethee_s = float(st.session_state.get("promethee_s", 0.20))
        fuzzy_spread = float(st.session_state.get("fuzzy_spread", 0.10))
        sensitivity_iterations = int(st.session_state.get("sensitivity_iterations", 400))
        sensitivity_iterations = max(0, min(MAX_SENSITIVITY_ITERATIONS, sensitivity_iterations))
        sensitivity_sigma = float(st.session_state.get("sensitivity_sigma", 0.12))
        run_heavy_robustness = bool(st.session_state.get("run_heavy_robustness", False))

        _is_manual_mode_local = (weight_mode_key == "manual")
        _weight_step_done_local = bool(weight_method) and (manual_weights_valid if _is_manual_mode_local else True)

        if _weight_step_done_local:
            st.divider()
            st.markdown(f"#### {tt('📊 2. Sıralama Yöntemi', '📊 2. Ranking Method')}")

        with st.container():
            if not _weight_step_done_local:
                pass  # Sıralama bölümü gizli — ağırlık seçilmeden gösterilmez
            elif needs_ranking:

                _show_step_caption(
                    "Sıralama yöntemlerini seçin.",
                    "Choose the ranking methods.",
                    "Birden fazla yöntem seçerseniz yöntem uyumu da karşılaştırılır.",
                    "If you choose multiple methods, method agreement is also compared.",
                )
                _layer_options = [
                    tt("Temel Yöntemler (Klasik Mantık)", "Core Methods (Classical Logic)"),
                    tt("İleri Düzey Yöntemler (Fuzzy)", "Advanced Methods (Fuzzy)"),
                ]
                layer_choice = st.radio(
                    f"**{tt('Analiz Katmanı', 'Analysis Layer')}**",
                    _layer_options,
                    horizontal=True,
                    help=tt("Klasik katman kesin sayısal değerlerle, Fuzzy katman ise belirsizlik toleransı (spread) ile çalışır.", "Classical uses crisp numeric values; fuzzy uses uncertainty tolerance (spread)."),
                )
                if layer_choice == _layer_options[0]:
                    all_ranks = me.CLASSICAL_MCDM_METHODS
                    _layer_key = "classical"
                else:
                    all_ranks = me.FUZZY_MCDM_METHODS
                    _layer_key = "fuzzy"

                _sugg_rank_raw = st.session_state.get("_sugg_rank", "TOPSIS, VIKOR, EDAS")
                _sugg_rank_list = st.session_state.get("_sugg_rank_methods") or ["TOPSIS", "VIKOR", "EDAS"]
                _fallback_recommendations = (
                    ["TOPSIS", "VIKOR", "EDAS"]
                    if _layer_key == "classical"
                    else ["Fuzzy TOPSIS", "Fuzzy VIKOR", "Fuzzy EDAS"]
                )

                def _resolve_rank_method_name(candidate: str) -> str | None:
                    cand = candidate.strip()
                    if not cand:
                        return None
                    cand_upper = cand.upper().replace("FUZZY ", "").strip()
                    for method_name in all_ranks:
                        method_upper = method_name.upper().replace("FUZZY ", "").strip()
                        if method_upper == cand_upper:
                            return method_name
                    return None

                recommended_rank_methods: List[str] = []
                for candidate in [*_sugg_rank_list, *_sugg_rank_raw.replace("/", ",").split(","), *_fallback_recommendations]:
                    resolved_method = _resolve_rank_method_name(candidate)
                    if resolved_method and resolved_method not in recommended_rank_methods:
                        recommended_rank_methods.append(resolved_method)
                    if len(recommended_rank_methods) >= 3:
                        break
                quick_pick_methods = recommended_rank_methods[:3]

                st.caption(
                    tt(
                        f"Bu katmanda toplam {len(all_ranks)} yöntem kullanılabilir.",
                        f"A total of {len(all_ranks)} methods are available in this layer.",
                    )
                )
                _mcol1, _mcol2 = st.columns(2)
                with _mcol1:
                    if st.button(tt("✅ Önerilen Yöntemleri Çalıştır", "✅ Run Recommended Methods"), use_container_width=True, key=f"btn_select_recommended_{_layer_key}"):
                        st.session_state["ranking_prefs"] = list(quick_pick_methods)
                        for rank_method in all_ranks:
                            st.session_state[f"rank_cb_{rank_method}"] = rank_method in quick_pick_methods
                        st.rerun()
                with _mcol2:
                    if st.button(tt("🗑️ Temizle", "🗑️ Clear"), use_container_width=True, key=f"btn_clear_all_{_layer_key}"):
                        st.session_state["ranking_prefs"] = []
                        for rank_method in all_ranks:
                            st.session_state[f"rank_cb_{rank_method}"] = False
                        st.rerun()
                st.caption(
                    tt(
                        f"Önerilen başlangıç yöntemleri: {', '.join(recommended_rank_methods[:3])}.",
                        f"Recommended starter methods: {', '.join(recommended_rank_methods[:3])}.",
                    )
                )
                st.caption(
                    tt(
                        f"Neden bu öneri? {_rank_reason}",
                        f"Why this recommendation? {_rank_reason}",
                    )
                )

                _pref_ranks = st.session_state.get("ranking_prefs") or []
                _has_pref_in_layer = any(m in all_ranks for m in _pref_ranks)
                for group_label, group_methods in ranking_method_groups(_layer_key):
                    _filtered_methods = [m for m in group_methods if m in all_ranks]
                    if not _filtered_methods:
                        continue
                    st.markdown(f"**{group_label}**")
                    rank_cols = st.columns(3)
                    for i, rank_method in enumerate(_filtered_methods):
                        with rank_cols[i % len(rank_cols)]:
                            if _pref_ranks and _has_pref_in_layer:
                                default_val = rank_method in _pref_ranks
                            else:
                                default_val = False
                            _rank_widget_key = f"rank_cb_{rank_method}"
                            _rank_help = tt(
                                f"{_method_help_text(rank_method)}\nBu yöntemi seçerseniz karşılaştırma çıktılarında da birlikte değerlendirilir.",
                                f"{_method_help_text(rank_method)}\nIf you select this method, it will also be evaluated in the comparison outputs.",
                            )
                            if _rank_widget_key in st.session_state:
                                _rank_checked = st.checkbox(rank_method, key=_rank_widget_key, help=_rank_help)
                            else:
                                _rank_checked = st.checkbox(rank_method, value=default_val, key=_rank_widget_key, help=_rank_help)
                            if _rank_checked:
                                ranking_methods_selected.append(rank_method)
                st.session_state["ranking_prefs"] = ranking_methods_selected

                if not ranking_methods_selected:
                    st.markdown(
                        f"<p style='color:#B91C1C;font-size:0.78rem;font-weight:700;margin:0.45rem 0 0 0;'>❗ {tt('Lütfen seçiminizi yapın.', 'Please make your selection.')}</p>",
                        unsafe_allow_html=True,
                    )
                    st.warning(tt("Lütfen en az bir sıralama yöntemi seçiniz.", "Please select at least one ranking method."))
                else:
                    _sugg_primary = next(
                        (m for m in ranking_methods_selected
                         if any(m == _resolve_rank_method_name(part) for part in [*_sugg_rank_list, *_sugg_rank_raw.replace("/", ",").split(",")])),
                        ranking_methods_selected[0] if ranking_methods_selected else None,
                    )
                    primary_rank_method = _sugg_primary or (ranking_methods_selected[0] if ranking_methods_selected else None)

                    uses_vikor = ("VIKOR" in ranking_methods_selected) or ("Fuzzy VIKOR" in ranking_methods_selected)
                    uses_waspas = ("WASPAS" in ranking_methods_selected) or ("Fuzzy WASPAS" in ranking_methods_selected)
                    uses_codas = ("CODAS" in ranking_methods_selected) or ("Fuzzy CODAS" in ranking_methods_selected)
                    uses_cocoso = ("CoCoSo" in ranking_methods_selected) or ("Fuzzy CoCoSo" in ranking_methods_selected)
                    uses_gra = ("GRA" in ranking_methods_selected) or ("Fuzzy GRA" in ranking_methods_selected)
                    uses_promethee = ("PROMETHEE" in ranking_methods_selected) or ("Fuzzy PROMETHEE" in ranking_methods_selected)
                    uses_fuzzy = any(m.startswith("Fuzzy") for m in ranking_methods_selected)
                    rec = recommend_parameter_defaults(ranking_methods_selected, len(working), len(criteria), uses_fuzzy)

                    with st.container():
                        _show_step_caption(
                            "Yalnız gerekli yöntem parametreleri burada yer alır.",
                            "Only required method parameters appear here.",
                            "Monte Carlo ve sağlamlık ayarları bir sonraki ayrı adımda yer alır.",
                            "Monte Carlo and robustness settings are placed in the next dedicated step.",
                        )
                        param_mode = st.radio(
                            tt("Parametre Modu", "Parameter Mode"),
                            [tt("🎯 Önerilen Varsayılan", "🎯 Recommended Default"), tt("✍️ Manuel Değiştir", "✍️ Manual Override")],
                            horizontal=True,
                            key="param_mode_choice",
                        )
                        if not any([uses_vikor, uses_waspas, uses_codas, uses_cocoso, uses_gra, uses_promethee, uses_fuzzy]):
                            st.caption(tt("Seçili sıralama yöntemleri ek parametre gerektirmiyor.", "Selected ranking methods do not require extra parameters."))
                        if ("Önerilen" in param_mode) or ("Recommended" in param_mode):
                            vikor_v = float(rec["vikor_v"])
                            waspas_lambda = float(rec["waspas_lambda"])
                            codas_tau = float(rec["codas_tau"])
                            cocoso_lambda = float(rec["cocoso_lambda"])
                            gra_rho = float(rec["gra_rho"])
                            promethee_pref_func = str(rec.get("promethee_pref_func", "linear"))
                            promethee_q = float(rec.get("promethee_q", 0.05))
                            promethee_p = float(rec.get("promethee_p", 0.30))
                            promethee_s = float(rec.get("promethee_s", 0.20))
                            fuzzy_spread = float(rec["fuzzy_spread"])
                            sensitivity_iterations = int(max(100, min(MAX_SENSITIVITY_ITERATIONS, int(rec["sensitivity_iterations"]))))
                            sensitivity_sigma = float(rec["sensitivity_sigma"])
                            st.caption(tt("Önerilen değerler otomatik uygulanır. İsterseniz Manuel Değiştir moduna geçebilirsiniz.", "Recommended values are applied automatically. Switch to Manual Override if needed."))
                            rec_rows = []
                            if uses_vikor:
                                rec_rows.append({tt("Parametre", "Parameter"): "VIKOR v", tt("Değer", "Value"): vikor_v})
                            if uses_waspas:
                                rec_rows.append({tt("Parametre", "Parameter"): "WASPAS λ", tt("Değer", "Value"): waspas_lambda})
                            if uses_codas:
                                rec_rows.append({tt("Parametre", "Parameter"): "CODAS τ", tt("Değer", "Value"): codas_tau})
                            if uses_cocoso:
                                rec_rows.append({tt("Parametre", "Parameter"): "CoCoSo λ", tt("Değer", "Value"): cocoso_lambda})
                            if uses_gra:
                                rec_rows.append({tt("Parametre", "Parameter"): "GRA ρ", tt("Değer", "Value"): gra_rho})
                            if uses_promethee:
                                rec_rows.extend([
                                    {tt("Parametre", "Parameter"): "PROMETHEE pref_func", tt("Değer", "Value"): promethee_pref_func},
                                    {tt("Parametre", "Parameter"): "PROMETHEE q", tt("Değer", "Value"): promethee_q},
                                    {tt("Parametre", "Parameter"): "PROMETHEE p", tt("Değer", "Value"): promethee_p},
                                    {tt("Parametre", "Parameter"): "PROMETHEE s", tt("Değer", "Value"): promethee_s},
                                ])
                            if uses_fuzzy:
                                rec_rows.append({tt("Parametre", "Parameter"): "Fuzzy spread", tt("Değer", "Value"): fuzzy_spread})
                            rec_df = pd.DataFrame(rec_rows)
                            render_table(rec_df)
                        else:
                            if uses_fuzzy:
                                fuzzy_spread = st.slider(tt("Fuzzy — spread (belirsizlik bandı)", "Fuzzy — spread (uncertainty band)"), 0.01, 0.50, float(fuzzy_spread), 0.01, help=tt("Crisp değerlerin TFN dönüşüm oranı. 0.10=±%10 belirsizlik.", "TFN conversion ratio. 0.10=±10% uncertainty."))
                            if uses_vikor:
                                vikor_v = st.slider(tt("VIKOR — v (uzlaşı katsayısı)", "VIKOR — v (compromise factor)"), 0.0, 1.0, float(vikor_v), 0.01, help=tt("v=0.5: denge. v→1: çoğunluk. v→0: azınlık koruması.", "v=0.5: balance. v→1: majority. v→0: minority protection."))
                            if uses_waspas:
                                waspas_lambda = st.slider(tt("WASPAS — λ (hibrit katsayı)", "WASPAS — λ (hybrid coefficient)"), 0.0, 1.0, float(waspas_lambda), 0.01, help=tt("λ=1: saf WSM. λ=0: saf WPM. λ=0.5: eşit karışım.", "λ=1: pure WSM. λ=0: pure WPM. λ=0.5: equal mix."))
                            if uses_codas:
                                codas_tau = st.slider(tt("CODAS — τ (eşik değeri)", "CODAS — τ (threshold)"), 0.0, 0.20, float(codas_tau), 0.005, help=tt("Öklid yakınsa Manhattan devreye girer. Varsayılan: 0.02.", "Manhattan kicks in when Euclidean is close. Default: 0.02."))
                            if uses_cocoso:
                                cocoso_lambda = st.slider(tt("CoCoSo — λ (birleşim katsayısı)", "CoCoSo — λ (aggregation coefficient)"), 0.0, 1.0, float(cocoso_lambda), 0.01, help=tt("λ→1: toplamsal. λ→0: üstel ağırlık.", "λ→1: additive. λ→0: exponential weight."))
                            if uses_gra:
                                gra_rho = st.slider(tt("GRA — ρ (ayırt edici katsayı)", "GRA — ρ (distinguishing coefficient)"), 0.10, 0.90, float(gra_rho), 0.01, help=tt("ρ=0.5: dengeli. ρ→0: uç farklara hassas. ρ→1: bastırır.", "ρ=0.5: balanced. ρ→0: sensitive to extremes. ρ→1: suppresses."))
                            if uses_promethee:
                                promethee_pref_func = st.selectbox(
                                    tt("PROMETHEE tercih fonksiyonu", "PROMETHEE preference function"),
                                    ["linear", "usual", "u_shape", "v_shape", "level", "gaussian"],
                                    index=["linear", "usual", "u_shape", "v_shape", "level", "gaussian"].index(str(promethee_pref_func) if str(promethee_pref_func) in ["linear", "usual", "u_shape", "v_shape", "level", "gaussian"] else "linear"),
                                )
                                if promethee_pref_func in {"u_shape", "level", "linear"}:
                                    promethee_q = st.slider("PROMETHEE q", 0.0, 1.0, float(promethee_q), 0.01)
                                if promethee_pref_func in {"v_shape", "level", "linear"}:
                                    promethee_p = st.slider("PROMETHEE p", 0.0, 1.0, float(promethee_p), 0.01)
                                if promethee_pref_func == "gaussian":
                                    promethee_s = st.slider("PROMETHEE s", 0.01, 1.0, float(promethee_s), 0.01)

        _ranking_step_done_local = bool(ranking_methods_selected) if needs_ranking else False
        if _weight_step_done_local and _ranking_step_done_local:
            st.divider()
            st.markdown(f"#### {tt('🛡️ 3. Dayanıklılık Testi', '🛡️ 3. Robustness Test')}")

        with st.container():
            
            if not _weight_step_done_local or (needs_ranking and not _ranking_step_done_local):
                st.warning(tt("⚠️ Dayanıklılık (Robustness) testlerini yapılandırabilmek için lütfen önce Ağırlık ve/veya Sıralama adımlarını tamamlayın.", "⚠️ To configure robustness testing, please first complete the Weighting and/or Ranking steps."))
            elif needs_ranking:
                # Removed nested step 3.4 expander, putting it inside TAB 3
                with st.container():
                    _show_step_caption(
                        "Sağlamlık testini burada ayarlayın.",
                        "Configure robustness testing here.",
                        "Monte Carlo ve analiz sonuç toleransı burada değerlendirilir.",
                        "Monte Carlo and result tolerance are tested here.",
                    )
                    run_heavy_robustness = st.checkbox(
                        tt(
                            "🛡️ Geniş sağlamlık testlerini çalıştır (Biraz uzun sürebilir.)",
                            "🛡️ Run full robustness tests (May take a little longer.)",
                        ),
                        key="run_heavy_robustness",
                    )
                    if run_heavy_robustness:
                        _robustness_mode = st.radio(
                            tt("Sağlamlık Parametre Modu", "Robustness Parameter Mode"),
                            [tt("🎯 Önerilen Varsayılan", "🎯 Recommended Default"), tt("✍️ Manuel Değiştir", "✍️ Manual Override")],
                            horizontal=True,
                            key="robustness_param_mode_choice",
                        )
                        if ("Önerilen" in _robustness_mode) or ("Recommended" in _robustness_mode):
                            sensitivity_iterations = int(max(100, min(MAX_SENSITIVITY_ITERATIONS, int(rec["sensitivity_iterations"]))))
                            sensitivity_sigma = float(rec["sensitivity_sigma"])
                            _rob_df = pd.DataFrame(
                                [
                                    {tt("Parametre", "Parameter"): tt("Monte Carlo iterasyon", "Monte Carlo iterations"), tt("Değer", "Value"): sensitivity_iterations},
                                    {tt("Parametre", "Parameter"): "Monte Carlo sigma", tt("Değer", "Value"): sensitivity_sigma},
                                ]
                            )
                            render_table(_rob_df)
                        else:
                            _sens_default = int(sensitivity_iterations) if int(sensitivity_iterations) > 0 else 400
                            _sens_default = int(max(100, min(MAX_SENSITIVITY_ITERATIONS, _sens_default)))
                            sensitivity_iterations = st.slider(
                                tt("Monte Carlo iterasyon", "Monte Carlo iterations"),
                                100,
                                MAX_SENSITIVITY_ITERATIONS,
                                _sens_default,
                                100,
                            )
                            sensitivity_sigma = st.slider(tt("Monte Carlo sigma", "Monte Carlo sigma"), 0.01, 0.50, float(sensitivity_sigma), 0.01, help=tt("Log-normal gürültü std sapma. Düşük=hafif, yüksek=agresif. Varsayılan: 0.12.", "Log-normal noise std dev. Low=mild, high=aggressive. Default: 0.12."))
                    else:
                        sensitivity_iterations = 0
                        st.caption(
                            tt(
                                "Sağlamlık testleri bu çalışmada kapatıldı (performans modu).",
                                "Robustness tests are disabled for this run (performance mode).",
                            )
                        )

                    st.session_state["sensitivity_iterations"] = int(sensitivity_iterations)
                    st.session_state["sensitivity_sigma"] = float(sensitivity_sigma)

_is_manual_mode = ("Manuel" in weight_mode) or ("Manual" in weight_mode)
_weight_step_done = bool(weight_method) and (manual_weights_valid if _is_manual_mode else True)
_ranking_step_done = bool(ranking_methods_selected) if needs_ranking else False
_method_step_done = _weight_step_done and ((not needs_ranking) or _ranking_step_done)
_weight_mode_summary = {
    "objective": tt("Objektif", "Objective"),
    "equal": tt("Eşit", "Equal"),
    "manual": tt("Manuel", "Manual"),
}.get(weight_mode_key, tt("Belirlenmedi", "Not set"))
_weight_method_summary = method_display_name(weight_method) if weight_method else tt("Henüz seçilmedi", "Not selected yet")
_ranking_methods_summary = (
    _preview_list_text([method_display_name(m) for m in ranking_methods_selected], 3)
    if ranking_methods_selected
    else (tt("Bu çalışmada sıralama yok", "No ranking in this run") if not needs_ranking else tt("Henüz seçilmedi", "Not selected yet"))
)
_primary_rank_summary = (
    method_display_name(primary_rank_method or ranking_methods_selected[0])
    if ranking_methods_selected
    else tt("Gerekli değil", "Not required")
)
_run_scope_summary = (
    f"{tt('Panel veri', 'Panel data')} · {len(st.session_state.get('panel_selected_years', []))} {tt('dönem', 'periods')}"
    if panel_mode
    else tt("Tek dönem", "Single-period")
)
_robustness_summary = tt("Açık", "On") if run_heavy_robustness else tt("Kapalı", "Off")
_render_tracking_panel(
    tt("Analiz Öncesi Çalışma Özeti", "Pre-Run Analysis Summary"),
    tt(
        "Çalıştırma anındaki kurulum özeti.",
        "Setup snapshot used at run time.",
    ),
    [
        (tt("Analiz amacı", "Objective"), _step1_objective_summary, "blue"),
        (tt("Analiz kapsamı", "Scope"), _run_scope_summary, "blue"),
        (tt("Kriter seti", "Criteria set"), f"{len(criteria)} · {_preview_list_text(criteria, 4)}", "green"),
        (tt("Ağırlık kurgusu", "Weighting setup"), f"{_weight_mode_summary} · {_weight_method_summary}", "amber" if weight_method else "rose"),
        (tt("Sıralama paketi", "Ranking package"), _ranking_methods_summary, "amber" if ranking_methods_selected or not needs_ranking else "rose"),
        (tt("Ana sıralama", "Primary ranking"), _primary_rank_summary, "slate"),
        (tt("Sağlamlık modu", "Robustness mode"), _robustness_summary, "green" if run_heavy_robustness else "slate"),
    ],
    note=(
        tt(
            "Sıradaki adım: analizi çalıştırın.",
            "Next step: run the analysis.",
        )
        if _method_step_done
        else tt(
            "Devam için ağırlık ve gerekiyorsa sıralama seçimini tamamlayın.",
            "Complete weighting and ranking selection before continuing.",
        )
    ),
    tone="green" if _method_step_done else "rose",
    expanded=_ui_stage == "step3",
    icon="🧪",
)
if _method_step_done:
    st.caption(tt("Hazırsanız analizi çalıştırın.", "Run the analysis when ready."))
    _show_step_hint_once(
        "stage_methods",
        "Yöntem seçimi tamamlandı. Analizi çalıştırabilirsiniz.",
        "Method selection is complete. You can now run the analysis.",
        icon="🧪",
    )

# ── İleri Düzey Analizler (Kapalı) ──
with st.expander(tt("🔬 İleri Düzey Analizler", "🔬 Advanced Analysis Options"), expanded=False):
    st.caption(tt(
        "Ön eleme filtreleri ve çoklu paydaş senaryoları gibi ileri analiz araçları. Temel analiz için bu bölümü açmanız gerekmez.",
        "Advanced tools such as pre-screening filters and multi-stakeholder scenarios. Not required for standard analysis."
    ))

    # ── Ön Eleme ──
    st.markdown(f"**{tt('🚧 Ön Eleme Eşikleri', '🚧 Pre-screening Thresholds')}**")
    st.caption(tt(
        "Minimum kabul değeri veya maksimum sınır belirleyin. Eşiği karşılamayan alternatifler sıralamadan çıkarılır.",
        "Set minimum acceptable or maximum limit values. Alternatives failing thresholds are excluded before ranking."
    ))
    _show_thresholds = st.checkbox(tt("Ön eleme filtresi kullan", "Use pre-screening filter"), key="show_thresholds")
    if _show_thresholds and criteria:
        if "crit_thresholds" not in st.session_state:
            st.session_state["crit_thresholds"] = {}
        for _c in criteria:
            _tc = st.columns([3, 1.5, 1.5], gap="small")
            _existing = st.session_state["crit_thresholds"].get(_c, {})
            with _tc[0]:
                st.markdown(f"**{_c}**")
            with _tc[1]:
                _min_v = st.number_input(f"Min {_c}", value=_existing.get("min"), key=f"th_min_{_c}", label_visibility="collapsed", format="%.2f")
            with _tc[2]:
                _max_v = st.number_input(f"Max {_c}", value=_existing.get("max"), key=f"th_max_{_c}", label_visibility="collapsed", format="%.2f")
            _th = {}
            if _min_v is not None and _min_v != 0.0:
                _th["min"] = float(_min_v)
            if _max_v is not None and _max_v != 0.0:
                _th["max"] = float(_max_v)
            if _th:
                st.session_state["crit_thresholds"][_c] = _th
            elif _c in st.session_state.get("crit_thresholds", {}):
                del st.session_state["crit_thresholds"][_c]

    st.divider()

    # ── Senaryo Modu ──
    st.markdown(f"**{tt('🎭 Çoklu Paydaş Senaryo Analizi', '🎭 Multi-Stakeholder Scenario Analysis')}**")
    st.caption(tt(
        "Farklı paydaş perspektiflerini (yönetim, mühendislik, müşteri vb.) modellemek için 2-5 ağırlık senaryosu tanımlayın. Her senaryo ayrı çalıştırılır ve sonuçlar karşılaştırılır.",
        "Define 2-5 weight scenarios to model different stakeholder perspectives (management, engineering, customer, etc.). Each runs separately and results are compared."
    ))
    _sc_enabled = st.checkbox(tt("Senaryo analizi kullan", "Use scenario analysis"), key="scenario_enabled")
    if _sc_enabled and criteria:
        _n_sc = st.slider(tt("Senaryo sayısı", "Number of scenarios"), 2, 5, 2, key="n_scenarios")
        if "scenario_weights" not in st.session_state:
            st.session_state["scenario_weights"] = {}
        for i in range(_n_sc):
            _sc_name = st.text_input(tt(f"Senaryo {i+1} adı", f"Scenario {i+1} name"), value=st.session_state.get(f"sc_name_{i}", tt(f"Senaryo {i+1}", f"Scenario {i+1}")), key=f"sc_name_{i}")
            _sc_w = {}
            _sc_cols = st.columns(min(len(criteria), 4))
            for j, c in enumerate(criteria):
                with _sc_cols[j % len(_sc_cols)]:
                    _prev = st.session_state.get("scenario_weights", {}).get(_sc_name, {}).get(c, round(1.0/len(criteria), 3))
                    _sc_w[c] = st.number_input(c, value=_prev, min_value=0.0, max_value=1.0, step=0.01, key=f"sc_w_{i}_{c}", format="%.3f")
            _ws = sum(_sc_w.values())
            if _ws > 0:
                _sc_w = {c: v/_ws for c, v in _sc_w.items()}
            st.session_state["scenario_weights"][_sc_name] = _sc_w

if st.button(tt("🚀 Analiz Zamanı", "🚀 Run Analysis"), use_container_width=True, key="btn_run_analysis_main"):
    st.session_state["panel_run_warnings"] = []
    if not weight_method:
        st.error(tt("❗ Lütfen bir ağırlıklandırma yöntemi seçin.", "❗ Please select a weighting method."))
        st.stop()
    if needs_ranking and not ranking_methods_selected:
        st.error(tt("❗ Lütfen en az bir sıralama yöntemi seçin.", "❗ Please select at least one ranking method."))
        st.stop()
    if ("Manuel" in weight_mode or "Manual" in weight_mode) and (not manual_weights_valid):
        st.error(tt("Manuel ağırlık için her seçili kriterde geçerli bir pozitif ham değer girin.", "Enter a valid positive raw value for every selected criterion in manual weighting."))
        st.stop()
    if panel_mode and (not panel_year_col or panel_year_col not in working.columns):
        st.error(tt("Panel veri için geçerli bir yıl sütunu seçmelisiniz.", "You must select a valid year column for panel data."))
        st.stop()
    if panel_mode:
        _panel_sel_count = len(st.session_state.get("panel_selected_years", []))
        if _panel_sel_count > MAX_PANEL_YEAR_SELECTION:
            st.error(
                tt(
                    f"Panel analizde en fazla {MAX_PANEL_YEAR_SELECTION} dönem seçebilirsiniz.",
                    f"You can select at most {MAX_PANEL_YEAR_SELECTION} periods for panel analysis.",
                )
            )
            st.stop()
    _analysis_event_payload = {
        "scope": "panel" if panel_mode else "single",
        "rows": int(working.shape[0]),
        "columns": int(working.shape[1]),
        "criteria_count": int(len(criteria)),
        "weight_method": str(weight_method),
        "ranking_methods": list(ranking_methods_selected),
        "panel_year_count": int(len(st.session_state.get("panel_selected_years", []))),
    }
    access.track_event("analysis_started", _analysis_event_payload)
    with st.spinner(tt("Laboratuvar çalışıyor, veriler işleniyor...", "Running analysis... processing data...")):
        st.session_state.docx_buffer = None
        main_rank = primary_rank_method if primary_rank_method else (ranking_methods_selected[0] if ranking_methods_selected else None)
        comp_ranks = ranking_methods_selected if len(ranking_methods_selected) > 1 else []
        if main_rank and main_rank not in ranking_methods_selected:
            st.error(tt("Ana sıralama yöntemi seçimi geçersiz. Lütfen yöntemleri yeniden seçin.", "Invalid primary ranking selection. Please reselect methods."))
            st.stop()
        st.session_state["weight_method_pref"] = weight_method
        st.session_state["ranking_prefs"] = ranking_methods_selected

        _config_kwargs: Dict[str, Any] = {
            "criteria": criteria,
            "criteria_types": criteria_types,
            "weight_method": weight_method,
            "weight_mode": weight_mode_key,
            "manual_weights": manual_weights,
            "ranking_method": main_rank,
            "compare_methods": comp_ranks,
            "vikor_v": vikor_v,
            "waspas_lambda": waspas_lambda,
            "codas_tau": codas_tau,
            "cocoso_lambda": cocoso_lambda,
            "gra_rho": gra_rho,
            "promethee_pref_func": promethee_pref_func,
            "promethee_q": promethee_q,
            "promethee_p": promethee_p,
            "promethee_s": promethee_s,
            "fuzzy_spread": fuzzy_spread,
            "sensitivity_iterations": int(max(0, min(MAX_SENSITIVITY_ITERATIONS, sensitivity_iterations))),
            "sensitivity_sigma": sensitivity_sigma,
        }
        try:
            _cfg_sig = inspect.signature(me.AnalysisConfig)
            if "run_heavy_robustness" in _cfg_sig.parameters:
                _config_kwargs["run_heavy_robustness"] = bool(run_heavy_robustness)
        except Exception:
            _config_kwargs["run_heavy_robustness"] = bool(run_heavy_robustness)

        try:
            config = me.AnalysisConfig(**_config_kwargs)
        except TypeError:
            _config_kwargs.pop("run_heavy_robustness", None)
            config = me.AnalysisConfig(**_config_kwargs)

        # ── Ön eleme filtrelemesi ──
        _thresholds = st.session_state.get("crit_thresholds") or {}
        if _thresholds:
            working, _eliminated = me.apply_threshold_filter(working, criteria, criteria_types, _thresholds)
            if _eliminated:
                _elim_names = sorted(set(e["alternative"] for e in _eliminated))
                st.warning(tt(f"🚧 Ön eleme: {len(_elim_names)} alternatif elenmiştir: {', '.join(_elim_names)}", f"🚧 Pre-screening: {len(_elim_names)} alternative(s) eliminated: {', '.join(_elim_names)}"))
            if len(working) < 2:
                st.error(tt("Ön eleme sonrası 2'den az alternatif kaldı.", "Fewer than 2 alternatives after pre-screening."))
                st.stop()
            st.session_state["eliminated_alternatives"] = _eliminated
        else:
            st.session_state["eliminated_alternatives"] = []

        start = time.time()
        try:
            if panel_mode:
                _all_year_labels = _sorted_panel_years(working[panel_year_col])
                _user_sel_years  = st.session_state.get("panel_selected_years", [])
                year_labels = [y for y in _all_year_labels if y in _user_sel_years]
                if not year_labels:
                    raise ValueError(tt("Panel analiz için en az bir dönem seçmelisiniz.", "Select at least one period for panel analysis."))
                panel_results: Dict[str, Dict[str, Any]] = {}
                panel_failures: List[str] = []
                panel_run_config = config
                panel_run_weight_mode_key = weight_mode_key
                panel_strategy_effective = panel_weight_strategy
                if len(year_labels) < 2:
                    st.session_state["panel_run_warnings"] = [
                        tt(
                            "Panel analizde yalnızca tek dönem seçildi. Trend yorumları sınırlı olacaktır.",
                            "Only one period is selected in panel analysis. Trend interpretation will be limited.",
                        )
                    ]
                if panel_weight_strategy == "global" and weight_mode_key == "objective":
                    _pool_mask = working[panel_year_col].map(_panel_label).isin(year_labels)
                    pooled_slice = working.loc[_pool_mask, criteria].copy()
                    if len(pooled_slice) < 2:
                        raise ValueError(
                            tt(
                                "Global ağırlık için seçili yıllarda en az 2 gözlem gerekli.",
                                "Global weighting requires at least 2 observations across selected years.",
                            )
                        )
                    pooled_weights, _ = me.compute_objective_weights(
                        pooled_slice,
                        criteria,
                        criteria_types,
                        weight_method,
                        fuzzy_spread=fuzzy_spread,
                    )
                    _panel_cfg_kwargs = dict(_config_kwargs)
                    _panel_cfg_kwargs["weight_mode"] = "manual"
                    _panel_cfg_kwargs["manual_weights"] = {c: float(pooled_weights.get(c, 0.0)) for c in criteria}
                    try:
                        _cfg_sig_local = inspect.signature(me.AnalysisConfig)
                        if "run_heavy_robustness" in _cfg_sig_local.parameters:
                            _panel_cfg_kwargs["run_heavy_robustness"] = False
                        else:
                            _panel_cfg_kwargs.pop("run_heavy_robustness", None)
                    except Exception:
                        _panel_cfg_kwargs.pop("run_heavy_robustness", None)
                    try:
                        panel_run_config = me.AnalysisConfig(**_panel_cfg_kwargs)
                    except TypeError:
                        _panel_cfg_kwargs.pop("run_heavy_robustness", None)
                        panel_run_config = me.AnalysisConfig(**_panel_cfg_kwargs)
                    panel_run_weight_mode_key = "manual"
                    st.session_state["panel_run_warnings"] = list(st.session_state.get("panel_run_warnings", [])) + [
                        tt(
                            "Global ağırlık stratejisi aktif: seçili tüm yıllardan tek ağırlık vektörü üretildi ve tüm yıllara uygulandı.",
                            "Global weight strategy active: one pooled weight vector was computed from selected years and applied to all years.",
                        )
                    ]
                elif panel_weight_strategy == "global":
                    panel_strategy_effective = "yearly"
                    st.session_state["panel_run_warnings"] = list(st.session_state.get("panel_run_warnings", [])) + [
                        tt(
                            "Global ağırlık stratejisi objektif modda çalışır. Mevcut ağırlık modu için varsayılan akış kullanıldı.",
                            "Global weight strategy works in objective mode. The default flow was used for the current weighting mode.",
                        )
                    ]
                for year_label in year_labels:
                    _year_mask = _panel_mask(working, panel_year_col, year_label)
                    year_slice = working.loc[_year_mask, criteria].copy()
                    if year_slice.empty:
                        panel_failures.append(
                            tt(
                                f"{year_label}: veri bulunamadı.",
                                f"{year_label}: no rows found.",
                            )
                        )
                        continue
                    if len(year_slice) < 2:
                        panel_failures.append(
                            tt(
                                f"{year_label}: en az 2 alternatif gerekli.",
                                f"{year_label}: at least 2 alternatives are required.",
                            )
                        )
                        continue
                    if panel_entity_col and panel_entity_col in working.columns and panel_entity_col != panel_year_col:
                        entity_vals = working.loc[_year_mask, panel_entity_col].tolist()
                        year_slice.index = _make_unique_labels(entity_vals, fallback_prefix="A")
                    else:
                        year_slice.index = [f"A{idx+1}" for idx in range(len(year_slice))]
                    try:
                        year_result = _run_single_analysis_bundle(
                            data_slice=year_slice,
                            criteria=criteria,
                            criteria_types=criteria_types,
                            config=panel_run_config,
                            weight_mode_key=panel_run_weight_mode_key,
                            weight_method=weight_method,
                            main_rank=main_rank,
                        )
                    except Exception as year_exc:
                        panel_failures.append(
                            tt(
                                f"{year_label}: hesaplama hatası ({_safe_error_code(year_exc)}).",
                                f"{year_label}: calculation error ({_safe_error_code(year_exc)}).",
                            )
                        )
                        continue
                    year_result["panel_year"] = year_label
                    year_result["panel_year_column"] = panel_year_col
                    year_result["panel_entity_column"] = panel_entity_col
                    year_result["panel_weight_strategy"] = panel_strategy_effective
                    panel_results[year_label] = year_result
                if panel_failures:
                    _n_fail = len(panel_failures)
                    _preview = "; ".join(panel_failures[:3])
                    _more = tt(f" (+{_n_fail-3} dönem daha)", f" (+{_n_fail-3} more periods)") if _n_fail > 3 else ""
                    st.session_state["panel_run_warnings"] = list(st.session_state.get("panel_run_warnings", [])) + [
                        tt(
                            f"Panel analizde {_n_fail} dönem atlandı: {_preview}{_more}",
                            f"{_n_fail} period(s) were skipped in panel analysis: {_preview}{_more}",
                        )
                    ]
                if not panel_results:
                    st.session_state["panel_results"] = None
                    st.session_state["panel_years"] = []
                    st.session_state["panel_view_choice"] = None
                    st.session_state["panel_run_warnings"] = list(st.session_state.get("panel_run_warnings", [])) + [
                        tt(
                            "Seçili yıllarda panel sonuç üretilemedi; analiz tek dönem modunda tüm veriyle çalıştırıldı.",
                            "No valid panel result could be produced for selected years; analysis was run in single-period mode on full data.",
                        )
                    ]
                    result = _run_single_analysis_bundle(
                        data_slice=working,
                        criteria=criteria,
                        criteria_types=criteria_types,
                        config=config,
                        weight_mode_key=weight_mode_key,
                        weight_method=weight_method,
                        main_rank=main_rank,
                    )
                else:
                    default_year = list(panel_results.keys())[-1]
                    st.session_state["panel_results"] = panel_results
                    st.session_state["panel_years"] = list(panel_results.keys())
                    st.session_state["panel_view_choice"] = default_year
                    result = panel_results[default_year]
            else:
                result = _run_single_analysis_bundle(
                    data_slice=working,
                    criteria=criteria,
                    criteria_types=criteria_types,
                    config=config,
                    weight_mode_key=weight_mode_key,
                    weight_method=weight_method,
                    main_rank=main_rank,
                )
                st.session_state["panel_results"] = None
                st.session_state["panel_years"] = []
                st.session_state["panel_view_choice"] = None
                st.session_state["panel_run_warnings"] = []
        except Exception as exc:
            access.track_event(
                "analysis_failed",
                {
                    **_analysis_event_payload,
                    "error_code": _safe_error_code(exc),
                },
            )
            st.error(tt("Analiz işlemi tamamlanamadı. Parametreleri kontrol edip tekrar deneyin.", "The analysis could not be completed. Please review parameters and try again."))
            st.caption(tt(f"Hata kodu: {_safe_error_code(exc)}", f"Error code: {_safe_error_code(exc)}"))
            st.stop()

        result["analysis_time"] = time.time() - start
        result["criteria"] = list(criteria)
        result["criteria_types"] = dict(criteria_types)
        result["data_source_id"] = st.session_state.get("data_source_id")
        if st.session_state.get("panel_results"):
            for year_result in st.session_state["panel_results"].values():
                year_result["analysis_time"] = result["analysis_time"]
                year_result["criteria"] = list(criteria)
                year_result["criteria_types"] = dict(criteria_types)
                year_result["data_source_id"] = st.session_state.get("data_source_id")
        st.session_state["analysis_result"] = result
        st.session_state["report_docx"] = None
        st.session_state["download_blob_cache"] = {}
        st.session_state["download_blob_sig"] = None
        # ── Senaryo analizi ──
        if st.session_state.get("scenario_enabled") and st.session_state.get("scenario_weights") and main_rank:
            try:
                st.session_state["scenario_results"] = me.run_scenario_analysis(
                    working, criteria, criteria_types, st.session_state["scenario_weights"], main_rank,
                    vikor_v=vikor_v, waspas_lambda=waspas_lambda, codas_tau=codas_tau,
                    cocoso_lambda=cocoso_lambda, gra_rho=gra_rho, promethee_pref_func=promethee_pref_func,
                    promethee_q=promethee_q, promethee_p=promethee_p, promethee_s=promethee_s, fuzzy_spread=fuzzy_spread)
            except Exception:
                st.session_state["scenario_results"] = None
        else:
            st.session_state["scenario_results"] = None
        access.track_event(
            "analysis_completed",
            {
                **_analysis_event_payload,
                "duration_seconds": round(float(result["analysis_time"]), 4),
                "effective_scope": "panel" if st.session_state.get("panel_results") else "single",
                "panel_result_count": int(len(st.session_state.get("panel_results") or {})),
            },
        )
        st.rerun()

# ---------------------------------------------------------
# SONUÇLARIN GÖSTERİMİ
# ---------------------------------------------------------
result = st.session_state.get("analysis_result")
panel_results = st.session_state.get("panel_results")
_panel_all_choice = "__ALL_YEARS__"
panel_all_requested = False
panel_active_year = None
if isinstance(panel_results, dict) and panel_results:
    panel_years = st.session_state.get("panel_years") or list(panel_results.keys())
    view_choice = st.session_state.get("panel_view_choice")
    if view_choice == _panel_all_choice:
        panel_all_requested = True
        view_choice = panel_years[-1]
    elif view_choice not in panel_results:
        view_choice = panel_years[-1]
        st.session_state["panel_view_choice"] = view_choice
    panel_active_year = view_choice
    result = panel_results[view_choice]
if result is None:
    st.stop()
if not _ensure_plotly_support():
    st.error(tt("Grafik modulu yuklenemedi. Lutfen Plotly kurulumunu kontrol edin.", "Chart module could not be loaded. Please check the Plotly installation."))
    st.stop()

_show_step_hint_once(
    "stage_results",
    "Sonuçlar hazır. Önce yapı özetine, ardından KPI bandına ve sekmelere bakabilirsiniz.",
    "Results are ready. Start with the setup summary, then use the KPI strip and tabs.",
    icon="✅",
)
st.markdown(f"<h3 style='font-size: 1.2rem; margin-top:1rem;'>📥 {tt('Analiz Sonuçları ve Raporlar', 'Analysis Results and Reports')}</h3>", unsafe_allow_html=True)
_panel_run_warnings = st.session_state.get("panel_run_warnings", [])
if isinstance(_panel_run_warnings, list) and _panel_run_warnings:
    for _warn in _panel_run_warnings:
        st.warning(str(_warn))

weight_df     = result["weights"]["table"]
ranking_table = result.get("ranking", {}).get("table")
alt_names     = st.session_state.get("alt_names", {})

_lead_alt = "—"
if isinstance(ranking_table, pd.DataFrame) and not ranking_table.empty:
    _alt_col = col_key(ranking_table, "Alternatif", "Alternative")
    _raw_lead = str(ranking_table.iloc[0][_alt_col])
    _lead_alt = alt_names.get(_raw_lead, _raw_lead)
_rank_method = result.get("ranking", {}).get("method") or "—"
_lead_alt_safe = _safe_html_text(_lead_alt)
_rank_method_safe = _safe_html_text(_rank_method)
_sensitivity_payload = result.get("sensitivity") or {}
_stability = _sensitivity_payload.get("top_stability")
_stability_txt = f"%{float(_stability) * 100:.1f}" if _stability is not None else "—"
_runtime = result.get("analysis_time")
_runtime_txt = f"{float(_runtime):.2f} {tt('sn', 'sec')}" if _runtime is not None else "—"
_shape = result.get("selected_data", pd.DataFrame()).shape
_result_ranking_prefs = st.session_state.get("ranking_prefs") or []
_result_scope_text = (
    f"{tt('Panel veri', 'Panel data')} · {len(panel_results)} {tt('dönem', 'periods')}"
    if isinstance(panel_results, dict) and panel_results
    else tt("Tek dönem", "Single-period")
)
_result_scope_note = (
    f"{_result_scope_text} · {tt('ekranda', 'showing')} {panel_active_year}"
    if panel_active_year is not None
    else _result_scope_text
)
_result_ranking_summary = (
    _preview_list_text([method_display_name(m) for m in _result_ranking_prefs], 3)
    if _result_ranking_prefs
    else (method_display_name(result.get("ranking", {}).get("method")) if result.get("ranking", {}).get("method") else tt("Sıralama uygulanmadı", "Ranking not applied"))
)
_render_tracking_panel(
    tt("Analiz Yapı Özeti", "Analysis Setup Snapshot"),
    tt(
        "KPI ve sekmeler bu kurulumdan üretildi.",
        "KPI strip and tabs were generated from this setup.",
    ),
    [
        (tt("Veri kaynağı", "Data source"), _data_source_label(result.get("data_source_id") or st.session_state.get("data_source_id")), "green"),
        (tt("Analiz amacı", "Objective"), _step1_objective_summary, "blue"),
        (tt("Kapsam", "Scope"), _result_scope_note, "blue"),
        (tt("Ağırlık yöntemi", "Weighting method"), method_display_name(result.get("weights", {}).get("method") or tt("Belirlenmedi", "Not set")), "amber"),
        (tt("Sıralama paketi", "Ranking package"), _result_ranking_summary, "rose" if result.get("ranking", {}).get("method") else "slate"),
        (tt("Kriterler", "Criteria"), f"{len(result.get('criteria', []) or [])} · {_preview_list_text(result.get('criteria', []), 4)}", "green"),
    ],
    note=(
        tt(
            f"Not: Ana ekran şu anda {panel_active_year} dönemi için gösteriliyor.",
            f"Note: The main screen is currently shown for period {panel_active_year}.",
        )
        if panel_active_year is not None
        else tt(
            "Not: Bu özet raporların dayandığı kurulumdur.",
            "Note: This snapshot is the setup behind the reports.",
        )
    ),
    tone="blue",
    expanded=True,
    icon="📌",
)
st.markdown(
    f"""
    <div class="kpi-strip">
        <div class="kpi-grid">
            <div class="kpi-item" title="{tt('En yüksek skoru alan alternatif','Highest scoring alternative')}"><div class="kpi-label">{tt('Lider Alternatif', 'Leading Alternative')}</div><div class="kpi-value">{_lead_alt_safe}</div></div>
            <div class="kpi-item" title="{tt('Alternatifleri sıralamak için kullanılan yöntem','Method used to rank alternatives')}"><div class="kpi-label">{tt('Sıralama Yöntemi', 'Ranking Method')}</div><div class="kpi-value">{_rank_method_safe}</div></div>
            <div class="kpi-item" title="{tt('MC simülasyonunda liderin birinciliğini koruma oranı. ≥%80: yüksek, %60-80: orta, <%60: düşük','% of MC runs where leader stays #1. ≥80%: high, 60-80%: moderate, <60%: low')}"><div class="kpi-label">{tt('Kararlılık', 'Stability')}</div><div class="kpi-value">{_stability_txt}</div></div>
            <div class="kpi-item" title="{tt('Analiz tamamlanma süresi','Analysis completion time')}"><div class="kpi-label">{tt('Analiz Süresi', 'Analysis Time')}</div><div class="kpi-value">{_runtime_txt}</div></div>
            <div class="kpi-item" title="{tt('Alternatif sayısı × kriter sayısı','Alternatives × criteria')}"><div class="kpi-label">{tt('Veri Boyutu', 'Data Size')}</div><div class="kpi-value">{_shape[0]}×{_shape[1]}</div></div>
        </div>
    </div>
    """,
    unsafe_allow_html=True,
)

with st.expander(tt("🧮 Hesaplama Detaylarını Gör", "🧮 View Calculation Details"), expanded=False):
    if isinstance(panel_results, dict) and panel_results:
        _all_years_label = tt("Tüm yılları görmek istiyorum", "I want to see all years")
        _year_options = list(st.session_state.get("panel_years") or panel_results.keys()) + [_all_years_label]
        _stored_choice = st.session_state.get("panel_view_choice")
        _display_choice = _all_years_label if _stored_choice == _panel_all_choice else (_stored_choice or _year_options[0])
        if _display_choice not in _year_options:
            _display_choice = _year_options[0]
        _selected_year_choice = st.selectbox(
            tt("Sonucu görmek istediğiniz yılı seçin", "Select the year whose results you want to view"),
            _year_options,
            index=_year_options.index(_display_choice),
            key="panel_year_result_selector",
        )
        _new_panel_choice = _panel_all_choice if _selected_year_choice == _all_years_label else _selected_year_choice
        if _new_panel_choice != st.session_state.get("panel_view_choice"):
            st.session_state["panel_view_choice"] = _new_panel_choice
            st.rerun()
        if st.session_state.get("panel_view_choice") == _panel_all_choice:
            st.info(
                tt(
                    "Tüm yıllar seçildi. Ana sekmelerde tek yıl gösterilir; yıllar arası toplu sıralama özeti paneli aşağıda sunulur.",
                    "All years selected. Main tabs still show one year at a time; the cross-year summary panel is shown below.",
                )
            )
        if panel_active_year is not None:
            st.caption(tt(f"Ekranda gösterilen yıl: {panel_active_year}", f"Year currently shown on screen: {panel_active_year}"))
    if "show_calc_details" not in st.session_state:
        st.session_state["show_calc_details"] = False
    if st.button(tt("🔍 Hesaplama Adımlarını Göster / Gizle", "🔍 Show / Hide Calculation Steps"), use_container_width=True, key="btn_toggle_calc_details"):
        st.session_state["show_calc_details"] = not st.session_state["show_calc_details"]
    if st.session_state["show_calc_details"]:
        dt1, dt2, dt3 = st.tabs([tt("⚖️ Ağırlık Adımları", "⚖️ Weighting Steps"), tt("🏆 Sıralama Adımları", "🏆 Ranking Steps"), tt("🎲 Duyarlılık Adımları", "🎲 Sensitivity Steps")])

        with dt1:
            wm = result["weights"]["method"]
            st.markdown(f"**{tt('Yöntem', 'Method')}:** `{method_display_name(wm)}`")
            st.code(get_method_steps(wm, "weight"), language="text")
            w_frames = _extract_detail_tables(result["weights"].get("details", {}))
            if w_frames:
                w_key = st.selectbox(tt("Ağırlık detay tablosu", "Weighting detail table"), list(w_frames.keys()), key="calc_weight_tbl")
                render_table(localize_df(w_frames[w_key].head(250)))
            else:
                st.info(tt("Ağırlık yöntemi için detay tablo bulunamadı.", "No detail table found for the weighting method."))

        with dt2:
            rm = result.get("ranking", {}).get("method")
            if rm:
                comp = result.get("comparison", {}) or {}
                method_details = dict(comp.get("method_details") or {})
                method_details.setdefault(str(rm), result.get("ranking", {}).get("details", {}) or {})
                method_names = [m for m in method_details.keys() if str(m).strip()]

                if len(method_names) > 1:
                    method_tabs = st.tabs([str(m) for m in method_names])
                    for meth_name, meth_tab in zip(method_names, method_tabs):
                        with meth_tab:
                            st.markdown(f"**{tt('Yöntem', 'Method')}:** `{meth_name}`")
                            st.code(get_method_steps(str(meth_name), "ranking"), language="text")
                            r_frames = _extract_detail_tables(method_details.get(str(meth_name), {}))
                            if r_frames:
                                r_key = st.selectbox(
                                    tt("Sıralama detay tablosu", "Ranking detail table"),
                                    list(r_frames.keys()),
                                    key=f"calc_rank_tbl_{str(meth_name)}",
                                )
                                render_table(localize_df(r_frames[r_key].head(250)))
                            else:
                                st.info(tt("Sıralama yöntemi için detay tablo bulunamadı.", "No detail table found for the ranking method."))
                else:
                    st.markdown(f"**{tt('Yöntem', 'Method')}:** `{rm}`")
                    st.code(get_method_steps(rm, "ranking"), language="text")
                    r_frames = _extract_detail_tables(result.get("ranking", {}).get("details", {}))
                    if r_frames:
                        r_key = st.selectbox(tt("Sıralama detay tablosu", "Ranking detail table"), list(r_frames.keys()), key="calc_rank_tbl")
                        render_table(localize_df(r_frames[r_key].head(250)))
                    else:
                        st.info(tt("Sıralama yöntemi için detay tablo bulunamadı.", "No detail table found for the ranking method."))
            else:
                st.info(tt("Bu analizde sıralama yöntemi seçilmedi.", "No ranking method was selected in this analysis."))

        with dt3:
            sens = result.get("sensitivity")
            if sens:
                st.markdown(f"**{tt('Duyarlılık yaklaşımı', 'Sensitivity approach')}:** {tt('Lokal senaryolar + Monte Carlo bozulma analizi', 'Local scenarios + Monte Carlo perturbation analysis')}")
                st.code(
                    tt(
                        "1) Baz ağırlıkla referans sıralama\n2) Her kriterde ±%10/±%20 lokal değişim\n3) Spearman ile sıra benzerliği\n4) Monte Carlo ile birincilik oranı",
                        "1) Baseline ranking with base weights\n2) ±10%/±20% local perturbations per criterion\n3) Rank similarity via Spearman\n4) First-place frequency via Monte Carlo",
                    ),
                    language="text",
                )
                local_df = sens.get("local_scenarios")
                mc_df = sens.get("monte_carlo_summary")
                if isinstance(local_df, pd.DataFrame) and not local_df.empty:
                    st.markdown(f"**{tt('Lokal senaryolar', 'Local scenarios')}**")
                    render_table(localize_df(local_df.head(250)))
                if isinstance(mc_df, pd.DataFrame) and not mc_df.empty:
                    st.markdown(f"**{tt('Monte Carlo özeti', 'Monte Carlo summary')}**")
                    render_table(localize_df(mc_df.head(250)))
            else:
                st.info(tt("Duyarlılık analizi bu çalışmada üretilmedi.", "Sensitivity analysis was not generated in this study."))

tabs_list = [
    tt("📊 Temel İstatistik", "📊 Basic Statistics"),
    tt("⚖️ Ağırlıklar", "⚖️ Weights"),
]
if needs_ranking:
    tabs_list.extend([
        tt("🏆 Sıralama", "🏆 Ranking"),
        tt("🔁 Karşılaştırma", "🔁 Comparison"),
        tt("🛡️ Sağlamlık", "🛡️ Robustness"),
        tt("🎲 Dayanıklılık", "🎲 Durability"),
    ])
else:
    tabs_list.append(tt("🛡️ Sağlamlık", "🛡️ Robustness"))
tabs_list.append(tt("📥 Çıktı", "📥 Output"))
tabs = st.tabs(tabs_list)

_tab_stats = 0
_tab_weights = 1
if needs_ranking:
    _tab_ranking = 2
    _tab_comparison = 3
    _tab_robustness = 4
    _tab_durability = 5
    _tab_output = 6
else:
    _tab_robustness = 2
    _tab_output = 3

with tabs[_tab_stats]:
    _stat_comment = gen_stat_commentary(result)
    render_tab_assistant(_stat_comment, key="stat")
    with st.expander(f"📋 {tt('Tablolar', 'Tables')}", expanded=True):
        st.markdown(f"##### 📊 {tt('Betimleyici İstatistikler', 'Descriptive Statistics')}")
        _stats_disp = localize_df(result["stats"]).copy()
        _num_cols = _stats_disp.select_dtypes(include=[np.number]).columns
        _stats_disp[_num_cols] = _stats_disp[_num_cols].round(3)
        render_table(_stats_disp)
        with st.container():
            st.markdown(f'<div class="commentary-box">{_safe_plain_commentary_html(_stat_comment)}</div>', unsafe_allow_html=True)

    with st.expander(f"📊 {tt('Şekiller', 'Figures')}", expanded=True):
        s1, s2 = st.columns([1, 1])
        with s1:
            st.plotly_chart(fig_box_plots(result["selected_data"], criteria), use_container_width=True)
            with st.container():
                st.markdown(f'<div class="commentary-box">{tt("Kutu grafikleri her kriterin dağılım profilini gösterir. Geniş kutular ve aykırı noktalar, o kriterin yüksek ayırt edici gücüne işaret eder.", "Box plots show the distribution profile of each criterion. Wide boxes and outlier points indicate high discriminative power.")}</div>', unsafe_allow_html=True)
        with s2:
            corr_df = result.get("correlation_matrix")
            if isinstance(corr_df, pd.DataFrame):
                fig_corr = px.imshow(
                    corr_df, text_auto=".2f", aspect="auto",
                    color_continuous_scale="RdBu_r", zmin=-1, zmax=1,
                    title=tt("Kriterler Arası Korelasyon Haritası", "Correlation Heatmap Among Criteria"),
                )
                fig_corr.update_layout(height=420, **_THEME)
                st.plotly_chart(fig_corr, use_container_width=True)
                with st.container():
                    st.markdown(f'<div class="commentary-box">{tt("Korelasyon ısı haritası, kriterler arasındaki doğrusal ilişkilerin yoğunluğunu gösterir. Kırmızı: güçlü pozitif, mavi: güçlü negatif ilişki. Yüksek korelasyon (|ρ|>0.75) bilgi tekrarına işaret eder.", "The correlation heatmap shows the intensity of linear relationships between criteria. Red: strong positive, blue: strong negative. High correlation (|ρ|>0.75) suggests information redundancy.")}</div>', unsafe_allow_html=True)
            else:
                st.info(tt("Korelasyon matrisi bulunamadı.", "Correlation matrix not available."))

with tabs[_tab_weights]:
    _weight_comment = gen_weight_commentary(result)
    render_tab_assistant(_weight_comment, key="weight")
    weight_col = col_key(weight_df, "Ağırlık", "Weight")
    top3_weights = weight_df.sort_values(weight_col, ascending=False).head(3).reset_index(drop=True)

    with st.expander(f"📋 {tt('Tablolar', 'Tables')}", expanded=True):
        w1, w2 = st.columns(2)
        with w1:
            st.markdown(f"##### 🥇 {tt('İlk 3 Önemli Kriter', 'Top 3 Important Criteria')}")
            _top3_disp = localize_df(top3_weights)
            _top3_w_col = col_key(_top3_disp, "Ağırlık", "Weight")
            if _top3_w_col in _top3_disp.columns:
                _top3_disp[_top3_w_col] = pd.to_numeric(_top3_disp[_top3_w_col], errors="coerce").round(6)
            render_table(_top3_disp)
        with w2:
            st.markdown(f"##### ⚖️ {tt('Tüm Kriter Ağırlıkları', 'All Criterion Weights')}")
            _weight_disp = localize_df(weight_df)
            _weight_disp_col = col_key(_weight_disp, "Ağırlık", "Weight")
            if _weight_disp_col in _weight_disp.columns:
                _weight_disp[_weight_disp_col] = pd.to_numeric(_weight_disp[_weight_disp_col], errors="coerce").round(6)
            render_table(_weight_disp)
        with st.container():
            st.markdown(f'<div class="commentary-box">{_safe_plain_commentary_html(_weight_comment)}</div>', unsafe_allow_html=True)

    with st.expander(f"📊 {tt('Şekiller', 'Figures')}", expanded=True):
        fig_col1, fig_col2 = st.columns(2)
        with fig_col1:
            st.plotly_chart(fig_weight_bar(weight_df), use_container_width=True)
            with st.container():
                st.markdown(f'<div class="commentary-box">{tt("Ağırlık çubuğu grafiği, kriterlerin göreli önem sıralamasını ve büyüklüklerini karşılaştırmalı olarak gösterir. En uzun çubuk, analiz kararını en güçlü etkileyen kriteri temsil eder.", "The weight bar chart compares the relative importance and magnitude of criteria. The longest bar represents the criterion with the highest influence on the analysis decision.")}</div>', unsafe_allow_html=True)
        with fig_col2:
            st.plotly_chart(fig_weight_radar(weight_df), use_container_width=True)
            with st.container():
                st.markdown(f'<div class="commentary-box">{tt("Radar grafiği, ağırlıkların çok boyutlu dağılımını görsel olarak sunar. Dengeli bir dağılım, kararın tek bir kritere bağımlı olmadığına işaret eder; asimetrik yapı ise belirli bir kriterin hakimiyetini gösterir.", "The radar chart visually presents the multi-dimensional distribution of weights. A balanced distribution indicates the decision is not dependent on a single criterion; asymmetric structure shows dominance of a specific criterion.")}</div>', unsafe_allow_html=True)

with tabs[_tab_robustness]:
    _wr = result.get("weight_robustness", {}) or {}
    _rob_comment = gen_weight_robustness_commentary(result)
    st.markdown(f"### 🛡️ {tt('Ağırlık Sağlamlık Testleri', 'Weight Robustness Tests')}")
    if _wr.get("info"):
        st.info(str(_wr.get("info")))
    elif _wr.get("error"):
        st.warning(tt("Ağırlık sağlamlık testleri üretilemedi.", "Weight robustness tests could not be generated."))
    else:
        _loo_mean = _wr.get("loo_mean_rho")
        _loo_min = _wr.get("loo_min_rho")
        _top_match = _wr.get("loo_top_match_rate")
        _eff_n = _wr.get("effective_criteria_n")
        _dom = _wr.get("dominance_ratio")

        m1, m2, m3, m4 = st.columns(4)
        m1.metric(tt("Ortalama Uyum (LOO)", "Mean Agreement (LOO)"), (f"{float(_loo_mean):.3f}" if _loo_mean is not None and np.isfinite(_loo_mean) else "—"))
        m2.metric(tt("Min. Uyum (LOO)", "Min. Agreement (LOO)"), (f"{float(_loo_min):.3f}" if _loo_min is not None and np.isfinite(_loo_min) else "—"))
        m3.metric(tt("Lider Kriter Tutarlılığı", "Top-Criterion Consistency"), (f"%{float(_top_match)*100:.1f}" if _top_match is not None and np.isfinite(_top_match) else "—"))
        m4.metric(tt("Etkin Kriter Sayısı", "Effective Criterion Count"), (f"{float(_eff_n):.2f}" if _eff_n is not None and np.isfinite(_eff_n) else "—"))

        with st.expander(f"💬 {tt('Sağlamlık Özet Yorumu', 'Robustness Summary Commentary')}", expanded=True):
            st.markdown(f'<div class="commentary-box">{_safe_plain_commentary_html(_rob_comment)}</div>', unsafe_allow_html=True)
        st.caption(tt("Detaylar için lütfen aşağıdaki tablo ve şekilleri inceleyin.", "For details, please review the tables and figures below."))

        _felsefe_msg = tt(
            "📌 Yöntem felsefesi: Bu testler 'seçtiğimiz objektif ağırlıklandırma, veri biraz oynadığında aynı hikâyeyi anlatıyor mu?' sorusunu yanıtlar.",
            "📌 Method philosophy: These tests answer: 'Does the chosen objective weighting tell the same story when data changes slightly?'"
        )
        st.caption(_felsefe_msg)
        if _dom is not None and np.isfinite(_dom):
            st.caption(tt(
                f"⚖️ Lider/ikinci kriter ağırlık oranı: {_dom:.2f}. Oran yükseldikçe karar tek kritere daha bağımlı hale gelir.",
                f"⚖️ Top/second criterion weight ratio: {_dom:.2f}. As this ratio grows, the decision depends more on a single criterion."
            ))

        with st.expander(f"📋 {tt('Tablolar — Detay Sağlamlık Verileri', 'Tables — Detailed Robustness Data')}", expanded=False):
            _loo_df = _wr.get("leave_one_out")
            if isinstance(_loo_df, pd.DataFrame) and not _loo_df.empty:
                st.markdown(f"##### 🔍 {tt('Leave-One-Out Analizi', 'Leave-One-Out Analysis')}")
                st.caption(tt(
                    "Her alternatif birer birer çıkarılarak ağırlık vektörünün ne kadar değiştiği ölçülür. Spearman ρ = 1.0 mükemmel kararlılık, düşük değer kırılganlık işareti.",
                    "Each alternative is removed one by one and the change in weight vector is measured. Spearman ρ = 1.0 indicates perfect stability; lower values signal fragility."
                ))
                _loo_disp = localize_df(_loo_df.copy())
                _rho_col = col_key(_loo_disp, "SpearmanRho", "SpearmanRho")
                _maxdiff_col = col_key(_loo_disp, "MaksMutlakFark", "MaxAbsoluteDiff")
                if _rho_col in _loo_disp.columns:
                    _loo_disp[_rho_col] = pd.to_numeric(_loo_disp[_rho_col], errors="coerce").round(3)
                if _maxdiff_col in _loo_disp.columns:
                    _loo_disp[_maxdiff_col] = pd.to_numeric(_loo_disp[_maxdiff_col], errors="coerce").round(4)
                render_table(_loo_disp)
                with st.container():
                    st.markdown(f'<div class="commentary-box">{tt("Leave-one-out tablosu, her alternatifin çıkarıldığında ağırlık düzeninin ne ölçüde korunduğunu gösterir. SpearmanRho sütunu 1\'e yakın değerler yüksek kararlılık anlamına gelir. MaksMutlakFark ise en fazla değişen kriter ağırlığının büyüklüğünü verir.", "The leave-one-out table shows how well the weight order is preserved when each alternative is removed. SpearmanRho values close to 1 indicate high stability. MaxAbsoluteDiff gives the magnitude of the most changed criterion weight.")}</div>', unsafe_allow_html=True)

            _boot_df = _wr.get("bootstrap_summary")
            if isinstance(_boot_df, pd.DataFrame) and not _boot_df.empty:
                st.markdown(f"##### 🎲 {tt('Bootstrap Kriter Güven Aralıkları', 'Bootstrap Criterion Confidence Intervals')}")
                st.caption(tt(
                    "Veri tekrar örneklenerek her kriterin ağırlık dağılımı simüle edilir. Dar aralık = güvenilir ağırlık, geniş aralık = belirsizlik yüksek.",
                    "Data is resampled to simulate the weight distribution for each criterion. Narrow interval = reliable weight, wide interval = high uncertainty."
                ))
                _boot_disp = localize_df(_boot_df.copy())
                _mean_col = col_key(_boot_disp, "OrtalamaAğırlık", "MeanWeight")
                _std_col = col_key(_boot_disp, "StdSapma", "StdDev")
                _l5_col = col_key(_boot_disp, "AltYuzde5", "Lower5Pct")
                _u95_col = col_key(_boot_disp, "UstYuzde95", "Upper95Pct")
                for _c in [_mean_col, _std_col, _l5_col, _u95_col]:
                    if _c in _boot_disp.columns:
                        _boot_disp[_c] = pd.to_numeric(_boot_disp[_c], errors="coerce").round(4)
                render_table(_boot_disp)
                with st.container():
                    st.markdown(f'<div class="commentary-box">{tt("Bootstrap özet tablosu, her kriterin ağırlığının istatistiksel güven aralığını gösterir. Standart sapması düşük kriterler kararlı ağırlıklara sahiptir ve analizin yapı taşını oluşturur. Geniş aralıklı kriterler için duyarlılık analizi ile çapraz doğrulama yapılması önerilir.", "The bootstrap summary table shows the statistical confidence interval for each criterion weight. Criteria with low standard deviation have stable weights and form the backbone of the analysis. Cross-validation with sensitivity analysis is recommended for criteria with wide intervals.")}</div>', unsafe_allow_html=True)

        with st.expander(f"📊 {tt('Şekiller — Sağlamlık Görselleştirmeleri', 'Figures — Robustness Visualizations')}", expanded=False):
            _boot_df2 = _wr.get("bootstrap_summary")
            if isinstance(_boot_df2, pd.DataFrame) and not _boot_df2.empty:
                _k_col = col_key(_boot_df2, "Kriter", "Criterion") if "Kriter" in _boot_df2.columns or "Criterion" in _boot_df2.columns else _boot_df2.columns[0]
                _m_col = col_key(_boot_df2, "OrtalamaAğırlık", "MeanWeight")
                _s_col = col_key(_boot_df2, "StdSapma", "StdDev")
                _l_col = col_key(_boot_df2, "AltYuzde5", "Lower5Pct")
                _u_col = col_key(_boot_df2, "UstYuzde95", "Upper95Pct")
                _boot_plot = _boot_df2.copy()
                for _c in [_m_col, _s_col, _l_col, _u_col]:
                    if _c in _boot_plot.columns:
                        _boot_plot[_c] = pd.to_numeric(_boot_plot[_c], errors="coerce")
                if all(c in _boot_plot.columns for c in [_k_col, _m_col, _l_col, _u_col]):
                    _boot_plot = _boot_plot.sort_values(_m_col, ascending=True)
                    fig_boot = go.Figure()
                    fig_boot.add_trace(go.Scatter(
                        x=_boot_plot[_m_col], y=_boot_plot[_k_col],
                        mode='markers',
                        marker=dict(color='#3D6491', size=10, symbol='circle'),
                        name=tt("Ortalama", "Mean"),
                        error_x=dict(
                            type='data',
                            symmetric=False,
                            array=(_boot_plot[_u_col] - _boot_plot[_m_col]).clip(lower=0).tolist(),
                            arrayminus=(_boot_plot[_m_col] - _boot_plot[_l_col]).clip(lower=0).tolist(),
                            color='rgba(61,100,145,0.5)',
                            thickness=2,
                            width=6,
                        ),
                    ))
                    fig_boot.update_layout(
                        height=max(280, len(_boot_plot) * 38 + 60),
                        title=tt("Bootstrap Ağırlık Güven Aralıkları (5%–95%)", "Bootstrap Weight Confidence Intervals (5%–95%)"),
                        xaxis_title=tt("Ağırlık", "Weight"),
                        yaxis_title=tt("Kriter", "Criterion"),
                        **_THEME,
                    )
                    st.plotly_chart(fig_boot, use_container_width=True)
                    with st.container():
                        st.markdown(f'<div class="commentary-box">{tt("Bootstrap güven aralığı grafiği, her kriterin ağırlığının olası aralığını gösterir. Yatay çizgi uzunluğu belirsizlik düzeyini temsil eder: kısa çizgi = güvenilir ağırlık, uzun çizgi = veri değişimine duyarlı ağırlık.", "The bootstrap confidence interval chart shows the plausible range for each criterion weight. Bar length represents uncertainty level: short bar = reliable weight, long bar = weight sensitive to data changes.")}</div>', unsafe_allow_html=True)

if needs_ranking:
    with tabs[_tab_ranking]:
        _rank_comment = gen_ranking_commentary(result, alt_names)
        render_tab_assistant(_rank_comment, key="rank")
        if ranking_table is None or ranking_table.empty:
            st.info(tt("Sıralama tablosu oluşturulamadı.", "Ranking table could not be generated."))
        else:
            _rt_disp = ranking_table.copy()
            _rt_alt_col = col_key(_rt_disp, "Alternatif", "Alternative")
            _rt_score_col = col_key(_rt_disp, "Skor", "Score")
            if alt_names:
                _rt_disp[_rt_alt_col] = _rt_disp[_rt_alt_col].astype(str).map(
                    lambda x: alt_names.get(x, x)
                )

            _rank_alt_col = col_key(ranking_table, "Alternatif", "Alternative")
            top_disp = alt_names.get(str(ranking_table.iloc[0][_rank_alt_col]), str(ranking_table.iloc[0][_rank_alt_col]))
            st.info(f"🏆 {tt('Lider Alternatif', 'Leading Alternative')}: **{top_disp}** — {tt('Yöntem', 'Method')}: {result.get('ranking', {}).get('method')}")

            if isinstance(panel_results, dict) and panel_results:
                _panel_strategies = {
                    str((yr.get("panel_weight_strategy") or "yearly")).strip().lower()
                    for yr in panel_results.values()
                    if isinstance(yr, dict)
                }
                _is_yearly_panel_view = ("yearly" in _panel_strategies) and ("global" not in _panel_strategies)
                with st.expander(tt("🗂️ Panel Yıl Bazlı Sıralama Özeti", "🗂️ Panel Year-Based Ranking Summary"), expanded=_is_yearly_panel_view):
                    st.caption(tt(
                        "Seçilen tüm yıllar birlikte özetlenir; böylece aynı alternatifin dönemler arasında kaçıncı sıraya düştüğü net görünür.",
                        "All selected years are summarized together so you can clearly see how each alternative ranks across periods.",
                    ))
                    _panel_summary_df = _panel_summary_rows(panel_results)
                    if isinstance(_panel_summary_df, pd.DataFrame) and not _panel_summary_df.empty:
                        _panel_summary_disp = localize_df(_panel_summary_df.copy())
                        _panel_score_col = col_key(_panel_summary_disp, "LiderSkor", "TopScore")
                        _panel_stab_col = col_key(_panel_summary_disp, "LiderKararlılığı", "LeaderStability")
                        _panel_strategy_col = col_key(_panel_summary_disp, "AğırlıkStratejisi", "WeightStrategy")
                        if _panel_score_col in _panel_summary_disp.columns:
                            _panel_summary_disp[_panel_score_col] = pd.to_numeric(_panel_summary_disp[_panel_score_col], errors="coerce").round(4)
                        if _panel_stab_col in _panel_summary_disp.columns:
                            _panel_summary_disp[_panel_stab_col] = pd.to_numeric(_panel_summary_disp[_panel_stab_col], errors="coerce").round(4)
                        if _panel_strategy_col in _panel_summary_disp.columns:
                            _panel_summary_disp[_panel_strategy_col] = (
                                _panel_summary_disp[_panel_strategy_col]
                                .astype(str)
                                .str.strip()
                                .str.lower()
                                .map(lambda x: tt("Global", "Global") if x == "global" else tt("Yıl Bazlı", "Year-Specific"))
                            )
                        render_table(_panel_summary_disp)
                    _panel_weight_df = _panel_weight_matrix(panel_results)
                    if isinstance(_panel_weight_df, pd.DataFrame) and not _panel_weight_df.empty:
                        st.markdown(f"##### ⚖️ {tt('Kriter Ağırlıkları (Yıl Sütunları)', 'Criterion Weights (Year Columns)')}")
                        render_table(localize_df(_panel_weight_df.head(500)))
                    _panel_rank_df = _panel_rank_matrix(panel_results)
                    if isinstance(_panel_rank_df, pd.DataFrame) and not _panel_rank_df.empty:
                        st.markdown(f"##### 📈 {tt('Alternatiflerin Yıllara Göre Sıraları', 'Alternative Ranks by Year')}")
                        render_table(localize_df(_panel_rank_df.head(500)))
                    _panel_score_df = _panel_score_matrix(panel_results)
                    if isinstance(_panel_score_df, pd.DataFrame) and not _panel_score_df.empty:
                        st.markdown(f"##### 📊 {tt('Alternatiflerin Yıllara Göre Skorları', 'Alternative Scores by Year')}")
                        render_table(localize_df(_panel_score_df.head(500)))

            with st.expander(f"📋 {tt('Tablolar', 'Tables')}", expanded=True):
                comp = result.get("comparison", {}) or {}
                _method_tables = dict(comp.get("method_tables") or {})
                _primary_method_name = str(result.get("ranking", {}).get("method") or "").strip()
                if _primary_method_name:
                    _method_tables.setdefault(_primary_method_name, ranking_table.copy())
                _method_names = [
                    str(_m) for _m, _tbl in _method_tables.items()
                    if str(_m).strip() and isinstance(_tbl, pd.DataFrame) and not _tbl.empty
                ]

                if len(_method_names) > 1:
                    st.caption(tt("Çoklu yöntem seçimi algılandı. Aşağıdaki alt sekmelerde her yöntemin sıralamasını görebilirsiniz.", "Multiple methods detected. Use the subtabs below to view each method's ranking table."))
                    _rank_method_tabs = st.tabs(_method_names)
                    for _meth_name, _meth_tab in zip(_method_names, _rank_method_tabs):
                        with _meth_tab:
                            _meth_df = _method_tables.get(_meth_name)
                            if not isinstance(_meth_df, pd.DataFrame) or _meth_df.empty:
                                st.info(tt("Bu yöntem için sıralama tablosu bulunamadı.", "Ranking table not found for this method."))
                                continue
                            _meth_disp = _meth_df.copy()
                            _meth_alt_col = col_key(_meth_disp, "Alternatif", "Alternative")
                            _meth_rank_col = col_key(_meth_disp, "Sıra", "Rank")
                            if _meth_rank_col in _meth_disp.columns:
                                _meth_disp = _meth_disp.sort_values(_meth_rank_col, ascending=True).reset_index(drop=True)
                            if alt_names and _meth_alt_col in _meth_disp.columns:
                                _meth_disp[_meth_alt_col] = _meth_disp[_meth_alt_col].astype(str).map(
                                    lambda x: alt_names.get(x, x)
                                )
                            _m1, _m2 = st.columns(2)
                            with _m1:
                                st.markdown(f"##### 🥇 {tt('İlk 5 Alternatif', 'Top 5 Alternatives')}")
                                _meth_top5 = localize_df(_meth_disp.head(5)).copy()
                                _meth_top5_score_col = col_key(_meth_top5, "Skor", "Score")
                                if _meth_top5_score_col in _meth_top5.columns:
                                    _meth_top5[_meth_top5_score_col] = pd.to_numeric(_meth_top5[_meth_top5_score_col], errors="coerce").round(4)
                                render_table(_meth_top5)
                            with _m2:
                                st.markdown(f"##### 📊 {tt('Nihai Sıralama', 'Final Ranking')}")
                                _meth_full_disp = localize_df(_meth_disp)
                                _meth_full_score_col = col_key(_meth_full_disp, "Skor", "Score")
                                if _meth_full_score_col in _meth_full_disp.columns:
                                    _meth_full_disp[_meth_full_score_col] = pd.to_numeric(_meth_full_disp[_meth_full_score_col], errors="coerce").round(4)
                                render_table(_meth_full_disp)
                else:
                    r1, r2 = st.columns(2)
                    with r1:
                        st.markdown(f"##### 🥇 {tt('İlk 5 Alternatif', 'Top 5 Alternatives')}")
                        _rt_top5 = localize_df(_rt_disp.head(5)).copy()
                        _rt_top5_score_col = col_key(_rt_top5, "Skor", "Score")
                        if _rt_top5_score_col in _rt_top5.columns:
                            _rt_top5[_rt_top5_score_col] = pd.to_numeric(_rt_top5[_rt_top5_score_col], errors="coerce").round(4)
                        render_table(_rt_top5)
                    with r2:
                        st.markdown(f"##### 📊 {tt('Nihai Sıralama', 'Final Ranking')}")
                        _rt_table_disp = localize_df(_rt_disp)
                        _rt_table_score_col = col_key(_rt_table_disp, "Skor", "Score")
                        if _rt_table_score_col in _rt_table_disp.columns:
                            _rt_table_disp[_rt_table_score_col] = pd.to_numeric(_rt_table_disp[_rt_table_score_col], errors="coerce").round(4)
                        render_table(_rt_table_disp)
                with st.container():
                    st.markdown(f'<div class="commentary-box">{_safe_plain_commentary_html(_rank_comment)}</div>', unsafe_allow_html=True)

            with st.expander(f"📊 {tt('Şekiller', 'Figures')}", expanded=True):
                fig_r1, fig_r2 = st.columns(2)
                with fig_r1:
                    st.plotly_chart(fig_rank_bar(_rt_disp), use_container_width=True)
                    with st.container():
                        st.markdown(f'<div class="commentary-box">{tt("Yatay çubuk grafik, her alternatifin sıralama skorunu karşılaştırmalı olarak gösterir. Uzun çubuk yüksek performansı simgeler. Barlar arasındaki mesafe, alternatiflerin birbirinden ne kadar ayrıştığını ortaya koyar.", "The horizontal bar chart comparatively shows the ranking score of each alternative. A longer bar symbolizes higher performance. The distance between bars reveals how much alternatives are differentiated from each other.")}</div>', unsafe_allow_html=True)
                with fig_r2:
                    st.plotly_chart(fig_network_alternatives(_rt_disp), use_container_width=True)
                    with st.container():
                        st.markdown(f'<div class="commentary-box">{tt("Ağ diyagramı, alternatiflerin merkezi referansa göre göreli konumunu dairesel yerleşimde gösterir. Merkeze yakın düğümler daha yüksek skor anlamına gelir. Düğüm boyutu ve rengi performans düzeyini yansıtır.", "The network diagram shows the relative position of alternatives in circular layout against the central reference. Nodes closer to the center mean higher scores. Node size and color reflect performance level.")}</div>', unsafe_allow_html=True)

                contrib = result.get("contribution_table")
                if isinstance(contrib, pd.DataFrame):
                    _contrib_disp = contrib.copy()
                    _contrib_alt_col = col_key(_contrib_disp, "Alternatif", "Alternative")
                    if alt_names:
                        _contrib_disp[_contrib_alt_col] = _contrib_disp[_contrib_alt_col].astype(str).map(
                            lambda x: alt_names.get(x, x)
                        )
                    st.plotly_chart(fig_parallel_coords(_rt_disp, _contrib_disp, criteria), use_container_width=True)
                    with st.container():
                        st.markdown(f'<div class="commentary-box">{tt("Paralel koordinat grafiği, her alternatifin tüm kriterler boyunca performans profilini tek bir görünümde sunar. Çizgilerin rengi ve yoğunluğu skor büyüklüğünü gösterir. Çizgilerin yoğun kesiştiği bölgeler, alternatiflerin benzer kriter değerlerine sahip olduğu alanları işaret eder.", "The parallel coordinates chart presents each alternative\'s performance profile across all criteria in a single view. Line color and intensity show score magnitude. Regions where lines densely intersect indicate areas where alternatives have similar criterion values.")}</div>', unsafe_allow_html=True)

            with st.expander(f"🧠 {tt('Yöntem-Özel İçgörüler', 'Method-Specific Insights')}", expanded=False):
                render_method_specific_insights(result, alt_names)

    with tabs[_tab_comparison]:
        _comp_comment = gen_comparison_commentary(result)
        render_tab_assistant(_comp_comment, key="comp")
        comp = result.get("comparison", {})
        if comp and "rank_table" in comp:
            with st.expander(f"📋 {tt('Tablolar', 'Tables')}", expanded=True):
                c1, c2 = st.columns(2)
                with c1:
                    st.markdown(f"##### 🔁 {tt('Yöntem Sıralama Karşılaştırması', 'Method Ranking Comparison')}")
                    _rank_comp = comp["rank_table"].copy()
                    _rank_comp_alt_col = col_key(_rank_comp, "Alternatif", "Alternative")
                    if alt_names:
                        if _rank_comp_alt_col in _rank_comp.columns:
                            _rank_comp[_rank_comp_alt_col] = _rank_comp[_rank_comp_alt_col].astype(str).map(
                                lambda x: alt_names.get(x, x)
                            )
                    render_table(localize_df(_rank_comp))
                with c2:
                    _method_rob_df = build_ranking_robustness_table(result)
                    if isinstance(_method_rob_df, pd.DataFrame) and not _method_rob_df.empty:
                        st.markdown(f"##### 🛡️ {tt('Yöntem Bazlı Sağlamlık', 'Method-Based Robustness')}")
                        st.caption(tt(
                            "Her yöntemin ana yöntemle ne kadar benzer düşündüğünü basit dille açıklar.",
                            "Explains in plain language how similarly each method thinks compared with the primary method."
                        ))
                        render_table(_method_rob_df)
                with st.container():
                    st.markdown(f'<div class="commentary-box">{_safe_plain_commentary_html(_comp_comment)}</div>', unsafe_allow_html=True)

            with st.expander(f"📊 {tt('Şekiller', 'Figures')}", expanded=True):
                if isinstance(comp.get("spearman_matrix"), pd.DataFrame):
                    spdf = comp["spearman_matrix"]
                    _method_col = col_key(spdf, "Yöntem", "Method")
                    fig_sp = px.imshow(
                        spdf.set_index(_method_col), text_auto=".2f", aspect="auto",
                        color_continuous_scale="RdYlGn", zmin=-1, zmax=1,
                        title=tt("Spearman Uyum Matrisi", "Spearman Agreement Matrix"),
                    )
                    fig_sp.update_layout(height=420, **_THEME)
                    st.plotly_chart(fig_sp, use_container_width=True)
                    with st.container():
                        st.markdown(f'<div class="commentary-box">{tt("Spearman ısı haritası, seçili yöntemlerin birbiriyle ne kadar tutarlı sıralama ürettiğini gösterir. Yeşil hücreler (ρ ≥ 0.85) yüksek yöntem uyumuna, sarı-kırmızı hücreler ise yöntemler arası görüş ayrılığına işaret eder. Üçgen matristeki tüm hücreler yeşilse sonuçlar metodolojiden bağımsız kararlı kabul edilebilir.", "The Spearman heatmap shows how consistently the selected methods produce rankings. Green cells (ρ ≥ 0.85) indicate high method agreement, yellow-red cells indicate disagreement between methods. If all cells in the triangular matrix are green, results can be considered stable regardless of methodology.")}</div>', unsafe_allow_html=True)
        else:
            st.info(tt("Karşılaştırma için birden fazla yöntem seçilmelidir.", "Select more than one method for comparison."))

    with tabs[_tab_durability]:
        render_tab_assistant(gen_mc_commentary(result), key="mc")
        sens = result.get("sensitivity")
        if not sens:
            st.info(tt("Sağlamlık analizi verisi oluşturulamadı. Lütfen sıralama yöntemi seçerek analizi çalıştırın.", "Robustness data could not be generated. Please run analysis with a ranking method."))
        else:
            mc_df  = sens.get("monte_carlo_summary")
            loc_df = sens.get("local_scenarios")
            stab   = float(sens.get("top_stability", 0.0))
            n_iter = int(sens.get("n_iterations", 0)) if "n_iterations" in sens else "—"

            mc1, mc2, mc3 = st.columns(3)
            mc1.metric(tt("Lider Kararlılığı", "Leader Stability"), f"%{stab*100:.1f}", delta="Monte Carlo")
            mc2.metric(tt("Simülasyon Sayısı", "Simulation Count"), f"{n_iter:,}" if isinstance(n_iter, int) else n_iter)
            if isinstance(mc_df, pd.DataFrame) and not mc_df.empty:
                _mc_alt_col = col_key(mc_df, "Alternatif", "Alternative")
                runner_up = mc_df.iloc[1][_mc_alt_col] if len(mc_df) > 1 else "—"
                runner_disp = alt_names.get(str(runner_up), str(runner_up))
                mc3.metric(tt("2. Sıra Alternatif", "Runner-up Alternative"), runner_disp)

            if isinstance(mc_df, pd.DataFrame) and not mc_df.empty:
                _mc_disp = mc_df.copy()
                _mc_disp_alt_col = col_key(_mc_disp, "Alternatif", "Alternative")
                if alt_names:
                    _mc_disp[_mc_disp_alt_col] = _mc_disp[_mc_disp_alt_col].astype(str).map(
                        lambda x: alt_names.get(x, x)
                    )

                _mc_comment = gen_mc_commentary(result)

                with st.expander(f"📋 {tt('Tablolar', 'Tables')}", expanded=True):
                    st.markdown(f"##### 🎲 {tt('Simülasyon Özet Tablosu', 'Simulation Summary Table')}")
                    _mc_table_disp = localize_df(_mc_disp)
                    _fp_col = col_key(_mc_table_disp, "BirincilikOranı", "FirstPlaceRate")
                    _mr_col = col_key(_mc_table_disp, "OrtalamaSıra", "MeanRank")
                    if _fp_col in _mc_table_disp.columns:
                        _mc_table_disp[_fp_col] = (pd.to_numeric(_mc_table_disp[_fp_col], errors="coerce") * 100.0).round(2)
                        _mc_table_disp = _mc_table_disp.rename(columns={_fp_col: f"{_fp_col} (%)"})
                    if _mr_col in _mc_table_disp.columns:
                        _mc_table_disp[_mr_col] = pd.to_numeric(_mc_table_disp[_mr_col], errors="coerce").round(2)
                    render_table(_mc_table_disp)
                    with st.container():
                        st.markdown(f'<div class="commentary-box">{_safe_plain_commentary_html(_mc_comment)}</div>', unsafe_allow_html=True)

                with st.expander(f"📊 {tt('Şekiller', 'Figures')}", expanded=True):
                    mg1, mg2 = st.columns(2)
                    with mg1:
                        st.plotly_chart(fig_mc_rank_bar(_mc_disp), use_container_width=True)
                        with st.container():
                            st.markdown(f'<div class="commentary-box">{tt("Birincilik oranı çubuğu, her alternatifin Monte Carlo simülasyonları boyunca kaç kez birinci sıraya oturduğunu yüzde olarak gösterir. %50 üzeri oran güçlü kararlılığı ifade eder; %20 altı oran alternatifin liderliğinin rastlantısal olduğuna işaret edebilir.", "The first-place rate bar shows what percentage of Monte Carlo simulations each alternative ranked first. A rate above 50% indicates strong stability; a rate below 20% may indicate the alternative\'s leadership is coincidental.")}</div>', unsafe_allow_html=True)
                    with mg2:
                        st.plotly_chart(fig_mc_stability_bubble(_mc_disp), use_container_width=True)
                        with st.container():
                            st.markdown(f'<div class="commentary-box">{tt("Balon grafiği, birincilik oranı ile ortalama sıra arasındaki ilişkiyi gösterir. Sol üst köşeye yakın (düşük ortalama sıra, yüksek birincilik oranı) alternatifler en sağlam adaylardır.", "The bubble chart shows the relationship between first-place rate and mean rank. Alternatives close to the upper-left corner (low mean rank, high first-place rate) are the most robust candidates.")}</div>', unsafe_allow_html=True)

                    if isinstance(loc_df, pd.DataFrame) and not loc_df.empty:
                        st.plotly_chart(fig_local_sensitivity(loc_df), use_container_width=True)
                        with st.container():
                            st.markdown(f'<div class="commentary-box">{tt("Lokal duyarlılık grafiği, her kriterin ağırlığı ±%10 ve ±%20 değiştirildiğinde sıralamanın ne kadar etkilendiğini Spearman korelasyonu ile ölçer. ρ = 1.0\'a yakın çizgiler, o kriterin ağırlık değişimine karşı sıralamanın kararlı kaldığını gösterir; aşağı düşen çizgiler kritik duyarlılık noktalarına işaret eder.", "The local sensitivity chart measures how much the ranking is affected when each criterion weight changes by ±10% and ±20%, using Spearman correlation. Lines close to ρ = 1.0 show the ranking remains stable against that criterion weight change; lines dropping down indicate critical sensitivity points.")}</div>', unsafe_allow_html=True)

with tabs[_tab_output]:
    _out_lang = st.session_state.get("ui_lang", "TR")
    _doc_sections = _build_doc_sections(result, _out_lang)
    _ref_heading = tt("Kaynakça", "References")
    _refs = _doc_sections.get(_ref_heading, [])
    st.text_input(
        tt("Excel çalışma başlığı", "Excel study title"),
        key="study_title",
        placeholder=tt("Örn. OECD Sağlık Dayanıklılığı Analizi", "E.g. OECD Health Resilience Analysis"),
    )
    st.caption(
        tt(
            "Bu başlık yalnızca Excel kapağında ve dosya adında kullanılır. Word çıktısı aynen korunur.",
            "This title is used only on the Excel cover and file name. The Word output remains unchanged.",
        )
    )
    # ── Sonra ne yapmalıyım rehberi ──
    _ar = st.session_state.get("analysis_result")
    if isinstance(_ar, dict):
        _guidance = tt(
            "✅ <b>Sonuçlarınız hazır.</b> (1) Karşılaştırma sekmesinde yöntem uyumunu kontrol edin. (2) Sağlamlık sekmesinde MC kararlılığını inceleyin. (3) Hedefinize uygun formatı indirin: iş raporu→Excel, sunum→APA Rapor, makale/tez→IMRAD.",
            "✅ <b>Results are ready.</b> (1) Check method agreement in Comparison tab. (2) Review MC stability in Robustness tab. (3) Download the right format: business→Excel, presentation→APA Report, paper/thesis→IMRAD."
        )
        st.markdown(f"<div class='assistant-box' style='margin-bottom:1rem;'>{_guidance}</div>", unsafe_allow_html=True)

    _render_report_download_controls(_out_lang)

    # ── Ek çıktılar: Yönetici Özeti + Dashboard ──
    _dl_extra_1, _dl_extra_2 = st.columns(2)
    with _dl_extra_1:
        if _ensure_docx_support():
            try:
                _exec_bytes = mcdm_article.generate_executive_summary_docx(_ar, _ar.get("selected_data", pd.DataFrame()), lang=_out_lang) if isinstance(_ar, dict) else None
            except Exception:
                _exec_bytes = None
            if _exec_bytes:
                st.download_button(
                    tt("📋 Yönetici Özeti (Word)", "📋 Executive Summary (Word)"),
                    _exec_bytes, tt("MCDM_Yonetici_Ozeti.docx", "MCDM_Executive_Summary.docx"),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key=f"dl_exec_{_out_lang}", use_container_width=True)
                st.caption(tt("Tek sayfalık özet. Yönetim kurulu ve sunum için.", "One-page summary for board meetings."))
    with _dl_extra_2:
        try:
            _dash_html = mcdm_article.generate_dashboard_html(_ar, _ar.get("selected_data", pd.DataFrame()), lang=_out_lang) if isinstance(_ar, dict) else None
        except Exception:
            _dash_html = None
        if _dash_html:
            st.download_button(
                tt("🌐 Dashboard (HTML)", "🌐 Dashboard (HTML)"),
                _dash_html.encode("utf-8"), tt("MCDM_Dashboard.html", "MCDM_Dashboard.html"),
                "text/html", key=f"dl_dash_{_out_lang}", use_container_width=True)
            st.caption(tt("Tarayıcıda açılabilen interaktif dashboard.", "Interactive dashboard viewable in any browser."))

    st.markdown(_reference_notice_html(_out_lang), unsafe_allow_html=True)
    st.markdown(f"### 📚 {_ref_heading}")
    if _refs:
        st.markdown("\n".join(f"{idx}. {ref}" for idx, ref in enumerate(_refs, start=1)))
    else:
        st.info(tt("Kaynakça bilgisi üretilemedi.", "Reference information could not be generated."))

st.markdown(
    f"""
    <div style="text-align:right; padding: 16px 20px 6px 20px; font-family:'Georgia',serif;
                color:#1d3557; font-size:15px; font-weight:600; font-style:italic;">
        Prof. Dr. Ömer Faruk Rençber
    </div>
    <div style="text-align:center; padding: 4px 0 24px 0; font-size:0.80rem; color:#718096;">
        <a href="https://www.ofrencber.com" target="_blank"
           style="color:#2E7D9E; text-decoration:none; font-weight:600;">
            🌐 www.ofrencber.com
        </a>
        &nbsp;·&nbsp; MCDM Karar Destek Sistemi Professional &nbsp;·&nbsp; {tt("Akademik Karar Destek Sistemi", "Academic Decision Support System")}
    </div>
    """,
    unsafe_allow_html=True,
)
