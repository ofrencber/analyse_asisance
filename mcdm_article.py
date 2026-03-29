"""
MCDM Karar Destek Sistemi — IMRAD Academic Article Generator
================================================
Generates a journal-style DOCX article in IMRAD format
(Introduction, Method, Results, and Discussion)
based on the analysis results and the methodology template.
"""
from __future__ import annotations

import io
import re
from pathlib import Path
from typing import Any, Dict, List, Optional, Sequence

import numpy as np
import pandas as pd

# ---------------------------------------------------------------------------
# Lazy imports (same pattern as mcdm_app.py)
# ---------------------------------------------------------------------------
Document = None
WD_ALIGN_PARAGRAPH = None
OxmlElement = None
qn = None
Inches = None
Pt = None
RGBColor = None

_DOCX_READY = False


def _ensure_docx() -> bool:
    global _DOCX_READY, Document, WD_ALIGN_PARAGRAPH, OxmlElement, qn, Inches, Pt, RGBColor
    if _DOCX_READY:
        return True
    try:
        from docx import Document as _Doc
        from docx.enum.text import WD_ALIGN_PARAGRAPH as _WD
        from docx.oxml.ns import qn as _qn
        from docx.oxml import OxmlElement as _Ox
        from docx.shared import Inches as _In, Pt as _Pt, RGBColor as _RGB
        Document = _Doc
        WD_ALIGN_PARAGRAPH = _WD
        OxmlElement = _Ox
        qn = _qn
        Inches = _In
        Pt = _Pt
        RGBColor = _RGB
        _DOCX_READY = True
    except Exception:
        _DOCX_READY = False
    return _DOCX_READY


# ---------------------------------------------------------------------------
# matplotlib availability
# ---------------------------------------------------------------------------
_MPL_READY: bool | None = None
_plt = None


def _ensure_matplotlib() -> bool:
    global _MPL_READY, _plt
    if _MPL_READY is not None:
        return _MPL_READY
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as mplt
        _plt = mplt
        _MPL_READY = True
    except Exception:
        _MPL_READY = False
    return _MPL_READY


# ---------------------------------------------------------------------------
# Flowchart generation
# ---------------------------------------------------------------------------

def _generate_flowchart_bytes(d: Dict[str, Any], lang: str) -> bytes | None:
    """Generate a detailed methodology flowchart as PNG bytes.

    Shows every computational step explicitly:
    - Where normalization happens
    - Where fuzzification enters the pipeline
    - Where crisp weights feed into fuzzy ranking
    - Robustness sub-components
    """
    if not _ensure_matplotlib():
        return None

    try:
        from matplotlib.patches import FancyBboxPatch

        wm = d["weight_method"] or "?"
        rm = d["ranking_method"] or ""
        is_fuzzy = d.get("is_fuzzy", False)
        spread = d.get("fuzzy_spread", 0.10)
        is_fuzzy_topsis = rm == "Fuzzy TOPSIS"

        # Determine steps dynamically
        # Each step: (color, title_line, detail_lines)
        if lang == "EN":
            steps = [
                ("#1565C0", "STAGE 1 — Decision Matrix",
                 [f"X(m×n) : {d['n_alt']} alternatives × {d['n_crit']} criteria",
                  f"Benefit: {', '.join(d['benefit'][:4])}" + ("..." if len(d['benefit']) > 4 else ""),
                  f"Cost: {', '.join(d['cost'][:4])}" + ("..." if len(d['cost']) > 4 else "") if d['cost'] else None]),

                ("#2E7D32", "STAGE 2 — Normalization & Criteria Weighting",
                 [f"Method: {wm}",
                  "Crisp weight vector w = (w₁, ..., wₙ) is computed",
                  "from the raw decision matrix (not from fuzzy values)"]),
            ]
            if rm:
                if is_fuzzy:
                    steps.append(
                        ("#E65100", "STAGE 3a — Fuzzification (TFN Conversion)",
                         [f"x̃ᵢⱼ = ( xᵢⱼ(1−s), xᵢⱼ, xᵢⱼ(1+s) )  where s = {_fmt(spread, lang)}",
                          "Applied AFTER weight calculation, BEFORE ranking",
                          "Weights remain crisp; uncertainty enters at scoring stage"]))
                    if is_fuzzy_topsis:
                        steps.append(
                            ("#BF360C", "STAGE 3b — Fuzzy Ranking: Fuzzy TOPSIS",
                             ["TFN normalization → Weighted fuzzy matrix",
                              "FPIS / FNIS → Vertex distances → CCᵢ ∈ [0,1]",
                              "Operates directly on TFN triplets (no scenario wrapper)"]))
                    else:
                        base_name = rm.replace("Fuzzy ", "")
                        steps.append(
                            ("#BF360C", f"STAGE 3b — Fuzzy Ranking: {rm}",
                             [f"Three scenarios: X^L = X(1−s), X^M = X, X^U = X(1+s)",
                              f"Classical {base_name} is executed on each scenario",
                              "Final score = mean of three scenario scores"]))
                else:
                    steps.append(
                        ("#E65100", f"STAGE 3 — Classical Ranking: {rm}",
                         ["Normalization → Weighted matrix → Method-specific aggregation",
                          "Crisp weights applied directly to crisp decision matrix"]))

                steps.append(
                    ("#6A1B9A", "STAGE 4 — Robustness & Sensitivity Analysis",
                     [f"Local sensitivity: ±10%, ±20% perturbation on each weight ({d['n_crit']}×4 = {d['n_crit']*4} scenarios)",
                      f"Monte Carlo: N = {d['mc_n']}, σ = {_fmt(d['mc_sigma'], lang)} log-normal noise" if d['mc_n'] > 0 else None,
                      f"Cross-method comparison: {', '.join(d['comp_methods'][:4])}" if d['comp_methods'] else None,
                      "Spearman ρ for pairwise rank agreement"]))
                steps.append(
                    ("#C62828", "STAGE 5 — Decision Confidence Assessment",
                     [f"Leader: {d['top_alt']}  |  Score: {_fmt(d['top_score'], lang)}" if d['top_alt'] else None,
                      f"MC stability: {_fmt(d['stability']*100, lang, 1)}% first-place rate" if d['stability'] > 0 else None,
                      f"Mean Spearman ρ̄ = {_fmt(d['mean_spearman'], lang)} → {'High' if d['mean_spearman'] >= 0.85 else ('Moderate' if d['mean_spearman'] >= 0.70 else 'Low')} consistency" if d['mean_spearman'] > 0 else None]))
            else:
                steps.append(
                    ("#C62828", "RESULT — Criteria Weight Vector",
                     [f"Top criterion: {d['top_criterion']} (w = {_fmt(d['top_weight'], lang)})" if d['top_criterion'] else None]))
        else:
            # TURKISH
            steps = [
                ("#1565C0", "AŞAMA 1 — Karar Matrisi",
                 [f"X(m×n) : {d['n_alt']} alternatif × {d['n_crit']} kriter",
                  f"Fayda: {', '.join(d['benefit'][:4])}" + ("..." if len(d['benefit']) > 4 else ""),
                  f"Maliyet: {', '.join(d['cost'][:4])}" + ("..." if len(d['cost']) > 4 else "") if d['cost'] else None]),

                ("#2E7D32", "AŞAMA 2 — Normalizasyon ve Kriter Ağırlıklandırma",
                 [f"Yöntem: {wm}",
                  "Crisp ağırlık vektörü w = (w₁, ..., wₙ) ham karar",
                  "matrisinden hesaplanır (bulanık değerlerden değil)"]),
            ]
            if rm:
                if is_fuzzy:
                    steps.append(
                        ("#E65100", "AŞAMA 3a — Bulanıklaştırma (TFN Dönüşümü)",
                         [f"x̃ᵢⱼ = ( xᵢⱼ(1−s),  xᵢⱼ,  xᵢⱼ(1+s) )   s = {_fmt(spread, lang)}",
                          "Ağırlık hesabından SONRA, sıralamadan ÖNCE uygulanır",
                          "Ağırlıklar crisp kalır; belirsizlik skorlama aşamasında girer"]))
                    if is_fuzzy_topsis:
                        steps.append(
                            ("#BF360C", "AŞAMA 3b — Bulanık Sıralama: Fuzzy TOPSIS",
                             ["TFN normalizasyonu → Ağırlıklı bulanık matris",
                              "FPIS / FNIS → Vertex uzaklıkları → CCᵢ ∈ [0,1]",
                              "Doğrudan TFN üçlüleri üzerinde çalışır (senaryo sarmalayıcısı yok)"]))
                    else:
                        base_name = rm.replace("Fuzzy ", "")
                        steps.append(
                            ("#BF360C", f"AŞAMA 3b — Bulanık Sıralama: {rm}",
                             [f"Üç senaryo: X^L = X(1−s),  X^M = X,  X^U = X(1+s)",
                              f"Klasik {base_name} her senaryoda ayrı çalıştırılır",
                              "Nihai skor = üç senaryo skorunun ortalaması"]))
                else:
                    steps.append(
                        ("#E65100", f"AŞAMA 3 — Klasik Sıralama: {rm}",
                         ["Normalizasyon → Ağırlıklı matris → Yönteme özgü toplulaştırma",
                          "Crisp ağırlıklar doğrudan crisp karar matrisine uygulanır"]))

                steps.append(
                    ("#6A1B9A", "AŞAMA 4 — Sağlamlık ve Duyarlılık Analizi",
                     [f"Lokal duyarlılık: her ağırlığa ±%10, ±%20 pertürbasyon ({d['n_crit']}×4 = {d['n_crit']*4} senaryo)",
                      f"Monte Carlo: N = {d['mc_n']}, σ = {_fmt(d['mc_sigma'], lang)} log-normal gürültü" if d['mc_n'] > 0 else None,
                      f"Çapraz yöntem karşılaştırması: {', '.join(d['comp_methods'][:4])}" if d['comp_methods'] else None,
                      "İkili sıra uyumu için Spearman ρ hesaplanır"]))
                steps.append(
                    ("#C62828", "AŞAMA 5 — Karar Güven Değerlendirmesi",
                     [f"Lider: {d['top_alt']}  |  Skor: {_fmt(d['top_score'], lang)}" if d['top_alt'] else None,
                      f"MC kararlılık: %{_fmt(d['stability']*100, lang, 1)} birincilik oranı" if d['stability'] > 0 else None,
                      f"Ortalama Spearman ρ̄ = {_fmt(d['mean_spearman'], lang)} → {'Yüksek' if d['mean_spearman'] >= 0.85 else ('Orta' if d['mean_spearman'] >= 0.70 else 'Düşük')} tutarlılık" if d['mean_spearman'] > 0 else None]))
            else:
                steps.append(
                    ("#C62828", "SONUÇ — Kriter Ağırlık Vektörü",
                     [f"En önemli: {d['top_criterion']} (w = {_fmt(d['top_weight'], lang)})" if d['top_criterion'] else None]))

        # Filter None lines
        for step in steps:
            step_lines = step[2]
            step_lines[:] = [l for l in step_lines if l]

        n_steps = len(steps)
        box_h = 0.85
        gap = 0.55
        total_h = n_steps * box_h + (n_steps - 1) * gap
        fig_h = max(8, total_h * 0.85 + 2.5)
        fig, ax = _plt.subplots(figsize=(8.5, fig_h))
        y_max = total_h + 2.0
        ax.set_xlim(0, 12)
        ax.set_ylim(0, y_max)
        ax.axis("off")

        cx = 6.0
        box_w = 10.5

        for i, (color, title, details) in enumerate(steps):
            cy = y_max - 1.0 - i * (box_h + gap)
            x0 = cx - box_w / 2
            y0 = cy - box_h / 2

            box = FancyBboxPatch(
                (x0, y0), box_w, box_h,
                boxstyle="round,pad=0.12",
                facecolor=color, edgecolor="white", linewidth=1.5,
            )
            ax.add_patch(box)

            # Step number on left
            ax.text(x0 + 0.35, cy, str(i + 1), ha="center", va="center",
                    fontsize=16, fontweight="bold", color="#FFFFFF80",
                    fontfamily="sans-serif")

            # Title
            all_lines = [title] + details
            n_lines = len(all_lines)
            line_spacing = min(0.18, box_h / (n_lines + 0.5))
            y_start = cy + (n_lines - 1) * line_spacing / 2

            for j, line in enumerate(all_lines):
                y_text = y_start - j * line_spacing
                fs = 9.0 if j == 0 else 7.5
                fw = "bold" if j == 0 else "normal"
                ax.text(cx + 0.15, y_text, line, ha="center", va="center",
                        fontsize=fs, fontweight=fw, color="white",
                        fontfamily="sans-serif")

            # Arrow
            if i < n_steps - 1:
                arr_y1 = cy - box_h / 2
                arr_y2 = y_max - 1.0 - (i + 1) * (box_h + gap) + box_h / 2
                ax.annotate(
                    "", xy=(cx, arr_y2 + 0.03), xytext=(cx, arr_y1 - 0.03),
                    arrowprops=dict(arrowstyle="-|>", color="#37474F", lw=2.5,
                                   mutation_scale=18))

        # Caption
        if lang == "EN":
            cap = "Figure 1. Proposed Methodology Flowchart"
        else:
            cap = "Şekil 1. Önerilen Yöntemin Akış Şeması"
        ax.text(cx, 0.35, cap, ha="center", va="center",
                fontsize=10, fontweight="bold", fontstyle="italic",
                color="#333333", fontfamily="sans-serif")

        _plt.tight_layout()
        buf = io.BytesIO()
        fig.savefig(buf, format="png", dpi=200, bbox_inches="tight",
                    facecolor="white", edgecolor="none")
        _plt.close(fig)
        return buf.getvalue()

    except Exception:
        return None


# ---------------------------------------------------------------------------
# latex2mathml availability
# ---------------------------------------------------------------------------
def _generate_weight_bar_bytes(weight_table, lang: str) -> bytes | None:
    if not _ensure_matplotlib() or weight_table is None or weight_table.empty:
        return None
    try:
        wt = weight_table.copy()
        w_col = "Ağırlık" if "Ağırlık" in wt.columns else ("Weight" if "Weight" in wt.columns else None)
        c_col = "Kriter" if "Kriter" in wt.columns else ("Criterion" if "Criterion" in wt.columns else None)
        if not w_col or not c_col:
            return None
        wt[w_col] = pd.to_numeric(wt[w_col], errors="coerce")
        wt = wt.dropna(subset=[w_col]).sort_values(w_col, ascending=True)
        fig, ax = _plt.subplots(figsize=(6, max(2.5, len(wt) * 0.4)))
        ax.barh(wt[c_col].astype(str), wt[w_col].astype(float), color="#2B6CA0")
        ax.set_xlabel("Ağırlık" if lang != "EN" else "Weight", fontsize=9)
        ax.spines["top"].set_visible(False); ax.spines["right"].set_visible(False)
        _plt.tight_layout()
        buf = io.BytesIO()
        fig.savefig(buf, format="png", dpi=200, bbox_inches="tight", facecolor="white")
        _plt.close(fig)
        return buf.getvalue()
    except Exception:
        return None


def _generate_ranking_bar_bytes(ranking_table, lang: str) -> bytes | None:
    if not _ensure_matplotlib() or ranking_table is None or ranking_table.empty:
        return None
    try:
        rt = ranking_table.copy()
        a_col = next((c for c in rt.columns if c in ("Alternatif", "Alternative")), None)
        s_col = next((c for c in rt.columns if c in ("Skor", "Score")), None)
        r_col = next((c for c in rt.columns if c in ("Sıra", "Rank")), None)
        if not a_col or not s_col:
            return None
        rt[s_col] = pd.to_numeric(rt[s_col], errors="coerce")
        if r_col:
            rt = rt.sort_values(r_col)
        fig, ax = _plt.subplots(figsize=(max(5, len(rt) * 0.7), 4))
        colors = ["#1B4D8F" if i == 0 else "#5B9BD5" for i in range(len(rt))]
        ax.bar(rt[a_col].astype(str), rt[s_col].astype(float), color=colors)
        ax.set_ylabel("Skor" if lang != "EN" else "Score", fontsize=9)
        ax.spines["top"].set_visible(False); ax.spines["right"].set_visible(False)
        _plt.tight_layout()
        buf = io.BytesIO()
        fig.savefig(buf, format="png", dpi=200, bbox_inches="tight", facecolor="white")
        _plt.close(fig)
        return buf.getvalue()
    except Exception:
        return None


def _generate_method_specific_chart(d: Dict[str, Any], lang: str) -> bytes | None:
    if not _ensure_matplotlib():
        return None
    details = d.get("ranking_details") or {}
    rm = d.get("ranking_method") or ""
    base = rm.replace("Fuzzy ", "")
    try:
        if base == "TOPSIS":
            df = details.get("distance_table")
            if not isinstance(df, pd.DataFrame) or df.empty:
                return None
            fig, ax = _plt.subplots(figsize=(6, 5))
            sc = ax.scatter(df["D+"], df["D-"], c=df["Skor"], cmap="Blues", s=80, edgecolors="#1B365D")
            for _, row in df.iterrows():
                ax.annotate(str(row["Alternatif"]), (row["D+"], row["D-"]), fontsize=7)
            ax.set_xlabel("D+"); ax.set_ylabel("D-"); ax.set_title("TOPSIS Distance Map")
            ax.invert_xaxis(); _plt.colorbar(sc, ax=ax, label="CC")
            buf = io.BytesIO()
            fig.savefig(buf, format="png", dpi=200, bbox_inches="tight", facecolor="white")
            _plt.close(fig)
            return buf.getvalue()
        elif base == "VIKOR":
            df = details.get("vikor_table")
            if not isinstance(df, pd.DataFrame) or df.empty:
                return None
            df = df.sort_values("Q")
            x = np.arange(len(df)); w = 0.25
            fig, ax = _plt.subplots(figsize=(max(5, len(df)*0.8), 4))
            ax.bar(x-w, df["S"], w, label="S", color="#7AA6D1")
            ax.bar(x, df["R"], w, label="R", color="#315D8A")
            ax.bar(x+w, df["Q"], w, label="Q", color="#17324D")
            ax.set_xticks(x); ax.set_xticklabels(df["Alternatif"].astype(str), fontsize=8)
            ax.legend(); ax.set_title("VIKOR S-R-Q")
            buf = io.BytesIO()
            fig.savefig(buf, format="png", dpi=200, bbox_inches="tight", facecolor="white")
            _plt.close(fig)
            return buf.getvalue()
    except Exception:
        pass
    return None


_L2M_READY: bool | None = None
_MML2OMML_XSL: str | None = None


def _ensure_latex2mathml() -> bool:
    """Return True if latex2mathml is available."""
    global _L2M_READY
    if _L2M_READY is not None:
        return _L2M_READY
    try:
        import latex2mathml.converter  # noqa: F401
        _L2M_READY = True
    except Exception:
        _L2M_READY = False
    return _L2M_READY


# ---------------------------------------------------------------------------
# Template loading
# ---------------------------------------------------------------------------
_TEMPLATE_PATH = Path(__file__).resolve().parent / "mcdm_sistem_metodoloji_sablonu.md"
_TEMPLATE_CACHE: Dict[str, Any] | None = None


def _load_template() -> Dict[str, Any]:
    """Parse the methodology template MD file into structured sections."""
    global _TEMPLATE_CACHE
    if _TEMPLATE_CACHE is not None:
        return _TEMPLATE_CACHE

    text = _TEMPLATE_PATH.read_text(encoding="utf-8")

    # Split into major sections
    bolum_splits = re.split(r"^## BÖLÜM \d+ — ", text, flags=re.MULTILINE)
    bolum_headers = re.findall(r"^## (BÖLÜM \d+ — .+)$", text, flags=re.MULTILINE)

    bolumler: Dict[str, str] = {}
    for i, header in enumerate(bolum_headers):
        if i + 1 < len(bolum_splits):
            bolumler[header] = bolum_splits[i + 1]

    # Parse weight methods (BÖLÜM 1)
    weight_methods: Dict[str, str] = {}
    ranking_methods: Dict[str, str] = {}
    robustness_text: str = ""
    imrad_template: str = ""

    for key, content in bolumler.items():
        if "AĞIRLIK" in key:
            weight_methods = _parse_method_blocks(content, r"####\s+\d+\.\w+\.\d+\s+")
        elif "SIRALAMA" in key:
            ranking_methods = _parse_method_blocks(content, r"####\s+\d+\.\w+\.\d+\s+")
            # Also capture "Ortak Altyapı" block
            ortak_match = re.search(
                r"####\s+Ortak Altyapı\s*\n(.*?)(?=####|\Z)",
                content, re.DOTALL
            )
            if ortak_match:
                ranking_methods["_fuzzy_common"] = ortak_match.group(1).strip()
            # Capture the summary table for other fuzzy methods
            tablo_match = re.search(
                r"####\s+2\.B\.2.+?\n(.*?)(?=---\s*$|\Z)",
                content, re.DOTALL | re.MULTILINE,
            )
            if tablo_match:
                ranking_methods["_fuzzy_others_table"] = tablo_match.group(0).strip()
        elif "SAĞLAMLIK" in key or "DAYANIKLILIK" in key:
            robustness_text = content.strip()
        elif "IMRAD" in key:
            imrad_template = content.strip()

    _TEMPLATE_CACHE = {
        "weight_methods": weight_methods,
        "ranking_methods": ranking_methods,
        "robustness": robustness_text,
        "imrad_template": imrad_template,
    }
    return _TEMPLATE_CACHE


def _parse_method_blocks(text: str, heading_pattern: str) -> Dict[str, str]:
    """Split text by #### headings and map method name → content block."""
    parts = re.split(r"(####\s+.+)", text)
    methods: Dict[str, str] = {}
    current_header = None
    for part in parts:
        if re.match(r"####\s+", part):
            current_header = part.strip()
        elif current_header:
            # Extract method name from header like "#### 1.A.1 Entropy (Shannon Bilgi Entropisi)"
            m = re.match(r"####\s+[\d.]+[A-Z]*\.?\d*\s+(.+)", current_header)
            if m:
                method_raw = m.group(1).strip()
                # Get the primary method name (before parenthesis)
                method_name = re.split(r"\s*\(", method_raw)[0].strip()
                methods[method_name] = part.strip()
            current_header = None
    return methods


# ---------------------------------------------------------------------------
# Method name mapping — template names → engine names
# ---------------------------------------------------------------------------
_TEMPLATE_TO_ENGINE: Dict[str, str] = {
    "Entropy": "Entropy",
    "CRITIC": "CRITIC",
    "Standart Sapma": "Standart Sapma",
    "MEREC": "MEREC",
    "LOPCOW": "LOPCOW",
    "PCA": "PCA",
    "CILOS": "CILOS",
    "IDOCRIW": "IDOCRIW",
    "Fuzzy IDOCRIW": "Fuzzy IDOCRIW",
    "AHP": "AHP",
    "BWM": "BWM",
    "SWARA": "SWARA",
    "DEMATEL": "DEMATEL",
    "SMART": "SMART",
    "TOPSIS": "TOPSIS",
    "VIKOR": "VIKOR",
    "EDAS": "EDAS",
    "CODAS": "CODAS",
    "COPRAS": "COPRAS",
    "OCRA": "OCRA",
    "ARAS": "ARAS",
    "SAW": "SAW",
    "WPM": "WPM",
    "MAUT": "MAUT",
    "WASPAS": "WASPAS",
    "MOORA": "MOORA",
    "MULTIMOORA": "MULTIMOORA",
    "MABAC": "MABAC",
    "MARCOS": "MARCOS",
    "CoCoSo": "CoCoSo",
    "PROMETHEE II": "PROMETHEE",
    "PROMETHEE": "PROMETHEE",
    "GRA": "GRA",
    "SPOTIS": "SPOTIS",
    "RAWEC": "RAWEC",
    "RAFSI": "RAFSI",
    "ROV": "ROV",
    "AROMAN": "AROMAN",
    "DNMA": "DNMA",
    "Fuzzy TOPSIS": "Fuzzy TOPSIS",
}


def _find_template_content(method: str, method_dict: Dict[str, str]) -> str:
    """Find the best matching template content for a given engine method name."""
    if method in method_dict:
        return method_dict[method]
    # Try without "Fuzzy " prefix — fuzzy methods share the classical base description
    base = method.replace("Fuzzy ", "")
    if base in method_dict:
        return method_dict[base]
    # Fuzzy search
    for key, content in method_dict.items():
        if method.lower() in key.lower() or key.lower() in method.lower():
            return content
    return ""


# ---------------------------------------------------------------------------
# Localization helpers
# ---------------------------------------------------------------------------
_IMRAD_HEADINGS_TR = {
    "title_prefix": "Çok Kriterli Karar Analizi:",
    "abstract": "ÖZET",
    "keywords_label": "Anahtar Kelimeler",
    "introduction": "1. GİRİŞ",
    "method": "2. YÖNTEM",
    "method_dm": "2.1 Karar Matrisi",
    "method_weight": "2.2 Kriter Ağırlıklarının Belirlenmesi",
    "method_rank": "2.3 Alternatif Sıralaması",
    "method_robust": "2.4 Sağlamlık Analizi",
    "results": "3. BULGULAR",
    "results_weight": "3.1 Kriter Ağırlıkları",
    "results_rank": "3.2 Sıralama Sonuçları",
    "results_robust": "3.3 Sağlamlık Bulguları",
    "discussion": "4. TARTIŞMA",
    "conclusion": "5. SONUÇ",
    "references": "KAYNAKLAR",
    "table": "Tablo",
    "figure": "Şekil",
    "benefit": "fayda",
    "cost": "maliyet",
    "criterion": "Kriter",
    "weight": "Ağırlık",
    "alternative": "Alternatif",
    "score": "Skor",
    "rank": "Sıra",
}

_IMRAD_HEADINGS_EN = {
    "title_prefix": "Multi-Criteria Decision Analysis:",
    "abstract": "ABSTRACT",
    "keywords_label": "Keywords",
    "introduction": "1. INTRODUCTION",
    "method": "2. METHOD",
    "method_dm": "2.1 Decision Matrix",
    "method_weight": "2.2 Determination of Criteria Weights",
    "method_rank": "2.3 Alternative Ranking",
    "method_robust": "2.4 Robustness Analysis",
    "results": "3. RESULTS",
    "results_weight": "3.1 Criteria Weights",
    "results_rank": "3.2 Ranking Results",
    "results_robust": "3.3 Robustness Findings",
    "discussion": "4. DISCUSSION",
    "conclusion": "5. CONCLUSION",
    "references": "REFERENCES",
    "table": "Table",
    "figure": "Figure",
    "benefit": "benefit",
    "cost": "cost",
    "criterion": "Criterion",
    "weight": "Weight",
    "alternative": "Alternative",
    "score": "Score",
    "rank": "Rank",
}


def _h(lang: str) -> Dict[str, str]:
    return _IMRAD_HEADINGS_EN if lang == "EN" else _IMRAD_HEADINGS_TR


def _fmt(value: float, lang: str, decimals: int = 3) -> str:
    s = f"{value:.{decimals}f}"
    if lang != "EN":
        s = s.replace(".", ",")
    return s


# ---------------------------------------------------------------------------
# Data extraction helpers
# ---------------------------------------------------------------------------

def _extract_analysis_data(result: dict, selected_data: pd.DataFrame) -> Dict[str, Any]:
    """Extract all needed values from the analysis result dict."""
    weights_info = result.get("weights") or {}
    ranking_info = result.get("ranking") or {}
    comparison = result.get("comparison") or {}
    sensitivity = result.get("sensitivity") or {}
    validation = result.get("validation") or {}

    weight_method = weights_info.get("method", "")
    ranking_method = ranking_info.get("method")
    weight_table = weights_info.get("weight_table") or weights_info.get("table", pd.DataFrame())
    ranking_table = ranking_info.get("ranking_table") or ranking_info.get("table", pd.DataFrame())
    ranking_details = ranking_info.get("ranking_details") or ranking_info.get("details") or {}

    criteria = list(validation.get("criteria", []))
    criteria_types = validation.get("criteria_types", {})
    # Fallback: if validation doesn't have criteria, derive from weight table or selected_data
    if not criteria:
        if isinstance(weight_table, pd.DataFrame) and not weight_table.empty:
            c_col = "Kriter" if "Kriter" in weight_table.columns else ("Criterion" if "Criterion" in weight_table.columns else None)
            if c_col:
                criteria = list(weight_table[c_col])
        elif selected_data is not None and not selected_data.empty:
            criteria = list(selected_data.select_dtypes(include=["number"]).columns)
    n_alt = len(selected_data) if selected_data is not None else 0
    n_crit = len(criteria)

    benefit = [c for c in criteria if criteria_types.get(c, "max") == "max"]
    cost = [c for c in criteria if criteria_types.get(c, "max") == "min"]

    # Top criterion
    top_criterion = ""
    top_weight = 0.0
    if isinstance(weight_table, pd.DataFrame) and not weight_table.empty:
        wt = weight_table.copy()
        w_col = "Ağırlık" if "Ağırlık" in wt.columns else ("Weight" if "Weight" in wt.columns else None)
        c_col = "Kriter" if "Kriter" in wt.columns else ("Criterion" if "Criterion" in wt.columns else None)
        if w_col and c_col:
            idx = wt[w_col].idxmax()
            top_criterion = str(wt.loc[idx, c_col])
            top_weight = float(wt.loc[idx, w_col])

    # Top alternative
    top_alt = ""
    top_score = 0.0
    second_score = 0.0
    if isinstance(ranking_table, pd.DataFrame) and not ranking_table.empty:
        rt = ranking_table.copy()
        a_col = next((c for c in rt.columns if c in ("Alternatif", "Alternative")), None)
        s_col = next((c for c in rt.columns if c in ("Skor", "Score")), None)
        r_col = next((c for c in rt.columns if c in ("Sıra", "Rank")), None)
        if r_col and a_col:
            sorted_rt = rt.sort_values(r_col)
            if len(sorted_rt) >= 1:
                top_alt = str(sorted_rt.iloc[0][a_col])
                if s_col:
                    top_score = float(sorted_rt.iloc[0][s_col])
            if len(sorted_rt) >= 2 and s_col:
                second_score = float(sorted_rt.iloc[1][s_col])

    # Spearman
    mean_spearman = 0.0
    comp_methods: List[str] = []
    if comparison:
        corr_matrix = comparison.get("corr_matrix") or comparison.get("spearman_matrix")
        if isinstance(corr_matrix, pd.DataFrame) and not corr_matrix.empty:
            vals = corr_matrix.values
            mask = np.triu(np.ones_like(vals, dtype=bool), k=1)
            upper = vals[mask]
            if len(upper) > 0:
                mean_spearman = float(np.nanmean(upper))
        rank_table_comp = comparison.get("rank_table")
        if isinstance(rank_table_comp, pd.DataFrame):
            comp_methods = [c for c in rank_table_comp.columns if c not in ("Alternatif", "Alternative")]

    # Sensitivity / Monte Carlo
    mc_n = 0
    stability = 0.0
    mc_sigma = 0.12
    if sensitivity:
        mc_n = int(sensitivity.get("iterations") or sensitivity.get("n_iterations", 0))
        mc_sigma = float(sensitivity.get("sigma", 0.12))
        top_stab = sensitivity.get("top_stability")
        if isinstance(top_stab, (int, float)):
            stability = float(top_stab)
        elif isinstance(top_stab, dict):
            stability = float(top_stab.get("rate", 0))

    fuzzy_spread = float(result.get("fuzzy_spread", 0.10))
    is_fuzzy = bool(ranking_method and "Fuzzy" in str(ranking_method))

    return {
        "weight_method": weight_method,
        "ranking_method": ranking_method,
        "weight_table": weight_table,
        "ranking_table": ranking_table,
        "ranking_details": ranking_details,
        "comparison": comparison,
        "sensitivity": sensitivity,
        "criteria": criteria,
        "criteria_types": criteria_types,
        "n_alt": n_alt,
        "n_crit": n_crit,
        "benefit": benefit,
        "cost": cost,
        "top_criterion": top_criterion,
        "top_weight": top_weight,
        "top_alt": top_alt,
        "top_score": top_score,
        "second_score": second_score,
        "mean_spearman": mean_spearman,
        "comp_methods": comp_methods,
        "mc_n": mc_n,
        "mc_sigma": mc_sigma,
        "stability": stability,
        "fuzzy_spread": fuzzy_spread,
        "is_fuzzy": is_fuzzy,
        "eliminated": result.get("eliminated_alternatives") or [],
    }


# ---------------------------------------------------------------------------
# DOCX building helpers
# ---------------------------------------------------------------------------

def _configure_doc(doc) -> None:
    """Apply APA/SSCI formatting: Times New Roman 12pt, justified, black text."""
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.font.color.rgb = RGBColor(0, 0, 0)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    style.paragraph_format.line_spacing = 1.15
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)


def _set_run_style(run, lang_code: str = "tr-TR", font_size: float = 12,
                   *, bold: bool | None = None, italic: bool | None = None) -> None:
    rpr = run._element.get_or_add_rPr()
    rfonts = OxmlElement("w:rFonts")
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), "Times New Roman")
    rpr.append(rfonts)
    run.font.name = "Times New Roman"
    run.font.size = Pt(font_size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    # language
    lang_el = OxmlElement("w:lang")
    lang_el.set(qn("w:val"), lang_code)
    rpr.append(lang_el)


def _add_heading(doc, text: str, level: int = 1, lang_code: str = "tr-TR") -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        _set_run_style(run, lang_code=lang_code, font_size=14 if level == 1 else 12, bold=True)
        run.font.color.rgb = RGBColor(0, 0, 0)


def _add_paragraph(doc, text: str, lang_code: str = "tr-TR", font_size: float = 12,
                   *, bold: bool = False, italic: bool = False,
                   alignment=None, space_after: float = 6) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_run_style(run, lang_code=lang_code, font_size=font_size, bold=bold, italic=italic)
    run.font.color.rgb = RGBColor(0, 0, 0)
    p.paragraph_format.space_after = Pt(space_after)
    p.alignment = alignment if alignment is not None else WD_ALIGN_PARAGRAPH.JUSTIFY


def _add_placeholder(doc, text: str, lang_code: str = "tr-TR") -> None:
    """Add a yellow-highlighted placeholder text that the user should fill in."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_run_style(run, lang_code=lang_code, font_size=12, bold=True)
    # Yellow highlight
    rpr = run._element.get_or_add_rPr()
    highlight = OxmlElement("w:highlight")
    highlight.set(qn("w:val"), "yellow")
    rpr.append(highlight)
    p.paragraph_format.space_after = Pt(6)


def _add_mixed_paragraph(doc, parts: List[Dict[str, Any]], lang_code: str = "tr-TR") -> None:
    """Add a paragraph with mixed normal text and yellow placeholder runs.
    Each part: {"text": "...", "placeholder": True/False}
    """
    p = doc.add_paragraph()
    for part in parts:
        run = p.add_run(part["text"])
        is_ph = part.get("placeholder", False)
        _set_run_style(run, lang_code=lang_code, font_size=12, bold=is_ph)
        if is_ph:
            rpr = run._element.get_or_add_rPr()
            highlight = OxmlElement("w:highlight")
            highlight.set(qn("w:val"), "yellow")
            rpr.append(highlight)
    p.paragraph_format.space_after = Pt(6)


def _add_table(doc, df: pd.DataFrame, lang_code: str = "tr-TR", max_rows: int = 30) -> None:
    """Add a DataFrame as a DOCX table."""
    if df is None or df.empty:
        return
    use_df = df.copy().head(max_rows)
    table = doc.add_table(rows=1, cols=len(use_df.columns))
    table.style = "Table Grid"
    is_tr = lang_code.startswith("tr")
    for i, col in enumerate(use_df.columns):
        table.rows[0].cells[i].text = str(col)
        for paragraph in table.rows[0].cells[i].paragraphs:
            for run in paragraph.runs:
                _set_run_style(run, lang_code=lang_code, bold=True)
    for _, row in use_df.iterrows():
        cells = table.add_row().cells
        for j, value in enumerate(row):
            if isinstance(value, (float, np.floating)) and pd.notna(value):
                s = f"{value:.3f}"
                cells[j].text = s.replace(".", ",") if is_tr else s
            elif isinstance(value, float) and pd.isna(value):
                cells[j].text = ""
            else:
                cells[j].text = str(value)
            for paragraph in cells[j].paragraphs:
                for run in paragraph.runs:
                    _set_run_style(run, lang_code=lang_code)
    doc.add_paragraph("")


def _add_figure(doc, image_bytes: bytes, caption: str, lang_code: str = "tr-TR",
                width_inches: float = 5.5) -> None:
    """Insert an image with a centered caption below it."""
    if not image_bytes:
        return
    buf = io.BytesIO(image_bytes)
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.add_run().add_picture(buf, width=Inches(width_inches))
    p_img.paragraph_format.space_after = Pt(2)
    # Caption
    p_cap = doc.add_paragraph()
    run = p_cap.add_run(caption)
    _set_run_style(run, lang_code=lang_code, font_size=10, bold=True, italic=True)
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_after = Pt(10)


# ---------------------------------------------------------------------------
# Math formula rendering
# ---------------------------------------------------------------------------

def _add_math_block(doc, latex_lines: List[str], lang_code: str = "tr-TR") -> None:
    """Render LaTeX formula lines into the DOCX.

    Attempts OMML conversion via latex2mathml; falls back to plain text.
    """
    if _ensure_latex2mathml():
        for line in latex_lines:
            _add_math_omml(doc, line, lang_code)
    else:
        # Fallback: formulas as styled italic text with math font
        for line in latex_lines:
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = "Cambria Math"
            run.font.size = Pt(11)
            run.italic = True
            run.font.color.rgb = RGBColor(0, 0, 0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(4)


def _add_math_omml(doc, latex: str, lang_code: str = "tr-TR") -> None:
    """Convert a single LaTeX expression to OMML and insert into doc."""
    try:
        import latex2mathml.converter
        from lxml import etree

        mathml = latex2mathml.converter.convert(latex)

        xslt_str = _get_mml2omml_xslt()
        if xslt_str:
            xslt_tree = etree.fromstring(xslt_str.encode("utf-8"))
            transform = etree.XSLT(xslt_tree)
            mathml_tree = etree.fromstring(mathml.encode("utf-8"))
            omml_tree = transform(mathml_tree)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p._element.append(omml_tree.getroot())
            p.paragraph_format.space_after = Pt(4)
            return

        # If no XSLT, fallback to plain
        _add_paragraph(doc, latex, lang_code=lang_code, font_size=11, italic=True,
                       alignment=WD_ALIGN_PARAGRAPH.CENTER)
    except Exception:
        _add_paragraph(doc, latex, lang_code=lang_code, font_size=11, italic=True,
                       alignment=WD_ALIGN_PARAGRAPH.CENTER)


def _get_mml2omml_xslt() -> str | None:
    """Return the MML2OMML XSLT stylesheet content if available.

    Looks for the XSLT file bundled with Microsoft Office or
    in a local resource. Returns None if not found.
    """
    global _MML2OMML_XSL
    if _MML2OMML_XSL is not None:
        return _MML2OMML_XSL if _MML2OMML_XSL else None

    # Common locations
    candidates = [
        Path("/Applications/Microsoft Word.app/Contents/Resources/MML2OMML.XSL"),
        Path.home() / "MML2OMML.XSL",
        Path(__file__).resolve().parent / "MML2OMML.XSL",
    ]
    # Windows paths
    import platform
    if platform.system() == "Windows":
        import glob
        candidates.extend(
            Path(p) for p in glob.glob(r"C:\Program Files*\Microsoft Office\**\MML2OMML.XSL", recursive=True)
        )

    for p in candidates:
        if p.is_file():
            _MML2OMML_XSL = p.read_text(encoding="utf-8")
            return _MML2OMML_XSL

    _MML2OMML_XSL = ""
    return None


def _parse_md_math_blocks(text: str) -> List[Dict[str, str]]:
    """Parse markdown text into a list of content blocks.

    Returns list of dicts with keys:
    - "type": "text" | "formula_display" | "formula_inline"
    - "content": the text or LaTeX content
    """
    blocks: List[Dict[str, str]] = []
    lines = text.split("\n")
    i = 0
    current_text_lines: List[str] = []

    while i < len(lines):
        line = lines[i]
        # Display math: $$...$$
        if line.strip().startswith("$$"):
            # Flush accumulated text
            if current_text_lines:
                blocks.append({"type": "text", "content": "\n".join(current_text_lines)})
                current_text_lines = []
            # Collect formula lines until closing $$
            formula_lines = []
            first_line = line.strip()
            if first_line == "$$":
                # Multi-line formula
                i += 1
                while i < len(lines) and not lines[i].strip().startswith("$$"):
                    formula_lines.append(lines[i])
                    i += 1
            else:
                # Single-line $$formula$$
                inner = first_line[2:]
                if inner.endswith("$$"):
                    inner = inner[:-2]
                formula_lines.append(inner)
            blocks.append({"type": "formula_display", "content": "\n".join(formula_lines)})
        elif line.strip().startswith("**") or line.strip().startswith("---"):
            # Keep formatting markers as text
            current_text_lines.append(line)
        else:
            current_text_lines.append(line)
        i += 1

    if current_text_lines:
        blocks.append({"type": "text", "content": "\n".join(current_text_lines)})

    return blocks


def _render_md_content(doc, md_text: str, lang_code: str = "tr-TR") -> None:
    """Render parsed markdown content (text + formulas) into the DOCX."""
    blocks = _parse_md_math_blocks(md_text)
    for block in blocks:
        if block["type"] == "formula_display":
            content = block["content"].strip()
            if content:
                _add_math_block(doc, [content], lang_code)
        elif block["type"] == "text":
            for line in block["content"].split("\n"):
                line = line.strip()
                if not line:
                    continue
                if line.startswith("---"):
                    continue
                # Bold markdown
                if line.startswith("**") and line.endswith("**"):
                    _add_paragraph(doc, line.strip("*").strip(), lang_code=lang_code, bold=True)
                elif line.startswith("Adım") or line.startswith("Step"):
                    _add_paragraph(doc, line, lang_code=lang_code)
                else:
                    # Handle inline math $...$
                    if "$" in line:
                        _add_paragraph(doc, _strip_inline_math(line), lang_code=lang_code, italic=False)
                    else:
                        _add_paragraph(doc, line, lang_code=lang_code)


def _strip_inline_math(text: str) -> str:
    """Remove $ delimiters from inline math, keeping the content as plain text."""
    return re.sub(r"\$([^$]+)\$", r"\1", text)


# ---------------------------------------------------------------------------
# IMRAD section builders
# ---------------------------------------------------------------------------

def _build_title(d: Dict[str, Any], lang: str) -> str:
    h = _h(lang)
    wm = d["weight_method"] or "?"
    rm = d["ranking_method"] or ""
    if rm:
        if lang == "EN":
            return f"Multi-Criteria Decision Analysis with {wm} and {rm} Integration"
        return f"{wm} ve {rm} Entegrasyonuyla Çok Kriterli Karar Analizi"
    else:
        if lang == "EN":
            return f"Criteria Weighting Analysis Using {wm} Method"
        return f"{wm} Yöntemiyle Kriter Ağırlıklandırma Analizi"


def _build_abstract(d: Dict[str, Any], lang: str) -> str:
    f = _fmt
    wm = d["weight_method"]
    rm = d["ranking_method"] or ""

    if lang == "EN":
        txt = (
            f"A decision matrix of {d['n_alt']} alternatives and {d['n_crit']} criteria was analyzed. "
            f"Criteria weights were determined using {wm}"
        )
        if rm:
            txt += f" and alternatives were ranked via {rm}."
        else:
            txt += "."
        if d["mc_n"] > 0:
            txt += f" Robustness was verified through {d['mc_n']}-iteration Monte Carlo simulation and local sensitivity analysis."
        if d["top_alt"] and d["stability"] > 0:
            txt += f" {d['top_alt']} ranks first with {f(d['stability'] * 100, lang, 1)}% stability"
        if d["mean_spearman"] > 0:
            txt += f" and mean Spearman ρ = {f(d['mean_spearman'], lang)}."
        else:
            txt += "."
    else:
        txt = (
            f"{d['n_alt']} alternatif ve {d['n_crit']} kriterden oluşan karar matrisi analiz edilmiştir. "
            f"Kriter ağırlıkları {wm} yöntemiyle belirlenmiş"
        )
        if rm:
            txt += f", alternatifler {rm} ile sıralanmıştır."
        else:
            txt += "tir."
        if d["mc_n"] > 0:
            txt += f" Sağlamlık {d['mc_n']} iterasyonlu Monte Carlo simülasyonu ve lokal duyarlılık analizi ile doğrulanmıştır."
        if d["top_alt"] and d["stability"] > 0:
            txt += f" {d['top_alt']} alternatifi %{f(d['stability'] * 100, lang, 1)} kararlılıkla birinci sıradadır"
        if d["mean_spearman"] > 0:
            txt += f", ortalama Spearman ρ = {f(d['mean_spearman'], lang)}."
        else:
            txt += "."

    return txt


def _write_introduction(doc, d: Dict[str, Any], lang: str, lang_code: str) -> None:
    """Write the Introduction section."""
    h = _h(lang)
    _add_heading(doc, h["introduction"], level=1, lang_code=lang_code)
    _add_paragraph(doc,
        "This section should present the research problem, literature summary, research gap, and contribution of the study." if lang == "EN"
        else "Bu bölümde araştırma probleminin tanımı, literatür özeti, araştırma boşluğu ve çalışmanın katkısı sunulmalıdır.",
        lang_code=lang_code, italic=True, font_size=10, space_after=10)

    wm = d["weight_method"]
    rm = d["ranking_method"] or ""

    if lang == "EN":
        _add_paragraph(doc, (
            f"This study employs {wm} for criteria weighting"
            + (f" and {rm} for alternative ranking" if rm else "")
            + f" on a decision matrix of {d['n_alt']} alternatives and {d['n_crit']} criteria."
            + " The methodology integrates objective weight determination with systematic ranking and multi-layered robustness verification."
        ), lang_code=lang_code)
    else:
        _add_paragraph(doc, (
            f"Bu çalışmada {d['n_alt']} alternatif ve {d['n_crit']} kriterden oluşan karar matrisi üzerinde "
            f"{wm} ile kriter ağırlıklandırma"
            + (f" ve {rm} ile alternatif sıralama" if rm else "")
            + " uygulanmıştır."
            + " Metodoloji, nesnel ağırlık belirlemeyi sistematik sıralama ve çok katmanlı sağlamlık doğrulamasıyla bütünleştirmektedir."
        ), lang_code=lang_code)


# ---------------------------------------------------------------------------
# Method justification builders — WHY a method was chosen
# ---------------------------------------------------------------------------

# Justification rationale per weight method
_WEIGHT_JUSTIFICATION_TR: Dict[str, str] = {
    "Entropy": (
        "Kriter ağırlıklarının belirlenmesinde Shannon bilgi entropisi yöntemi tercih edilmiştir. "
        "Bu tercihin temel gerekçesi, yöntemin karar vericinin öznel yargısına ihtiyaç duymaksızın "
        "yalnızca verinin kendi bilgi yapısından hareket etmesidir. Entropi, alternatifler arasında "
        "yüksek ayrışma gösteren kriterlere daha yüksek ağırlık atarken, tüm alternatiflerde benzer "
        "değer alan — dolayısıyla karar sürecine katkısı sınırlı olan — kriterlerin ağırlığını otomatik "
        "olarak bastırır. Bu özellik, özellikle uzman görüşüne erişimin kısıtlı olduğu veya farklı uzmanlar "
        "arasında sistematik yargı tutarsızlığı riski taşıyan durumlarda nesnel ve tekrar üretilebilir "
        "bir ağırlıklandırma çerçevesi sunmaktadır."
    ),
    "CRITIC": (
        "Kriter ağırlıklarının belirlenmesinde CRITIC (Criteria Importance Through Intercriteria Correlation) "
        "yöntemi tercih edilmiştir. Bu tercihin temel gerekçesi, CRITIC'in kriter önemini yalnızca varyansla değil, "
        "aynı zamanda kriterler arası çatışma yapısını da dikkate alarak belirlemesidir. Bir kriter hem yüksek "
        "varyansa sahipse hem de diğer kriterlerle düşük korelasyon gösteriyorsa, karar sürecine benzersiz bilgi "
        "katkısı sağladığı anlaşılır. Bu çift boyutlu değerlendirme, birbirine paralel hareket eden kriterlerin "
        "toplam ağırlık payını sınırlandırarak çoklu sayım (double counting) riskini azaltır."
    ),
    "Standart Sapma": (
        "Kriter ağırlıklarının belirlenmesinde standart sapma yöntemi tercih edilmiştir. Bu yöntem, "
        "alternatifleri birbirinden daha fazla ayıran kriterlere daha yüksek ağırlık atar. Diğer objektif "
        "yöntemlere kıyasla hesaplama açısından daha basit ve şeffaftır; ancak kriterler arası korelasyon "
        "yapısını göz ardı ettiği için yüksek korelasyonlu kriter gruplarının varlığında dikkatli yorumlanmalıdır."
    ),
    "MEREC": (
        "Kriter ağırlıklarının belirlenmesinde MEREC (Method Based on the Removal Effects of Criteria) "
        "yöntemi tercih edilmiştir. MEREC'in ayırt edici özelliği, her kriterin sisteme katkısını o kriter "
        "çıkarıldığında toplam performansın ne kadar bozulduğu üzerinden ölçmesidir. Bu 'olmasa ne olurdu' "
        "mantığı, kriterin bilgi içeriğini doğrudan sistemik bağlamda değerlendirir ve kriter bağımsızlığı "
        "varsayımına ihtiyaç duymaz."
    ),
    "LOPCOW": (
        "Kriter ağırlıklarının belirlenmesinde LOPCOW (Logarithmic Percentage Change-driven Objective Weighting) "
        "yöntemi tercih edilmiştir. LOPCOW, RMS/standart sapma oranını logaritmik ölçekte değerlendirerek "
        "ölçek farklarının ağırlıklar üzerindeki bozucu etkisini bastırır. Bu özellik, kriterlerin farklı "
        "birimlerde veya büyüklük derecelerinde ölçüldüğü veri setlerinde özellikle avantajlıdır."
    ),
    "PCA": (
        "Kriter ağırlıklarının belirlenmesinde Temel Bileşen Analizi (PCA) tercih edilmiştir. PCA, "
        "kriterler arası korelasyon yapısını ortaya çıkararak bilgi taşıyan ana eksenleri belirler ve "
        "bu eksenlerdeki yük dağılımlarını ağırlıklara dönüştürür. Bu yaklaşım, yüksek boyutlu ve güçlü "
        "ilişkili kriter setlerinde çoklu sayım riskini azaltırken temel varyans kaynaklarını ön plana çıkarır."
    ),
    "IDOCRIW": (
        "Kriter ağırlıklarının belirlenmesinde IDOCRIW (Integrated Determination of Objective Criteria Weights) "
        "yöntemi tercih edilmiştir. IDOCRIW, entropi temelli bilgi içeriği ile CILOS tabanlı göreli etki kaybını "
        "eleman çarpımıyla bütünleştirir. Bu hibrit yapı, hem bilgi çeşitliliğini hem de kriterin çıkarılmasının "
        "yarattığı performans kaybını eş zamanlı değerlendirerek tek başına entropiden veya CILOS'tan daha "
        "dengeli bir ağırlık vektörü üretir."
    ),
    "Fuzzy IDOCRIW": (
        "Kriter ağırlıklarının belirlenmesinde Fuzzy IDOCRIW yöntemi tercih edilmiştir. Bu yöntem, IDOCRIW "
        "mantığını belirsizlik altında üç senaryolu (alt-orta-üst) bulanık bir çerçeveye taşıyarak veri "
        "kesinliksizliğinin ağırlık hesabı üzerindeki etkisini de modele dahil eder."
    ),
    "CILOS": (
        "Kriter ağırlıklarının belirlenmesinde CILOS (Criterion Impact Loss) yöntemi tercih edilmiştir. "
        "CILOS, her kriter çıkarıldığında en iyi alternatifin diğer kriterler üzerinde yarattığı kayıp "
        "etkisini ölçerek ağırlıklandırma yapar. Bu doğrusal denklem sistemi temelli yaklaşım, kriter "
        "etkileşim yapısını doğrudan modele yansıtır."
    ),
    "AHP": (
        "Kriter ağırlıklarının belirlenmesinde Analitik Hiyerarşi Süreci (AHP) tercih edilmiştir. "
        "AHP'nin tercih edilme gerekçesi, alandaki uzman bilgisinin sistematik biçimde yapılandırılmasına "
        "olanak tanımasıdır. İkili karşılaştırma matrisi, uzmanın kriterler arasındaki göreli önem "
        "yargısını tutarlılık kontrolü (CR) altında ağırlık vektörüne dönüştürür. Tutarlılık oranının "
        "kabul edilebilir düzeyde çıkması, uzman yargısının içsel çelişki taşımadığını doğrular."
    ),
    "BWM": (
        "Kriter ağırlıklarının belirlenmesinde En İyi-En Kötü Yöntemi (BWM) tercih edilmiştir. "
        "BWM, AHP'ye kıyasla çok daha az karşılaştırma (2n-3 yerine n(n-1)/2) gerektirmesi ve optimize "
        "edilmiş tutarlılık yapısıyla uzman yükünü azaltır. Bu özellik, kriter sayısı arttığında veya "
        "uzman zamanının kısıtlı olduğu durumlarda belirleyici bir avantaj sağlar."
    ),
    "SWARA": (
        "Kriter ağırlıklarının belirlenmesinde SWARA (Stepwise Weight Assessment Ratio Analysis) tercih "
        "edilmiştir. SWARA, uzmanın kriterleri önem sırasına göre dizdikten sonra ardışık göreli önem "
        "oranları belirlemesini sağlar. Bu adımsal yapı, ikili karşılaştırma matrislerine göre daha sezgisel "
        "ve hızlıdır; özellikle uzmanın kriterler arasında net bir önem hiyerarşisi gördüğü durumlarda uygundur."
    ),
    "DEMATEL": (
        "Kriter ağırlıklarının belirlenmesinde DEMATEL (Decision Making Trial and Evaluation Laboratory) "
        "tercih edilmiştir. DEMATEL'in ayırt edici özelliği, kriterleri yalnızca önem düzeyiyle değil, "
        "aynı zamanda birbirlerini ne ölçüde etkiledikleriyle değerlendirmesidir. Toplam ilişki matrisinden "
        "elde edilen 'neden-sonuç' yapısı, karar probleminin yapısal dinamiğini ortaya koyar."
    ),
    "SMART": (
        "Kriter ağırlıklarının belirlenmesinde SMART (Simple Multi-Attribute Rating Technique) tercih "
        "edilmiştir. SMART, uzmanın her kritere doğrudan puan atamasına dayanan en basit subjektif "
        "ağırlıklandırma yöntemidir. Karmaşık tutarlılık kontrolleri gerektirmez; ağırlıkların şeffaf "
        "ve doğrudan yorumlanabilir olması istenen durumlarda tercih edilir."
    ),
}

_WEIGHT_JUSTIFICATION_EN: Dict[str, str] = {
    "Entropy": (
        "Shannon information entropy was selected for criteria weighting. The rationale is that the method derives "
        "weights solely from the data's intrinsic information structure without requiring subjective expert judgment. "
        "Entropy assigns higher weights to criteria that exhibit greater discrimination among alternatives while "
        "automatically suppressing criteria with near-uniform values — which contribute little to the decision "
        "process. This property provides an objective and reproducible weighting framework, particularly valuable "
        "when expert access is limited or when systematic judgment inconsistency across experts poses a risk."
    ),
    "CRITIC": (
        "CRITIC (Criteria Importance Through Intercriteria Correlation) was selected for criteria weighting. "
        "The rationale is that CRITIC determines criteria importance not only through variance but also by "
        "accounting for the conflict structure among criteria. A criterion that has both high variance and low "
        "correlation with other criteria provides unique informational contributions to the decision process. "
        "This dual-dimensional evaluation limits the combined weight share of parallel-moving criteria, "
        "thereby reducing double-counting risk."
    ),
    "Standart Sapma": (
        "The standard deviation (SD) method was selected for criteria weighting. SD assigns higher weights to "
        "criteria with greater dispersion across alternatives, on the principle that a criterion that does not "
        "vary cannot discriminate. Compared to CRITIC, SD does not adjust for intercriteria correlation; this "
        "makes it computationally simpler and fully transparent, but users should be aware that highly correlated "
        "criteria may receive redundant weight shares."
    ),
    "MEREC": (
        "MEREC (Method Based on the Removal Effects of Criteria) was selected for criteria weighting. "
        "The distinctive feature of MEREC is that it measures each criterion's contribution by quantifying "
        "how much the overall performance deteriorates when that criterion is removed. This 'what if it were absent' "
        "logic evaluates information content directly in the systemic context without requiring a criterion "
        "independence assumption."
    ),
    "LOPCOW": (
        "LOPCOW (Logarithmic Percentage Change-driven Objective Weighting) was selected for criteria weighting. "
        "LOPCOW evaluates the RMS-to-standard-deviation ratio on a logarithmic scale, which effectively dampens "
        "the distorting effect of scale differences among criteria measured in different units or orders of magnitude. "
        "This property makes LOPCOW particularly suitable for heterogeneous criterion sets."
    ),
    "PCA": (
        "Principal Component Analysis (PCA) was selected for criteria weighting. PCA uncovers the latent variance "
        "structure among criteria and converts factor loadings on the retained components into criterion weights. "
        "This approach reduces double-counting risk in high-dimensional, strongly correlated criterion sets by "
        "channeling weight through the principal directions of variance."
    ),
    "CILOS": (
        "CILOS (Criterion Impact Loss) was selected for criteria weighting. CILOS quantifies how much the best "
        "alternative's performance on other criteria deteriorates when a given criterion is removed. The resulting "
        "linear equation system directly encodes the intercriteria impact structure into the weight vector."
    ),
    "IDOCRIW": (
        "IDOCRIW (Integrated Determination of Objective Criteria Weights) was selected for criteria weighting. "
        "IDOCRIW integrates Entropy-based information content with CILOS-based relative impact loss through "
        "element-wise multiplication. This hybrid structure simultaneously accounts for both information diversity "
        "and systemic criterion impact, producing a more balanced weight vector than either component alone."
    ),
    "Fuzzy IDOCRIW": (
        "Fuzzy IDOCRIW was selected for criteria weighting. This method extends the IDOCRIW framework to a "
        "three-scenario fuzzy setting (lower-middle-upper), thereby incorporating data imprecision into the "
        "weight calculation itself. The averaged scenario weights reflect the weight vector's sensitivity "
        "to measurement uncertainty at the weighting stage."
    ),
    "AHP": (
        "The Analytic Hierarchy Process (AHP) was selected for criteria weighting. AHP enables systematic "
        "structuring of domain expert knowledge through a pairwise comparison matrix that converts relative "
        "importance judgments into a weight vector under consistency control (CR). An acceptable consistency "
        "ratio confirms that the expert judgment is free of internal contradictions."
    ),
    "BWM": (
        "The Best-Worst Method (BWM) was selected for criteria weighting. Compared to AHP, BWM requires "
        "substantially fewer comparisons (2n−3 versus n(n−1)/2) while maintaining an optimized consistency "
        "structure, reducing expert burden. This is a decisive advantage when the number of criteria increases "
        "or expert time is limited."
    ),
    "SWARA": (
        "SWARA (Stepwise Weight Assessment Ratio Analysis) was selected for criteria weighting. SWARA asks the "
        "expert to rank criteria by importance and then assign stepwise relative importance ratios. This sequential "
        "structure is more intuitive than pairwise comparison matrices and is particularly suitable when the expert "
        "perceives a clear importance hierarchy among criteria."
    ),
    "DEMATEL": (
        "DEMATEL (Decision Making Trial and Evaluation Laboratory) was selected for criteria weighting. "
        "The distinctive feature of DEMATEL is that it evaluates criteria not only by importance but also by "
        "their mutual influence — identifying which criteria are causes and which are effects. The total relation "
        "matrix reveals the structural dynamics of the decision problem, and the prominence scores (D+R) serve "
        "as influence-aware weights."
    ),
    "SMART": (
        "SMART (Simple Multi-Attribute Rating Technique) was selected for criteria weighting. SMART converts "
        "direct expert point scores into normalized weights without requiring complex consistency checks. "
        "This transparency makes the method appropriate when straightforward, easily interpretable weights are preferred."
    ),
}

# Ranking method justification rationale
_RANKING_JUSTIFICATION_TR: Dict[str, str] = {
    "TOPSIS": (
        "Alternatiflerin sıralanmasında TOPSIS yöntemi tercih edilmiştir. TOPSIS'in tercih gerekçesi, "
        "alternatifleri hem pozitif ideal çözüme yakınlık hem de negatif ideal çözüme uzaklık bakımından "
        "eş zamanlı değerlendirmesidir. Bu çift referanslı yapı, yalnızca en iyi senaryoya yakınlığı "
        "ölçen yöntemlere göre daha dengeleyici bir sıralama üretir."
    ),
    "VIKOR": (
        "Alternatiflerin sıralanmasında VIKOR yöntemi tercih edilmiştir. VIKOR, çoğunluğun faydasını "
        "en yüksek düzeye çıkarırken en büyük bireysel pişmanlığı en aza indiren bir uzlaşı çözümü arar. "
        "v parametresi aracılığıyla grup faydası ile bireysel pişmanlık arasındaki denge ayarlanabilir. "
        "C1 ve C2 uzlaşı koşulları, lider alternatifin savunulabilir bir üstünlük marjına sahip olup "
        "olmadığını doğrudan test eder."
    ),
    "EDAS": (
        "Alternatiflerin sıralanmasında EDAS yöntemi tercih edilmiştir. EDAS, ortalama çözümü referans alarak "
        "pozitif ve negatif sapmaları değerlendirir. İdeal çözüm referansına göre aşırı uç değerlere daha az "
        "duyarlıdır; bu da aykırı değer içeren veri setlerinde daha kararlı sıralama üretir."
    ),
    "CODAS": (
        "Alternatiflerin sıralanmasında CODAS yöntemi tercih edilmiştir. CODAS, negatif ideal çözümden Öklid "
        "uzaklığını birincil, Manhattan uzaklığını ikincil ayırt edici olarak kullanır. Bu çift mesafe yapısı, "
        "Öklid uzaklıkları birbirine yakın alternatifleri ayırt etmede ek bir ayrışma katmanı sağlar."
    ),
    "COPRAS": (
        "Alternatiflerin sıralanmasında COPRAS yöntemi tercih edilmiştir. COPRAS, fayda ve maliyet katkılarını "
        "ayrı ayrı izleyerek göreli önem skorlarını üretir. Bu ayrıştırma, kriter yönlerinin sıralama üzerindeki "
        "etkisini şeffaf biçimde takip etmeye olanak tanır."
    ),
    "OCRA": (
        "Alternatiflerin sıralanmasında OCRA yöntemi tercih edilmiştir. OCRA, fayda ve maliyet kriterlerindeki "
        "göreli rekabet üstünlüğünü ayrı bileşenler olarak toplayarak operasyonel performans temelli bir sıralama "
        "sunar."
    ),
    "ARAS": (
        "Alternatiflerin sıralanmasında ARAS yöntemi tercih edilmiştir. ARAS, her alternatifin performansını "
        "yapay ideal alternatifle karşılaştırarak fayda derecesi üretir. Bu oran temelli yorum, karar vericiye "
        "her alternatifin ideale ne kadar yakın olduğunu doğrudan gösterir."
    ),
    "SAW": (
        "Alternatiflerin sıralanmasında SAW (Simple Additive Weighting) yöntemi tercih edilmiştir. SAW, en temel "
        "telafi edici toplama modelidir: bir kriterdeki düşük performans diğer kriterlerdeki yüksek performansla "
        "telafi edilebilir. Bu şeffaf yapı, sonuçların yorumlanabilirliğini en üst düzeye çıkarır."
    ),
    "WPM": (
        "Alternatiflerin sıralanmasında WPM (Weighted Product Model) yöntemi tercih edilmiştir. WPM, kriter "
        "değerlerini ağırlıklı üsler altında çarpar; bu çarpımsal yapı ölçekten bağımsızdır ve zayıf boyutları "
        "toplamsal modellere göre daha güçlü cezalandırır."
    ),
    "MAUT": (
        "Alternatiflerin sıralanmasında MAUT (Multi-Attribute Utility Theory) yöntemi tercih edilmiştir. "
        "MAUT, her kriterin fayda fonksiyonunu doğrusal olarak modelleyerek toplam beklenen faydayı hesaplar. "
        "Normatif karar teorisine dayanan yapısı, aksiyomatik olarak tutarlı bir sıralama garantisi sunar."
    ),
    "WASPAS": (
        "Alternatiflerin sıralanmasında WASPAS yöntemi tercih edilmiştir. WASPAS, toplamsal (WSM) ve çarpımsal "
        "(WPM) modelleri λ parametresiyle birleştirerek her iki yaklaşımın güçlü yönlerinden yararlanır. "
        "Bu hibrit yapı, tek bir toplulaştırma stratejisine bağımlılığı azaltır."
    ),
    "MOORA": (
        "Alternatiflerin sıralanmasında MOORA yöntemi tercih edilmiştir. MOORA, vektör normalizasyonu ardından "
        "fayda ve maliyet katkılarının net farkını hesaplar. Hesaplama sadeliği ve sonuçların doğrudan "
        "yorumlanabilirliği yöntemin temel avantajlarıdır."
    ),
    "MULTIMOORA": (
        "Alternatiflerin sıralanmasında MULTIMOORA yöntemi tercih edilmiştir. MULTIMOORA, oran sistemi, referans "
        "nokta ve tam çarpımsal form olmak üzere üç bağımsız bileşeni Borda toplamıyla birleştirir. Bu çok "
        "bileşenli yapı, tek bir toplulaştırma mantığına güvenmenin riskini azaltarak yapısal sağlamlık sağlar."
    ),
    "MABAC": (
        "Alternatiflerin sıralanmasında MABAC yöntemi tercih edilmiştir. MABAC, alternatiflerin sınır yaklaşım "
        "alanına (geometrik ortalama) göre mesafesini ölçer. Pozitif mesafe üst yarıya, negatif mesafe alt yarıya "
        "ait olduğunu gösterir; bu sezgisel ayrım karar vericiye net bir üstünlük/zayıflık haritası sunar."
    ),
    "MARCOS": (
        "Alternatiflerin sıralanmasında MARCOS yöntemi tercih edilmiştir. MARCOS, ideal ve anti-ideal referans "
        "noktalarına göre fayda dereceleri hesaplayarak yararlılık fonksiyonu üretir. Bu çift referanslı oran "
        "yapısı, alternatiflerin hem en iyi hem de en kötü senaryoya göre konumunu birlikte değerlendirir."
    ),
    "CoCoSo": (
        "Alternatiflerin sıralanmasında CoCoSo yöntemi tercih edilmiştir. CoCoSo, toplamsal ve üstel "
        "toplulaştırma stratejilerini üç bileşik strateji aracılığıyla birleştirir. λ parametresiyle ayarlanabilen "
        "bu çoklu bileşim, tek bir toplulaştırma mantığının potansiyel sapmasını dengeleyerek daha tutarlı "
        "sıralama üretir."
    ),
    "PROMETHEE": (
        "Alternatiflerin sıralanmasında PROMETHEE II yöntemi tercih edilmiştir. PROMETHEE, ikili karşılaştırma "
        "temelli tercih fonksiyonları aracılığıyla pozitif ve negatif akışları hesaplar. Net akış (φ) değeri, "
        "her alternatifin diğerlerine göre toplam üstünlüğünü ölçer. Tercih fonksiyonunun q ve p parametreleri, "
        "karar vericinin kayıtsızlık ve tercih eşiklerini açıkça tanımlamasına olanak verir."
    ),
    "GRA": (
        "Alternatiflerin sıralanmasında GRA (Grey Relational Analysis) yöntemi tercih edilmiştir. GRA, her "
        "alternatifin ideal referans diziye ne kadar benzer olduğunu gri ilişkisel katsayı üzerinden ölçer. "
        "ρ parametresi, uç değerlerin etkisini kontrol altında tutar. Yöntem, eksik bilgi veya küçük örneklem "
        "boyutlarında teorik olarak güçlüdür."
    ),
    "SPOTIS": (
        "Alternatiflerin sıralanmasında SPOTIS yöntemi tercih edilmiştir. SPOTIS'in tercih gerekçesi, sabit ideal "
        "referans noktası kullanarak rank reversal problemine yapısal düzeyde direnmesidir. Yeni alternatif "
        "eklendiğinde veya çıkarıldığında mevcut alternatiflerin skorları değişmez; bu özellik, dinamik karar "
        "ortamlarında sıralamanın güvenilirliğini artırır."
    ),
    "RAWEC": (
        "Alternatiflerin sıralanmasında RAWEC yöntemi tercih edilmiştir. RAWEC, her kriter içindeki alternatif "
        "sırasını o kriterin ağırlığıyla ağırlıklandırarak harmonik toplulaştırma yapar. Önemli kriterlerde üst "
        "sırada olmak belirleyici avantaj sağlarken, zayıf kriterlerdeki düşük sıra daha az cezalandırılır."
    ),
    "RAFSI": (
        "Alternatiflerin sıralanmasında RAFSI yöntemi tercih edilmiştir. RAFSI, kriter değerlerini sabit bir "
        "referans aralığa [1–9] eşleyerek rank reversal problemini yapısal olarak önler. SPOTIS gibi, yeni "
        "alternatif eklenmesi mevcut skorları değiştirmez."
    ),
    "ROV": (
        "Alternatiflerin sıralanmasında ROV yöntemi tercih edilmiştir. ROV, min-max normalleştirilmiş performans "
        "değerlerini ağırlıklı toplamsal modelde birleştirerek fayda aralığını tam kullanır. Basit ve şeffaf "
        "yapısı, sonuçların doğrudan yorumlanmasını kolaylaştırır."
    ),
    "AROMAN": (
        "Alternatiflerin sıralanmasında AROMAN yöntemi tercih edilmiştir. AROMAN, sum ve min-max normalizasyonu "
        "adımlarının geometrik bileşimini kullanarak tek bir normalizasyon yöntemine bağımlılığı azaltır. "
        "Bu çift normalleştirme stratejisi, normalleştirme kaynaklı sapmayı sınırlar."
    ),
    "DNMA": (
        "Alternatiflerin sıralanmasında DNMA yöntemi tercih edilmiştir. DNMA, min-max ve sum normalizasyonlarından "
        "elde edilen skorları eşit ağırlıkla ortalar. Hiçbir tek normalizasyon yaklaşımına tam güvenilmediği "
        "durumlarda, bu çift normalleştirme stratejisi normalleştirme duyarlılığını azaltır."
    ),
}

_RANKING_JUSTIFICATION_EN: Dict[str, str] = {
    "TOPSIS": (
        "TOPSIS was selected for alternative ranking. TOPSIS simultaneously evaluates alternatives based on "
        "proximity to the positive ideal solution and distance from the negative ideal solution. This dual-reference "
        "structure produces a more balanced ranking compared to methods that only measure proximity to the best scenario."
    ),
    "VIKOR": (
        "VIKOR was selected for alternative ranking. VIKOR seeks a compromise solution that maximizes group utility "
        "while minimizing individual regret. The v parameter adjusts the balance between utility and regret, and "
        "the C1/C2 compromise conditions directly test whether the leader has a defensible superiority margin."
    ),
    "EDAS": (
        "EDAS was selected for alternative ranking. EDAS uses the average solution as reference, making it less "
        "sensitive to extreme values than ideal-solution-based methods. This yields more stable rankings in data "
        "sets with outliers or small sample sizes."
    ),
    "CODAS": (
        "CODAS was selected for alternative ranking. CODAS uses Euclidean distance as the primary and Manhattan "
        "distance as the secondary discriminator from the negative ideal solution. This dual-distance structure "
        "provides an additional layer of discrimination when Euclidean distances are close."
    ),
    "COPRAS": (
        "COPRAS was selected for alternative ranking. COPRAS tracks benefit and cost contributions separately, "
        "making the directional effect of each criterion on the ranking transparently traceable."
    ),
    "OCRA": (
        "OCRA was selected for alternative ranking. OCRA aggregates relative competitive advantage across benefit "
        "and cost components, producing an operational performance-based ranking."
    ),
    "ARAS": (
        "ARAS was selected for alternative ranking. ARAS compares each alternative's performance to an artificial "
        "ideal alternative, producing a utility degree ratio that directly shows how close each alternative is "
        "to the ideal."
    ),
    "SAW": (
        "SAW (Simple Additive Weighting) was selected for alternative ranking. As the most fundamental compensatory "
        "aggregation model, SAW maximizes result interpretability: poor performance on one criterion can be "
        "compensated by strong performance on others."
    ),
    "WPM": (
        "WPM (Weighted Product Model) was selected for alternative ranking. WPM multiplies criterion values "
        "under weighted exponents; this multiplicative structure is scale-invariant and penalizes weak dimensions "
        "more heavily than additive models."
    ),
    "MAUT": (
        "MAUT (Multi-Attribute Utility Theory) was selected for alternative ranking. Rooted in normative decision "
        "theory, MAUT provides an axiomatically consistent ranking guarantee through linear utility functions."
    ),
    "WASPAS": (
        "WASPAS was selected for alternative ranking. WASPAS combines additive (WSM) and multiplicative (WPM) "
        "models via the λ parameter, leveraging the strengths of both approaches and reducing dependence on a "
        "single aggregation strategy."
    ),
    "MOORA": (
        "MOORA was selected for alternative ranking. MOORA computes the net difference between benefit and cost "
        "contributions after vector normalization. Its computational simplicity and direct interpretability are "
        "key advantages."
    ),
    "MULTIMOORA": (
        "MULTIMOORA was selected for alternative ranking. MULTIMOORA combines three independent components — "
        "ratio system, reference point, and full multiplicative form — via Borda aggregation. This multi-component "
        "structure reduces the risk of relying on a single aggregation logic."
    ),
    "MABAC": (
        "MABAC was selected for alternative ranking. MABAC measures each alternative's distance from the border "
        "approximation area (geometric mean). Positive distance indicates upper-half performance; negative indicates "
        "lower-half — providing an intuitive superiority/weakness map."
    ),
    "MARCOS": (
        "MARCOS was selected for alternative ranking. MARCOS computes utility degrees relative to both ideal and "
        "anti-ideal reference points, jointly evaluating each alternative's position against best-case and "
        "worst-case scenarios."
    ),
    "CoCoSo": (
        "CoCoSo was selected for alternative ranking. CoCoSo combines additive and exponential aggregation "
        "strategies through three composite strategies, reducing the bias of any single aggregation logic."
    ),
    "PROMETHEE": (
        "PROMETHEE II was selected for alternative ranking. PROMETHEE computes positive and negative preference "
        "flows through pairwise comparison-based preference functions. The q and p parameters allow explicit "
        "definition of indifference and preference thresholds."
    ),
    "GRA": (
        "GRA (Grey Relational Analysis) was selected for alternative ranking. GRA measures how similar each "
        "alternative is to the ideal reference series via grey relational coefficients. The method is theoretically "
        "robust under incomplete information or small sample sizes."
    ),
    "SPOTIS": (
        "SPOTIS was selected for alternative ranking. SPOTIS uses a fixed ideal reference point, making it "
        "structurally resistant to rank reversal. Adding or removing alternatives does not change existing scores, "
        "enhancing ranking reliability in dynamic decision environments."
    ),
    "RAWEC": (
        "RAWEC was selected for alternative ranking. RAWEC weights each alternative's within-criterion rank by "
        "the criterion's importance through harmonic aggregation. High ranks on important criteria provide "
        "decisive advantage."
    ),
    "RAFSI": (
        "RAFSI was selected for alternative ranking. RAFSI maps criterion values to a fixed reference interval "
        "[1–9], structurally preventing rank reversal — like SPOTIS, adding alternatives does not alter existing scores."
    ),
    "ROV": (
        "ROV was selected for alternative ranking. ROV combines min-max normalized performance values in a weighted "
        "additive model, fully utilizing the range of each criterion. Its simplicity facilitates direct interpretation."
    ),
    "AROMAN": (
        "AROMAN was selected for alternative ranking. AROMAN uses a geometric blend of sum and min-max normalization "
        "to reduce dependence on a single normalization approach, limiting normalization-induced bias."
    ),
    "DNMA": (
        "DNMA was selected for alternative ranking. DNMA equally averages scores from min-max and sum normalizations. "
        "When no single normalization approach is fully trusted, this dual-normalization strategy reduces "
        "normalization sensitivity."
    ),
}


def _write_weight_justification(doc, wm: str, d: Dict[str, Any], lang: str, lang_code: str) -> None:
    """Write the justification paragraph for the chosen weight method."""
    if lang == "EN":
        just_dict = _WEIGHT_JUSTIFICATION_EN
    else:
        just_dict = _WEIGHT_JUSTIFICATION_TR

    text = just_dict.get(wm)
    if text:
        _add_paragraph(doc, text, lang_code=lang_code)
    else:
        # Fallback: use METHOD_PHILOSOPHY + generic justification
        try:
            from mcdm_engine import METHOD_PHILOSOPHY
            phil = METHOD_PHILOSOPHY.get(wm, {})
            phil_text = phil.get("academic", "")
            if phil_text:
                if lang == "EN":
                    _add_paragraph(doc, (
                        f"The {wm} method was selected for criteria weighting. {phil_text} "
                        "This method was deemed appropriate for the current problem structure."
                    ), lang_code=lang_code)
                else:
                    _add_paragraph(doc, (
                        f"Kriter ağırlıklarının belirlenmesinde {wm} yöntemi tercih edilmiştir. {phil_text} "
                        "Bu yöntem, mevcut problem yapısı için uygun görülmüştür."
                    ), lang_code=lang_code)
        except ImportError:
            if lang == "EN":
                _add_paragraph(doc, f"The {wm} method was selected for criteria weighting.",
                               lang_code=lang_code)
            else:
                _add_paragraph(doc, f"Kriter ağırlıklarının belirlenmesinde {wm} yöntemi tercih edilmiştir.",
                               lang_code=lang_code)


def _write_ranking_justification(doc, rm: str, d: Dict[str, Any], lang: str, lang_code: str) -> None:
    """Write the justification paragraph for the chosen ranking method."""
    base_rm = rm.replace("Fuzzy ", "")
    is_fuzzy = rm.startswith("Fuzzy ")

    if lang == "EN":
        just_dict = _RANKING_JUSTIFICATION_EN
    else:
        just_dict = _RANKING_JUSTIFICATION_TR

    text = just_dict.get(rm) or just_dict.get(base_rm)

    if text:
        _add_paragraph(doc, text, lang_code=lang_code)
    else:
        try:
            from mcdm_engine import METHOD_PHILOSOPHY
            phil = METHOD_PHILOSOPHY.get(rm, {}) or METHOD_PHILOSOPHY.get(base_rm, {})
            phil_text = phil.get("academic", "")
            if phil_text:
                if lang == "EN":
                    _add_paragraph(doc, (
                        f"{rm} was selected for alternative ranking. {phil_text}"
                    ), lang_code=lang_code)
                else:
                    _add_paragraph(doc, (
                        f"Alternatiflerin sıralanmasında {rm} yöntemi tercih edilmiştir. {phil_text}"
                    ), lang_code=lang_code)
        except ImportError:
            pass

    # Fuzzy justification
    if is_fuzzy:
        if lang == "EN":
            _add_paragraph(doc, (
                f"The fuzzy variant of {base_rm} was specifically chosen to model measurement uncertainty "
                "and data imprecision inherent in the evaluation process. Rather than treating criterion values "
                "as deterministic quantities, the triangular fuzzy number (TFN) representation acknowledges that "
                "each observed value carries an implicit confidence interval. This epistemic humility strengthens "
                "the defensibility of the ranking under reviewer scrutiny, as the results explicitly account for "
                "the effect of input uncertainty on the final ordering."
            ), lang_code=lang_code)
        else:
            _add_paragraph(doc, (
                f"{base_rm} yönteminin bulanık varyantı, değerlendirme sürecine içkin ölçüm belirsizliğini ve "
                "veri kesinliksizliğini modellemek amacıyla özellikle tercih edilmiştir. Kriter değerlerini "
                "deterministik büyüklükler olarak ele almak yerine, üçgensel bulanık sayı (TFN) temsili her "
                "gözlenen değerin örtük bir güven aralığı taşıdığını kabul eder. Bu epistemik ihtiyatlılık, "
                "sonuçların girdi belirsizliğinin nihai sıralama üzerindeki etkisini açıkça hesaba katması "
                "nedeniyle sıralamanın hakem denetimi karşısındaki savunulabilirliğini güçlendirir."
            ), lang_code=lang_code)


def _write_method_section(doc, d: Dict[str, Any], selected_data: pd.DataFrame,
                          lang: str, lang_code: str) -> None:
    """Write the Method section with mathematical formulas."""
    h = _h(lang)
    template = _load_template()

    _add_heading(doc, h["method"], level=1, lang_code=lang_code)

    # Methodology flowchart
    flowchart_bytes = _generate_flowchart_bytes(d, lang)
    if flowchart_bytes:
        caption = "Şekil 1. Önerilen Yöntemin Akış Şeması" if lang != "EN" else "Figure 1. Proposed Methodology Flowchart"
        _add_figure(doc, flowchart_bytes, caption, lang_code)

    # Methodology overview paragraph
    if lang == "EN":
        overview = (
            "The proposed methodology consists of three integrated layers. "
            "In the first layer, criteria weights are determined objectively from the raw decision matrix; "
            "the resulting crisp weight vector serves as the input for all subsequent stages. "
        )
        if d["is_fuzzy"]:
            overview += (
                "In the second layer, the crisp decision matrix is converted to triangular fuzzy numbers (TFN) "
                "to account for measurement uncertainty and data imprecision; the fuzzification step is deliberately "
                "placed after weight calculation so that the weight vector remains free of artificial uncertainty inflation. "
            )
        if d["ranking_method"]:
            overview += (
                "The ranking algorithm then combines the crisp weights with the "
                + ("fuzzified " if d["is_fuzzy"] else "")
                + "decision matrix to produce a final preference ordering. "
            )
        overview += (
            "In the third layer, the robustness of the results is validated through local sensitivity analysis, "
            "Monte Carlo simulation, and cross-method comparison."
        )
        _add_paragraph(doc, overview, lang_code=lang_code)
    else:
        overview = (
            "Önerilen metodoloji birbiriyle bütünleşik üç katmandan oluşmaktadır. "
            "Birinci katmanda, kriter ağırlıkları ham karar matrisinden nesnel olarak hesaplanır; "
            "elde edilen crisp ağırlık vektörü sonraki tüm aşamalara girdi sağlar. "
        )
        if d["is_fuzzy"]:
            overview += (
                "İkinci katmanda, ölçüm belirsizliğini ve veri kesinliksizliğini modellemek amacıyla crisp karar matrisi "
                "üçgensel bulanık sayılara (TFN) dönüştürülür. Bulanıklaştırma adımı bilinçli olarak ağırlık hesabından "
                "sonraya konumlandırılmıştır; böylece ağırlık vektörü yapay belirsizlik şişirmesinden arındırılmış olur. "
            )
        if d["ranking_method"]:
            overview += (
                "Sıralama algoritması, crisp ağırlıkları "
                + ("bulanıklaştırılmış " if d["is_fuzzy"] else "")
                + "karar matrisiyle birleştirerek nihai tercih sıralamasını üretir. "
            )
        overview += (
            "Üçüncü katmanda, sonuçların sağlamlığı lokal duyarlılık analizi, "
            "Monte Carlo simülasyonu ve çapraz yöntem karşılaştırması ile doğrulanır."
        )
        _add_paragraph(doc, overview, lang_code=lang_code)

    # 2.1 Decision Matrix
    _add_heading(doc, h["method_dm"], level=2, lang_code=lang_code)
    benefit_str = ", ".join(d["benefit"]) if d["benefit"] else "-"
    cost_str = ", ".join(d["cost"]) if d["cost"] else "-"

    if lang == "EN":
        _add_mixed_paragraph(doc, [
            {"text": (
                f"In this study, {d['n_alt']} alternatives and {d['n_crit']} criteria were considered. "
                f"The criteria were defined based on the problem context"
            )},

            {"text": (
                f". The decision matrix was formulated as X(m×n) where m = {d['n_alt']} and n = {d['n_crit']}. "
                f"Each criterion was classified as either benefit-type (higher is better) or cost-type (lower is better) "
                f"based on the problem's domain requirements. "
                f"Benefit criteria: {benefit_str}. "
                f"Cost criteria: {cost_str}."
            )},
        ], lang_code)
    else:
        _add_mixed_paragraph(doc, [
            {"text": (
                f"Bu çalışmada {d['n_alt']} alternatif ve {d['n_crit']} kriter dikkate alınmıştır. "
                f"Kriterler problem bağlamına göre belirlenmiştir"
            )},

            {"text": (
                f" kaynağından elde edilmiştir. "
                f"Karar matrisi X(m×n) olarak formüle edilmiş olup m = {d['n_alt']} ve n = {d['n_crit']}'dir. "
                f"Her kriter, problemin alan gereksinimlerine göre fayda yönlü (büyük olan tercih edilir) "
                f"veya maliyet yönlü (küçük olan tercih edilir) olarak sınıflandırılmıştır. "
                f"Fayda yönlü kriterler: {benefit_str}. "
                f"Maliyet yönlü kriterler: {cost_str}."
            )},
        ], lang_code)

    # Pre-screening paragraph (if alternatives were eliminated)
    eliminated = d.get("eliminated") or []
    if eliminated:
        elim_names = sorted(set(e.get("alternative", "?") for e in eliminated))
        n_elim = len(elim_names)
        if lang == "EN":
            _add_paragraph(doc, (
                f"Prior to the main analysis, a threshold-based pre-screening filter was applied. "
                f"{n_elim} alternative(s) ({', '.join(elim_names)}) were eliminated for failing to meet "
                f"the minimum acceptable performance thresholds defined for the relevant criteria. "
                f"The remaining {d['n_alt']} alternatives proceeded to the weighting and ranking stages."
            ), lang_code=lang_code)
        else:
            _add_paragraph(doc, (
                f"Ana analizden önce eşik tabanlı bir ön eleme filtresi uygulanmıştır. "
                f"{n_elim} alternatif ({', '.join(elim_names)}) ilgili kriterler için tanımlanan minimum "
                f"kabul edilebilir performans eşiklerini karşılayamadığı için elenmiştir. "
                f"Kalan {d['n_alt']} alternatif ağırlıklandırma ve sıralama aşamalarına devam etmiştir."
            ), lang_code=lang_code)

    # 2.2 Weight Method
    _add_heading(doc, h["method_weight"], level=2, lang_code=lang_code)
    wm = d["weight_method"]

    # Justification paragraph — WHY this method
    _write_weight_justification(doc, wm, d, lang, lang_code)

    # Get mathematical content from template
    weight_content = _find_template_content(wm, template["weight_methods"])
    if weight_content:
        _render_md_content(doc, weight_content, lang_code)
    else:
        if lang == "EN":
            _add_paragraph(doc, f"The {wm} method was used for criteria weighting.", lang_code=lang_code)
        else:
            _add_paragraph(doc, f"Kriter ağırlıkları {wm} yöntemiyle belirlenmiştir.", lang_code=lang_code)

    # 2.3 Ranking Method
    if d["ranking_method"]:
        _add_heading(doc, h["method_rank"], level=2, lang_code=lang_code)
        rm = d["ranking_method"]

        # Justification paragraph — WHY this ranking method
        _write_ranking_justification(doc, rm, d, lang, lang_code)

        # Fuzzy common infrastructure — with explicit pipeline position
        if d["is_fuzzy"]:
            fuzzy_common = template["ranking_methods"].get("_fuzzy_common", "")
            if fuzzy_common:
                if lang == "EN":
                    _add_paragraph(doc, (
                        "An important architectural decision in this study is the placement of the fuzzification step. "
                        f"The crisp decision matrix is converted to triangular fuzzy numbers (TFN) with "
                        f"a spread parameter of s = {_fmt(d['fuzzy_spread'], lang)} only after the weight vector has been "
                        "computed from the original crisp data. This sequencing ensures that the weight calculation "
                        "operates on the actual observed values without artificial uncertainty amplification, "
                        "while the ranking stage benefits from the TFN-based robustness to measurement noise. "
                        "The TFN conversion is defined as follows:"
                    ), lang_code=lang_code)
                else:
                    _add_paragraph(doc, (
                        "Bu çalışmadaki önemli bir mimari karar, bulanıklaştırma adımının pipeline'daki konumudur. "
                        f"Crisp karar matrisi, ağırlık vektörü orijinal crisp veriden hesaplandıktan sonra "
                        f"s = {_fmt(d['fuzzy_spread'], lang)} spread parametresiyle üçgensel bulanık sayılara (TFN) "
                        "dönüştürülür. Bu sıralama, ağırlık hesabının yapay belirsizlik şişirmesinden uzak, "
                        "gerçek gözlem değerleri üzerinde çalışmasını garanti ederken sıralama aşamasının "
                        "TFN tabanlı ölçüm gürültüsü dayanıklılığından yararlanmasını sağlar. "
                        "TFN dönüşümü şu şekilde tanımlanmıştır:"
                    ), lang_code=lang_code)
                _render_md_content(doc, fuzzy_common, lang_code)

        rank_content = _find_template_content(rm, template["ranking_methods"])
        if rank_content:
            _render_md_content(doc, rank_content, lang_code)
        else:
            try:
                from mcdm_engine import METHOD_PHILOSOPHY
                phil = METHOD_PHILOSOPHY.get(rm, {})
                phil_text = phil.get("academic", "")
                if phil_text:
                    _add_paragraph(doc, phil_text, lang_code=lang_code)
            except ImportError:
                pass

    # 2.4 Robustness
    _add_heading(doc, h["method_robust"], level=2, lang_code=lang_code)

    if lang == "EN":
        _add_paragraph(doc, (
            "A ranking result is only as credible as its sensitivity to the assumptions underlying it. "
            "For this reason, the reliability of the findings was verified through a multi-layered robustness protocol "
            "consisting of three independent tests."
        ), lang_code=lang_code)
        _add_paragraph(doc, (
            f"Local sensitivity: A perturbation of \u03b4 \u2208 {{-0.20, -0.10, +0.10, +0.20}} was applied "
            f"to each criterion weight; Spearman rank correlation \u03c1s was computed for each scenario."
        ), lang_code=lang_code, bold=True)
        if d["mc_n"] > 0:
            _add_paragraph(doc, (
                f"Monte Carlo simulation: Over N = {d['mc_n']} iterations, log-normal noise with "
                f"\u03c3 = {_fmt(d['mc_sigma'], lang)} was injected into the weight vector."
            ), lang_code=lang_code, bold=True)
        if d["comp_methods"]:
            _add_paragraph(doc, (
                f"Method comparison: Cross-validation was performed with the following methods: "
                f"{', '.join(d['comp_methods'])}."
            ), lang_code=lang_code, bold=True)
    else:
        _add_paragraph(doc, (
            "Bir sıralama sonucu, ancak dayandığı varsayımlara karşı duyarlılığı ölçüldüğünde güvenilir kabul edilebilir. "
            "Bu nedenle bulguların güvenilirliği, birbirinden bağımsız üç testten oluşan çok katmanlı bir "
            "sağlamlık protokolü ile doğrulanmıştır."
        ), lang_code=lang_code)
        _add_paragraph(doc, (
            f"Lokal duyarlılık: Her kriter ağırlığına \u03b4 \u2208 {{-0,20; -0,10; +0,10; +0,20}} "
            f"pertürbasyonu uygulanmış; her senaryoda Spearman sıra korelasyonu \u03c1s hesaplanmıştır."
        ), lang_code=lang_code, bold=True)
        if d["mc_n"] > 0:
            _add_paragraph(doc, (
                f"Monte Carlo simülasyonu: N = {d['mc_n']} iterasyon boyunca ağırlık vektörüne "
                f"\u03c3 = {_fmt(d['mc_sigma'], lang)} standart sapmalı log-normal gürültü enjekte edilmiştir."
            ), lang_code=lang_code, bold=True)
        if d["comp_methods"]:
            _add_paragraph(doc, (
                f"Yöntem karşılaştırması: {', '.join(d['comp_methods'])} yöntemleriyle çapraz doğrulama "
                f"yapılmıştır."
            ), lang_code=lang_code, bold=True)

    # Render robustness formulas from template
    robustness_content = template.get("robustness", "")
    if robustness_content:
        _render_md_content(doc, robustness_content, lang_code)


def _write_results_section(doc, d: Dict[str, Any], lang: str, lang_code: str) -> None:
    """Write the Results section."""
    h = _h(lang)
    f = _fmt

    _add_heading(doc, h["results"], level=1, lang_code=lang_code)

    # 3.1 Weights
    _add_heading(doc, h["results_weight"], level=2, lang_code=lang_code)

    if isinstance(d["weight_table"], pd.DataFrame) and not d["weight_table"].empty:
        if lang == "EN":
            _add_paragraph(doc, f"Table 1. Criteria weights determined by the {d['weight_method']} method.",
                           lang_code=lang_code, bold=True)
        else:
            _add_paragraph(doc, f"Tablo 1. {d['weight_method']} yöntemiyle belirlenen kriter ağırlıkları.",
                           lang_code=lang_code, bold=True)
        _add_table(doc, d["weight_table"], lang_code=lang_code)
        # Weight bar chart
        wb = _generate_weight_bar_bytes(d["weight_table"], lang)
        if wb:
            cap = "Şekil 2. Kriter Ağırlık Dağılımı" if lang != "EN" else "Figure 2. Criteria Weight Distribution"
            _add_figure(doc, wb, cap, lang_code)

    if d["top_criterion"]:
        if lang == "EN":
            _add_paragraph(doc, (
                f"The highest weight was assigned to {d['top_criterion']} "
                f"with w = {f(d['top_weight'], lang)}."
            ), lang_code=lang_code)
        else:
            _add_paragraph(doc, (
                f"En yüksek ağırlık w = {f(d['top_weight'], lang)} ile {d['top_criterion']} "
                f"kriterine aittir."
            ), lang_code=lang_code)

    # 3.2 Ranking
    if d["ranking_method"]:
        _add_heading(doc, h["results_rank"], level=2, lang_code=lang_code)

        if isinstance(d["ranking_table"], pd.DataFrame) and not d["ranking_table"].empty:
            if lang == "EN":
                _add_paragraph(doc, f"Table 2. Alternative ranking by {d['ranking_method']} method.",
                               lang_code=lang_code, bold=True)
            else:
                _add_paragraph(doc, f"Tablo 2. {d['ranking_method']} yöntemiyle alternatif sıralaması.",
                               lang_code=lang_code, bold=True)
            _add_table(doc, d["ranking_table"], lang_code=lang_code)
            # Ranking bar chart
            rb = _generate_ranking_bar_bytes(d["ranking_table"], lang)
            if rb:
                cap = "Şekil 3. Alternatif Sıralama Skorları" if lang != "EN" else "Figure 3. Alternative Ranking Scores"
                _add_figure(doc, rb, cap, lang_code)
            # Method-specific chart
            method_chart = _generate_method_specific_chart(d, lang)
            if method_chart:
                base = (d["ranking_method"] or "").replace("Fuzzy ", "")
                fig_num = 4
                cap = f"Şekil {fig_num}. {d['ranking_method']} Yöntem-Özel Görselleştirme" if lang != "EN" \
                    else f"Figure {fig_num}. {d['ranking_method']} Method-Specific Visualization"
                _add_figure(doc, method_chart, cap, lang_code)

        if d["top_alt"]:
            score_gap = abs(d["top_score"] - d["second_score"])
            if lang == "EN":
                _add_paragraph(doc, (
                    f"The analysis reveals the ranking as {d['top_alt']} in first place "
                    f"with a score of {f(d['top_score'], lang)}. "
                    f"The score gap between the first and second alternatives is {f(score_gap, lang)}, "
                    f"indicating {'a clear' if score_gap > 0.05 else 'a narrow'} decision margin."
                ), lang_code=lang_code)
            else:
                _add_paragraph(doc, (
                    f"Analiz sonucunda {d['top_alt']} alternatifi "
                    f"{f(d['top_score'], lang)} skoruyla birinci sıraya yerleşmiştir. "
                    f"Birinci ve ikinci alternatif arasındaki skor farkı {f(score_gap, lang)} olup "
                    f"{'net bir karar ayrışması' if score_gap > 0.05 else 'dar bir karar marjı'} "
                    f"söz konusudur."
                ), lang_code=lang_code)

    # 3.x Intermediate Calculation Tables (Appendix-style)
    _write_intermediate_tables(doc, d, lang, lang_code)

    # 3.3 Robustness
    _add_heading(doc, h["results_robust"], level=2, lang_code=lang_code)

    n_crit = d["n_crit"]
    n_scenarios = n_crit * 4

    if lang == "EN":
        if d["mc_n"] > 0 and d["stability"] > 0:
            _add_paragraph(doc, (
                f"Local sensitivity analysis produced Spearman \u03c1s values across all "
                f"{n_scenarios} scenarios ({n_crit} criteria \u00d7 4 perturbation levels)."
            ), lang_code=lang_code)
            _add_paragraph(doc, (
                f"Monte Carlo simulation confirmed the leader alternative's stability with "
                f"Stab({d['top_alt']}) = {f(d['stability'] * 100, lang, 1)}% first-place rate "
                f"over {d['mc_n']} iterations."
            ), lang_code=lang_code)
        if d["mean_spearman"] > 0:
            level = "high" if d["mean_spearman"] >= 0.85 else ("moderate" if d["mean_spearman"] >= 0.70 else "low")
            _add_paragraph(doc, (
                f"Method comparison yielded an average Spearman agreement of "
                f"\u03c1\u0305 = {f(d['mean_spearman'], lang)}, indicating {level} methodological consistency."
            ), lang_code=lang_code)
    else:
        if d["mc_n"] > 0 and d["stability"] > 0:
            _add_paragraph(doc, (
                f"Lokal duyarlılık analizi tüm {n_scenarios} senaryoda ({n_crit} kriter \u00d7 4 pertürbasyon düzeyi) "
                f"Spearman \u03c1s değerleri üretmiştir."
            ), lang_code=lang_code)
            _add_paragraph(doc, (
                f"Monte Carlo simülasyonu Stab({d['top_alt']}) = %{f(d['stability'] * 100, lang, 1)} "
                f"birincilik oranıyla lider alternatifin kararlılığını teyit etmiştir."
            ), lang_code=lang_code)
        if d["mean_spearman"] > 0:
            level = "yüksek" if d["mean_spearman"] >= 0.85 else ("orta" if d["mean_spearman"] >= 0.70 else "düşük")
            _add_paragraph(doc, (
                f"Yöntem karşılaştırmasında \u03c1\u0305 = {f(d['mean_spearman'], lang)} değeri "
                f"metodolojik tutarlılığın {level} düzeyde olduğunu göstermektedir."
            ), lang_code=lang_code)


def _write_discussion(doc, d: Dict[str, Any], lang: str, lang_code: str) -> None:
    """Write the Discussion section."""
    h = _h(lang)
    _add_heading(doc, h["discussion"], level=1, lang_code=lang_code)
    _add_paragraph(doc,
        "This section should compare findings with the literature, interpret the leader's superiority, and state study limitations." if lang == "EN"
        else "Bu bölümde bulgular literatürle karşılaştırılmalı, lider alternatifin üstünlüğü yorumlanmalı ve çalışmanın kısıtları belirtilmelidir.",
        lang_code=lang_code, italic=True, font_size=10, space_after=10)

    wm = d["weight_method"]
    rm = d["ranking_method"] or ""
    top = d["top_alt"] or "?"

    if lang == "EN":
        if d["mean_spearman"] >= 0.85:
            _add_paragraph(doc, f"Cross-method agreement (mean ρ = {_fmt(d['mean_spearman'], lang)}) indicates that the ranking is robust to method selection.", lang_code=lang_code)
        elif d["mean_spearman"] >= 0.70:
            _add_paragraph(doc, f"Moderate cross-method agreement (mean ρ = {_fmt(d['mean_spearman'], lang)}) suggests that the ranking is partially sensitive to the choice of method.", lang_code=lang_code)
        if d["top_alt"]:
            _add_paragraph(doc, f"The leading position of {top} is supported by the weight distribution and the scoring logic of {rm or wm}.", lang_code=lang_code)
    else:
        if d["mean_spearman"] >= 0.85:
            _add_paragraph(doc, f"Yöntemler arası uyum (ortalama ρ = {_fmt(d['mean_spearman'], lang)}) sıralamanın yöntem seçimine karşı dayanıklı olduğunu göstermektedir.", lang_code=lang_code)
        elif d["mean_spearman"] >= 0.70:
            _add_paragraph(doc, f"Orta düzey yöntem uyumu (ortalama ρ = {_fmt(d['mean_spearman'], lang)}) sıralamanın yöntem seçimine kısmen duyarlı olduğunu düşündürmektedir.", lang_code=lang_code)
        if d["top_alt"]:
            _add_paragraph(doc, f"{top} alternatifinin lider konumu, ağırlık dağılımı ve {rm or wm} yönteminin skorlama mantığıyla desteklenmektedir.", lang_code=lang_code)


def _write_conclusion(doc, d: Dict[str, Any], lang: str, lang_code: str) -> None:
    """Write the Conclusion section."""
    h = _h(lang)
    _add_heading(doc, h["conclusion"], level=1, lang_code=lang_code)
    _add_paragraph(doc,
        "This section should summarize main findings, highlight the study's contribution, and present future research recommendations." if lang == "EN"
        else "Bu bölümde ana bulgular özetlenmeli, çalışmanın katkısı vurgulanmalı ve gelecek araştırma önerileri sunulmalıdır.",
        lang_code=lang_code, italic=True, font_size=10, space_after=10)

    wm = d["weight_method"]
    rm = d["ranking_method"] or ""

    if lang == "EN":
        _add_paragraph(doc, (
            f"This study presented a multi-criteria evaluation framework using {wm}"
            + (f" and {rm}" if rm else "")
            + f" on {d['n_alt']} alternatives across {d['n_crit']} criteria."
        ), lang_code=lang_code)
        if d["top_alt"]:
            stab = d.get("stability", 0)
            if stab >= 0.80:
                _add_paragraph(doc, (
                    f"Alternative {d['top_alt']} ranked first consistently across all robustness tests "
                    f"(MC stability: {_fmt(stab * 100, lang, 1)}%), establishing a reliable reference "
                    f"point for decision-makers."
                ), lang_code=lang_code)
            elif stab >= 0.60:
                _add_paragraph(doc, (
                    f"Alternative {d['top_alt']} showed relative dominance with a moderate stability rate "
                    f"of {_fmt(stab * 100, lang, 1)}%. While it ranked first in the majority of robustness "
                    f"scenarios, the result should be interpreted alongside the sensitivity findings."
                ), lang_code=lang_code)
            elif stab > 0:
                _add_paragraph(doc, (
                    f"Alternative {d['top_alt']} emerged as the top-ranked option in the base analysis; however, "
                    f"the Monte Carlo stability rate of {_fmt(stab * 100, lang, 1)}% indicates that the ranking "
                    f"is sensitive to weight perturbations. Decision-makers should consider this uncertainty "
                    f"and may benefit from examining the full stability distribution."
                ), lang_code=lang_code)
            else:
                _add_paragraph(doc, (
                    f"Alternative {d['top_alt']} achieved the highest score in the base analysis."
                ), lang_code=lang_code)
    else:
        _add_paragraph(doc, (
            f"Bu çalışmada {wm}"
            + (f" ve {rm}" if rm else "")
            + f" entegrasyonuyla {d['n_alt']} alternatif ve {d['n_crit']} kriter üzerinde çok kriterli değerlendirme gerçekleştirilmiştir."
        ), lang_code=lang_code)
        if d["top_alt"]:
            stab = d.get("stability", 0)
            if stab >= 0.80:
                _add_paragraph(doc, (
                    f"{d['top_alt']} alternatifi tüm sağlamlık testlerinde tutarlı biçimde birinci sırada "
                    f"yer almış (MC kararlılık: %{_fmt(stab * 100, lang, 1)}) ve karar vericiler için "
                    f"güvenilir bir referans noktası oluşturmuştur."
                ), lang_code=lang_code)
            elif stab >= 0.60:
                _add_paragraph(doc, (
                    f"{d['top_alt']} alternatifi %{_fmt(stab * 100, lang, 1)} kararlılık oranıyla göreli "
                    f"bir üstünlük sergilemiştir. Sağlamlık senaryolarının çoğunluğunda birinci sırada yer "
                    f"almakla birlikte, sonuç duyarlılık bulguları eşliğinde değerlendirilmelidir."
                ), lang_code=lang_code)
            elif stab > 0:
                _add_paragraph(doc, (
                    f"{d['top_alt']} alternatifi temel analizde en yüksek skoru elde etmiştir; ancak "
                    f"Monte Carlo kararlılık oranı (%{_fmt(stab * 100, lang, 1)}) sıralamanın ağırlık "
                    f"pertürbasyonlarına duyarlı olduğuna işaret etmektedir. Karar vericilerin bu "
                    f"belirsizliği göz önünde bulundurması ve tam kararlılık dağılımını incelemesi önerilir."
                ), lang_code=lang_code)
            else:
                _add_paragraph(doc, (
                    f"{d['top_alt']} alternatifi temel analizde en yüksek skoru elde etmiştir."
                ), lang_code=lang_code)
        pass  # Gelecek çalışma önerisi kullanıcıya bırakılmıştır


def _write_references(doc, d: Dict[str, Any], lang: str, lang_code: str) -> None:
    """Write the References section."""
    h = _h(lang)
    _add_heading(doc, h["references"], level=1, lang_code=lang_code)

    try:
        from mcdm_engine import _report_references
        refs = _report_references(d["weight_method"], d["ranking_method"])
    except ImportError:
        refs = []

    if not refs:
        # Fallback: use template references
        template = _load_template()
        imrad = template.get("imrad_template", "")
        ref_match = re.search(r"###\s+KAYNAKLAR\s*\n(.*)", imrad, re.DOTALL)
        if ref_match:
            ref_text = ref_match.group(1)
            for line in ref_text.strip().split("\n"):
                line = line.strip().lstrip("- ").strip()
                if line:
                    refs.append(line)

    for ref in refs:
        p = doc.add_paragraph()
        run = p.add_run(ref)
        _set_run_style(run, lang_code=lang_code, font_size=11)
        p.paragraph_format.space_after = Pt(4)
        # Hanging indent
        p.paragraph_format.first_line_indent = Inches(-0.5)
        p.paragraph_format.left_indent = Inches(0.5)


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------


def generate_imrad_docx(
    result: Dict[str, Any],
    selected_data: pd.DataFrame,
    lang: str = "TR",
) -> bytes | None:
    """Generate IMRAD article as DOCX via pandoc with native Word equations.

    APA/SSCI style via reference.docx template.
    Includes flowchart, weight/ranking charts, method-specific visuals.
    """
    import subprocess, tempfile

    d = _extract_analysis_data(result, selected_data)
    f = _fmt
    template = _load_template()
    h = _h(lang)
    wm = d["weight_method"] or "?"
    rm = d["ranking_method"] or ""
    is_fuzzy = d.get("is_fuzzy", False)

    md = []
    _tmp_files = []  # track temp files for cleanup

    def _save_png(png_bytes):
        import tempfile as _tm
        p = _tm.mktemp(suffix=".png")
        with open(p, "wb") as _f:
            _f.write(png_bytes)
        _tmp_files.append(p)
        return p

    # ── TITLE ──
    title = _build_title(d, lang)
    md.append(f"# {title}\n")

    # ── ABSTRACT ──
    md.append(f"## {h['abstract']}\n")
    md.append(_build_abstract(d, lang) + "\n")
    kw_label = "Keywords" if lang == "EN" else "Anahtar Kelimeler"
    kw_val = ("MCDM" if lang == "EN" else "ÇKKV") + f", {wm}" + (f", {rm}" if rm else "") + ", " + ("Robustness Analysis" if lang == "EN" else "Sağlamlık Analizi")
    md.append(f"***{kw_label}:*** *{kw_val}*\n")

    # ── 1. INTRODUCTION ──
    md.append(f"## {h['introduction']}\n")
    guide_intro = "This section should present the research problem, literature review, research gap, and contribution." if lang == "EN" else "Bu bölümde araştırma problemi, literatür taraması, araştırma boşluğu ve çalışmanın katkısı sunulmalıdır."
    md.append(f"*{guide_intro}*\n")
    if lang == "EN":
        md.append(f"This study applies {wm} for criteria weighting" + (f" and {rm} for ranking" if rm else "") + f" on {d['n_alt']} alternatives and {d['n_crit']} criteria. The methodology integrates objective weight determination with systematic ranking and multi-layered robustness verification.\n")
    else:
        md.append(f"Bu çalışmada {d['n_alt']} alternatif ve {d['n_crit']} kriter üzerinde {wm} ile ağırlıklandırma" + (f" ve {rm} ile sıralama" if rm else "") + " uygulanmaktadır. Önerilen metodoloji, nesnel ağırlık belirlemeyi sistematik sıralama ve çok katmanlı sağlamlık doğrulamasıyla bütünleştirmektedir.\n")

    # ── 2. METHOD ──
    md.append(f"## {h['method']}\n")

    # Overview
    if lang == "EN":
        md.append("The proposed methodology consists of three integrated layers: (1) objective criteria weighting from the raw decision matrix, " + ("(2) fuzzification of the decision matrix via TFN, " if is_fuzzy else "") + (f"({'3' if is_fuzzy else '2'}) systematic ranking, and ({'4' if is_fuzzy else '3'}) multi-layered robustness validation.\n"))
    else:
        md.append("Önerilen metodoloji birbiriyle bütünleşik üç katmandan oluşmaktadır: (1) ham karar matrisinden nesnel kriter ağırlıklandırma, " + ("(2) karar matrisinin TFN ile bulanıklaştırılması, " if is_fuzzy else "") + (f"({'3' if is_fuzzy else '2'}) sistematik sıralama ve ({'4' if is_fuzzy else '3'}) çok katmanlı sağlamlık doğrulaması.\n"))

    # Flowchart
    try:
        fc = _generate_flowchart_bytes(d, lang)
    except Exception:
        fc = None
    if fc:
        fc_path = _save_png(fc)
        cap = "Şekil 1. Önerilen Yöntemin Akış Şeması" if lang != "EN" else "Figure 1. Proposed Methodology Flowchart"
        md.append(f"![{cap}]({fc_path}){{width=85%}}\n")
        if lang != "EN":
            md.append("Şekil 1, çalışmada uygulanan metodolojik akışı özetlemektedir. Her aşama bir önceki aşamanın çıktısını girdi olarak kullanmakta olup süreç bütünsel bir karar destek çerçevesi oluşturmaktadır.\n")
        else:
            md.append("Figure 1 summarizes the methodological flow applied in this study. Each stage uses the output of the previous stage as input, forming a holistic decision support framework.\n")

    # 2.1 Decision Matrix
    md.append(f"### {h['method_dm']}\n")
    ben = ", ".join(d["benefit"]) if d["benefit"] else "-"
    cos = ", ".join(d["cost"]) if d["cost"] else "-"
    if lang == "EN":
        md.append(f"The decision matrix consists of {d['n_alt']} alternatives across {d['n_crit']} criteria. Benefit: {ben}. Cost: {cos}.\n")
    else:
        md.append(f"Karar matrisi {d['n_alt']} alternatif ve {d['n_crit']} kriterden oluşmaktadır. Fayda yönlü (büyük olan tercih edilir): {ben}. Maliyet yönlü (küçük olan tercih edilir): {cos}.\n")

    elim = d.get("eliminated") or []
    if elim:
        en = sorted(set(e.get("alternative", "?") for e in elim))
        if lang == "EN":
            md.append(f"{len(en)} alternative(s) ({', '.join(en)}) were eliminated by pre-screening.\n")
        else:
            md.append(f"{len(en)} alternatif ({', '.join(en)}) eşik tabanlı ön eleme prosedürü ile analiz dışı bırakılmıştır.\n")

    # 2.2 Weight — Philosophy + Justification + Formulas
    md.append(f"### {h['method_weight']}\n")
    try:
        from mcdm_engine import METHOD_PHILOSOPHY
        phil = METHOD_PHILOSOPHY.get(wm, {})
        if phil.get("academic"):
            md.append(phil["academic"] + "\n")
    except ImportError:
        pass
    wj = (_WEIGHT_JUSTIFICATION_EN if lang == "EN" else _WEIGHT_JUSTIFICATION_TR).get(wm, "")
    if wj:
        md.append(wj + "\n")
    wc = _find_template_content(wm, template["weight_methods"])
    if wc:
        md.append(wc + "\n")

    # 2.3 Ranking
    if rm:
        md.append(f"### {h['method_rank']}\n")
        try:
            rphil = METHOD_PHILOSOPHY.get(rm, {}) or METHOD_PHILOSOPHY.get(rm.replace("Fuzzy ", ""), {})
            if rphil.get("academic"):
                md.append(rphil["academic"] + "\n")
        except Exception:
            pass
        base_rm = rm.replace("Fuzzy ", "")
        rj = (_RANKING_JUSTIFICATION_EN if lang == "EN" else _RANKING_JUSTIFICATION_TR).get(rm, "") or (_RANKING_JUSTIFICATION_EN if lang == "EN" else _RANKING_JUSTIFICATION_TR).get(base_rm, "")
        if rj:
            md.append(rj + "\n")
        if is_fuzzy:
            fc_common = template["ranking_methods"].get("_fuzzy_common", "")
            if fc_common:
                spread_txt = f"s = {f(d['fuzzy_spread'], lang)}"
                if lang == "EN":
                    md.append(f"The decision matrix was converted to TFN with {spread_txt}. Fuzzification is applied after weighting to preserve weight objectivity.\n")
                else:
                    md.append(f"Karar matrisi {spread_txt} spread ile TFN'ye dönüştürülmüştür. Bulanıklaştırma adımı, ağırlık vektörünün nesnelliğini korumak amacıyla bilinçli olarak ağırlık hesabından sonraya konumlandırılmaktadır.\n")
                md.append(fc_common + "\n")
        rc = _find_template_content(rm, template["ranking_methods"])
        if rc:
            md.append(rc + "\n")

    # 2.4 Robustness
    md.append(f"### {h['method_robust']}\n")
    if lang == "EN":
        rob = "Reliability was verified through: (1) local sensitivity (±10%/±20% per criterion, " + str(d['n_crit']*4) + " scenarios)"
        if d["mc_n"] > 0:
            rob += f", (2) Monte Carlo (N={d['mc_n']}, σ={f(d['mc_sigma'], lang)})"
        if d["comp_methods"]:
            rob += f", (3) cross-method comparison ({', '.join(d['comp_methods'][:3])})"
        md.append(rob + ".\n")
    else:
        rob = "Sonuçların güvenilirliği aşağıdaki çok katmanlı sağlamlık protokolü ile doğrulanmaktadır: (1) lokal duyarlılık (kriter başına ±%10/±%20, " + str(d['n_crit']*4) + " senaryo)"
        if d["mc_n"] > 0:
            rob += f", (2) Monte Carlo (N={d['mc_n']}, σ={f(d['mc_sigma'], lang)})"
        if d["comp_methods"]:
            rob += f", (3) çapraz yöntem karşılaştırması ({', '.join(d['comp_methods'][:3])})"
        md.append(rob + ".\n")

    # ── 3. RESULTS ──
    md.append(f"## {h['results']}\n")

    # 3.1 Weights
    md.append(f"### {h['results_weight']}\n")
    if isinstance(d["weight_table"], pd.DataFrame) and not d["weight_table"].empty:
        wt = d["weight_table"].copy()
        for c in wt.select_dtypes(include=["float", "number"]).columns:
            wt[c] = wt[c].apply(lambda x: f"{x:.3f}" if pd.notna(x) else "")
        tbl_cap = f"**{'Table' if lang == 'EN' else 'Tablo'} 1.** {wm} {'weights' if lang == 'EN' else 'ağırlıkları'}."
        md.append(tbl_cap + "\n")
        md.append(wt.to_markdown(index=False) + "\n")
    if lang != "EN":
        md.append(f"Tablo 1, {wm} yöntemiyle hesaplanan kriter ağırlıklarını sunmaktadır. Ağırlık değerleri, her kriterin karar sürecine olan göreli katkısını yansıtmaktadır.\n")
    else:
        md.append(f"Table 1 presents the criteria weights calculated by {wm}. Weight values reflect the relative contribution of each criterion to the decision process.\n")
    if d["top_criterion"]:
        if lang == "EN":
            md.append(f"The highest weight belongs to **{d['top_criterion']}** (w = {f(d['top_weight'], lang)}), indicating it carries the greatest discriminatory power among alternatives.\n")
        else:
            md.append(f"En yüksek ağırlık **{d['top_criterion']}** kriterine aittir (w = {f(d['top_weight'], lang)}). Bu bulgu, söz konusu kriterin mevcut veri setinde alternatifleri birbirinden en güçlü biçimde ayırt eden bilgi yapısına sahip olduğuna işaret etmektedir.\n")

    try:
        wb = _generate_weight_bar_bytes(d["weight_table"], lang)
    except Exception:
        wb = None
    if wb:
        wb_path = _save_png(wb)
        cap_w = ("Şekil 2. Kriter Ağırlık Dağılımı" if lang != "EN" else "Figure 2. Criteria Weight Distribution")
        md.append(f"![{cap_w}]({wb_path}){{width=70%}}\n")
        if lang != "EN":
            md.append("Şekil 2 incelendiğinde, kriter ağırlıklarının dağılımı görsel olarak ortaya konulmaktadır. Ağırlık yoğunlaşmasının belirli kriterlerde toplanması, bu kriterlerin karar sürecindeki belirleyici rolüne işaret etmektedir.\n")
        else:
            md.append("Figure 2 visually presents the distribution of criteria weights. Weight concentration on specific criteria indicates their decisive role in the decision process.\n")

    # 3.2 Ranking
    if rm:
        md.append(f"### {h['results_rank']}\n")
        if isinstance(d["ranking_table"], pd.DataFrame) and not d["ranking_table"].empty:
            rt = d["ranking_table"].copy()
            for c in rt.select_dtypes(include=["float", "number"]).columns:
                rt[c] = rt[c].apply(lambda x: f"{x:.3f}" if pd.notna(x) else "")
            md.append(f"**{'Table' if lang == 'EN' else 'Tablo'} 2.** {rm} {'ranking' if lang == 'EN' else 'sıralaması'}.\n")
            md.append(rt.to_markdown(index=False) + "\n")
        if lang != "EN":
            md.append(f"Tablo 2, {rm} yöntemiyle elde edilen alternatif sıralamasını göstermektedir. Skor değerleri, her alternatifin tüm kriterler bazındaki bütüncül performansını temsil etmektedir.\n")
        else:
            md.append(f"Table 2 shows the alternative ranking obtained by {rm}. Score values represent the holistic performance of each alternative across all criteria.\n")
        if d["top_alt"]:
            gap = abs(d["top_score"] - d["second_score"])
            if lang == "EN":
                gw = "a clear margin" if gap > 0.05 else "a narrow margin"
                md.append(f"**{d['top_alt']}** ranks first (score: {f(d['top_score'], lang)}). Gap to second: {f(gap, lang)} ({gw}).\n")
            else:
                gw = "net bir fark" if gap > 0.05 else "dar bir marj"
                md.append(f"**{d['top_alt']}** birinci sıradadır (skor: {f(d['top_score'], lang)}). İkinciyle fark: {f(gap, lang)} ({gw}).\n")

        try:
            rb = _generate_ranking_bar_bytes(d["ranking_table"], lang)
        except Exception:
            rb = None
        if rb:
            rb_path = _save_png(rb)
            cap_r = "Şekil 3. Alternatif Skorları" if lang != "EN" else "Figure 3. Alternative Scores"
            md.append(f"![{cap_r}]({rb_path}){{width=70%}}\n")
        if lang != "EN":
            md.append("Şekil 3, alternatiflerin nihai skorlarını karşılaştırmalı olarak sunmaktadır. Lider alternatifin diğerlerinden ayrışma düzeyi, karar güvenilirliğinin görsel bir göstergesi olarak değerlendirilebilmektedir.\n")
        else:
            md.append("Figure 3 presents the final scores of alternatives comparatively. The degree of separation of the leading alternative serves as a visual indicator of decision reliability.\n")

        try:
            mc = _generate_method_specific_chart(d, lang)
        except Exception:
            mc = None
        if mc:
            mc_path = _save_png(mc)
            cap_m = f"Şekil 4. {rm} Görselleştirme" if lang != "EN" else f"Figure 4. {rm} Visualization"
            md.append(f"![{cap_m}]({mc_path}){{width=70%}}\n")
        if lang != "EN":
            md.append(f"Şekil 4, {rm} yönteminin kendine özgü analitik çıktısını görselleştirmektedir. Bu görsel, yöntemin karar mantığını ve alternatiflerin yöntem-özel performans profilini ortaya koymaktadır.\n")
        else:
            md.append(f"Figure 4 visualizes the method-specific analytical output of {rm}. This visual reveals the decision logic and the method-specific performance profile of alternatives.\n")

    # 3.3 Robustness findings
    md.append(f"### {h['results_robust']}\n")
    stab = d.get("stability", 0)
    if stab > 0 and d["top_alt"]:
        if lang == "EN":
            lev = "high" if stab >= 0.80 else ("moderate" if stab >= 0.60 else "low")
            md.append(f"MC stability: **{d['top_alt']}** first in **{f(stab*100, lang, 1)}%** of iterations (**{lev}**).\n")
        else:
            lev = "yüksek" if stab >= 0.80 else ("orta" if stab >= 0.60 else "düşük")
            md.append(f"MC kararlılık: **{d['top_alt']}** iterasyonların **%{f(stab*100, lang, 1)}**'inde birinci (**{lev}**).\n")
    if d["mean_spearman"] > 0:
        lev_en = "high" if d["mean_spearman"] >= 0.85 else ("moderate" if d["mean_spearman"] >= 0.70 else "low")
        lev_tr = "yüksek" if d["mean_spearman"] >= 0.85 else ("orta" if d["mean_spearman"] >= 0.70 else "düşük")
        if lang == "EN":
            md.append(f"Cross-method Spearman: **ρ̄ = {f(d['mean_spearman'], lang)}** ({lev_en}).\n")
        else:
            md.append(f"Çapraz yöntem Spearman: **ρ̄ = {f(d['mean_spearman'], lang)}** ({lev_tr}).\n")

    # ── 4. DISCUSSION ──
    md.append(f"## {h['discussion']}\n")
    guide_disc = "Compare findings with literature, interpret the leader's advantage, discuss method strengths and limitations." if lang == "EN" else "Bulgular literatürle karşılaştırılmalı, liderin avantajı yorumlanmalı, yöntemin güçlü ve zayıf yönleri tartışılmalıdır."
    md.append(f"*{guide_disc}*\n")
    if d["top_alt"]:
        if lang == "EN":
            md.append(f"The leading position of **{d['top_alt']}** is supported by the weight distribution ({wm}) and scoring logic ({rm or wm}).\n")
        else:
            md.append(f"**{d['top_alt']}** alternatifinin lider konumu {wm} ağırlık dağılımı ve {rm or wm} skorlama mantığıyla desteklenmektedir.\n")

    # ── 5. CONCLUSION ──
    md.append(f"## {h['conclusion']}\n")
    guide_conc = "Summarize key findings, state contribution, suggest future work." if lang == "EN" else "Ana bulguları özetleyin, katkıyı belirtin, gelecek çalışma önerin."
    md.append(f"*{guide_conc}*\n")
    if lang == "EN":
        md.append(f"This study evaluated {d['n_alt']} alternatives across {d['n_crit']} criteria using {wm}" + (f" and {rm}" if rm else "") + ".")
    else:
        md.append(f"Bu çalışmada {wm}" + (f" ve {rm}" if rm else "") + f" ile {d['n_alt']} alternatif ve {d['n_crit']} kriter değerlendirilmiştir.")
    if d["top_alt"] and stab >= 0.60:
        if lang == "EN":
            md.append(f" **{d['top_alt']}** emerged as the leader ({f(stab*100, lang, 1)}% stability).\n")
        else:
            md.append(f" **{d['top_alt']}** lider olarak belirlenmiştir (%{f(stab*100, lang, 1)} kararlılık).\n")
    else:
        md.append("\n")

    # ── REFERENCES ──
    md.append(f"## {h['references']}\n")
    try:
        from mcdm_engine import _report_references
        for ref in _report_references(d["weight_method"], d["ranking_method"]):
            md.append(f"- {ref}\n")
    except ImportError:
        pass

    # ── DISCLAIMER ──
    md.append("\n---\n")
    disc = "This output is for informational and academic evaluation purposes only. Final decision responsibility rests with the decision-maker." if lang == "EN" else "Bu çıktı yalnızca bilgilendirme ve akademik değerlendirme amacıyla üretilmiştir. Nihai karar sorumluluğu karar vericiye aittir."
    md.append(f"*{disc}*\n")
    md.append("\n*MCDM Karar Destek Sistemi — Prof. Dr. Ömer Faruk Rençber*\n")

    # ── CONVERT MD → HTML ──
    md_text = "\n".join(md)
    import base64 as _b64

    def _md_to_html_body(md: str) -> str:
        """Convert markdown to HTML with basic formatting."""
        lines = md.split("\n")
        html_parts = []
        in_table = False
        for line in lines:
            s = line.strip()
            if not s:
                if in_table:
                    html_parts.append("</table>")
                    in_table = False
                html_parts.append("<p>&nbsp;</p>")
                continue
            if s.startswith("# "):
                html_parts.append(f"<h1>{s[2:]}</h1>")
            elif s.startswith("### "):
                html_parts.append(f"<h3>{s[4:]}</h3>")
            elif s.startswith("## "):
                html_parts.append(f"<h2>{s[3:]}</h2>")
            elif s.startswith("!["):
                m = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", s)
                if m:
                    cap, img_path = m.group(1), m.group(2)
                    try:
                        with open(img_path, "rb") as _imgf:
                            img_b64 = _b64.b64encode(_imgf.read()).decode()
                        html_parts.append(f'<div class="fig"><img src="data:image/png;base64,{img_b64}" style="max-width:85%"><p class="cap">{cap}</p></div>')
                    except Exception:
                        pass
            elif s.startswith("$$"):
                formula = s.strip("$").strip()
                if formula:
                    html_parts.append(f'<div class="math">$${formula}$$</div>')
            elif s.startswith("|"):
                if not in_table:
                    html_parts.append('<table class="dt">')
                    in_table = True
                if s.startswith("|--") or s.startswith("| --") or s.replace(" ","").replace("-","").replace("|","") == "":
                    continue
                cells = [c.strip() for c in s.split("|")[1:-1]]
                if not any(c for c in cells):
                    continue
                tag = "th" if not any(c.replace("*","").strip() for c in cells if c.replace("*","").strip() and c.replace("*","").strip()[0].isdigit()) and html_parts[-1] == '<table class="dt">' else "td"
                row_html = "".join(f"<{tag}>{c.replace('**','')}</{tag}>" for c in cells)
                html_parts.append(f"<tr>{row_html}</tr>")
            elif s.startswith("---"):
                if in_table:
                    html_parts.append("</table>")
                    in_table = False
                html_parts.append("<hr>")
            elif s.startswith("- "):
                html_parts.append(f"<p style='margin-left:1.5em;'>• {s[2:]}</p>")
            elif s.startswith("***") and s.endswith("***"):
                html_parts.append(f"<p><b><i>{s[3:-3]}</i></b></p>")
            elif s.startswith("*") and s.endswith("*") and not s.startswith("**"):
                html_parts.append(f"<p class='guide'><i>{s.strip('*')}</i></p>")
            elif s.startswith("**") and s.endswith("**"):
                html_parts.append(f"<p><b>{s.strip('*')}</b></p>")
            else:
                # Inline bold/italic
                processed = re.sub(r'\*\*([^*]+)\*\*', r'<b>\1</b>', s)
                processed = re.sub(r'\*([^*]+)\*', r'<i>\1</i>', processed)
                # Inline math
                processed = re.sub(r'\$([^$]+)\$', r'\\(\1\\)', processed)
                html_parts.append(f"<p>{processed}</p>")
        if in_table:
            html_parts.append("</table>")
        return "\n".join(html_parts)

    body = _md_to_html_body(md_text)

    html = f"""<!DOCTYPE html>
<html lang="{'en' if lang == 'EN' else 'tr'}">
<head>
<meta charset="utf-8">
<title>{_build_title(d, lang)}</title>
<script>
MathJax = {{tex: {{inlineMath: [['\\\\(','\\\\)']]}}, displayMath: [['$$','$$']]}};
</script>
<script src="https://cdn.jsdelivr.net/npm/mathjax@3/es5/tex-mml-chtml.js" async></script>
<style>
@page {{ size: A4; margin: 2.5cm; }}
body {{ font-family: 'Times New Roman', serif; font-size: 12pt; color: #000; line-height: 1.6; text-align: justify; max-width: 800px; margin: 0 auto; padding: 2em; }}
h1 {{ font-size: 15pt; font-weight: bold; text-align: center; margin: 1em 0 0.5em; color: #000; }}
h2 {{ font-size: 13pt; font-weight: bold; margin: 1em 0 0.4em; color: #000; border-bottom: 1px solid #ddd; padding-bottom: 0.2em; }}
h3 {{ font-size: 12pt; font-weight: bold; margin: 0.8em 0 0.3em; color: #000; }}
p {{ margin: 0.4em 0; }}
.guide {{ font-style: italic; font-size: 10pt; color: #666; margin: 0.3em 0 0.8em; }}
.math {{ text-align: center; margin: 0.8em 0; font-size: 11pt; }}
.fig {{ text-align: center; margin: 1em 0; }}
.fig img {{ max-width: 85%; border: 1px solid #eee; }}
.cap {{ font-style: italic; font-size: 10pt; color: #333; margin-top: 0.3em; }}
.dt {{ width: 100%; border-collapse: collapse; font-size: 10pt; margin: 0.8em 0; }}
.dt th {{ background: #f0f0f0; border: 1px solid #ccc; padding: 5px 10px; text-align: left; font-weight: bold; }}
.dt td {{ border: 1px solid #ddd; padding: 4px 10px; }}
hr {{ border: none; border-top: 1px solid #ccc; margin: 1.5em 0; }}
.footer {{ font-size: 9pt; font-style: italic; color: #666; margin-top: 2em; border-top: 1px solid #ccc; padding-top: 0.5em; }}
</style>
</head>
<body>
{body}
</body>
</html>"""

    # Cleanup temp image files
    for tp in _tmp_files:
        try:
            import os
            os.unlink(tp)
        except Exception:
            pass

    return html.encode("utf-8")



def generate_panel_imrad_docx(
    panel_results: Dict[str, Dict[str, Any]],
    lang: str = "TR",
) -> bytes | None:
    """Generate an IMRAD-format article for panel (multi-period) analysis.

    Parameters
    ----------
    panel_results : dict
        Mapping of period label → analysis result dict.
    lang : str
        "TR" or "EN".

    Returns
    -------
    bytes or None
    """
    if not _ensure_docx():
        return None
    if not panel_results:
        return None

    lang_code = "en-US" if lang == "EN" else "tr-TR"
    h = _h(lang)

    # Use the first period's data for common info
    first_key = next(iter(panel_results))
    first_result = panel_results[first_key]
    first_data = first_result.get("selected_data", pd.DataFrame())
    d = _extract_analysis_data(first_result, first_data)

    doc = Document()
    _configure_doc(doc)

    # TITLE
    n_periods = len(panel_results)
    if lang == "EN":
        title = f"Panel Multi-Criteria Decision Analysis ({n_periods} Periods): {d['weight_method']}"
        if d["ranking_method"]:
            title += f" and {d['ranking_method']}"
    else:
        title = f"Panel Çok Kriterli Karar Analizi ({n_periods} Dönem): {d['weight_method']}"
        if d["ranking_method"]:
            title += f" ve {d['ranking_method']}"

    p_title = doc.add_paragraph()
    run = p_title.add_run(title)
    _set_run_style(run, lang_code=lang_code, font_size=14, bold=True)
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtitle with placeholders
    # Subtitle omitted — title is sufficient

    doc.add_paragraph("")

    # ABSTRACT
    _add_heading(doc, h["abstract"], level=1, lang_code=lang_code)
    abstract = _build_abstract(d, lang)
    if lang == "EN":
        abstract += f" The analysis covers {n_periods} time periods."
    else:
        abstract += f" Analiz {n_periods} zaman dönemini kapsamaktadır."
    _add_paragraph(doc, abstract, lang_code=lang_code)

    doc.add_paragraph("")

    # INTRODUCTION (common)
    _write_introduction(doc, d, lang, lang_code)

    # METHOD (common — uses first period's data as reference)
    _write_method_section(doc, d, first_data, lang, lang_code)

    # RESULTS
    _add_heading(doc, h["results"], level=1, lang_code=lang_code)

    period_keys = list(panel_results.keys())
    all_period_data: Dict[str, Dict[str, Any]] = {}
    for pk in period_keys:
        pr = panel_results[pk]
        pd_sel = pr.get("selected_data", pd.DataFrame())
        all_period_data[pk] = _extract_analysis_data(pr, pd_sel)

    # --- 3.1 Reference period: full results for the first period ---
    ref_key = period_keys[0]
    ref_d = all_period_data[ref_key]

    if lang == "EN":
        _add_heading(doc, f"3.1 Reference Period Results: {ref_key}", level=2, lang_code=lang_code)
        _add_paragraph(doc, (
            f"To illustrate the computational steps in detail, the results for the {ref_key} period are "
            f"presented as the reference case. The same methodology was applied identically to all "
            f"{n_periods} periods; therefore, the mathematical procedure is documented once below and "
            f"the remaining periods are summarized in the cross-period comparison tables that follow."
        ), lang_code=lang_code)
    else:
        _add_heading(doc, f"3.1 Referans Dönem Sonuçları: {ref_key}", level=2, lang_code=lang_code)
        _add_paragraph(doc, (
            f"Hesaplama adımlarını ayrıntılı göstermek amacıyla {ref_key} dönemi referans olarak sunulmuştur. "
            f"Aynı metodoloji {n_periods} dönemin tamamına birebir uygulanmıştır; dolayısıyla matematiksel "
            f"prosedür aşağıda bir kez belgelenmiş, kalan dönemler ilerleyen karşılaştırma tablolarında "
            f"özetlenmiştir."
        ), lang_code=lang_code)

    # Weight table for reference period
    if isinstance(ref_d["weight_table"], pd.DataFrame) and not ref_d["weight_table"].empty:
        tbl_label = f"Tablo 1. {ref_key} dönemi kriter ağırlıkları ({ref_d['weight_method']})." if lang != "EN" \
            else f"Table 1. Criteria weights for {ref_key} period ({ref_d['weight_method']})."
        _add_paragraph(doc, tbl_label, lang_code=lang_code, bold=True)
        _add_table(doc, ref_d["weight_table"], lang_code=lang_code)

    # Ranking table for reference period
    if ref_d["ranking_method"] and isinstance(ref_d["ranking_table"], pd.DataFrame) and not ref_d["ranking_table"].empty:
        tbl_label = f"Tablo 2. {ref_key} dönemi sıralama sonuçları ({ref_d['ranking_method']})." if lang != "EN" \
            else f"Table 2. Ranking results for {ref_key} period ({ref_d['ranking_method']})."
        _add_paragraph(doc, tbl_label, lang_code=lang_code, bold=True)
        _add_table(doc, ref_d["ranking_table"], lang_code=lang_code)

    # Robustness for reference period
    if ref_d["stability"] > 0 or ref_d["mean_spearman"] > 0:
        if lang == "EN":
            _add_paragraph(doc, "Robustness findings for the reference period:", lang_code=lang_code, bold=True)
        else:
            _add_paragraph(doc, "Referans dönem sağlamlık bulguları:", lang_code=lang_code, bold=True)
        _write_results_section_robustness(doc, ref_d, lang, lang_code)

    # --- 3.2 Cross-period comparison ---
    if lang == "EN":
        _add_heading(doc, "3.2 Cross-Period Comparison", level=2, lang_code=lang_code)
    else:
        _add_heading(doc, "3.2 Dönemler Arası Karşılaştırma", level=2, lang_code=lang_code)

    # Build cross-period weight comparison table
    _write_panel_weight_comparison(doc, period_keys, all_period_data, lang, lang_code)

    # Build cross-period ranking comparison table
    if d["ranking_method"]:
        _write_panel_ranking_comparison(doc, period_keys, all_period_data, lang, lang_code)

    # Cross-period leader summary
    _write_panel_leader_summary(doc, period_keys, all_period_data, lang, lang_code)

    # Cross-period comparison chart
    chart_bytes = _generate_panel_comparison_chart(period_keys, all_period_data, lang)
    if chart_bytes:
        cap = "Şekil 2. Dönemler Arası Lider Alternatif Skor Değişimi" if lang != "EN" \
            else "Figure 2. Leader Alternative Score Trends Across Periods"
        _add_figure(doc, chart_bytes, cap, lang_code)

    # DISCUSSION & CONCLUSION
    _write_discussion(doc, d, lang, lang_code)
    _write_conclusion(doc, d, lang, lang_code)

    # REFERENCES (consolidated)
    _write_references(doc, d, lang, lang_code)

    # Signature
    doc.add_paragraph("")
    # Disclaimer + Signature
    doc.add_paragraph("")
    _add_paragraph(doc,
        "This output is generated for informational and academic evaluation purposes only. The results presented are analytical recommendations based on the selected method and data; final decision responsibility rests with the decision-maker." if lang == "EN" else "Bu çıktı yalnızca bilgilendirme ve akademik değerlendirme amacıyla üretilmiştir. Sunulan sonuçlar, seçilen yöntem ve veriye dayalı analitik öneriler niteliğindedir; nihai karar sorumluluğu karar vericiye aittir.",
        lang_code=lang_code, italic=True, font_size=9, space_after=4)
    _add_paragraph(doc, "MCDM Karar Destek Sistemi — Prof. Dr. Ömer Faruk Rençber",
        lang_code=lang_code, italic=True, font_size=9,
        alignment=WD_ALIGN_PARAGRAPH.RIGHT)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Executive Summary (1-page DOCX)
# ---------------------------------------------------------------------------

def generate_executive_summary_docx(
    result: Dict[str, Any],
    selected_data: pd.DataFrame,
    lang: str = "TR",
) -> bytes | None:
    """Generate a single-page Executive Summary DOCX for non-technical decision-makers."""
    if not _ensure_docx():
        return None

    lang_code = "en-US" if lang == "EN" else "tr-TR"
    d = _extract_analysis_data(result, selected_data)
    f = _fmt

    doc = Document()
    _configure_doc(doc)

    # Title
    title = "EXECUTIVE SUMMARY — Multi-Criteria Decision Analysis" if lang == "EN" \
        else "YÖNETİCİ ÖZETİ — Çok Kriterli Karar Analizi"
    p = doc.add_paragraph()
    run = p.add_run(title)
    _set_run_style(run, lang_code=lang_code, font_size=14, bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    import datetime
    _add_paragraph(doc, datetime.date.today().strftime("%d.%m.%Y"), lang_code=lang_code,
                   font_size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph("")

    # Problem summary
    wm = d["weight_method"] or "?"
    rm = d["ranking_method"] or ""
    if lang == "EN":
        _add_paragraph(doc, (
            f"This analysis evaluated {d['n_alt']} alternatives across {d['n_crit']} criteria "
            f"using {wm} for weighting" + (f" and {rm} for ranking." if rm else ".")
        ), lang_code=lang_code)
    else:
        _add_paragraph(doc, (
            f"Bu analizde {d['n_alt']} alternatif, {d['n_crit']} kriter üzerinden "
            f"{wm} ağırlıklandırma" + (f" ve {rm} sıralama yöntemiyle değerlendirilmiştir." if rm else " yöntemiyle değerlendirilmiştir.")
        ), lang_code=lang_code)
    doc.add_paragraph("")

    # Key finding
    if d["top_alt"]:
        _add_paragraph(doc, "KEY FINDING" if lang == "EN" else "ANA BULGU",
                       lang_code=lang_code, bold=True, font_size=11)
        if lang == "EN":
            finding = f"The leading alternative is {d['top_alt']} with a score of {f(d['top_score'], lang)}."
        else:
            finding = f"Lider alternatif {d['top_alt']}, skor: {f(d['top_score'], lang)}."
        if d["stability"] >= 0.80:
            finding += f" {'Highly stable' if lang == 'EN' else 'Yüksek kararlılık'} (MC: {'%' if lang != 'EN' else ''}{f(d['stability']*100, lang, 1)}{'%' if lang == 'EN' else ''})."
        elif d["stability"] >= 0.60:
            finding += f" {'Moderately stable' if lang == 'EN' else 'Orta kararlılık'} ({'%' if lang != 'EN' else ''}{f(d['stability']*100, lang, 1)}{'%' if lang == 'EN' else ''})."
        elif d["stability"] > 0:
            finding += f" {'Caution: low stability' if lang == 'EN' else 'Dikkat: düşük kararlılık'} ({'%' if lang != 'EN' else ''}{f(d['stability']*100, lang, 1)}{'%' if lang == 'EN' else ''})."
        _add_paragraph(doc, finding, lang_code=lang_code, font_size=12)
    doc.add_paragraph("")

    # Top-3 tables
    if isinstance(d["weight_table"], pd.DataFrame) and not d["weight_table"].empty:
        wt = d["weight_table"].copy()
        w_col = "Ağırlık" if "Ağırlık" in wt.columns else ("Weight" if "Weight" in wt.columns else None)
        if w_col:
            top3_w = wt.nlargest(3, w_col).reset_index(drop=True)
            _add_paragraph(doc, "Top 3 Criteria" if lang == "EN" else "En Önemli 3 Kriter",
                           lang_code=lang_code, bold=True, font_size=10)
            _add_table(doc, top3_w, lang_code=lang_code)

    if isinstance(d["ranking_table"], pd.DataFrame) and not d["ranking_table"].empty:
        r_col = next((c for c in d["ranking_table"].columns if c in ("Sıra", "Rank")), None)
        if r_col:
            top3_r = d["ranking_table"].nsmallest(3, r_col).reset_index(drop=True)
            _add_paragraph(doc, "Top 3 Alternatives" if lang == "EN" else "En İyi 3 Alternatif",
                           lang_code=lang_code, bold=True, font_size=10)
            _add_table(doc, top3_r, lang_code=lang_code)

    # Mini chart
    try:
        wb = _generate_weight_bar_bytes(d["weight_table"], lang)
    except Exception:
        wb = None
    if wb:
        _add_figure(doc, wb, "Criteria Weights" if lang == "EN" else "Kriter Ağırlıkları", lang_code, width_inches=4.5)

    # Recommendation
    doc.add_paragraph("")
    _add_paragraph(doc, "RECOMMENDATION" if lang == "EN" else "ÖNERİ",
                   lang_code=lang_code, bold=True, font_size=11)
    if d["stability"] >= 0.80 and d["mean_spearman"] >= 0.85:
        rec = (f"{d['top_alt']} can be confidently recommended." if lang == "EN"
               else f"{d['top_alt']} güvenle önerilebilir.")
    elif d["stability"] >= 0.60:
        rec = (f"{d['top_alt']} shows advantage but review sensitivity findings." if lang == "EN"
               else f"{d['top_alt']} avantajlıdır ancak duyarlılık bulguları ile birlikte değerlendirilmelidir.")
    else:
        rec = ("No clearly dominant alternative. Further investigation recommended." if lang == "EN"
               else "Açık bir baskın alternatif üretilmemiştir. Ek inceleme önerilir.")
    _add_paragraph(doc, rec, lang_code=lang_code)

    doc.add_paragraph("")
    # Disclaimer + Signature
    _add_paragraph(doc,
        "This output is generated for informational and academic evaluation purposes only. The results presented are analytical recommendations based on the selected method and data; final decision responsibility rests with the decision-maker." if lang == "EN" else "Bu çıktı yalnızca bilgilendirme ve akademik değerlendirme amacıyla üretilmiştir. Sunulan sonuçlar, seçilen yöntem ve veriye dayalı analitik öneriler niteliğindedir; nihai karar sorumluluğu karar vericiye aittir.",
        lang_code=lang_code, italic=True, font_size=8, space_after=4)
    _add_paragraph(doc, "MCDM Karar Destek Sistemi — Prof. Dr. Ömer Faruk Rençber",
        lang_code=lang_code, italic=True, font_size=8,
        alignment=WD_ALIGN_PARAGRAPH.RIGHT)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Dashboard HTML
# ---------------------------------------------------------------------------

def generate_dashboard_html(
    result: Dict[str, Any],
    selected_data: pd.DataFrame,
    lang: str = "TR",
) -> str | None:
    """Generate a self-contained HTML dashboard for browser viewing."""
    d = _extract_analysis_data(result, selected_data)
    f = _fmt

    wm = d["weight_method"] or "?"
    rm = d["ranking_method"] or ""
    stab_pct = f(d["stability"] * 100, lang, 1) if d["stability"] > 0 else "—"
    spearman_txt = f(d["mean_spearman"], lang) if d["mean_spearman"] > 0 else "—"

    weight_rows = ""
    if isinstance(d["weight_table"], pd.DataFrame) and not d["weight_table"].empty:
        wt = d["weight_table"].copy()
        w_col = "Ağırlık" if "Ağırlık" in wt.columns else ("Weight" if "Weight" in wt.columns else None)
        c_col = "Kriter" if "Kriter" in wt.columns else ("Criterion" if "Criterion" in wt.columns else None)
        if w_col and c_col:
            wt = wt.sort_values(w_col, ascending=False)
            for _, row in wt.iterrows():
                pct = float(row[w_col]) * 300
                weight_rows += f'<tr><td>{row[c_col]}</td><td>{f(float(row[w_col]), lang, 4)}</td><td><div class="bar" style="width:{pct:.0f}px"></div></td></tr>\n'

    ranking_rows = ""
    if isinstance(d["ranking_table"], pd.DataFrame) and not d["ranking_table"].empty:
        rt = d["ranking_table"].copy()
        a_col = next((c for c in rt.columns if c in ("Alternatif", "Alternative")), None)
        s_col = next((c for c in rt.columns if c in ("Skor", "Score")), None)
        r_col = next((c for c in rt.columns if c in ("Sıra", "Rank")), None)
        if a_col and s_col and r_col:
            rt = rt.sort_values(r_col)
            max_s = rt[s_col].max() if not rt.empty else 1
            for _, row in rt.iterrows():
                pct = float(row[s_col]) / max_s * 200 if max_s > 0 else 0
                cls = ' class="leader"' if int(row[r_col]) == 1 else ""
                ranking_rows += f'<tr{cls}><td>#{int(row[r_col])}</td><td>{row[a_col]}</td><td>{f(float(row[s_col]), lang, 4)}</td><td><div class="bar" style="width:{pct:.0f}px"></div></td></tr>\n'

    if d["stability"] >= 0.80:
        stab_cls, stab_lbl = "good", ("Yüksek" if lang != "EN" else "High")
    elif d["stability"] >= 0.60:
        stab_cls, stab_lbl = "mid", ("Orta" if lang != "EN" else "Moderate")
    elif d["stability"] > 0:
        stab_cls, stab_lbl = "low", ("Düşük" if lang != "EN" else "Low")
    else:
        stab_cls, stab_lbl = "na", "—"

    title = "Çok Kriterli Karar Analizi — Dashboard" if lang != "EN" else "Multi-Criteria Decision Analysis — Dashboard"

    return f"""<!DOCTYPE html>
<html lang="{'tr' if lang != 'EN' else 'en'}">
<head>
<meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1">
<title>{title}</title>
<style>
*{{margin:0;padding:0;box-sizing:border-box}}
body{{font-family:'Segoe UI',system-ui,sans-serif;background:#f0f4f8;color:#1a202c;padding:1.5rem}}
.header{{text-align:center;margin-bottom:1.5rem}}
.header h1{{font-size:1.4rem;color:#1a365d}}.header p{{font-size:0.85rem;color:#718096}}
.kpi-grid{{display:grid;grid-template-columns:repeat(auto-fit,minmax(170px,1fr));gap:0.8rem;margin-bottom:1.5rem}}
.kpi{{background:#fff;border-radius:10px;padding:1rem;text-align:center;border:1px solid #e2e8f0}}
.kpi .value{{font-size:1.5rem;font-weight:800;color:#2b6cb0}}.kpi .label{{font-size:0.68rem;text-transform:uppercase;color:#a0aec0}}
.kpi.leader .value{{color:#2f855a}}.kpi.stability .value{{color:{'#2f855a' if stab_cls=='good' else '#c05621' if stab_cls=='mid' else '#c53030'}}}
.section{{background:#fff;border-radius:10px;padding:1.2rem;margin-bottom:1rem;border:1px solid #e2e8f0}}
.section h2{{font-size:1rem;color:#2d3748;margin-bottom:0.6rem;border-bottom:2px solid #edf2f7;padding-bottom:0.3rem}}
table{{width:100%;border-collapse:collapse;font-size:0.82rem}}
th{{text-align:left;padding:0.4rem;background:#edf2f7;color:#4a5568;font-weight:600}}
td{{padding:0.35rem 0.4rem;border-bottom:1px solid #f7fafc}}
tr.leader td{{background:#f0fff4;font-weight:700}}
.bar{{height:14px;background:linear-gradient(90deg,#3182ce,#63b3ed);border-radius:3px;min-width:2px}}
.badge{{display:inline-block;padding:0.12rem 0.45rem;border-radius:999px;font-size:0.62rem;font-weight:700}}
.badge.good{{background:#c6f6d5;color:#22543d}}.badge.mid{{background:#fefcbf;color:#744210}}.badge.low{{background:#fed7d7;color:#742a2a}}
.footer{{text-align:center;margin-top:1.5rem;font-size:0.68rem;color:#a0aec0}}
</style>
</head>
<body>
<div class="header"><h1>{title}</h1>
<p>{wm}{' + ' + rm if rm else ''} | {d['n_alt']} {'alternatif' if lang != 'EN' else 'alternatives'} × {d['n_crit']} {'kriter' if lang != 'EN' else 'criteria'}</p></div>
<div class="kpi-grid">
<div class="kpi leader"><div class="value">{d['top_alt'] or '—'}</div><div class="label">{'Lider' if lang != 'EN' else 'Leader'}</div></div>
<div class="kpi"><div class="value">{f(d['top_score'], lang) if d['top_score'] > 0 else '—'}</div><div class="label">{'Skor' if lang != 'EN' else 'Score'}</div></div>
<div class="kpi stability"><div class="value">%{stab_pct} <span class="badge {stab_cls}">{stab_lbl}</span></div><div class="label">{'Kararlılık' if lang != 'EN' else 'Stability'}</div></div>
<div class="kpi"><div class="value">{spearman_txt}</div><div class="label">Spearman ρ</div></div>
</div>
<div class="section"><h2>{'Kriter Ağırlıkları' if lang != 'EN' else 'Criteria Weights'}</h2>
<table><tr><th>{'Kriter' if lang != 'EN' else 'Criterion'}</th><th>{'Ağırlık' if lang != 'EN' else 'Weight'}</th><th></th></tr>{weight_rows}</table></div>
<div class="section"><h2>{'Alternatif Sıralaması' if lang != 'EN' else 'Alternative Ranking'}</h2>
<table><tr><th>#</th><th>{'Alternatif' if lang != 'EN' else 'Alternative'}</th><th>{'Skor' if lang != 'EN' else 'Score'}</th><th></th></tr>{ranking_rows}</table></div>
<div class="footer"><em>Bu çıktı yalnızca bilgilendirme ve akademik değerlendirme amacıyla üretilmiştir. Sunulan sonuçlar, seçilen yöntem ve veriye dayalı analitik öneriler niteliğindedir; nihai karar sorumluluğu karar vericiye aittir.</em><br>MCDM Karar Destek Sistemi — Prof. Dr. Ömer Faruk Rençber</div>
</body></html>"""

# ---------------------------------------------------------------------------
# IMRAD PDF Generator (via fpdf2)
# ---------------------------------------------------------------------------

def _pdf_render_md_block(pdf, md_text: str, font_name: str) -> None:
    """Render markdown method content into PDF, with LaTeX as rendered images."""
    _LATEX_INDICATORS = ["\\frac", "\\sum", "\\prod", "\\sqrt", "\\text", "\\cdot",
                         "\\quad", "\\geq", "\\leq", "\\neq", "\\in", "\\to",
                         "_{", "^{", "\\left", "\\right", "\\max", "\\min",
                         "\\ln", "\\log", "\\bar", "\\hat", "\\tilde"]

    def _has_latex(line: str) -> bool:
        return any(ind in line for ind in _LATEX_INDICATORS)

    for line in md_text.split("\n"):
        line = line.strip()
        if not line or line == "---":
            continue
        if line.startswith("$$"):
            # Display math
            pdf.formula_text(line.strip("$").strip())
        elif _has_latex(line):
            # Line contains LaTeX — render as formula
            # Extract just the math part if mixed with text
            clean = line.strip("$").strip()
            pdf.formula_text(clean)
        elif line.startswith("**") and line.endswith("**"):
            pdf.set_font(font_name, "B", 11)
            pdf.multi_cell(0, 5.5, line.strip("*").strip())
            pdf.ln(2)
        elif line.startswith("Adım") or line.startswith("Step"):
            pdf.body_text(line)
        else:
            # Remove inline $ delimiters for plain text
            clean = re.sub(r'\$([^$]+)\$', r'\1', line)
            if _has_latex(clean):
                pdf.formula_text(clean)
            else:
                pdf.body_text(clean)


def generate_imrad_pdf(
    result: Dict[str, Any],
    selected_data: pd.DataFrame,
    lang: str = "TR",
) -> bytes | None:
    """Generate IMRAD article as PDF with APA formatting.

    Uses fpdf2 for platform-independent PDF generation.
    Formulas rendered as UTF-8 text with Cambria Math fallback.
    """
    try:
        from fpdf import FPDF
    except ImportError:
        return None

    d = _extract_analysis_data(result, selected_data)
    f = _fmt
    template = _load_template()
    wm = d["weight_method"] or "?"
    rm = d["ranking_method"] or ""

    class ArticlePDF(FPDF):
        def header(self):
            self.set_font(self._fn, "I", 8)
            self.set_text_color(128, 128, 128)
            self.cell(0, 5, "MCDM Karar Destek Sistemi", align="R")
            self.ln(8)

        def footer(self):
            self.set_y(-15)
            self.set_font(self._fn, "I", 8)
            self.set_text_color(128, 128, 128)
            self.cell(0, 10, f"Sayfa {self.page_no()}/{{nb}}", align="C")

        def section_heading(self, text):
            self.set_font(self._fn, "B", 13)
            self.set_text_color(0, 0, 0)
            self.cell(0, 8, text)
            self.ln(10)

        def sub_heading(self, text):
            self.set_font(self._fn, "B", 12)
            self.set_text_color(0, 0, 0)
            self.cell(0, 7, text)
            self.ln(9)

        def body_text(self, text):
            self.set_font(self._fn, "", 12)
            self.set_text_color(0, 0, 0)
            self.multi_cell(0, 6, text)
            self.ln(3)

        def italic_guidance(self, text):
            self.set_font(self._fn, "I", 10)
            self.set_text_color(100, 100, 100)
            self.multi_cell(0, 5, text)
            self.ln(4)

        def formula_text(self, latex):
            """Render LaTeX formula as PNG via matplotlib and embed in PDF."""
            if _ensure_matplotlib():
                try:
                    fig, ax = _plt.subplots(figsize=(6, 0.6))
                    ax.axis("off")
                    ax.text(0.5, 0.5, f"${latex}$", transform=ax.transAxes,
                            ha="center", va="center", fontsize=12,
                            fontfamily="serif", math_fontfamily="cm")
                    buf = io.BytesIO()
                    fig.savefig(buf, format="png", dpi=200, bbox_inches="tight",
                                facecolor="white", edgecolor="none", pad_inches=0.05)
                    _plt.close(fig)
                    buf.seek(0)
                    x = (210 - 120) / 2  # center on A4
                    self.image(buf, x=x, w=120)
                    self.ln(3)
                    return
                except Exception:
                    pass
            # Fallback: plain text
            self.set_font("Courier", "", 9)
            self.set_text_color(0, 0, 0)
            self.multi_cell(0, 5, latex, align="C")
            self.ln(2)

        def add_data_table(self, df, max_rows=20):
            if df is None or df.empty:
                return
            df = df.head(max_rows)
            col_widths = [max(25, min(45, 190 // len(df.columns)))] * len(df.columns)
            total_w = sum(col_widths[:len(df.columns)])
            col_widths = [w * 190 / total_w for w in col_widths]

            self.set_font(self._fn, "B", 9)
            self.set_fill_color(230, 230, 230)
            for i, col in enumerate(df.columns):
                self.cell(col_widths[i], 6, str(col)[:15], border=1, fill=True, align="C")
            self.ln()

            self.set_font(self._fn, "", 9)
            self.set_fill_color(255, 255, 255)
            for _, row in df.iterrows():
                for j, val in enumerate(row):
                    if isinstance(val, (float, np.floating)) and pd.notna(val):
                        txt = f"{val:.3f}"
                    else:
                        txt = str(val)[:15]
                    self.cell(col_widths[j], 5, txt, border=1, align="C")
                self.ln()
            self.ln(4)

    pdf = ArticlePDF()
    pdf.alias_nb_pages()
    # Load Unicode TTF fonts
    _font_dir = "/System/Library/Fonts/Supplemental/"
    try:
        pdf.add_font("TNR", "", _font_dir + "Times New Roman.ttf", uni=True)
        pdf.add_font("TNR", "B", _font_dir + "Times New Roman Bold.ttf", uni=True)
        pdf.add_font("TNR", "I", _font_dir + "Times New Roman Italic.ttf", uni=True)
        pdf.add_font("TNR", "BI", _font_dir + "Times New Roman Bold Italic.ttf", uni=True)
        _font_name = "TNR"
    except Exception:
        _font_name = "Times"
    pdf._fn = _font_name
    pdf.add_page()

    # Title
    pdf.set_font(_font_name, "B", 15)
    pdf.set_text_color(0, 0, 0)
    title = _build_title(d, lang)
    pdf.multi_cell(0, 8, title, align="C")
    pdf.ln(8)

    # Abstract
    h = _h(lang)
    pdf.section_heading(h["abstract"])
    abstract = _build_abstract(d, lang)
    pdf.set_font(_font_name, "", 11)
    pdf.multi_cell(0, 5.5, abstract)
    pdf.ln(3)
    kw = f"{h['keywords_label']}: {'MCDM' if lang == 'EN' else 'CKKV'}, {wm}" + (f", {rm}" if rm else "") + ", " + ("Robustness" if lang == "EN" else "Saglamlik")
    pdf.set_font(_font_name, "BI", 11)
    pdf.cell(0, 6, kw)
    pdf.ln(10)

    # Introduction
    pdf.section_heading(h["introduction"])
    pdf.italic_guidance(
        "This section should present the research problem, literature summary, and contribution."
        if lang == "EN" else
        "Bu bolumde arastirma problemi, literatur ozeti ve calismanin katkisi sunulmalidir."
    )
    pdf.body_text(
        f"{d['n_alt']} {'alternatives and' if lang == 'EN' else 'alternatif ve'} {d['n_crit']} "
        f"{'criteria were analyzed using' if lang == 'EN' else 'kriter uzerinde'} {wm}"
        + (f" {'and' if lang == 'EN' else 've'} {rm}" if rm else "")
        + f" {'were applied.' if lang == 'EN' else 'uygulanmistir.'}"
    )

    # Method
    pdf.section_heading(h["method"])

    # Weight method
    pdf.sub_heading(h["method_weight"])
    # Justification
    w_just = (_WEIGHT_JUSTIFICATION_EN if lang == "EN" else _WEIGHT_JUSTIFICATION_TR).get(wm, "")
    if w_just:
        pdf.body_text(w_just)

    # Weight formulas from template
    weight_content = _find_template_content(wm, template["weight_methods"])
    if weight_content:
        _pdf_render_md_block(pdf, weight_content, _font_name)

    # Ranking method
    if rm:
        pdf.sub_heading(h["method_rank"])
        base_rm = rm.replace("Fuzzy ", "")
        r_just = (_RANKING_JUSTIFICATION_EN if lang == "EN" else _RANKING_JUSTIFICATION_TR).get(rm, "")
        if not r_just:
            r_just = (_RANKING_JUSTIFICATION_EN if lang == "EN" else _RANKING_JUSTIFICATION_TR).get(base_rm, "")
        if r_just:
            pdf.body_text(r_just)

        rank_content = _find_template_content(rm, template["ranking_methods"])
        if rank_content:
            _pdf_render_md_block(pdf, rank_content, _font_name)

    # Results
    pdf.add_page()
    pdf.section_heading(h["results"])

    pdf.sub_heading(h["results_weight"])
    if isinstance(d["weight_table"], pd.DataFrame) and not d["weight_table"].empty:
        pdf.add_data_table(d["weight_table"])
    if d["top_criterion"]:
        pdf.body_text(
            f"{'Highest weight:' if lang == 'EN' else 'En yuksek agirlik:'} "
            f"{d['top_criterion']} (w = {f(d['top_weight'], lang)})"
        )

    if rm:
        pdf.sub_heading(h["results_rank"])
        if isinstance(d["ranking_table"], pd.DataFrame) and not d["ranking_table"].empty:
            pdf.add_data_table(d["ranking_table"])

        pdf.sub_heading(h["results_robust"])
        stab = d.get("stability", 0)
        if stab > 0 and d["top_alt"]:
            if lang == "EN":
                pdf.body_text(f"MC stability: {d['top_alt']} maintains first place in {f(stab*100, lang, 1)}% of iterations.")
            else:
                pdf.body_text(f"MC kararliligi: {d['top_alt']} iterasyonlarin %{f(stab*100, lang, 1)}'inde birinciligi korumaktadir.")
        if d["mean_spearman"] > 0:
            pdf.body_text(f"Spearman rho = {f(d['mean_spearman'], lang)}")

    # Discussion
    pdf.section_heading(h["discussion"])
    pdf.italic_guidance(
        "Compare findings with literature, interpret results, state limitations."
        if lang == "EN" else
        "Bulgular literaturle karsilastirilmali, sonuclar yorumlanmali, kisitlar belirtilmelidir."
    )

    # Conclusion
    pdf.section_heading(h["conclusion"])
    pdf.italic_guidance(
        "Summarize findings, highlight contribution, suggest future work."
        if lang == "EN" else
        "Bulgulari ozetleyin, katkiyi vurgulayin, gelecek calisma onerin."
    )

    # Disclaimer
    pdf.ln(10)
    pdf.set_font(_font_name, "I", 9)
    pdf.set_text_color(100, 100, 100)
    disclaimer = (
        "This output is for informational and academic evaluation purposes only. "
        "Final decision responsibility rests with the decision-maker."
    ) if lang == "EN" else (
        "Bu cikti yalnizca bilgilendirme ve akademik degerlendirme amaciyla uretilmistir. "
        "Nihai karar sorumlulugu karar vericiye aittir."
    )
    pdf.multi_cell(0, 4.5, disclaimer)
    pdf.ln(2)
    pdf.set_font(_font_name, "I", 9)
    pdf.cell(0, 5, "MCDM Karar Destek Sistemi -- Prof. Dr. Omer Faruk Rencber", align="R")

    return pdf.output()
